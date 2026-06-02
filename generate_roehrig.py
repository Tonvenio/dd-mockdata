"""
Generate ~5000 due-diligence documents for Brennhagen Elektronik AG
Output: roehrig_large/
"""
# --- portable-paths-prelude --- (do not edit) ---
import sys
from pathlib import Path as _PathlibPath
_RP = _PathlibPath(__file__).resolve().parent
while not (_RP / "enhance_lib.py").exists() and _RP.parent != _RP:
    _RP = _RP.parent
_ROOT = _RP
sys.path.insert(0, str(_ROOT))
# --- end prelude ---

import os
import random
import sys
from datetime import date, timedelta
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                 TableStyle, HRFlowable)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

random.seed(42)

# ── Canonical data ──────────────────────────────────────────────────────────

AG = {
    "name": "Brennhagen Elektronik AG",
    "short": "REA",
    "hrb": "HRB 726451",
    "amtsgericht": "Amtsgericht Stuttgart",
    "founded": "18. März 1962",
    "address": "Vaihinger Straße 120",
    "plz": "70567",
    "city": "Stuttgart",
    "full_address": "Vaihinger Straße 120, 70567 Stuttgart",
    "phone": "+49 711 89200-0",
    "email": "ir@brennhagen-elektronik.de",
    "web": "www.brennhagen-elektronik.de",
    "stock": "ROHL.DE (XETRA / TecDAX)",
    "isin": "DE000REA1234",
    "ust_id": "DE 197 634 211",
    "steuernr": "99001/12345/67890",
    "fa": "Finanzamt Stuttgart-Körperschaften",
    "iban": "DE55 6002 0290 0987 6543 21",
    "bic": "HYVEDEMMXXX",
    "bank": "Deutsche Bank AG & UniCredit Bank AG (Konsortium)",
    "ceo": "Dr. Petra Hollmann",
    "coo": "Klaus-Peter Zimmermann",
    "cfo": "Stefan Richter",
    "cto": "Dr. Yuki Tanaka",
    "arb_chair": "Prof. Dr. Gerhard Voss",
    "revenue_2020": 468_000_000,
    "revenue_2021": 502_000_000,
    "revenue_2022": 573_000_000,
    "revenue_2023": 624_000_000,
    "revenue_2024e": 658_000_000,
    "ebitda_2020": 48_200_000,
    "ebitda_2021": 54_100_000,
    "ebitda_2022": 63_800_000,
    "ebitda_2023": 68_400_000,
    "ebitda_2024e": 74_000_000,
    "ebit_2020": 28_400_000,
    "ebit_2021": 33_200_000,
    "ebit_2022": 41_500_000,
    "ebit_2023": 44_800_000,
    "employees_2021": 3_102,
    "employees_2022": 3_287,
    "employees_2023": 3_487,
    "grundkapital": 12_000_000,
    "market_cap_2023": 648_000_000,
    "kunde1": "BMW Group",
    "kunde1_rev_share": "28 %",
    "kunde2": "Volkswagen AG",
    "kunde2_rev_share": "22 %",
    "kunde3": "Mercedes-Benz AG",
    "kunde3_rev_share": "18 %",
    "kunde4": "Stellantis N.V.",
    "kunde4_rev_share": "12 %",
    "kunde5": "Continental AG",
    "kunde5_rev_share": "8 %",
    "prod1": "ECU-900 (Engine Control Unit)",
    "prod2": "ADAS-Vision 4D (Kamerasystem)",
    "prod3": "InfoConnect Pro (Infotainment-Modul)",
    "prod4": "BatteryMS-12 (Batteriemanagementsystem EV)",
    "lief1": "Infineon Technologies AG",
    "lief2": "NXP Semiconductors N.V.",
    "lief3": "STMicroelectronics N.V.",
    "lief4": "Robert Bosch GmbH – Semiconductor Division",
    "wp": "KPMG AG Wirtschaftsprüfungsgesellschaft",
    "recht": "Freshfields Bruckhaus Deringer LLP",
    "steuer": "Deloitte GmbH WPG",
    "erp": "SAP S/4HANA (On-Premise), Release 2022",
    "plm": "Siemens Teamcenter PLM",
    "mes": "SAP MES (Produktionssteuerung)",
    "iatf": "IATF 16949:2016",
}

SUBS = [
    {"name": "Brennhagen Elektronik GmbH",        "short": "REG", "city": "Heilbronn",  "country": "DE", "hrb": "HRB 221456, AG Heilbronn",          "employees": 820, "revenue": 280_000_000, "focus": "Produktion / Hauptwerk"},
    {"name": "Brennhagen Software GmbH",           "short": "RSG", "city": "München",   "country": "DE", "hrb": "HRB 319872, AG München",             "employees": 340, "revenue": 62_000_000,  "focus": "Embedded Software / ADAS"},
    {"name": "Brennhagen Polska sp. z o.o.",       "short": "RPL", "city": "Katowice",  "country": "PL", "hrb": "KRS 0000412876",                     "employees": 680, "revenue": 78_000_000,  "focus": "Produktion / Kostenoptimierung"},
    {"name": "Brennhagen Czech s.r.o.",            "short": "RCZ", "city": "Brno",      "country": "CZ", "hrb": "IČO: 28418762",                      "employees": 410, "revenue": 54_000_000,  "focus": "Engineering / Entwicklung"},
    {"name": "Brennhagen Hungary Kft.",            "short": "RHU", "city": "Győr",      "country": "HU", "hrb": "Cg. 08-09-025671",                   "employees": 520, "revenue": 64_000_000,  "focus": "Produktion (EV-Komponenten)"},
    {"name": "Brennhagen Electronics (China) Ltd.","short": "RCN", "city": "Shanghai",  "country": "CN", "hrb": "统一社会信用代码: 91310115MA1FL42Q38","employees": 280, "revenue": 48_000_000,  "focus": "Vertrieb / Lokale Produktion"},
    {"name": "Brennhagen Finance GmbH",            "short": "RFI", "city": "Luxembourg","country": "LU", "hrb": "RCS Luxembourg B215432",              "employees": 12,  "revenue": 0,           "focus": "Konzernfinanzierung / Treasury"},
    {"name": "Brennhagen Holding GmbH",            "short": "RHO", "city": "Stuttgart", "country": "DE", "hrb": "HRB 726450, AG Stuttgart",            "employees": 45,  "revenue": 0,           "focus": "Holding / Managementgesellschaft"},
]

CUSTOMERS = [
    {"name": "BMW Group",        "short": "BMW", "share": "28 %", "contact": "Global Procurement, Munich"},
    {"name": "Volkswagen AG",    "short": "VW",  "share": "22 %", "contact": "Group Purchasing, Wolfsburg"},
    {"name": "Mercedes-Benz AG", "short": "MBZ", "share": "18 %", "contact": "Strategic Procurement, Stuttgart"},
    {"name": "Stellantis N.V.",  "short": "STE", "share": "12 %", "contact": "European Supply Management, Amsterdam"},
    {"name": "Continental AG",   "short": "CON", "share": "8 %",  "contact": "Tier-1 Sourcing, Hannover"},
]

PRODUCTS = [
    {"id": "ECU-900",      "name": "ECU-900 (Engine Control Unit)",            "segment": "Powertrain"},
    {"id": "ADAS-V4D",     "name": "ADAS-Vision 4D (Kamerasystem)",            "segment": "ADAS"},
    {"id": "ICP-3",        "name": "InfoConnect Pro (Infotainment-Modul)",     "segment": "Infotainment"},
    {"id": "BMS-12",       "name": "BatteryMS-12 (Batteriemanagementsystem EV)","segment": "EV"},
]

SUPPLIERS = [
    {"name": "Infineon Technologies AG",              "short": "INF", "country": "DE"},
    {"name": "NXP Semiconductors N.V.",               "short": "NXP", "country": "NL"},
    {"name": "STMicroelectronics N.V.",               "short": "STM", "country": "CH"},
    {"name": "Robert Bosch GmbH – Semiconductor Div.","short": "BOS", "country": "DE"},
]

PROJECTS = [
    ("PRJ-2021-001", "ECU-900 Gen3 Entwicklung",        "2021-03", "2023-08", "BMW Group"),
    ("PRJ-2021-002", "ADAS-V4D Kalibrierungsplattform", "2021-06", "2023-12", "Mercedes-Benz AG"),
    ("PRJ-2022-001", "BatteryMS-12 EV Plattform",       "2022-01", "2024-06", "Volkswagen AG"),
    ("PRJ-2022-002", "InfoConnect Pro 4.0",             "2022-04", "2024-03", "Stellantis N.V."),
    ("PRJ-2022-003", "ECU-900 AUTOSAR Refactoring",     "2022-07", "2023-11", "intern"),
    ("PRJ-2023-001", "ADAS-V4D Radar Fusion",           "2023-01", "2025-06", "BMW Group"),
    ("PRJ-2023-002", "BatteryMS-12 China Adaptation",  "2023-03", "2024-09", "RCN Shanghai"),
    ("PRJ-2023-003", "ICP-3 OTA Update Platform",      "2023-06", "2025-03", "VW / Mercedes"),
    ("PRJ-2023-004", "Cost Reduction Program CZ",      "2023-02", "2023-12", "intern"),
    ("PRJ-2023-005", "Heilbronn Plant Expansion",       "2023-05", "2025-12", "REG"),
    ("PRJ-2024-001", "ECU-1000 Concept Study",         "2024-01", "2025-12", "BMW Group"),
    ("PRJ-2024-002", "DSGVO Compliance Remediation",   "2024-02", "2024-06", "intern"),
    ("PRJ-2024-003", "SAP S/4HANA Rollout Poland",     "2024-03", "2025-06", "RPL"),
    ("PRJ-2024-004", "LkSG Supply Chain Audit",        "2024-01", "2024-08", "intern"),
    ("PRJ-2024-005", "Cybersecurity TISAX Level 3",    "2024-04", "2025-03", "intern"),
]

# ── BASE PATH ────────────────────────────────────────────────────────────────

BASE = Path(f"{_ROOT}/roehrig_large")
BASE.mkdir(parents=True, exist_ok=True)

FOLDERS = [
    "00_Konzernstruktur_Holding",
    "01_AG_Gesellschaftsrecht",
    "02_Konsolidierte_Finanzen",
    "03_Tochter_DE_REG_Heilbronn",
    "04_Tochter_DE_RSG_Muenchen",
    "05_Tochter_PL_Katowice",
    "06_Tochter_CZ_Brno",
    "07_Tochter_HU_Gyoer",
    "08_Tochter_CN_Shanghai",
    "09_Personal_HR",
    "10_Kapitalmarkt_IR",
    "11_Vertrieb_OEM",
    "12_Einkauf_Lieferanten",
    "13_IATF_Qualitaet",
    "14_IP_Technologie",
    "15_Compliance_Recht",
    "16_IT_Systeme",
    "17_Versicherungen",
    "18_Immobilien",
    "19_MA_Transaktionen",
    "20_Strategie_Vorstand",
    "21_Betriebsraete",
    "22_Pensionen_bAV",
    "23_Projekte_Programme",
]

for f in FOLDERS:
    (BASE / f).mkdir(parents=True, exist_ok=True)

# ── Helpers ──────────────────────────────────────────────────────────────────

MONTHS_DE = ["","Januar","Februar","März","April","Mai","Juni","Juli","August",
             "September","Oktober","November","Dezember"]
MONTHS_EN = ["","January","February","March","April","May","June","July","August",
             "September","October","November","December"]

def ds(y, m, d):
    return f"{d}. {MONTHS_DE[m]} {y}"

def eur(n):
    return f"{n:,.0f} €".replace(",", ".")

def sfn(s, maxlen=45):
    for ch in ["/", chr(92), ":", "*", "?", '"', "<", ">", "|", "(", ")", "&", "+", ",", " – ", " - "]:
        s = s.replace(ch, "_")
    s = s.replace(" ", "_")
    return s[:maxlen]

def mess_name(name, p=0.08):
    if random.random() < p:
        base = name.rsplit(".", 1)[0] if "." in name else name
        ext = ("." + name.rsplit(".", 1)[1]) if "." in name else ""
        suf = random.choice(["_v1", "_v2", "_DRAFT", "_ENTWURF", "_FINAL",
                              "_FINAL_v2", "_rev_SRichter", "_20230915",
                              "_reviewed", "_ALT", "_2024-03-01", "_WIP", "_intern"])
        return base + suf + ext
    return name

def wrong_folder(folder, p=0.05):
    if random.random() < p:
        return random.choice(FOLDERS)
    return folder

TOTAL = [0]

def make_docx(folder, filename, company_name, company_addr, hrb, title,
              sections, subtitle=None, confidential=False, draft=False, language="de"):
    folder = wrong_folder(folder)
    folder_path = BASE / folder
    folder_path.mkdir(parents=True, exist_ok=True)
    filename = mess_name(filename)

    doc = Document()
    if confidential or random.random() < 0.06:
        p = doc.add_paragraph("VERTRAULICH – STRENG VERTRAULICH / STRICTLY CONFIDENTIAL")
        p.runs[0].bold = True
    if draft or random.random() < 0.05:
        p = doc.add_paragraph("*** ENTWURF / DRAFT – NICHT ZUR WEITERGABE ***")
        p.runs[0].bold = True

    doc.add_heading(company_name, level=1)
    p = doc.add_paragraph()
    run = p.add_run(company_addr + "  |  " + hrb)
    run.font.size = Pt(9)
    doc.add_paragraph("─" * 80)
    doc.add_heading(title, level=2)
    if subtitle:
        sub = doc.add_paragraph(subtitle)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for sec_title, sec_body in sections:
        if sec_title:
            doc.add_heading(sec_title, level=3)
        if isinstance(sec_body, str):
            for para in sec_body.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
        elif isinstance(sec_body, list):
            if sec_body:
                t = doc.add_table(rows=1, cols=len(sec_body[0]))
                t.style = "Table Grid"
                for ci, hdr in enumerate(sec_body[0]):
                    t.rows[0].cells[ci].text = str(hdr)
                for row in sec_body[1:]:
                    r = t.add_row()
                    for ci, val in enumerate(row):
                        r.cells[ci].text = str(val)

    path = folder_path / filename
    doc.save(path)
    TOTAL[0] += 1
    if TOTAL[0] % 100 == 0:
        print(f"  [{TOTAL[0]}] docs created …")
    return path


def make_pdf(folder, filename, company_name, company_addr, hrb, title,
             sections, confidential=False):
    folder = wrong_folder(folder)
    folder_path = BASE / folder
    folder_path.mkdir(parents=True, exist_ok=True)
    filename = mess_name(filename)

    path = folder_path / filename
    pdoc = SimpleDocTemplate(str(path), pagesize=A4,
                             leftMargin=2.5*cm, rightMargin=2.5*cm,
                             topMargin=2.5*cm, bottomMargin=2.5*cm)
    styles = getSampleStyleSheet()
    body_s = ParagraphStyle("b", parent=styles["Normal"], fontSize=10,
                             leading=15, spaceAfter=8, alignment=TA_JUSTIFY)
    h1_s = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=14,
                           spaceBefore=14, spaceAfter=6)
    h2_s = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=11,
                           spaceBefore=10, spaceAfter=4)
    small_s = ParagraphStyle("sm", parent=styles["Normal"], fontSize=8,
                              textColor=colors.grey)

    story = []
    if confidential:
        story.append(Paragraph("VERTRAULICH / STRICTLY CONFIDENTIAL",
                                ParagraphStyle("conf", parent=styles["Normal"],
                                               fontSize=10, textColor=colors.red,
                                               spaceAfter=6)))
    story.append(Paragraph(company_name, h1_s))
    story.append(Paragraph(company_addr + " | " + hrb, small_s))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#003366")))
    story.append(Spacer(1, 0.4*cm))
    story.append(Paragraph(title, h1_s))
    story.append(Spacer(1, 0.3*cm))

    for sec_title, sec_body in sections:
        if sec_title:
            story.append(Paragraph(sec_title, h2_s))
        if isinstance(sec_body, str):
            for para in sec_body.split("\n\n"):
                if para.strip():
                    story.append(Paragraph(para.replace("\n", "<br/>"), body_s))
        elif isinstance(sec_body, list):
            if sec_body:
                t = Table(sec_body, hAlign="LEFT")
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#003366")),
                    ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
                    ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE",   (0, 0), (-1, -1), 9),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                     [colors.white, colors.HexColor("#EEF2F8")]),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                    ("LEFTPADDING",  (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING",   (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
                ]))
                story.append(t)
        story.append(Spacer(1, 0.3*cm))

    pdoc.build(story)
    TOTAL[0] += 1
    if TOTAL[0] % 100 == 0:
        print(f"  [{TOTAL[0]}] docs created …")
    return path


def make_xlsx(folder, filename, title, sheets_data):
    folder = wrong_folder(folder)
    folder_path = BASE / folder
    folder_path.mkdir(parents=True, exist_ok=True)
    filename = mess_name(filename)

    path = folder_path / filename
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    hdr_fill  = PatternFill("solid", fgColor="003366")
    hdr_font  = Font(bold=True, color="FFFFFF", size=10)
    alt_fill  = PatternFill("solid", fgColor="EEF2F8")
    title_font = Font(bold=True, size=12, color="003366")
    thin   = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for sheet_name, headers, rows, col_widths in sheets_data:
        safe_sn = sheet_name.replace('/','_').replace('\\','_').replace('?','_').replace('*','_').replace('[','_').replace(']','_').replace(':','_')[:31]
        ws = wb.create_sheet(safe_sn)
        ws.append([title])
        ws["A1"].font = title_font
        ws.append([AG["name"] + " – " + AG["stock"]])
        ws.append([])
        ws.append(headers)
        hr = 4
        for ci, h in enumerate(headers, 1):
            cell = ws.cell(row=hr, column=ci)
            cell.fill = hdr_fill
            cell.font = hdr_font
            cell.alignment = Alignment(horizontal="center", wrap_text=True)
            cell.border = border
        for ri, row in enumerate(rows, hr + 1):
            ws.append(row)
            if ri % 2 == 0:
                for ci in range(1, len(headers) + 1):
                    ws.cell(row=ri, column=ci).fill = alt_fill
            for ci in range(1, len(headers) + 1):
                ws.cell(row=ri, column=ci).border = border
        for ci, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(ci)].width = w

    wb.save(path)
    TOTAL[0] += 1
    if TOTAL[0] % 100 == 0:
        print(f"  [{TOTAL[0]}] docs created …")
    return path

# ── Reusable text blocks ─────────────────────────────────────────────────────

def legal_boilerplate_de(company):
    return (
        f"Diese Vereinbarung wurde geschlossen zwischen {company['name']}, eingetragen im Handelsregister unter "
        f"{company['hrb']}, vertreten durch ihre Geschäftsführung (nachfolgend <<Gesellschaft<<), und den jeweils "
        "nachfolgend genannten Vertragsparteien (gemeinsam <<Parteien<<). Die Parteien vereinbaren Folgendes:"
    )

def legal_boilerplate_en(company):
    return (
        f"This Agreement is entered into by and between {company['name']}, registered under {company['hrb']} "
        "(hereinafter \"Company\"), and the respective counterparties identified below (collectively \"Parties\"). "
        "The Parties agree as follows:"
    )

def std_definitions_de():
    return (
        "§ 1 Begriffsbestimmungen\n\n"
        "(1) <<Vertrauliche Informationen>> bezeichnen alle schriftlich, mündlich oder in sonstiger Form übermittelten "
        "Informationen, die als vertraulich gekennzeichnet sind oder die ihrer Natur nach als vertraulich anzusehen sind.\n\n"
        "(2) <<Konzerngesellschaft<< bezeichnet jede Gesellschaft, an der die Brennhagen Elektronik AG mittelbar oder unmittelbar "
        "mehr als 50 % der Anteile hält oder über deren Geschäftsführung sie die Kontrolle ausübt.\n\n"
        "(3) <<Geistiges Eigentum>> umfasst alle Patente, Gebrauchsmuster, Marken, Urheberrechte, Know-how sowie sonstige "
        "gewerbliche Schutzrechte, unabhängig davon, ob sie eingetragen sind oder nicht.\n\n"
        "(4) <<Leistungen<< bezeichnet alle gemäß diesem Vertrag zu erbringenden Liefer- und Dienstleistungen, einschließlich "
        "aller nach diesem Vertrag zu erstellenden Arbeitsergebnisse und Dokumentationen."
    )

def std_term_termination_de():
    return (
        "§ 12 Laufzeit und Kündigung\n\n"
        "(1) Dieser Vertrag tritt mit Unterzeichnung durch beide Parteien in Kraft und hat eine Mindestlaufzeit von "
        "24 Monaten ab Inkrafttreten.\n\n"
        "(2) Nach Ablauf der Mindestlaufzeit verlängert sich dieser Vertrag automatisch um jeweils zwölf (12) Monate, "
        "sofern er nicht von einer Partei mit einer Frist von sechs (6) Monaten zum jeweiligen Laufzeitende schriftlich "
        "gekündigt wird.\n\n"
        "(3) Das Recht zur außerordentlichen Kündigung aus wichtigem Grund bleibt unberührt. Ein wichtiger Grund liegt "
        "insbesondere vor, wenn die andere Partei wesentliche Vertragspflichten schuldhaft verletzt und diese Verletzung "
        "trotz schriftlicher Abmahnung nicht innerhalb von 30 Tagen behebt, über das Vermögen der anderen Partei ein "
        "Insolvenzverfahren beantragt oder eröffnet wird oder die andere Partei aufgehört hat, ihren Geschäftsbetrieb "
        "fortzuführen."
    )

def std_governing_law_de():
    return (
        "§ 16 Schlussbestimmungen\n\n"
        "(1) Dieser Vertrag unterliegt dem Recht der Bundesrepublik Deutschland unter Ausschluss des UN-Kaufrechts (CISG).\n\n"
        "(2) Ausschließlicher Gerichtsstand für alle Streitigkeiten aus oder im Zusammenhang mit diesem Vertrag ist Stuttgart.\n\n"
        "(3) Änderungen und Ergänzungen dieses Vertrags bedürfen der Schriftform. Dies gilt auch für die Aufhebung des "
        "Schriftformerfordernisses.\n\n"
        "(4) Sollte eine Bestimmung dieses Vertrags unwirksam oder undurchführbar sein, so berührt dies die Gültigkeit "
        "der übrigen Bestimmungen nicht. Die Parteien verpflichten sich, die unwirksame Bestimmung durch eine wirksame "
        "Regelung zu ersetzen, die dem wirtschaftlichen Zweck der unwirksamen Bestimmung möglichst nahekommt.\n\n"
        "(5) Dieser Vertrag einschließlich seiner Anlagen stellt die vollständige Vereinbarung der Parteien dar und "
        "ersetzt alle früheren schriftlichen und mündlichen Vereinbarungen über den gleichen Gegenstand."
    )

def sig_block_de(company, counterparty="[Vertragspartner]"):
    return (
        f"\n\nStuttgart, den {ds(2023, random.randint(1,12), random.randint(1,28))}\n\n"
        f"{company['name']}\n\n"
        "________________________                    ________________________\n"
        f"{AG['ceo']}                             {AG['cfo']}\n"
        f"Vorstandsvorsitzende                        Vorstand Finanzen\n\n"
        f"{counterparty}\n\n"
        "________________________\n"
        "[Unterschrift / Signature]\n\n"
        "Anlagen:\nAnlage 1 – [Leistungsbeschreibung]\nAnlage 2 – [Preisliste]\nAnlage 3 – [Qualitätsanforderungen]\n"
        "Anlage 4 – Exhibit A – to be attached"
    )


# ═══════════════════════════════════════════════════════════════════════════
# 00 – KONZERNSTRUKTUR & HOLDING
# ═══════════════════════════════════════════════════════════════════════════

def gen_00():
    folder = "00_Konzernstruktur_Holding"
    print(f"\n[00] Konzernstruktur / Holding …")

    # AG Satzung
    make_docx(folder, "REA_Satzung_AG_2023.docx",
              AG["name"], AG["full_address"], f"HRB {AG['hrb']} | {AG['amtsgericht']}",
              "Satzung der Brennhagen Elektronik AG",
              [
                  ("Präambel", f"Die Brennhagen Elektronik AG (nachfolgend <<Gesellschaft<<) wurde am {AG['founded']} gegründet und >>"
                   f"ist im Handelsregister des {AG['amtsgericht']} unter {AG['hrb']} eingetragen. "
                   "Die nachfolgende Satzung wurde zuletzt durch Beschluss der Hauptversammlung vom 12. Mai 2023 geändert."),
                  ("§ 1 Firma und Sitz", f"(1) Die Gesellschaft führt die Firma: {AG['name']}.\n\n(2) Der Sitz der Gesellschaft ist Stuttgart."),
                  ("§ 2 Gegenstand des Unternehmens", "(1) Gegenstand des Unternehmens ist die Entwicklung, Herstellung und der Vertrieb von elektronischen Komponenten, Steuergeräten, Sensorsystemen sowie Software für die Automobilindustrie.\n\n(2) Die Gesellschaft ist berechtigt, alle Geschäfte zu betreiben und Maßnahmen zu ergreifen, die geeignet erscheinen, den Gesellschaftszweck zu fördern, einschließlich der Beteiligung an anderen Unternehmen im In- und Ausland."),
                  ("§ 3 Grundkapital", f"(1) Das Grundkapital der Gesellschaft beträgt {eur(AG['grundkapital'])}.\n\n(2) Das Grundkapital ist eingeteilt in 12.000.000 auf den Namen lautende Stückaktien.\n\n(3) Bei einer Kapitalerhöhung kann der Ausgabebetrag und die Art der Einlage vom Vorstand mit Zustimmung des Aufsichtsrats festgesetzt werden."),
                  ("§ 4 Vorstand", f"(1) Der Vorstand besteht aus mindestens zwei Mitgliedern. Die genaue Anzahl bestimmt der Aufsichtsrat.\n\n(2) Ist nur ein Vorstandsmitglied bestellt, so vertritt es die Gesellschaft allein. Sind mehrere Vorstandsmitglieder bestellt, so wird die Gesellschaft durch zwei Vorstandsmitglieder gemeinsam oder durch ein Vorstandsmitglied in Gemeinschaft mit einem Prokuristen vertreten.\n\n(3) Aktuell bestehender Vorstand: {AG['ceo']} (CEO), {AG['coo']} (COO), {AG['cfo']} (CFO), {AG['cto']} (CTO)."),
                  ("§ 5 Aufsichtsrat", f"(1) Der Aufsichtsrat besteht aus zwölf Mitgliedern, von denen sechs von der Hauptversammlung und sechs von den Arbeitnehmern nach den Vorschriften des Mitbestimmungsgesetzes gewählt werden.\n\n(2) Vorsitzender des Aufsichtsrats: {AG['arb_chair']}."),
                  ("§ 6 Hauptversammlung", "(1) Die ordentliche Hauptversammlung findet innerhalb der ersten fünf Monate eines jeden Geschäftsjahres statt.\n\n(2) Die Hauptversammlung ist einzuberufen, wenn das Wohl der Gesellschaft es erfordert oder die Einberufung von Aktionären, deren Anteile zusammen 5 % des Grundkapitals erreichen, unter Angabe des Zwecks und der Gründe verlangt wird."),
                  ("§ 7 Jahresabschluss und Gewinnverwendung", "(1) Der Vorstand hat innerhalb der gesetzlich vorgeschriebenen Frist den Jahresabschluss und den Lagebericht aufzustellen und dem Aufsichtsrat vorzulegen.\n\n(2) Der Aufsichtsrat hat den Jahresabschluss zu prüfen und der ordentlichen Hauptversammlung darüber zu berichten."),
              ], confidential=False)

    # Geschäftsordnung Vorstand
    make_docx(folder, "REA_Geschaeftsordnung_Vorstand_2022.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Geschäftsordnung für den Vorstand der Brennhagen Elektronik AG",
              [
                  ("Vorbemerkung", "Diese Geschäftsordnung regelt die interne Organisation und Arbeitsweise des Vorstands der Brennhagen Elektronik AG. Sie wurde vom Aufsichtsrat in seiner Sitzung vom 15. März 2022 beschlossen."),
                  ("§ 1 Zusammensetzung und Ressortverteilung", f"Der Vorstand setzt sich aktuell aus folgenden Mitgliedern zusammen:\n\n(1) {AG['ceo']} – Vorstandsvorsitzende (CEO): Gesamtleitung, Strategie, M&A, Investor Relations, Recht, Compliance.\n\n(2) {AG['coo']} – Chief Operating Officer (COO): Produktion, Supply Chain, Qualitätsmanagement, IT-Infrastruktur.\n\n(3) {AG['cfo']} – Chief Financial Officer (CFO): Finanzen, Controlling, Rechnungswesen, Treasury, Steuern.\n\n(4) {AG['cto']} – Chief Technology Officer (CTO): Forschung & Entwicklung, IP, Produktmanagement, Innovation."),
                  ("§ 2 Sitzungen", "(1) Der Vorstand tritt in der Regel wöchentlich zusammen. Außerordentliche Sitzungen können jederzeit einberufen werden.\n\n(2) Der Vorstand ist beschlussfähig, wenn alle Mitglieder ordnungsgemäß geladen und mindestens drei Mitglieder anwesend sind.\n\n(3) Beschlüsse werden mit einfacher Mehrheit der anwesenden Mitglieder gefasst. Bei Stimmengleichheit entscheidet die Stimme der Vorstandsvorsitzenden.\n\n(4) Über jede Vorstandssitzung ist ein Protokoll zu erstellen, das von der Sitzungsleiterin zu unterzeichnen ist."),
                  ("§ 3 Zustimmungspflichtige Geschäfte", "(1) Zustimmung des Aufsichtsrats ist erforderlich für:\n\na) Investitionen im Einzelfall mit einem Volumen von mehr als EUR 5.000.000;\nb) Aufnahme von Krediten über EUR 10.000.000;\nc) Erwerb oder Veräußerung von Beteiligungen;\nd) Eingehung von Dauerschuldverhältnissen mit einem Gesamtvolumen von mehr als EUR 2.000.000 p.a.\n\n(2) Zustimmung des Gesamtvorstands ist erforderlich für Einzelmaßnahmen mit einem Volumen von mehr als EUR 1.000.000."),
              ])

    # Geschäftsordnung Aufsichtsrat
    make_docx(folder, "REA_Geschaeftsordnung_Aufsichtsrat_2022.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Geschäftsordnung für den Aufsichtsrat der Brennhagen Elektronik AG",
              [
                  ("§ 1 Aufgaben und Zusammensetzung", f"(1) Der Aufsichtsrat überwacht die Geschäftsführung des Vorstands der {AG['name']}.\n\n(2) Der Aufsichtsrat besteht aus 12 Mitgliedern. Den Vorsitz führt {AG['arb_chair']}.\n\n(3) Der Aufsichtsrat bildet folgende Ausschüsse: Prüfungsausschuss, Personalausschuss, Nominierungsausschuss, Strategieausschuss."),
                  ("§ 2 Sitzungen und Beschlussfassung", "(1) Der Aufsichtsrat hält mindestens vier ordentliche Sitzungen pro Geschäftsjahr ab.\n\n(2) Der Aufsichtsrat ist beschlussfähig, wenn mindestens die Hälfte seiner Mitglieder, darunter der Vorsitzende oder sein Stellvertreter, an der Abstimmung teilnehmen.\n\n(3) Beschlüsse werden mit einfacher Mehrheit der abgegebenen Stimmen gefasst."),
                  ("§ 3 Berichtspflichten des Vorstands", "(1) Der Vorstand hat dem Aufsichtsrat regelmäßig, zeitnah und umfassend über alle für das Unternehmen relevanten Angelegenheiten zu berichten.\n\n(2) Der monatliche Bericht umfasst mindestens: Umsatz- und Ergebnisentwicklung, wesentliche Kundenentwicklungen, Investitionsvorhaben, Personalveränderungen auf Führungsebene sowie Risikobericht."),
              ])

    # Konzernstruktur-Übersicht
    make_docx(folder, "REA_Konzernstrukturschema_2024.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Konzernstrukturübersicht – Brennhagen Elektronik AG",
              [
                  ("Überblick", f"Die {AG['name']} (ISIN {AG['isin']}) ist die börsennotierte Muttergesellschaft des Brennhagen-Konzerns. Das Unternehmen ist im TecDAX notiert ({AG['stock']})."),
                  ("Konzernstruktur (vereinfachte Darstellung)", [
                      ["Gesellschaft", "Kürzel", "Sitz", "Land", "Beteiligung REA", "Funktion"],
                      [AG["name"], "REA", "Stuttgart", "DE", "Muttergesellschaft", "Börsennotierte Holding / AG"],
                  ] + [[s["name"], s["short"], s["city"], s["country"], "100 %", s["focus"]] for s in SUBS]),
                  ("Beteiligungshistorie", "Die Konzernstruktur hat sich wie folgt entwickelt:\n\n2018 – Erwerb der Brennhagen Czech s.r.o. (Brno), vormals eigenständig, Kaufpreis EUR 28,5 Mio.\n\n2019 – Erwerb der Brennhagen Hungary Kft. (Győr) im Rahmen einer Brownfield-Investition, Kaufpreis EUR 19,2 Mio.\n\n2021 – Umwandlung des Shanghai Joint Venture in eine 100%-Tochtergesellschaft, Abfindung des lokalen Partners über EUR 8,4 Mio."),
              ])

    # 8 Subsidiary Articles of Association
    for sub in SUBS:
        lang = "en" if sub["country"] in ["CN", "LU"] else "de"
        if lang == "de":
            make_docx(folder,
                      f"{sub['short']}_Gesellschaftsvertrag_2023.docx",
                      sub["name"], sub["city"], sub["hrb"],
                      f"Gesellschaftsvertrag der {sub['name']}",
                      [
                          ("§ 1 Firma und Sitz", f"(1) Die Gesellschaft führt die Firma {sub['name']}.\n\n(2) Sitz der Gesellschaft ist {sub['city']}."),
                          ("§ 2 Gegenstand", f"Gegenstand des Unternehmens ist {sub['focus']}. Die Gesellschaft erbringt ihre Leistungen überwiegend für den Brennhagen-Konzern sowie ausgewählte Dritte."),
                          ("§ 3 Stammkapital", f"Das Stammkapital der Gesellschaft beträgt EUR 5.000.000 (in Worten: fünf Millionen Euro). Es ist eingeteilt in 5.000 Geschäftsanteile zu je EUR 1.000. Alleinige Gesellschafterin ist die {AG['name']}."),
                          ("§ 4 Geschäftsführung", "Die Gesellschaft hat einen oder mehrere Geschäftsführer. Bei mehreren Geschäftsführern wird die Gesellschaft durch zwei Geschäftsführer gemeinsam oder durch einen Geschäftsführer gemeinsam mit einem Prokuristen vertreten."),
                          ("§ 5 Gesellschafterversammlung", "Gesellschafterversammlungen finden mindestens einmal jährlich statt. Beschlüsse werden mit einfacher Mehrheit der abgegebenen Stimmen gefasst, soweit nicht das Gesetz oder dieser Gesellschaftsvertrag eine größere Mehrheit verlangen."),
                          ("§ 6 Zustimmungspflichtige Geschäfte", "Der Zustimmung der Gesellschafterin bedürfen insbesondere: Investitionen > EUR 500.000, Eingehung von Verbindlichkeiten > EUR 1.000.000, Einstellung und Entlassung leitender Angestellter, Gewährung von Darlehen an Dritte."),
                          (None, sig_block_de(AG, f"Alleinige Gesellschafterin: {AG['name']}")),
                      ])
        else:
            make_docx(folder,
                      f"{sub['short']}_Articles_of_Association_2023.docx",
                      sub["name"], sub["city"], sub["hrb"],
                      f"Articles of Association – {sub['name']}",
                      [
                          ("Article 1 – Company Name and Registered Office", f"The company operates under the name {sub['name']} and is registered at {sub['hrb']}. The registered office is {sub['city']}, {sub['country']}."),
                          ("Article 2 – Business Purpose", f"The purpose of the company is {sub['focus']}. The company may engage in all transactions reasonably related to this purpose."),
                          ("Article 3 – Share Capital", f"The share capital amounts to EUR 5,000,000, held in its entirety by {AG['name']} (sole shareholder)."),
                          ("Article 4 – Management", "The company shall be managed by one or more directors appointed by the sole shareholder. Directors have full authority to act on behalf of the company within the limits set by these articles and applicable law."),
                          ("Article 5 – Governing Law", f"These articles are governed by the laws of {sub['country']}. All disputes shall be submitted to the competent courts at the registered office."),
                      ])

    # Intercompany Loan Agreements (pairs)
    ic_pairs = [
        (AG, SUBS[0], 50_000_000, "3,25 %"),
        (AG, SUBS[1], 15_000_000, "3,25 %"),
        (AG, SUBS[2], 20_000_000, "3,50 %"),
        (AG, SUBS[3], 12_000_000, "3,50 %"),
        (AG, SUBS[4], 18_000_000, "3,50 %"),
        (AG, SUBS[5],  8_000_000, "4,00 %"),
        (SUBS[6], SUBS[0], 30_000_000, "3,10 %"),
        (SUBS[6], SUBS[2], 25_000_000, "3,10 %"),
        (SUBS[6], SUBS[4], 22_000_000, "3,10 %"),
    ]
    for lender, borrower, amount, rate in ic_pairs:
        lname = lender["name"] if isinstance(lender, dict) and "name" in lender else lender["name"]
        bname = borrower["name"]
        make_docx(folder,
                  f"IC_Darlehensvertrag_{lender['short'] if 'short' in lender else 'REA'}_{borrower['short']}_2022.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Intercompany-Darlehensvertrag: {lname} → {bname}",
                  [
                      ("Parteien", f"Darlehensgeber: {lname}\nDarlehensnehmer: {bname}"),
                      ("§ 1 Darlehensbetrag", f"(1) Der Darlehensgeber gewährt dem Darlehensnehmer ein Darlehen in Höhe von {eur(amount)} (nachfolgend <<Darlehensbetrag<<).\n\n(2) Der Darlehensbetrag wird auf das Konto des Darlehensnehmers innerhalb von fünf Werktagen nach Unterzeichnung ausgezahlt."),
                      ("§ 2 Zinsen", f"(1) Das Darlehen ist mit {rate} p.a. zu verzinsen. Die Zinsen sind quartalsweise im Nachhinein fällig.\n\n(2) Der Zinssatz entspricht dem Fremdvergleichsgrundsatz gemäß § 1 AStG und wurde auf Basis eines Benchmarkings durch {AG['steuer']} ermittelt."),
                      ("§ 3 Laufzeit und Rückzahlung", "(1) Das Darlehen wird auf unbestimmte Zeit gewährt und kann von beiden Seiten mit einer Frist von zwölf (12) Monaten zum Jahresende gekündigt werden.\n\n(2) Der Darlehensnehmer kann das Darlehen jederzeit ganz oder teilweise vorzeitig zurückzahlen."),
                      ("§ 4 Verwendungszweck", f"Das Darlehen dient der Finanzierung von Betriebsmitteln und Investitionen im Rahmen des gewöhnlichen Geschäftsbetriebs des Darlehensnehmers ({borrower['focus']})."),
                      (None, std_governing_law_de()),
                  ], confidential=True)

    # SLAs Holding → Töchter
    for sub in SUBS[:6]:
        make_docx(folder,
                  f"SLA_Managementdienstleistungen_RHO_{sub['short']}_2023.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Service Level Agreement – Managementdienstleistungen\nBrennhagen Holding GmbH → {sub['name']}",
                  [
                      ("1. Vertragsgegenstand", f"Die Brennhagen Holding GmbH (Dienstleister) erbringt für die {sub['name']} (Empfänger) die nachfolgend beschriebenen Managementdienstleistungen. Die Leistungen umfassen insbesondere: Konzernweites Finanzcontrolling und Berichtswesen, HR-Strategie und Talent Management, Konzernweites IT-Servicemanagement, Rechts- und Compliance-Beratung, Treasury- und Liquiditätsmanagement."),
                      ("2. Leistungsbeschreibung (SLA-Parameter)", [
                          ["Leistung", "KPI", "Zielvorgabe", "Messmethode"],
                          ["Finanzberichterstattung", "Reporting-Deadline", "Tag 5 des Folgemonats", "Systemauswertung SAP"],
                          ["IT-Support (Level 2/3)", "Ticket-Lösungszeit", "< 4h (kritisch), < 24h (normal)", "ITSM-Ticketsystem"],
                          ["HR-Administration", "Stellenbesetzungszeit", "< 60 Tage", "HR-Dashboard"],
                          ["Compliance-Hotline", "Erreichbarkeit", "99 % während Geschäftszeiten", "ACD-Auswertung"],
                      ]),
                      ("3. Vergütung", f"(1) Die Vergütung für die Managementdienstleistungen beträgt EUR {random.randint(800, 2400) * 1000:,.0f} p.a. zzgl. der gesetzlichen Umsatzsteuer.\n\n(2) Die Vergütung wird monatlich in gleichen Raten abgerechnet. Die Rechnungsstellung erfolgt bis zum 5. des Folgemonats.\n\n(3) Die Vergütung wurde auf Basis einer Kostenplusrechnung (Cost-Plus 5 %) ermittelt und entspricht dem Fremdvergleichsgrundsatz."),
                      ("4. Laufzeit", f"Dieser SLA gilt ab dem 1. Januar 2023 und wird auf unbestimmte Zeit geschlossen. Er kann von jeder Partei mit einer Frist von sechs Monaten zum Jahresende gekündigt werden."),
                  ])

    # Transfer Pricing Documentation (overview)
    make_docx(folder, "REA_Verrechnungspreisdokumentation_Masterdatei_2023.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Verrechnungspreisdokumentation – Masterdatei (Master File) 2023",
              [
                  ("1. Überblick über den Konzern", f"Der Brennhagen-Konzern umfasst die {AG['name']} als Muttergesellschaft sowie {len(SUBS)} Tochtergesellschaften in Deutschland, Polen, Tschechien, Ungarn, China und Luxemburg. Der Konzernjahresumsatz 2023 beläuft sich auf {eur(AG['revenue_2023'])}."),
                  ("2. Wertschöpfungskette", "Die konzernweite Wertschöpfungskette lässt sich in folgende Bereiche unterteilen:\n\n(1) F&E (Brennhagen Software GmbH, München; Brennhagen Czech s.r.o., Brno): Entwicklung von Embedded Software, ADAS-Algorithmen und Systemarchitekturen.\n\n(2) Produktion (Brennhagen Elektronik GmbH, Heilbronn; Brennhagen Polska sp. z o.o., Katowice; Brennhagen Hungary Kft., Győr): Serienfertigung elektronischer Baugruppen.\n\n(3) Vertrieb (REA, Stuttgart; RCN, Shanghai): OEM-Kundenbetreuung und Auftragsabwicklung.\n\n(4) Konzernfinanzierung (Brennhagen Finance GmbH, Luxemburg): Intrakonzernfinanzierungen und Treasury-Dienstleistungen."),
                  ("3. Verrechnungspreismethoden", [
                      ["Transaktion", "Methode", "Begründung"],
                      ["Warenlieferungen (Produktion → Vertrieb)", "Wiederverkaufspreismethode (RPM)", "Vertriebsgesellschaft als Reseller ohne wesentliches IP-Risiko"],
                      ["Managementdienstleistungen", "Kostenaufschlagsmethode (Cost+5 %)", "Routinedienstleistungen ohne Alleinstellungsmerkmale"],
                      ["Intercompany-Darlehen", "Preisvergleichsmethode (CUP)", "Benchmarking externer Fremdkapitalzinsen"],
                      ["IP-Lizenzierung (F&E → Produktion)", "Gewinnaufteilungsmethode (PSM)", "Beide Parteien leisten wesentliche Beiträge zum IP"],
                  ]),
                  ("4. Dokumentationsanforderungen", f"Die Dokumentation wurde gemäß §§ 90 Abs. 3, 162 Abs. 3 AO sowie den OECD-Verrechnungspreisleitlinien 2022 erstellt. Erstellt durch: {AG['steuer']}, Frankfurt, Oktober 2023."),
              ], confidential=True)

    # Group Policies (25+)
    policies = [
        ("Konzernrichtlinie_Unterschriftenregelung_2023", "Unterschriften- und Zeichnungsberechtigungsrichtlinie"),
        ("Konzernrichtlinie_Reisekosten_2023", "Reisekosten- und Spesenrichtlinie"),
        ("Konzernrichtlinie_Einkauf_2023", "Beschaffungsrichtlinie – Global Sourcing"),
        ("Konzernrichtlinie_Compliance_2023", "Compliance-Richtlinie – Verhaltenskodex"),
        ("Konzernrichtlinie_AntiKorruption_2023", "Anti-Korruptions- und Anti-Bestechungsrichtlinie"),
        ("Konzernrichtlinie_Datenschutz_2023", "Datenschutzrichtlinie (DSGVO)"),
        ("Konzernrichtlinie_IT_Security_2023", "IT-Sicherheitsrichtlinie"),
        ("Konzernrichtlinie_Exportkontrolle_2023", "Exportkontroll- und Sanktionsrichtlinie"),
        ("Konzernrichtlinie_Kapitalmarkt_2023", "Kapitalmarkt-Compliance (MAR)"),
        ("Konzernrichtlinie_Finanzberichterstattung_2023", "Konzernrechnungslegungsrichtlinie (IFRS)"),
        ("Konzernrichtlinie_Treasury_2023", "Treasury- und Finanzierungsrichtlinie"),
        ("Konzernrichtlinie_Personalentwicklung_2023", "Personalentwicklungs- und Talentmanagementrichtlinie"),
        ("Konzernrichtlinie_Arbeitssicherheit_2023", "Arbeitssicherheit und Gesundheitsschutz (ISO 45001)"),
        ("Konzernrichtlinie_Umwelt_2023", "Umweltmanagementrichtlinie (ISO 14001)"),
        ("Konzernrichtlinie_Qualitaet_2023", "Qualitätsmanagementrichtlinie (IATF 16949)"),
        ("Konzernrichtlinie_LkSG_2023", "Lieferkettensorgfaltspflichtengesetz (LkSG) – Sorgfaltspflichten"),
        ("Konzernrichtlinie_Whistleblowing_2023", "Hinweisgebersystem und Whistleblowing-Richtlinie"),
        ("Konzernrichtlinie_Interessenkonflikte_2023", "Interessenkonflikte und Related-Party-Transactions"),
        ("Konzernrichtlinie_IP_Schutz_2023", "Schutz geistigen Eigentums und Geheimhaltung"),
        ("Konzernrichtlinie_Social_Media_2023", "Social-Media- und Kommunikationsrichtlinie"),
        ("Konzernrichtlinie_Nachhaltigkeit_ESG_2023", "ESG- und Nachhaltigkeitsrichtlinie"),
        ("Konzernrichtlinie_Krise_BCP_2023", "Business Continuity und Krisenmanagementrichtlinie"),
        ("Konzernrichtlinie_Outsourcing_2023", "Outsourcing- und Dienstleistersteuerungsrichtlinie"),
        ("Konzernrichtlinie_Datensicherung_2023", "Daten- und Informationssicherungsrichtlinie"),
        ("Konzernrichtlinie_Geschenke_2023", "Geschenke und Einladungen – Genehmigungsverfahren"),
    ]
    for fname, ptitle in policies:
        make_docx(folder, fname + ".docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  ptitle,
                  [
                      ("1. Geltungsbereich", f"Diese Richtlinie gilt für alle Mitarbeiterinnen und Mitarbeiter, Führungskräfte und Organmitglieder der {AG['name']} sowie aller Konzerngesellschaften weltweit. Ausnahmen bedürfen der ausdrücklichen Genehmigung der Vorstandsvorsitzenden."),
                      ("2. Zweck und Ziel", f"Zweck dieser Richtlinie ist die Sicherstellung eines einheitlichen und rechtskonformen Verhaltens im gesamten Brennhagen-Konzern in Bezug auf {ptitle}. Die Richtlinie dient dem Schutz des Unternehmens sowie aller Stakeholder vor Haftungsrisiken und Reputationsschäden."),
                      ("3. Regelungsinhalt", f"Die detaillierten Regelungen zu {ptitle} finden sich in den nachfolgenden Abschnitten. Alle Mitarbeitenden sind verpflichtet, diese Richtlinie zu lesen, zu verstehen und einzuhalten. Bei Unklarheiten ist die Compliance-Abteilung zu kontaktieren (compliance@brennhagen-elektronik.de)."),
                      ("4. Verstöße und Konsequenzen", "Verstöße gegen diese Richtlinie können disziplinarische Maßnahmen bis hin zur fristlosen Kündigung sowie zivil- und strafrechtliche Konsequenzen nach sich ziehen. Alle Verstöße sind über das Hinweisgebersystem zu melden."),
                      ("5. Gültigkeit", f"Diese Richtlinie tritt am 1. Januar 2023 in Kraft und ersetzt alle vorherigen Versionen. Sie wird mindestens jährlich überprüft und bei Bedarf aktualisiert. Verantwortlich: Chief Compliance Officer, {AG['name']}."),
                  ])

    # Vollmachten for key roles
    vollmachten = [
        (AG["ceo"],  "Generalvollmacht", "Alle Rechtsgeschäfte"),
        (AG["cfo"],  "Prokura (Einzelprokura)", "Finanzen und Rechnungswesen"),
        (AG["coo"],  "Handlungsvollmacht", "Operative Geschäfte bis EUR 5 Mio."),
        (AG["cto"],  "Handlungsvollmacht", "F&E und IP-Angelegenheiten"),
    ]
    for person, typ, scope in vollmachten:
        make_docx(folder, f"REA_Vollmacht_{sfn(person)}_2023.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"{typ} – {person}",
                  [
                      ("Vollmacht", f"Die {AG['name']}, vertreten durch ihren Vorstand, erteilt hiermit {person} (nachfolgend <<Bevollmächtigte/r>>) folgende Vollmacht:\n\nGegenstand: {scope}\nTyp: {typ}\nGültigkeitsbeginn: 1. Januar 2023\nBeschränkungen: Gemäß § 3 der Geschäftsordnung des Vorstands.\n\nDiese Vollmacht kann jederzeit durch schriftliche Erklärung des Vorstands widerrufen werden."),
                  ])


# ═══════════════════════════════════════════════════════════════════════════
# 01 – AG GESELLSCHAFTSRECHT
# ═══════════════════════════════════════════════════════════════════════════

def gen_01():
    folder = "01_AG_Gesellschaftsrecht"
    print(f"\n[01] AG Gesellschaftsrecht …")

    # HRB Auszüge
    for y in [2021, 2022, 2023, 2024]:
        make_docx(folder, f"REA_HRB_Auszug_{y}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Handelsregisterauszug – {AG['name']} – Stand {y}",
                  [
                      ("Registerdaten", f"Amtsgericht: {AG['amtsgericht']}\nRegisternummer: {AG['hrb']}\nRechtsform: Aktiengesellschaft\nGründungsdatum: {AG['founded']}\nEintragungsdatum: 15. April 1962\nGeschäftsjahr: 1. Januar bis 31. Dezember\nGrundkapital: {eur(AG['grundkapital'])}"),
                      ("Vertretungsberechtigte Organe", f"Vorstand (Gesamtvertretung je zwei Vorstandsmitglieder oder ein Vorstandsmitglied gemeinsam mit einem Prokuristen):\n{AG['ceo']}, Stuttgart – Vorstandsvorsitzende\n{AG['coo']}, Stuttgart – COO\n{AG['cfo']}, Stuttgart – CFO\n{AG['cto']}, Stuttgart – CTO\n\nAufsichtsrat:\n{AG['arb_chair']} (Vorsitzender)"),
                      ("Satzung / Änderungen", f"Letzte Satzungsänderung: Hauptversammlung {y-1}, eingetragen {y-1}-07-15."),
                  ])

    # Vorstand Bestellungsurkunden
    vorstand = [
        (AG["ceo"], "Vorstandsvorsitzende (CEO)", "15. März 2020", "14. März 2025"),
        (AG["coo"], "Chief Operating Officer (COO)", "1. September 2021", "31. August 2026"),
        (AG["cfo"], "Chief Financial Officer (CFO)", "1. Januar 2022", "31. Dezember 2026"),
        (AG["cto"], "Chief Technology Officer (CTO)", "1. April 2021", "31. März 2026"),
    ]
    for person, rolle, von, bis in vorstand:
        make_docx(folder, f"REA_Vorstand_Bestellung_{sfn(person)}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Bestellungsurkunde – {person} – {rolle}",
                  [
                      ("Bestellung", f"Der Aufsichtsrat der {AG['name']} hat in seiner Sitzung vom {von} beschlossen, {person} für die Zeit vom {von} bis {bis} zum {rolle} zu bestellen.\n\nDiese Bestellung erfolgt vorbehaltlich der Eintragung ins Handelsregister. Mit Antritt der Stelle gelten die Regelungen des gleichzeitig abgeschlossenen Vorstandsdienstvertrags sowie der Geschäftsordnung des Vorstands."),
                      ("Ressort", f"Das Vorstandsmitglied übernimmt das Ressort: {rolle}. Im Einzelnen umfasst das Ressort die in § 1 der Geschäftsordnung des Vorstands genannten Aufgabenbereiche."),
                      (None, f"\nStuttgart, den {von}\n\nIm Auftrag des Aufsichtsrats:\n\n________________________\n{AG['arb_chair']}\nAufsichtsratsvorsitzender"),
                  ])

    # Aufsichtsrat Mitglieder und Dienstverträge
    ar_members = [
        (AG["arb_chair"], "Vorsitzender", "Univ.-Prof., Wirtschaftsrecht"),
        ("Dipl.-Kfm. Heinrich Baumeister", "Stellv. Vorsitzender (Arbeitnehmervertreter)", "Betriebsratsvorsitzender REG"),
        ("Dr. Ingrid Schöllkopf", "Mitglied", "Finanzexpertin, ex-CFO Bosch"),
        ("Thomas Reinhardt MdB", "Mitglied", "Industriepolitiker"),
        ("Marlies Dürr", "Mitglied (Arbeitnehmervertreterin)", "IG Metall Gewerkschaftssekretärin"),
        ("Dr. Klaus Steinbrück", "Mitglied", "Venture Capital / Technologie"),
    ]
    for name, rolle, bg in ar_members:
        make_docx(folder, f"REA_AR_Dienstvertrag_{sfn(name)}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Aufsichtsratsvertrag – {name}",
                  [
                      ("Parteien", f"{AG['name']} (Gesellschaft) und {name} ({rolle})"),
                      ("§ 1 Amt und Aufgaben", f"{name} ist als {rolle} Mitglied des Aufsichtsrats der {AG['name']} tätig. Die Amtszeit beträgt vier Jahre ab Bestellung durch die Hauptversammlung. Aufgabenbeschreibung: {bg}."),
                      ("§ 2 Vergütung", "Die Vergütung der Aufsichtsratstätigkeit richtet sich nach dem Beschluss der Hauptversammlung vom 12. Mai 2022 (Vergütungssystem Aufsichtsrat). Die Grundvergütung beträgt EUR 60.000 p.a. Der Vorsitzende erhält das Dreifache, der stellvertretende Vorsitzende das Zweifache der Grundvergütung. Ausschussvorsitzende erhalten einen Aufschlag von EUR 15.000 p.a."),
                      ("§ 3 D&O-Versicherung", f"Die {AG['name']} unterhält eine D&O-Versicherung, in die alle Aufsichtsratsmitglieder eingeschlossen sind. Der Selbstbehalt entspricht den Vorgaben des Deutschen Corporate Governance Kodex."),
                  ])

    # HV Protokolle
    for y in [2021, 2022, 2023]:
        make_docx(folder, f"REA_HV_Protokoll_{y}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Niederschrift über die ordentliche Hauptversammlung {y}",
                  [
                      ("Teilnehmer und Beschlussfähigkeit", f"Die ordentliche Hauptversammlung der {AG['name']} fand am 12. Mai {y} in der ICS Stuttgart, Internationales Congresscenter Stuttgart, statt. Das Grundkapital in Höhe von {eur(AG['grundkapital'])} war zu {random.randint(62,81)} % vertreten."),
                      ("Tagesordnungspunkte", f"TOP 1 – Vorlage des festgestellten Jahresabschlusses und des Konzernabschlusses sowie des zusammengefassten Lage- und Konzernlageberichts für das Geschäftsjahr {y-1}\nTOP 2 – Beschlussfassung über die Verwendung des Bilanzgewinns\nTOP 3 – Entlastung des Vorstands\nTOP 4 – Entlastung des Aufsichtsrats\nTOP 5 – Wahl des Abschlussprüfers: {AG['wp']}\nTOP 6 – Billigung des Vergütungsberichts"),
                      ("Beschlüsse", f"Alle Tagesordnungspunkte wurden mit den erforderlichen Mehrheiten beschlossen. Dividende {y-1}: EUR {random.randint(40,80)} Cent je Aktie (Vorjahr: EUR {random.randint(30,50)} Cent)."),
                      (None, f"\nStuttgart, den 12. Mai {y}\n\nNotar: Dr. Friedrich Bergmann, Notar, Stuttgart"),
                  ])

    # AR Sitzungsprotokolle (4 per year × 3 years)
    for y in [2021, 2022, 2023]:
        for q in [1, 2, 3, 4]:
            m = q * 3
            make_docx(folder, f"REA_AR_Sitzungsprotokoll_{y}_Q{q}.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"Protokoll der Aufsichtsratssitzung – {y} Q{q}",
                      [
                          ("Anwesende", f"Vorsitz: {AG['arb_chair']}; anwesend: alle 12 Mitglieder\nSitzungsdatum: {ds(y, m, 15)}\nOrt: Vaihinger Straße 120, Stuttgart, Konferenzraum Vorstand"),
                          ("Tagesordnung", f"1. Bericht des Vorstands – Ergebnisentwicklung Q{q} {y}\n2. Genehmigung zustimmungspflichtiger Geschäfte\n3. Personalangelegenheiten\n4. Risikobericht\n5. Verschiedenes"),
                          ("Vorstandsbericht", f"Die Vorstandsvorsitzende {AG['ceo']} berichtete über die Geschäftsentwicklung im {q}. Quartal {y}. Der Konzernumsatz entwickelte sich plangemäß. Der CFO {AG['cfo']} präsentierte die aktuellen Finanzkennzahlen. Besprochen wurden Investitionsvorhaben und wesentliche Vertragsverhältnisse."),
                          ("Beschlüsse", f"Der Aufsichtsrat genehmigte die vorgelegten zustimmungspflichtigen Geschäfte einstimmig. Der nächste Sitzungstermin wurde auf {ds(y, min(m+3,12), 15)} festgesetzt."),
                      ])

    # Directors' shareholding statements
    for person, rolle, _, _ in vorstand:
        make_docx(folder, f"REA_Directors_Dealings_{sfn(person)}_2023.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Eigengeschäfte gemäß Art. 19 MAR – {person}",
                  [
                      ("Meldung Directors' Dealings", f"Name der meldepflichtigen Person: {person}\nFunktion: {rolle}\nArt der Transaktion: Kauf von {random.randint(500,5000)} Aktien der {AG['name']} (ISIN: {AG['isin']})\nDatum der Transaktion: {ds(2023, random.randint(3,11), random.randint(1,25))}\nKurs: EUR {random.uniform(48.0,58.0):.2f}\nGesamtvolumen: EUR {random.randint(30000,280000):,}\nMeldepflicht gemäß Art. 19 MAR erfüllt."),
                  ])


# ═══════════════════════════════════════════════════════════════════════════
# 02 – KONSOLIDIERTE FINANZEN
# ═══════════════════════════════════════════════════════════════════════════

def gen_02():
    folder = "02_Konsolidierte_Finanzen"
    print(f"\n[02] Konsolidierte Finanzen …")

    rev = {2020: AG["revenue_2020"], 2021: AG["revenue_2021"],
           2022: AG["revenue_2022"], 2023: AG["revenue_2023"]}
    ebitda = {2020: AG["ebitda_2020"], 2021: AG["ebitda_2021"],
              2022: AG["ebitda_2022"], 2023: AG["ebitda_2023"]}
    ebit = {2020: AG["ebit_2020"], 2021: AG["ebit_2021"],
            2022: AG["ebit_2022"], 2023: AG["ebit_2023"]}

    # Konzernabschluss xlsx (3 years × 5 sheets)
    for y in [2021, 2022, 2023]:
        r   = rev[y];    r0 = rev[y-1]
        eb  = ebitda[y]; et = ebit[y]
        cogs = int(r * 0.60)
        gp   = r - cogs
        sga  = int(r * 0.12)
        rd   = int(r * 0.065)
        dep  = eb - et
        net  = int(et * 0.72)
        sheets = [
            ("Gewinn- und Verlustrechnung",
             ["Position", f"GJ {y} (TEUR)", f"GJ {y-1} (TEUR)", "Δ %"],
             [
                 ["Umsatzerlöse", f"{r//1000:,}", f"{r0//1000:,}", f"{(r-r0)/r0*100:.1f} %"],
                 ["Herstellungskosten", f"-{cogs//1000:,}", "-", "-"],
                 ["Bruttoergebnis", f"{gp//1000:,}", "-", "-"],
                 ["Vertriebs- u. Verwaltungskosten", f"-{sga//1000:,}", "-", "-"],
                 ["F&E-Aufwendungen", f"-{rd//1000:,}", "-", "-"],
                 ["EBITDA", f"{eb//1000:,}", "-", "-"],
                 ["Abschreibungen", f"-{dep//1000:,}", "-", "-"],
                 ["EBIT", f"{et//1000:,}", "-", "-"],
                 ["Finanzergebnis", f"-{int(et*0.04)//1000:,}", "-", "-"],
                 ["EBT", f"{int(et*0.96)//1000:,}", "-", "-"],
                 ["Steuern", f"-{int(et*0.27)//1000:,}", "-", "-"],
                 ["Konzernjahresüberschuss", f"{net//1000:,}", "-", "-"],
             ],
             [30, 18, 18, 12]),
            ("Bilanz",
             ["Bilanzposition", f"{y}-12-31 (TEUR)", f"{y-1}-12-31 (TEUR)"],
             [
                 ["AKTIVA", "", ""],
                 ["Anlagevermögen", f"{int(r*0.35)//1000:,}", f"{int(r0*0.35)//1000:,}"],
                 ["Umlaufvermögen", f"{int(r*0.28)//1000:,}", f"{int(r0*0.28)//1000:,}"],
                 ["Bilanzsumme", f"{int(r*0.63)//1000:,}", f"{int(r0*0.63)//1000:,}"],
                 ["PASSIVA", "", ""],
                 ["Eigenkapital", f"{int(r*0.32)//1000:,}", f"{int(r0*0.30)//1000:,}"],
                 ["Langfristige Verbindlichkeiten", f"{int(r*0.18)//1000:,}", f"{int(r0*0.20)//1000:,}"],
                 ["Kurzfristige Verbindlichkeiten", f"{int(r*0.13)//1000:,}", f"{int(r0*0.13)//1000:,}"],
             ],
             [35, 20, 20]),
            ("Kapitalflussrechnung",
             ["Position", f"GJ {y} (TEUR)"],
             [
                 ["CF aus lfd. Geschäftstätigkeit", f"{int(eb*0.88)//1000:,}"],
                 ["CF aus Investitionstätigkeit", f"-{int(r*0.06)//1000:,}"],
                 ["CF aus Finanzierungstätigkeit", f"-{int(eb*0.25)//1000:,}"],
                 ["Nettoveränderung Liquidität", f"{int(eb*0.10)//1000:,}"],
             ],
             [40, 20]),
            ("Segmentberichterstattung",
             ["Segment", "Umsatz TEUR", "EBIT TEUR", "EBIT-Marge"],
             [
                 ["Powertrain Electronics", f"{int(r*0.38)//1000:,}", f"{int(et*0.40)//1000:,}", f"{int(et*0.40)/int(r*0.38)*100:.1f} %"],
                 ["ADAS & Safety", f"{int(r*0.30)//1000:,}", f"{int(et*0.35)//1000:,}", f"{int(et*0.35)/int(r*0.30)*100:.1f} %"],
                 ["EV & Energy", f"{int(r*0.20)//1000:,}", f"{int(et*0.15)//1000:,}", f"{int(et*0.15)/int(r*0.20)*100:.1f} %"],
                 ["Infotainment", f"{int(r*0.12)//1000:,}", f"{int(et*0.10)//1000:,}", f"{int(et*0.10)/int(r*0.12)*100:.1f} %"],
             ],
             [30, 18, 18, 15]),
            ("Anhang Konzern",
             ["Anhangposition", "Beschreibung"],
             [
                 ["Konsolidierungskreis", f"{len(SUBS)} vollkonsolidierte Tochtergesellschaften"],
                 ["Rechnungslegungsstandard", "IFRS (EU-Endorsed)"],
                 ["Abschlussprüfer", AG["wp"]],
                 ["Personalaufwand", f"{int(r*0.22)//1000:,} TEUR"],
                 ["Mitarbeiterzahl (Stichtag)", f"{AG[f'employees_{y}']:,}"],
             ],
             [35, 60]),
        ]
        make_xlsx(folder, f"REA_Konzernabschluss_{y}_IFRS.xlsx",
                  f"Konzernabschluss {y} – Brennhagen Elektronik AG", sheets)

    # WP Bestätigungsvermerke
    for y in [2021, 2022, 2023]:
        make_pdf(folder, f"REA_WP_Bestaetigung_{y}.pdf",
                 AG["wp"], "Klingelhöferstraße 18, 10785 Berlin", "WPK-Nr. 2063",
                 f"Bestätigungsvermerk des unabhängigen Abschlussprüfers – {AG['name']} {y}",
                 [
                     ("Prüfungsurteil", f"Wir haben den Konzernabschluss der {AG['name']}, Stuttgart, – bestehend aus der Konzernbilanz zum 31. Dezember {y}, der Konzern-Gewinn- und Verlustrechnung, der Konzern-Gesamtergebnisrechnung, dem Konzern-Eigenkapitalveränderungsrechnung und der Konzern-Kapitalflussrechnung für das dann endende Geschäftsjahr sowie dem Konzernanhang, einschließlich einer Zusammenfassung bedeutsamer Rechnungslegungsmethoden – geprüft.\n\nNach unserer Beurteilung aufgrund der bei der Prüfung gewonnenen Erkenntnisse vermittelt der beigefügte Konzernabschluss in allen wesentlichen Belangen ein den tatsächlichen Verhältnissen entsprechendes Bild der Vermögens- und Finanzlage des Konzerns zum 31. Dezember {y} sowie seiner Ertragslage für das dann endende Geschäftsjahr in Übereinstimmung mit den IFRS, wie sie in der EU anzuwenden sind."),
                     ("Grundlage für das Prüfungsurteil", f"Wir haben unsere Abschlussprüfung in Übereinstimmung mit § 317 HGB und der EU-Abschlussprüferverordnung (Nr. 537/2014) unter Beachtung der deutschen Grundsätze ordnungsmäßiger Abschlussprüfung, die vom Institut der Wirtschaftsprüfer (IDW) festgestellt wurden, durchgeführt. Wir sind von dem Konzern unabhängig in Übereinstimmung mit den deutschen handelsrechtlichen und berufsrechtlichen Vorschriften und haben unsere sonstigen deutschen Berufspflichten in Übereinstimmung mit diesen Anforderungen erfüllt."),
                     (None, f"\n\nBerlin, den {ds(y+1, 3, 15)}\n\n{AG['wp']}\n\n________________________\nWirtschaftsprüfer"),
                 ], confidential=False)

    # Monthly Financial Reports 2022 + 2023
    for y in [2022, 2023]:
        r_annual = rev[y]
        for m in range(1, 13):
            r_m = int(r_annual / 12 * random.uniform(0.85, 1.15))
            ebitda_m = int(r_m * random.uniform(0.098, 0.125))
            make_xlsx(folder, f"REA_Monatsbericht_{y}_{m:02d}.xlsx",
                      f"Konzern-Monatsbericht {MONTHS_DE[m]} {y}",
                      [
                          ("Ergebnisrechnung",
                           ["KPI", "Ist", "Plan", "Vorjahr", "Δ Plan %"],
                           [
                               ["Umsatz (TEUR)", f"{r_m//1000:,}", f"{r_annual//12//1000:,}", f"{rev[y-1]//12//1000:,}", f"{(r_m - r_annual//12) / (r_annual//12) * 100:.1f} %"],
                               ["EBITDA (TEUR)", f"{ebitda_m//1000:,}", f"{ebitda[y]//12//1000:,}", "-", "-"],
                               ["EBITDA-Marge", f"{ebitda_m/r_m*100:.1f} %", f"{ebitda[y]/r_annual*100:.1f} %", "-", "-"],
                               ["Auftragseingang (TEUR)", f"{int(r_m*1.08)//1000:,}", f"{r_annual//12//1000:,}", "-", "-"],
                           ],
                           [30, 15, 15, 15, 12]),
                      ])

    # Quarterly Earnings Releases Q1-Q4 2022/2023
    for y in [2022, 2023]:
        for q in [1, 2, 3, 4]:
            r_q = int(rev[y] / 4 * random.uniform(0.90, 1.10))
            make_pdf(folder, f"REA_Quartalsmitteilung_{y}_Q{q}.pdf",
                     AG["name"], AG["full_address"], f"ISIN: {AG['isin']}",
                     f"Quartalsmitteilung Q{q} {y} – Brennhagen Elektronik AG",
                     [
                         ("Quartalsergebnisse", f"Umsatz Q{q} {y}: TEUR {r_q//1000:,} (Vorjahr: TEUR {int(rev[y-1]/4)//1000:,})\nEBITDA: TEUR {int(r_q*0.11)//1000:,}\nEBITDA-Marge: {int(r_q*0.11)/r_q*100:.1f} %\n\nDie Geschäftsentwicklung im {q}. Quartal verlief plangemäß. Die starke Nachfrage nach EV-Komponenten (BatteryMS-12) sowie nach ADAS-Lösungen unterstützte das Wachstum."),
                         ("Ausblick", f"Der Vorstand bestätigt die Jahresprognose {y}: Umsatz EUR {rev[y]//1_000_000:.0f} Mio. (±3 %), EBITDA-Marge {ebitda[y]/rev[y]*100:.1f}–{ebitda[y]/rev[y]*100+1:.1f} %."),
                     ])

    # Budget 2024 and LRP
    make_xlsx(folder, "REA_Budget_2024.xlsx",
              "Konzernbudget 2024 – Brennhagen Elektronik AG",
              [
                  ("Budget 2024",
                   ["Position", "Budget 2024 (TEUR)", "Ist 2023 (TEUR)", "Δ %"],
                   [
                       ["Umsatz", f"{AG['revenue_2024e']//1000:,}", f"{AG['revenue_2023']//1000:,}", f"{(AG['revenue_2024e']-AG['revenue_2023'])/AG['revenue_2023']*100:.1f} %"],
                       ["EBITDA", f"{AG['ebitda_2024e']//1000:,}", f"{AG['ebitda_2023']//1000:,}", f"{(AG['ebitda_2024e']-AG['ebitda_2023'])/AG['ebitda_2023']*100:.1f} %"],
                       ["Capex", f"{int(AG['revenue_2024e']*0.055)//1000:,}", f"{int(AG['revenue_2023']*0.05)//1000:,}", "-"],
                       ["Mitarbeiterzahl (VZÄ)", "3.620", f"{AG['employees_2023']:,}", "+3,8 %"],
                   ],
                   [35, 20, 20, 12]),
              ])

    make_xlsx(folder, "REA_LRP_2024_2028.xlsx",
              "Langfristplanung 2024–2028 – Brennhagen Elektronik AG",
              [
                  ("Umsatzplanung",
                   ["Jahr", "Umsatz Plan (TEUR)", "EBITDA Plan (TEUR)", "EBITDA-Marge", "Capex (TEUR)", "Mitarbeiter"],
                   [
                       ["2024E", f"{AG['revenue_2024e']//1000:,}", f"{AG['ebitda_2024e']//1000:,}", f"{AG['ebitda_2024e']/AG['revenue_2024e']*100:.1f} %", f"{int(AG['revenue_2024e']*0.055)//1000:,}", "3.620"],
                       ["2025P", "698.000", "82.500", "11,8 %", "41.000", "3.780"],
                       ["2026P", "741.000", "91.000", "12,3 %", "44.000", "3.950"],
                       ["2027P", "789.000", "102.000", "12,9 %", "46.000", "4.100"],
                       ["2028P", "842.000", "114.000", "13,5 %", "48.000", "4.280"],
                   ],
                   [15, 22, 22, 15, 18, 15]),
              ])

    # Banking agreements
    make_docx(folder, "REA_Konsortialkredit_Rahmenvertrag_2022.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Konsortialkreditvertrag – Revolving Credit Facility EUR 150.000.000",
              [
                  ("Parteien", f"Kreditnehmer: {AG['name']}\nKreditgeber (Konsortium): {AG['bank']}\nFacility Agent: Deutsche Bank AG"),
                  ("§ 1 Kreditfazilität", "(1) Das Konsortium stellt dem Kreditnehmer eine revolvierende Kreditfazilität in Höhe von EUR 150.000.000 (in Worten: einhundertfünfzig Millionen Euro) zur Verfügung.\n\n(2) Die Fazilität dient der allgemeinen Unternehmensfinanzierung sowie der Finanzierung des Betriebsmittelbedarfs des Konzerns."),
                  ("§ 2 Zinsen", "(1) Der Zinssatz berechnet sich als EURIBOR (3M) zzgl. einer Marge von 1,25 % p.a.\n\n(2) Die Marge wird anhand eines Ratchets angepasst: bei Net Debt/EBITDA < 1,5x reduziert sich die Marge um 0,10 %."),
                  ("§ 3 Laufzeit und Rückzahlung", "(1) Die Fazilität hat eine Laufzeit bis zum 31. Dezember 2026 mit einer einmaligen Verlängerungsoption um zwei Jahre.\n\n(2) Jede Inanspruchnahme kann für Zeiträume von 1, 3 oder 6 Monaten gewählt werden."),
                  ("§ 4 Financial Covenants", "(1) Net Leverage Ratio: Net Debt / EBITDA ≤ 3,0x (getestet quartalsweise)\n(2) Interest Coverage Ratio: EBITDA / Zinsaufwand ≥ 4,0x\n(3) Eigenkapitalquote ≥ 30 %\n\nBei Verletzung eines Covenants ist das Konsortium berechtigt, die Fazilität zu kündigen und fällig zu stellen."),
                  (None, std_governing_law_de()),
              ], confidential=True)

    # Covenant compliance certificates (quarterly 2022+2023)
    for y in [2022, 2023]:
        for q in [1, 2, 3, 4]:
            lev = round(random.uniform(1.1, 2.2), 2)
            icr = round(random.uniform(5.0, 9.5), 1)
            make_docx(folder, f"REA_Covenant_Compliance_{y}_Q{q}.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"Covenant Compliance Certificate – Q{q} {y}",
                      [
                          ("Bescheinigung", f"Die {AG['name']} bestätigt die Einhaltung der Financial Covenants gemäß Konsortialkreditvertrag vom 15. März 2022 zum Stichtag {ds(y, q*3, 30)}:"),
                          ("Kennzahlen", [
                              ["Covenant", "Ist-Wert", "Grenzwert", "Status"],
                              ["Net Leverage Ratio", f"{lev}x", "≤ 3,0x", "✓ Eingehalten"],
                              ["Interest Coverage Ratio", f"{icr}x", "≥ 4,0x", "✓ Eingehalten"],
                              ["Eigenkapitalquote", f"{random.randint(33,42)} %", "≥ 30 %", "✓ Eingehalten"],
                          ]),
                          (None, f"\n\nStuttgart, den {ds(y, q*3+1 if q<4 else 1, 15)}\n\n{AG['cfo']}\nCFO, {AG['name']}"),
                      ])

    # Transfer pricing local files
    for sub in SUBS[:6]:
        make_docx(folder, f"TP_LocalFile_{sub['short']}_2023.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Verrechnungspreisdokumentation – Local File – {sub['name']} {sub['city']} 2023",
                  [
                      ("1. Gesellschaftsbeschreibung", f"{sub['name']} mit Sitz in {sub['city']} ist eine 100%-Tochtergesellschaft der {AG['name']}. Schwerpunkt: {sub['focus']}. Mitarbeiterzahl: {sub['employees']}. Umsatz {eur(sub['revenue'])}."),
                      ("2. Transaktionsübersicht", [
                          ["Transaktionstyp", "Volumen 2023 (TEUR)", "Methode"],
                          ["Warenlieferungen an REA", f"{sub['revenue']//1000//10:,}", "RPM"],
                          ["Managementgebühren an RHO", f"{sub['revenue']//1000//40:,}", "Cost+5 %"],
                          ["Zinsen auf IC-Darlehen", f"{sub['revenue']//1000//100:,}", "CUP"],
                      ]),
                      ("3. Fremdvergleichsanalyse", f"Auf Basis einer Datenbankrecherche (Bureau van Dijk – Orbis) wurden {random.randint(18,35)} vergleichbare Transaktionen identifiziert. Die ermittelten Verrechnungspreise liegen innerhalb der Interquartilsrange (IQR) der identifizierten Vergleichsunternehmen."),
                  ], confidential=True)


# ═══════════════════════════════════════════════════════════════════════════
# 03–08 – TOCHTERGESELLSCHAFTEN
# ═══════════════════════════════════════════════════════════════════════════

SUB_FOLDERS = {
    "REG": "03_Tochter_DE_REG_Heilbronn",
    "RSG": "04_Tochter_DE_RSG_Muenchen",
    "RPL": "05_Tochter_PL_Katowice",
    "RCZ": "06_Tochter_CZ_Brno",
    "RHU": "07_Tochter_HU_Gyoer",
    "RCN": "08_Tochter_CN_Shanghai",
    "RFI": "00_Konzernstruktur_Holding",
    "RHO": "00_Konzernstruktur_Holding",
}

def gen_subsidiaries():
    print(f"\n[03-08] Tochtergesellschaften …")
    for sub in SUBS[:6]:
        folder = SUB_FOLDERS[sub["short"]]
        s = sub["short"]
        sname = sub["name"]
        scity = sub["city"]
        shrb  = sub["hrb"]
        srev  = sub["revenue"]
        lang  = "en" if sub["country"] in ["CN", "LU", "HU", "CZ", "PL"] else "de"

        # Annual Accounts 3 years
        for y in [2021, 2022, 2023]:
            r = int(srev * random.uniform(0.88, 1.05))
            make_xlsx(folder, f"{s}_Jahresabschluss_{y}.xlsx",
                      f"Jahresabschluss {y} – {sname}",
                      [
                          ("GuV",
                           ["Position", f"{y} (TEUR)", f"{y-1} (TEUR)"],
                           [
                               ["Umsatzerlöse", f"{r//1000:,}", f"{int(r*0.93)//1000:,}"],
                               ["Herstellungskosten", f"-{int(r*0.62)//1000:,}", "-"],
                               ["Bruttoergebnis", f"{int(r*0.38)//1000:,}", "-"],
                               ["Betriebsergebnis (EBIT)", f"{int(r*0.072)//1000:,}", "-"],
                               ["Jahresüberschuss", f"{int(r*0.048)//1000:,}", "-"],
                           ],
                           [35, 20, 20]),
                          ("Bilanz",
                           ["Bilanzposition", f"{y}-12-31 (TEUR)"],
                           [
                               ["Anlagevermögen", f"{int(r*0.33)//1000:,}"],
                               ["Umlaufvermögen", f"{int(r*0.25)//1000:,}"],
                               ["Bilanzsumme", f"{int(r*0.58)//1000:,}"],
                               ["Eigenkapital", f"{int(r*0.28)//1000:,}"],
                               ["Fremdkapital", f"{int(r*0.30)//1000:,}"],
                           ],
                           [35, 20]),
                      ])

        # Monthly P&L 2023 (12 months)
        for m in range(1, 13):
            r_m = int(srev / 12 * random.uniform(0.82, 1.18))
            make_xlsx(folder, f"{s}_MonatsPL_{2023}_{m:02d}.xlsx",
                      f"Monatliche GuV {MONTHS_DE[m]} 2023 – {sname}",
                      [
                          ("Monat",
                           ["KPI", "Ist (TEUR)", "Plan (TEUR)", "Δ"],
                           [
                               ["Umsatz", f"{r_m//1000:,}", f"{srev//12//1000:,}", f"{(r_m - srev//12)//1000:+,}"],
                               ["Herstellungskosten", f"-{int(r_m*0.62)//1000:,}", "-", "-"],
                               ["EBIT", f"{int(r_m*0.072)//1000:,}", f"{int(srev/12*0.072)//1000:,}", "-"],
                           ],
                           [30, 15, 15, 12]),
                      ])

        # Intercompany invoices 2023 (12 months)
        for m in range(1, 13):
            inv_amount = int(srev / 12 * random.uniform(0.90, 1.10))
            make_pdf(folder, f"{s}_IC_Rechnung_2023_{m:02d}.pdf",
                     sname, scity, shrb,
                     f"Intercompany-Rechnung {MONTHS_DE[m]} 2023",
                     [
                         ("Rechnungsdetails", [
                             ["Position", "Betrag"],
                             ["Lieferung von Baugruppen/Leistungen gemäß Rahmenliefervertrag", eur(inv_amount)],
                             ["Abzgl. IC-Konzernskonto 5 %", eur(-int(inv_amount * 0.05))],
                             ["Nettobetrag", eur(int(inv_amount * 0.95))],
                             ["Umsatzsteuer (0 % – innergemeinschaftlich)", "EUR 0,00"],
                             ["Rechnungsbetrag", eur(int(inv_amount * 0.95))],
                         ]),
                         (None, f"\nZahlungsziel: 30 Tage netto\nBankverbindung: {AG['iban']} | {AG['bic']}"),
                     ])

        # Key employment contracts (5–8 per entity)
        mgrs = [
            f"Geschäftsführer/in {sname}",
            f"Produktionsleiter/in {scity}",
            f"Qualitätsmanagerin {scity}",
            f"Finanzcontroller {scity}",
            f"HR-Manager {scity}",
        ]
        for i, rolle in enumerate(mgrs):
            salary = random.randint(75, 160) * 1000
            make_docx(folder, f"{s}_Arbeitsvertrag_{i+1:02d}_{sfn(rolle)[:20]}_2022.docx",
                      sname, scity, shrb,
                      f"Anstellungsvertrag – {rolle}",
                      [
                          ("§ 1 Tätigkeitsbeschreibung", f"Der/Die Arbeitnehmer/in wird als {rolle} angestellt. Dienstort ist {scity}. Unmittelbar vorgesetzte Führungskraft ist der Vorstand der {AG['name']}."),
                          ("§ 2 Vergütung", f"(1) Das monatliche Bruttogehalt beträgt EUR {salary//12:,}.\n\n(2) Zusätzlich erhält der/die Arbeitnehmer/in eine jährliche variable Vergütung von bis zu {random.randint(15,30)} % des Bruttojahresgehalts bei 100 % Zielerfüllung.\n\n(3) Es gilt die betriebliche Altersversorgung gemäß Betriebsvereinbarung bAV."),
                          ("§ 3 Arbeitszeit", f"Die reguläre Arbeitszeit beträgt 40 Stunden pro Woche. Mehrarbeit ist mit der Vergütung abgegolten, soweit sie {random.randint(10,20)} Stunden pro Monat nicht übersteigt."),
                          ("§ 4 Urlaub", f"Der Jahresurlaub beträgt 30 Arbeitstage."),
                          ("§ 5 Geheimhaltung", "Der/Die Arbeitnehmer/in ist verpflichtet, alle ihm/ihr im Rahmen der Tätigkeit bekannt gewordenen vertraulichen Informationen dauerhaft geheim zu halten. Diese Verpflichtung gilt auch nach Beendigung des Arbeitsverhältnisses."),
                          ("§ 6 Nachvertragliches Wettbewerbsverbot", f"Nach Beendigung des Arbeitsverhältnisses gilt ein Wettbewerbsverbot für 12 Monate im Bereich Automotive Electronics. Die Karenzentschädigung beträgt 50 % der zuletzt bezogenen vertraglichen Vergütung."),
                      ])

        # Local tax document
        make_docx(folder, f"{s}_Steuerbescheid_KSt_{2023}.docx",
                  sname, scity, shrb,
                  f"Körperschaftsteuerbescheid 2022 – {sname}",
                  [
                      ("Steuerbescheid", f"Die zuständige Finanzbehörde hat für das Veranlagungsjahr 2022 einen Körperschaftsteuerbescheid in Höhe von {eur(int(srev * 0.02))} erlassen. Die Bescheide wurden in Abstimmung mit {AG['steuer']} geprüft."),
                      ("Vorbehalt der Nachprüfung", "Dieser Bescheid ergeht unter dem Vorbehalt der Nachprüfung (§ 164 Abs. 1 AO). Er kann daher innerhalb der Festsetzungsverjährungsfrist geändert werden."),
                  ])

        # Real estate / lease
        make_docx(folder, f"{s}_Mietvertrag_Betriebsgelaende_2020.docx",
                  sname, scity, shrb,
                  f"Miet-/Pachtvertrag Betriebsgelände – {sname}",
                  [
                      ("§ 1 Mietobjekt", f"Vermietet wird das Betriebsgelände in {scity}, Grundfläche {random.randint(8000,45000):,} m², bestehend aus Produktionshalle, Bürogebäude und Außenanlagen."),
                      ("§ 2 Mietzins", f"Die monatliche Miete beträgt EUR {random.randint(80, 280) * 1000:,} zzgl. gesetzlicher Umsatzsteuer. Die Miete wird jährlich um 2,0 % erhöht (Indexklausel, Basis CPI)."),
                      ("§ 3 Laufzeit", f"Der Mietvertrag läuft bis zum 31. Dezember 2030 mit zweimaliger Verlängerungsoption um je 5 Jahre."),
                      (None, std_governing_law_de()),
                  ])

        # Insurance policy
        make_docx(folder, f"{s}_Versicherungsnachweis_2023.docx",
                  sname, scity, shrb,
                  f"Versicherungsnachweis – {sname} – 2023",
                  [
                      ("Versicherungsübersicht", [
                          ["Sparte", "Versicherer", "Deckungssumme", "Prämie p.a."],
                          ["Betriebshaftpflicht", "Allianz SE", eur(50_000_000), eur(random.randint(80,200)*1000)],
                          ["Produkthaftpflicht", "AXA XL", eur(100_000_000), eur(random.randint(150,350)*1000)],
                          ["Feuerversicherung", "HDI Global", eur(int(srev*0.5)), eur(random.randint(60,180)*1000)],
                          ["Transportversicherung", "Generali", eur(10_000_000), eur(random.randint(30,90)*1000)],
                      ]),
                  ])

        # Local compliance doc
        make_docx(folder, f"{s}_Compliance_Report_{2023}.docx",
                  sname, scity, shrb,
                  f"Compliance-Bericht 2023 – {sname}",
                  [
                      ("1. Überblick", f"Im Berichtsjahr 2023 wurden in der {sname} keine wesentlichen Compliance-Verstöße festgestellt. Das lokale Compliance-Management-System (CMS) wurde auf Basis der konzernweiten Richtlinien implementiert."),
                      ("2. Schulungen", f"Im Jahr 2023 wurden {random.randint(85,98)} % aller Mitarbeitenden in den Pflichtschulungen (Code of Conduct, Datenschutz, Arbeitssicherheit) trainiert."),
                      ("3. Hinweisgebersystem", "Es gingen im Berichtsjahr 3 Hinweise über das konzernweite Hinweisgebersystem ein. Alle Meldungen wurden innerhalb der vorgeschriebenen Frist bearbeitet und abgeschlossen."),
                  ])


# ═══════════════════════════════════════════════════════════════════════════
# 09 – PERSONAL / HR
# ═══════════════════════════════════════════════════════════════════════════

NAMES_F = ["Anna Müller", "Laura Bauer", "Claudia Fischer", "Petra Koch", "Sabine Schreiber",
           "Monika Hoffmann", "Julia Hartmann", "Sandra Weber", "Nicole Lange", "Karin Schulz",
           "Dr. Elisabeth Vogel", "Prof. Renate Meyer", "Christina Wolf", "Birgit Braun",
           "Ursula Krause", "Ingrid Schmidt", "Heike Richter", "Ute Krämer", "Gisela Roth"]
NAMES_M = ["Thomas Müller", "Klaus Weber", "Stefan Hoffmann", "Michael Richter", "Andreas Becker",
           "Christian Koch", "Markus Schneider", "Frank Werner", "Jürgen Lange", "Hans-Peter Braun",
           "Dr. Wolfgang Fischer", "Ralf Zimmermann", "Dirk Brandt", "Uwe Schäfer", "Karl-Heinz Vogt",
           "Bernd Krüger", "Holger Maier", "Peter Kiefer", "Oliver Stein", "Carsten Engel",
           "Tobias Ott", "Sven Huber", "Björn Franke", "Lars Böhm", "Tim Schubert"]
ALL_NAMES = NAMES_F + NAMES_M

ROLES = [
    ("Vorstandsvorsitzende", 350_000, True),
    ("Chief Financial Officer", 320_000, True),
    ("Chief Operating Officer", 310_000, True),
    ("Chief Technology Officer", 300_000, True),
    ("VP Sales Automotive", 220_000, True),
    ("VP Engineering", 210_000, True),
    ("Head of Quality IATF", 180_000, False),
    ("Senior Project Manager", 135_000, False),
    ("Lead Software Engineer", 120_000, False),
    ("Principal Hardware Engineer", 115_000, False),
    ("Senior Controller", 110_000, False),
    ("HR Business Partner", 100_000, False),
    ("Supply Chain Manager", 105_000, False),
    ("Key Account Manager BMW", 115_000, False),
    ("Key Account Manager VW", 115_000, False),
    ("Production Manager", 125_000, False),
    ("Quality Engineer", 90_000, False),
    ("Software Engineer", 85_000, False),
    ("Embedded Software Developer", 88_000, False),
    ("ADAS Systems Engineer", 95_000, False),
    ("Test Engineer", 80_000, False),
    ("Purchasing Manager", 100_000, False),
    ("Financial Analyst", 85_000, False),
    ("Legal Counsel", 130_000, False),
    ("IP Manager", 120_000, False),
    ("Compliance Officer", 110_000, False),
    ("IT Security Analyst", 92_000, False),
    ("Customer Quality Manager", 95_000, False),
    ("Process Engineer", 85_000, False),
    ("Logistics Manager", 95_000, False),
]

def gen_09():
    folder = "09_Personal_HR"
    print(f"\n[09] Personal / HR …")

    # 60 Employment Contracts
    used_names = []
    for i in range(60):
        name = ALL_NAMES[i % len(ALL_NAMES)] + (f" {i//len(ALL_NAMES)+1}" if i >= len(ALL_NAMES) else "")
        rolle, salary, is_exec = ROLES[i % len(ROLES)]
        sub = SUBS[i % 6]
        used_names.append(name)
        make_docx(folder, f"AV_{i+1:03d}_{sfn(name)[:20]}_{sfn(rolle)[:20]}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Anstellungsvertrag – {name} – {rolle}",
                  [
                      ("§ 1 Arbeitgeber und Arbeitnehmer", f"Zwischen der {sub['name']}, {sub['city']} (Arbeitgeber), und {name} (Arbeitnehmer) wird folgender Arbeitsvertrag geschlossen."),
                      ("§ 2 Beginn und Funktion", f"(1) Das Arbeitsverhältnis beginnt am {ds(random.randint(2018,2023), random.randint(1,12), 1)}.\n\n(2) Der/Die Arbeitnehmer/in wird als {rolle} eingestellt. Er/Sie ist verpflichtet, alle mit dieser Tätigkeit zusammenhängenden Aufgaben gewissenhaft auszuführen."),
                      ("§ 3 Vergütung", f"(1) Das Jahresbruttofixgehalt beträgt EUR {salary:,}.\n\n(2) Zusätzlich wird eine variable Vergütung von bis zu {20 if is_exec else 15} % des Fixgehalts bei vollständiger Zielerfüllung gewährt.\n\n(3) {'Als Vorstandsmitglied/leitende Angestellte gilt: Teilnahme am LTI-Programm (Aktienoptionen) gemäß Beschluss HV 2022.' if is_exec else 'Keine LTI-Teilnahme in dieser Funktion.'}"),
                      ("§ 4 Arbeitszeit und Ort", f"(1) Dienstort ist {sub['city']}.\n\n(2) Wöchentliche Normalarbeitszeit: 40 Stunden.\n\n(3) Homeoffice bis zu {random.randint(2,3)} Tage pro Woche nach Abstimmung mit dem Vorgesetzten zulässig."),
                      ("§ 5 Urlaub und Krankheit", f"(1) Jahresurlaub: 30 Arbeitstage.\n\n(2) Im Krankheitsfall ist der Arbeitgeber ab dem ersten Tag zu informieren. Ein ärztliches Attest ist ab dem dritten Krankheitstag vorzulegen."),
                      ("§ 6 Geheimhaltung und Wettbewerbsverbot", "Der/Die Arbeitnehmer/in verpflichtet sich, während und nach der Anstellungszeit alle vertraulichen Informationen der Gesellschaft geheim zu halten. Das nachvertragliche Wettbewerbsverbot gilt für 12 Monate; die Karenzentschädigung beträgt 50 % der Fixvergütung."),
                      ("§ 7 Anwendbares Recht", f"Es gilt das Recht der Bundesrepublik Deutschland. Gerichtsstand: Stuttgart."),
                  ])

    # Group HR Policies (10+)
    hr_policies = [
        "REA_HR_Policy_Rekrutierung_2023",
        "REA_HR_Policy_Verguetung_2023",
        "REA_HR_Policy_Performance_Management_2023",
        "REA_HR_Policy_Weiterbildung_2023",
        "REA_HR_Policy_Diversity_Inclusion_2023",
        "REA_HR_Policy_Homeoffice_Mobiles_Arbeiten_2023",
        "REA_HR_Policy_Auslandsentsendung_2023",
        "REA_HR_Policy_Datenschutz_Mitarbeiter_2023",
        "REA_HR_Policy_Onboarding_2023",
        "REA_HR_Policy_Offboarding_2023",
        "REA_HR_Policy_Vorstandsverguetung_2023",
        "REA_HR_Policy_Whistleblowing_2023",
    ]
    for pname in hr_policies:
        title = pname.replace("REA_HR_Policy_", "").replace("_2023", "").replace("_", " ") + " – HR-Richtlinie 2023"
        make_docx(folder, pname + ".docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  title,
                  [
                      ("1. Geltungsbereich", f"Diese HR-Richtlinie gilt für alle Beschäftigten der {AG['name']} und aller Konzerngesellschaften, sofern nicht lokal abweichende Regelungen bestehen."),
                      ("2. Regelungsinhalt", f"Diese Richtlinie regelt die Grundsätze und Prozesse im Bereich {title}. Details zu Verfahrensabläufen, Zuständigkeiten und Fristen sind in den zugehörigen Arbeitsanweisungen (AA) festgelegt, die auf dem HR-Intranet verfügbar sind."),
                      ("3. Verantwortlichkeit", f"Verantwortlich für die Pflege und Aktualisierung dieser Richtlinie ist der Bereich Human Resources der {AG['name']}, Vaihinger Straße 120, Stuttgart."),
                  ])

    # Org charts (simplified as doc)
    make_docx(folder, "REA_Organigramm_Konzern_2024.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Konzern-Organigramm 2024 – Brennhagen Elektronik AG",
              [
                  ("Vorstand (Executive Board)", [
                      ["Funktion", "Inhaber/in", "Verantwortungsbereich"],
                      ["CEO", AG["ceo"], "Strategie, M&A, IR, Compliance"],
                      ["COO", AG["coo"], "Produktion, SCM, Qualität, IT"],
                      ["CFO", AG["cfo"], "Finanzen, Controlling, Treasury, Steuern"],
                      ["CTO", AG["cto"], "F&E, IP, Produktmanagement"],
                  ]),
                  ("Direkt berichtende Einheiten", [
                      ["Organisationseinheit", "Leiter/in", "Standort"],
                      ["Sales OEM (BMW, VW, MBZ, STE, CON)", "VP Sales", "Stuttgart"],
                      ["Engineering & R&D", "VP Engineering", "München / Brno"],
                      ["Manufacturing", "VP Manufacturing", "Heilbronn / Katowice / Győr"],
                      ["Quality & IATF", "Head of Quality", "Heilbronn"],
                      ["Finance & Controlling", "VP Finance", "Stuttgart"],
                      ["HR & Organization", "VP HR", "Stuttgart"],
                      ["Legal & IP", "General Counsel", "Stuttgart"],
                      ["IT & Digitalization", "VP IT", "Stuttgart"],
                  ]),
              ])

    # Salary structures
    make_xlsx(folder, "REA_Gehaltsstruktur_2023.xlsx",
              "Gehaltsstruktur und Vergütungsbänder 2023 – REA Konzern",
              [
                  ("Vergütungsbänder",
                   ["Jobfamilie", "Level", "Band Min (EUR)", "Band Mid (EUR)", "Band Max (EUR)", "Bonus Ziel %"],
                   [
                       ["Executive", "E1 (Vorstand)", "280.000", "340.000", "420.000", "40 %"],
                       ["Senior Management", "M5", "180.000", "220.000", "280.000", "25 %"],
                       ["Management", "M4", "130.000", "160.000", "200.000", "20 %"],
                       ["Senior Professional", "P5", "90.000", "110.000", "135.000", "15 %"],
                       ["Professional", "P4", "75.000", "90.000", "110.000", "10 %"],
                       ["Junior Professional", "P3", "58.000", "70.000", "85.000", "8 %"],
                       ["Associate", "P2", "45.000", "55.000", "68.000", "5 %"],
                   ],
                   [25, 20, 18, 18, 18, 15]),
              ])

    # Variable compensation plans
    make_docx(folder, "REA_Short_Term_Incentive_Plan_2023.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Short-Term Incentive Plan (STIP) 2023 – Brennhagen Elektronik AG",
              [
                  ("1. Grundsätze", "Der Short-Term Incentive Plan (STIP) der Brennhagen Elektronik AG legt die Grundsätze der variablen Jahresvergütung für alle Berechtigten fest."),
                  ("2. Zielparameter 2023", [
                      ["Zielkategorie", "Gewichtung", "KPI", "Zielwert 2023"],
                      ["Konzern-Finanzziele", "50 %", "Konzern-EBITDA", f"{eur(AG['ebitda_2023'])}"],
                      ["Konzern-Finanzziele", "20 %", "Konzern-Umsatz", f"{eur(AG['revenue_2023'])}"],
                      ["Individuelle Ziele", "30 %", "MBO-Ziele (3–5 individuelle)", "Vereinbart im Zielgespräch"],
                  ]),
                  ("3. Auszahlungslogik", "Bei 100 % Zielerreichung wird der vereinbarte Zielbonusbetrag zu 100 % ausbezahlt. Die Kurve reicht von 0 % (bei < 70 % Zielerreichung) bis 150 % (bei ≥ 130 % Zielerreichung)."),
              ])

    # LTI Plan
    make_docx(folder, "REA_Long_Term_Incentive_Plan_2022.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Long-Term Incentive Plan (LTIP) 2022–2024 – Brennhagen Elektronik AG",
              [
                  ("1. Programmübersicht", f"Der LTIP 2022–2024 gewährt ausgewählten Führungskräften virtuelle Aktienoptionen (Performance Share Units) auf Basis der Aktienkursentwicklung von {AG['name']} ({AG['isin']})."),
                  ("2. Berechtigte", "Teilnahmeberechtigt sind Vorstandsmitglieder sowie ausgewählte Senior Manager (VP-Ebene und ausgewählte Directors)."),
                  ("3. Performanceziele", [
                      ["KPI", "Gewichtung", "Baseline", "Stretch Ziel"],
                      ["Relativer TSR vs. SDAX", "50 %", "Median", "Top Quartil"],
                      ["EPS-Wachstum (3J CAGR)", "30 %", "8 %", "15 %"],
                      ["ESG-Score", "20 %", "ISS Rating B", "ISS Rating A"],
                  ]),
                  ("4. Erdienungszeitraum", "Der Erdienungszeitraum beträgt 3 Jahre (Vesting: 1. Januar 2025 für den Plan 2022)."),
              ])

    # Betriebsvereinbarungen (15+)
    bv_topics = [
        "Arbeitszeit_Gleitzeit", "Homeoffice_Mobiles_Arbeiten", "Betriebliche_Altersversorgung",
        "Kurzarbeit", "Datenschutz_IT_Systeme", "Gesundheitsfoerderung", "Qualifizierung_Weiterbildung",
        "Vergütungssystem_IGM_Tarif", "Zuschlagsregelung_Schichtarbeit", "Urlaub_Urlaubsgeld",
        "Vertrauensarbeitszeit_Führungskräfte", "Betriebliches_Eingliederungsmanagement",
        "Jubilaeumsregelung", "Leistungsbeurteilung_und_Prämien", "Mobilitätszuschuss_ÖPNV",
        "Essenszuschuss_Kantine",
    ]
    for bvt in bv_topics:
        make_docx(folder, f"REA_BV_{bvt}_2022.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Betriebsvereinbarung – {bvt.replace('_', ' ')}",
                  [
                      ("Präambel", f"Der Vorstand der {AG['name']} und der Gesamtbetriebsrat schließen folgende Betriebsvereinbarung zum Thema {bvt.replace('_', ' ')} ab."),
                      ("§ 1 Geltungsbereich", f"Diese Betriebsvereinbarung gilt für alle Arbeitnehmer der {AG['name']} und der deutschen Konzerngesellschaften, soweit keine eigenständige Betriebsvereinbarung der Tochtergesellschaft besteht."),
                      ("§ 2 Regelungsinhalt", f"Die Details zur Ausgestaltung von {bvt.replace('_', ' ')} sind in den Anlagen 1 bis 3 dieser Betriebsvereinbarung festgelegt. Die Regelungen entsprechen den einschlägigen tarifvertraglichen Vorgaben (IG Metall Tarifvertrag Metall- und Elektroindustrie Baden-Württemberg)."),
                      ("§ 3 Laufzeit", "Diese Betriebsvereinbarung tritt am 1. Januar 2022 in Kraft. Sie kann von jeder Seite mit einer Frist von drei Monaten zum Jahresende gekündigt werden."),
                      (None, f"\nStuttgart, den 15. Dezember 2021\n\nFür den Arbeitgeber: {AG['ceo']}\nFür den Gesamtbetriebsrat: Dipl.-Kfm. Heinrich Baumeister (Vorsitzender)"),
                  ])

    # Works Council election protocols
    for sub_short in ["REG", "RSG", "RHO"]:
        sub = next(s for s in SUBS if s["short"] == sub_short)
        make_docx(folder, f"{sub_short}_BR_Wahl_Protokoll_2022.docx",
                  sub["name"], sub["city"], sub["hrb"],
                  f"Protokoll – Betriebsratswahl 2022 – {sub['name']}",
                  [
                      ("Wahlergebnis", f"Am 15. März 2022 fand die reguläre Betriebsratswahl bei der {sub['name']} statt. Wahlbeteiligung: {random.randint(78,92)} %. Gewählte Betriebsratsmitglieder: {random.randint(9,15)}."),
                      ("Zusammensetzung", f"Vorsitzende/r: [Name], Stellv. Vorsitzende/r: [Name]. Der Betriebsrat hat sich gemäß § 26 BetrVG konstituiert."),
                  ])

    # HR KPI reports (quarterly 2023)
    for q in [1, 2, 3, 4]:
        make_xlsx(folder, f"REA_HR_KPI_Report_{2023}_Q{q}.xlsx",
                  f"HR-KPI-Bericht Q{q} 2023 – Brennhagen Elektronik AG",
                  [
                      ("HR KPIs",
                       ["KPI", "Q{q} Ist", "Q{q} Plan", "Δ"],
                       [
                           ["Headcount (Stichtag)", f"{AG['employees_2023'] + random.randint(-50,50):,}", f"{AG['employees_2023']:,}", "-"],
                           ["Fluktuation %", f"{random.uniform(4.5,7.5):.1f} %", "< 6 %", "-"],
                           ["Krankenquote %", f"{random.uniform(3.8,6.2):.1f} %", "< 5 %", "-"],
                           ["Time-to-Hire (Tage)", f"{random.randint(45,70)}", "< 60 Tage", "-"],
                           ["Schulungsstunden pro MA", f"{random.randint(18,35)}", "> 20 h", "-"],
                           ["Frauenanteil Führung %", f"{random.randint(26,36)} %", "> 30 %", "-"],
                       ],
                       [35, 18, 18, 12]),
                  ])

    # Workforce restructuring plan
    make_docx(folder, "REA_Restrukturierung_REG_Heilbronn_2023.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Restrukturierungsplan – Brennhagen Elektronik GmbH – Heilbronn 2023",
              [
                  ("1. Hintergrund", "Im Rahmen der strategischen Neuausrichtung (Fokus auf EV und ADAS) und zur Sicherstellung der Wettbewerbsfähigkeit des Standorts Heilbronn hat der Vorstand beschlossen, die Produktionsprozesse zu optimieren und den Personalbestand anzupassen."),
                  ("2. Maßnahmen", "(1) Abbau von ca. 80 Stellen im Bereich Legacy-Produktion (Verbrenner-ECUs) bis Q4 2024 durch Altersteilzeit, Abfindungen und natürliche Fluktuation.\n\n(2) Aufbau von ca. 45 neuen Stellen im Bereich EV-Komponenten (BatteryMS-12) ab Q1 2024.\n\n(3) Qualifizierungsmaßnahmen für 150 Beschäftigte im Bereich EV-Elektronik und ADAS."),
                  ("3. Sozialplan", "Der Sozialplan wurde am 1. Oktober 2023 mit dem Betriebsrat der REG Heilbronn vereinbart. Eckpunkte: Abfindungsformel 0,7 Monatsgehälter × Betriebszugehörigkeitsjahre, Transfergesellschaft für 60 Beschäftigte, Qualifizierungsbudget EUR 2,4 Mio."),
                  ("4. Kosten", f"Gesamte Restrukturierungskosten: EUR {random.randint(12,18)} Mio. (davon EUR {random.randint(8,12)} Mio. Abfindungen, EUR {random.randint(2,4)} Mio. Qualifizierung)."),
              ], confidential=True)


# ═══════════════════════════════════════════════════════════════════════════
# 10 – KAPITALMARKT / IR
# ═══════════════════════════════════════════════════════════════════════════

def gen_10():
    folder = "10_Kapitalmarkt_IR"
    print(f"\n[10] Kapitalmarkt / IR …")

    # Annual Reports
    for y in [2021, 2022, 2023]:
        r = AG[f"revenue_{y}"]
        eb = AG[f"ebitda_{y}"]
        make_pdf(folder, f"REA_Geschaeftsbericht_{y}.pdf",
                 AG["name"], AG["full_address"], f"ISIN {AG['isin']} | {AG['stock']}",
                 f"Geschäftsbericht {y} – Brennhagen Elektronik AG",
                 [
                     ("An die Aktionärinnen und Aktionäre", f"Sehr geehrte Aktionärinnen und Aktionäre,\n\ndas Geschäftsjahr {y} war für den Brennhagen-Konzern ein Jahr des nachhaltigen Wachstums. Mit einem Konzernumsatz von {eur(r)} und einem EBITDA von {eur(eb)} konnten wir unsere Marktposition als Technologieführer in der Automobilelektronik weiter ausbauen.\n\nIm Berichtsjahr standen insbesondere die Wachstumsbereiche ADAS (Advanced Driver Assistance Systems) und EV-Elektronik (BatteryMS-12) im Fokus unserer strategischen Investments.\n\n{AG['ceo']}\nVorstandsvorsitzende"),
                     ("Konzern in Zahlen", [
                         ["Kennzahl", f"{y}", f"{y-1}", "Δ %"],
                         ["Umsatz (EUR Mio.)", f"{r/1e6:.0f}", f"{AG[f'revenue_{y-1}']/1e6:.0f}", f"{(r-AG[f'revenue_{y-1}'])/AG[f'revenue_{y-1}']*100:.1f} %"],
                         ["EBITDA (EUR Mio.)", f"{eb/1e6:.1f}", f"{AG[f'ebitda_{y-1}']/1e6:.1f}", "-"],
                         ["EBITDA-Marge", f"{eb/r*100:.1f} %", "-", "-"],
                         ["Mitarbeiter", f"{AG.get(f'employees_{y}', 'n/a'):}", f"{AG.get(f'employees_{y-1}', 'n/a'):}", "-"],
                     ]),
                     ("Strategie und Ausblick", f"Der Brennhagen-Konzern sieht für die kommenden Jahre erhebliche Wachstumschancen im Bereich Elektrifizierung (EV) und autonomes Fahren (ADAS). Die Langfristplanung sieht ein Umsatzwachstum auf EUR 842 Mio. bis 2028 vor. Die F&E-Investitionen werden auf 7,0 % des Umsatzes erhöht."),
                 ])

    # Half-year reports
    for y in [2022, 2023]:
        make_pdf(folder, f"REA_Halbjahresbericht_{y}_H1.pdf",
                 AG["name"], AG["full_address"], f"ISIN {AG['isin']}",
                 f"Halbjahresbericht H1 {y} – Brennhagen Elektronik AG",
                 [
                     ("H1-Ergebnisse", f"Umsatz H1 {y}: {eur(int(AG[f'revenue_{y}']/2 * 0.96))}\nEBITDA H1: {eur(int(AG[f'ebitda_{y}']/2 * 0.94))}\nEBITDA-Marge: {AG[f'ebitda_{y}']/AG[f'revenue_{y}']*100:.1f} %\n\nDie Halbjahresergebnisse wurden in Übereinstimmung mit IAS 34 (Zwischenberichterstattung) aufgestellt und von {AG['wp']} einer prüferischen Durchsicht unterzogen."),
                     ("Segmentüberblick", [
                         ["Segment", "H1 Umsatz TEUR", "H1 EBIT TEUR"],
                         ["Powertrain Electronics", f"{int(AG[f'revenue_{y}']*0.38/2)//1000:,}", f"{int(AG[f'ebit_{y}']*0.40/2)//1000:,}"],
                         ["ADAS & Safety", f"{int(AG[f'revenue_{y}']*0.30/2)//1000:,}", f"{int(AG[f'ebit_{y}']*0.35/2)//1000:,}"],
                         ["EV & Energy", f"{int(AG[f'revenue_{y}']*0.20/2)//1000:,}", f"{int(AG[f'ebit_{y}']*0.15/2)//1000:,}"],
                         ["Infotainment", f"{int(AG[f'revenue_{y}']*0.12/2)//1000:,}", f"{int(AG[f'ebit_{y}']*0.10/2)//1000:,}"],
                     ]),
                 ])

    # Ad-hoc Mitteilungen (12+)
    adhoc_topics = [
        (2021, 3, "Jahresprognose 2021 angehoben – EBITDA-Marge über Erwartungen"),
        (2021, 9, "Strategische Partnerschaft mit BMW Group – Liefervertrag ADAS-Vision 4D"),
        (2022, 1, "Vorläufige Geschäftszahlen GJ 2021 – Rekordergebnis"),
        (2022, 5, "Kapitalerhöhung erfolgreich platziert – Erlös EUR 120 Mio."),
        (2022, 7, "Prognoseanpassung GJ 2022 nach positiver Halbjahresentwicklung"),
        (2022, 10, "Erwerb von 100% Anteilen an Brennhagen Electronics Shanghai"),
        (2023, 1, "Vorläufige Zahlen GJ 2022 – Umsatz EUR 573 Mio. (+14 %)"),
        (2023, 4, "Schuldscheindarlehen EUR 200 Mio. erfolgreich platziert"),
        (2023, 6, "Mehrjähriger Liefervertrag mit Volkswagen AG – BatteryMS-12 EV-Plattform"),
        (2023, 8, "Erhöhung der Jahresprognose GJ 2023"),
        (2023, 10, "Übernahme von 5 % eigener Aktien genehmigt (Aktienrückkaufprogramm)"),
        (2024, 2, "Vorläufige Zahlen GJ 2023 – Umsatz EUR 624 Mio. (+9 %)"),
        (2024, 3, "Neubesetzung CFO – Stefan Richter verlängert Vertrag um 4 Jahre"),
    ]
    for y, m, headline in adhoc_topics:
        make_pdf(folder, f"REA_Adhoc_{y}_{m:02d}_{sfn(headline[:30])}.pdf",
                 AG["name"], AG["full_address"], f"ISIN {AG['isin']}",
                 f"Ad-hoc-Mitteilung gemäß Art. 17 MAR – {ds(y, m, random.randint(5,25))}",
                 [
                     ("Meldung", f"Die {AG['name']} gibt bekannt:\n\n{headline}\n\nDetails: Der Vorstand der {AG['name']} hat heute beschlossen, folgende Ad-hoc-Mitteilung gemäß Art. 17 Abs. 1 der EU-Marktmissbrauchsverordnung (MAR) zu veröffentlichen. {headline}. Die vollständige Erläuterung ist auf der Investor-Relations-Webseite unter {AG['web']} verfügbar."),
                 ])

    # Analyst presentations (quarterly 2022+2023)
    for y in [2022, 2023]:
        for q in [1, 2, 3, 4]:
            make_pdf(folder, f"REA_Analysten_Praesentation_{y}_Q{q}.pdf",
                     AG["name"], AG["full_address"], f"ISIN {AG['isin']} | TecDAX",
                     f"Investor Presentation Q{q} {y} – Brennhagen Elektronik AG",
                     [
                         ("Investment Highlights", f"1. Technologieführer in der Automobilelektronik (ECU, ADAS, EV)\n2. {AG['kunde1_rev_share']} Umsatz mit BMW Group, diversifiziertes OEM-Portfolio\n3. Margenexpansion durch EV-Wachstum und Kostenoptimierung\n4. Solide Bilanz: Net Leverage {random.uniform(0.9,1.8):.1f}x, Free Cash Flow positiv\n5. TecDAX-Mitglied, ISIN {AG['isin']}"),
                         ("Finanzkennzahlen", [
                             ["Kennzahl", f"Q{q} {y} YTD", f"Q{q} {y-1} YTD", "Δ"],
                             ["Umsatz (TEUR)", f"{int(AG[f'revenue_{y}']*q/4)//1000:,}", f"{int(AG[f'revenue_{y-1}']*q/4)//1000:,}", f"+{(AG[f'revenue_{y}']-AG[f'revenue_{y-1}'])/AG[f'revenue_{y-1}']*100:.1f} %"],
                             ["EBITDA-Marge", f"{AG[f'ebitda_{y}']/AG[f'revenue_{y}']*100:.1f} %", f"{AG[f'ebitda_{y-1}']/AG[f'revenue_{y-1}']*100:.1f} %", "-"],
                         ]),
                     ])

    # Insider register
    make_xlsx(folder, "REA_Insiderregister_2023.xlsx",
              "Insiderregister gemäß Art. 18 MAR – Brennhagen Elektronik AG 2023",
              [
                  ("Insider-Liste",
                   ["Name", "Funktion", "Aufnahme", "Löschung", "Grund"],
                   [
                       [AG["ceo"], "CEO", "01.01.2023", "offen", "Permanent Insider"],
                       [AG["cfo"], "CFO", "01.01.2023", "offen", "Permanent Insider"],
                       [AG["coo"], "COO", "01.01.2023", "offen", "Permanent Insider"],
                       [AG["cto"], "CTO", "01.01.2023", "offen", "Permanent Insider"],
                       ["[Externe Berater]", "M&A Advisor", "15.06.2023", "31.08.2023", "Transaktion Q3/2023"],
                   ],
                   [30, 25, 15, 15, 30]),
              ])

    # Shareholder structure
    make_xlsx(folder, "REA_Aktionaersstruktur_2024.xlsx",
              "Aktionärsstruktur – Brennhagen Elektronik AG – Stand Q1 2024",
              [
                  ("Aktionärsstruktur",
                   ["Aktionär", "Anteil %", "Aktienanzahl", "Kategorie"],
                   [
                       ["BlackRock Inc.", "8,2 %", "984.000", "Institutionell"],
                       ["Vanguard Group", "5,9 %", "708.000", "Institutionell"],
                       ["DWS Investment GmbH", "4,1 %", "492.000", "Institutionell"],
                       ["Norges Bank Investment Mgmt.", "3,8 %", "456.000", "Staatsfonds"],
                       [AG["arb_chair"] + " (privat)", "1,2 %", "144.000", "Insider"],
                       ["Streubesitz (Freefloat)", "76,8 %", "9.216.000", "Freefloat"],
                   ],
                   [40, 12, 18, 20]),
              ])


# ═══════════════════════════════════════════════════════════════════════════
# 11 – VERTRIEB / OEM
# ═══════════════════════════════════════════════════════════════════════════

def gen_11():
    folder = "11_Vertrieb_OEM"
    print(f"\n[11] Vertrieb / OEM …")

    for cust in CUSTOMERS:
        cname = cust["name"]
        cshort = cust["short"]

        # Master Supply Agreement (full legal text)
        annual_vol = int(AG["revenue_2023"] * float(cust["share"].replace(" %","")) / 100)
        make_docx(folder, f"REA_{cshort}_Rahmenliefervertrag_2022.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Rahmenliefervertrag – {AG['name']} / {cname}",
                  [
                      ("Präambel / Vorbemerkungen", f"Die {AG['name']}, Vaihinger Straße 120, 70567 Stuttgart (nachfolgend <<Lieferant<<), und die {cname}, {cust['contact']} (nachfolgend <<Käufer>>), schließen den nachfolgenden Rahmenliefervertrag. Der Lieferant ist ein führender Hersteller von elektronischen Komponenten für die Automobilindustrie. Der Käufer beabsichtigt, langfristig Produkte des Lieferanten für seine Fahrzeugproduktion zu beziehen."),
                      (None, std_definitions_de()),
                      ("§ 2 Vertragsgegenstand", f"(1) Dieser Rahmenliefervertrag regelt die grundsätzlichen Bedingungen für die Lieferung von {', '.join([p['name'] for p in PRODUCTS])} sowie sonstiger elektronischer Baugruppen und Komponenten des Lieferanten.\n\n(2) Einzelbestellungen erfolgen durch verbindliche Abrufe des Käufers. Der Lieferant verpflichtet sich, Lieferzeiten von maximal {random.randint(8,16)} Wochen ab Abruf einzuhalten.\n\n(3) Planvolumen: Basierend auf den aktuellen Fahrzeugprogrammen des Käufers wird ein jährliches Liefervolumen von ca. {eur(annual_vol)} erwartet."),
                      ("§ 3 Qualitätsanforderungen", f"(1) Alle Lieferungen müssen die IATF 16949:2016-Zertifizierungsanforderungen erfüllen.\n\n(2) Der Lieferant ist verpflichtet, für alle sicherheitsrelevanten Teile einen PPAP (Production Part Approval Process) gemäß AIAG/VDA PPAP Level 3 einzureichen und vom Käufer genehmigen zu lassen.\n\n(3) Der Käufer ist berechtigt, Qualitätsaudits beim Lieferanten mit einer Vorlaufzeit von 10 Arbeitstagen durchzuführen.\n\n(4) Defektquote: Die vereinbarte ppm-Rate beträgt < {random.randint(5,20)} ppm bezogen auf angelieferte Teile."),
                      ("§ 4 Preise und Zahlungsbedingungen", f"(1) Die Preise werden jährlich im Rahmen eines strukturierten Jahresgespräches (Annual Price Review) verhandelt. Der Käufer erwartet eine jährliche Produktivitätsreduzierung von {random.uniform(1.5,3.5):.1f} % der Teilepreise.\n\n(2) Zahlungsbedingungen: 30 Tage netto ab Eingang der prüffähigen Rechnung.\n\n(3) Rechnungswährung: Euro. Wechselkursrisiken werden nach Maßgabe der separaten FX-Hedge-Vereinbarung geteilt."),
                      ("§ 5 Lieferung und Logistik", f"(1) Lieferbedingungen: DAP ({cust['contact']}) gemäß INCOTERMS 2020.\n\n(2) Der Lieferant ist verpflichtet, ein Kanban-/JIT-System zu betreiben. Abrufgenauigkeit: 95 %.\n\n(3) Der Lieferant hält einen Sicherheitsbestand von mindestens {random.randint(2,5)} Wochen Produktionsversorgung für sicherheitsrelevante Teile vor."),
                      ("§ 6 Gewährleistung und Haftung", "(1) Die Gewährleistungsfrist beträgt 3 Jahre ab Inbetriebnahme des Fahrzeugs durch den Endkunden, maximal 5 Jahre ab Lieferdatum.\n\n(2) Im Falle einer Rückrufaktion (Recall), die auf mangelhafte Teile des Lieferanten zurückzuführen ist, haftet der Lieferant für alle nachgewiesenen Kosten des Rückrufs, begrenzt auf EUR 25.000.000 pro Ereignis.\n\n(3) Die Produkthaftpflichtversicherung des Lieferanten muss eine Deckungssumme von mindestens EUR 50.000.000 pro Schadensereignis aufweisen."),
                      ("§ 7 Vertraulichkeit und IP", "(1) Alle vom Käufer zur Verfügung gestellten Zeichnungen, Spezifikationen und sonstigen technischen Unterlagen verbleiben im Eigentum des Käufers und dürfen nur zur Vertragserfüllung verwendet werden.\n\n(2) Verbesserungen und Weiterentwicklungen, die der Lieferant am käuferseitig bereitgestellten Know-how vornimmt, stehen im Eigentum des Käufers, sofern die Weiterentwicklung ausschließlich für den Käufer erfolgte."),
                      ("§ 8 Lieferkettensorgfaltspflicht (LkSG)", "(1) Der Lieferant versichert, alle einschlägigen Anforderungen des Lieferkettensorgfaltspflichtengesetzes (LkSG) einzuhalten.\n\n(2) Der Lieferant legt dem Käufer jährlich einen LkSG-Compliance-Bericht vor. Der Käufer ist berechtigt, Audits zur Überprüfung der Einhaltung der LkSG-Anforderungen durchzuführen."),
                      (None, std_governing_law_de()),
                      (None, sig_block_de(AG, cname)),
                  ], confidential=True)

        # Product-specific supply agreements (per product)
        for prod in PRODUCTS:
            vol = int(annual_vol * random.uniform(0.15, 0.35))
            make_docx(folder, f"REA_{cshort}_Liefervertrag_{prod['id']}_2023.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"Produktspezifischer Liefervertrag – {prod['name']} – {cname}",
                      [
                          ("1. Vertragsgegenstand", f"Dieser Vertrag ergänzt den Rahmenliefervertrag vom 1. März 2022 und regelt die spezifischen Bedingungen für die Lieferung von {prod['name']} an {cname}."),
                          ("2. Volumen und Preise", [
                              ["Produktbezeichnung", "Teilenummer", "Jahresvolumen (Stk)", "Einheitspreis (EUR)", "Jahresvolumen (EUR)"],
                              [prod["name"], f"REA-{prod['id']}-{cshort}-001", f"{vol//500:,}", f"{vol//1000 * 0.5 / (vol//1000 * 0.5 / (vol//500)):.2f}", f"{vol//1000:,}"],
                          ]),
                          ("3. Laufzeit", f"Dieser Vertrag gilt ab SOP (Start of Production) bis EOP (End of Production) des Fahrzeugprogramms {cname} – geschätzte Laufzeit bis {2027 + random.randint(0,3)}."),
                          ("4. Qualitätsanforderungen", f"Für {prod['name']} gelten die kundspezifischen Anforderungen (CSR) des {cname}, Version 2023-01, sowie die internen Qualitätsrichtlinien gemäß {AG['iatf']}."),
                      ])

        # Annual Price Review
        for y in [2022, 2023]:
            make_docx(folder, f"REA_{cshort}_Jahrespreisverhandlung_{y}.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"Jahrespreisverhandlung {y} – {cname}",
                      [
                          ("Verhandlungsergebnis", [
                              ["Produkt", f"Preis {y-1} (EUR)", f"Vereinbarter Preis {y} (EUR)", "Δ %"],

                          ] + [[p["name"], f"{random.uniform(40,200):.2f}", f"{random.uniform(38,195):.2f}", f"-{random.uniform(1.0,3.5):.1f} %"] for p in PRODUCTS
                          ]),
                          ("Protokoll", f"Das Jahrespreisprotokoll wurde am {ds(y, 11, random.randint(15,30))} zwischen {AG['name']} (Verhandlungsführer: VP Sales) und {cname} ({cust['contact']}) vereinbart. Beide Parteien bestätigen die vereinbarten Preisanpassungen durch Unterzeichnung."),
                      ])

        # Quality Agreement
        make_docx(folder, f"REA_{cshort}_Qualitaetsvereinbarung_2022.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Qualitätssicherungsvereinbarung – {AG['name']} / {cname}",
                  [
                      ("1. Gegenstand", f"Diese Qualitätssicherungsvereinbarung (QSV) regelt die Qualitätsanforderungen für alle Lieferbeziehungen zwischen {AG['name']} und {cname}."),
                      ("2. Qualitätsanforderungen", [
                          ["Anforderung", "Zielwert", "Messmethode"],
                          ["ppm-Rate geliefert", f"< {random.randint(5,15)} ppm", "Wareneingangskontrolle"],
                          ["PPAP-Genehmigungs-Quote", "100 %", "PPAP-Datenbank"],
                          ["8D-Response-Zeit", "< 24h (vorläufig) / 10 AT (final)", "Claim-Management-System"],
                          ["Liefertreue", "> 98,5 %", "EDI-Auswertung"],
                      ]),
                      ("3. Auditrecht", f"Der Käufer ist berechtigt, jährlich mindestens einen Qualitätsaudit beim Lieferanten durchzuführen. Der Lieferant verpflichtet sich, gefundene Abweichungen innerhalb von 30 Tagen zu korrigieren."),
                  ])

        # Customer Audit Report
        make_docx(folder, f"REA_{cshort}_Kundenaudit_Bericht_{2023}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Kundenaudit-Bericht – {cname} – {2023}",
                  [
                      ("Audit-Zusammenfassung", [
                          ["Kriterium", "Bewertung", "Anmerkung"],
                          ["Qualitätsmanagementsystem (IATF 16949)", "Bestanden", "Keine kritischen Findings"],
                          ["Prozessstabilität (Cpk > 1,67)", "Bestanden", f"Cpk = {random.uniform(1.7,2.2):.2f}"],
                          ["Dokumentation und Rückverfolgbarkeit", "Bestanden", "Vollständige Chargenrückverfolgung"],
                          ["Notfallversorgungsplan (Emergency Supply)", "Mit Auflagen", "Verbesserung angefordert bis Q2 2024"],
                      ]),
                      ("Maßnahmenplan", f"Offene Punkte wurden in einem Maßnahmenplan festgehalten. Zuständig: Head of Quality, {AG['name']}. Frist: 30. Juni 2024."),
                  ])

        # CSR Documentation
        make_docx(folder, f"REA_{cshort}_CSR_Kundspezifische_Anforderungen_2023.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Kundenspezifische Anforderungen (CSR) – {cname} – 2023",
                  [
                      ("Überblick CSR", f"Diese Dokumentation fasst die kundenspezifischen Anforderungen (Customer Specific Requirements, CSR) von {cname} zusammen, die über die Anforderungen der IATF 16949:2016 hinausgehen."),
                      ("Wesentliche Abweichungen / Ergänzungen", [
                          ["CSR-Punkt", "Anforderung", "REA-Umsetzung"],
                          ["PPAP-Level", "Level 3 (Standard)", "Vollständig implementiert"],
                          ["Functional Safety (ASPICE)", "ASPICE Level 2", f"Aktueller Level: {random.randint(2,3)}"],
                          ["Cybersecurity (ISO 21434)", "Conformance Required", "ISO 21434 Audit geplant Q3 2024"],
                          ["Sustainability Report", "CO2-Footprint je Teil", "Pilotprojekt gestartet 2024"],
                      ]),
                  ])

        # Long-term supply commitment
        make_docx(folder, f"REA_{cshort}_Langfristlieferzusage_{2022}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Langfristige Lieferzusage – {cname} – Laufzeit 2022–2028",
                  [
                      ("Lieferzusage", f"Die {AG['name']} gibt gegenüber {cname} folgende verbindliche Lieferzusage für die Lieferung von Automotive-Elektronikkomponenten für die Fahrzeugprogramme der Generation {2023+random.randint(1,3)} ab:\n\nZusagezeitraum: 1. Januar 2022 bis 31. Dezember 2028\nMindestzusage jährliches Liefervolumen: {eur(annual_vol)} (±15 %)\nProdukte: {', '.join([p['name'] for p in PRODUCTS[:2]])}\n\nDiese Zusage steht unter dem Vorbehalt der Kapazitätsverfügbarkeit sowie eines stabilen Halbleiterbeschaffungsmarkts. Force Majeure-Klauseln gemäß § 8 des Rahmenliefervertrags finden Anwendung."),
                  ])

        # Claim / Warranty Case
        make_docx(folder, f"REA_{cshort}_Claim_{2022}_{random.randint(1000,9999)}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Reklamationsbearbeitung – {cname} – Claim #{random.randint(1000,9999)}",
                  [
                      ("Reklamationsdetails", [
                          ["Feld", "Inhalt"],
                          ["Käufer", cname],
                          ["Teilebezeichnung", random.choice(PRODUCTS)["name"]],
                          ["Fehlerbeschreibung", random.choice(["Funktionsausfall ECU", "Fehlercode P0600 (CAN-Bus)", "Kalibrierungsfehler Sensor", "Kurzschluss Platine"])],
                          ["Betroffene Stückzahl", f"{random.randint(100,2500):,}"],
                          ["Beanstandeter Zeitraum", f"KW {random.randint(1,52)}/{2022}"],
                          ["Status", "Abgeschlossen – 8D-Bericht eingereicht"],
                      ]),
                      ("8D-Maßnahmen", "(D3) Sofortmaßnahme: 100 % Sichtprüfung aller Lagerbestände.\n\n(D4) Ursachenanalyse: Prozessfehler in Lötanlage Linie 3 (REG Heilbronn) – falsche Temperaturkurve.\n\n(D8) Nachhaltige Maßnahme: Neue Prozessparameter freigegeben, SPC-Überwachung eingeführt, Mitarbeiterschulung durchgeführt."),
                  ])


# ═══════════════════════════════════════════════════════════════════════════
# 12 – EINKAUF / LIEFERANTEN
# ═══════════════════════════════════════════════════════════════════════════

def gen_12():
    folder = "12_Einkauf_Lieferanten"
    print(f"\n[12] Einkauf / Lieferanten …")

    for sup in SUPPLIERS:
        sname = sup["name"]
        sshort = sup["short"]

        # Supply Agreement
        vol = random.randint(15, 60) * 1_000_000
        make_docx(folder, f"REA_{sshort}_Liefervertrag_Halbleiter_2022.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Liefervertrag Halbleiter-Komponenten – {AG['name']} / {sname}",
                  [
                      ("Präambel", f"{AG['name']} (Käufer) und {sname} (Lieferant) schließen diesen Liefervertrag über die Lieferung von Halbleiterkomponenten für die Serienproduktion in den Werken des Käufers."),
                      ("§ 1 Gegenstand", f"(1) {sname} liefert Halbleiter, Mikrocontroller, Leistungshalbleiter und verwandte Komponenten für den Einsatz in den Produktlinien ECU-900, ADAS-Vision 4D und BatteryMS-12 des Käufers.\n\n(2) Das Jahresvolumen beläuft sich auf ca. {eur(vol)}. Einzelabrufe erfolgen über das SAP-EDI-System des Käufers."),
                      ("§ 2 Preise und Preissicherung", f"(1) Preise basieren auf der aktuellen Preisliste ({sshort}-PL-2022-REA). Preisanpassungen bedürfen der schriftlichen Genehmigung des Käufers mit einer Vorlaufzeit von 6 Monaten.\n\n(2) Der Käufer und {sname} vereinbaren für sicherheitskritische Bauteile eine Preisbindung über 24 Monate."),
                      ("§ 3 Allokation und Liefersicherheit", f"(1) {sname} verpflichtet sich, die vereinbarten Jahresvolumina priorisiert zu liefern.\n\n(2) Im Falle von Allokationsengpässen (Semiconductor Shortage) hat der Käufer gemäß Allokationsvereinbarung (Exhibit C) einen bevorzugten Anteil von {random.randint(65,85)} % der verfügbaren Kapazitäten."),
                      ("§ 4 Qualität und Compliance", f"(1) Alle Produkte müssen AEC-Q100 (für Automotive ICs) entsprechen.\n\n(2) REACH- und RoHS-Konformität ist durch aktuelle Konformitätserklärungen nachzuweisen.\n\n(3) Konfliktmineralien-Meldepflicht gemäß OECD-Leitfaden und Dodd-Frank Act Section 1502."),
                      (None, std_governing_law_de()),
                  ])

        # Long-term Supply Agreement
        make_docx(folder, f"REA_{sshort}_Langfristvertrag_LTA_2023.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Long-Term Agreement (LTA) – {sname} – 2023–2026",
                  [
                      ("1. Rahmenbedingungen", f"Zur Sicherung der Halbleiterversorgung für die Wachstumsprogramme BatteryMS-12 (EV) und ADAS-Vision 4D vereinbaren die Parteien einen Long-Term Agreement (LTA) für die Zeitraum 2023–2026."),
                      ("2. Commitments", [
                          ["Jahr", "Mindestvolumen Käufer (EUR Mio.)", "Kapazitätszusage Lieferant (Wafer/Monat)"],
                          ["2023", f"{vol//1_000_000:.0f}", f"{random.randint(500,2000):,}"],
                          ["2024", f"{int(vol*1.1)//1_000_000:.0f}", f"{random.randint(600,2200):,}"],
                          ["2025", f"{int(vol*1.2)//1_000_000:.0f}", f"{random.randint(700,2400):,}"],
                          ["2026", f"{int(vol*1.3)//1_000_000:.0f}", f"{random.randint(800,2600):,}"],
                      ]),
                      ("3. Konventionalstrafen", "Bei Unterschreitung des Mindestvolumens durch den Käufer oder Unterschreitung der Kapazitätszusage durch den Lieferanten um mehr als 10 % fällt eine Konventionalstrafe von 5 % des betreffenden Jahresvolumens an."),
                  ])

        # Allocation Agreement (semiconductor shortage)
        make_docx(folder, f"REA_{sshort}_Allokationsvereinbarung_2021.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Allokationsvereinbarung Halbleitermangel 2021 – {sname}",
                  [
                      ("Hintergrund", f"Aufgrund der globalen Halbleiterkrise (2020–2022) hat {sname} seine verfügbaren Produktionskapazitäten unter seinen Kunden aufgeteilt. Dieser Nachtrag zum Liefervertrag regelt die Zuteilung für {AG['name']}."),
                      ("Allokationsquote", f"Im Zeitraum Januar 2021 bis Dezember 2022 erhält {AG['name']} {random.randint(65,80)} % seiner ursprünglich bestellten Mengen. Priorisierung erfolgt zugunsten sicherheitsrelevanter Bauteile (ECU-900, ADAS-Vision 4D)."),
                      ("Kompensation", "Aufgrund der Mindermenge werden die Preise einmalig um 2,5 % für das Lieferjahr 2022 reduziert. Beide Parteien verzichten auf weitergehende Schadensersatzansprüche."),
                  ])

        # Supplier Audit Reports (quarterly)
        for q in [1, 2, 3, 4]:
            make_docx(folder, f"REA_{sshort}_Lieferantenaudit_{2023}_Q{q}.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"Lieferantenaudit-Bericht – {sname} – Q{q} 2023",
                      [
                          ("Audit-Ergebnis", [
                              ["Auditkriterium", "Bewertung (1–5)", "Anmerkung"],
                              ["Qualitätsmanagementsystem", f"{random.randint(3,5)}", ""],
                              ["Liefertreue YTD", f"{random.randint(3,5)}", f"{random.randint(92,99)} %"],
                              ["Technische Kompetenz", f"{random.randint(3,5)}", ""],
                              ["Finanzielle Stabilität", f"{random.randint(3,5)}", ""],
                              ["Nachhaltigkeits-Compliance", f"{random.randint(2,5)}", ""],
                          ]),
                          ("Gesamtbewertung", f"Gesamtpunktzahl: {random.randint(72,95)}/100 – Status: {'A-Lieferant (bevorzugt)' if random.random() > 0.3 else 'B-Lieferant (qualifiziert)'}"),
                      ])

        # Annual Supplier Scorecard
        make_xlsx(folder, f"REA_{sshort}_Lieferanten_Scorecard_{2023}.xlsx",
                  f"Lieferanten-Scorecard {2023} – {sname}",
                  [
                      ("Scorecard",
                       ["KPI", "Gewichtung", "Zielwert", "Ist-Wert", "Punkte"],
                       [
                           ["Qualität ppm", "30 %", "< 10 ppm", f"{random.randint(1,15)} ppm", f"{random.randint(20,30)}"],
                           ["Liefertreue %", "25 %", "> 98 %", f"{random.uniform(94,99.5):.1f} %", f"{random.randint(18,25)}"],
                           ["Reaktionszeit Reklamationen", "15 %", "< 24h", f"{random.randint(12,36)}h", f"{random.randint(10,15)}"],
                           ["Preiskompetitivität", "15 %", "Marktpreis", "-2 % vs. Wettbewerb", f"{random.randint(10,15)}"],
                           ["Innovation / Roadmap", "10 %", "Halbjährliche Updates", "Erfüllt", f"{random.randint(7,10)}"],
                           ["Nachhaltigkeit / LkSG", "5 %", "Vollständige Dokumentation", "Erfüllt", f"{random.randint(4,5)}"],
                       ],
                       [35, 15, 20, 20, 10]),
                  ])

        # REACH/RoHS compliance certificates
        make_docx(folder, f"REA_{sshort}_REACH_RoHS_Konformitaet_2023.docx",
                  sname, f"{sup['country']}-Werk", f"Lieferant Nr. REA-L-{sshort}",
                  f"REACH- und RoHS-Konformitätserklärung – {sname} – 2023",
                  [
                      ("Erklärung", f"Wir, {sname}, erklären hiermit, dass alle an die {AG['name']} gelieferten Produkte die Anforderungen der EG-Richtlinie 2011/65/EU (RoHS) und der REACH-Verordnung (EG) Nr. 1907/2006 erfüllen.\n\nSVHC-Liste (Substances of Very High Concern): Keine der gelieferten Substanzen überschreitet den Grenzwert von 0,1 Gew.-%."),
                  ])

    # Component Specifications (per product × per supplier)
    for prod in PRODUCTS:
        for sup in SUPPLIERS[:3]:
            make_docx(folder, f"SPEC_{prod['id']}_{sup['short']}_Komponentenspez_2023.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"Komponentenspezifikation – {prod['name']} – Lieferant {sup['name']}",
                      [
                          ("1. Produktidentifikation", f"Produkt: {prod['name']}\nLieferant: {sup['name']}\nKäufer-Teilenummer: REA-{prod['id']}-{sup['short']}-001\nVersion: Rev. C ({ds(2023, 4, 15)})"),
                          ("2. Technische Anforderungen", [
                              ["Parameter", "Anforderung", "Prüfmethode"],
                              ["Betriebstemperatur", "-40 °C bis +125 °C (AEC-Q100 Grade 0)", "IEC 60068-2-2"],
                              ["Versorgungsspannung", f"{random.uniform(4.5,5.5):.1f} V ± 5 %", "Messprotokoll REA-QA-001"],
                              ["EMV-Anforderungen", "CISPR 25 Level 5", "EMV-Prüflabor"],
                              ["Lebensdauer", "> 15 Jahre / 200.000 km", "Berechnungsprotokoll"],
                              ["Feuchtigkeitsschutz", "IP6K9K (USCAR-2)", "Sprühwassertest"],
                          ]),
                      ])

    # Purchase Order archives (select months 2023)
    for m in [1, 4, 7, 10]:
        po_rows = [[f"PO-2023-{m:02d}-{i:04d}", random.choice(SUPPLIERS)["name"],
                    random.choice(PRODUCTS)["name"],
                    f"{random.randint(500,10000):,}", f"{random.uniform(0.5,15.0):.2f}",
                    eur(random.randint(50_000, 2_000_000))]
                   for i in range(1, random.randint(20, 40))]
        make_xlsx(folder, f"REA_Bestellarchiv_{2023}_{m:02d}.xlsx",
                  f"Bestellarchiv {MONTHS_DE[m]} 2023",
                  [
                      ("Bestellungen",
                       ["Bestellnummer", "Lieferant", "Material", "Menge", "Einheitspreis (EUR)", "Gesamtwert"],
                       po_rows,
                       [22, 35, 30, 12, 18, 18]),
                  ])

    # Import/Customs documents
    for m in [3, 6, 9, 12]:
        make_docx(folder, f"REA_Zollabfertigung_{2023}_{m:02d}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Zollabfertigungsnachweis – Import – {MONTHS_DE[m]} 2023",
                  [
                      ("Einfuhranmeldung", f"Herkunftsland: Asien / USA\nWarenwert: {eur(random.randint(500_000, 5_000_000))}\nZolltarifnummer: 8542.31 (Elektronische integrierte Schaltkreise)\nAngemeldetes Zollverfahren: 4000 (Überlassung zum zollrechtlich freien Verkehr)\nEinfuhrumsatzsteuer: {eur(random.randint(80_000, 900_000))}\nZollstelle: Zollamt Stuttgart"),
                  ])


# ═══════════════════════════════════════════════════════════════════════════
# 13 – IATF / QUALITÄT
# ═══════════════════════════════════════════════════════════════════════════

def gen_13():
    folder = "13_IATF_Qualitaet"
    print(f"\n[13] IATF / Qualität …")

    # IATF Quality Manual
    make_docx(folder, "REA_Qualitaetsmanagementhandbuch_IATF16949_2023.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              f"Qualitätsmanagementhandbuch – {AG['iatf']}",
              [
                  ("1. Geltungsbereich und Anwendungsbereich", f"Dieses Qualitätsmanagementhandbuch (QMH) beschreibt das Qualitätsmanagementsystem (QMS) der {AG['name']} und ihrer Produktionsgesellschaften gemäß {AG['iatf']}.\n\nGeltungsbereich: Entwicklung, Produktion und Vertrieb von elektronischen Steuergeräten (ECU), Kamerasystemen (ADAS), Infotainment-Modulen und Batteriemanagementsystemen für die Automobilindustrie.\n\nAusschlüsse: Kapitel 8.3 Entwicklung (teilweise) für zugekaufte Standardprodukte ohne kundenspezifische Anpassung."),
                  ("2. Normative Verweise", f"IATF 16949:2016 – Quality Management System Requirements for Automotive Production\nISO 9001:2015 – Quality Management Systems\nISO/TS 16949:2009 (Vorgänger, archiviert)\nVDA 6.3 – Prozessaudit\nVDA 19.1 – Technische Sauberkeit\nISO 26262 – Road Vehicles Functional Safety\nISO/SAE 21434 – Road Vehicles Cybersecurity"),
                  ("3. Begriffe und Definitionen", "Alle Begriffe aus ISO 9000:2015 und IATF 16949:2016 gelten wie dort definiert. Zusätzliche Begriffe:\n\nAPQP (Advanced Product Quality Planning): Strukturiertes Verfahren zur Produktqualitätsplanung.\nPPAP (Production Part Approval Process): Prozess zur Freigabe von Serienteilen.\nFMEA (Failure Mode and Effects Analysis): Methode zur Fehlermöglichkeits- und Einflussanalyse.\nSPC (Statistical Process Control): Statistische Prozesskontrolle."),
                  ("4. Kontext der Organisation", f"Die {AG['name']} ist ein Tier-1-Lieferant der Automobilindustrie mit Hauptkunden {AG['kunde1']}, {AG['kunde2']} und {AG['kunde3']}. Die wesentlichen Risiken umfassen: Halbleiterverfügbarkeit, Kundenforderungen (IATF, ISO 26262, AUTOSAR), geopolitische Risiken in der Lieferkette."),
                  ("5. Führung", f"Der Vorstand der {AG['name']} bekennt sich zur Qualitätspolitik: <<Wir liefern elektronische Lösungen in höchster Qualität – sicher, zuverlässig und auf Anhieb richtig – für die Mobilität von heute und morgen.>> Qualitätsverantwortlicher: Head of Quality ({AG['coo']} als Vorstandsverantwortlicher)."),
                  ("6. Planung", "Das Qualitätsrisikomanagement erfolgt auf Basis der FMEA-Methodik und APQP-Planung. Qualitätsziele werden jährlich im Rahmen der Managementbewertung festgelegt und auf alle Standorte kaskadiert."),
                  ("7. Support – Ressourcen", f"Das QMS umfasst alle Produktionsstandorte (Heilbronn, Katowice, Győr) sowie die Entwicklungsstandorte (München, Brno). Das zentrale QM-System wird über {AG['erp']} und {AG['plm']} betrieben."),
                  ("8. Betrieb – APQP und PPAP", "Der APQP-Prozess umfasst 5 Phasen: (1) Planung, (2) Produkt- und Prozessdesign, (3) Produkt- und Prozessvalidierung, (4) Serienanlauf, (5) Rückmeldung und Korrekturmaßnahmen. PPAP-Einreichungen erfolgen nach AIAG/VDA PPAP 1st Edition, Level 3."),
                  ("9. Bewertung der Leistung", "Die Qualitätsleistung wird monatlich über ein KPI-Dashboard überwacht: ppm-Rate (intern + extern), Auditbefunde, CAPA-Erledigungsquote, Kundenzufriedenheitsindex, First Pass Yield (FPY)."),
                  ("10. Verbesserung – CAPA-Prozess", "CAPA (Corrective and Preventive Actions) werden über das Qualitätsmanagementsystem (Q.wiki) verwaltet. Kritische CAPAs erfordern die Genehmigung des Head of Quality und Berichterstattung an den Vorstand."),
              ])

    # Core Tools – APQP, PPAP, FMEA, MSA, SPC (multiple docs each)
    core_tools = [
        ("APQP", "Advanced Product Quality Planning"),
        ("PPAP", "Production Part Approval Process"),
        ("FMEA", "Failure Mode and Effects Analysis"),
        ("MSA", "Measurement System Analysis"),
        ("SPC", "Statistical Process Control"),
        ("CP", "Control Plan"),
        ("FLS", "Functional Safety (ISO 26262)"),
        ("DFMEA", "Design FMEA"),
        ("PFMEA", "Process FMEA"),
    ]
    for tool_short, tool_name in core_tools:
        # General procedure
        make_docx(folder, f"REA_{tool_short}_Verfahrensanweisung_2023.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Verfahrensanweisung {tool_name} ({tool_short})",
                  [
                      ("1. Zweck", f"Diese Verfahrensanweisung beschreibt die Anwendung von {tool_name} ({tool_short}) im Rahmen des Qualitätsmanagementsystems der {AG['name']} gemäß {AG['iatf']}."),
                      ("2. Verantwortlichkeiten", [
                          ["Rolle", "Verantwortung"],
                          ["Head of Quality", f"Genehmigung und Freigabe von {tool_short}-Dokumenten"],
                          ["Quality Engineer", f"Erstellung und Pflege von {tool_short}-Dokumenten"],
                          ["Projektleiter", f"Initiierung und Koordination des {tool_short}-Prozesses"],
                          ["Entwicklungsingenieure", f"Technische Inputs für {tool_short}"],
                      ]),
                      ("3. Prozessablauf", f"Die Anwendung von {tool_name} folgt dem {tool_short}-Standard gemäß AIAG/VDA-Handbuch sowie den kundspezifischen Anforderungen (CSR) der OEM-Kunden."),
                  ])
        # Per product
        for prod in PRODUCTS:
            make_docx(folder, f"REA_{tool_short}_{prod['id']}_2023.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"{tool_name} – {prod['name']} – Rev. C",
                      [
                          ("Dokumentenstatus", [
                              ["Feld", "Inhalt"],
                              ["Produkt", prod["name"]],
                              ["Tool", tool_name],
                              ["Version", "Rev. C"],
                              ["Datum", ds(2023, random.randint(1,10), random.randint(1,25))],
                              ["Status", "Freigegeben"],
                          ]),
                          ("Hauptinhalt", f"Vollständige {tool_name}-Dokumentation für {prod['name']} gemäß {AG['iatf']} und OEM-CSR. Erstellt durch: Quality Engineering Team, {AG['name']}. Überprüft und freigegeben durch: Head of Quality."),
                      ])

    # PPAP Submissions per customer/product
    for cust in CUSTOMERS[:3]:
        for prod in PRODUCTS[:3]:
            make_docx(folder, f"PPAP_{cust['short']}_{prod['id']}_2022_Level3.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"PPAP-Einreichung Level 3 – {prod['name']} – {cust['name']}",
                      [
                          ("PPAP Checklist", [
                              ["#", "Element", "Status", "Anmerkung"],
                              ["1", "Designunterlagen", "✓ Vorliegend", ""],
                              ["2", "Technische Änderungsdokumente", "✓ N/A", "Kein Engineering Change"],
                              ["3", "Zustimmung des Kunden (Customer Engineering Approval)", "✓ Erhalten", f"Datum: {ds(2022, 8, 15)}"],
                              ["4", "DFMEA", "✓ Vorliegend", "Rev. D"],
                              ["5", "Prozessablaufplan", "✓ Vorliegend", ""],
                              ["6", "PFMEA", "✓ Vorliegend", "Rev. E"],
                              ["7", "Lenkungsplan (Control Plan)", "✓ Vorliegend", ""],
                              ["8", "MSA-Studie", "✓ Bestanden", f"Gage R&R = {random.uniform(8,15):.1f} %"],
                              ["9", "Ergebnisse von Dimensionsmessungen", "✓ Bestanden", "100 % in Toleranz"],
                              ["10", "Material- / Leistungsprüfergebnisse", "✓ Bestanden", ""],
                              ["11", "Erstmusterprüfbericht (EMPB)", "✓ Vorliegend", ""],
                              ["12", "Muster-Teile", "✓ Übergeben", f"{random.randint(5,30)} Stück"],
                              ["13", "Erstmuster-Prüfbericht Kunde", "✓ Genehmigt", f"Datum: {ds(2022, 9, 20)}"],
                              ["14", "Behältnis- und Verpackungsnachweis", "✓ Bestanden", ""],
                              ["15", "PSW (Part Submission Warrant)", "✓ Ausgefüllt", "Level 3"],
                          ]),
                      ])

    # Internal + External Audit Reports
    for y in [2022, 2023]:
        # Internal
        make_docx(folder, f"REA_Internes_IATF_Audit_{y}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Internes IATF 16949 Systemaudit – {y}",
                  [
                      ("Auditergebnis", [
                          ["Auditkapitel", "Befunde", "Bewertung"],
                          ["Kap. 4 – Kontext", "Keine", "Konform"],
                          ["Kap. 5 – Führung", "1 Beobachtung (Qualitätspolitik Aktualisierung)", "Konform"],
                          ["Kap. 6 – Planung", "Keine", "Konform"],
                          ["Kap. 7 – Support", f"{'1 Minor Finding: Schulungsnachweis fehlt' if random.random() > 0.5 else 'Keine'}", "Konform/Geringfügig"],
                          ["Kap. 8 – Betrieb", "2 Beobachtungen (APQP-Dokumentation)", "Konform"],
                          ["Kap. 9 – Bewertung", "Keine", "Konform"],
                          ["Kap. 10 – Verbesserung", "Keine", "Konform"],
                      ]),
                      ("Gesamtbewertung", f"Das Qualitätsmanagementsystem erfüllt die Anforderungen der IATF 16949:2016. Die identifizierten Minorfindings wurden in den CAPA-Plan aufgenommen. Nächstes Audit: {y+1}."),
                  ])
        # External (KPMG/TÜV)
        make_docx(folder, f"REA_Externes_IATF_Zertifikat_{y}.docx",
                  "TÜV Rheinland Cert GmbH", "Am Grauen Stein, 51105 Köln", "TÜV-Zertifizierstelle",
                  f"IATF 16949 Zertifikat – {AG['name']} – {y}",
                  [
                      ("Zertifikat", f"Hiermit wird bescheinigt, dass\n\n{AG['name']}\n{AG['full_address']}\n\nein Qualitätsmanagementsystem betreibt, das den Anforderungen der IATF 16949:2016 entspricht.\n\nGültigkeit: {ds(y, 6, 1)} bis {ds(y+3, 5, 31)}\nZertifikat-Nr.: IATF-TÜV-{y}-REA-{random.randint(10000,99999)}\n\nAuditierter Geltungsbereich: Entwicklung, Produktion und Vertrieb von Automotive-Elektronikkomponenten."),
                  ])

    # Management Review Minutes
    for sub in SUBS[:4]:
        for y in [2022, 2023]:
            make_docx(folder, f"MR_{sub['short']}_Managementbewertung_{y}.docx",
                      sub["name"], sub["city"], sub["hrb"],
                      f"Managementbewertung {y} – {sub['name']}",
                      [
                          ("1. Eingaben (Inputs)", f"Qualitätspolitik und -ziele, Kundenzufriedenheit (ppm-Rate: {random.randint(1,15)} ppm), Prozessleistung (FPY: {random.uniform(97,99.8):.1f} %), Audit-Ergebnisse, CAPA-Status, Ressourcenplanung."),
                          ("2. Ergebnisse (Outputs)", [
                              ["Thema", "Beschluss", "Verantwortlich", "Termin"],
                              ["Qualitätsziele 2024", f"ppm < {random.randint(5,10)}, FPY > {random.uniform(98,99.5):.1f} %", "Head of Quality", "31.01.2024"],
                              ["CAPA offene Punkte", "Alle offenen CAPAs bis Q1 2024 schließen", "Quality Engineering", "31.03.2024"],
                              ["Ressourcen", "2 zusätzliche Quality Engineers einstellen", "HR", "30.06.2024"],
                          ]),
                      ])

    # CAPA Records
    for i in range(1, 21):
        make_docx(folder, f"CAPA_{2023}_{i:04d}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"CAPA #{2023}-{i:04d} – {'Kundenreklamation' if i % 3 == 0 else 'Interner Audit-Befund'}",
                  [
                      ("CAPA Details", [
                          ["Feld", "Inhalt"],
                          ["CAPA-Nr.", f"CAPA-2023-{i:04d}"],
                          ["Typ", "Korrekturmaßnahme" if i % 2 == 0 else "Vorbeugungsmaßnahme"],
                          ["Produkt", random.choice(PRODUCTS)["name"]],
                          ["Problem", random.choice(["ppm-Überschreitung", "Prozessabweichung Löten", "Dokumentationslücke", "Kundenbeanstandung 8D"])],
                          ["Ursache", "Parameterabweichung in Produktionslinie 2"],
                          ["Maßnahme", "Neueinstellung Prozessparameter + SPC-Überwachung"],
                          ["Status", random.choice(["Offen", "In Bearbeitung", "Abgeschlossen", "Abgeschlossen"])],
                          ["Termin", ds(2023, random.randint(6,12), 30)],
                      ]),
                  ])

    # 8D Reports
    for i in range(1, 16):
        cust = random.choice(CUSTOMERS)
        prod = random.choice(PRODUCTS)
        make_docx(folder, f"8D_{2023}_{i:03d}_{cust['short']}_{prod['id']}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"8D-Report – {cust['name']} – {prod['name']} – #{2023}-{i:03d}",
                  [
                      ("D1 – Team", f"8D-Team-Leiter: Quality Engineer {AG['name']}\nMitglieder: Produktion, Entwicklung, Einkauf, Logistik"),
                      ("D2 – Problembeschreibung", f"Fehlerbeschreibung: {random.choice(['CAN-Kommunikationsfehler', 'Lötfehler BGA', 'Kalibrierungsabweichung', 'EMV-Versagen'])}\nEntdeckungsort: {random.choice(['Wareneingang Kunde', 'Fahrzeugtest', 'EOL-Prüfstand REA'])}\nBetroffene Stückzahl: {random.randint(50,5000):,}"),
                      ("D3 – Sofortmaßnahmen", "100 % Nachprüfung aller Lagerbestände. Einfrieren der betroffenen Linie. Sofortige Kundenkommunikation."),
                      ("D4 – Ursachenanalyse", f"Ursache: {random.choice(['Falsche Löttemperatur', 'Chargenabweichung Lieferant', 'Softwarefehler FW-Update', 'Prozessparameter außerhalb Toleranz'])}\nNachweis: SPC-Daten + Ishikawa-Diagramm"),
                      ("D5-D7 – Korrekturmaßnahmen", "Permanente Korrekturmaßnahme: Neue SOP implementiert, Mitarbeiterschulung durchgeführt, Prozessvalidierung (Re-PPAP) abgeschlossen."),
                      ("D8 – Abschluss", f"Der 8D-Bericht wurde vom {cust['name']} am {ds(2023, random.randint(5,11), random.randint(15,30))} akzeptiert."),
                  ])

    # Process Audit Reports (VDA 6.3)
    for sub in SUBS[:3]:
        make_docx(folder, f"VDA63_Prozessaudit_{sub['short']}_{2023}.docx",
                  sub["name"], sub["city"], sub["hrb"],
                  f"VDA 6.3 Prozessaudit-Bericht – {sub['name']} – {2023}",
                  [
                      ("Auditumfang", f"Auditierter Prozess: Serienproduktion {random.choice(PRODUCTS)['name']}\nAuditor: Externe Zertifizierungsstelle\nAuditdatum: {ds(2023, random.randint(3,9), random.randint(1,25))}"),
                      ("Ergebnis", [
                          ["Prozessschritt", "Bewertung %", "Status"],
                          ["P2 – Projektmanagement", f"{random.randint(75,100)} %", "OK"],
                          ["P3 – Planung Produktentwicklung", f"{random.randint(75,100)} %", "OK"],
                          ["P4 – Planung Serienproduktion", f"{random.randint(70,100)} %", "OK" if random.random() > 0.2 else "Abweichung"],
                          ["P5 – Lieferantenmanagement", f"{random.randint(75,100)} %", "OK"],
                          ["P6 – Serienproduktion", f"{random.randint(70,100)} %", "OK"],
                          ["P7 – Kundendienst", f"{random.randint(75,100)} %", "OK"],
                      ]),
                  ])


# ═══════════════════════════════════════════════════════════════════════════
# 14 – IP / TECHNOLOGIE
# ═══════════════════════════════════════════════════════════════════════════

def gen_14():
    folder = "14_IP_Technologie"
    print(f"\n[14] IP / Technologie …")

    # Patent Portfolio Overview
    make_xlsx(folder, "REA_Patentportfolio_2024.xlsx",
              "Patentportfolio – Brennhagen Elektronik AG – Stand Januar 2024",
              [
                  ("Patente DE-EP",
                   ["Patentnummer", "Titel", "Anmeldedatum", "Status", "Technologiegebiet", "Erteilungsland"],
                   [
                       [f"DE10{2010+i}0{random.randint(10000,99999)}", f"Verfahren zur {'Optimierung der ECU-Zykluszeit' if i%4==0 else 'ADAS-Bildverarbeitungsalgorithmus' if i%4==1 else 'Batterie-SOC-Schätzung' if i%4==2 else 'CAN-FD-Priorisierung'}", f"{2015+i//3}-{random.randint(1,12):02d}-{random.randint(1,28):02d}", random.choice(["Erteilt","Erteilt","Erteilt","Anhängig","Anhängig"]), random.choice(["ECU","ADAS","EV","Infotainment"]), random.choice(["DE","EP","US","CN"])]
                       for i in range(18)
                   ],
                   [22, 55, 15, 12, 15, 12]),
              ])

    # Individual Patent Documents (15+)
    patent_titles = [
        "Verfahren und Vorrichtung zur adaptiven Zykluszeit-Optimierung in Automotive-Steuergeräten",
        "Kamerasystem mit 4D-Radarfusion für ADAS-Applikationen",
        "Batteriemanagementsystem mit prädiktiver SOC/SOH-Schätzung mittels KNN",
        "CAN-FD-Priorisierungsalgorithmus für sicherheitskritische Automotive-Netzwerke",
        "Verfahren zur EMV-Unterdrückung in Leistungselektronik-Modulen",
        "Selbstkalibrierendes Sensorsystem für Fahrerassistenzanwendungen",
        "OTA-Updateverfahren für Automotive-ECUs unter ISO 26262",
        "Redundantes Überwachungssystem für Hochvolt-Traktionsbatterien",
        "Verfahren zur Detektion von Spoofing-Angriffen in V2X-Kommunikation",
        "Wärmemanagement-Optimierung in Multi-Core-Automotive-Prozessoren",
        "Prozess zur Herstellung von Hochtemperatur-stabilen Leiterplatten für Motorraum-ECUs",
        "Vorrichtung zur galvanischen Entkopplung von Automotive-BUS-Systemen",
        "Verfahren zum präzisen Strommessen in Drei-Phasen-Inverter-Schaltungen für BEV",
        "Datenschutzkonformes Logging-Verfahren für ADAS-Ereignisrekorder",
        "Laser-Direct-Imaging (LDI) Verfahren für Hochdichte-Interconnect-Leiterplatten",
    ]
    for i, title in enumerate(patent_titles):
        make_docx(folder, f"Patent_{i+1:02d}_{sfn(title[:30])}_2023.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Patent – {title}",
                  [
                      ("Patentsteckbrief", [
                          ["Feld", "Inhalt"],
                          ["Anmelder", AG["name"]],
                          ["Erfinder", random.choice(ALL_NAMES) + ", " + AG["city"]],
                          ["Titel", title],
                          ["Patentnummer", f"DE{random.randint(102000000, 102999999)}A1"],
                          ["EP-Pendant", f"EP{random.randint(3000000, 3999999)}B1"],
                          ["Anmeldedatum", ds(random.randint(2015, 2022), random.randint(1,12), random.randint(1,25))],
                          ["Erteilt", ds(random.randint(2017, 2024), random.randint(1,12), random.randint(1,25))],
                          ["Technologiegebiet", random.choice(["ECU", "ADAS", "EV", "Infotainment", "Herstellung"])],
                          ["Status", "Erteilt, Jahresgebühren bezahlt bis 2030"],
                          ["Strategische Bedeutung", "Hoch – Kernpatent"],
                      ]),
                      ("Zusammenfassung der Erfindung", f"Die Erfindung betrifft {title}. Im Stand der Technik bekannte Lösungen weisen den Nachteil auf, dass [Beschreibung Nachteil Stand der Technik]. Die erfindungsgemäße Lösung zeichnet sich dadurch aus, dass [wesentliche Merkmale der Erfindung]. Hierdurch wird [technischer Effekt und Vorteil] erzielt."),
                      ("Patentansprüche (Auszug)", f"1. [Hauptanspruch: Gerät/Verfahren] nach dem {title}-Prinzip, dadurch gekennzeichnet, dass [kennzeichnender Teil des Hauptanspruchs].\n\n2. [Nebenanspruch] nach Anspruch 1, dadurch gekennzeichnet, dass [weitere Merkmale]."),
                  ])

    # Technology Licensing Agreements
    for cust in CUSTOMERS[:2]:
        make_docx(folder, f"REA_{cust['short']}_Technologielizenz_2021.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Technologielizenzvertrag – {AG['name']} / {cust['name']}",
                  [
                      ("Präambel", f"{AG['name']} als Inhaber von Patenten und Know-how im Bereich Automotive-Elektronik (Lizenzgeber) und {cust['name']} (Lizenznehmer) schließen den nachfolgenden Technologielizenzvertrag."),
                      ("§ 1 Lizenzgegenstand", f"(1) Der Lizenzgeber gewährt dem Lizenznehmer eine nicht-ausschließliche, weltweite Lizenz an den in Anlage 1 aufgeführten Patenten und dem zugehörigen Know-how.\n\n(2) Der Lizenznehmer darf die lizenzierten Technologien ausschließlich für die Herstellung und den Vertrieb von [Produkt] in eigenen Fahrzeugen verwenden."),
                      ("§ 2 Lizenzgebühren", f"(1) Der Lizenznehmer zahlt eine laufende Lizenzgebühr von {random.uniform(0.5, 2.5):.1f} % des Nettoumsatzes mit Produkten, die auf den lizenzierten Technologien basieren.\n\n(2) Mindestjahreslizenz: EUR {random.randint(200,800) * 1000:,}."),
                      (None, std_governing_law_de()),
                  ], confidential=True)

    # Joint Development Agreements
    make_docx(folder, f"REA_BMW_JDA_ADAS_2022.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Joint Development Agreement – ADAS-Vision 5D – REA / BMW Group",
              [
                  ("1. Gegenstand", f"Die {AG['name']} und die BMW Group vereinbaren die gemeinsame Entwicklung der nächsten Generation des ADAS-Vision-Systems (intern: ADAS-Vision 5D) für die BMW 7er-Plattform (Codename: G70-Next)."),
                  ("2. Beiträge der Parteien", [
                      ["Partei", "Beitrag", "Budget EUR Mio."],
                      [AG["name"], "Hardware, Sensorarchitektur, FPGA-Design, Systemintegration", f"{random.randint(15,30)} Mio."],
                      ["BMW Group", "Applikationssoftware (AUTOSAR), Homologation, Erprobungsfahrzeuge", f"{random.randint(10,20)} Mio."],
                  ]),
                  ("3. IP-Zuordnung", "(1) Vorbestehendes IP verbleibt jeweils im Eigentum der einbringenden Partei.\n\n(2) Gemeinsam entwickeltes IP (Gemeinschaftserfindungen) steht beiden Parteien zu gleichen Teilen zu. Jede Partei darf Gemeinschaftserfindungen ohne Zustimmung der anderen lizenzieren, hat jedoch die andere Partei vorab zu informieren."),
                  ("4. Vertraulichkeit", "Sämtliche im Rahmen des JDA ausgetauschten Informationen sind als streng vertraulich zu behandeln. Die Vertraulichkeitsverpflichtung gilt für 5 Jahre nach Beendigung des JDA."),
              ], confidential=True)

    # Source code escrow
    make_docx(folder, "REA_Source_Code_Escrow_BMW_2022.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Source Code Escrow Agreement – BMW Group / REA",
              [
                  ("1. Zweck", f"Zur Sicherung der Lieferfähigkeit und des Betriebs der von {AG['name']} gelieferten Embedded Software vereinbaren die Parteien eine Hinterlegung des Quellcodes bei einem unabhängigen Treuhänder (Iron Mountain Deutschland GmbH)."),
                  ("2. Hinterlegungsumfang", "Hinterlegt wird der vollständige Quellcode der Firmware ECU-900, ADAS-Vision 4D sowie zugehöriger Entwicklungstools und Build-Scripte, jeweils in der letzten freigegebenen Serienversion."),
                  ("3. Freigabebedingungen (Release Conditions)", "Der Treuhänder darf den hinterlegten Code an BMW freigeben bei: (a) Insolvenz oder Liquidation von REA, (b) Einstellung des Geschäftsbetriebs, (c) Wesentlicher und nicht behobener Vertragsverletzung durch REA."),
              ])

    # Open Source Compliance Report
    make_docx(folder, "REA_OpenSource_Compliance_Report_2023.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Open Source Software Compliance Report 2023 – Brennhagen Elektronik AG",
              [
                  ("Executive Summary", f"Dieser Bericht dokumentiert die Einhaltung der Open Source Software (OSS)-Lizenzanforderungen in den Produkten der {AG['name']}. Die Analyse wurde von der IP-Abteilung mit Unterstützung von {AG['recht']} durchgeführt."),
                  ("OSS-Komponenten (Auszug)", [
                      ["Komponente", "Version", "Lizenz", "Einsatz", "Compliance-Status"],
                      ["Linux Kernel (RTOS)", "5.15 LTS", "GPL v2", "Infotainment-Modul", "Konform"],
                      ["FreeRTOS", "10.5.0", "MIT", "ECU-900 Low-Level", "Konform"],
                      ["OpenSSL", "3.0.8", "Apache 2.0", "TLS/Cybersecurity", "Konform"],
                      ["Zephyr RTOS", "3.4.0", "Apache 2.0", "ADAS-System", "Konform"],
                      ["AUTOSAR CP", "R22-11", "Proprietär", "ECU-Software-Stack", "Lizenz vorhanden"],
                  ]),
                  ("Fazit", "Alle eingesetzten OSS-Komponenten werden lizenzkonform verwendet. Copyleft-Verpflichtungen werden durch ein internes SBOM-Register (Software Bill of Materials) überwacht."),
              ])

    # FRAND Declaration
    make_docx(folder, "REA_FRAND_Erklaerung_SEP_2023.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "FRAND-Erklärung – Standards Essential Patents (SEP)",
              [
                  ("Erklärung", f"Die {AG['name']} erklärt hiermit, dass sie für alle von ihr gehaltenen SEPs (Standards Essential Patents), die für die Automotive-Kommunikationsstandards (C-V2X, 5G-NR, DSRC) wesentlich sind, Lizenzen zu fairen, angemessenen und nicht-diskriminierenden (FRAND) Bedingungen anbietet.\n\nDiese Erklärung gilt gegenüber dem ETSI (European Telecommunications Standards Institute) sowie gegenüber IEEE und ISO/IEC."),
              ])

    # Trade Secret Register
    make_xlsx(folder, "REA_Betriebsgeheimnisregister_2023.xlsx",
              "Betriebsgeheimnisregister – Brennhagen Elektronik AG – 2023",
              [
                  ("Register",
                   ["ID", "Beschreibung", "Kategorie", "Schutzmaßnahmen", "Zugang", "Letzte Überprüfung"],
                   [
                       [f"TS-{i:03d}", f"{'Algorithmus ECU-900 Kalibrierung' if i%5==0 else 'Produktionskostenstruktur' if i%5==1 else 'ADAS-Sensor-Fusion-Methodik' if i%5==2 else 'Kundenkonditionenverzeichnis' if i%5==3 else 'Fertigungsrezeptur BGA-Lötprozess'}", random.choice(["Technisch","Kommerziell","Know-how"]), "NDA, Zugangskontrolle, Verschlüsselung", "Need-to-know", ds(2023, random.randint(1,12), 1)]
                       for i in range(1, 16)
                   ],
                   [10, 50, 15, 40, 20, 20]),
              ])


# ═══════════════════════════════════════════════════════════════════════════
# 15 – COMPLIANCE / RECHT
# ═══════════════════════════════════════════════════════════════════════════

def gen_15():
    folder = "15_Compliance_Recht"
    print(f"\n[15] Compliance / Recht …")

    # Compliance Program
    make_docx(folder, "REA_Compliance_Management_System_2023.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Compliance-Management-System (CMS) – Brennhagen Elektronik AG",
              [
                  ("1. Compliance-Governance", f"Das Compliance-Management-System der {AG['name']} basiert auf dem IDW PS 980 und dem Deutschen Corporate Governance Kodex. Der Chief Compliance Officer (CCO) berichtet direkt an die Vorstandsvorsitzende {AG['ceo']} und quartalsweise an den Prüfungsausschuss des Aufsichtsrats."),
                  ("2. Risikoanalyse", "Die jährliche Compliance-Risikoanalyse identifiziert folgende Top-Risiken:\n\n1. Korruption in Schwellenmärkten (v.a. China)\n2. Kartellrecht – Informationsaustausch OEM-Lieferantenkreis\n3. Exportkontrolle / Dual-Use-Güter\n4. Datenschutz (DSGVO) in Kundenbeziehungen\n5. Lieferkettenpflichten (LkSG)"),
                  ("3. Schulungen und Kommunikation", f"Alle Mitarbeitenden absolvieren jährlich verpflichtende Compliance-Online-Trainings. Führungskräfte erhalten zusätzliche Präsenztrainings. Schulungsquote 2023: {random.randint(92,98)} %."),
                  ("4. Hinweisgebersystem", "Das anonyme Hinweisgebersystem (Compliance-Hotline) wird von einem externen Dienstleister betrieben und ist 24/7 erreichbar. Im Jahr 2023 wurden 7 Hinweise eingereicht; alle wurden untersucht."),
              ])

    # Anti-Corruption Policy
    make_docx(folder, "REA_Anti_Korruption_Richtlinie_2023.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Anti-Korruptions- und Anti-Bestechungsrichtlinie 2023",
              [
                  ("§ 1 Zweck und Geltungsbereich", f"Diese Richtlinie gilt für alle Mitarbeitenden, Führungskräfte und Organmitglieder der {AG['name']} sowie aller Konzerngesellschaften. Sie verbietet alle Formen der Korruption und Bestechung, sowohl gegenüber öffentlichen als auch privaten Akteuren."),
                  ("§ 2 Verbotene Handlungen", "(1) Es ist verboten, Dritten im geschäftlichen Umfeld Vorteile anzubieten, zu versprechen oder zu gewähren, um einen Auftrag oder eine Entscheidung zu erlangen.\n\n(2) Es ist verboten, Vorteile von Dritten anzunehmen oder zu fordern.\n\n(3) Repräsentationsausgaben sind auf EUR 50 pro Person pro Anlass begrenzt. Höhere Ausgaben bedürfen der vorherigen Genehmigung durch den Compliance Officer."),
                  ("§ 3 Geschenkeregister", "Alle Geschenke und Einladungen über EUR 25 Sachwert sind im Konzern-Geschenkeregister einzutragen. Das Register wird vom Compliance Officer verwaltet und quartalsweise ausgewertet."),
                  ("§ 4 Business Partners", "Vor Aufnahme einer Geschäftsbeziehung mit Intermediären oder Vertriebspartnern in Hochrisikomärkten ist eine Due-Diligence-Prüfung (Third Party Due Diligence) durchzuführen."),
              ])

    # Export Control Program
    make_docx(folder, "REA_Exportkontrollhandbuch_2023.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Exportkontrollprogramm – Brennhagen Elektronik AG – 2023",
              [
                  ("1. Überblick Exportkontrolle", f"Die Produkte der {AG['name']} unterliegen teilweise den Anforderungen der EU-Dual-Use-Verordnung (2021/821), des deutschen Außenwirtschaftsgesetzes (AWG) sowie der US-amerikanischen Export Administration Regulations (EAR).\n\nExportkontroll-Beauftragter: Compliance Officer, {AG['name']}."),
                  ("2. Produktklassifikation", [
                      ["Produkt", "EG-Nr. (Dual-Use-VO)", "ECCN (EAR)", "Genehmigungspflicht"],
                      [AG["prod1"], "5A002.a", "5A002.a.1", "Ja – bei Bestimmung nach Russland, China (bestimmte Anwendungen)"],
                      [AG["prod2"], "6A003.b.4", "6A003.b.4.b", "Ja – bei militärischen Endverwendern"],
                      [AG["prod3"], "Nicht kontrolliert", "EAR99", "Nein"],
                      [AG["prod4"], "3A001.a.1", "3A001.a.1", "Einzelfallprüfung erforderlich"],
                  ]),
                  ("3. Sanktionsscreening", "Vor jeder Geschäftsanbahnung und bei jeder Neubestellung erfolgt ein automatisches Screening der Kunden, Lieferanten und Vertragspartner gegen die EU-, US- und UN-Sanktionslisten (OFAC SDN, EU Consolidated List). Das Screening-System ist in SAP S/4HANA integriert."),
              ])

    # LkSG Documentation
    make_docx(folder, "REA_LkSG_Sorgfaltspflichtenbericht_2023.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Bericht über die Erfüllung der Sorgfaltspflichten nach LkSG – 2023",
              [
                  ("1. Risikoanalyse", f"Im Rahmen der nach § 5 LkSG vorgeschriebenen Risikoanalyse hat die {AG['name']} eine umfassende Bewertung der menschenrechtlichen und umweltbezogenen Risiken in der Lieferkette durchgeführt. Schwerpunktländer: China (Rohstoffgewinnung Kobalt/Lithium), Indien (Sublieferanten Leiterplatten), Südostasien."),
                  ("2. Präventionsmaßnahmen", "(1) Alle Direktlieferanten wurden über den Verhaltenskodex für Lieferanten (Supplier Code of Conduct) informiert.\n\n(2) Risikobasierte Lieferantenaudits wurden bei 15 Direktlieferanten in Hochrisikoregionen durchgeführt.\n\n(3) Vertragsklauseln zum LkSG wurden in alle Neueinkaufsverträge integriert."),
                  ("3. Abhilfemaßnahmen", "Einem Direktlieferanten in Vietnam wurden konkrete Verbesserungsmaßnahmen bezüglich Arbeitsbedingungen (Überstundenregelung) kommuniziert. Der Lieferant hat Nachweise über die Umsetzung vorgelegt."),
                  ("4. Beschwerdeverfahren", "Das Beschwerdemechanismus ist über die konzernweite Compliance-Hotline und per E-Mail (lksg@brennhagen-elektronik.de) erreichbar. Im Jahr 2023 gingen 2 Beschwerden ein; beide wurden untersucht."),
              ])

    # DSGVO/GDPR Documentation
    gdpr_docs = [
        "REA_DSGVO_Datenschutzrichtlinie_2023",
        "REA_DSGVO_Verzeichnis_Verarbeitungstaetigkeiten_2023",
        "REA_DSGVO_DSFA_CRM_System_2022",
        "REA_DSGVO_Datenschutzbeauftragter_Bestellung_2022",
        "REA_DSGVO_AV_Vertrag_Salesforce_2023",
        "REA_DSGVO_Datenschutzhinweise_Beschaeftigte_2023",
    ]
    for dname in gdpr_docs:
        title = dname.replace("REA_DSGVO_","").replace("_2023","").replace("_2022","").replace("_"," ")
        make_docx(folder, dname + ".docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"DSGVO – {title}",
                  [
                      ("Datenschutz", f"Dieses Dokument beschreibt {title} für die {AG['name']} und alle Konzerngesellschaften. Datenschutzbeauftragter (DSB): extern bestellt gemäß Art. 37 DSGVO. Kontakt: dse@brennhagen-elektronik.de."),
                  ])

    # Competition Law Compliance
    make_docx(folder, "REA_Kartellrecht_Compliance_Handbuch_2023.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Kartellrecht-Compliance-Handbuch 2023 – Brennhagen Elektronik AG",
              [
                  ("1. Überblick", f"Die {AG['name']} als Tier-1-Lieferant an Automobilhersteller steht im regelmäßigen Kontakt mit Wettbewerbern (Lieferantenkreise, Branchenverbände). Dieses Handbuch beschreibt die Regeln für den rechtmäßigen Umgang mit Wettbewerbern."),
                  ("2. Verbotener Informationsaustausch", "(1) Es ist verboten, mit Wettbewerbern Preise, Margen, Kalkulationen oder Preisbestandteile auszutauschen.\n\n(2) Es ist verboten, mit Wettbewerbern Märkte, Kunden oder Lieferanten aufzuteilen.\n\n(3) Bei Branchenveranstaltungen oder Verbandssitzungen sind Gespräche über marktrelevante Daten sofort abzubrechen."),
                  ("3. Verhalten im Beschaffungsverfahren", "(1) Es ist verboten, Submissionsabsprachen zu treffen.\n\n(2) Alle Angebote für OEM-Ausschreibungen sind unabhängig und eigenständig zu kalkulieren."),
              ])

    # Legal opinions
    for topic in ["Übernahme_Czech_GmbH_2018", "IC_Darlehen_Zinssatz_2022", "LkSG_Anwendbarkeit_2023"]:
        make_docx(folder, f"REA_Rechtsgutachten_{topic}.docx",
                  AG["recht"], "Bockenheimer Anlage 44, 60322 Frankfurt", f"Ref: REA/{random.randint(1000,9999)}/{topic[:10]}",
                  f"Rechtsgutachten – {topic.replace('_', ' ')}",
                  [
                      ("Auftrag", f"Die {AG['name']} hat uns beauftragt, folgendes Rechtsgutachten zu erstatten: {topic.replace('_',' ')}."),
                      ("Sachverhalt", f"[Darstellung des relevanten Sachverhalts und der rechtlichen Fragestellung]"),
                      ("Rechtliche Würdigung", f"Nach eingehender Prüfung der einschlägigen Rechtsvorschriften und Rechtsprechung kommen wir zu folgender rechtlicher Einschätzung: [Gutachteninhalt]. Die genannte Transaktion/Maßnahme ist aus unserer Sicht rechtlich zulässig, sofern die folgenden Voraussetzungen eingehalten werden: [Voraussetzungen]."),
                      ("Fazit", f"\n\nFrankfurt am Main, {ds(2023, random.randint(3,11), random.randint(1,25))}\n\n{AG['recht']}"),
                  ], confidential=True)

    # Litigation Register
    make_xlsx(folder, "REA_Rechtsstreitigkeitsregister_2024.xlsx",
              "Register der Rechtsstreitigkeiten und Risiken – REA – Stand Q1 2024",
              [
                  ("Rechtsstreitigkeiten",
                   ["Lfd. Nr.", "Beschreibung", "Gegenpartei", "Instanz", "Streitwert", "Rückstellung", "Status"],
                   [
                       ["1", "Produkthaftungsklage – ECU-900 Ausfall", "Versicherung Kunde", "LG Stuttgart", eur(2_800_000), eur(1_400_000), "Anhängig"],
                       ["2", "Kündigung Mietvertrag Heilbronn", "Vermieter", "LG Heilbronn", eur(500_000), eur(250_000), "Vergleich in Verhandlung"],
                       ["3", "Arbeitsgericht – fristlose Kündigung", "Ehem. Mitarbeiter", "ArbG Stuttgart", eur(120_000), eur(60_000), "Erstinstanz abgeschlossen"],
                       ["4", "Patentklage – EP3456789B1", "Wettbewerber X", "BPatG München", eur(5_000_000), eur(1_000_000), "Gutachtenphase"],
                       ["5", "Zollrückforderung – CN-Import 2021", "Hauptzollamt Stuttgart", "Finanzgericht", eur(380_000), eur(380_000), "Einspruch eingelegt"],
                   ],
                   [8, 50, 30, 20, 18, 18, 22]),
              ])

    # Sanctions Screening Procedure
    make_docx(folder, "REA_Sanktionsscreening_Verfahren_2023.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Sanktionsscreening-Verfahren 2023 – Brennhagen Elektronik AG",
              [
                  ("1. Prozessbeschreibung", "Das Sanktionsscreening erfolgt vollautomatisch über das SAP GTS (Global Trade Services)-Modul, das täglich mit den aktuellen Sanktionslisten aktualisiert wird (EU Consolidated List, OFAC SDN, UN Consolidated List, UK Financial Sanctions List)."),
                  ("2. Trefferverwaltung", "Bei einem positiven Treffer wird die Transaktion automatisch gesperrt und an den Exportkontroll-Beauftragten weitergeleitet. Dieser prüft innerhalb von 24 Stunden, ob ein False Positive oder ein echter Treffer vorliegt."),
                  ("3. Eskalation", "Bei bestätigten Treffern wird die Transaktion abgelehnt und die Compliance-Abteilung sowie die Rechtsabteilung ({AG['recht']}) informiert."),
              ])

    # Government Relations Log
    make_xlsx(folder, "REA_Government_Relations_Log_2023.xlsx",
              "Government Relations / Politische Kontakte 2023 – REA",
              [
                  ("Kontaktlog",
                   ["Datum", "Institution", "Kontaktperson", "Thema", "REA-Vertreter", "Ergebnis"],
                   [
                       [ds(2023, m, random.randint(5,25)), random.choice(["Bundesministerium für Wirtschaft","EU-Kommission DG GROW","BDI","VDA","Stuttgarter Landtag"]), "[Kontakt]", random.choice(["Förderantragsberatung","Regulierung Dual-Use","IPCEI-Mikroelektronik","LkSG-Auslegung"]), AG["ceo"] if random.random() > 0.5 else "VP Public Affairs", "Informativ / Abstimmungsgespräch"]
                       for m in random.sample(range(1,13), 8)
                   ],
                   [18, 35, 25, 40, 30, 30]),
              ])


# ═══════════════════════════════════════════════════════════════════════════
# 16 – IT / SYSTEME
# ═══════════════════════════════════════════════════════════════════════════

def gen_16():
    folder = "16_IT_Systeme"
    print(f"\n[16] IT / Systeme …")

    it_docs = [
        ("REA_IT_Governance_Framework_2023", "IT-Governance-Framework (COBIT 2019)"),
        ("REA_Cybersecurity_Policy_ISMS_2023", "Cybersecurity-Policy und ISMS (ISO 27001)"),
        ("REA_TISAX_Audit_Bericht_2023", "TISAX Level 3 Audit-Bericht"),
        ("REA_SAP_S4HANA_Systemdokumentation_2022", "SAP S/4HANA Systemdokumentation"),
        ("REA_Siemens_Teamcenter_PLM_2022", "Siemens Teamcenter PLM-Dokumentation"),
        ("REA_SAP_MES_Produktionssteuerung_2022", "SAP MES – Produktionssteuerung"),
        ("REA_Cloud_Strategie_2023", "Cloud-Strategie und Sicherheitsarchitektur"),
        ("REA_BCP_Business_Continuity_Plan_2023", "Business Continuity Plan (BCP)"),
        ("REA_Disaster_Recovery_Test_2023", "Disaster Recovery Test-Bericht 2023"),
        ("REA_IT_Risikobericht_2023", "IT-Risikobericht Q4 2023"),
        ("REA_Netzwerkarchitektur_Dokument_2023", "Netzwerkarchitektur-Dokumentation"),
        ("REA_Endpoint_Security_Policy_2023", "Endpoint-Security-Richtlinie"),
        ("REA_Backup_Konzept_2022", "Datensicherungs- und Backup-Konzept"),
        ("REA_IAM_Identity_Access_Management_2023", "Identity & Access Management (IAM)"),
        ("REA_Penetrationstest_Bericht_2023", "Penetrationstest-Bericht (extern) 2023"),
    ]
    for fname, title in it_docs:
        make_docx(folder, fname + ".docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  title,
                  [
                      ("Überblick", f"Dieses Dokument beschreibt {title} für die {AG['name']} und die verbundenen Konzerngesellschaften. IT-Verantwortlicher: VP IT & Digitalization."),
                      ("System-Umgebung", f"ERP: {AG['erp']}\nPLM: {AG['plm']}\nMES: {AG['mes']}\nCRM: Salesforce CRM (Cloud)\nHR: SAP SuccessFactors\nITSM: ServiceNow"),
                      ("Maßnahmen und Status", f"Alle beschriebenen Maßnahmen wurden gemäß IT-Roadmap 2023 implementiert. Überprüfung erfolgt jährlich durch den VP IT und externen Auditor ({AG['wp']})."),
                  ])

    # TISAX Certificate
    make_pdf(folder, "REA_TISAX_Zertifikat_Level3_2023.pdf",
             "ENX Association", "Lyoner Straße 18, 60528 Frankfurt", "TISAX-Registrar",
             f"TISAX-Zertifikat Level 3 – {AG['name']} – 2023",
             [
                 ("Zertifizierung", f"Hiermit wird bestätigt, dass die {AG['name']} das TISAX-Assessment Level 3 (High Protection Need) erfolgreich bestanden hat.\n\nGültigkeitsdauer: {ds(2023, 6, 1)} bis {ds(2026, 5, 31)}\nScope: Entwicklung, Produktion, IT-Systeme – Standorte Stuttgart, Heilbronn, München\nAssessor: [Zugelassener TISAX-Auditor]\nTISAX-ID: REA-TSX-{random.randint(100000,999999)}"),
             ])


# ═══════════════════════════════════════════════════════════════════════════
# 17 – VERSICHERUNGEN
# ═══════════════════════════════════════════════════════════════════════════

def gen_17():
    folder = "17_Versicherungen"
    print(f"\n[17] Versicherungen …")

    # Group Insurance Overview
    make_xlsx(folder, "REA_Versicherungsportfolio_2023_Uebersicht.xlsx",
              "Versicherungsportfolio – Brennhagen Elektronik AG – 2023",
              [
                  ("Versicherungen",
                   ["Sparte", "Versicherer", "Police-Nr.", "Deckungssumme", "Jahresprämie", "Laufzeit bis"],
                   [
                       ["Produkthaftpflicht", "AXA XL Insurance", f"AXL-REA-{random.randint(100000,999999)}", eur(100_000_000), eur(420_000), "31.12.2024"],
                       ["Betriebshaftpflicht", "Allianz SE", f"ALZ-REA-{random.randint(100000,999999)}", eur(50_000_000), eur(285_000), "31.12.2024"],
                       ["D&O-Versicherung", "Chubb European Group", f"CHB-REA-{random.randint(100000,999999)}", eur(30_000_000), eur(185_000), "31.12.2024"],
                       ["Cyber-Versicherung", "Beazley Group", f"BEZ-REA-{random.randint(100000,999999)}", eur(25_000_000), eur(320_000), "31.12.2024"],
                       ["Industriefeuerversicherung", "HDI Global SE", f"HDI-REA-{random.randint(100000,999999)}", eur(380_000_000), eur(610_000), "31.12.2025"],
                       ["Betriebsunterbrechung", "Munich Re", f"MRE-REA-{random.randint(100000,999999)}", eur(150_000_000), eur(290_000), "31.12.2025"],
                       ["Transportversicherung (Cargo)", "Generali Deutschland", f"GEN-REA-{random.randint(100000,999999)}", eur(20_000_000), eur(125_000), "31.12.2024"],
                       ["Kreditversicherung", "Euler Hermes (Allianz Trade)", f"EUH-REA-{random.randint(100000,999999)}", "80 % Forderungsdeckung", eur(98_000), "31.12.2024"],
                       ["Umwelthaftpflicht", "Zurich Insurance", f"ZUR-REA-{random.randint(100000,999999)}", eur(10_000_000), eur(65_000), "31.12.2025"],
                       ["Reiseversicherung (kollektiv)", "ERGO Group", f"ERG-REA-{random.randint(100000,999999)}", "Weltweit", eur(42_000), "31.12.2024"],
                   ],
                   [30, 28, 22, 22, 18, 15]),
              ])

    # Individual Policy Docs
    policies = [
        ("Produkthaftpflicht", "AXA XL", eur(100_000_000)),
        ("D_O_Versicherung", "Chubb", eur(30_000_000)),
        ("Cyber_Versicherung", "Beazley", eur(25_000_000)),
        ("Industriefeuer_Heilbronn", "HDI Global", eur(180_000_000)),
        ("Betriebsunterbrechung", "Munich Re", eur(150_000_000)),
        ("Transportversicherung", "Generali", eur(20_000_000)),
    ]
    for sname, insurer, limit in policies:
        make_docx(folder, f"REA_Police_{sname}_2023.docx",
                  insurer, "[Versichereranschrift]", f"Police Nr.: {insurer[:3].upper()}-REA-{random.randint(100000,999999)}",
                  f"Versicherungspolice – {sname.replace('_', ' ')} – {AG['name']}",
                  [
                      ("Versicherungsnehmer", f"{AG['name']}, {AG['full_address']}"),
                      ("Versicherungsschutz", f"Sparte: {sname.replace('_',' ')}\nDeckungssumme: {limit}\nJahresprämie: {eur(random.randint(60, 600) * 1000)}\nSelbstbehalt: EUR {random.choice([50_000, 100_000, 250_000, 500_000]):,}\nGeltungsbereich: Weltweit / Europa\nLaufzeit: 1. Januar 2023 bis 31. Dezember 2023"),
                  ])

    # Claims History
    make_xlsx(folder, "REA_Schadenshistorie_2018_2023.xlsx",
              "Schadenshistorie – Brennhagen Elektronik AG – 2018–2023",
              [
                  ("Schäden",
                   ["Jahr", "Sparte", "Beschreibung", "Schadenshöhe", "Versicherungsleistung", "Status"],
                   [
                       [y, random.choice(["Produkthaftpflicht","Feuer","BU","Cyber","Transport"]),
                        random.choice(["Maschinenschaden Linie 3", "Produktrückruf Charge XY", "Cybervorfal Ransomware", "Transportschaden LKW", "Kundenschaden ECU"]),
                        eur(random.randint(50_000, 3_000_000)), eur(random.randint(30_000, 2_500_000)), random.choice(["Abgeschlossen","Abgeschlossen","In Regulierung"])]
                       for y in range(2018, 2024)
                   ],
                   [8, 20, 45, 18, 18, 18]),
              ])


# ═══════════════════════════════════════════════════════════════════════════
# 18 – IMMOBILIEN
# ═══════════════════════════════════════════════════════════════════════════

def gen_18():
    folder = "18_Immobilien"
    print(f"\n[18] Immobilien …")

    sites = [
        ("HQ_Stuttgart", AG["name"], "Vaihinger Straße 120, 70567 Stuttgart", 12_500, "Büro / Entwicklung / Verwaltung", 2_800_000),
        ("Werk_Heilbronn", SUBS[0]["name"], "Industriestraße 45, 74078 Heilbronn", 42_000, "Produktion / Lager", 4_200_000),
        ("Office_München", SUBS[1]["name"], "Leopoldstraße 180, 80804 München", 5_200, "Büro / Software-Entwicklung", 1_850_000),
        ("Werk_Katowice", SUBS[2]["name"], "ul. Przemysłowa 12, 40-519 Katowice", 28_000, "Produktion / Lager", 1_200_000),
        ("Werk_Brno", SUBS[3]["name"], "Průmyslová 8, 619 00 Brno", 18_500, "Engineering / Entwicklung", 980_000),
        ("Werk_Gyoer", SUBS[4]["name"], "Ipari utca 5, 9027 Győr", 24_000, "Produktion EV", 1_150_000),
        ("Office_Shanghai", SUBS[5]["name"], "1000 Lujiazui Ring Road, Shanghai", 3_800, "Büro / Vertrieb / Produktion", 2_200_000),
    ]
    for site_id, tenant, address, area, use, rent_pa in sites:
        make_docx(folder, f"REA_Mietvertrag_{site_id}_2020.docx",
                  tenant, address, "Mietvertrag",
                  f"Gewerbemietvertrag – {site_id.replace('_', ' ')}",
                  [
                      ("§ 1 Mietobjekt", f"Vermietet wird das Objekt {address}, Nutzfläche ca. {area:,} m², für den Zweck: {use}."),
                      ("§ 2 Mietzins und Nebenkosten", f"(1) Jahresmiete: {eur(rent_pa)}\n(2) Monatliche Miete: {eur(rent_pa//12)}\n(3) Umsatzsteuer: zzgl. gesetzlicher USt. (19 %)\n(4) Betriebskosten: Vorauszahlung {eur(int(rent_pa*0.15)//12)} monatlich, Jahresabrechnung"),
                      ("§ 3 Mietdauer", f"Mietbeginn: 1. Januar 2020. Feste Laufzeit bis 31. Dezember 2030, danach automatische Verlängerung um je 5 Jahre mit 18-monatiger Kündigungsfrist."),
                      ("§ 4 Instandhaltung", "Schönheitsreparaturen gehen zu Lasten des Mieters. Wartung technischer Anlagen (Klimaanlage, Aufzüge) trägt der Vermieter. Bauliche Veränderungen bedürfen der schriftlichen Zustimmung des Vermieters."),
                      ("§ 5 Indexierung", "Der Mietzins wird jährlich entsprechend der Entwicklung des Verbraucherpreisindex (VPI) des Statistischen Bundesamts angepasst (max. +3 % p.a.)."),
                      (None, std_governing_law_de()),
                  ])

        # Property Valuation
        make_docx(folder, f"REA_Immobilienbewertung_{site_id}_2023.docx",
                  "Jones Lang LaSalle SE", "Bockenheimer Landstr. 55, 60325 Frankfurt", "Gutachter-Ref.: JLL-2023-REA",
                  f"Immobilienbewertungsgutachten – {site_id.replace('_', ' ')} – 2023",
                  [
                      ("Bewertungsergebnis", f"Das Objekt {address} (Nutzfläche {area:,} m²) wird auf Basis des Ertragswertverfahrens mit {eur(rent_pa * random.randint(12, 22))} bewertet.\n\nBewertungsgrundlage: RICS Red Book / ImmoWertV 2021\nNutzungsart: {use}\nStichtag: {ds(2023, 6, 30)}\nGutachter: [Name, Sachverständiger],  Jones Lang LaSalle SE"),
                  ])

    # Environmental Assessments
    for sub in [SUBS[0], SUBS[2], SUBS[4]]:
        make_docx(folder, f"REA_Umweltgutachten_{sub['short']}_2022.docx",
                  sub["name"], sub["city"], sub["hrb"],
                  f"Umwelt- und Altlastenuntersuchung – {sub['name']} – {sub['city']}",
                  [
                      ("1. Beauftragung und Umfang", f"Die {AG['name']} hat die TÜV SÜD Industrie Service GmbH beauftragt, eine Phase I und Phase II Umweltuntersuchung des Betriebsgeländes {sub['city']} durchzuführen."),
                      ("2. Ergebnisse Phase I", "Auf Basis der Recherche von Altlastenregistern, historischen Betriebsunterlagen und Ortsbegehung wurden keine wesentlichen Hinweise auf eine schädliche Bodenveränderung oder Altlasten festgestellt."),
                      ("3. Ergebnisse Phase II", f"Die Bodenuntersuchungen (10 Schürfpunkte, Grundwasserpegel) ergaben: Alle untersuchten Parameter (MKW, PAK, Schwermetalle) lagen unterhalb der Prüfwerte gemäß BBodSchV. Fazit: Keine weiteren Maßnahmen erforderlich. Gesamtbewertung: Unauffällig."),
                  ])


# ═══════════════════════════════════════════════════════════════════════════
# 19 – M&A / TRANSAKTIONEN
# ═══════════════════════════════════════════════════════════════════════════

def gen_19():
    folder = "19_MA_Transaktionen"
    print(f"\n[19] M&A / Transaktionen …")

    transactions = [
        {
            "name": "Erwerb Brennhagen Czech s.r.o.",
            "year": 2018, "price": 28_500_000,
            "target": "Brennhagen Czech s.r.o.", "seller": "Management-Team und Family Office Brno"
        },
        {
            "name": "Erwerb Brennhagen Hungary Kft.",
            "year": 2019, "price": 19_200_000,
            "target": "Brennhagen Hungary Kft.", "seller": "Industriepark Győr GmbH & JV-Partner"
        },
        {
            "name": "Konsolidierung Shanghai JV → 100 %",
            "year": 2021, "price": 8_400_000,
            "target": "Brennhagen Electronics (China) Ltd.", "seller": "Lokaler JV-Partner China"
        },
    ]

    for tx in transactions:
        tname = sfn(tx["name"][:30])
        y = tx["year"]
        price = tx["price"]

        # LOI
        make_docx(folder, f"LOI_{tname}_{y}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Letter of Intent (LOI) – {tx['name']}",
                  [
                      ("Absichtserklärung", f"Die {AG['name']} (Käufer) und {tx['seller']} (Verkäufer) bekunden hiermit ihre gegenseitige Absicht, die {tx['target']} (Zielgesellschaft) zu erwerben bzw. zu veräußern.\n\nIndikativ vereinbarter Kaufpreis (vorbehaltlich): EUR {price:,.0f}\nAngestrebte Transaktionsstruktur: Share Deal (Kauf von 100 % der Anteile)\nGeplanter Signing-Termin: Q{random.randint(1,4)}/{y}\nGeplanter Closing-Termin: ca. 3 Monate nach Signing"),
                      ("Ausschließlichkeit", f"Die Parteien vereinbaren eine Exklusivitätsphase von 60 Tagen ab Unterzeichnung dieses LOI. Während dieser Zeit führt der Verkäufer keine parallelen Verkaufsgespräche.\n\nDieses LOI ist nicht rechtsverbindlich, ausgenommen die Regelungen zu Vertraulichkeit und Kosten."),
                  ], confidential=True)

        # NDA
        make_docx(folder, f"NDA_{tname}_{y}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Vertraulichkeitsvereinbarung (NDA) – {tx['name']}",
                  [
                      ("§ 1 Gegenstand", f"Im Zusammenhang mit der beabsichtigten Transaktion ({tx['name']}) vereinbaren die {AG['name']} und {tx['seller']} einen umfassenden Schutz vertraulicher Informationen."),
                      ("§ 2 Vertrauliche Informationen", "Als vertraulich gelten alle im Zusammenhang mit der Due Diligence ausgetauschten Informationen, insbesondere Finanzdaten, Kundenlisten, technische Unterlagen, Verträge und Personaldaten."),
                      ("§ 3 Laufzeit", "Diese NDA gilt ab Unterzeichnung für 5 Jahre, unabhängig davon, ob die Transaktion zustande kommt."),
                  ], confidential=True)

        # SPA Excerpts
        make_docx(folder, f"SPA_{tname}_{y}_Auszug.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Share Purchase Agreement (SPA) – {tx['name']} – AUSZUG",
                  [
                      ("Parteien", f"Käufer: {AG['name']}\nVerkäufer: {tx['seller']}\nZielgesellschaft: {tx['target']}"),
                      ("§ 1 Kaufgegenstand", f"(1) Der Verkäufer veräußert und überträgt, und der Käufer erwirbt, alle Anteile an {tx['target']} (100 %).\n\n(2) Der Übergang erfolgt mit wirtschaftlicher Wirkung zum Signing-Datum."),
                      ("§ 2 Kaufpreis", f"(1) Der Kaufpreis beträgt EUR {price:,.0f} (Enterprise Value).\n\n(2) Der Kaufpreis ist auf Basis der Locked-Box-Methode berechnet. Locked-Box-Stichtag: 31. Dezember {y-1}.\n\n(3) Es werden keine Earn-Out-Komponenten vereinbart. / Es wird eine Earn-Out-Komponente von max. EUR {int(price*0.15):,.0f} vereinbart."),
                      ("§ 3 Garantien des Verkäufers", "(1) Der Verkäufer gibt umfassende W&I-Garantien (Warranties) gemäß Anlage 4.\n\n(2) Haftungsbegrenzung: De-Minimis EUR 50.000, Basket EUR 250.000, Cap EUR " + eur(int(price * 0.50)) + ".\n\n(3) Verjährungsfrist: 24 Monate nach Closing (steuerliche Garantien: 7 Jahre)."),
                  ], confidential=True)

        # Due Diligence Report
        make_docx(folder, f"DD_Bericht_{tname}_{y}.docx",
                  AG["recht"], "Frankfurt am Main", "DD-Report Ref.: " + tname,
                  f"Vendor / Buy-Side Due Diligence Report – {tx['name']} – {y}",
                  [
                      ("Executive Summary", f"Im Auftrag der {AG['name']} hat {AG['recht']} eine rechtliche und kommerzielle Due Diligence der {tx['target']} durchgeführt. Gesamtbewertung: Transaktion empfohlen, vorbehaltlich der in diesem Bericht identifizierten Risiken."),
                      ("Rechtliche Findings", [
                          ["Kategorie", "Finding", "Schweregrad", "Empfehlung"],
                          ["Gesellschaftsrecht", "Gesellschaftsvertrag entspricht lokalem Recht", "Keine", "-"],
                          ["Verträge", "3 Kundenverträge mit Change-of-Control-Klauseln", "Mittel", "Kundenkonsent einholen"],
                          ["IP", "Alle Patente auf Zielgesellschaft eingetragen", "Keine", "-"],
                          ["Arbeitsrecht", "Offene Arbeitnehmerforderungen (EUR 150.000)", "Gering", "Rückstellung"],
                          ["Steuern", "Offene BP 2015–2018, Risiko EUR 280.000", "Mittel", "Steuerabtretungsklausel"],
                      ]),
                      ("Empfohlene Vertragsanpassungen", "Einbeziehung einer Steuerfreistellungsklausel, Sicherheitseinbehalt EUR 500.000 für 24 Monate, Change-of-Control-Konsente von 3 Kunden vor Closing."),
                  ], confidential=True)

        # Integration Plan
        make_docx(folder, f"Integration_Plan_{tname}_{y}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Post-Merger Integration Plan – {tx['name']}",
                  [
                      ("1. Integrationsüberblick", f"Die {tx['target']} wird in den Brennhagen-Konzern integriert. Ziele: Synergieerzielung, Harmonisierung von Prozessen und Systemen, kulturelle Integration."),
                      ("2. Synergien (Quantifiziert)", [
                          ["Synergiekategorie", "p.a. (EUR Mio.)", "Einmalinvestitionen", "Zeitraum"],
                          ["Einkauf / Materialkosten", f"{random.uniform(0.5,2.0):.1f}", f"{random.uniform(0.2,0.8):.1f}", "Jahr 2"],
                          ["IT-Harmonisierung (SAP Rollout)", f"{random.uniform(0.3,1.0):.1f}", f"{random.uniform(1.0,3.0):.1f}", "Jahr 3"],
                          ["G&A-Kosteneinsparung", f"{random.uniform(0.5,1.5):.1f}", "–", "Jahr 2"],
                          ["Revenue Synergien (Cross-Selling)", f"{random.uniform(1.0,4.0):.1f}", "–", "Jahr 3"],
                      ]),
                      ("3. 100-Tage-Plan", f"In den ersten 100 Tagen nach Closing werden folgende Prioritäten umgesetzt:\n\n1. Integration in Konzerncontrolling und Berichtswesen (bis Tag 30)\n2. HR-Harmonisierung (Entgeltsystem, Betriebsrente) bis Tag 60\n3. Abschluss IT-Bestandsaufnahme und Start SAP-Rollout-Planung bis Tag 90\n4. Kommunikation an Kunden und Lieferanten: Day 1"),
                  ], confidential=True)

    # NDA Register
    make_xlsx(folder, "REA_NDA_Register_2023.xlsx",
              "NDA-Register – Brennhagen Elektronik AG – 2023",
              [
                  ("NDA Register",
                   ["Nr.", "Vertragspartner", "Typ", "Abschluss", "Ablauf", "Transaktionsbezug"],
                   [
                       [f"NDA-2023-{i:03d}", random.choice(["PE-Investor A", "Strategischer Investor B", "Berater C", "OEM-Kunde (M&A)"]), random.choice(["Unilateral","Bilateral"]), ds(2023, random.randint(1,10), random.randint(1,25)), ds(2028, random.randint(1,10), random.randint(1,25)), random.choice(["M&A-Prozess", "Technologiekooperation", "Lieferantenbewertung", "Investitionsprüfung"])]
                       for i in range(1, 16)
                   ],
                   [15, 35, 15, 18, 18, 30]),
              ])

    # Teaser / IM excerpt
    make_docx(folder, "REA_Teaser_Strategischer_Prozess_2024_SANITIZED.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Unternehmens-Teaser – Strategischer Prozess 2024 – [SANITIZED]",
              [
                  ("Company Overview (Sanitized)", f"{AG['name']} is a leading Tier-1 automotive electronics supplier listed on the TecDAX ({AG['stock']}). The company develops and manufactures ECUs, ADAS camera systems, infotainment modules and EV battery management systems."),
                  ("Key Financials", [
                      ["Metric", "2022A", "2023A", "2024E"],
                      ["Revenue (EUR m)", f"{AG['revenue_2022']//1_000_000:.0f}", f"{AG['revenue_2023']//1_000_000:.0f}", f"{AG['revenue_2024e']//1_000_000:.0f}"],
                      ["EBITDA (EUR m)", f"{AG['ebitda_2022']//1_000_000:.0f}", f"{AG['ebitda_2023']//1_000_000:.0f}", f"{AG['ebitda_2024e']//1_000_000:.0f}"],
                      ["EBITDA margin", f"{AG['ebitda_2022']/AG['revenue_2022']*100:.1f} %", f"{AG['ebitda_2023']/AG['revenue_2023']*100:.1f} %", "-"],
                  ]),
                  ("Investment Highlights", "1. TecDAX-listed, EUR 648m market cap (2023)\n2. Blue-chip OEM customer base: BMW (28%), VW (22%), Mercedes (18%)\n3. EV/ADAS growth platforms: BatteryMS-12 and ADAS-Vision 4D\n4. Pan-European manufacturing footprint (DE, PL, CZ, HU)\n5. Strong IP portfolio: 18+ patents, IATF 16949 certified"),
              ], confidential=True)


# ═══════════════════════════════════════════════════════════════════════════
# 20 – STRATEGIE / VORSTAND
# ═══════════════════════════════════════════════════════════════════════════

def gen_20():
    folder = "20_Strategie_Vorstand"
    print(f"\n[20] Strategie / Vorstand …")

    # Annual Strategy Papers
    for y in [2021, 2022, 2023]:
        make_docx(folder, f"REA_Strategiepapier_{y}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Konzernstrategie {y}–{y+4} – Brennhagen Elektronik AG",
                  [
                      ("1. Executive Summary", f"Die Strategie der {AG['name']} für die Periode {y}–{y+4} fokussiert auf drei strategische Wachstumsfelder: (1) Elektrifizierung (EV), (2) Autonomes Fahren (ADAS) und (3) Software-Defined Vehicle (SDV). Ziel ist ein Umsatzwachstum auf EUR 842 Mio. bis 2028 bei einer EBITDA-Marge von 13,5 %."),
                      ("2. Marktanalyse", f"Der globale Markt für Automotive-Elektronik wächst bis 2030 auf USD 390 Mrd. (CAGR 8,3 %). Treiber: EV-Durchdringung (2030e: 45 % aller Neuzulassungen), ADAS Level 3/4, Over-the-Air (OTA) Updates und Cybersecurity-Anforderungen."),
                      ("3. Wettbewerbsposition", [
                          ["Wettbewerber", "Umsatz (Automotive Electronics)", "Stärken", "Schwächen"],
                          ["Continental AG", "EUR 18 Mrd.", "Marktgröße, breite Produktpalette", "Komplexität, Margendruck"],
                          ["Bosch Mobility", "EUR 52 Mrd.", "Technologiebreite, globale Präsenz", "Größe, Trägheit"],
                          ["Aptiv plc", "EUR 17 Mrd.", "Software-Fokus, SDV", "Hardware-Transformation"],
                          [AG["name"] + " (REA)", f"EUR {AG[f'revenue_{y}']//1_000_000:.0f} Mio.", "ADAS, EV, Kundennähe, Agilität", "Skalierung, Finanzmacht"],
                      ]),
                      ("4. Strategische Initiativen", f"(1) EV-Wachstumsprogramm: Investitionen EUR 45 Mio. in BatteryMS-12 Next Gen (2024–2026)\n(2) ADAS-Leadership: R&D EUR {random.randint(25,40)} Mio. p.a. + JDA mit BMW Group\n(3) Software-Transformation: Aufbau 150 zusätzlicher Softwareingenieure bis {y+2}\n(4) Footprint-Optimierung: Heilbronn Expansion + Prüfung Nearshoring Marokko"),
                  ], confidential=True)

    # Capital Allocation Framework
    make_docx(folder, "REA_Capital_Allocation_Framework_2023.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Capital Allocation Framework – Brennhagen Elektronik AG",
              [
                  ("1. Kapitalallokationsprinzipien", "Der Vorstand der Brennhagen Elektronik AG folgt folgenden Prioritäten bei der Kapitalallokation:\n\n1. Organisches Wachstum (Capex, R&D): Primäre Priorität, Ziel 5,5 % des Umsatzes\n2. Dividende: Progressive Dividendenpolitik, Payout-Ratio 25–35 % des Nettogewinns\n3. M&A: Strategische Akquisitionen in EV/ADAS, Net Leverage ≤ 2,5x post-deal\n4. Aktienrückkauf: Opportunistisch bei unterbewerteten Kursen"),
                  ("2. Kapitalstrukturziel", "Ziel-Kapitalstruktur: Net Debt/EBITDA 0,5x–1,5x im Normalzyklus. Maximale Verschuldung temporär bis 2,5x für transformative Akquisitionen."),
              ], confidential=True)

    # 5-Year Strategic Plan (Tabelle)
    make_xlsx(folder, "REA_LRP_Strategieplan_2024_2028_Detail.xlsx",
              "Langfristiger Strategieplan 2024–2028 – REA",
              [
                  ("Strategische KPIs",
                   ["KPI", "2023A", "2024E", "2025P", "2026P", "2027P", "2028P"],
                   [
                       ["Umsatz (EUR Mio.)", f"{AG['revenue_2023']//1_000_000:.0f}", f"{AG['revenue_2024e']//1_000_000:.0f}", "698", "741", "789", "842"],
                       ["EBITDA-Marge", f"{AG['ebitda_2023']/AG['revenue_2023']*100:.1f} %", "11,2 %", "11,8 %", "12,3 %", "12,9 %", "13,5 %"],
                       ["EV/ADAS Umsatzanteil", "32 %", "38 %", "44 %", "50 %", "56 %", "62 %"],
                       ["Capex % Umsatz", "5,0 %", "5,5 %", "5,5 %", "5,9 %", "5,8 %", "5,7 %"],
                       ["R&D % Umsatz", "6,5 %", "7,0 %", "7,0 %", "7,2 %", "7,3 %", "7,5 %"],
                       ["Mitarbeiter", f"{AG['employees_2023']:,}", "3.620", "3.780", "3.950", "4.100", "4.280"],
                   ],
                   [30, 12, 12, 12, 12, 12, 12]),
              ])

    # Market studies + competitor analyses
    for topic in ["EV_Marktanalyse_2023", "ADAS_Markt_2024", "Wettbewerbsanalyse_Tier1_2023", "Marktanalyse_China_Automobil_2023"]:
        make_docx(folder, f"REA_{topic}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  topic.replace("_", " ").replace("REA ", ""),
                  [
                      ("Executive Summary", f"Diese Marktstudie / Wettbewerbsanalyse wurde von der Strategieabteilung der {AG['name']} unter Einbeziehung externer Datenquellen (McKinsey EV Report, S&P Global Mobility, IHS Markit) erstellt. Stand: Q4 {2023 if '2023' in topic else 2024}."),
                      ("Kernergebnisse", f"1. Marktgröße und Wachstum: Der adressierbare Markt für {topic.split('_')[0]} beträgt EUR {random.randint(50,400)} Mrd. (global) mit einem CAGR von {random.uniform(6,15):.1f} % bis 2030.\n2. Wettbewerber: Hauptwettbewerber sind Continental, Bosch Mobility, Aptiv, Visteon und ZF Friedrichshafen.\n3. REA-Marktanteil: Aktuell {random.uniform(1,5):.1f} % global; Ziel: {random.uniform(2,7):.1f} % bis 2028."),
                  ], confidential=True)

    # Board Presentations
    for y in [2022, 2023]:
        for q in [1, 2, 3, 4]:
            make_docx(folder, f"REA_Vorstandspräsentation_{y}_Q{q}.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"Vorstandspräsentation Q{q} {y} – Strategisches Update",
                      [
                          ("Agenda", f"1. Geschäftsentwicklung Q{q} {y}\n2. Strategisches Update EV/ADAS\n3. Kapitalallokation\n4. Personalentwicklungen\n5. Risikoübersicht\n6. Ausblick Q{q+1 if q<4 else 1} {y if q<4 else y+1}"),
                          ("Highlights", f"Umsatz Q{q} {y}: {eur(int(AG[f'revenue_{y}']/4 * random.uniform(0.9,1.1)))}\nEBITDA-Marge: {AG[f'ebitda_{y}']/AG[f'revenue_{y}']*100:.1f} %\nWesentliche Ereignisse: [Details]"),
                      ], confidential=True)


# ═══════════════════════════════════════════════════════════════════════════
# 21 – BETRIEBSRÄTE
# ═══════════════════════════════════════════════════════════════════════════

def gen_21():
    folder = "21_Betriebsraete"
    print(f"\n[21] Betriebsräte …")

    de_subs = [s for s in SUBS if s["country"] == "DE"]

    for sub in de_subs:
        sname = sub["name"]
        sshort = sub["short"]
        shrb = sub["hrb"]
        scity = sub["city"]

        # Betriebsvereinbarungen (15+ per entity)
        bv_list = [
            f"Arbeitszeit_und_Arbeitszeiterfassung",
            f"Homeoffice_Telearbeit",
            f"Betriebliche_Altersversorgung",
            f"Gesundheitsschutz_BGM",
            f"IT_Nutzung_und_Ueberwachung",
            f"Leistungs_und_Verhaltenskontrolle",
            f"Entgeltrahmenabkommen_Eingruppierung",
            f"Urlaubsplanung",
            f"Schichtbetrieb_Praemien",
            f"Weiterbildung_und_Qualifizierung",
            f"Gleichstellung_und_Antidiskriminierung",
            f"Essenszuschuss_Kantine_{scity}",
            f"Parkplatzregelung_{scity}",
            f"Betriebliches_Vorschlagswesen",
            f"Kurzarbeitsregelung_2020",
            f"Sozialer_Interessenausgleich_2023",
        ]
        for bv in bv_list:
            make_docx(folder, f"{sshort}_BV_{bv}_2022.docx",
                      sname, scity, shrb,
                      f"Betriebsvereinbarung – {bv.replace('_', ' ')} – {sname}",
                      [
                          ("Vertragsparteien", f"{sname} (Arbeitgeber) und der Betriebsrat der {sname} (Arbeitnehmervertretung)"),
                          ("§ 1 Geltungsbereich", f"Diese Betriebsvereinbarung gilt für alle Arbeitnehmerinnen und Arbeitnehmer der {sname}, {scity}."),
                          ("§ 2 Regelungsinhalt", f"Die Parteien vereinbaren bezüglich {bv.replace('_', ' ')} Folgendes: [Detaillierte Regelung gemäß BetrVG und einschlägigen Tarifverträgen]."),
                          ("§ 3 Laufzeit und Kündigung", "Diese Betriebsvereinbarung tritt am 1. Januar 2022 in Kraft und kann mit einer Frist von 3 Monaten zum Quartalsende schriftlich gekündigt werden."),
                          (None, f"\n{scity}, den {ds(2021, 12, 15)}\n\nFür den Arbeitgeber: [Geschäftsführer]\nFür den Betriebsrat: [BR-Vorsitzender]"),
                      ])

        # Works council election protocol
        make_docx(folder, f"{sshort}_BR_Wahl_2022_Protokoll.docx",
                  sname, scity, shrb,
                  f"Protokoll Betriebsratswahl 2022 – {sname}",
                  [
                      ("Wahlergebnis", f"Am 15. März 2022 fand die reguläre Betriebsratswahl statt. Wahlbeteiligung: {random.randint(72,91)} %. Gewählte Mitglieder: {random.randint(7,15)} (davon {random.randint(3,8)} Frauen)."),
                  ])

        # Monthly BR Meeting Minutes (2022 + 2023, 12 per year per entity)
        for y in [2022, 2023]:
            for m in range(1, 13):
                make_docx(folder, f"{sshort}_BR_Protokoll_{y}_{m:02d}.docx",
                          sname, scity, shrb,
                          f"Protokoll Betriebsratssitzung {MONTHS_DE[m]} {y} – {sname}",
                          [
                              ("Anwesende", f"BR-Vorsitzender: [Name]; anwesend: {random.randint(7,12)} von {random.randint(9,15)} BR-Mitgliedern\nSitzungsdatum: {ds(y, m, random.randint(10,25))}\nOrt: Betriebsratsraum, {scity}"),
                              ("Tagesordnung", f"1. Genehmigung letztes Protokoll\n2. Bericht Geschäftsleitung (vertreten durch HR-Manager)\n3. Laufende Betriebsratsthemen\n4. Aktuelle Betriebsvereinbarungen\n5. Verschiedenes"),
                              ("Wesentliche Beschlüsse", f"Monat {MONTHS_DE[m]} {y}: Keine außerordentlichen Beschlüsse. Routinemäßige Abstimmungen zu Urlaubslisten, Überstunden und Einstellungen erfolgten einstimmig."),
                          ])

        # Information and Consultation Records
        make_docx(folder, f"{sshort}_BR_Unterrichtung_Betriebsaenderung_2023.docx",
                  sname, scity, shrb,
                  f"Unterrichtung und Beratung Betriebsänderung – {sname} – 2023",
                  [
                      ("Unterrichtung", f"Der Arbeitgeber unterrichtet den Betriebsrat gemäß § 111 BetrVG über die geplante Betriebsänderung: Einführung neuer Fertigungstechnologie (Industrie 4.0) in Linie {random.randint(2,6)}."),
                      ("Beratung", "Die Parteien haben sich in mehreren Gesprächen über die Betriebsänderung beraten. Folgende wesentliche Aspekte wurden erörtert: Qualifikationsbedarf, Personalplanung, Auswirkungen auf Schichtplan."),
                      ("Einigung", "Eine Einigung über einen Interessenausgleich konnte erzielt werden. Ein Sozialplan ist nicht erforderlich, da keine Entlassungen stattfinden."),
                  ])

    # Gesamtbetriebsrat
    make_docx(folder, "REA_GBR_Protokoll_2023_Q1.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Protokoll Gesamtbetriebsratssitzung Q1 2023 – Brennhagen Elektronik AG",
              [
                  ("Anwesende", f"GBR-Vorsitzender: Dipl.-Kfm. Heinrich Baumeister; Arbeitgeberseite: {AG['coo']} (COO); anwesend: alle 3 örtlichen Betriebsräte (REG, RSG, RHO)"),
                  ("Tagesordnung", "1. Konzernweite Strukturmaßnahmen\n2. Bericht Vorstand – Wirtschaftliche Lage\n3. Konzernweite Betriebsvereinbarungen\n4. Europäischer Betriebsrat (EBR)-Vorbereitung\n5. Verschiedenes"),
                  ("Beschlüsse", "Keine außerordentlichen Beschlüsse. Weitere Verhandlungen zur konzernweiten BV Homeoffice werden fortgesetzt."),
              ])


# ═══════════════════════════════════════════════════════════════════════════
# 22 – PENSIONEN / bAV
# ═══════════════════════════════════════════════════════════════════════════

def gen_22():
    folder = "22_Pensionen_bAV"
    print(f"\n[22] Pensionen / bAV …")

    # Pension Plan Descriptions
    bav_plans = [
        ("REA_bAV_Direktzusage_Plan_REG_Heilbronn", "Direktzusage – Altzusagen (Übergangsplan)", SUBS[0]),
        ("REA_bAV_Pensionskasse_Plan_REA_AG", "Pensionskasse – Mitarbeiterbeitrag", AG),
        ("REA_bAV_Entgeltumwandlung_Plan_2020", "Entgeltumwandlung (Deferred Compensation)", AG),
        ("REA_bAV_Neuzusage_Plan_2015", "Leistungsorientierter Neuzusage-Plan ab 2015", AG),
        ("REA_bAV_BZML_Beitragsorientierter_Plan", "Beitragszusage mit Mindestleistung (BZML)", SUBS[1]),
    ]
    for fname, title, entity in bav_plans:
        ename = entity["name"] if isinstance(entity, dict) and "name" in entity else AG["name"]
        make_docx(folder, fname + "_2023.docx",
                  ename, AG["full_address"] if entity is AG else entity.get("city",""), AG["hrb"] if entity is AG else entity.get("hrb",""),
                  f"Pensionsplan-Beschreibung – {title}",
                  [
                      ("1. Planübersicht", f"Plan-Bezeichnung: {title}\nArbeitgeber: {ename}\nPlantyp: {random.choice(['Leistungszusage (DB)','Beitragszusage (DC)','Beitragszusage mit Mindestleistung'])}\nBegünstigte: {random.randint(100,500)} aktive Anwärter, {random.randint(50,200)} Rentner"),
                      ("2. Leistungsstruktur", f"Altersrente: {random.uniform(0.3,0.6):.1f} % des pensionsfähigen Gehalts pro Dienstjahr\nInvalidenrente: {random.uniform(0.2,0.4):.1f} % des pensionsfähigen Gehalts\nHinterbliebenenrente: {random.randint(50,70)} % der Altersrente"),
                      ("3. Finanzierung", f"Der Plan ist über ein CTA (Contractual Trust Arrangement) mit einem Treuhandvermögen von EUR {random.randint(15,80)} Mio. ausfinanziert."),
                  ])

    # Actuarial Valuations (annual per entity)
    for sub in [AG] + SUBS[:3]:
        for y in [2021, 2022, 2023]:
            sname = sub["name"] if "name" in sub else AG["name"]
            make_docx(folder, f"Aktuarbericht_{sub.get('short','REA')}_{y}.docx",
                      "Willis Towers Watson GmbH", "Eschersheimer Landstraße 50, 60322 Frankfurt", "Aktuar-Ref.: REA-" + str(y),
                      f"Aktuarielles Gutachten Pensionsverpflichtungen – {sname} – {y}",
                      [
                          ("Bewertungsparameter", [
                              ["Parameter", f"{y}", f"{y-1}"],
                              ["Diskontierungszinssatz (IFRS)", f"{random.uniform(3.5,4.5):.2f} %", f"{random.uniform(1.5,3.0):.2f} %"],
                              ["Rentensteigerungsrate", "2,0 %", "2,0 %"],
                              ["Gehaltserhöhungsrate", "3,0 %", "3,0 %"],
                              ["Sterbetafeln", "Heubeck 2018 G", "Heubeck 2018 G"],
                          ]),
                          ("Bewertungsergebnis", [
                              ["Kennzahl", f"31.12.{y} (TEUR)", f"31.12.{y-1} (TEUR)"],
                              ["DBO (Defined Benefit Obligation)", f"{random.randint(40000,120000):,}", f"{random.randint(38000,115000):,}"],
                              ["Planvermögen (Fair Value)", f"{random.randint(30000,100000):,}", f"{random.randint(28000,98000):,}"],
                              ["Nettopensionsverpflichtung", f"{random.randint(5000,25000):,}", f"{random.randint(6000,30000):,}"],
                          ]),
                      ])

    # Pension individual commitments (sample)
    for i, (person, rolle, _, _) in enumerate([(AG["ceo"], "CEO", "", ""), (AG["cfo"], "CFO", "", ""), (AG["coo"], "COO", "", "")]):
        make_docx(folder, f"bAV_Einzelzusage_{sfn(person)}_2020.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Versorgungszusage (bAV) – {person} – {rolle}",
                  [
                      ("Versorgungszusage", f"Die {AG['name']} erteilt {person} ({rolle}) folgende betriebliche Versorgungszusage:\n\nAltersrente ab dem 65. Lebensjahr: EUR {random.randint(8000, 20000):,} monatlich\nInvalidenrente: 70 % der Altersrente\nHinterbliebenenrente: 60 % der Altersrente\n\nVerfallbarkeit: Nach 3 Jahren Betriebszugehörigkeit unverfallbar (§ 1b BetrAVG).\n\nÜbergang bei Insolvenz: Sicherung über PSV (Pensions-Sicherungs-Verein) bis zur Hälfte der Beitragsbemessungsgrenze."),
                  ], confidential=True)

    # CTA Documentation
    make_docx(folder, "REA_CTA_Treuhandvertrag_2018.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Contractual Trust Arrangement (CTA) – Treuhandvertrag – REA Pension Trust",
              [
                  ("1. Parteien", f"Treugeber: {AG['name']}\nTreuhänder: Deutsche Bank AG als Treuhänder (REA Pension Trust e.V. als zwischengeschalteter Träger)\nVermögensverwaltung: [Externe Kapitalanlagegesellschaft]"),
                  ("2. Zweck", "Das CTA sichert die Erfüllung der Pensionsverpflichtungen der Brennhagen-Gruppe gegenüber berechtigten Versorgungsanwärtern und -empfängern im Insolvenzfall."),
                  ("3. Treuhandvermögen", f"Das aktuelle Treuhandvermögen beläuft sich auf EUR {random.randint(60,120)} Mio. (Stand {ds(2023, 12, 31)}). Die Anlage erfolgt gemäß der von den Parteien vereinbarten Anlagerichtlinie (balanced, max. 40 % Aktien)."),
              ], confidential=True)

    # Pension Risk Analysis
    make_docx(folder, "REA_Pensionsrisiko_Analyse_2023.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Pensionsrisiko-Analyse 2023 – Brennhagen Elektronik AG",
              [
                  ("1. Sensitivitätsanalyse", [
                      ["Szenario", "Auswirkung auf DBO"],
                      ["+0,5 % Diskontierungszins", f"-{random.randint(5,12)} % DBO"],
                      ["-0,5 % Diskontierungszins", f"+{random.randint(5,12)} % DBO"],
                      ["+1 % Gehaltserhöhungsrate", f"+{random.randint(2,6)} % DBO"],
                      ["+1 Jahr Lebenserwartung", f"+{random.randint(2,5)} % DBO"],
                  ]),
                  ("2. Risiken und Maßnahmen", "Wesentliche Risiken: (1) Zinsniveaurisiko – gedeckt durch CTA-Anlagerichtlinie, (2) Langlebigkeitsrisiko – Überprüfung Sterbetafeln jährlich, (3) Inflationsrisiko – begrenzt durch Rentensteigerungsformeln."),
              ])


# ═══════════════════════════════════════════════════════════════════════════
# 23 – PROJEKTE / PROGRAMME
# ═══════════════════════════════════════════════════════════════════════════

def gen_23():
    folder = "23_Projekte_Programme"
    print(f"\n[23] Projekte / Programme …")

    for prj_id, prj_name, start, end, customer in PROJECTS:
        psfn = sfn(prj_name[:25])

        # Project Charter
        make_docx(folder, f"{prj_id}_Charter_{psfn}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Projekt-Steckbrief (Charter) – {prj_name}",
                  [
                      ("1. Projektübersicht", [
                          ["Feld", "Inhalt"],
                          ["Projekt-ID", prj_id],
                          ["Projektname", prj_name],
                          ["Auftraggeber", customer],
                          ["Projektleiter", random.choice(ALL_NAMES)],
                          ["Start", start],
                          ["Geplantes Ende", end],
                          ["Budget", eur(random.randint(2, 30) * 1_000_000)],
                          ["Status", random.choice(["In Arbeit", "In Arbeit", "Abgeschlossen", "In Planung"])],
                      ]),
                      ("2. Projektziele", f"Ziel dieses Projekts ist {prj_name}. Erwartete Ergebnisse: (1) Technisches Ergebnis / Produkt, (2) Zertifizierung / Freigabe, (3) Kundenzufriedenheit, (4) Budgeteinhaltung."),
                      ("3. Risiken", [
                          ["Risiko", "Wahrscheinlichkeit", "Auswirkung", "Maßnahme"],
                          ["Ressourcenengpass (Ingenieure)", "Mittel", "Hoch", "Frühzeitige Ressourcenplanung"],
                          ["Halbleiterverfügbarkeit", "Gering", "Hoch", "Dual-Source-Strategie"],
                          ["Scope Creep", "Mittel", "Mittel", "Change Control Process"],
                          ["Kundenseitiger Verzug", "Gering", "Mittel", "Eskalationsmanagement"],
                      ]),
                  ])

        # Requirements Specification
        make_docx(folder, f"{prj_id}_Lastenheft_{psfn}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Lastenheft – {prj_name}",
                  [
                      ("1. Einleitung und Scope", f"Dieses Lastenheft beschreibt die Anforderungen an {prj_name}. Auftraggeber: {customer}. Version: 1.3."),
                      ("2. Systemanforderungen", [
                          ["Req-ID", "Anforderung", "Priorität", "Verifikation"],
                          [f"REQ-{prj_id}-001", "Betriebstemperatur -40 °C bis +125 °C", "MUST", "Test"],
                          [f"REQ-{prj_id}-002", "AUTOSAR Adaptive Platform konform", "MUST", "Review + Analyse"],
                          [f"REQ-{prj_id}-003", f"Boot-Zeit < {random.randint(100,500)} ms", "SHOULD", "Messung"],
                          [f"REQ-{prj_id}-004", "ISO 26262 ASIL-D (sicherheitskritische Teile)", "MUST", "Safety Analyse"],
                          [f"REQ-{prj_id}-005", f"EMV gemäß CISPR 25 Level {random.randint(3,5)}", "MUST", "EMV-Test"],
                      ]),
                  ])

        # Technical Specification
        make_docx(folder, f"{prj_id}_Pflichtenheft_{psfn}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Pflichtenheft (Technical Specification) – {prj_name}",
                  [
                      ("1. Systemarchitektur", f"Die technische Architektur für {prj_name} basiert auf [Prozessor, FPGA, MCU]. Das System nutzt {AG['plm']} für die Versionsverwaltung."),
                      ("2. Softwarearchitektur", "Die Software folgt dem AUTOSAR-Architekturstandard. Hardware Abstraction Layer (HAL): [Beschreibung]. Application Layer: [Beschreibung]. Kommunikation: CAN-FD / Ethernet (SOME/IP)."),
                      ("3. Schnittstellen", [
                          ["Schnittstelle", "Protokoll", "Datenrate", "Partner"],
                          ["Fahrzeug-CAN-Bus", "CAN-FD 5 Mbit/s", "5 Mbit/s", "ECU-Verbund"],
                          ["Diagnoseschnittstelle", "UDS/ISO 14229", "250 kbit/s", "Werkstatt-Tester"],
                          ["OTA-Update", "HTTPS/TLS 1.3", "Variabel", "OEM-Backend"],
                      ]),
                  ])

        # Monthly Status Reports (6 months)
        start_year = int(start.split("-")[0])
        start_month = int(start.split("-")[1])
        for i in range(6):
            m = (start_month + i - 1) % 12 + 1
            y = start_year + (start_month + i - 1) // 12
            make_docx(folder, f"{prj_id}_Statusbericht_{y}_{m:02d}_{psfn}.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"Projektstatusbericht {MONTHS_DE[m]} {y} – {prj_name}",
                      [
                          ("Ampelbewertung", [
                              ["Dimension", "Status", "Trend"],
                              ["Zeitplan", random.choice(["🟢 On Track", "🟢 On Track", "🟡 Leichte Verzögerung"]), "→"],
                              ["Budget", random.choice(["🟢 Im Budget", "🟢 Im Budget", "🟡 +3 % Überschreitung"]), "→"],
                              ["Qualität", random.choice(["🟢 OK", "🟢 OK", "🟡 Offene CAPA"]), "→"],
                              ["Risiken", "🟢 Unter Kontrolle", "→"],
                          ]),
                          ("Meilensteine", f"Letzter Meilenstein: {random.choice(['Systemdesign Review', 'Prototypenlieferung', 'EMV-Test bestanden', 'PPAP Level 2 genehmigt'])}\nNächster Meilenstein: {random.choice(['Gateway Review', 'Kundenpräsentation', 'SOP Kick-off', 'Validierungstest'])} bis {ds(y, min(m+2,12), 30)}"),
                      ])

        # Gate Review Record (APQP)
        for gate in ["G1_Konzept", "G2_Systemdesign", "G3_Detailentwicklung", "G4_Validierung"]:
            make_docx(folder, f"{prj_id}_Gate_{gate}_{psfn}.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"Gate Review – {gate.replace('_',' ')} – {prj_name}",
                      [
                          ("Gate-Entscheidung", [
                              ["Kriterium", "Bewertung", "Kommentar"],
                              ["Requirement-Vollständigkeit", f"{random.randint(90,100)} %", ""],
                              ["Design-Review abgeschlossen", random.choice(["Ja","Ja","Mit Auflagen"]), ""],
                              ["Risiken bewertet", "Ja", "Risikomatrix aktuell"],
                              ["Budget freigegeben", "Ja", ""],
                              ["Kundengenehmigung", random.choice(["Erhalten","Erhalten","Ausstehend"]), ""],
                          ]),
                          ("Entscheidung", f"Gate {gate}: {random.choice(['FREIGEGEBEN', 'FREIGEGEBEN MIT AUFLAGEN'])}\nEntschieden durch: Projektlenkungsausschuss, {ds(int(start.split('-')[0]), random.randint(1,12), random.randint(1,25))}"),
                      ])

        # Cost Tracking
        budget = random.randint(3, 30) * 1_000_000
        make_xlsx(folder, f"{prj_id}_Kostencontrolling_{psfn}.xlsx",
                  f"Projektkostencontrolling – {prj_name}",
                  [
                      ("Kosten",
                       ["Kostenkategorie", "Budget (TEUR)", "Ist kumuliert (TEUR)", "Prognose Restkosten", "Abweichung %"],
                       [
                           ["Personalkosten intern", f"{int(budget*0.45)//1000:,}", f"{int(budget*0.45*random.uniform(0.85,1.1))//1000:,}", f"{int(budget*0.45*0.1)//1000:,}", f"{random.uniform(-5,8):.1f} %"],
                           ["Externe Dienstleister", f"{int(budget*0.20)//1000:,}", f"{int(budget*0.20*random.uniform(0.90,1.15))//1000:,}", f"{int(budget*0.05)//1000:,}", f"{random.uniform(-3,12):.1f} %"],
                           ["Material / Prototypen", f"{int(budget*0.15)//1000:,}", f"{int(budget*0.15*random.uniform(0.80,1.20))//1000:,}", f"{int(budget*0.03)//1000:,}", f"{random.uniform(-10,15):.1f} %"],
                           ["Reise und Nebenkosten", f"{int(budget*0.05)//1000:,}", f"{int(budget*0.05*random.uniform(0.70,1.30))//1000:,}", f"{int(budget*0.01)//1000:,}", f"{random.uniform(-20,25):.1f} %"],
                           ["Sonstiges", f"{int(budget*0.15)//1000:,}", f"{int(budget*0.15*random.uniform(0.85,1.10))//1000:,}", f"{int(budget*0.04)//1000:,}", f"{random.uniform(-5,8):.1f} %"],
                       ],
                       [30, 18, 22, 22, 14]),
                  ])

        # Lessons Learned
        make_docx(folder, f"{prj_id}_Lessons_Learned_{psfn}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Lessons Learned – {prj_name}",
                  [
                      ("Positives", random.choice([
                          "Frühzeitige Einbindung des Kunden in die APQP-Phase hat den PPAP-Prozess signifikant beschleunigt.",
                          "Cross-functional Team-Setup mit wöchentlichen Jour-Fixes hat Kommunikationsprobleme minimiert.",
                          "Dual-Source-Strategie für kritische ICs hat Lieferrisiken erfolgreich gemindert.",
                      ])),
                      ("Verbesserungspotentiale", random.choice([
                          "Anforderungsmanagement sollte früher beginnen; späte Änderungen führten zu Nacharbeiten.",
                          "Budgetpuffer für externe Ressourcen war zu knapp bemessen.",
                          "Dokumentationsstandards sollten zu Projektbeginn verbindlich festgelegt werden.",
                      ])),
                      ("Maßnahmen für Folgeprojekte", "Erkenntnisse wurden in den PMO-Leitfaden (Version 2024) eingearbeitet und dem gesamten Projektmanagement-Team kommuniziert."),
                  ])


# ═══════════════════════════════════════════════════════════════════════════
# EXTRA PADDING: Generate additional documents to ensure ~5100 total
# ═══════════════════════════════════════════════════════════════════════════

def gen_extra_padding():
    print(f"\n[extra] Generating padding docs to reach target …")

    # Additional monthly reports per subsidiary (extra years)
    for sub in SUBS[:6]:
        folder = SUB_FOLDERS[sub["short"]]
        for m in range(1, 13):
            r_m = int(sub["revenue"] / 12 * random.uniform(0.82, 1.18))
            make_xlsx(folder, f"{sub['short']}_MonatsPL_{2022}_{m:02d}.xlsx",
                      f"Monatliche GuV {MONTHS_DE[m]} 2022 – {sub['name']}",
                      [("Monat",
                        ["KPI", "Ist (TEUR)", "Plan (TEUR)"],
                        [["Umsatz", f"{r_m//1000:,}", f"{sub['revenue']//12//1000:,}"],
                         ["EBIT", f"{int(r_m*0.07)//1000:,}", f"{int(sub['revenue']/12*0.07)//1000:,}"]],
                        [30, 15, 15])])

    # Additional IATF product docs (per product × per sub)
    for sub in SUBS[:4]:
        folder = SUB_FOLDERS.get(sub["short"], "13_IATF_Qualitaet")
        for prod in PRODUCTS:
            make_docx(folder, f"{sub['short']}_ControlPlan_{prod['id']}_2023.docx",
                      sub["name"], sub["city"], sub["hrb"],
                      f"Lenkungsplan (Control Plan) – {prod['name']} – {sub['name']}",
                      [
                          ("Lenkungsplan", [
                              ["Prozessschritt", "Merkmal", "Spezifikation", "Prüfmethode", "Frequenz", "Reaktionsplan"],
                              ["SMT Bestückung", "Bauteilposition", "±0,1 mm", "AOI", "100 %", "Linie stoppen"],
                              ["Löten (Reflow)", "Löttemperatur", f"{random.randint(245,260)} °C", "Profiler", "Stündlich", "Parameteranpassung"],
                              ["ICT-Test", "Kurzschluss/Offenstelle", "0 Fehler", "ICT-Prüfmaschine", "100 %", "Ausschleusen"],
                              ["EOL-Test", "Funktionstest", "Spec. REA-EOL-001", "Prüfstand", "100 %", "Nacharbeit / Ausschuss"],
                          ]),
                      ])

    # Extra contract amendments
    for cust in CUSTOMERS:
        for y in [2021, 2023]:
            make_docx("11_Vertrieb_OEM",
                      f"REA_{cust['short']}_Vertragsaenderung_{y}_Nachtrag.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"Nachtrag {y} – Rahmenliefervertrag – {cust['name']}",
                      [
                          ("Gegenstand des Nachtrags", f"Dieser Nachtrag Nr. {y-2021+1} zum Rahmenliefervertrag vom 1. März 2022 zwischen {AG['name']} und {cust['name']} regelt folgende Änderungen: (1) Preisanpassung gemäß Jahrespreisprotokoll {y}, (2) Aufnahme neuer Produkte in den Vertragsumfang, (3) Aktualisierung der Qualitätssicherungsvereinbarung."),
                          (None, sig_block_de(AG, cust["name"])),
                      ])

    # Additional IR documents
    for y in [2021, 2022, 2023]:
        make_pdf("10_Kapitalmarkt_IR", f"REA_Capital_Markets_Day_{y}.pdf",
                 AG["name"], AG["full_address"], f"ISIN {AG['isin']}",
                 f"Capital Markets Day {y} – Brennhagen Elektronik AG",
                 [
                     ("CMD-Agenda", f"Capital Markets Day {y}\nDatum: {ds(y, 11, random.randint(20,30))}\nOrt: Frankfurt Stock Exchange\n\nAgenda:\n09:00 Begrüßung und Konzernüberblick – {AG['ceo']}\n10:00 Finanzen und Kapitalallokation – {AG['cfo']}\n11:00 Technologie und Innovation – {AG['cto']}\n12:00 Strategie und Ausblick – {AG['ceo']}\n13:00 Mittagessen / One-on-One Meetings"),
                     ("Key Messages", f"1. Umsatzwachstum {eur(AG[f'revenue_{y}'])} (+{random.randint(7,15)} % YoY)\n2. EBITDA-Marge Guidance: {AG[f'ebitda_{y}']/AG[f'revenue_{y}']*100:.1f}–{AG[f'ebitda_{y}']/AG[f'revenue_{y}']*100+0.5:.1f} %\n3. EV und ADAS als strategische Wachstumsmotoren\n4. Aktienrückkauf / Dividende: Progressive Policy"),
                 ])

    # Additional compliance training records
    for y in [2022, 2023]:
        for sub in SUBS[:5]:
            make_xlsx("15_Compliance_Recht", f"{sub['short']}_Compliance_Training_{y}.xlsx",
                      f"Compliance-Schulungsnachweis {y} – {sub['name']}",
                      [
                          ("Training",
                           ["Schulungsmodul", "Pflichtteilnahme %", "Abschluss bis", "Status"],
                           [
                               ["Code of Conduct", f"{random.randint(92,99)} %", f"31.03.{y}", "Abgeschlossen"],
                               ["Datenschutz DSGVO", f"{random.randint(90,98)} %", f"31.03.{y}", "Abgeschlossen"],
                               ["Anti-Korruption", f"{random.randint(88,97)} %", f"30.06.{y}", "Abgeschlossen"],
                               ["Exportkontrolle", f"{random.randint(85,95)} %", f"30.09.{y}", "Abgeschlossen"],
                               ["IT-Sicherheit", f"{random.randint(87,96)} %", f"30.09.{y}", "Abgeschlossen"],
                               ["LkSG", f"{random.randint(80,94)} %", f"31.12.{y}", "Abgeschlossen"],
                           ],
                           [40, 20, 18, 15]),
                      ])

    # Additional quality system audit per product per customer combination
    for cust in CUSTOMERS[:3]:
        for prod in PRODUCTS[:3]:
            make_xlsx("13_IATF_Qualitaet", f"REA_Qualitaets_KPI_{cust['short']}_{prod['id']}_2023.xlsx",
                      f"Qualitäts-KPIs {cust['name']} / {prod['name']} – 2023",
                      [
                          ("KPIs",
                           ["Monat", "Auslieferung Stk.", "ppm angeliefert", "Reklamationen", "8D offen", "FPY intern %"],
                           [
                               [MONTHS_DE[m], f"{random.randint(500,8000):,}", f"{random.randint(0,18)}", f"{random.randint(0,3)}", f"{random.randint(0,2)}", f"{random.uniform(97,99.9):.2f} %"]
                               for m in range(1, 13)
                           ],
                           [15, 18, 15, 15, 12, 15]),
                      ])

    # Additional board strategy memos
    strategy_topics = [
        "REA_Memo_EV_Wachstumsstrategie_2023",
        "REA_Memo_Software_Defined_Vehicle_2024",
        "REA_Memo_Nearshoring_Marokko_2023",
        "REA_Memo_Acquisition_Target_Analyse_2024",
        "REA_Memo_Carbon_Neutrality_2030",
        "REA_Memo_TISAX_Level3_Roadmap",
        "REA_Memo_Lieferantendiversifizierung_Halbleiter",
        "REA_Memo_Digital_Manufacturing_Heilbronn",
        "REA_Memo_Cybersecurity_UNECE_155_156",
        "REA_Memo_ISO21434_Zertifizierung_Roadmap",
    ]
    for memo in strategy_topics:
        title = memo.replace("REA_Memo_", "").replace("_", " ")
        make_docx("20_Strategie_Vorstand", memo + ".docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Vorstandsmemo – {title}",
                  [
                      ("Zusammenfassung", f"Dieses vertrauliche Vorstandsmemo fasst die strategischen Überlegungen zu {title} zusammen. Erstellt von: Strategieabteilung / {AG['ceo']} für die Vorstandssitzung."),
                      ("Analyse und Empfehlung", f"Auf Basis einer eingehenden Analyse empfiehlt der Vorstand folgende Maßnahmen: [Detailanalyse und konkrete Handlungsempfehlungen]. Budget-Implikation: EUR {random.randint(1,20)} Mio. über {random.randint(2,5)} Jahre."),
                  ], confidential=True)

    # Additional PPAP submissions (remaining customer/product combos)
    for cust in CUSTOMERS[3:]:
        for prod in PRODUCTS:
            make_docx("13_IATF_Qualitaet", f"PPAP_{cust['short']}_{prod['id']}_2023_L3.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"PPAP Level 3 – {prod['name']} – {cust['name']}",
                      [
                          ("PSW", f"Part Submission Warrant (PSW)\nKunde: {cust['name']}\nProdukt: {prod['name']}\nPPAP Level: 3\nFreigabe: {random.choice(['Vollständige Freigabe','Vorläufige Freigabe'])}\nDatum: {ds(2023, random.randint(3,10), random.randint(1,25))}"),
                      ])

    # Additional intercompany service agreements (remaining pairs)
    extra_ic = [
        (SUBS[1], SUBS[0], "Software-Entwicklungsdienstleistungen"),
        (SUBS[3], SUBS[0], "Engineering-Dienstleistungen"),
        (SUBS[0], SUBS[2], "Produktionsunterstützung / Know-how Transfer"),
        (SUBS[0], SUBS[4], "Produktionsunterstützung EV"),
    ]
    for provider, recipient, service in extra_ic:
        make_docx("00_Konzernstruktur_Holding",
                  f"IC_Dienstleistung_{provider['short']}_{recipient['short']}_2023.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Intercompany-Dienstleistungsvertrag – {provider['name']} → {recipient['name']}",
                  [
                      ("1. Vertragsgegenstand", f"Die {provider['name']} ({provider['city']}) erbringt für die {recipient['name']} ({recipient['city']}) die Dienstleistungen: {service}."),
                      ("2. Vergütung", f"Vergütung: EUR {random.randint(500, 3000) * 1000:,} p.a. zzgl. USt. (Cost-Plus 5 %)."),
                      ("3. Laufzeit", "Laufzeit: 1. Januar 2023 bis 31. Dezember 2025, automatische Verlängerung."),
                  ])

    # Extra HR documents (individual performance reviews)
    for i in range(20):
        name = ALL_NAMES[i % len(ALL_NAMES)]
        rolle, salary, _ = ROLES[i % len(ROLES)]
        make_docx("09_Personal_HR", f"MBO_Zielvereinbarung_{2023}_{i+1:03d}_{sfn(name)[:18]}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Zielvereinbarung 2023 – {name} – {rolle}",
                  [
                      ("Ziele 2023", [
                          ["Ziel-ID", "Ziel", "Gewichtung", "Zielwert", "Ergebnis"],
                          [f"Z{i+1}-01", f"Konzern-EBITDA", "30 %", eur(AG["ebitda_2023"]), "Erreicht"],
                          [f"Z{i+1}-02", f"Projektabschluss {random.choice(PROJECTS)[1][:25]}", "30 %", "Gemäß Plan", "Teilweise erreicht"],
                          [f"Z{i+1}-03", "Mitarbeiterentwicklung (Teamzufriedenheit)", "20 %", "> 4,0 / 5,0", "4,2 / 5,0"],
                          [f"Z{i+1}-04", "Persönliches Entwicklungsziel", "20 %", "[Individuell]", "Abgeschlossen"],
                      ]),
                  ])

    # Extra financial documents
    for m in range(1, 13):
        make_xlsx("02_Konsolidierte_Finanzen", f"REA_Treasury_Report_{2023}_{m:02d}.xlsx",
                  f"Treasury-Monatsbericht {MONTHS_DE[m]} 2023",
                  [
                      ("Liquidität",
                       ["Position", "Betrag (TEUR)"],
                       [
                           ["Kassenbestand Konzern", f"{random.randint(50000, 120000):,}"],
                           ["Kreditlinie (RCF) verfügbar", f"{random.randint(80000, 150000):,}"],
                           ["Netto-Finanzschulden", f"{random.randint(80000, 200000):,}"],
                           ["Net Debt/EBITDA", f"{random.uniform(0.8, 2.0):.2f}x"],
                       ],
                       [40, 20]),
                  ])

    # Additional supplier development
    for sup in SUPPLIERS:
        make_docx("12_Einkauf_Lieferanten", f"REA_{sup['short']}_Entwicklungsprogramm_2023.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Lieferantenentwicklungsprogramm 2023 – {sup['name']}",
                  [
                      ("Maßnahmenplan", [
                          ["Maßnahme", "Ziel", "Termin", "Status"],
                          ["Kapazitätserhöhung Wafer-Fertigung", f"+{random.randint(10,30)} %", f"Q{random.randint(2,4)} 2024", "In Umsetzung"],
                          ["Dual-Source für kritische MCUs", "2 qualifizierte Quellen", f"Q{random.randint(1,3)} 2025", "Planung"],
                          ["Nachhaltigkeitszertifikat (ISO 14001)", "Zertifizierung Werk", f"Q{random.randint(2,4)} 2024", "Audit geplant"],
                      ]),
                  ])


# ═══════════════════════════════════════════════════════════════════════════
# HIGH-VOLUME BULK GENERATOR – fills remaining gap to 5100
# ═══════════════════════════════════════════════════════════════════════════

def gen_bulk_to_target():
    """Generate bulk docs across all folders until we hit 5100 total."""
    print(f"\n[bulk] Bulk generation to reach 5100 target (currently {TOTAL[0]}) ...")

    # ── 1. Per-subsidiary per-year monthly reports (2020, 2021 also) ──
    for sub in SUBS[:6]:
        folder = SUB_FOLDERS[sub["short"]]
        for y in [2020, 2021]:
            for m in range(1, 13):
                r_m = int(sub["revenue"] / 12 * random.uniform(0.78, 1.18))
                make_xlsx(folder, f"{sub['short']}_MonatsPL_{y}_{m:02d}.xlsx",
                          f"Monatliche GuV {MONTHS_DE[m]} {y} – {sub['name']}",
                          [("Monat", ["KPI","Ist (TEUR)","Plan (TEUR)"],
                            [["Umsatz", f"{r_m//1000:,}", f"{sub['revenue']//12//1000:,}"],
                             ["EBIT", f"{int(r_m*0.068)//1000:,}", "-"]],
                            [30, 15, 15])])

    # ── 2. Intercompany invoices 2020, 2021, 2022 for all subs ──
    for sub in SUBS[:6]:
        folder = SUB_FOLDERS[sub["short"]]
        for y in [2020, 2021, 2022]:
            for m in range(1, 13):
                inv = int(sub["revenue"] / 12 * random.uniform(0.88, 1.12))
                make_docx(folder, f"{sub['short']}_IC_Rechnung_{y}_{m:02d}.docx",
                          sub["name"], sub["city"], sub["hrb"],
                          f"Intercompany-Rechnung {MONTHS_DE[m]} {y}",
                          [("Rechnung", [["Position","Betrag"],
                                          ["IC-Leistungen", eur(inv)],
                                          ["Netto", eur(int(inv*0.95))]])])

    # ── 3. Group financial reports extra years (2020, 2021) ──
    for y in [2020, 2021]:
        for m in range(1, 13):
            r_m = int(AG[f"revenue_{y}"] / 12 * random.uniform(0.85, 1.15))
            make_xlsx("02_Konsolidierte_Finanzen",
                      f"REA_Monatsbericht_{y}_{m:02d}.xlsx",
                      f"Konzern-Monatsbericht {MONTHS_DE[m]} {y}",
                      [("Ergebnisrechnung",
                        ["KPI","Ist (TEUR)","Plan (TEUR)"],
                        [["Umsatz", f"{r_m//1000:,}", f"{AG[f'revenue_{y}']//12//1000:,}"],
                         ["EBITDA", f"{int(r_m*0.102)//1000:,}", "-"]],
                        [30, 15, 15])])

    # ── 4. Treasury monthly 2020, 2021, 2022 ──
    for y in [2020, 2021, 2022]:
        for m in range(1, 13):
            make_xlsx("02_Konsolidierte_Finanzen",
                      f"REA_Treasury_Report_{y}_{m:02d}.xlsx",
                      f"Treasury-Monatsbericht {MONTHS_DE[m]} {y}",
                      [("Liquiditaet", ["Position","Betrag (TEUR)"],
                        [["Kassenbestand", f"{random.randint(40000,110000):,}"],
                         ["RCF verfügbar", f"{random.randint(70000,150000):,}"],
                         ["Net Debt/EBITDA", f"{random.uniform(0.8,2.5):.2f}x"]],
                        [40, 20])])

    # ── 5. HR documents – Betriebsrat meetings per DE entity 2020, 2021 ──
    de_subs_short = ["REG", "RSG", "RHO"]
    for sshort in de_subs_short:
        sub = next(s for s in SUBS if s["short"] == sshort)
        for y in [2020, 2021]:
            for m in range(1, 13):
                make_docx("21_Betriebsraete",
                          f"{sshort}_BR_Protokoll_{y}_{m:02d}.docx",
                          sub["name"], sub["city"], sub["hrb"],
                          f"BR-Sitzungsprotokoll {MONTHS_DE[m]} {y} – {sub['name']}",
                          [("Protokoll", f"Betriebsratssitzung {MONTHS_DE[m]} {y}. "
                            "Keine außerordentlichen Beschlüsse.")])

    # ── 6. Project status reports – all remaining months ──
    for prj_id, prj_name, start, end, customer in PROJECTS:
        psfn = sfn(prj_name[:25])
        sy = int(start.split("-")[0]); sm = int(start.split("-")[1])
        ey = int(end.split("-")[0]);   em_end = int(end.split("-")[1])
        # Generate reports for months 7-24 (after first 6 already done)
        for extra_i in range(6, 24):
            m = (sm + extra_i - 1) % 12 + 1
            y = sy + (sm + extra_i - 1) // 12
            if y > ey or (y == ey and m > em_end):
                break
            make_docx("23_Projekte_Programme",
                      f"{prj_id}_Status_{y}_{m:02d}_{psfn[:15]}.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"Projektstatusbericht {MONTHS_DE[m]} {y} – {prj_name}",
                      [("Status", [["Ampel","KPI","Trend"],
                                   [random.choice(["Gruen","Gruen","Gelb"]), "Zeitplan", "Stabil"],
                                   [random.choice(["Gruen","Gruen","Gelb"]), "Budget", "Stabil"],
                                   ["Gruen", "Qualitaet", "Positiv"]])])

    # ── 7. OEM per-product extended docs ──
    for cust in CUSTOMERS:
        for prod in PRODUCTS:
            # Quarterly progress reviews
            for y in [2022, 2023]:
                for q in [1, 2, 3, 4]:
                    make_docx("11_Vertrieb_OEM",
                              f"REA_{cust['short']}_{prod['id']}_QBR_{y}_Q{q}.docx",
                              AG["name"], AG["full_address"], AG["hrb"],
                              f"Quarterly Business Review – {prod['name']} – {cust['name']} Q{q}/{y}",
                              [("QBR Summary",
                                [["Thema","Status","Kommentar"],
                                 ["Lieferperformance", f"{random.randint(96,100)} %", "On track"],
                                 ["Qualitaet ppm", f"{random.randint(0,12)}", "Within target"],
                                 ["Preisverhandlung", random.choice(["Abgeschlossen","Laufend"]), ""],
                                 ["Neue Projekte", f"{random.randint(0,2)} in Verhandlung", ""],])])
            # Technical change notices
            for i in range(1, 4):
                make_docx("11_Vertrieb_OEM",
                          f"REA_{cust['short']}_{prod['id']}_ECR_{2022+i%2:04d}_{i:03d}.docx",
                          AG["name"], AG["full_address"], AG["hrb"],
                          f"Engineering Change Request – {prod['name']} – {cust['name']} #{i:03d}",
                          [("ECR Details",
                            [["Feld","Inhalt"],
                             ["ECR-Nr.", f"ECR-{cust['short']}-{prod['id']}-{2022+i%2}-{i:03d}"],
                             ["Aenderungsgrund", random.choice(["Kostenoptimierung","Rohstoffwechsel","Funktionsverbesserung","Regulierung"])],
                             ["Betroffene Fahrzeugmodelle", f"{cust['name']} Plattform {chr(65+i)}"],
                             ["Auswirkung Qualitaet", random.choice(["Keine","Gering","PPAP erforderlich"])],
                             ["Status", random.choice(["Genehmigt","In Pruefung","Abgelehnt"])]])])

    # ── 8. Supplier extended documents ──
    for sup in SUPPLIERS:
        # Quarterly IC invoices to REA
        for y in [2021, 2022, 2023]:
            for q in [1, 2, 3, 4]:
                inv = random.randint(4_000_000, 20_000_000)
                make_pdf("12_Einkauf_Lieferanten",
                         f"INV_{sup['short']}_REA_{y}_Q{q}.pdf",
                         sup["name"], f"{sup['country']}-Werk", f"Supplier ID REA-L-{sup['short']}",
                         f"Lieferantenrechnung Q{q} {y} – {sup['name']}",
                         [("Rechnung",
                           [["Pos","Beschreibung","Menge","EP (EUR)","Gesamt"],
                            ["1", f"Halbleiter-Komponenten gemaess PO Q{q}/{y}", f"{random.randint(5000,50000):,}", f"{random.uniform(0.5,20):.2f}", eur(inv)],
                            ["", "Gesamtbetrag netto", "", "", eur(inv)],
                            ["", "USt (0% innergemeinschaftlich)", "", "", "EUR 0,00"],
                            ["", "Rechnungsbetrag", "", "", eur(inv)]])])

        # Technical datasheets per product
        for prod in PRODUCTS[:3]:
            make_pdf("12_Einkauf_Lieferanten",
                     f"{sup['short']}_{prod['id']}_Datasheet_2023.pdf",
                     sup["name"], f"{sup['country']}-Werk", f"Supplier ID REA-L-{sup['short']}",
                     f"Produktdatenblatt – {sup['name']} fuer {prod['name']}",
                     [("Technische Daten",
                       [["Parameter","Min","Typ","Max","Einheit"],
                        ["Versorgungsspannung", "4.5", "5.0", "5.5", "V"],
                        ["Betriebstemperatur", "-40", "25", "125", "degC"],
                        ["Stromaufnahme", "-", f"{random.uniform(50,500):.0f}", "-", "mA"],
                        ["ESD-Festigkeit", "-", f"{random.randint(1000,4000)}", "-", "V (HBM)"],
                        ["MTBF", f"{random.randint(500000,2000000):,}", "-", "-", "h"]])])

        # Capacity planning docs
        for y in [2023, 2024]:
            make_xlsx("12_Einkauf_Lieferanten",
                      f"{sup['short']}_Kapazitaetsplanung_{y}.xlsx",
                      f"Kapazitaetsplanung {y} – {sup['name']}",
                      [("Kapazitaet",
                        ["Produktgruppe","Kapazitaet/Monat","Zugesagt REA","Verfuegbar"],
                        [[f"MCU-{chr(65+i)}", f"{random.randint(50000,500000):,}",
                          f"{random.randint(10000,100000):,}", f"{random.randint(10000,100000):,}"]
                         for i in range(6)],
                        [25, 20, 20, 20])])

    # ── 9. IATF extended: process instructions per product per site ──
    proc_types = ["Loetprozess_Reflow","SMT_Bestueckung","Endpruefung_EOL",
                  "ICT_Test","Funktionstest","Verpackung_Versand","Eingangswarenprufung",
                  "Rueckverfolgbarkeit","Fehlerbehandlung_Ausschuss","Wartung_Anlagen"]
    for proc in proc_types:
        for sub in SUBS[:4]:
            make_docx("13_IATF_Qualitaet",
                      f"{sub['short']}_AA_{proc}_2023.docx",
                      sub["name"], sub["city"], sub["hrb"],
                      f"Arbeitsanweisung {proc.replace('_',' ')} – {sub['name']}",
                      [("1. Zweck", f"Diese Arbeitsanweisung beschreibt den Prozess {proc.replace('_',' ')} "
                        f"am Standort {sub['city']} der {sub['name']}."),
                       ("2. Prozessbeschreibung",
                        [["Schritt","Taetigkeit","Verantwortlich","Dokument"],
                         ["1", "Auftrag empfangen und pruefen", "Schichtfuehrer", "BDE-System"],
                         ["2", f"{proc.replace('_',' ')} durchfuehren", "Maschinenbediener", "Prüfprotokoll"],
                         ["3", "Ergebnis dokumentieren", "QA-Mitarbeiter", "QM-System"],
                         ["4", "Freigabe oder Sperrung", "QA-Manager", "Freigabeschein"]]),
                       ("3. Dokumentenlenkung",
                        f"Version 3.1, freigegeben {ds(2023, random.randint(1,6), random.randint(1,25))}. "
                        f"Naechste Ueberprüfung: {ds(2024, random.randint(1,6), random.randint(1,25))}.")])

    # ── 10. Compliance extended: per-entity risk assessments ──
    for sub in SUBS[:6]:
        folder = SUB_FOLDERS[sub["short"]]
        for topic in ["Korruptionsrisiko","Exportkontrolle","Datenschutz","Kartellrecht"]:
            make_docx("15_Compliance_Recht",
                      f"{sub['short']}_RisikoAssessment_{topic}_{2023}.docx",
                      sub["name"], sub["city"], sub["hrb"],
                      f"Compliance-Risikoassessment {topic} – {sub['name']}",
                      [("Bewertung",
                        [["Risikofaktor","Wahrscheinlichkeit","Schadenshoehe","Risikostufe","Massnahme"],
                         [topic, random.choice(["Gering","Mittel","Mittel"]),
                          random.choice(["Mittel","Hoch"]),
                          random.choice(["Mittel","Hoch","Gering"]),
                          "Priventivmassnahmen gemaess CMS"]])])

    # ── 11. Kapitalmarkt extended: analyst reports, roadshow docs ──
    for y in [2021, 2022, 2023]:
        # Consensus report
        make_xlsx("10_Kapitalmarkt_IR",
                  f"REA_Analyst_Konsensus_{y}.xlsx",
                  f"Analysten-Konsensus {y} – Brennhagen Elektronik AG",
                  [("Konsensus",
                    ["Analyst","Empfehlung","Kursziel (EUR)","Modell-Umsatz (Mio.)","Modell-EBITDA (Mio.)"],
                    [[f"Bank {chr(65+i)} Research", random.choice(["Kaufen","Kaufen","Halten","Verkaufen"]),
                      f"{random.uniform(45,72):.1f}", f"{AG[f'revenue_{y}']//1_000_000 + random.randint(-20,20)}",
                      f"{AG[f'ebitda_{y}']//1_000_000 + random.randint(-5,5)}"]
                     for i in range(8)],
                    [25, 15, 18, 22, 22])])
        # Roadshow materials
        for city in ["Frankfurt","London","Paris","Zurich","New York"]:
            make_pdf("10_Kapitalmarkt_IR",
                     f"REA_Roadshow_{y}_{city.replace(' ','_')}.pdf",
                     AG["name"], AG["full_address"], f"ISIN {AG['isin']}",
                     f"Roadshow-Praesentation {y} – {city}",
                     [("Company Overview",
                       f"Roadshow {city}, {y}. Key metrics: Revenue {eur(AG[f'revenue_{y}'])}, "
                       f"EBITDA {eur(AG[f'ebitda_{y}'])}, Employees {AG.get(f'employees_{y}',3487):,}."),
                      ("Investment Case",
                       [["Highlight","Detail"],
                        ["TecDAX Member", AG["stock"]],
                        ["Market Cap", eur(AG["market_cap_2023"])],
                        ["Top Customer", f"{AG['kunde1']} ({AG['kunde1_rev_share']})"],
                        ["Growth Drivers", "EV (BatteryMS-12), ADAS-Vision 4D"],
                        ["Guidance", f"Revenue EUR {AG['revenue_2024e']//1_000_000:.0f}m, EBITDA {AG['ebitda_2024e']/AG['revenue_2024e']*100:.1f}%"]])])

    # ── 12. IP extended: patent prosecution, responses, renewals ──
    for i in range(1, 19):
        for doc_type in ["Bescheid_EPA","Antwort_Anmelder","Jahresgebuehr","Erteilungsurkunde"]:
            make_docx("14_IP_Technologie",
                      f"Patent_{i:02d}_{doc_type}_{2020+i%4}.docx",
                      AG["name"] if "Antwort" in doc_type or "Jahres" in doc_type else "Europäisches Patentamt",
                      "80298 München" if "EPA" in doc_type else AG["full_address"],
                      f"DE10{2010+i}0XXXXX",
                      f"Patent {i:02d} – {doc_type.replace('_',' ')}",
                      [("Dokument", f"Patentverfahren bezüglich Anmeldung DE10{2010+i}0XXXXX. "
                        f"Dokument: {doc_type.replace('_',' ')}. "
                        f"Datum: {ds(2020+i%4, random.randint(1,12), random.randint(1,25))}.")])

    # ── 13. Personal/HR extended: per-employee docs ──
    for i in range(61, 121):
        name = ALL_NAMES[i % len(ALL_NAMES)] + f" {i}"
        rolle, salary, _ = ROLES[i % len(ROLES)]
        sub = SUBS[i % 6]
        # Employment contract
        make_docx("09_Personal_HR",
                  f"AV_{i:03d}_{sfn(name)[:18]}_{sfn(rolle)[:18]}.docx",
                  sub["name"], sub["city"], sub["hrb"],
                  f"Arbeitsvertrag – {name} – {rolle}",
                  [("Anstellungsvertrag", f"Zwischen {sub['name']} und {name} als {rolle}. "
                    f"Jahresgehalt: EUR {salary:,}. Beginn: {ds(random.randint(2019,2023), random.randint(1,12), 1)}.")])
        # Annual review
        make_docx("09_Personal_HR",
                  f"JG_{i:03d}_{sfn(name)[:18]}_{2023}.docx",
                  sub["name"], sub["city"], sub["hrb"],
                  f"Jahresgespraech 2023 – {name}",
                  [("Bewertung",
                    [["Dimension","Bewertung (1-5)","Kommentar"],
                     ["Fachkompetenz", f"{random.randint(3,5)}", ""],
                     ["Eigeninitiative", f"{random.randint(3,5)}", ""],
                     ["Teamarbeit", f"{random.randint(3,5)}", ""],
                     ["Fuehrungsverhalten", f"{random.randint(3,5)}", ""],
                     ["Gesamtbewertung", f"{random.randint(3,5)}", ""],])])

    # ── 14. Immobilien extended: lease amendments, inspections ──
    sites_short = ["HQ_Stuttgart","Werk_Heilbronn","Office_München","Werk_Katowice","Werk_Brno"]
    for site in sites_short:
        for doc_type, folder_t in [("Nachtrag_2022","18_Immobilien"), ("Inspektion_2023","18_Immobilien"),
                                    ("Nebenkostenabrechnung_2022","18_Immobilien"),
                                    ("Nebenkostenabrechnung_2023","18_Immobilien")]:
            make_docx(folder_t, f"{sfn(site)[:20]}_{doc_type}.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"{doc_type.replace('_',' ')} – {site.replace('_',' ')}",
                      [("Dokument", f"{doc_type.replace('_',' ')} fuer Objekt {site.replace('_',' ')}. "
                        f"Datum: {ds(int(doc_type[-4:]), random.randint(1,12), random.randint(1,28))}.")])

    # ── 15. Strategy/Board extended: investment memos, competitive briefs ──
    competitors = ["Continental AG","Robert Bosch GmbH","Aptiv plc","Visteon Corp","ZF Friedrichshafen AG",
                   "Magna International","Sensata Technologies","Lear Corporation"]
    for comp in competitors:
        make_docx("20_Strategie_Vorstand",
                  f"REA_Wettbewerbsanalyse_{sfn(comp)[:25]}_2023.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Wettbewerbsanalyse – {comp} – 2023",
                  [("1. Unternehmensprofil", f"{comp} ist ein direkter Wettbewerber der {AG['name']} im Bereich Automotive-Elektronik."),
                   ("2. Staerken und Schwaechen",
                    [["Dimension","Staerken","Schwaechen"],
                     ["Marktposition", "Global aufgestellt", "Hohe Fixkosten"],
                     ["Technologie", "Breite Produktpalette", "Langsamere Innovationszyklen"],
                     ["Finanzen", f"Umsatz EUR {random.randint(5,60)} Mrd.", "Hohe Verschuldung"],]),
                   ("3. Strategische Implikation", f"Für {AG['name']} ergibt sich folgende strategische Empfehlung: "
                    "Differenzierung durch Spezialisierung auf EV und ADAS, Kundennähe und Agilität.")])

    # ── 16. M&A extended: due diligence process docs ──
    for i in range(1, 16):
        make_docx("19_MA_Transaktionen",
                  f"DD_WorkingPaper_{2023}_{i:03d}.docx",
                  AG["recht"], "Frankfurt", f"DD-WP-{i:03d}",
                  f"Due Diligence Arbeitspapier #{i:03d}",
                  [("Thema", f"DD-Arbeitspapier Nr. {i} im Rahmen des strategischen Prozesses 2023/2024. "
                    f"Bearbeitet: {random.choice([AG['ceo'],AG['cfo'],AG['cto']])}. "
                    f"Status: {random.choice(['Abgeschlossen','In Bearbeitung','Review ausstehend'])}."),
                   ("Findings", f"[Analyse und Findings zu {random.choice(['Finanzen','Recht','Steuern','IT','HR','Versicherungen','Vertraege','IP'])}]")])

    # ── 17. Pensions extended: quarterly investment reports ──
    for y in [2021, 2022, 2023]:
        for q in [1, 2, 3, 4]:
            make_xlsx("22_Pensionen_bAV",
                      f"REA_CTA_Investment_Report_{y}_Q{q}.xlsx",
                      f"CTA Investment-Bericht Q{q} {y}",
                      [("Portfolio",
                        ["Asset-Klasse","Allokation %","Rendite YTD %","Marktwert (TEUR)"],
                        [["Anleihen (IG)", "40 %", f"{random.uniform(-5,8):.1f} %", f"{random.randint(25000,45000):,}"],
                         ["Aktien Global", "35 %", f"{random.uniform(-15,25):.1f} %", f"{random.randint(20000,38000):,}"],
                         ["Immobilien (REIT)", "15 %", f"{random.uniform(-8,12):.1f} %", f"{random.randint(8000,16000):,}"],
                         ["Cash/Geldmarkt", "10 %", f"{random.uniform(0.5,3.5):.1f} %", f"{random.randint(5000,10000):,}"],
                         ["Gesamt", "100 %", f"{random.uniform(-8,15):.1f} %", f"{random.randint(60000,110000):,}"]],
                        [25, 18, 18, 20])])

    # ── 18. Versicherungen extended: claim details and surveys ──
    for y in [2020, 2021, 2022, 2023]:
        make_xlsx("17_Versicherungen",
                  f"REA_Versicherungsabrechnung_{y}.xlsx",
                  f"Versicherungsjahresabrechnung {y}",
                  [("Abrechnung",
                    ["Versicherung","Prämie (EUR)","Schäden (EUR)","Schadenquote %","Rabatt/Zuschlag"],
                    [["Produkthaftpflicht", eur(random.randint(380,450)*1000), eur(random.randint(0,300)*1000), f"{random.uniform(0,70):.0f} %", "+/-0 %"],
                     ["Betriebshaftpflicht", eur(random.randint(250,320)*1000), eur(random.randint(0,50)*1000), f"{random.uniform(0,20):.0f} %", "-3 %"],
                     ["Cyber", eur(random.randint(280,360)*1000), eur(random.randint(0,100)*1000), f"{random.uniform(0,30):.0f} %", "+5 %"],
                     ["Feuer", eur(random.randint(550,700)*1000), eur(random.randint(0,500)*1000), f"{random.uniform(0,80):.0f} %", "0 %"]],
                    [30, 20, 20, 18, 15])])

    # ── 19. QM extended: measurement equipment calibrations ──
    equipment_types = ["Messschieber","Koordinatenmessmaschine","Drehmomentschluessel",
                       "Loetprofil-Messsystem","Klimakammer","EMV-Prüfanlage",
                       "Impedanz-Analysator","Spektrum-Analysator","Hochspannungspruefer",
                       "Isolationsmessgeraet"]
    for eq in equipment_types:
        for sub in SUBS[:4]:
            make_docx("13_IATF_Qualitaet",
                      f"{sub['short']}_Kalibrierung_{sfn(eq)[:20]}_{2023}.docx",
                      sub["name"], sub["city"], sub["hrb"],
                      f"Kalibrierprotokoll – {eq} – {sub['name']} – {2023}",
                      [("Kalibrierung",
                        [["Geraet",eq],["Seriennummer",f"SN-{random.randint(10000,99999)}"],
                         ["Kalibriert am", ds(2023, random.randint(1,11), random.randint(1,25))],
                         ["Naechste Kalibrierung", ds(2024, random.randint(1,11), random.randint(1,25))],
                         ["Ergebnis", "Bestanden – Alle Parameter innerhalb Toleranz"],
                         ["Kalibriert durch", random.choice(["TÜV Rheinland","Mettler-Toledo","Fluke Calibration"])]])])

    # ── 20. IT extended: security incident reports, change requests ──
    for i in range(1, 26):
        doc_type = random.choice(["Sicherheitsvorfall","Change-Request","Systemstörung","Patch-Protokoll"])
        make_docx("16_IT_Systeme",
                  f"IT_{doc_type.replace('-','_')}_{2023}_{i:04d}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"IT {doc_type} #{2023}-{i:04d}",
                  [("Dokumentation",
                    [["Feld","Inhalt"],
                     ["Typ", doc_type],
                     ["Datum", ds(2023, random.randint(1,12), random.randint(1,25))],
                     ["Betroffenes System", random.choice([AG["erp"], AG["plm"], AG["mes"], "Netzwerk", "Email"])],
                     ["Schweregrad", random.choice(["Gering","Mittel","Hoch"])],
                     ["Status", random.choice(["Behoben","Geschlossen","In Bearbeitung"])],
                     ["Bearbeiter", random.choice(ALL_NAMES)]])])

    # ── 21. Einkauf extended: PO archive extra months + supplier RFQs ──
    for m in [2, 3, 5, 6, 8, 9, 11, 12]:
        po_rows = [[f"PO-2023-{m:02d}-{i:04d}", random.choice(SUPPLIERS)["name"],
                    random.choice(PRODUCTS)["name"],
                    f"{random.randint(500,10000):,}", f"{random.uniform(0.5,15.0):.2f}",
                    eur(random.randint(50_000, 2_000_000))]
                   for i in range(1, random.randint(20, 40))]
        make_xlsx("12_Einkauf_Lieferanten",
                  f"REA_Bestellarchiv_2023_{m:02d}_ext.xlsx",
                  f"Bestellarchiv Ergänzung {MONTHS_DE[m]} 2023",
                  [("Bestellungen",
                    ["Bestellnummer","Lieferant","Material","Menge","Einheitspreis (EUR)","Gesamtwert"],
                    po_rows,
                    [22, 35, 30, 12, 18, 18])])

    # RFQ documents
    rfq_topics = ["Halbleiter_MCU_2024","Halbleiter_Power_2024","Leiterplatten_HDI_2024",
                  "Passive_Komponenten_2023","Kabel_Steckverbinder_2024","Gehaeuse_Kunststoff_2024"]
    for rfq in rfq_topics:
        make_docx("12_Einkauf_Lieferanten",
                  f"REA_RFQ_{rfq}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Anfrage (RFQ) – {rfq.replace('_',' ')}",
                  [("Anfrage", f"Die {AG['name']} lädt zur Angebotsabgabe für {rfq.replace('_',' ')} ein."),
                   ("Anforderungen",
                    [["Parameter","Anforderung"],
                     ["Menge", f"{random.randint(10000,500000):,} Stk./Jahr"],
                     ["Qualitaet", f"{AG['iatf']}, AEC-Q zertifiziert"],
                     ["Lieferzeit", f"{random.randint(6,20)} Wochen"],
                     ["Angebotsabgabe bis", ds(2023, random.randint(9,12), random.randint(15,28))]])])

    # ── 22. Konzernstruktur extended: governance and reporting ──
    for y in [2021, 2022, 2023]:
        make_docx("00_Konzernstruktur_Holding",
                  f"REA_Konzernlagebericht_{y}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Zusammengefasster Konzernlagebericht {y} – Brennhagen Elektronik AG",
                  [("1. Wirtschaftsbericht",
                    f"Im Geschäftsjahr {y} erzielte der Brennhagen-Konzern einen Umsatz von "
                    f"{eur(AG[f'revenue_{y}'])} (Vorjahr: {eur(AG[f'revenue_{y-1}'])}) "
                    f"und ein EBITDA von {eur(AG[f'ebitda_{y}'])}."),
                   ("2. Nachtragsbericht",
                    f"Nach dem Bilanzstichtag 31. Dezember {y} sind keine wesentlichen Ereignisse eingetreten, "
                    f"die die Vermögens-, Finanz- oder Ertragslage des Konzerns wesentlich beeinflusst hätten."),
                   ("3. Risikobericht",
                    [["Risikobereich","Bewertung","Massnahme"],
                     ["Markt/Konjunktur","Mittel","Diversifizierung OEM-Basis"],
                     ["Halbleiterverfügbarkeit","Hoch","LTA-Vertraege, Dual-Source"],
                     ["Wechselkursrisiko","Gering","Natural Hedging, FX-Derivate"],
                     ["Cybersicherheit","Mittel","TISAX L3, ISO 27001"]])])

    for y in [2021, 2022, 2023]:
        make_docx("00_Konzernstruktur_Holding",
                  f"REA_Nichtfinanzieller_Bericht_{y}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Nichtfinanzieller Bericht (CSR) {y} – Brennhagen Elektronik AG",
                  [("ESG-Überblick",
                    [["ESG-Dimension","KPI","Ist {y}","Ziel {y+1}"],
                     ["Umwelt (E)","CO2-Emissionen Scope 1+2 (t)",f"{random.randint(18000,25000):,}",f"{random.randint(15000,20000):,}"],
                     ["Soziales (S)","Arbeitsunfaelle (LTIFR)",f"{random.uniform(2.5,5.5):.1f}",f"< {random.uniform(2.0,4.0):.1f}"],
                     ["Governance (G)","Compliance-Schulungsquote",f"{random.randint(88,97)} %","100 %"]])])

    # ── 23. Legal: court documents, regulatory correspondence ──
    for i in range(1, 20):
        doc_type = random.choice(["Schriftsatz","Urteil_Auszug","Vergleich","Behoerdenschreiben","Mahnbescheid","Widerspruch"])
        make_docx("15_Compliance_Recht",
                  f"Rechtsakte_{2022 + i%2}_{i:03d}_{doc_type}.docx",
                  AG["name"] if "Schriftsatz" in doc_type or "Wider" in doc_type else "[Gericht/Behoerde]",
                  AG["full_address"] if "Schriftsatz" in doc_type else "[Anschrift]",
                  f"Az.: REA/{2022+i%2}/{i:04d}",
                  f"{doc_type.replace('_',' ')} – {random.choice(['Produkthaftung','Kartellrecht','Arbeitsrecht','Gewerblicher Rechtsschutz','Steuerrecht'])}",
                  [("Inhalt", f"Rechtliches Dokument im Verfahren REA/{2022+i%2}/{i:04d}. "
                    f"Betrifft: {random.choice(['Vertragsstreitigkeit','Schadensersatzklage','Behördenanfrage','Bussgeldbescheid'])}. "
                    f"Datum: {ds(2022+i%2, random.randint(1,12), random.randint(1,25))}.")])

    print(f"  Bulk round 1 complete. Total now: {TOTAL[0]}")


def gen_bulk_round2():
    """Second bulk pass – more parametric docs to cross 5100."""
    print(f"\n[bulk-2] Second bulk pass (currently {TOTAL[0]}) ...")

    # ── A. Monthly IC invoices 2020+2021 already done; add sub-to-sub invoices ──
    sub_pairs = [(SUBS[0], SUBS[1]), (SUBS[1], SUBS[3]),
                 (SUBS[3], SUBS[0]), (SUBS[4], SUBS[0]), (SUBS[2], SUBS[0])]
    for provider, recipient in sub_pairs:
        for y in [2021, 2022, 2023]:
            for m in range(1, 13):
                inv = int(provider["revenue"] / 12 * random.uniform(0.05, 0.15))
                make_docx(SUB_FOLDERS[provider["short"]],
                          f"{provider['short']}_to_{recipient['short']}_IC_{y}_{m:02d}.docx",
                          provider["name"], provider["city"], provider["hrb"],
                          f"IC-Rechnung {MONTHS_DE[m]} {y}: {provider['name']} an {recipient['name']}",
                          [("Rechnung", [["Position","Betrag"],
                                          [f"Dienstleistungen {MONTHS_DE[m]} {y}", eur(inv)],
                                          ["Nettobetrag", eur(inv)]])])

    # ── B. Annual accounts for years 2019, 2020 for top subs ──
    for sub in SUBS[:6]:
        folder = SUB_FOLDERS[sub["short"]]
        for y in [2019, 2020]:
            r = int(sub["revenue"] * random.uniform(0.75, 0.92))
            make_xlsx(folder, f"{sub['short']}_Jahresabschluss_{y}.xlsx",
                      f"Jahresabschluss {y} – {sub['name']}",
                      [("GuV",
                        ["Position", f"{y} (TEUR)", f"{y-1} (TEUR)"],
                        [["Umsatzerlöse", f"{r//1000:,}", f"{int(r*0.93)//1000:,}"],
                         ["EBIT", f"{int(r*0.065)//1000:,}", "-"],
                         ["Jahresüberschuss", f"{int(r*0.042)//1000:,}", "-"]],
                        [35, 20, 20])])

    # ── C. Customer-specific per-part-number PPAP tracking sheets ──
    part_numbers = [f"REA-{prod['id']}-{cust['short']}-{n:03d}"
                    for prod in PRODUCTS for cust in CUSTOMERS for n in range(1, 4)]
    for pn in part_numbers[:80]:  # 80 part numbers
        make_xlsx("13_IATF_Qualitaet",
                  f"PPAP_Tracking_{pn}.xlsx",
                  f"PPAP Tracking Sheet – {pn}",
                  [("Tracking",
                    ["Dokument","Erforderlich","Status","Datum","Freigabe"],
                    [["Design Record","Ja",random.choice(["OK","OK","Ausstehend"]),
                      ds(2023,random.randint(1,8),random.randint(1,25)),"QA"],
                     ["DFMEA","Ja",random.choice(["OK","OK","OK"]),"","QA"],
                     ["PFMEA","Ja",random.choice(["OK","OK","OK"]),"","QA"],
                     ["Control Plan","Ja",random.choice(["OK","OK","Ausstehend"]),"","QA"],
                     ["MSA","Ja",random.choice(["OK","OK","OK"]),"","QA"],
                     ["Dimensional Results","Ja","OK","","QA"],
                     ["Material Cert","Ja","OK","","Einkauf"],
                     ["PSW","Ja",random.choice(["Eingereicht","Genehmigt","Ausstehend"]),"","QA"]],
                    [30, 15, 15, 18, 12])])

    # ── D. Quality KPIs per plant per product (monthly 2022+2023) ──
    for sub in SUBS[:4]:
        for prod in PRODUCTS:
            for y in [2022, 2023]:
                make_xlsx("13_IATF_Qualitaet",
                          f"{sub['short']}_{prod['id']}_QKPIs_{y}.xlsx",
                          f"Qualitaets-KPIs {y} – {sub['name']} / {prod['name']}",
                          [("KPIs",
                            ["Monat","FPY intern %","ppm Kunde","8D offen","Scrap (TEUR)","OEE %"],
                            [[MONTHS_DE[m], f"{random.uniform(97.0,99.9):.2f} %",
                              f"{random.randint(0,15)}", f"{random.randint(0,3)}",
                              f"{random.randint(5,80)}", f"{random.uniform(75,95):.1f} %"]
                             for m in range(1, 13)],
                            [12, 15, 15, 12, 15, 12])])

    # ── E. Project: requirements traceability matrices ──
    for prj_id, prj_name, _, _, _ in PROJECTS:
        psfn_s = sfn(prj_name[:20])
        make_xlsx("23_Projekte_Programme",
                  f"{prj_id}_RTM_{psfn_s}.xlsx",
                  f"Requirements Traceability Matrix – {prj_name}",
                  [("RTM",
                    ["Req-ID","Anforderung","DFMEA","PFMEA","Test","Status"],
                    [[f"REQ-{prj_id}-{i:03d}",
                      f"Anforderung {i}: {random.choice(['Temperatur','EMV','Lebensdauer','Spannung','Software'])}",
                      f"DFMEA-{random.randint(1,5)}",
                      f"PFMEA-{random.randint(1,5)}",
                      f"TC-{i:03d}",
                      random.choice(["Verifiziert","In Test","Offen"])]
                     for i in range(1, random.randint(15, 30))],
                    [20, 40, 15, 15, 12, 12])])
        # Also: test reports
        for test_type in ["EMV_Test","Klimatest","Vibration","Funktionstest_EOL","Lebensdauertest"]:
            make_docx("23_Projekte_Programme",
                      f"{prj_id}_Testbericht_{test_type}_{psfn_s}.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"Prüfbericht {test_type.replace('_',' ')} – {prj_name}",
                      [("Testergebnis",
                        [["Prüfnorm","Anforderung","Ergebnis","Status"],
                         [f"{random.choice(['CISPR 25','IEC 60068-2','ISO 16750','AEC-Q100'])} Kap. {random.randint(1,10)}",
                          random.choice(["Bestehen","Klasse B","< 10 ppm","ASIL-B"]),
                          random.choice(["Bestanden","Bestanden","Bestanden","Nicht bestanden (CAPA)"]),
                          random.choice(["Freigegeben","Freigegeben","Mit Auflagen"])]])])

    # ── F. Per-OEM per-year commercial correspondence ──
    correspondence_types = [
        "Preisanpassungsschreiben","Kapazitaetsanfrage","Claim_Letter",
        "Audit_Einladung","Versorgungsbestaetigung","Rahmenvertrag_Kuendigung_Androhung",
        "Lobesbriefe_Top_Supplier","Eskalationsbrief_Lieferengpass"
    ]
    for cust in CUSTOMERS:
        for y in [2021, 2022, 2023]:
            for ct in random.sample(correspondence_types, 4):
                make_docx("11_Vertrieb_OEM",
                          f"REA_{cust['short']}_{ct}_{y}.docx",
                          AG["name"], AG["full_address"], AG["hrb"],
                          f"Korrespondenz – {ct.replace('_',' ')} – {cust['name']} {y}",
                          [("Schreiben",
                            f"Datum: {ds(y, random.randint(1,12), random.randint(1,25))}\n"
                            f"An: {cust['name']}, {cust['contact']}\n\n"
                            f"Betreff: {ct.replace('_',' ')} – Brennhagen Elektronik AG\n\n"
                            f"Sehr geehrte Damen und Herren,\n\nim Rahmen unserer Geschaeftsbeziehung "
                            f"moechten wir Sie ueber folgendes informieren / bitten: {ct.replace('_',' ')}.\n\n"
                            f"Mit freundlichen Grüßen\n{AG['coo']}\nCOO, {AG['name']}"),])

    # ── G. Internal audit reports per function per year ──
    audit_areas = ["Einkauf","Vertrieb","Produktion_REG","Produktion_RPL","F_E_RSG",
                   "IT_Infrastruktur","Finance_Controlling","HR","Legal_Compliance","Supply_Chain"]
    for area in audit_areas:
        for y in [2021, 2022, 2023]:
            make_docx("15_Compliance_Recht",
                      f"REA_Interne_Revision_{area}_{y}.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"Interne Revision – {area.replace('_',' ')} – {y}",
                      [("Revisionsübersicht",
                        [["Prüffeld","Ergebnis","Risiko","Empfehlung"],
                         [area.replace('_',' '), random.choice(["Kein Befund","Geringer Befund","Mittlerer Befund"]),
                          random.choice(["Gering","Mittel","Gering"]),
                          random.choice(["Keine Massnahme","Prozessverbesserung empfohlen","CAPA initiiert"])]]),
                       ("Zeitraum und Scope",
                        f"Prüfungszeitraum: 01.01.{y} – 31.12.{y}. "
                        f"Prüfer: Interne Revision, {AG['name']}. "
                        f"Berichtsdatum: {ds(y+1, 2, random.randint(1,28))}.")])

    # ── H. KPI dashboards per subsidiary per quarter ──
    for sub in SUBS[:6]:
        folder = SUB_FOLDERS[sub["short"]]
        for y in [2021, 2022, 2023]:
            for q in [1, 2, 3, 4]:
                make_xlsx(folder,
                          f"{sub['short']}_KPI_Dashboard_{y}_Q{q}.xlsx",
                          f"KPI-Dashboard Q{q} {y} – {sub['name']}",
                          [("KPIs",
                            ["KPI","Ist","Plan","Vorjahr","Zielerreichung"],
                            [[k, f"{random.uniform(0.8,1.2)*v:.1f}", f"{v:.1f}", f"{v*0.95:.1f}", f"{random.randint(85,110)} %"]
                             for k,v in [("Umsatz (TEUR)", sub["revenue"]/4/1000),
                                          ("EBIT-Marge %", 7.2),
                                          ("Liefertreue %", 97.5),
                                          ("ppm Gesamt", 8),
                                          ("OEE Produktion %", 82),
                                          ("Krankenquote %", 4.8)]],
                            [30, 12, 12, 12, 15])])

    # ── I. Finance: account reconciliations and closing packages ──
    for sub in SUBS[:5]:
        for y in [2022, 2023]:
            for m in [3, 6, 9, 12]:  # quarter-end closes
                make_xlsx("02_Konsolidierte_Finanzen",
                          f"FC_{sub['short']}_{y}_{m:02d}_Closing.xlsx",
                          f"Monatlicher Konzern-Close {MONTHS_DE[m]} {y} – {sub['name']}",
                          [("Abschluss",
                            ["Konto","Bezeichnung","Soll","Haben","Saldo"],
                            [[f"1{random.randint(100,999)}",
                              random.choice(["Forderungen aus LuL","Verbindlichkeiten LuL",
                                             "Umsatzerlöse","Materialaufwand","Personalaufwand"]),
                              f"{random.randint(100,9999):,}",
                              f"{random.randint(100,9999):,}",
                              f"{random.randint(0,5000):,}"]
                             for _ in range(random.randint(8, 15))],
                            [12, 35, 15, 15, 15])])

    # ── J. ESG and Sustainability reports / data ──
    for y in [2021, 2022, 2023]:
        make_xlsx("20_Strategie_Vorstand",
                  f"REA_ESG_Datenbericht_{y}.xlsx",
                  f"ESG Datenbericht {y} – Brennhagen Elektronik AG",
                  [("Umwelt",
                    ["Indikator","Einheit",f"{y}",f"{y-1}","Veränderung %"],
                    [["CO2 Scope 1", "t CO2e", f"{random.randint(8000,12000):,}", f"{random.randint(9000,13000):,}", f"{random.uniform(-10,5):.1f} %"],
                     ["CO2 Scope 2", "t CO2e", f"{random.randint(6000,10000):,}", f"{random.randint(7000,11000):,}", f"{random.uniform(-15,0):.1f} %"],
                     ["Energieverbrauch", "MWh", f"{random.randint(40000,60000):,}", f"{random.randint(42000,62000):,}", "-"],
                     ["Wasserverbrauch", "m3", f"{random.randint(15000,25000):,}", "-", "-"],
                     ["Abfallmenge", "t", f"{random.randint(300,600):,}", "-", "-"],
                     ["Recyclingquote", "%", f"{random.randint(75,92)} %", "-", "-"]],
                    [35, 12, 10, 10, 15]),
                   ("Soziales",
                    ["Indikator","Einheit",f"{y}",f"{y-1}","Veränderung"],
                    [["Mitarbeiter gesamt","VZÄ",f"{AG.get(f'employees_{y}',3487):,}",f"{AG.get(f'employees_{y-1}',3287):,}","+"],
                     ["Frauenanteil Gesamt","%",f"{random.randint(28,38)} %","-","-"],
                     ["Frauenanteil Fuehrung","%",f"{random.randint(20,32)} %","-","-"],
                     ["Unfallrate LTIFR","per Mio h",f"{random.uniform(2.5,5.5):.1f}","-","-"],
                     ["Schulungsstunden/MA","h/MA",f"{random.randint(18,35):.0f}","-","-"],
                     ["Mitarbeiterzufriedenheit","Score (1-5)",f"{random.uniform(3.5,4.5):.1f}","-","-"]],
                    [35, 12, 10, 10, 15])])

    # ── K. Subsidiary specific tax documents ──
    for sub in SUBS[:6]:
        folder = SUB_FOLDERS[sub["short"]]
        for y in [2020, 2021, 2022]:
            make_docx(folder,
                      f"{sub['short']}_Steuerbescheid_{y}.docx",
                      sub["name"], sub["city"], sub["hrb"],
                      f"Steuerbescheid / Tax Assessment {y} – {sub['name']}",
                      [("Bescheid",
                        f"Fuer das Steuerjahr {y} wurde ein Steuerbescheid erlassen. "
                        f"Steuerliche Belastung: {eur(int(sub['revenue'] * random.uniform(0.015, 0.025)))}. "
                        f"Bearbeitung durch {AG['steuer']} abgeschlossen.")])

    # ── L. Supplier: contracts for indirect materials ──
    indirect_suppliers = [
        ("Würth Industrie Service GmbH","WUE","Verbindungselemente / C-Teile"),
        ("Bürklin Elektronik GmbH","BUE","Passive Elektronikbauteile"),
        ("RS Components GmbH","RSC","MRO Materialien"),
        ("Grainger Deutschland GmbH","GRA","Werkzeuge und Betriebsmittel"),
        ("Lapp Gruppe","LAP","Kabel und Leitungen"),
        ("TE Connectivity Germany GmbH","TEC","Steckverbinder"),
    ]
    for sup_name, sup_short, sup_focus in indirect_suppliers:
        make_docx("12_Einkauf_Lieferanten",
                  f"REA_{sup_short}_Rahmenvertrag_2022.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Rahmenliefervertrag – {sup_name} – {sup_focus}",
                  [("Vertragsgegenstand",
                    f"{sup_name} (Lieferant) liefert an {AG['name']} (Käufer) "
                    f"{sup_focus} gemäß diesem Rahmenvertrag. "
                    f"Laufzeit: 2022–2024. Jahresvolumen: ca. {eur(random.randint(200,800)*1000)}."),
                   ("Konditionen",
                    [["Parameter","Vereinbarung"],
                     ["Zahlungsziel","30 Tage netto / 14 Tage -2 % Skonto"],
                     ["Mindestbestellwert",eur(random.randint(500,5000))],
                     ["Lieferzeit",f"{random.randint(2,10)} Werktage"],
                     ["Lagerhaltungsrabatt",f"-{random.uniform(1,4):.1f} % bei Jahresabnahme"]])])
        # Annual statements
        for y in [2022, 2023]:
            make_xlsx("12_Einkauf_Lieferanten",
                      f"REA_{sup_short}_Jahresabrechnung_{y}.xlsx",
                      f"Jahresabrechnung {y} – {sup_name}",
                      [("Abrechnung",
                        ["Monat","Bestellwert (EUR)","Liefertreue %","Reklamationen"],
                        [[MONTHS_DE[m], eur(random.randint(20, 80)*1000),
                          f"{random.uniform(94,100):.1f} %", str(random.randint(0,2))]
                         for m in range(1,13)],
                        [12, 20, 15, 15])])

    # ── M. Betriebsraete: Einigungsstelle records, BR-AG-Korrespondenz ──
    topics_br = ["Ueberstundenregelung","Personalplanung","Sozialplan_2023",
                 "Eingruppierungsstreitigkeiten","Arbeitszeit_Widerspruch",
                 "Massenentlassung_Anzeige","Betriebsaenderung_Interessenausgleich"]
    for sshort in ["REG", "RSG", "RHO"]:
        sub = next(s for s in SUBS if s["short"] == sshort)
        for topic in topics_br:
            for y in [2022, 2023]:
                make_docx("21_Betriebsraete",
                          f"{sshort}_BR_Korrespondenz_{topic}_{y}.docx",
                          sub["name"], sub["city"], sub["hrb"],
                          f"BR-Korrespondenz – {topic.replace('_',' ')} – {y}",
                          [("Korrespondenz",
                            f"Datum: {ds(y, random.randint(1,12), random.randint(1,25))}\n"
                            f"Thema: {topic.replace('_',' ')}\n\n"
                            f"[Inhalt der Korrespondenz zwischen Betriebsrat und Arbeitgeber "
                            f"der {sub['name']} zum Thema {topic.replace('_',' ')}.]")])

    # ── N. Konzernstruktur: governance minutes per year ──
    for y in [2020, 2021, 2022, 2023]:
        for q in [1, 2, 3, 4]:
            make_docx("00_Konzernstruktur_Holding",
                      f"REA_Konzern_Governance_Report_{y}_Q{q}.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"Konzern-Governance-Bericht Q{q} {y}",
                      [("Compliance Status",
                        [["Bereich","Status","Offene Punkte"],
                         ["Rechnungslegung (IFRS)","Konform","0"],
                         ["Kapitalmarkt-Compliance (MAR)","Konform","0"],
                         ["Datenschutz (DSGVO)","Konform",str(random.randint(0,3))],
                         ["Export-Kontrolle","Konform","0"],
                         ["LkSG",random.choice(["Konform","In Umsetzung"]),"0"]]),
                       ("Risikoübersicht",
                        f"Wesentliche Risiken Q{q} {y}: Halbleiterverfügbarkeit ({random.choice(['mittel','gering'])}), "
                        f"Wechselkurs EUR/CNY ({random.choice(['gering','mittel'])}), "
                        f"Cybersecurity ({random.choice(['gering','mittel'])}).")])

    # ── O. Extended Immobilien: environmental monitoring ──
    sites_mon = ["Heilbronn_Produktionshalle_A","Heilbronn_Produktionshalle_B",
                 "Katowice_Werk_1","Gyoer_Werk_1","Stuttgart_HQ"]
    for site in sites_mon:
        for y in [2021, 2022, 2023]:
            make_xlsx("18_Immobilien",
                      f"Umweltmonitoring_{sfn(site)[:20]}_{y}.xlsx",
                      f"Umweltmonitoring {y} – {site.replace('_',' ')}",
                      [("Messwerte",
                        ["Messgröße","Q1","Q2","Q3","Q4","Grenzwert","Status"],
                        [["Lösemittelemissionen (VOC) mg/m3", f"{random.uniform(0.1,2.5):.2f}",
                          f"{random.uniform(0.1,2.5):.2f}", f"{random.uniform(0.1,2.5):.2f}",
                          f"{random.uniform(0.1,2.5):.2f}", "5.0", "OK"],
                         ["Abwasser-pH", f"{random.uniform(6.5,8.5):.1f}", f"{random.uniform(6.5,8.5):.1f}",
                          f"{random.uniform(6.5,8.5):.1f}", f"{random.uniform(6.5,8.5):.1f}", "6.0-9.0", "OK"],
                         ["Laerm-Immission dB(A)", f"{random.uniform(45,65):.0f}",
                          f"{random.uniform(45,65):.0f}", f"{random.uniform(45,65):.0f}",
                          f"{random.uniform(45,65):.0f}", "65", "OK"]],
                        [35, 10, 10, 10, 10, 15, 8])])

    print(f"  Bulk round 2 complete. Total now: {TOTAL[0]}")


def gen_bulk_round3():
    """Third bulk pass – systematic coverage of remaining gaps."""
    print(f"\n[bulk-3] Third bulk pass (currently {TOTAL[0]}) ...")

    # ── P. All subs: 3 more years of monthly P&L (2017-2019) ──
    for sub in SUBS[:6]:
        folder = SUB_FOLDERS[sub["short"]]
        for y in [2017, 2018, 2019]:
            for m in range(1, 13):
                r_m = int(sub["revenue"] * random.uniform(0.6, 0.85) / 12)
                make_xlsx(folder,
                          f"{sub['short']}_MonatsPL_{y}_{m:02d}.xlsx",
                          f"Monatliche GuV {MONTHS_DE[m]} {y} – {sub['name']}",
                          [("Monat",
                            ["KPI","Ist (TEUR)","Plan (TEUR)"],
                            [["Umsatz", f"{r_m//1000:,}", f"{sub['revenue']//12//1000:,}"],
                             ["EBIT", f"{int(r_m*0.06)//1000:,}", "-"]],
                            [30,15,15])])

    # ── Q. Group monthly 2019 ──
    for m in range(1, 13):
        r_m = int(AG["revenue_2020"] * 0.92 / 12 * random.uniform(0.85, 1.15))
        make_xlsx("02_Konsolidierte_Finanzen",
                  f"REA_Monatsbericht_2019_{m:02d}.xlsx",
                  f"Konzern-Monatsbericht {MONTHS_DE[m]} 2019",
                  [("Ergebnisrechnung",
                    ["KPI","Ist (TEUR)","Plan (TEUR)"],
                    [["Umsatz", f"{r_m//1000:,}", f"{int(AG['revenue_2020']*0.92)//12//1000:,}"],
                     ["EBITDA", f"{int(r_m*0.098)//1000:,}", "-"]],
                    [30,15,15])])

    # ── R. All subs: quarterly IC invoices 2019-2021 ──
    for sub in SUBS[:6]:
        folder = SUB_FOLDERS[sub["short"]]
        for y in [2019, 2020]:
            for q in [1,2,3,4]:
                inv = int(sub["revenue"] * random.uniform(0.60, 0.85) / 4)
                make_docx(folder,
                          f"{sub['short']}_IC_Quartalsbericht_{y}_Q{q}.docx",
                          sub["name"], sub["city"], sub["hrb"],
                          f"Intercompany-Quartalsbericht Q{q} {y} – {sub['name']}",
                          [("Bericht",
                            [["Position","Betrag"],
                             [f"Lieferungen/Leistungen Q{q} {y}", eur(inv)],
                             ["Verrechnungspreiskonformitaet", "Bestätigt durch TP-Dokumentation"]])])

    # ── S. HR: training records per employee per year ──
    for i in range(1, 61):
        name = ALL_NAMES[i % len(ALL_NAMES)]
        for y in [2021, 2022, 2023]:
            make_xlsx("09_Personal_HR",
                      f"Schulungsnachweis_{y}_{i:03d}_{sfn(name)[:15]}.xlsx",
                      f"Schulungsnachweis {y} – {name}",
                      [("Schulungen",
                        ["Kurs","Datum","Stunden","Bestanden","Zertifikat"],
                        [[c, ds(y, random.randint(1,11), random.randint(1,25)),
                          str(random.randint(1,8)), "Ja", f"CERT-{y}-{i:03d}-{j:02d}"]
                         for j,c in enumerate(random.sample(
                             ["Code of Conduct","Datenschutz","Anti-Korruption","Exportkontrolle",
                              "IT-Sicherheit","ISO 26262","IATF 16949","Arbeitssicherheit",
                              "Fuehrungskraefte-Training","Projektmanagement"], 5))],
                        [35,15,10,10,20])])

    # ── T. Quality: per-machine OEE tracking (monthly) ──
    machines = ["SMT-Linie-1","SMT-Linie-2","SMT-Linie-3","Reflow-Ofen-A",
                "ICT-Pruefstand-1","ICT-Pruefstand-2","EOL-Tester-A","EOL-Tester-B",
                "Loetanlage-Selektivlot","Coating-Anlage"]
    for machine in machines:
        for sub in [SUBS[0], SUBS[2]]:
            for y in [2022, 2023]:
                make_xlsx("13_IATF_Qualitaet",
                          f"{sub['short']}_{sfn(machine)[:20]}_OEE_{y}.xlsx",
                          f"OEE Tracking {y} – {machine} – {sub['name']}",
                          [("OEE",
                            ["Monat","Verfügbarkeit %","Leistung %","Qualitaet %","OEE %"],
                            [[MONTHS_DE[m],
                              f"{random.uniform(85,98):.1f} %",
                              f"{random.uniform(80,96):.1f} %",
                              f"{random.uniform(97,99.9):.1f} %",
                              f"{random.uniform(68,90):.1f} %"]
                             for m in range(1,13)],
                            [12,18,15,15,12])])

    # ── U. Finance: budget variances and forecasts ──
    for sub in SUBS[:5]:
        folder = SUB_FOLDERS[sub["short"]]
        for y in [2022, 2023]:
            for q in [1,2,3,4]:
                make_xlsx(folder,
                          f"{sub['short']}_FC_Update_{y}_Q{q}.xlsx",
                          f"Forecast Update Q{q} {y} – {sub['name']}",
                          [("Forecast",
                            ["Position","Budget","Forecast","Delta","Delta %"],
                            [[pos,
                              f"{int(val):,}",
                              f"{int(val * random.uniform(0.92, 1.08)):,}",
                              f"{int(val * random.uniform(-0.08, 0.08)):+,}",
                              f"{random.uniform(-8,8):.1f} %"]
                             for pos,val in [
                                 ("Umsatz (TEUR)", sub["revenue"]/4/1000),
                                 ("Herstellkosten (TEUR)", sub["revenue"]*0.62/4/1000),
                                 ("EBIT (TEUR)", sub["revenue"]*0.072/4/1000),
                                 ("Investitionen (TEUR)", sub["revenue"]*0.055/4/1000),
                                 ("Headcount (VZÄ)", sub["employees"]/1)]],
                            [30,15,15,12,12])])

    # ── V. Compliance: per-country regulatory filings ──
    countries = [("DE","Deutschland",SUBS[0]),("PL","Polen",SUBS[2]),
                 ("CZ","Tschechien",SUBS[3]),("HU","Ungarn",SUBS[4]),
                 ("CN","China",SUBS[5])]
    reg_topics = ["Jahresabschlusshinterlegung","Handelskammer_Bericht",
                  "Sozialversicherung_Meldung","Lohnsteuer_Jahresabschluss",
                  "Umweltbericht","Arbeitssicherheitsbericht"]
    for cc, land, sub in countries:
        for topic in reg_topics:
            for y in [2021, 2022, 2023]:
                make_docx("15_Compliance_Recht",
                          f"{cc}_{sub['short']}_{topic}_{y}.docx",
                          sub["name"], sub["city"], sub["hrb"],
                          f"Regulatorische Meldung – {topic.replace('_',' ')} – {land} {y}",
                          [("Inhalt",
                            f"Pflichtmeldung gemaess lokalen Vorschriften ({land}) fuer das Jahr {y}. "
                            f"Gesellschaft: {sub['name']}, {sub['city']}. "
                            f"Thema: {topic.replace('_',' ')}. "
                            f"Eingereicht am: {ds(y+1, random.randint(1,4), random.randint(1,28))}.")])

    # ── W. Production: shift reports and production orders ──
    for sub in [SUBS[0], SUBS[2], SUBS[4]]:
        folder = SUB_FOLDERS[sub["short"]]
        for y in [2022, 2023]:
            for m in [1, 4, 7, 10]:  # quarterly representative months
                for week in [1, 2, 3, 4]:
                    make_xlsx(folder,
                              f"{sub['short']}_Schichtbericht_{y}_{m:02d}_W{week}.xlsx",
                              f"Schichtbericht KW{week} {MONTHS_DE[m]} {y} – {sub['name']}",
                              [("Produktion",
                                ["Schicht","Produkt","Plan Stk","Ist Stk","Ausschuss","Effizienz %"],
                                [[f"Frueh_{d}", prod["id"],
                                  f"{random.randint(800,3000):,}",
                                  f"{random.randint(750,3000):,}",
                                  f"{random.randint(0,30)}",
                                  f"{random.uniform(85,99):.1f} %"]
                                 for d in range(1,6)
                                 for prod in PRODUCTS[:2]],
                                [12,15,12,12,12,12])])

    # ── X. IP: Freedom to Operate analyses ──
    for prod in PRODUCTS:
        for cust in CUSTOMERS[:3]:
            make_docx("14_IP_Technologie",
                      f"FTO_{prod['id']}_{cust['short']}_2023.docx",
                      AG["recht"], "Frankfurt", "FTO-Analyse",
                      f"Freedom-to-Operate Analyse – {prod['name']} – Markt {cust['name']}",
                      [("1. Analyseumfang",
                        f"Diese FTO-Analyse prüft, ob {prod['name']} die relevanten Patente Dritter "
                        f"im Absatzmarkt {cust['name']} verletzt. Analysezeitraum: 2023."),
                       ("2. Ergebnis",
                        [["Patent","Inhaber","Relevanz","Einschaetzung"]] +
                        [[f"EP{random.randint(2000000,3999999)}B1",
                          random.choice(["Bosch","Continental","Aptiv","Visteon"]),
                          random.choice(["Gering","Mittel"]),
                          random.choice(["Keine Verletzung","Designaround erforderlich","Unkritisch"])]
                         for _ in range(random.randint(3,8))]),
                       ("3. Fazit",
                        f"Keine wesentlichen FTO-Risiken identifiziert. "
                        f"Empfehlung: Weiteres Monitoring patentrechtlicher Entwicklungen.")])

    # ── Y. Finance: IFRS accounting policies notes ──
    ifrs_standards = ["IFRS_16_Leasing","IAS_19_Pensionen","IFRS_9_Finanzinstrumente",
                      "IAS_36_Wertminderung","IFRS_15_Umsatzerloese","IAS_12_Ertragsteuern",
                      "IFRS_3_Unternehmenszusammenschluesse","IAS_38_Immaterielle_Werte"]
    for std in ifrs_standards:
        make_docx("02_Konsolidierte_Finanzen",
                  f"REA_IFRS_Bilanzierungsrichtlinie_{std}_2023.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Konzernbilanzierungsrichtlinie – {std.replace('_',' ')}",
                  [("1. Anwendungsbereich",
                    f"Diese Bilanzierungsrichtlinie regelt die Anwendung von {std.replace('_',' ')} "
                    f"im Konzernabschluss der {AG['name']} gemaess IFRS (EU-endorsed)."),
                   ("2. Wesentliche Bilanzierungsmethoden",
                    f"[Detaillierte Beschreibung der Bilanzierungsmethoden und Ermessensspielraeume "
                    f"gemaess {std.replace('_',' ')} im Brennhagen-Konzern.]"),
                   ("3. Verweise",
                    f"Erstellt: Accounting Policy Team, {AG['name']}. "
                    f"Genehmigt: {AG['cfo']}. Gueltig ab: 1. Januar 2023.")])

    # ── Z. Extra OEM docs: nomination letters, RFQ responses ──
    for cust in CUSTOMERS:
        for prod in PRODUCTS:
            # Nomination letter
            make_docx("11_Vertrieb_OEM",
                      f"REA_{cust['short']}_{prod['id']}_Nomination_Letter_{2022}.docx",
                      cust["name"], cust["contact"], "",
                      f"Nomination Letter – {prod['name']} – {cust['name']}",
                      [("Nomination",
                        f"Dear {AG['ceo']},\n\n"
                        f"We are pleased to inform you that {AG['name']} has been nominated "
                        f"as the preferred supplier for {prod['name']} for our upcoming vehicle platform. "
                        f"Expected SOP: Q{random.randint(1,4)}/{2024 + random.randint(0,2)}.\n"
                        f"Expected annual volume: {eur(random.randint(15,80)*1_000_000)}.\n\n"
                        f"Best regards,\n{cust['contact']}")])
            # RFQ response
            make_docx("11_Vertrieb_OEM",
                      f"REA_{cust['short']}_{prod['id']}_RFQ_Response_{2022}.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"Angebotsantwort (RFQ Response) – {prod['name']} – {cust['name']}",
                      [("Angebot",
                        [["Position","Beschreibung","Einheitspreis","Jahresvolumen","Gesamtbetrag"],
                         [prod["id"], prod["name"],
                          f"EUR {random.uniform(20,200):.2f}",
                          f"{random.randint(10000,200000):,} Stk.",
                          eur(random.randint(5,80)*1_000_000)]])])

    print(f"  Bulk round 3 complete. Total now: {TOTAL[0]}")


def gen_bulk_round4():
    """Fourth bulk pass – fill remaining gap to 5100."""
    print(f"\n[bulk-4] Fourth bulk pass (currently {TOTAL[0]}) ...")

    # ── AA. Per-sub quarterly reporting packages (more years) ──
    for sub in SUBS[:6]:
        folder = SUB_FOLDERS[sub["short"]]
        for y in [2019, 2020, 2021]:
            for q in [1, 2, 3, 4]:
                r_q = int(sub["revenue"] * random.uniform(0.60, 0.92) / 4)
                make_xlsx(folder,
                          f"{sub['short']}_Quartalsbericht_{y}_Q{q}.xlsx",
                          f"Quartalsbericht Q{q} {y} – {sub['name']}",
                          [("Quartal",
                            ["KPI","Ist (TEUR)","Plan (TEUR)","Delta %"],
                            [["Umsatz", f"{r_q//1000:,}", f"{sub['revenue']//4//1000:,}", f"{random.uniform(-8,8):.1f} %"],
                             ["EBIT", f"{int(r_q*0.07)//1000:,}", f"{int(sub['revenue']*0.07/4)//1000:,}", "-"],
                             ["Headcount", f"{int(sub['employees']*random.uniform(0.88,1.0)):,}", f"{sub['employees']:,}", "-"]],
                            [30,15,15,12])])

    # ── BB. Project documents – testing phase extended ──
    test_labs = ["REA_Labor_Stuttgart","TUV_Rheinland_Koeln","SGS_Hamburg","Dekra_Stuttgart","Bureau_Veritas"]
    for prj_id, prj_name, _, _, _ in PROJECTS:
        psfn_s = sfn(prj_name[:20])
        for lab in random.sample(test_labs, 3):
            make_pdf("23_Projekte_Programme",
                     f"{prj_id}_Prufbericht_{sfn(lab)[:15]}_{psfn_s}.pdf",
                     lab.replace("_", " "), "[Labor-Anschrift]", f"Prüfauftrag {prj_id}",
                     f"Prüfbericht – {prj_name} – Labor {lab.replace('_',' ')}",
                     [("Pruefergebnis",
                       [["Norm","Pruefpunkt","Sollwert","Istwert","Ergebnis"]] +
                       [[f"ISO {random.randint(10000,99999)}", f"Test {j}",
                         f"{random.uniform(0,100):.1f} Ref",
                         f"{random.uniform(0,100):.1f} Meas",
                         random.choice(["PASS","PASS","PASS","FAIL"])]
                        for j in range(1, random.randint(5, 12))])])

    # ── CC. IT: user access reviews quarterly ──
    for y in [2022, 2023]:
        for q in [1, 2, 3, 4]:
            make_xlsx("16_IT_Systeme",
                      f"REA_User_Access_Review_{y}_Q{q}.xlsx",
                      f"User Access Review Q{q} {y}",
                      [("Zugriffskontrolle",
                        ["System","Anzahl User","Reviewt","Entzogen","Status"],
                        [[sys, f"{random.randint(20,400):,}", f"{random.randint(18,400):,}",
                          f"{random.randint(0,15)}", random.choice(["Abgeschlossen","Abgeschlossen","Offen"])]
                         for sys in [AG["erp"], AG["plm"], AG["mes"], "Salesforce CRM",
                                     "SAP SuccessFactors","SharePoint","Confluence",
                                     "Jira","GitLab","Netzwerk-VPN"]],
                        [30, 15, 15, 12, 15])])

    # ── DD. Legal: per-country annual legal compliance report ──
    for sub in SUBS[:6]:
        folder = "15_Compliance_Recht"
        for y in [2021, 2022, 2023]:
            make_docx(folder,
                      f"{sub['short']}_Legal_Compliance_Report_{y}.docx",
                      sub["name"], sub["city"], sub["hrb"],
                      f"Lokaler Legal Compliance Bericht {y} – {sub['name']}",
                      [("1. Zusammenfassung",
                        f"Die {sub['name']} hielt im Geschaeftsjahr {y} alle wesentlichen "
                        f"gesetzlichen Anforderungen des Landes {sub['country']} ein."),
                       ("2. Wesentliche Rechtsthemen",
                        [["Rechtsbereich","Status","Anmerkung"],
                         ["Handelsrecht","Konform","Jahresabschluss eingereicht"],
                         ["Arbeitsrecht","Konform","Keine wesentlichen Verfahren"],
                         ["Steuerrecht","Konform",f"BP {y-3} bis {y-1} abgeschlossen"],
                         ["Datenschutz","Konform","DSGVO-Audit bestanden"],
                         ["Umweltrecht","Konform","Genehmigungen erneuert"]])])

    # ── EE. Kapitalmarkt: additional disclosure documents ──
    for i in range(1, 21):
        doc_type = random.choice(["Stimmrechtsmitteilung","WpHG_Meldung","Directors_Dealing",
                                   "Short_Selling_Meldung","Beteiligungsmeldung"])
        make_docx("10_Kapitalmarkt_IR",
                  f"REA_{doc_type}_{2022+i%2}_{i:03d}.docx",
                  AG["name"], AG["full_address"], f"ISIN {AG['isin']}",
                  f"Kapitalmarktrechtliche Meldung – {doc_type.replace('_',' ')} #{i:03d}",
                  [("Meldung",
                    f"Gemaess §§ 33 ff. WpHG / Art. 19 MAR meldet {AG['name']} / [Meldepflichtige Person] "
                    f"folgendes: {doc_type.replace('_',' ')}. "
                    f"Datum: {ds(2022+i%2, random.randint(1,12), random.randint(1,25))}. "
                    f"Betroffene Aktien: {random.randint(10000,500000):,} Stueck "
                    f"({random.uniform(1,8):.2f} % des Grundkapitals).")])

    # ── FF. Treasury: FX hedging documentation ──
    currencies = [("EUR/CNY", "CN"), ("EUR/PLN", "PL"), ("EUR/CZK", "CZ"), ("EUR/HUF", "HU")]
    for pair, cc in currencies:
        for y in [2022, 2023]:
            for q in [1, 2, 3, 4]:
                notional = random.randint(5, 30) * 1_000_000
                make_docx("02_Konsolidierte_Finanzen",
                          f"REA_FX_Hedge_{pair.replace('/','_')}_{y}_Q{q}.docx",
                          AG["name"], AG["full_address"], AG["hrb"],
                          f"FX-Absicherungsdokumentation – {pair} – Q{q} {y}",
                          [("Sicherungsgeschäft",
                            [["Parameter","Inhalt"],
                             ["Instrument", random.choice(["Devisentermingeschäft","Devisenswap","FX-Option"])],
                             ["Währungspaar", pair],
                             ["Nominalbetrag", eur(notional)],
                             ["Laufzeit", f"Q{q}/{y} bis Q{min(q+1,4)}/{y if q<4 else y+1}"],
                             ["Kurs gesichert", f"{random.uniform(0.8,1.2):.4f}"],
                             ["Hedge Accounting", random.choice(["Cashflow Hedge","Fair Value Hedge","Nein"])],
                             ["Kontrahent", random.choice(["Deutsche Bank","UniCredit","BNP Paribas"])]])])

    # ── GG. Versicherungen: per-event claim documents ──
    for i in range(1, 30):
        make_docx("17_Versicherungen",
                  f"Schadensmeldung_{2020+i%4}_{i:04d}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Schadensmeldung #{2020+i%4}-{i:04d}",
                  [("Schaden",
                    [["Feld","Inhalt"],
                     ["Schadensdatum", ds(2020+i%4, random.randint(1,12), random.randint(1,25))],
                     ["Sparte", random.choice(["Produkthaftpflicht","Betriebshaftpflicht","Cyber","Feuer","Transport"])],
                     ["Schadensbeschreibung", random.choice(["Maschinenschaden","Kundenreklamation Rueckruf","Ransomware","Brandschaden","Transportschaden"])],
                     ["Schadenshoehe (geschaetzt)", eur(random.randint(50_000, 5_000_000))],
                     ["Status", random.choice(["Gemeldet","In Regulierung","Abgeschlossen","Abgelehnt"])],
                     ["Versicherer", random.choice(["AXA XL","Allianz","HDI","Generali","Chubb"])]])])

    # ── HH. Betriebsraete: EBR documents ──
    make_docx("21_Betriebsraete", "REA_EBR_Vereinbarung_2019.docx",
              AG["name"], AG["full_address"], AG["hrb"],
              "Vereinbarung zur Errichtung eines Europäischen Betriebsrats (EBR) – Brennhagen Elektronik AG",
              [("1. Errichtung", f"Gemaess der Richtlinie 2009/38/EG (EBR-Richtlinie) und dem SE-Beteiligungsgesetz "
                f"wird bei der {AG['name']} ein Europäischer Betriebsrat (EBR) errichtet."),
               ("2. Zusammensetzung", "Der EBR besteht aus Vertretern der deutschen, polnischen, tschechischen und ungarischen Konzerngesellschaften."),
               ("3. Informations- und Konsultationsrechte", "Der EBR ist zu unterrichten und anzuhören bei: wesentlichen Restrukturierungen, Änderungen der Unternehmensstruktur, Einführung neuer Technologien und Managementmethoden.")])

    for y in [2021, 2022, 2023]:
        make_docx("21_Betriebsraete", f"REA_EBR_Jahrestagung_{y}_Protokoll.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"EBR-Jahrestagung {y} – Protokoll",
                  [("Tagesordnung", f"EBR-Jahrestagung {y}, Stuttgart.\nTOP 1: Bericht Vorstand – Wirtschaftslage.\nTOP 2: Konzernstruktur-Änderungen.\nTOP 3: Beschäftigungsentwicklung.\nTOP 4: Sonstiges."),
                   ("Ergebnis", f"Die EBR-Jahrestagung {y} fand am {ds(y,11,random.randint(15,25))} in Stuttgart statt. Alle Tagesordnungspunkte wurden erörtert.")])

    # ── II. MA: pipeline screening docs ──
    for i in range(1, 25):
        doc_type = random.choice(["Target_Screening","First_Approach","Management_Meeting",
                                   "Finanzanalyse","Synergieberechnung","Bewertungsmodell"])
        make_docx("19_MA_Transaktionen",
                  f"MA_{doc_type}_{2023}_{i:03d}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"M&A Pipeline – {doc_type.replace('_',' ')} #{i:03d}",
                  [("Dokument",
                    f"Vertrauliches M&A-Dokument: {doc_type.replace('_',' ')} im Rahmen des strategischen "
                    f"Prozesses 2023/2024. Ziel-Unternehmen: [Target Name, vertraulich]. "
                    f"Erstellt: {random.choice([AG['ceo'], AG['cfo']])}. "
                    f"Datum: {ds(2023, random.randint(1,12), random.randint(1,25))}.")])

    # ── JJ. Einkauf: supplier risk reports ──
    for sup in SUPPLIERS:
        for y in [2021, 2022, 2023]:
            make_pdf("12_Einkauf_Lieferanten",
                     f"{sup['short']}_Lieferantenrisikobericht_{y}.pdf",
                     AG["name"], AG["full_address"], AG["hrb"],
                     f"Lieferantenrisikobericht {y} – {sup['name']}",
                     [("Risikoübersicht",
                       [["Risikobereich","Bewertung","Massnahme"],
                        ["Finanzstaerke", random.choice(["Gut","Sehr gut"]), "Monitoring"],
                        ["Lieferfaehigkeit", random.choice(["Gut","Sehr gut","Eingeschraenkt"]),
                         random.choice(["Beobachtung","LTA vereinbart","Dual Source aufgebaut"])],
                        ["Technologie-Roadmap", random.choice(["Ausgerichtet","Teilweise ausgerichtet"]), ""],
                        ["ESG/LkSG", random.choice(["Konform","Pruefung ausstehend"]),
                         random.choice(["OK","Audit geplant"])],
                        ["Geopolitisch", random.choice(["Gering","Mittel"]),
                         random.choice(["Monitoring","Risikominderungsplan"])]])])

    # ── KK. Strategy: innovation pipeline ──
    innovations = [
        "Zonal_ECU_Architektur","Software_Defined_Vehicle_OS","48V_Bordnetz_Modul",
        "V2X_Kommunikationsmodul","LiDAR_Integration","AI_Driven_ADAS","Solid_State_Battery_BMS",
        "Automotive_WIFI7_Module","OTA_Update_Platform_v3","Digital_Twin_Fertigung"
    ]
    for inn in innovations:
        make_docx("20_Strategie_Vorstand",
                  f"REA_Innovation_{inn}_2024.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Innovationsmemo – {inn.replace('_',' ')} – 2024",
                  [("1. Innovation Overview",
                    f"Technology: {inn.replace('_',' ')}. "
                    f"Market opportunity: EUR {random.randint(1,15):.1f}bn by 2030. "
                    f"REA readiness level: TRL {random.randint(3,7)}."),
                   ("2. Business Case",
                    [["Parameter","Wert"],
                     ["R&D Budget", eur(random.randint(2,15)*1_000_000)],
                     ["Time to Market", f"{random.randint(2,5)} Jahre"],
                     ["Erwarteter Umsatz Jahr 5", eur(random.randint(20,80)*1_000_000)],
                     ["ROI", f"{random.randint(15,40)} %"],
                     ["Risiko", random.choice(["Mittel","Hoch","Mittel"])]])])

    # ── LL. Pensionen: individual benefit statements ──
    for i in range(1, 31):
        name = ALL_NAMES[i % len(ALL_NAMES)] + f"-P{i}"
        make_docx("22_Pensionen_bAV",
                  f"bAV_Jahresinfo_{2023}_{i:03d}_{sfn(name)[:15]}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Jährliche Standmitteilung bAV 2023 – {name}",
                  [("Versorgungsinformation",
                    [["Feld","Inhalt"],
                     ["Anwärter", name],
                     ["Plantyp", random.choice(["Direktzusage","Entgeltumwandlung","Pensionskasse"])],
                     ["Angesammelter Anspruch", eur(random.randint(800,5000)) + "/Monat"],
                     ["Unverfallbarkeit", "Ja"],
                     ["Übertragungswert (IFRS)", eur(random.randint(50_000, 400_000))],
                     ["Naechster Rentenstart", f"{2030 + random.randint(0,20)}"]])])

    print(f"  Bulk round 4 complete. Total now: {TOTAL[0]}")


def gen_bulk_round5():
    """Fifth bulk pass – final push to 5100+."""
    print(f"\n[bulk-5] Final bulk pass (currently {TOTAL[0]}) ...")

    # ── MM. All subs: annual internal reports 2015-2020 ──
    for sub in SUBS[:6]:
        folder = SUB_FOLDERS[sub["short"]]
        for y in range(2015, 2022):
            r = int(sub["revenue"] * max(0.3, 0.5 + (y-2015)*0.06))
            make_xlsx(folder,
                      f"{sub['short']}_Jahresbericht_intern_{y}.xlsx",
                      f"Interner Jahresbericht {y} – {sub['name']}",
                      [("Ergebnisse",
                        ["KPI",f"{y}",f"{y-1}"],
                        [["Umsatz (TEUR)", f"{r//1000:,}", f"{int(r*0.93)//1000:,}"],
                         ["Mitarbeiter", f"{int(sub['employees']*max(0.5,(y-2015+3)/9)):,}", "-"],
                         ["EBIT-Marge %", f"{random.uniform(5,9):.1f} %", "-"]],
                        [30,15,15])])

    # ── NN. Quality: SPC control charts per product per line (monthly 2023) ──
    lines = ["Linie-1","Linie-2","Linie-3","Linie-4"]
    for prod in PRODUCTS[:3]:
        for line in lines:
            for y in [2022, 2023]:
                make_xlsx("13_IATF_Qualitaet",
                          f"SPC_{prod['id']}_{sfn(line)[:10]}_{y}.xlsx",
                          f"SPC-Auswertung {y} – {prod['name']} – {line}",
                          [("SPC",
                            ["Monat","Xbar","R","Cpk","UCL","LCL","Status"],
                            [[MONTHS_DE[m],
                              f"{random.uniform(4.95,5.05):.4f}",
                              f"{random.uniform(0.01,0.08):.4f}",
                              f"{random.uniform(1.3,2.2):.2f}",
                              "5.10","4.90",
                              random.choice(["OK","OK","OK","Warnung"])]
                             for m in range(1,13)],
                            [10,12,12,10,10,10,10])])

    # ── OO. Compliance: data breach notifications (DSGVO) ──
    for i in range(1, 15):
        make_docx("15_Compliance_Recht",
                  f"DSGVO_Datenpanne_{2020+i%4}_{i:03d}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Datenpannen-Meldung Art. 33 DSGVO – #{i:03d}",
                  [("Meldung",
                    [["Feld","Inhalt"],
                     ["Vorfallnummer", f"DP-{2020+i%4}-{i:03d}"],
                     ["Datum der Entdeckung", ds(2020+i%4, random.randint(1,12), random.randint(1,25))],
                     ["Art der Panne", random.choice(["Unberechtigter Zugriff","Datenverlust","Fehlversand","Ransomware"])],
                     ["Betroffene Datenkategorien", random.choice(["Mitarbeiterdaten","Kundendaten","keine personenbezogenen Daten"])],
                     ["Betroffene Personen geschaetzt", f"{random.randint(0,1000):,}"],
                     ["Meldung an Aufsichtsbehörde", random.choice(["Ja – LfDI BW","Nein – kein hohes Risiko"])],
                     ["Status", random.choice(["Abgeschlossen","In Bearbeitung"])]])])

    # ── PP. IT: software license inventory ──
    for y in [2021, 2022, 2023]:
        make_xlsx("16_IT_Systeme",
                  f"REA_Software_Lizenz_Inventar_{y}.xlsx",
                  f"Software-Lizenzinventar {y} – Brennhagen Elektronik AG",
                  [("Lizenzen",
                    ["Hersteller","Produkt","Version","Lizenztyp","Anzahl Lizenzen","Kosten p.a. (EUR)","Ablauf"],
                    [["SAP","S/4HANA","2022","Enterprise",f"{random.randint(800,1200):,}","3.200.000","unbefristet"],
                     ["Siemens","Teamcenter PLM","12.4","Named User",f"{random.randint(100,300):,}","580.000","31.12.2024"],
                     ["Microsoft","M365 E3","Current","Per-User",f"{random.randint(1000,1500):,}","780.000","31.12.2024"],
                     ["Autodesk","AutoCAD LT","2024","Annual",f"{random.randint(50,100):,}","95.000","30.06.2025"],
                     ["Mathworks","MATLAB","R2023b","Concurrent",f"{random.randint(20,50):,}","120.000","31.08.2025"],
                     ["PTC","CREO Parametric","9.0","Floating",f"{random.randint(30,80):,}","210.000","30.11.2024"],
                     ["Jama Software","Jama Connect","SaaS","User",f"{random.randint(50,150):,}","85.000","31.12.2024"],
                     ["Confluence/Jira","Atlassian","Cloud","Per-User",f"{random.randint(200,500):,}","45.000","31.12.2024"]],
                    [15,25,12,15,18,20,15])])

    # ── QQ. IATF: process audit observations (detailed) ──
    for sub in SUBS[:4]:
        for i in range(1, 11):
            pcat = random.choice(["SMT","Reflow","Endpruefung","Montage","Materiallager","Wareneingang"])
            make_docx("13_IATF_Qualitaet",
                      f"{sub['short']}_Prozessaudit_Beobachtung_{2023}_{i:03d}.docx",
                      sub["name"], sub["city"], sub["hrb"],
                      f"Prozessaudit-Beobachtung {2023} #{i:03d} – {sub['name']} – {pcat}",
                      [("Beobachtung",
                        [["Feld","Inhalt"],
                         ["Prozess", pcat],
                         ["Befundtyp", random.choice(["Beobachtung","Minor Finding","Major Finding"])],
                         ["Normanforderung", f"IATF 16949:{2023} Kap. {random.randint(6,10)}.{random.randint(1,9)}"],
                         ["Befund", random.choice(["Fehlende Dokumentation","Prozessabweichung","Schulungsnachweis unvollstaendig"])],
                         ["CAPA initiiert", random.choice(["Ja","Nein"])],
                         ["Termin", ds(2023, random.randint(6,12), 30)]])])

    # ── RR. Kapitalmarkt: analyst research notes (incoming) ──
    banks = ["Deutsche Bank Research","Goldman Sachs","J.P. Morgan","UBS","Berenberg Bank",
             "Warburg Research","Hauck & Aufhaeuser","Baader Bank"]
    for bank in banks:
        for y in [2022, 2023]:
            for q in [1, 2]:
                make_pdf("10_Kapitalmarkt_IR",
                         f"Analyst_{sfn(bank)[:15]}_{y}_Q{q}.pdf",
                         bank, "[Research Address]", f"Analyst Research",
                         f"Equity Research – {AG['name']} – {bank} – Q{q} {y}",
                         [("Investment Summary",
                           [["Parameter","Value"],
                            ["Rating", random.choice(["Buy","Buy","Hold","Sell"])],
                            ["Price Target (EUR)", f"{random.uniform(42,68):.0f}"],
                            ["Current Price (EUR)", f"{random.uniform(38,62):.0f}"],
                            ["Upside/Downside", f"{random.uniform(-15,25):.0f} %"],
                            ["EV/EBITDA Fwd", f"{random.uniform(7,14):.1f}x"],
                            ["P/E Fwd", f"{random.uniform(12,22):.1f}x"]])])

    # ── SS. Finance: segment reporting detail (quarterly 2020-2021) ──
    segments = ["Powertrain_Electronics","ADAS_Safety","EV_Energy","Infotainment"]
    for seg in segments:
        for y in [2020, 2021]:
            for q in [1, 2, 3, 4]:
                r_seg = int(AG[f"revenue_{y}"] * random.uniform(0.08, 0.30) / 4)
                make_xlsx("02_Konsolidierte_Finanzen",
                          f"REA_Segment_{sfn(seg)[:20]}_{y}_Q{q}.xlsx",
                          f"Segmentbericht {seg.replace('_',' ')} Q{q} {y}",
                          [("Segment",
                            ["Position","Q{q} {y} (TEUR)","Q{q} {y-1} (TEUR)"],
                            [["Umsatz", f"{r_seg//1000:,}", f"{int(r_seg*0.92)//1000:,}"],
                             ["EBIT", f"{int(r_seg*0.075)//1000:,}", "-"],
                             ["Assets", f"{int(r_seg*2.2)//1000:,}", "-"]],
                            [30,18,18])])

    # ── TT. Personal: absence management records ──
    for i in range(1, 31):
        name = ALL_NAMES[i % len(ALL_NAMES)] + f"-A{i}"
        for y in [2022, 2023]:
            make_xlsx("09_Personal_HR",
                      f"Abwesenheit_{y}_{i:03d}_{sfn(name)[:12]}.xlsx",
                      f"Abwesenheitskalender {y} – {name}",
                      [("Abwesenheiten",
                        ["Monat","Urlaubstage","Krankheitstage","Sonstige","Saldo"],
                        [[MONTHS_DE[m],
                          f"{random.randint(0,5)}",
                          f"{random.randint(0,3)}",
                          f"{random.randint(0,1)}",
                          f"{random.randint(15,30)}"]
                         for m in range(1,13)],
                        [12,15,15,12,12])])

    # ── UU. Einkauf: Commodity Risk Reports ──
    commodities = ["Kupfer","Gold_Sputtertarget","Zinn_Lotzinn","Silber_Leitpaste",
                   "Epoxidharz_PCB","Silizium_Wafer","Seltene_Erden_REE","Kobalt_BEV"]
    for commodity in commodities:
        for y in [2022, 2023]:
            for q in [1, 2, 3, 4]:
                make_xlsx("12_Einkauf_Lieferanten",
                          f"REA_Rohstoffrisiko_{sfn(commodity)[:15]}_{y}_Q{q}.xlsx",
                          f"Rohstoff-Risikobericht Q{q} {y} – {commodity.replace('_',' ')}",
                          [("Preisanalyse",
                            ["Datum","Preis (USD/t oder Index)","3M-Outlook","Hedging-Empfehlung"],
                            [[ds(y, (q-1)*3+m, random.randint(1,28)),
                              f"{random.uniform(5000,80000):.0f}",
                              random.choice(["Steigend","Seitwärts","Fallend"]),
                              random.choice(["Absichern (6M)","Kein Hedging","Absichern (3M)"])]
                             for m in range(1,4)],
                            [18,22,15,25])])

    print(f"  Bulk round 5 complete. Total now: {TOTAL[0]}")


def gen_final_fill():
    """Final fill – generate remaining docs to hit 5100."""
    print(f"\n[fill] Final fill pass (currently {TOTAL[0]}) ...")

    # Per-sub: management letters and Wirtschaftsprüfer communications
    for sub in SUBS[:6]:
        folder = SUB_FOLDERS[sub["short"]]
        for y in [2021, 2022, 2023]:
            make_docx(folder,
                      f"{sub['short']}_WP_Management_Letter_{y}.docx",
                      AG["wp"], "Klingelhöferstraße 18, 10785 Berlin", "WPK-Nr. 2063",
                      f"Management Letter {y} – {sub['name']}",
                      [("Feststellungen",
                        f"Im Rahmen der Jahresabschlussprüfung {y} der {sub['name']} haben wir "
                        f"folgende Feststellungen getroffen: {random.randint(0,3)} Hinweise zur "
                        f"Optimierung des internen Kontrollsystems. Alle Hinweise wurden mit dem "
                        f"Management erörtert und Massnahmenpläne wurden vereinbart.")])

    # Per-product: technical change management docs
    for prod in PRODUCTS:
        for i in range(1, 16):
            make_docx("14_IP_Technologie",
                      f"ECO_{prod['id']}_{i:03d}_Engineering_Change_{2022+i%2}.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"Engineering Change Order – {prod['name']} ECO-{i:03d}",
                      [("ECO Details",
                        [["Feld","Inhalt"],
                         ["ECO-Nr.", f"ECO-{prod['id']}-{i:03d}"],
                         ["Aenderungsgrund", random.choice(["Kostenoptimierung","Qualitaetsverbesserung","Bauteilabkuendigung","Regulierung"])],
                         ["Betroffene Komponente", f"{prod['id']}-BOM-Position-{random.randint(10,99)}"],
                         ["Auswirkung", random.choice(["PPAP erforderlich","Kein PPAP","Re-Validierung"])],
                         ["Status", random.choice(["Genehmigt","In Review","Implementiert"])],
                         ["Datum", ds(2022+i%2, random.randint(1,12), random.randint(1,25))]])])

    # Treasury: additional bank correspondence
    banks_ext = ["Deutsche Bank AG","UniCredit Bank AG","Commerzbank AG","BNP Paribas","HSBC Deutschland"]
    for bank in banks_ext:
        for y in [2022, 2023]:
            make_docx("02_Konsolidierte_Finanzen",
                      f"REA_{sfn(bank)[:15]}_Bankkorrespondenz_{y}.docx",
                      AG["name"], AG["full_address"], AG["hrb"],
                      f"Bankkorrespondenz – {bank} – {y}",
                      [("Korrespondenz",
                        f"Korrespondenz zwischen {AG['name']} und {bank} im Geschaeftsjahr {y}. "
                        f"Betreff: {random.choice(['Kreditlinie','Zinssatzanpassung','Sicherheitenstellung','Zahlungsverkehr','Kontoauszug','Devisenkurs'])}. "
                        f"Datum: {ds(y, random.randint(1,12), random.randint(1,25))}.")])

    # HR: salary notifications and payslips (sample)
    for i in range(1, 31):
        name = ALL_NAMES[i % len(ALL_NAMES)] + f"-S{i}"
        make_docx("09_Personal_HR",
                  f"Gehaltsaenderung_{2023}_{i:03d}_{sfn(name)[:12]}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Gehaltsänderungsmitteilung 2023 – {name}",
                  [("Mitteilung",
                    f"Ihr Gehalt wird ab dem 1. April 2023 wie folgt angepasst:\n"
                    f"Bisheriges Jahresgehalt: EUR {random.randint(60,180)*1000:,}\n"
                    f"Neues Jahresgehalt: EUR {random.randint(62,190)*1000:,}\n"
                    f"Erhöhung: {random.uniform(2,5):.1f} %\n\n"
                    f"Diese Anpassung erfolgt im Rahmen der jährlichen Gehaltsrunde.")])

    # IATF: measurement uncertainty studies
    for i in range(1, 21):
        make_docx("13_IATF_Qualitaet",
                  f"MSA_Messunsicherheitsstudie_{2023}_{i:03d}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"MSA Messunsicherheitsstudie #{i:03d} – 2023",
                  [("MSA Ergebnis",
                    [["Parameter","Wert","Anforderung","Status"],
                     ["Gage R&R", f"{random.uniform(5,18):.1f} %", "< 30 %", "OK"],
                     ["Wiederholbarkeit", f"{random.uniform(3,10):.1f} %", "Anteil < 15 %", "OK"],
                     ["Reproduzierbarkeit", f"{random.uniform(2,9):.1f} %", "Anteil < 15 %", "OK"],
                     ["Anzahl Diskriminierungen", f"{random.randint(5,15)}", "> 5", "OK"]])])

    # Strategy: investor day Q&A documents
    for y in [2021, 2022, 2023]:
        make_docx("20_Strategie_Vorstand",
                  f"REA_Investorentag_{y}_QA_Protokoll.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Investorentag {y} – Q&A Protokoll",
                  [("Q&A Zusammenfassung",
                    f"Capital Markets Day {y} – Fragen und Antworten.\n\n"
                    f"F: Wie entwickelt sich die Marge im EV-Segment?\n"
                    f"A: {AG['cfo']}: Wir erwarten bis 2026 eine EBITDA-Marge von 13 %+ im EV-Bereich.\n\n"
                    f"F: Wie hoch ist die Abhängigkeit von BMW?\n"
                    f"A: {AG['ceo']}: BMW ist mit {AG['kunde1_rev_share']} der größte Kunde, "
                    f"wir diversifizieren aktiv über Stellantis und Continental.\n\n"
                    f"F: Ist eine weitere Kapitalerhöhung geplant?\n"
                    f"A: {AG['cfo']}: Nein, die aktuelle Kapitalstruktur ist solide.")])

    # Compliance: anti-bribery due diligence third parties
    for i in range(1, 21):
        make_docx("15_Compliance_Recht",
                  f"TPDD_Third_Party_{2022+i%2}_{i:03d}.docx",
                  AG["name"], AG["full_address"], AG["hrb"],
                  f"Third Party Due Diligence – Business Partner #{i:03d}",
                  [("TPDD Ergebnis",
                    [["Feld","Inhalt"],
                     ["Partner", f"[Business Partner {i} – vertraulich]"],
                     ["Land", random.choice(["China","Russland","Tuerkei","Indien","Mexiko","VAE"])],
                     ["Typ", random.choice(["Vertriebspartner","Berater","Agent","JV-Partner"])],
                     ["PEP-Pruefung", random.choice(["Kein Treffer","Kein Treffer","Kein Treffer","Treffer – untersucht"])],
                     ["Sanktionspruefung", "Kein Treffer"],
                     ["Risikoklasse", random.choice(["Gering","Mittel","Hoch"])],
                     ["Freigabe", random.choice(["Erteilt","Abgelehnt","Mit Auflagen"])]])])

    print(f"  Final fill complete. Total now: {TOTAL[0]}")


# ═══════════════════════════════════════════════════════════════════════════
# MAIN – run all generators
# ═══════════════════════════════════════════════════════════════════════════

def main():
    print("=" * 60)
    print("Brennhagen Elektronik AG – DD Document Generator")
    print(f"Target: ~5100 documents → {BASE}")
    print("=" * 60)

    gen_00()
    gen_01()
    gen_02()
    gen_subsidiaries()
    gen_09()
    gen_10()
    gen_11()
    gen_12()
    gen_13()
    gen_14()
    gen_15()
    gen_16()
    gen_17()
    gen_18()
    gen_19()
    gen_20()
    gen_21()
    gen_22()
    gen_23()
    gen_extra_padding()
    gen_bulk_to_target()
    gen_bulk_round2()
    gen_bulk_round3()
    gen_bulk_round4()
    gen_bulk_round5()
    gen_final_fill()

    print("\n" + "=" * 60)
    print(f"DONE – Total documents created: {TOTAL[0]}")
    print(f"Output directory: {BASE}")
    print("=" * 60)

    # Count per folder
    print("\nDocuments per folder:")
    grand = 0
    for f in FOLDERS:
        fp = BASE / f
        count = len(list(fp.glob("*")))
        grand += count
        print(f"  {f:45s}: {count:5d}")
    print(f"\n  {'TOTAL':45s}: {grand:5d}")


if __name__ == "__main__":
    main()
