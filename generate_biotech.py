#!/usr/bin/env python3
"""
generate_company2.py
Generates ~1000 realistic due-diligence documents for Sentavia Precision GmbH.
Output: biotech_medium/
"""
# --- portable-paths-prelude --- (do not edit) ---
import sys
from pathlib import Path as _PathlibPath
_RP = _PathlibPath(__file__).resolve().parent
while not (_RP / "enhance_lib.py").exists() and _RP.parent != _RP:
    _RP = _RP.parent
_ROOT = _RP
sys.path.insert(0, str(_ROOT))
# --- end prelude ---

import os
import random
from datetime import date, timedelta, datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                 TableStyle, HRFlowable)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

random.seed(42)

# ─────────────────────────────────────────────
# CANONICAL COMPANY DATA
# ─────────────────────────────────────────────
C = {
    "name": "Sentavia Precision GmbH",
    "short": "BTP",
    "hrb": "HRB 218934",
    "amtsgericht": "Amtsgericht München",
    "founded": "12. Oktober 2008",
    "founded_year": 2008,
    "address": "Freimannstraße 45",
    "plz": "80939",
    "city": "München",
    "full_address": "Freimannstraße 45, 80939 München",
    "phone": "+49 89 32178-0",
    "email": "info@sentavia-precision.de",
    "web": "www.sentavia-precision.de",
    "ust_id": "DE 287 432 109",
    "steuernr": "143/201/89231",
    "fa": "Finanzamt München-Schwabing",
    "iban": "DE42 7002 0270 0012 3456 00",
    "bic": "HYVEDEMMXXX",
    "bank": "UniCredit Bank AG (HypoVereinsbank)",
    # Management
    "ceo": "Dr. Katharina Berger",
    "cto": "Dr. Marcus Vogt",
    "cfo": "Thomas Müller",
    "cmo": "Dr. Annika Schmidt",
    "qra": "Dipl.-Ing. Stefan Hoffmann",
    # Shareholders
    "investor1": "Sofinnova Partners (Paris)", "anteil1": "35 %",
    "investor2": "Lakestar (Zürich)", "anteil2": "25 %",
    "investor3": "Management-Beteiligungs-GbR", "anteil3": "15 %",
    "investor4": "HTGF (High-Tech Gründerfonds)", "anteil4": "10 %",
    "investor5": "Weiterer Streubesitz", "anteil5": "15 %",
    # Financials
    "stammkapital": 500_000,
    "revenue_2021": 58_100_000,
    "revenue_2022": 71_400_000,
    "revenue_2023": 87_200_000,
    "revenue_2024e": 104_000_000,
    "ebitda_2021": 7_200_000,
    "ebitda_2022": 9_800_000,
    "ebitda_2023": 12_400_000,
    "ebit_2021": 4_900_000,
    "ebit_2022": 7_100_000,
    "ebit_2023": 9_200_000,
    "employees_2021": 418,
    "employees_2022": 512,
    "employees_2023": 612,
    # Products
    "prod1_name": "Cardevio Pro",
    "prod1_class": "Klasse IIb (MDR Anhang VIII, Regel 10)",
    "prod1_ref": "MDR-UDI-DE/BTP/CS-PRO-001",
    "prod1_ce": "CE 0123-BTP-Cardevio-2022",
    "prod2_name": "Ostevo Navigator",
    "prod2_class": "Klasse IIa (MDR Anhang VIII, Regel 7)",
    "prod2_ref": "MDR-UDI-DE/BTP/OF-NAV-002",
    "prod3_name": "Veridiq SARS-Flex",
    "prod3_class": "IVD Klasse C (IVDR Anhang VIII, Regel 3)",
    "prod3_ref": "IVDR-UDI-DE/BTP/DK-SF-003",
    # Regulatory
    "nb": "TÜV SÜD Product Service GmbH",
    "nb_id": "0123",
    "nb_cert1": "G1 20 12 87024 018",
    "eudamed_reg": "SRN-DE-000012789",
    "iso_cert": "DIN EN ISO 13485:2016",
    "iso_cert_nr": "Q1 20 12 87024 012",
    # Key customers / distributors
    "dist1": "Siemens Healthineers AG",
    "dist2": "Fresenius Medical Care AG & Co. KGaA",
    "dist3": "B. Braun SE",
    "hosp1": "Charité – Universitätsmedizin Berlin",
    "hosp2": "Universitätsklinikum Hamburg-Eppendorf (UKE)",
    "hosp3": "LMU Klinikum München",
    "hosp4": "Universitätsklinikum Heidelberg",
    # Key suppliers
    "lief1": "Texas Instruments Deutschland GmbH",
    "lief2": "Sensirion AG (Schweiz)",
    "lief3": "Medi-Tec Supply GmbH",
    "lief4": "Schreiner MediPharm GmbH",
    # Advisors
    "wp": "PricewaterhouseCoopers GmbH WPG",
    "steuerber": "Deloitte GmbH Wirtschaftsprüfungsgesellschaft",
    "anwalt": "Noerr PartGmbB",
    "patent_anwalt": "Boehmert & Boehmert Partnerschaft",
    # IT
    "erp": "SAP S/4HANA Cloud (Life Sciences Edition)",
    "qms_tool": "MasterControl Quality Excellence Suite",
    "clinical_db": "REDCap Clinical Database",
}

BASE = Path(f"{_ROOT}/biotech_medium")

MONTHS_DE = ["", "Januar", "Februar", "März", "April", "Mai", "Juni",
             "Juli", "August", "September", "Oktober", "November", "Dezember"]

GENERATED_FILES = []


def ds(y, m, d):
    return f"{d}. {MONTHS_DE[m]} {y}"


def eur(n):
    return f"{n:,.0f} EUR".replace(",", ".")


def sfn(s, maxlen=60):
    for ch in ['/', '\\', ':', '*', '?', '"', '<', '>', '|', '(', ')', '&', '+', ',']:
        s = s.replace(ch, '_')
    return s[:maxlen]


def rdate(year, month_start=1, month_end=12):
    m = random.randint(month_start, month_end)
    d = random.randint(1, 28)
    return date(year, m, d)


# ─────────────────────────────────────────────
# MESS INJECTION
# ─────────────────────────────────────────────
def mess_filename(name, probability=0.08):
    if random.random() < probability:
        if '.' in name:
            base, ext = name.rsplit('.', 1)
        else:
            base, ext = name, ''
        suffix = random.choice(['_v1', '_v2', '_DRAFT', '_ENTWURF', '_FINAL',
                                '_FINAL_v2', '_rev_KBerger', '_20231015',
                                '_reviewed', '_2024-01-15'])
        return f"{base}{suffix}.{ext}" if ext else f"{base}{suffix}"
    return name


def maybe_draft_header(doc, probability=0.06):
    if random.random() < probability:
        p = doc.add_paragraph("*** ENTWURF – NICHT ZUR WEITERGABE ***")
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xCC, 0, 0)
        return True
    return False


def missing_exhibit_note():
    if random.random() < 0.12:
        return "\n\n[Anlage " + random.choice(['1', '2', '3', 'A', 'B']) + " – noch beizufügen / to be attached]"
    return ""


# ─────────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────────
def make_docx_header(doc, title, subtitle=None, confidential=False):
    maybe_draft_header(doc)
    if confidential:
        p = doc.add_paragraph("VERTRAULICH / CONFIDENTIAL")
        p.runs[0].font.color.rgb = RGBColor(0xCC, 0, 0)
        p.runs[0].bold = True
    h = doc.add_heading(C["name"], level=1)
    p = doc.add_paragraph()
    r1 = p.add_run(C["full_address"])
    r1.font.size = Pt(9)
    r2 = p.add_run("  |  " + C["hrb"] + "  |  " + C["amtsgericht"])
    r2.font.size = Pt(9)
    doc.add_paragraph("─" * 80)
    doc.add_heading(title, level=2)
    if subtitle:
        sub = doc.add_paragraph(subtitle)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def save_docx(doc, folder, filename):
    filename = mess_filename(filename)
    path = BASE / folder / filename
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    GENERATED_FILES.append(str(path))
    return path


def add_clause(doc, number, title, paragraphs):
    doc.add_heading(f"§ {number} {title}", level=3)
    for i, text in enumerate(paragraphs, 1):
        p = doc.add_paragraph()
        p.add_run(f"({i}) ").bold = True
        p.add_run(text)


def add_section(doc, title, paragraphs):
    doc.add_heading(title, level=3)
    for text in paragraphs:
        doc.add_paragraph(text)


# ─────────────────────────────────────────────
# PDF HELPERS
# ─────────────────────────────────────────────
def make_pdf_doc(folder, filename, title, sections, subtitle=None, confidential=False, scan_prefix=False):
    if scan_prefix and random.random() < 0.15:
        filename = "SCAN_" + filename
    filename = mess_filename(filename)
    path = BASE / folder / filename
    path.parent.mkdir(parents=True, exist_ok=True)
    pdoc = SimpleDocTemplate(str(path), pagesize=A4,
                             leftMargin=2.5 * cm, rightMargin=2.5 * cm,
                             topMargin=2.5 * cm, bottomMargin=2.5 * cm)
    styles = getSampleStyleSheet()
    body_s = ParagraphStyle('b', parent=styles['Normal'], fontSize=10, leading=15,
                            spaceAfter=8, alignment=TA_JUSTIFY)
    h1_s = ParagraphStyle('h1', parent=styles['Heading1'], fontSize=14,
                          spaceBefore=14, spaceAfter=6, textColor=colors.HexColor("#1A3A6B"))
    h2_s = ParagraphStyle('h2', parent=styles['Heading2'], fontSize=11,
                          spaceBefore=10, spaceAfter=4, textColor=colors.HexColor("#1A3A6B"))
    small_s = ParagraphStyle('sm', parent=styles['Normal'], fontSize=8,
                             textColor=colors.grey)
    red_s = ParagraphStyle('red', parent=styles['Normal'], fontSize=11,
                           textColor=colors.red, fontName='Helvetica-Bold')
    story = []
    if confidential:
        story.append(Paragraph("VERTRAULICH / CONFIDENTIAL", red_s))
        story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph(C["name"], h1_s))
    story.append(Paragraph(C["full_address"] + " | " + C["hrb"], small_s))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#1A3A6B")))
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph(title, h1_s))
    if subtitle:
        story.append(Paragraph(subtitle, ParagraphStyle('sub', parent=styles['Normal'],
                                                         fontSize=10, textColor=colors.grey,
                                                         alignment=TA_CENTER)))
    story.append(Spacer(1, 0.3 * cm))
    for sec_title, sec_body in sections:
        if sec_title:
            story.append(Paragraph(sec_title, h2_s))
        if isinstance(sec_body, str):
            for para in sec_body.split("\n\n"):
                if para.strip():
                    story.append(Paragraph(para.replace("\n", "<br/>"), body_s))
        elif isinstance(sec_body, list):
            t = Table(sec_body, hAlign='LEFT', repeatRows=1)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#1A3A6B")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor("#F0F4FA")]),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ]))
            story.append(t)
        story.append(Spacer(1, 0.3 * cm))
    pdoc.build(story)
    GENERATED_FILES.append(str(path))
    return path


def save_pdf(folder, filename, title, sections, **kwargs):
    return make_pdf_doc(folder, filename, title, sections, **kwargs)


# ─────────────────────────────────────────────
# XLSX HELPERS
# ─────────────────────────────────────────────
def make_xlsx_doc(folder, filename, title, sheets_data):
    filename = mess_filename(filename)
    path = BASE / folder / filename
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    hdr_fill = PatternFill("solid", fgColor="1A3A6B")
    hdr_font = Font(bold=True, color="FFFFFF", size=10)
    alt_fill = PatternFill("solid", fgColor="EFF4FB")
    title_font = Font(bold=True, size=12, color="1A3A6B")
    thin = Side(style='thin', color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for sheet_name, headers, rows, col_widths in sheets_data:
        ws = wb.create_sheet(sheet_name[:31])
        ws.append([title])
        ws["A1"].font = title_font
        ws.append([C["name"]])
        ws.append([])
        ws.append(headers)
        hr = 4
        for ci, h in enumerate(headers, 1):
            cell = ws.cell(row=hr, column=ci)
            cell.fill = hdr_fill
            cell.font = hdr_font
            cell.alignment = Alignment(horizontal='center', wrap_text=True)
            cell.border = border
        for ri, row in enumerate(rows, hr + 1):
            ws.append(row)
            fill = alt_fill if ri % 2 == 0 else None
            for ci in range(1, len(headers) + 1):
                cell = ws.cell(row=ri, column=ci)
                cell.border = border
                if fill:
                    cell.fill = fill
        for ci, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(ci)].width = w
    wb.save(path)
    GENERATED_FILES.append(str(path))
    return path


# ══════════════════════════════════════════════════════════════════════
# FOLDER 01 – GESELLSCHAFTSRECHT
# ══════════════════════════════════════════════════════════════════════
FOLDER_GES = "01_Gesellschaftsrecht"


def gen_gesellschaftsvertrag():
    doc = Document()
    make_docx_header(doc, "Gesellschaftsvertrag", "in der Fassung vom 15. März 2020", confidential=True)
    doc.add_paragraph(
        "Der vorliegende Gesellschaftsvertrag der Sentavia Precision GmbH wurde zuletzt durch Beschluss "
        "der Gesellschafterversammlung vom 15. März 2020 in der nachfolgenden Fassung neu gefasst. "
        "Er tritt an die Stelle aller vorherigen Fassungen des Gesellschaftsvertrags."
    )
    add_clause(doc, 1, "Firma und Sitz der Gesellschaft", [
        f"Die Gesellschaft führt die Firma {C['name']}.",
        f"Der Sitz der Gesellschaft ist {C['city']}.",
        f"Die Gesellschaft ist im Handelsregister des {C['amtsgericht']} unter der Nummer {C['hrb']} eingetragen.",
    ])
    add_clause(doc, 2, "Gegenstand des Unternehmens", [
        "Gegenstand des Unternehmens ist die Entwicklung, Herstellung, der Vertrieb und die Vermarktung von Medizinprodukten "
        "und In-vitro-Diagnostika, insbesondere von Systemen zur kardiologischen Diagnostik, orthopädischen Navigation "
        "und diagnostischen Schnelltests, sowie verwandter Software, Dienstleistungen und Zubehör.",
        "Die Gesellschaft ist berechtigt, alle Geschäfte zu betreiben, die geeignet sind, den Gesellschaftszweck "
        "unmittelbar oder mittelbar zu fördern, einschließlich des Erwerbs und der Verwaltung von Beteiligungen "
        "an anderen Unternehmen.",
        "Die Gesellschaft kann Zweigniederlassungen im In- und Ausland errichten sowie Tochtergesellschaften gründen "
        "oder sich an solchen beteiligen.",
    ])
    add_clause(doc, 3, "Stammkapital und Geschäftsanteile", [
        f"Das Stammkapital der Gesellschaft beträgt {eur(C['stammkapital'])} (in Worten: Fünfhunderttausend Euro).",
        "Das Stammkapital ist in Geschäftsanteile eingeteilt. Die Gesellschafter halten die folgenden Geschäftsanteile: "
        f"{C['investor1']}: {C['anteil1']} (laufende Nummern 1–3.500); "
        f"{C['investor2']}: {C['anteil2']} (laufende Nummern 3.501–6.000); "
        f"{C['investor3']}: {C['anteil3']} (laufende Nummern 6.001–7.500); "
        f"{C['investor4']}: {C['anteil4']} (laufende Nummern 7.501–8.500); "
        f"{C['investor5']}: {C['anteil5']} (laufende Nummern 8.501–10.000).",
        "Jeder Gesellschafter ist verpflichtet, seinen Geschäftsanteil vollständig einzuzahlen. Die Einlagen sind "
        "in bar oder durch Sacheinlage zu erbringen, sofern die Gesellschafterversammlung dies genehmigt hat.",
    ])
    add_clause(doc, 4, "Geschäftsführung", [
        "Die Gesellschaft hat einen oder mehrere Geschäftsführer. Die Bestellung und Abberufung der Geschäftsführer "
        "erfolgt durch Beschluss der Gesellschafterversammlung mit einfacher Mehrheit der abgegebenen Stimmen.",
        f"Zum alleinvertretungsberechtigten Geschäftsführer ist bestellt: {C['ceo']} (CEO). "
        f"Weitere Geschäftsführer sind: {C['cto']} (CTO), {C['cfo']} (CFO).",
        "Die Geschäftsführer sind von den Beschränkungen des § 181 BGB befreit, soweit die Gesellschafterversammlung "
        "dies mit Dreiviertelmehrheit beschlossen hat. Dies gilt insbesondere für die Vornahme von Rechtsgeschäften "
        "zwischen der Gesellschaft und Gesellschaften, an denen ein Geschäftsführer beteiligt ist.",
        "Für Rechtsgeschäfte, die im Rahmen des gewöhnlichen Geschäftsbetriebs liegen und den Betrag von "
        "EUR 500.000 nicht überschreiten, ist jeder Geschäftsführer alleinvertretungsberechtigt. "
        "Für darüber hinausgehende Geschäfte bedarf es der vorherigen Zustimmung der Gesellschafterversammlung.",
    ])
    add_clause(doc, 5, "Gesellschafterversammlung", [
        "Die ordentliche Gesellschafterversammlung findet einmal jährlich innerhalb der ersten sechs Monate "
        "des Geschäftsjahres statt. Außerordentliche Gesellschafterversammlungen sind einzuberufen, wenn es "
        "das Interesse der Gesellschaft erfordert oder wenn Gesellschafter, die zusammen mindestens 10 % des "
        "Stammkapitals halten, dies verlangen.",
        "Die Einberufung der Gesellschafterversammlung erfolgt durch die Geschäftsführung schriftlich oder per "
        "E-Mail mit einer Frist von mindestens 14 Tagen. Die Einberufung muss Ort, Zeit und Tagesordnung der "
        "Versammlung enthalten. In dringenden Fällen kann die Frist auf 7 Tage verkürzt werden.",
        "Beschlüsse der Gesellschafterversammlung werden mit einfacher Mehrheit der abgegebenen Stimmen gefasst, "
        "soweit Gesetz oder dieser Gesellschaftsvertrag keine andere Mehrheit vorschreiben. Je EUR 1,00 des "
        "Stammkapitals gewährt eine Stimme. Stimmenthaltungen gelten als nicht abgegebene Stimmen.",
        "Folgende Maßnahmen bedürfen der Zustimmung von mindestens 75 % der abgegebenen Stimmen: "
        "(a) Änderungen des Gesellschaftsvertrags; (b) Erhöhung oder Herabsetzung des Stammkapitals; "
        "(c) Auflösung der Gesellschaft; (d) Abschluss von Unternehmensverträgen; "
        "(e) Erwerb, Veräußerung oder Belastung von Grundstücken über EUR 2.000.000.",
    ])
    add_clause(doc, 6, "Beirat", [
        "Die Gesellschaft kann einen Beirat einrichten. Der Beirat besteht aus drei bis fünf Mitgliedern, "
        "die von der Gesellschafterversammlung bestellt werden.",
        "Der Beirat berät und unterstützt die Geschäftsführung in strategischen Fragen. Er hat keine "
        "Weisungsbefugnis gegenüber der Geschäftsführung, soweit nicht im Einzelfall eine solche Befugnis "
        "ausdrücklich vereinbart ist.",
        "Die Mitglieder des Beirats erhalten eine angemessene Vergütung sowie Aufwandsentschädigung, "
        "die von der Gesellschafterversammlung festgesetzt wird.",
    ])
    add_clause(doc, 7, "Gewinnverwendung", [
        "Der Jahresüberschuss, vermindert um einen etwaigen Verlustvortrag und um die in die gesetzliche "
        "Rücklage einzustellenden Beträge, steht zur Verfügung der Gesellschafterversammlung.",
        "Die Gesellschafterversammlung beschließt über die Verwendung des Bilanzgewinns, insbesondere "
        "über die Ausschüttung an die Gesellschafter. Ausschüttungen erfolgen im Verhältnis der "
        "Geschäftsanteile, sofern die Gesellschafterversammlung keine abweichende Regelung beschließt.",
        "Vorabausschüttungen im Laufe des Geschäftsjahres sind zulässig, wenn und soweit die "
        "Liquidität der Gesellschaft dies erlaubt und die Geschäftsführung dem zustimmt.",
    ])
    add_clause(doc, 8, "Informations- und Kontrollrechte", [
        "Jeder Gesellschafter hat das Recht, von der Geschäftsführung Auskunft über die Angelegenheiten "
        "der Gesellschaft zu verlangen und die Bücher und Schriften der Gesellschaft einzusehen.",
        "Die Geschäftsführung ist verpflichtet, den Gesellschaftern quartalsweise schriftlich über den "
        "Gang der Geschäfte zu berichten (Quartalsreport). Darüber hinaus sind Gesellschafter, die "
        "mindestens 10 % des Stammkapitals halten, berechtigt, jederzeit Einsicht in die Bücher zu nehmen.",
    ])
    add_clause(doc, 9, "Verfügungen über Geschäftsanteile", [
        "Die Übertragung von Geschäftsanteilen sowie die Begründung von Pfandrechten oder sonstigen "
        "Belastungen bedürfen der vorherigen Zustimmung der Gesellschafterversammlung. Die Zustimmung "
        "ist mit einer Mehrheit von 75 % der Stimmen aller Gesellschafter zu erteilen.",
        "Den Gesellschaftern steht bei beabsichtigter Veräußerung eines Geschäftsanteils ein Vorkaufsrecht "
        "im Verhältnis ihrer bestehenden Beteiligungen zu. Das Vorkaufsrecht ist innerhalb von 30 Tagen "
        "nach Mitteilung der beabsichtigten Veräußerung auszuüben.",
        "Drag-Along-Recht: Sofern Gesellschafter, die zusammen mindestens 75 % der Geschäftsanteile halten, "
        "die Gesamtheit der Geschäftsanteile an einen Dritten veräußern möchten, sind die übrigen Gesellschafter "
        "verpflichtet, ihre Geschäftsanteile zu gleichen Bedingungen auf den Dritten zu übertragen.",
        "Tag-Along-Recht: Wenn Gesellschafter, die zusammen mehr als 50 % der Geschäftsanteile halten, "
        "ihre Anteile an einen Dritten veräußern, haben die übrigen Gesellschafter das Recht, ihre Anteile "
        "zu gleichen Konditionen ebenfalls zu veräußern (Mitveräußerungsrecht).",
    ])
    add_clause(doc, 10, "Wettbewerbsverbot", [
        "Die Geschäftsführer und Gesellschafter, die mehr als 25 % der Geschäftsanteile halten, sind "
        "während der Dauer ihrer Stellung bzw. Beteiligung verpflichtet, der Gesellschaft keine Konkurrenz "
        "zu machen. Sie dürfen weder unmittelbar noch mittelbar Unternehmen betreiben oder sich an solchen "
        "beteiligen, die mit der Gesellschaft in Wettbewerb stehen.",
        "Das Wettbewerbsverbot gilt für die Dauer von 2 Jahren nach Beendigung der Tätigkeit als "
        "Geschäftsführer bzw. nach Veräußerung der Beteiligung und erstreckt sich auf das Bundesgebiet "
        "sowie alle Länder, in denen die Gesellschaft zum Zeitpunkt des Ausscheidens tätig ist.",
    ])
    add_clause(doc, 11, "Auflösung und Liquidation", [
        "Die Gesellschaft wird aufgelöst durch: (a) Beschluss der Gesellschafterversammlung mit einer "
        "Mehrheit von 75 % des Stammkapitals; (b) Eröffnung des Insolvenzverfahrens über das Vermögen "
        "der Gesellschaft; (c) rechtskräftige gerichtliche Entscheidung.",
        "Die Liquidation erfolgt durch die Geschäftsführer, sofern die Gesellschafterversammlung "
        "keine anderen Liquidatoren bestellt. Die Liquidatoren vertreten die Gesellschaft gemeinschaftlich.",
        "Nach Begleichung aller Verbindlichkeiten und nach Ablauf der Sperrjahre ist das verbleibende "
        "Gesellschaftsvermögen im Verhältnis der Geschäftsanteile an die Gesellschafter zu verteilen.",
    ])
    add_clause(doc, 12, "Salvatorische Klausel und Schlussbestimmungen", [
        "Sollten einzelne Bestimmungen dieses Gesellschaftsvertrags unwirksam oder undurchführbar sein "
        "oder werden oder sollte der Gesellschaftsvertrag eine Lücke enthalten, so berührt dies die "
        "Wirksamkeit der übrigen Bestimmungen nicht.",
        "Änderungen und Ergänzungen dieses Gesellschaftsvertrags bedürfen der notariellen Beurkundung "
        "und der Eintragung in das Handelsregister.",
        f"Gerichtsstand ist {C['city']}. Es gilt ausschließlich deutsches Recht.",
        f"Dieser Gesellschaftsvertrag wurde am 15. März 2020 notariell beurkundet durch Notar Dr. Wolfgang "
        f"Reichelt, Notariat München, Urkundenrolle Nr. 2020/456.",
    ])
    doc.add_heading("Unterschriften", level=3)
    doc.add_paragraph(f"\n\nMünchen, den 15. März 2020\n\n")
    doc.add_paragraph(f"_________________________\t\t_________________________")
    doc.add_paragraph(f"{C['investor1']}\t\t{C['investor2']}")
    doc.add_paragraph(f"\n_________________________\t\t_________________________")
    doc.add_paragraph(f"{C['investor3']}\t\t{C['investor4']}")
    save_docx(doc, FOLDER_GES, "GV_001_Gesellschaftsvertrag_2020_Fassung.docx")


def gen_handelsregisterauszug(year):
    d = rdate(year, 3, 4)
    sections = [
        ("Registerinformation", f"Amtsgericht: {C['amtsgericht']}\n"
         f"Registernummer: {C['hrb']}\n"
         f"Registerblatt: {C['hrb'].split()[1]}\n"
         f"Registerart: HRB (Handelsregister Abteilung B)\n"
         f"Datum des Auszugs: {ds(d.year, d.month, d.day)}"),
        ("Firmenbezeichnung", f"Firma: {C['name']}\n"
         f"Sitz: {C['city']}\n"
         f"Gegenstand: Entwicklung, Herstellung und Vertrieb von Medizinprodukten und In-vitro-Diagnostika."),
        ("Stammkapital", f"Stammkapital: {eur(C['stammkapital'])}"),
        ("Geschäftsführer", [
            ["Name", "Funktion", "Vertretungsbefugnis", "Geburtsdatum"],
            ["Dr. Katharina Berger", "Geschäftsführerin (CEO)", "Alleinvertretungsberechtigt", "12.03.1975"],
            ["Dr. Marcus Vogt", "Geschäftsführer (CTO)", "Alleinvertretungsberechtigt", "07.09.1978"],
            ["Thomas Müller", "Geschäftsführer (CFO)", "Alleinvertretungsberechtigt", "22.11.1980"],
        ]),
        ("Gesellschafter", [
            ["Gesellschafter", "Beteiligung", "Anzahl Anteile"],
            [C["investor1"], C["anteil1"], "3.500"],
            [C["investor2"], C["anteil2"], "2.500"],
            [C["investor3"], C["anteil3"], "1.500"],
            [C["investor4"], C["anteil4"], "1.000"],
            [C["investor5"], C["anteil5"], "1.500"],
        ]),
        ("Eintragungen", f"Gründungsdatum: {C['founded']}\n"
         f"Ursprüngliche Eintragung: {ds(2008, 11, 3)}\n"
         f"Letzte Änderung: {ds(d.year, d.month, d.day)}\n"
         "Prokuristen: Stefan Hoffmann (Einzelprokura), Maria Schneider (Einzelprokura)"),
        ("Hinweis", "Dieser Ausdruck aus dem elektronischen Handelsregister ist ein amtliches Dokument. "
         f"Er wurde am {ds(d.year, d.month, d.day)} abgerufen. Maßgeblich ist ausschließlich der Inhalt "
         "des Handelsregisters, das beim zuständigen Registergericht geführt wird."),
    ]
    save_pdf(FOLDER_GES, f"GR_00{year - 2019}_Handelsregisterauszug_{year}.pdf",
             f"Handelsregisterauszug {year}", sections)


def gen_gesellschafterliste():
    doc = Document()
    make_docx_header(doc, "Gesellschafterliste gemäß § 40 GmbHG",
                     f"Stand: {ds(2024, 1, 15)}", confidential=True)
    doc.add_paragraph(
        f"Hiermit wird die aktuelle Gesellschafterliste der {C['name']}, {C['hrb']}, "
        f"eingetragen beim {C['amtsgericht']}, gemäß § 40 GmbHG in der folgenden Fassung aufgestellt. "
        f"Die Liste ist unverzüglich beim Handelsregister einzureichen."
    )
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(["Gesellschafter", "Sitz/Wohnort", "Beteiligung (%)", "Betrag (EUR)", "Anteilsnummern"]):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    data = [
        (C["investor1"], "Paris, Frankreich", "35,00", "175.000", "1–3.500"),
        (C["investor2"], "Zürich, Schweiz", "25,00", "125.000", "3.501–6.000"),
        (C["investor3"], "München", "15,00", "75.000", "6.001–7.500"),
        (C["investor4"], "Bonn", "10,00", "50.000", "7.501–8.500"),
        (C["investor5"], "Diverse", "15,00", "75.000", "8.501–10.000"),
    ]
    for row_data in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
    doc.add_paragraph()
    doc.add_paragraph(
        f"München, den {ds(2024, 1, 15)}\n\n"
        f"Die vorstehende Gesellschafterliste wird von den Geschäftsführern als richtig und vollständig bestätigt.\n\n"
        f"_______________________\t\t_______________________\n"
        f"{C['ceo']}\t\t{C['cto']}\n"
        f"Geschäftsführerin\t\t\tGeschäftsführer"
    )
    save_docx(doc, FOLDER_GES, "GR_004_Gesellschafterliste_2024.docx")


def gen_gf_anstellungsvertrag(name, title, short):
    doc = Document()
    make_docx_header(doc, f"Geschäftsführer-Anstellungsvertrag – {name}",
                     f"zwischen {C['name']} und {name}", confidential=True)
    doc.add_paragraph("ANSTELLUNGSVERTRAG", style='Heading 2').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"zwischen\n\n{C['name']}, {C['full_address']}, {C['hrb']}\n"
                      f"– nachfolgend \"Gesellschaft\" –\n\nund\n\n{name}, geboren am "
                      f"{random.randint(1, 28)}. {MONTHS_DE[random.randint(1,12)]} "
                      f"{random.randint(1970, 1982)}, wohnhaft in München\n"
                      f"– nachfolgend \"Geschäftsführer\" –")
    add_clause(doc, 1, "Bestellung und Aufgaben", [
        f"Der Geschäftsführer wird mit Wirkung ab dem 1. Januar {random.randint(2018, 2022)} "
        f"als {title} der Gesellschaft bestellt. Die Bestellung erfolgt durch Gesellschafterbeschluss.",
        "Der Geschäftsführer ist für die Leitung, Organisation und strategische Ausrichtung des ihm "
        "zugewiesenen Geschäftsbereichs verantwortlich. Er handelt im Rahmen der gesetzlichen Vorschriften, "
        "des Gesellschaftsvertrags und der Beschlüsse der Gesellschafterversammlung.",
        "Der Geschäftsführer ist verpflichtet, die Gesellschaft mit der Sorgfalt eines ordentlichen "
        "Kaufmanns zu führen (§ 43 GmbHG) und die Interessen der Gesellschaft gegenüber Dritten zu wahren.",
    ])
    add_clause(doc, 2, "Vergütung", [
        f"Der Geschäftsführer erhält ein festes Jahresbruttogehalt von "
        f"{eur(random.choice([280_000, 320_000, 360_000, 300_000]))} (brutto), "
        f"zahlbar in zwölf gleichen Monatsraten jeweils zum letzten Werktag des Monats.",
        "Zusätzlich zum Festgehalt erhält der Geschäftsführer eine jährliche variable Vergütung "
        "(Bonus) von bis zu 30 % des Jahresbruttogehalts, die sich nach dem Erreichen individuell "
        "festzulegender KPIs und des Unternehmensziels richtet.",
        "Der Gesellschaft steht es frei, dem Geschäftsführer einen Dienstwagen der Mittelklasse "
        "(Wert bis zu EUR 80.000 brutto) zur Verfügung zu stellen, der auch privat genutzt werden darf. "
        "Der geldwerte Vorteil ist vom Geschäftsführer zu versteuern.",
        "Der Geschäftsführer nimmt an dem Long-Term Incentive Plan (LTIP) der Gesellschaft teil. "
        "Einzelheiten werden in einer gesonderten LTIP-Vereinbarung geregelt.",
    ])
    add_clause(doc, 3, "Arbeitszeit und Urlaub", [
        "Der Geschäftsführer ist verpflichtet, seine volle Arbeitskraft der Gesellschaft zu widmen. "
        "Er hat die für die ordnungsgemäße Erfüllung seiner Aufgaben erforderliche Arbeitszeit aufzuwenden "
        "und ist nicht an feste Arbeitszeiten gebunden.",
        "Der Geschäftsführer hat Anspruch auf 30 Arbeitstage Erholungsurlaub pro Kalenderjahr. "
        "Die zeitliche Lage des Urlaubs ist mit der Gesellschafterversammlung abzustimmen und darf die "
        "Geschäftstätigkeit der Gesellschaft nicht wesentlich beeinträchtigen.",
    ])
    add_clause(doc, 4, "Nebentätigkeiten", [
        "Entgeltliche oder zeitintensive Nebentätigkeiten bedürfen der vorherigen schriftlichen Zustimmung "
        "der Gesellschafterversammlung. Die Zustimmung kann ohne Angabe von Gründen verweigert werden, "
        "wenn die Nebentätigkeit die Interessen der Gesellschaft beeinträchtigen könnte.",
        "Unentgeltliche ehrenamtliche Tätigkeiten sowie Tätigkeiten in Berufsverbänden und gemeinnützigen "
        "Organisationen sind zustimmungsfrei, sofern sie nicht mehr als 5 Stunden pro Woche in Anspruch nehmen.",
    ])
    add_clause(doc, 5, "Verschwiegenheitspflicht und Datenschutz", [
        "Der Geschäftsführer ist verpflichtet, über alle Geschäfts- und Betriebsgeheimnisse sowie "
        "vertraulichen Informationen der Gesellschaft, die ihm in Ausübung seiner Tätigkeit bekannt werden, "
        "gegenüber Dritten Stillschweigen zu bewahren. Diese Pflicht gilt auch nach Beendigung des Vertrags.",
        "Als vertraulich gelten insbesondere: Kundendaten, Lieferantenbeziehungen, Produktentwicklungen, "
        "Geschäftsstrategien, Finanzdaten sowie Informationen über laufende Verhandlungen und Transaktionen.",
    ])
    add_clause(doc, 6, "Wettbewerbsverbot", [
        "Dem Geschäftsführer ist es während der Dauer dieses Vertrags untersagt, unmittelbar oder mittelbar "
        "für ein mit der Gesellschaft konkurrierendes Unternehmen tätig zu werden oder ein solches zu gründen "
        "oder zu erwerben.",
        "Nach Beendigung des Vertragsverhältnisses gilt ein nachvertragliches Wettbewerbsverbot von "
        "24 Monaten im Bereich der Entwicklung und des Vertriebs von Medizinprodukten der Klassen IIa, "
        "IIb und III sowie von In-vitro-Diagnostika für den europäischen Wirtschaftsraum. "
        "Als Karenzentschädigung erhält der Geschäftsführer 50 % seiner zuletzt bezogenen Jahresvergütung.",
    ])
    add_clause(doc, 7, "Haftung", [
        "Der Geschäftsführer haftet der Gesellschaft für Schäden, die er durch die Verletzung seiner "
        "Pflichten verursacht hat, gemäß § 43 GmbHG. Im Innenverhältnis haftet er nur für Vorsatz und "
        "grobe Fahrlässigkeit, soweit die Gesellschaft keinen Anspruch auf Schadensersatz aus einer "
        "D&O-Versicherung geltend machen kann.",
        "Die Gesellschaft schließt für den Geschäftsführer eine D&O-Versicherung (Directors & Officers) "
        "mit einem Selbstbehalt gemäß § 93 Abs. 2 AktG in entsprechender Anwendung ab.",
    ])
    add_clause(doc, 8, "Vertragsdauer und Kündigung", [
        "Dieser Anstellungsvertrag wird auf unbestimmte Zeit geschlossen. Er kann von beiden Seiten mit "
        "einer Frist von sechs Monaten zum Monatsende gekündigt werden.",
        "Das Recht zur außerordentlichen Kündigung aus wichtigem Grund bleibt unberührt. Als wichtiger "
        "Grund gilt insbesondere: grobe Pflichtverletzungen, strafbare Handlungen zu Lasten der Gesellschaft, "
        "dauernde Arbeitsunfähigkeit von mehr als 12 Monaten.",
        "Die Abberufung als Geschäftsführer durch die Gesellschafterversammlung gilt gleichzeitig als "
        "Kündigung dieses Anstellungsvertrags zum nächstmöglichen Termin, es sei denn, die Kündigung "
        "widerspräche zwingendem Recht.",
    ])
    add_clause(doc, 9, "Schlussbestimmungen", [
        "Änderungen und Ergänzungen dieses Vertrags bedürfen der Schriftform. Dies gilt auch für die "
        "Abbedingung dieses Schriftformerfordernisses.",
        f"Es gilt das Recht der Bundesrepublik Deutschland. Gerichtsstand ist {C['city']}.",
        "Sollten einzelne Bestimmungen dieses Vertrags unwirksam oder undurchführbar sein, so berührt "
        "dies die Wirksamkeit der übrigen Bestimmungen nicht.",
    ])
    doc.add_paragraph(f"\n\nMünchen, den {ds(2022, 1, 1)}\n\n")
    doc.add_paragraph("_______________________\t\t_______________________")
    doc.add_paragraph(f"{C['name']}\t\t\t{name}")
    doc.add_paragraph("(Gesellschafterin)\t\t\t(Geschäftsführer)")
    save_docx(doc, FOLDER_GES, f"GF_{short}_Anstellungsvertrag_{sfn(name.replace(' ','_'))}.docx")


def gen_gesellschafterbeschluss(year):
    doc = Document()
    d = rdate(year, 4, 6)
    make_docx_header(doc, f"Protokoll der ordentlichen Gesellschafterversammlung {year}",
                     f"Datum: {ds(d.year, d.month, d.day)}")
    doc.add_paragraph(
        f"Ort: {C['full_address']}, Konferenzraum \"Innovation\"\n"
        f"Datum: {ds(d.year, d.month, d.day)}, 10:00 Uhr\n"
        f"Protokollführer: {C['cfo']}"
    )
    doc.add_heading("Anwesende", level=3)
    doc.add_paragraph(
        f"• {C['investor1']} (vertreten durch Dr. Antoine Leblanc), {C['anteil1']} der Anteile\n"
        f"• {C['investor2']} (vertreten durch Henrik Brüggemann), {C['anteil2']} der Anteile\n"
        f"• {C['investor3']} (vertreten durch {C['ceo']}), {C['anteil3']} der Anteile\n"
        f"• {C['investor4']} (vertreten durch Dr. Franz Schmidt), {C['anteil4']} der Anteile\n"
        f"• Geschäftsführung: {C['ceo']}, {C['cto']}, {C['cfo']}\n"
        f"• Protokollführer: {C['cfo']}\n"
        f"• Wirtschaftsprüfer: Vertreter der {C['wp']}"
    )
    doc.add_heading("Feststellung der ordnungsgemäßen Einberufung", level=3)
    doc.add_paragraph(
        "Die Versammlungsleitung stellt fest, dass die Gesellschafterversammlung ordnungsgemäß und "
        f"fristgerecht mit Schreiben vom {ds(d.year, d.month, d.day - 14)} einberufen wurde. "
        "Das Stammkapital der Gesellschaft ist vollständig vertreten. Die Gesellschafterversammlung "
        "ist damit beschlussfähig."
    )
    doc.add_heading("Tagesordnung", level=3)
    doc.add_paragraph(
        f"TOP 1: Vorlage und Genehmigung des Jahresabschlusses {year - 1}\n"
        f"TOP 2: Verwendung des Bilanzgewinns {year - 1}\n"
        f"TOP 3: Entlastung der Geschäftsführung\n"
        f"TOP 4: Bericht der Geschäftsführung über die Geschäftsentwicklung {year}\n"
        f"TOP 5: Genehmigung des Wirtschaftsplans {year}\n"
        f"TOP 6: Verschiedenes"
    )
    rev = C[f"revenue_{year-1}"] if f"revenue_{year-1}" in C else 87_200_000
    ebitda = C[f"ebitda_{year-1}"] if f"ebitda_{year-1}" in C else 12_400_000
    doc.add_heading("Beschlüsse", level=3)
    doc.add_paragraph(
        f"Zu TOP 1: Die Gesellschafterversammlung genehmigt einstimmig den geprüften Jahresabschluss "
        f"(Bilanz, Gewinn- und Verlustrechnung, Anhang) für das Geschäftsjahr {year - 1}. "
        f"Der Jahresüberschuss beläuft sich auf {eur(int(ebitda * 0.72))}. "
        f"Der Umsatz betrug {eur(rev)}. Der Jahresabschluss wurde von {C['wp']} geprüft und "
        f"mit einem uneingeschränkten Bestätigungsvermerk versehen."
    )
    doc.add_paragraph(
        "Zu TOP 2: Die Gesellschafterversammlung beschließt einstimmig, den Bilanzgewinn vollständig "
        "in die freien Rücklagen einzustellen. Eine Ausschüttung an die Gesellschafter erfolgt im "
        f"Geschäftsjahr {year - 1} nicht."
    )
    doc.add_paragraph(
        "Zu TOP 3: Die Gesellschafterversammlung erteilt der Geschäftsführung einstimmig Entlastung "
        f"für das Geschäftsjahr {year - 1}."
    )
    doc.add_paragraph(
        f"Zu TOP 4: {C['ceo']} berichtet über die Geschäftsentwicklung des laufenden Jahres {year}. "
        "Die Gesellschaft liegt im Plan. Besondere Highlights sind: Markteinführung der erweiterten "
        f"{C['prod1_name']}-Plattform in drei neuen EU-Märkten, Abschluss der PMCF-Studie für "
        f"{C['prod2_name']} mit positivem Ergebnis sowie Aufnahme von Gesprächen mit einem "
        "potenziellen strategischen Partner aus den USA."
    )
    doc.add_paragraph(
        f"Zu TOP 5: Die Gesellschafterversammlung genehmigt einstimmig den Wirtschaftsplan {year}. "
        f"Geplanter Umsatz: {eur(int(rev * 1.18))}. Geplantes EBITDA: {eur(int(ebitda * 1.15))}. "
        f"Investitionsvolumen: {eur(int(rev * 0.07))}."
    )
    doc.add_paragraph(
        f"\nMünchen, den {ds(d.year, d.month, d.day)}\n\n"
        "________________________________\n"
        f"{C['cfo']}, Protokollführer"
    )
    save_docx(doc, FOLDER_GES, f"GV_{year}_Protokoll_Gesellschafterversammlung.docx")


def gen_prokura():
    doc = Document()
    make_docx_header(doc, "Erteilung von Prokura")
    doc.add_paragraph(
        f"Die Gesellschafterversammlung der {C['name']}, {C['hrb']}, hat in der Sitzung vom "
        f"{ds(2022, 3, 15)} beschlossen, folgenden Personen Prokura zu erteilen:"
    )
    add_section(doc, "1. Stefan Hoffmann (Einzelprokura)", [
        f"Dipl.-Ing. Stefan Hoffmann, geboren am 14. April 1979, wohnhaft in München, "
        f"wird mit sofortiger Wirkung Einzelprokura für die {C['name']} erteilt.",
        "Die Prokura berechtigt zur Vertretung der Gesellschaft in allen Geschäften, die der Betrieb "
        "eines Handelsgewerbes mit sich bringt, mit Ausnahme der Veräußerung und Belastung von "
        "Grundstücken (§ 49 Abs. 2 HGB).",
        "Herr Hoffmann ist als Head of Quality & Regulatory Affairs tätig und zeichnet verantwortlich "
        "für alle regulatorischen Eingaben, CE-Kennzeichnungsverfahren und Qualitätsmanagementsystem-Audits.",
    ])
    add_section(doc, "2. Maria Schneider (Einzelprokura)", [
        "Maria Schneider, M.Sc., geboren am 3. September 1983, wohnhaft in München, "
        f"wird mit sofortiger Wirkung Einzelprokura für die {C['name']} erteilt.",
        "Die Prokura berechtigt zur Vertretung der Gesellschaft in allen kaufmännischen und "
        "rechtlichen Angelegenheiten des operativen Geschäftsbetriebs, insbesondere im Bereich "
        "Vertrieb, Einkauf und HR-Management.",
    ])
    add_section(doc, "Handelsregistereintragung", [
        "Die Prokuraerteilung ist unverzüglich zum Handelsregister anzumelden. "
        f"Die Geschäftsführung ({C['ceo']}) ist beauftragt, die notwendigen Anmeldungen "
        "beim Registergericht vorzunehmen.",
    ])
    doc.add_paragraph(
        f"\nMünchen, den {ds(2022, 3, 15)}\n\n"
        f"_______________________\n{C['ceo']}\nGeschäftsführerin"
    )
    save_docx(doc, FOLDER_GES, "Prokura_Erteilung_Hoffmann_Schneider.docx")


def gen_sha():
    """Shareholders Agreement"""
    doc = Document()
    make_docx_header(doc, "Shareholders' Agreement", "Sentavia Precision GmbH", confidential=True)
    doc.add_paragraph(
        "This Shareholders' Agreement (\"Agreement\") is entered into as of January 10, 2020, "
        f"by and among the shareholders of {C['name']} (\"Company\"), a limited liability company "
        f"(Gesellschaft mit beschränkter Haftung) incorporated under the laws of Germany, registered "
        f"with the {C['amtsgericht']} under registration number {C['hrb']}, with registered office "
        f"at {C['full_address']}.\n\n"
        "The parties to this Agreement are:\n"
        f"(1) {C['investor1']}, a société en commandite par actions incorporated under French law (\"Sofinnova\");\n"
        f"(2) {C['investor2']}, a private limited company incorporated under Swiss law (\"Lakestar\");\n"
        f"(3) {C['investor3']}, a civil law partnership (Gesellschaft bürgerlichen Rechts) under German law "
        f"(\"Management GbR\");\n"
        f"(4) {C['investor4']}, a public limited fund under German law (\"HTGF\").\n\n"
        "WHEREAS, the parties hereto are shareholders of the Company and desire to set forth certain "
        "agreements regarding the governance, management, and transfer of shares in the Company."
    )
    for clause_num, clause_title, clause_texts in [
        (1, "Definitions", [
            "\"Shares\" means the shares (Geschäftsanteile) in the Company as registered in the commercial register.",
            "\"Exit Event\" means an IPO, a Trade Sale, or a Secondary Sale involving more than 50% of the Shares.",
            "\"Investor Director\" means a member of the Advisory Board nominated by an Investor pursuant to Section 4.",
        ]),
        (2, "Business Plan and Budget", [
            "The Company shall operate in accordance with an annual business plan and budget approved by the Shareholders' Meeting. "
            "The Geschäftsführung shall prepare and submit a proposed business plan and budget to the Shareholders' Meeting "
            "for approval no later than 30 November of each year.",
            "Material deviations from the approved business plan in excess of 15% of planned revenues or EBITDA shall "
            "require prior approval of a Shareholders' Meeting.",
        ]),
        (3, "Reporting", [
            "The Company shall provide each Investor holding more than 10% of the Shares with: (i) monthly management "
            "accounts within 20 business days of month-end; (ii) quarterly investor reports within 30 days of quarter-end; "
            "(iii) audited annual financial statements within 90 days of year-end.",
            "The Geschäftsführung shall hold quarterly investor calls with all Investors at which the performance of the "
            "Company against its business plan shall be discussed.",
        ]),
        (4, "Advisory Board", [
            "An Advisory Board shall be established comprising up to five members. Sofinnova shall be entitled to nominate "
            "two members; Lakestar shall be entitled to nominate one member; the Management GbR shall be entitled to "
            "nominate one member; HTGF shall be entitled to nominate one member.",
            "The Advisory Board shall meet at least quarterly and shall advise the Geschäftsführung on strategic matters. "
            "The Advisory Board shall not have any binding authority over the Geschäftsführung except as expressly "
            "set forth in this Agreement.",
        ]),
        (5, "Pre-emption Rights", [
            "If any Shareholder (\"Selling Shareholder\") wishes to transfer any of its Shares to a third party, it shall "
            "first offer such Shares to the other Shareholders pro rata to their respective shareholdings at the same price "
            "and on the same terms as those offered by the third party.",
            "The pre-emption right must be exercised within 30 business days of receipt of the transfer notice. "
            "If the pre-emption right is not exercised in full, the Selling Shareholder may transfer the Shares to "
            "the third party on terms no more favorable than those offered to the other Shareholders.",
        ]),
        (6, "Tag-Along and Drag-Along", [
            "Tag-Along: If any Shareholder(s) holding more than 50% of the Shares propose to sell their Shares to a "
            "third party, each other Shareholder shall have the right to sell its pro rata share of such Shares to such "
            "third party at the same price and on the same terms.",
            "Drag-Along: If Shareholders holding at least 75% of the Shares approve a sale of all Shares to a third "
            "party, the remaining Shareholders shall be required to sell their Shares on the same terms.",
        ]),
        (7, "Anti-Dilution", [
            "In case of a down-round financing, Investors holding Preferred Shares shall benefit from broad-based "
            "weighted-average anti-dilution protection as set out in Schedule 1 to this Agreement.",
        ]),
        (8, "Liquidation Preference", [
            "Upon any liquidation, dissolution, or winding up of the Company, or upon an Exit Event, the proceeds "
            "shall be distributed as follows: first, to Preferred Shareholders in accordance with their respective "
            "liquidation preferences (1x non-participating); second, to all Shareholders pro rata.",
        ]),
        (9, "Governing Law and Dispute Resolution", [
            "This Agreement shall be governed by and construed in accordance with the laws of Germany.",
            "Any dispute arising out of or in connection with this Agreement shall be finally settled by arbitration "
            "under the Rules of the German Institution of Arbitration (DIS) by three arbitrators. The seat of "
            f"arbitration shall be {C['city']}. The language of the arbitration shall be English or German.",
        ]),
    ]:
        add_clause(doc, clause_num, clause_title, clause_texts)
    doc.add_paragraph(
        f"\nIN WITNESS WHEREOF, the parties have executed this Agreement as of January 10, 2020.\n\n"
        f"_______________________\t\t_______________________\n"
        f"{C['investor1']}\t\t{C['investor2']}\n\n"
        f"_______________________\t\t_______________________\n"
        f"{C['investor3']}\t\t{C['investor4']}"
    )
    save_docx(doc, FOLDER_GES, "SHA_001_Shareholders_Agreement_2020.docx")


def gen_cap_table():
    sheets = [
        ("Cap Table", ["Gesellschafter", "Anteil %", "Anteile", "Nominalwert EUR",
                       "Investment EUR", "Implied Value EUR"],
         [
             (C["investor1"], "35,00%", "3.500", "175.000", "12.250.000", "36.400.000"),
             (C["investor2"], "25,00%", "2.500", "125.000", "8.750.000", "26.000.000"),
             (C["investor3"], "15,00%", "1.500", "75.000", "1.500.000 (Mgmt)", "15.600.000"),
             (C["investor4"], "10,00%", "1.000", "50.000", "2.000.000", "10.400.000"),
             (C["investor5"], "15,00%", "1.500", "75.000", "Diverse", "15.600.000"),
             ("GESAMT", "100,00%", "10.000", "500.000", "~24.500.000", "104.000.000"),
         ], [25, 12, 10, 18, 20, 20]),
        ("Funding History", ["Runde", "Datum", "Bewertung (pre)", "Betrag", "Lead Investor", "Kumulativ"],
         [
             ("Seed", "Okt 2008", "n/a", "500.000", "Gründer", "500.000"),
             ("Series A", "März 2012", "4.000.000", "3.500.000", C["investor4"], "4.000.000"),
             ("Series B", "Jul 2016", "18.000.000", "9.000.000", C["investor2"], "13.000.000"),
             ("Series C", "Jan 2020", "62.000.000", "15.000.000", C["investor1"], "28.000.000"),
         ], [12, 14, 20, 14, 28, 16]),
    ]
    make_xlsx_doc(FOLDER_GES, "SHA_002_Cap_Table_2024.xlsx", "Kapitaltabelle / Cap Table", sheets)


def gen_investor_rights():
    doc = Document()
    make_docx_header(doc, "Investor Rights Agreement", "Sentavia Precision GmbH", confidential=True)
    add_clause(doc, 1, "Information Rights", [
        f"Each Investor holding at least 5% of the outstanding shares of {C['name']} (each, a \"Major Investor\") "
        "shall have the right to receive from the Company: (a) as soon as practicable after the end of each "
        "fiscal year, and in any event within ninety (90) days thereof, an audited balance sheet and statements "
        "of income and cash flows for such fiscal year; (b) as soon as practicable after the end of each "
        "calendar quarter, but in any event within thirty (30) days thereafter, an unaudited income statement "
        "and balance sheet as of the end of such quarter.",
        "The Company shall also provide each Major Investor with an annual budget no later than 30 days "
        "before the beginning of each fiscal year, and with any material revisions to such budget within "
        "5 business days of adoption.",
    ])
    add_clause(doc, 2, "Registration Rights", [
        "In the event of an IPO of the Company, each Major Investor shall have the right to participate "
        "in such offering with respect to a pro rata share of the shares to be offered.",
        "Demand Registration Rights: At any time following the date that is 180 days after the effective date "
        "of the IPO, holders of at least 30% of the Registrable Securities may request that the Company "
        "file a registration statement covering the resale of their Registrable Securities.",
    ])
    add_clause(doc, 3, "Right of First Refusal and Co-Sale Agreement", [
        "The Company and each of the Key Holders grants to the Investors a right of first refusal with "
        "respect to any proposed transfer of Company securities held by the Key Holders, other than "
        "Exempt Transfers.",
        "Co-Sale Right: If any Key Holder proposes to transfer Company securities to a third party and "
        "the right of first refusal described in Section 3.1 is not exercised in full, each Investor "
        "shall have the right to participate in such transfer.",
    ])
    add_clause(doc, 4, "Protective Provisions", [
        "The Company shall not, without the approval of at least 67% of the then outstanding Preferred Shares: "
        "(a) alter or change the rights, preferences, or privileges of the Preferred Shares; "
        "(b) create any new class or series of shares having rights, preferences, or privileges senior "
        "to or on a parity with the Preferred Shares; (c) authorize or effect any merger, acquisition, "
        "or change of control transaction.",
    ])
    save_docx(doc, FOLDER_GES, "IRA_001_Investor_Rights_Agreement_2020.docx")


def gen_board_resolution(year, num, topic):
    doc = Document()
    d = rdate(year)
    make_docx_header(doc, f"Gesellschafterbeschluss im Umlaufverfahren – {topic}",
                     f"Datum: {ds(d.year, d.month, d.day)}")
    doc.add_paragraph(
        f"Die Gesellschafter der {C['name']}, {C['hrb']}, haben im schriftlichen Umlaufverfahren "
        f"gemäß § 48 Abs. 2 GmbHG am {ds(d.year, d.month, d.day)} folgenden Beschluss gefasst:\n\n"
        "Sämtliche Gesellschafter, die gemeinsam 100 % der Stimmrechte auf sich vereinigen, "
        "haben dem nachfolgenden Beschluss zugestimmt. Das erforderliche Quorum ist erreicht."
    )
    doc.add_heading("Beschluss", level=3)
    doc.add_paragraph(topic + "\n\nDer Beschluss wird einstimmig angenommen.\n\n"
                      "Abstimmungsergebnis: JA: 10.000 Stimmen / NEIN: 0 Stimmen / ENTHALTUNG: 0 Stimmen")
    doc.add_paragraph(
        f"\nMünchen, den {ds(d.year, d.month, d.day)}\n\n"
        f"_______________________\n{C['ceo']}\nGeschäftsführerin, zugleich Vertreterin der {C['investor3']}"
    )
    save_docx(doc, FOLDER_GES, f"BR_{year}_{num:02d}_Beschluss_{sfn(topic[:25])}.docx")


def generate_01_gesellschaftsrecht():
    print("  Generating 01_Gesellschaftsrecht...")
    gen_gesellschaftsvertrag()
    for yr in [2022, 2023, 2024]:
        gen_handelsregisterauszug(yr)
    gen_gesellschafterliste()
    gen_gf_anstellungsvertrag(C["ceo"], "Geschäftsführerin (CEO)", "001")
    gen_gf_anstellungsvertrag(C["cto"], "Geschäftsführer (CTO)", "002")
    gen_gf_anstellungsvertrag(C["cfo"], "Geschäftsführer (CFO)", "003")
    for yr in [2020, 2021, 2022, 2023]:
        gen_gesellschafterbeschluss(yr)
    gen_prokura()
    gen_sha()
    gen_cap_table()
    gen_investor_rights()
    resolutions = [
        (2021, 1, "Genehmigung der Kapitalkehöhung um EUR 50.000 durch Ausgabe neuer Geschäftsanteile"),
        (2021, 2, "Bestellung von Dr. Marcus Vogt zum Geschäftsführer (CTO) mit sofortiger Wirkung"),
        (2022, 1, "Genehmigung des Abschlusses eines Kreditvertrags mit der UniCredit Bank AG bis EUR 15.000.000"),
        (2022, 2, "Genehmigung der Eröffnung einer Zweigniederlassung in Wien, Österreich"),
        (2022, 3, "Zustimmung zum Abschluss des Distributionsvertrags mit Siemens Healthineers AG"),
        (2023, 1, "Genehmigung des Erwerbs der Minderheitsbeteiligung an MedSoft Analytics GmbH"),
        (2023, 2, "Bestellung von Dr. Annika Schmidt zur Chief Medical Officer"),
        (2023, 3, "Genehmigung eines Aktienoptionsplans (ESOP) für Schlüsselmitarbeiter"),
        (2023, 4, "Genehmigung des Jahresabschlusses 2022 und Entlastung der Geschäftsführung"),
        (2024, 1, "Genehmigung der Aufnahme von Gesprächen mit potenziellen strategischen Investoren"),
        (2024, 2, "Beauftragung der Noerr PartGmbB mit rechtlicher Due-Diligence-Begleitung"),
        (2024, 3, "Genehmigung der Anpassung der Vergütungsstruktur der Geschäftsführung"),
    ]
    for yr, n, topic in resolutions:
        gen_board_resolution(yr, n, topic)
    # Notarielle Urkunden
    for i, topic in enumerate(["Gründungsurkunde 2008", "Kapitalerhöhung 2012", "Satzungsänderung 2016",
                                "Anteilsübertragung 2020 Series C", "Prokuraeintragung 2022"], 1):
        sections = [
            ("Urkundentext", f"Vor mir, dem unterzeichnenden Notar Dr. Wolfgang Reichelt, Notariat München, "
             f"erschienen heute am {ds(random.randint(2008, 2022), random.randint(1,12), random.randint(1,28))} "
             f"die nachfolgend genannten Personen und erklärten Folgendes:\n\n"
             f"Gegenstand dieser Urkunde: {topic}\n\nDie vorstehenden Erklärungen wurden den Erschienenen "
             "vorgelesen, von ihnen genehmigt und eigenhändig unterschrieben."),
            ("Beglaubigungsvermerk", f"Notariell beurkundet. Urkundenrolle Nr. 2020/{1000+i}.\n"
             f"Dr. Wolfgang Reichelt, Notar, München"),
        ]
        save_pdf(FOLDER_GES, f"NOT_{i:03d}_Urkunde_{sfn(topic[:25])}.pdf",
                 f"Notarielle Urkunde – {topic}", sections)


# ══════════════════════════════════════════════════════════════════════
# FOLDER 02 – FINANZEN
# ══════════════════════════════════════════════════════════════════════
FOLDER_FIN = "02_Finanzen"


def gen_jahresabschluss(year):
    rev = C.get(f"revenue_{year}", 87_200_000)
    ebitda = C.get(f"ebitda_{year}", 12_400_000)
    ebit = C.get(f"ebit_{year}", 9_200_000)
    emp = C.get(f"employees_{year}", 612)
    cogs = int(rev * 0.42)
    gross = rev - cogs
    rd = int(rev * 0.13)
    sga = int(rev * 0.30)
    depr = ebitda - ebit
    interest = int(rev * 0.008)
    ebt = ebit - interest
    tax = int(ebt * 0.297)
    net = ebt - tax
    ta = int(rev * 1.12)
    equity = int(ta * 0.42)
    debt = int(ta * 0.28)
    current = int(ta * 0.30)
    sheets = [
        ("GuV", ["Position", f"GJ {year} EUR", f"GJ {year-1} EUR", "Veränderung %"],
         [
             ("1. Umsatzerlöse", f"{rev:,}", f"{int(rev*0.85):,}", "+17,6%"),
             ("2. Herstellungskosten", f"-{cogs:,}", f"-{int(cogs*0.85):,}", "+17,6%"),
             ("3. Bruttoergebnis", f"{gross:,}", f"{int(gross*0.85):,}", "+17,7%"),
             ("4. F&E-Aufwendungen", f"-{rd:,}", f"-{int(rd*0.85):,}", "+17,6%"),
             ("5. Vertrieb & Marketing", f"-{int(sga*0.55):,}", f"-{int(sga*0.55*0.85):,}", "+17,6%"),
             ("6. Allg. Verwaltung", f"-{int(sga*0.45):,}", f"-{int(sga*0.45*0.85):,}", "+17,6%"),
             ("7. EBITDA", f"{ebitda:,}", f"{int(ebitda*0.85):,}", "+17,6%"),
             ("8. Abschreibungen", f"-{depr:,}", f"-{int(depr*0.85):,}", "+17,6%"),
             ("9. EBIT", f"{ebit:,}", f"{int(ebit*0.85):,}", "+17,6%"),
             ("10. Zinsergebnis", f"-{interest:,}", f"-{int(interest*0.85):,}", "+17,6%"),
             ("11. EBT", f"{ebt:,}", f"{int(ebt*0.85):,}", "+17,6%"),
             ("12. Steuern", f"-{tax:,}", f"-{int(tax*0.85):,}", "+17,6%"),
             ("13. Jahresüberschuss", f"{net:,}", f"{int(net*0.85):,}", "+17,6%"),
         ], [35, 18, 18, 15]),
        ("Bilanz", ["Position", f"{year}-12-31 EUR", f"{year-1}-12-31 EUR"],
         [
             ("AKTIVA", "", ""),
             ("Anlagevermögen gesamt", f"{int(ta*0.70):,}", f"{int(ta*0.70*0.85):,}"),
             ("  Immaterielle Vermögensgegenstände", f"{int(ta*0.18):,}", f"{int(ta*0.18*0.85):,}"),
             ("  Sachanlagen", f"{int(ta*0.35):,}", f"{int(ta*0.35*0.85):,}"),
             ("  Finanzanlagen", f"{int(ta*0.17):,}", f"{int(ta*0.17*0.85):,}"),
             ("Umlaufvermögen gesamt", f"{current:,}", f"{int(current*0.85):,}"),
             ("  Vorräte", f"{int(current*0.28):,}", f"{int(current*0.28*0.85):,}"),
             ("  Forderungen LuL", f"{int(current*0.38):,}", f"{int(current*0.38*0.85):,}"),
             ("  Liquide Mittel", f"{int(current*0.34):,}", f"{int(current*0.34*0.85):,}"),
             ("Bilanzsumme", f"{ta:,}", f"{int(ta*0.85):,}"),
             ("PASSIVA", "", ""),
             ("Eigenkapital gesamt", f"{equity:,}", f"{int(equity*0.85):,}"),
             ("  Stammkapital", "500.000", "500.000"),
             ("  Kapitalrücklage", f"{int(equity*0.52):,}", f"{int(equity*0.52):,}"),
             ("  Gewinnrücklagen", f"{int(equity*0.28):,}", f"{int(equity*0.28*0.85):,}"),
             ("  Jahresüberschuss", f"{net:,}", f"{int(net*0.85):,}"),
             ("Fremdkapital gesamt", f"{ta-equity:,}", f"{int((ta-equity)*0.85):,}"),
             ("  Langfristige Verbindlichkeiten", f"{debt:,}", f"{int(debt*0.85):,}"),
             ("  Kurzfristige Verbindlichkeiten", f"{ta-equity-debt:,}", f"{int((ta-equity-debt)*0.85):,}"),
         ], [35, 18, 18]),
        ("Kennzahlen", ["Kennzahl", f"GJ {year}", f"GJ {year-1}"],
         [
             ("Umsatz (EUR)", f"{rev:,}", f"{int(rev*0.85):,}"),
             ("Umsatzwachstum", "+17,6%", "+22,9%"),
             ("EBITDA (EUR)", f"{ebitda:,}", f"{int(ebitda*0.85):,}"),
             ("EBITDA-Marge", f"{ebitda/rev*100:.1f}%", f"{ebitda/rev*100:.1f}%"),
             ("EBIT (EUR)", f"{ebit:,}", f"{int(ebit*0.85):,}"),
             ("EBIT-Marge", f"{ebit/rev*100:.1f}%", f"{ebit/rev*100:.1f}%"),
             ("Jahresüberschuss (EUR)", f"{net:,}", f"{int(net*0.85):,}"),
             ("Mitarbeiterzahl", str(emp), str(int(emp*0.85))),
             ("Umsatz je Mitarbeiter (EUR)", f"{int(rev/emp):,}", f"{int(rev*0.85/int(emp*0.85)):,}"),
             ("Eigenkapitalquote", f"{equity/ta*100:.1f}%", f"{equity/ta*100:.1f}%"),
         ], [35, 18, 18]),
    ]
    make_xlsx_doc(FOLDER_FIN, f"JA_{year}_Jahresabschluss.xlsx",
                  f"Jahresabschluss {year} – {C['name']}", sheets)


def gen_bwa_monatlich(year):
    months = ["Jan", "Feb", "Mär", "Apr", "Mai", "Jun",
              "Jul", "Aug", "Sep", "Okt", "Nov", "Dez"]
    rev = C.get(f"revenue_{year}", 71_400_000)
    base_monthly = rev / 12
    rows = []
    for i, m in enumerate(months):
        factor = 1 + random.uniform(-0.08, 0.12) + (0.05 if i >= 9 else 0)
        m_rev = int(base_monthly * factor)
        m_cogs = int(m_rev * 0.42)
        m_ebitda = int(m_rev * 0.142)
        m_ebit = int(m_rev * 0.10)
        m_net = int(m_ebit * 0.703)
        rows.append((m, f"{m_rev:,}", f"{m_cogs:,}", f"{m_rev-m_cogs:,}",
                     f"{m_ebitda:,}", f"{m_ebit:,}", f"{m_net:,}"))
    sheets = [
        ("BWA Monatsübersicht", ["Monat", "Umsatz EUR", "COGS EUR", "Bruttoergebnis EUR",
                                  "EBITDA EUR", "EBIT EUR", "Ergebnis EUR"],
         rows, [10, 15, 15, 18, 15, 15, 15]),
    ]
    make_xlsx_doc(FOLDER_FIN, f"BWA_{year}_Monatsuebersicht.xlsx",
                  f"Betriebswirtschaftliche Auswertung {year} – Monatsübersicht", sheets)


def gen_steuerbescheid(year):
    d = rdate(year + 1, 7, 9)
    ebt = int(C.get(f"ebit_{year}", 9_200_000) * 0.95)
    koerperschaftsteuer = int(ebt * 0.15)
    solidaritaetszuschlag = int(koerperschaftsteuer * 0.055)
    gewerbesteuer = int(ebt * 0.1435)
    gesamt = koerperschaftsteuer + solidaritaetszuschlag + gewerbesteuer
    sections = [
        ("Steuerbescheid – Körperschaftsteuer und Gewerbesteuer", None),
        ("Bescheid", f"Steuerpflichtiger: {C['name']}, {C['full_address']}\n"
         f"Steuernummer: {C['steuernr']}\nFinanzamt: {C['fa']}\n"
         f"Veranlagungszeitraum: 1. Januar {year} bis 31. Dezember {year}\n"
         f"Datum des Bescheids: {ds(d.year, d.month, d.day)}"),
        ("Festsetzung Körperschaftsteuer", [
            ["Position", "Betrag EUR"],
            ["Zu versteuerndes Einkommen", f"{ebt:,}"],
            ["Körperschaftsteuer 15 %", f"{koerperschaftsteuer:,}"],
            ["Solidaritätszuschlag 5,5 %", f"{solidaritaetszuschlag:,}"],
            ["Anrechenbare Vorauszahlungen", f"-{int(koerperschaftsteuer * 0.92):,}"],
            ["Verbleibende Zahlung", f"{int((koerperschaftsteuer + solidaritaetszuschlag) * 0.08):,}"],
        ]),
        ("Festsetzung Gewerbesteuer", [
            ["Position", "Betrag EUR"],
            ["Gewerbeertrag", f"{ebt:,}"],
            ["Gewerbesteuermessbetrag (3,5 %)", f"{int(ebt*0.035):,}"],
            ["Hebesatz München (490 %)", "490 %"],
            ["Gewerbesteuer", f"{gewerbesteuer:,}"],
            ["Anrechenbare Vorauszahlungen", f"-{int(gewerbesteuer * 0.90):,}"],
        ]),
        ("Rechtsbehelfsbelehrung",
         "Gegen diesen Bescheid kann innerhalb eines Monats nach seiner Bekanntgabe beim Finanzamt "
         f"{C['fa']} schriftlich oder zur Niederschrift Einspruch eingelegt werden."),
    ]
    save_pdf(FOLDER_FIN, f"STEUER_{year}_Steuerbescheid_KSt_GewSt.pdf",
             f"Steuerbescheid {year} – Körperschaftsteuer und Gewerbesteuer", sections)


def gen_quarterly_mgmt_report(year, quarter):
    qrevs = {
        "Q1": int(C.get(f"revenue_{year}", 87_200_000) * 0.22),
        "Q2": int(C.get(f"revenue_{year}", 87_200_000) * 0.24),
        "Q3": int(C.get(f"revenue_{year}", 87_200_000) * 0.25),
        "Q4": int(C.get(f"revenue_{year}", 87_200_000) * 0.29),
    }
    qrev = qrevs[quarter]
    qebitda = int(qrev * 0.142)
    sections = [
        ("Zusammenfassung",
         f"Das {quarter} {year} zeigt eine solide Entwicklung. Die Sentavia Precision GmbH erzielte im {quarter} "
         f"Umsatzerlöse von {eur(qrev)} (Vorjahr: {eur(int(qrev*0.85))}), was einem Wachstum von +17,6 % entspricht. "
         f"Das EBITDA beläuft sich auf {eur(qebitda)} ({qebitda/qrev*100:.1f} % Marge).\n\n"
         f"Besondere Entwicklungen im {quarter} {year}:\n"
         f"• {C['prod1_name']}: Neue Installationen bei {C['hosp1']} und {C['hosp2']}\n"
         f"• {C['prod2_name']}: CE-Rezertifizierung erfolgreich abgeschlossen\n"
         f"• Mitarbeiterzahl: {int(C.get(f'employees_{year}', 612) * (0.85 + 0.05*['Q1','Q2','Q3','Q4'].index(quarter)))} FTE"),
        ("Vertrieb und Märkte", [
            ["Produkt", "Umsatz EUR", "Einheiten", "Wachstum YoY", "Hauptmarkt"],
            [C["prod1_name"], f"{int(qrev*0.52):,}", str(random.randint(180, 240)), "+22%", "D-A-CH, BeNeLux"],
            [C["prod2_name"], f"{int(qrev*0.33):,}", str(random.randint(300, 450)), "+15%", "D-A-CH, Frankreich"],
            [C["prod3_name"], f"{int(qrev*0.15):,}", str(random.randint(5000, 8000)), "+12%", "Gesamte EU"],
        ]),
        ("Operative Highlights",
         f"Im Bereich Qualität und Regulatorik wurden im {quarter} {year} keine Serious Adverse Events (SAE) "
         "gemeldet. Die Anzahl offener Kundenbeschwerden (Complaints) ist auf ein Minimum reduziert worden. "
         f"Die Post-Market Surveillance Aktivitäten für {C['prod1_name']} verlaufen planmäßig.\n\n"
         f"Im Personalbereich wurden {random.randint(8, 20)} neue Mitarbeiter eingestellt, "
         f"vorrangig in den Bereichen F&E und Regulatory Affairs."),
        ("Ausblick",
         f"Für das nächste Quartal erwartet die Geschäftsführung ein weiteres Umsatzwachstum. "
         f"Die Jahresprognose von {eur(C.get('revenue_2024e', 104_000_000))} wird bestätigt."),
    ]
    save_pdf(FOLDER_FIN, f"MR_{year}_{quarter}_Management_Report.pdf",
             f"Management Report {quarter} {year}", sections)


def gen_budget_2024():
    months = ["Jan", "Feb", "Mär", "Apr", "Mai", "Jun",
              "Jul", "Aug", "Sep", "Okt", "Nov", "Dez", "GESAMT"]
    rev_annual = C["revenue_2024e"]
    rows = []
    ytd_rev = 0
    for i, m in enumerate(months[:-1]):
        factor = 1 + random.uniform(-0.05, 0.08) + (0.04 if i >= 8 else 0)
        m_rev = int(rev_annual / 12 * factor)
        ytd_rev += m_rev
        m_cogs = int(m_rev * 0.41)
        m_gross = m_rev - m_cogs
        m_opex = int(m_rev * 0.44)
        m_ebitda = m_gross - m_opex
        rows.append((m, f"{m_rev:,}", f"{m_cogs:,}", f"{m_gross:,}", f"{m_opex:,}", f"{m_ebitda:,}"))
    rows.append(("GESAMT", f"{rev_annual:,}", f"{int(rev_annual*0.41):,}",
                 f"{int(rev_annual*0.59):,}", f"{int(rev_annual*0.44):,}",
                 f"{int(rev_annual*0.15):,}"))
    sheets = [
        ("Budget 2024", ["Monat", "Umsatz EUR", "COGS EUR", "Bruttoergebnis EUR", "OpEx EUR", "EBITDA EUR"],
         rows, [10, 15, 13, 18, 13, 13]),
        ("Investitionen", ["Projekt", "Q1 EUR", "Q2 EUR", "Q3 EUR", "Q4 EUR", "Gesamt EUR"],
         [
             ("Reinraumausbau Phase 3", "400.000", "1.200.000", "800.000", "0", "2.400.000"),
             ("SAP S/4HANA Rollout", "300.000", "300.000", "200.000", "200.000", "1.000.000"),
             ("F&E Cardevio Pro Next", "500.000", "500.000", "500.000", "500.000", "2.000.000"),
             ("Laborausstattung", "200.000", "150.000", "100.000", "100.000", "550.000"),
             ("IT-Sicherheit & Cybersecurity", "150.000", "100.000", "100.000", "150.000", "500.000"),
             ("GESAMT", "1.550.000", "2.250.000", "1.700.000", "950.000", "6.450.000"),
         ], [30, 12, 12, 12, 12, 14]),
    ]
    make_xlsx_doc(FOLDER_FIN, "BUDGET_2024_Jahresplanung.xlsx",
                  "Budget und Jahresplanung 2024", sheets)


def gen_five_year_model():
    years = [2022, 2023, 2024, 2025, 2026, 2027]
    rev_base = C["revenue_2023"]
    rows = []
    prev_rev = rev_base
    for y in years:
        growth = random.uniform(0.14, 0.22) if y > 2023 else (0.226 if y == 2023 else 0.228)
        rev = int(prev_rev * (1 + growth)) if y > 2023 else C.get(f"revenue_{y}", prev_rev)
        ebitda_m = random.uniform(0.145, 0.185) if y > 2023 else (C.get(f"ebitda_{y}", 0) / rev)
        ebitda = C.get(f"ebitda_{y}", int(rev * ebitda_m)) if y <= 2023 else int(rev * ebitda_m)
        ebit_m = ebitda_m - random.uniform(0.025, 0.04)
        ebit = C.get(f"ebit_{y}", int(rev * ebit_m)) if y <= 2023 else int(rev * ebit_m)
        capex = int(rev * random.uniform(0.058, 0.075))
        fcf = ebitda - capex - int(rev * 0.015)
        emp = C.get(f"employees_{y}", int(prev_rev / 142_000))
        rows.append((str(y), f"{rev:,}", f"{(rev/prev_rev-1)*100:.1f}%",
                     f"{ebitda:,}", f"{ebitda/rev*100:.1f}%",
                     f"{ebit:,}", f"{ebit/rev*100:.1f}%",
                     f"{capex:,}", f"{fcf:,}", str(emp)))
        prev_rev = rev
    sheets = [
        ("5-Jahres-Modell", ["Jahr", "Umsatz EUR", "Wachstum", "EBITDA EUR", "EBITDA-%",
                              "EBIT EUR", "EBIT-%", "CapEx EUR", "FCF EUR", "Mitarbeiter"],
         rows, [8, 15, 10, 14, 10, 14, 10, 13, 13, 12]),
    ]
    make_xlsx_doc(FOLDER_FIN, "FIN_5_Jahres_Finanzmodell_2022_2027.xlsx",
                  "5-Jahres-Finanzmodell 2022–2027", sheets)


def gen_wp_bericht(year):
    rev = C.get(f"revenue_{year}", 87_200_000)
    sections = [
        ("Bestätigungsvermerk des unabhängigen Abschlussprüfers",
         f"An die Gesellschafter der {C['name']}, München.\n\n"
         f"Prüfungsurteil: Wir haben den Jahresabschluss der {C['name']} – bestehend aus der Bilanz "
         f"zum 31. Dezember {year}, der Gewinn- und Verlustrechnung, dem Anhang sowie dem Lagebericht – geprüft.\n\n"
         "Nach unserer Beurteilung vermittelt der beigefügte Jahresabschluss aufgrund der nach den deutschen "
         "handelsrechtlichen Vorschriften ein den tatsächlichen Verhältnissen entsprechendes Bild der "
         "Vermögens-, Finanz- und Ertragslage der Gesellschaft."),
        ("Grundlage für das Prüfungsurteil",
         "Wir haben unsere Abschlussprüfung in Übereinstimmung mit § 317 HGB und der EU-Abschlussprüfer-"
         "Verordnung (Nr. 537/2014) sowie unter Beachtung der vom Institut der Wirtschaftsprüfer (IDW) "
         "festgestellten deutschen Grundsätze ordnungsmäßiger Abschlussprüfung (ISA in der für Deutschland "
         "geltenden Fassung) durchgeführt. Unsere Verantwortung nach diesen Vorschriften und Grundsätzen "
         "ist im Abschnitt 'Verantwortung des Abschlussprüfers' weiter beschrieben."),
        ("Wesentliche Prüfungssachverhalte", [
            ["Prüfungssachverhalt", "Prüferische Handlungen"],
            ["Umsatzrealisierung und Periodenabgrenzung",
             "Analytische Prüfungshandlungen, Stichprobenprüfung von Lieferscheinen und Rechnungen"],
            ["Bewertung von Entwicklungskosten (IAS 38)",
             "Überprüfung der Aktivierungsvoraussetzungen, Bewertung der Nutzungsdauern"],
            ["Regulatorische Rückstellungen (MDR-Compliance)",
             "Einschätzung der Rückstellungshöhe durch Befragung des Regulatory-Teams, Einholung externer Gutachten"],
            ["Bewertung von Kundenforderungen",
             "Überprüfung der Wertberichtigungslogik, Bestätigung bei wesentlichen Kunden"],
        ]),
        ("Ergebnis", f"Der Jahresabschluss zum 31. Dezember {year} wurde mit einem uneingeschränkten "
         f"Bestätigungsvermerk versehen. Umsatz: {eur(rev)}. "
         f"EBITDA: {eur(C.get(f'ebitda_{year}', 12_400_000))}.\n\n"
         f"München, den {ds(year+1, 4, 15)}\n\n"
         f"{C['wp']}\n"
         "Wirtschaftsprüfungsgesellschaft\n\n"
         "Dr. Stefan Weber\t\tChristina Bauer\n"
         "Wirtschaftsprüfer\t\tWirtschaftsprüferin"),
    ]
    save_pdf(FOLDER_FIN, f"WP_{year}_Prüfungsbericht_Jahresabschluss.pdf",
             f"Bestätigungsvermerk Jahresabschluss {year}", sections)


def gen_kreditvertrag():
    doc = Document()
    make_docx_header(doc, "Kreditrahmenvertrag", f"{C['name']} / {C['bank']}", confidential=True)
    add_clause(doc, 1, "Krediteinräumung", [
        f"Die {C['bank']} (nachfolgend \"Bank\") räumt der {C['name']}, {C['full_address']}, "
        f"{C['hrb']}, (nachfolgend \"Kreditnehmer\") einen revolvierenden Kreditrahmen "
        f"(\"Kreditlinie\") in Höhe von bis zu EUR 15.000.000 (in Worten: Fünfzehn Millionen Euro) ein.",
        "Die Kreditlinie kann für folgende Zwecke in Anspruch genommen werden: (a) allgemeine "
        "Betriebsmittelfinanzierung; (b) kurzfristige Zwischenfinanzierung von Investitionen; "
        "(c) Überbrückung saisonaler Liquiditätsschwankungen; (d) Finanzierung von Warenlagern.",
        "Der Kreditrahmen ist revolverend ausgestaltet und kann innerhalb der vereinbarten Laufzeit "
        "jederzeit in Anspruch genommen, zurückgeführt und erneut abgerufen werden.",
    ])
    add_clause(doc, 2, "Zinsen und Gebühren", [
        "Für in Anspruch genommene Beträge ist ein Zinssatz von EURIBOR 3M + 1,85 % p.a. zu entrichten, "
        "mindestens jedoch 2,50 % p.a. Der Zinssatz wird vierteljährlich angepasst.",
        "Für nicht in Anspruch genommene Teile der Kreditlinie ist eine Bereitstellungsprovision von "
        "0,25 % p.a. auf den nicht genutzten Betrag zu zahlen.",
        "Eine einmalige Bearbeitungsgebühr von EUR 37.500 (0,25 % der Kreditlinie) ist bei "
        "Vertragsabschluss fällig.",
    ])
    add_clause(doc, 3, "Laufzeit und Kündigung", [
        "Die Kreditlinie wird für einen Zeitraum von 3 Jahren ab Vertragsschluss eingeräumt. "
        "Sie verlängert sich automatisch um jeweils ein weiteres Jahr, sofern sie nicht von einer "
        "Partei mit einer Frist von 6 Monaten zum Jahresende gekündigt wird.",
        "Die Bank ist zur außerordentlichen Kündigung berechtigt, wenn: (a) wesentliche "
        "Covenants verletzt werden; (b) eine wesentliche Verschlechterung der Bonität eintritt; "
        "(c) ein Insolvenzantrag gestellt wird.",
    ])
    add_clause(doc, 4, "Sicherheiten", [
        f"Der Kreditnehmer bestellt folgende Sicherheiten: (a) Globalzession aller gegenwärtigen "
        "und zukünftigen Kundenforderungen; (b) Sicherungsübereignung der Warenvorräte; "
        "(c) persönliche Bürgschaft der Gesellschafter für einen Betrag von insgesamt bis zu "
        "EUR 5.000.000.",
    ])
    add_clause(doc, 5, "Financial Covenants", [
        "Der Kreditnehmer verpflichtet sich, folgende Financial Covenants einzuhalten, jeweils "
        "gemessen auf Basis des geprüften Jahresabschlusses: (a) Verschuldungsgrad (Net Debt / "
        "EBITDA) kleiner als 3,0x; (b) DSCR (Debt Service Coverage Ratio) größer als 1,25x; "
        "(c) Eigenkapitalquote mindestens 35 %.",
        "Die Financial Covenants sind durch den Kreditnehmer jährlich innerhalb von 120 Tagen "
        "nach Geschäftsjahresende mittels geprüften Jahresabschlusses nachzuweisen. "
        "Bei Verletzung eines Covenants ist unverzüglich die Bank zu informieren.",
    ])
    save_docx(doc, FOLDER_FIN, f"KREDIT_001_Kreditrahmenvertrag_HypoVereinsbank.docx")


def gen_forschungsfoerderung():
    sections = [
        ("Antrag auf steuerliche Forschungsförderung nach FZulG",
         f"Antragsteller: {C['name']}, {C['full_address']}\n"
         f"Steuernummer: {C['steuernr']}, Finanzamt: {C['fa']}\n"
         f"Antragsjahr: 2023\n"
         f"Zuständige Bescheinigungsstelle: BSFZ (Bescheinigungsstelle Forschungszulage)"),
        ("F&E-Aktivitäten", [
            ["Vorhaben", "Personal-FTE", "Sachkosten EUR", "Gesamt EUR"],
            [f"NextGen {C['prod1_name']} Sensorik", "4,5", "180.000", "1.024.000"],
            [f"AI-basierte Bildanalyse {C['prod2_name']}", "3,2", "95.000", "748.000"],
            [f"Biomarker-Plattform {C['prod3_name']}", "2,8", "120.000", "622.000"],
            ["Grundlagenforschung Biosensor-Arrays", "2,0", "85.000", "430.000"],
            ["GESAMT", "12,5", "480.000", "2.824.000"],
        ]),
        ("Berechnungsgrundlage",
         "Förderungsfähige Aufwendungen gemäß § 3 FZulG: EUR 2.824.000\n"
         "Fördersatz: 25 % (KMU)\n"
         "Maximale jährliche Förderbasis: EUR 4.000.000\n"
         "Förderbetrag beantragt: EUR 706.000"),
    ]
    save_pdf(FOLDER_FIN, "FOERD_2023_Antrag_Forschungsfoerderung_FZulG.pdf",
             "Antrag auf Forschungszulage nach FZulG – Antragsjahr 2023", sections)


def generate_02_finanzen():
    print("  Generating 02_Finanzen...")
    for yr in [2021, 2022, 2023]:
        gen_jahresabschluss(yr)
        gen_steuerbescheid(yr)
        gen_wp_bericht(yr)
    for yr in [2022, 2023]:
        gen_bwa_monatlich(yr)
    for yr in [2023]:
        for q in ["Q1", "Q2", "Q3", "Q4"]:
            gen_quarterly_mgmt_report(yr, q)
    gen_budget_2024()
    gen_five_year_model()
    gen_kreditvertrag()
    gen_forschungsfoerderung()
    # Investor reporting packages
    for yr in [2022, 2023]:
        for q in ["Q1", "Q2", "Q3", "Q4"]:
            rev = C.get(f"revenue_{yr}", 87_200_000) // 4
            sheets = [
                ("Investor Report", ["Kennzahl", "Wert", "vs. Plan", "vs. Vorjahr"],
                 [
                     ("Umsatz EUR", f"{int(rev*(1+random.uniform(-0.05,0.1))):,}",
                      f"+{random.randint(1,8)}%", f"+{random.randint(12,25)}%"),
                     ("EBITDA EUR", f"{int(rev*0.142):,}", f"+{random.randint(-2,5)}%", f"+{random.randint(8,18)}%"),
                     ("Cash-Bestand EUR", f"{int(random.uniform(8,18)*1e6):,}", "n/a", "n/a"),
                     ("Mitarbeiter FTE", str(C.get(f"employees_{yr}", 512)), f"+{random.randint(3,10)}", f"+{random.randint(15,25)}%"),
                 ], [25, 15, 12, 15]),
            ]
            make_xlsx_doc(FOLDER_FIN, f"INV_{yr}_{q}_Investor_Reporting_Package.xlsx",
                          f"Investor Reporting Package {q} {yr}", sheets)
    # Additional financial docs
    for title, fname in [
        ("Liquiditätsplanung 2024 – Cashflow-Vorschau", "LIQUID_2024_Cashflow_Planung.xlsx"),
        ("Kreditoren-Debitorenliste Q4 2023", "KREDITOREN_DEBITOREN_Q4_2023.xlsx"),
        ("Anlagenspiegel 2023", "ANLAGENSPIEGEL_2023.xlsx"),
        ("Rückstellungsspiegel 2023", "RUECKSTELLUNGEN_2023.xlsx"),
        ("Intercompany-Transaktionen 2023", "INTERCOMPANY_2023.xlsx"),
    ]:
        sheets = [("Übersicht", ["Position", "Betrag EUR", "Anmerkung"],
                   [(f"Posten {i}", f"{random.randint(50,5000)*1000:,}", "Gemäß Buchhaltung") for i in range(1, 12)],
                   [30, 15, 30])]
        make_xlsx_doc(FOLDER_FIN, fname, title, sheets)


# ══════════════════════════════════════════════════════════════════════
# FOLDER 03 – PERSONAL / HR
# ══════════════════════════════════════════════════════════════════════
FOLDER_HR = "03_Personal_HR"

EMPLOYEES = [
    ("Dr. Lena Fischer", "Senior R&D Engineer", "F&E", 95_000),
    ("Markus Bauer", "Regulatory Affairs Manager", "Regulatory", 88_000),
    ("Sarah Krause", "Quality Assurance Lead", "Qualität", 82_000),
    ("Jan Hoffmann", "Software Engineer – Medical", "IT/Software", 78_000),
    ("Anna Weber", "Clinical Affairs Specialist", "Klinik", 85_000),
    ("Felix Schäfer", "Head of Sales DACH", "Vertrieb", 105_000),
    ("Julia Zimmermann", "Produktmanagerin Cardevio", "Produktmanagement", 90_000),
    ("Patrick Richter", "Supply Chain Manager", "Einkauf", 80_000),
    ("Claudia Steinberg", "HR Business Partner", "HR", 72_000),
    ("Michael Lange", "Finance Controller", "Finanzen", 82_000),
    ("Dr. Tobias Keller", "Medical Affairs Director", "Medical", 115_000),
    ("Laura Neumann", "IP & Patent Counsel", "Legal", 98_000),
    ("Benjamin Wolf", "IT Security Manager", "IT", 88_000),
    ("Monika Fuchs", "Production Manager", "Produktion", 92_000),
    ("Christoph Braun", "Business Development", "Strategie", 105_000),
]


def gen_employment_contract(name, title, dept, salary, idx):
    doc = Document()
    make_docx_header(doc, f"Arbeitsvertrag – {name}",
                     f"{C['name']} und {name}", confidential=True)
    start_year = random.randint(2015, 2023)
    start_month = random.randint(1, 12)
    doc.add_paragraph(
        f"ARBEITSVERTRAG\n\nzwischen\n\n{C['name']}, {C['full_address']}, {C['hrb']}\n"
        f"– nachfolgend \"Arbeitgeber\" –\n\nund\n\n{name}, wohnhaft in München\n"
        f"– nachfolgend \"Arbeitnehmer\" –"
    )
    add_clause(doc, 1, "Beginn und Art der Tätigkeit", [
        f"Das Arbeitsverhältnis beginnt am {ds(start_year, start_month, 1)}. "
        f"Der Arbeitnehmer wird als {title} im Bereich {dept} eingestellt.",
        "Der Arbeitgeber ist berechtigt, dem Arbeitnehmer eine andere zumutbare Tätigkeit zuzuweisen, "
        "sofern dies betrieblich erforderlich ist und die neue Tätigkeit der Qualifikation und dem "
        "Erfahrungsstand des Arbeitnehmers entspricht.",
    ])
    add_clause(doc, 2, "Vergütung", [
        f"Der Arbeitnehmer erhält ein monatliches Bruttogehalt von "
        f"{eur(salary // 12)} (Jahresgehalt: {eur(salary)}), zahlbar "
        "am letzten Werktag eines jeden Kalendermonats.",
        "Zusätzlich kann der Arbeitnehmer eine leistungsbezogene Jahresprämie (Bonus) von bis zu "
        "15 % des Jahresbruttogehalts erhalten. Die Gewährung erfolgt nach freiem Ermessen des "
        "Arbeitgebers auf Basis der individuellen Leistungsbeurteilung und der Zielerreichung.",
        "Der Arbeitnehmer erhält einen Zuschuss zur betrieblichen Altersvorsorge (bAV) in Höhe "
        "von 20 % des vom Arbeitnehmer eingezahlten Betrags, maximal EUR 1.200 pro Jahr.",
    ])
    add_clause(doc, 3, "Arbeitszeit", [
        "Die regelmäßige wöchentliche Arbeitszeit beträgt 40 Stunden. Beginn und Ende der "
        "täglichen Arbeitszeit richten sich nach den betrieblichen Erfordernissen und sind "
        "mit dem Vorgesetzten abzustimmen.",
        "Mehrarbeit ist mit dem Gehalt abgegolten, soweit sie 10 % der vertraglich vereinbarten "
        "Arbeitszeit nicht übersteigt. Darüber hinausgehende Mehrarbeit ist durch Freizeitausgleich "
        "oder eine gesonderte schriftliche Vereinbarung zu kompensieren.",
    ])
    add_clause(doc, 4, "Urlaub", [
        "Der Arbeitnehmer hat Anspruch auf 28 Arbeitstage Erholungsurlaub im Kalenderjahr "
        "(Urlaubsjahr = Kalenderjahr). Bei Ein- oder Austritt im Laufe des Jahres wird der Urlaub "
        "anteilig gewährt.",
        "Der Urlaub ist in Absprache mit dem Vorgesetzten zu nehmen und möglichst bis zum Ende "
        "des Kalenderjahres zu verbrauchen. Ein Übertrag ins Folgejahr ist nur bei dringenden "
        "betrieblichen Gründen oder persönlichen Verhinderungen und nur bis zum 31. März möglich.",
    ])
    add_clause(doc, 5, "Geheimhaltung", [
        "Der Arbeitnehmer verpflichtet sich, über alle ihm im Rahmen seiner Tätigkeit bekannt "
        "werdenden Geschäfts- und Betriebsgeheimnisse sowie vertraulichen Informationen der "
        "Gesellschaft und ihrer Kunden Stillschweigen zu bewahren. Diese Verpflichtung gilt "
        "auch nach Beendigung des Arbeitsverhältnisses.",
        "Zu den vertraulichen Informationen zählen insbesondere: technische Spezifikationen, "
        "klinische Daten, Kundenlisten, Lieferantenkonditionen, Finanzinformationen sowie "
        "Informationen über laufende Produktentwicklungen und Zulassungsverfahren.",
    ])
    add_clause(doc, 6, "Erfindungen und Urheberrecht", [
        "Alle Erfindungen, Entwicklungen, Software und sonstigen schutzfähigen Leistungen, "
        "die der Arbeitnehmer im Rahmen seiner Tätigkeit erbringt (\"Arbeitnehmererfindungen\"), "
        "werden dem Arbeitgeber gemäß Arbeitnehmererfindungsgesetz (ArbnErfG) gemeldet und "
        "von diesem in Anspruch genommen.",
        "Der Arbeitnehmer erhält für in Anspruch genommene Diensterfindungen eine angemessene "
        "Vergütung gemäß den Richtlinien für die Vergütung von Arbeitnehmererfindungen "
        "(Vergütungsrichtlinien Nr. 12 des DPMA).",
    ])
    add_clause(doc, 7, "Kündigung", [
        "Das Arbeitsverhältnis kann nach Ablauf der Probezeit (6 Monate) von beiden Seiten "
        "mit einer Frist von 3 Monaten zum Monatsende gekündigt werden. Bei längerer "
        "Betriebszugehörigkeit verlängern sich die Kündigungsfristen gemäß § 622 BGB.",
        "Das Recht zur außerordentlichen Kündigung aus wichtigem Grund (§ 626 BGB) bleibt unberührt.",
    ])
    add_clause(doc, 8, "Schlussbestimmungen", [
        "Änderungen und Ergänzungen dieses Vertrags bedürfen der Schriftform. "
        "Nebenabreden wurden nicht getroffen.",
        "Sollten einzelne Bestimmungen dieses Vertrags unwirksam sein, bleibt die Wirksamkeit "
        "des übrigen Vertrags unberührt. An die Stelle unwirksamer Klauseln tritt eine "
        "wirksame Regelung, die dem wirtschaftlichen Zweck am nächsten kommt.",
        f"Es gilt deutsches Recht. Gerichtsstand: {C['city']}.",
    ])
    doc.add_paragraph(
        f"\n\nMünchen, den {ds(start_year, start_month, 1)}\n\n"
        f"_______________________\t\t_______________________\n"
        f"{C['name']}\t\t{name}\n"
        f"(Arbeitgeber)\t\t\t(Arbeitnehmer)"
    )
    save_docx(doc, FOLDER_HR, f"AV_{idx:02d}_Arbeitsvertrag_{sfn(name.replace(' ','_'))}.docx")


def gen_organigramm():
    sections = [
        ("Organisationsstruktur Sentavia Precision GmbH",
         f"Stand: Januar 2024 | Gesamtmitarbeiter: {C['employees_2023']} FTE"),
        ("Geschäftsführung",
         f"CEO: {C['ceo']}\nCTO: {C['cto']}\nCFO: {C['cfo']}\nCMO: {C['cmo']}\nQRA: {C['qra']}"),
        ("Abteilungsstruktur", [
            ["Bereich", "Leitung", "Mitarbeiter", "Untereinheiten"],
            ["Forschung & Entwicklung", C["cto"], "128", "Hardwareentwicklung, Softwareentwicklung, Validierung"],
            ["Regulatory & Quality", C["qra"], "52", "Regulatory Affairs, QMS, Klinische Bewertung"],
            ["Medical Affairs", C["cmo"], "38", "Klinische Studien, Medical Education"],
            ["Vertrieb & Marketing", "Felix Schäfer (VP Sales)", "95", "DACH, BeNeLux, Frankreich, Export"],
            ["Produktion", "Monika Fuchs (Produktionsleitung)", "148", "Fertigung, Logistik, Supply Chain"],
            ["Finanzen & Controlling", C["cfo"], "42", "Buchhaltung, Controlling, Treasury"],
            ["HR & Organisation", "Claudia Steinberg (HR-Leitung)", "28", "Recruiting, Personalentwicklung"],
            ["IT & Digital", "Benjamin Wolf (IT-Leitung)", "35", "ERP, QMS-IT, Cybersecurity"],
            ["Legal & IP", "Laura Neumann (Legal)", "12", "Verträge, Patente, Compliance"],
            ["Geschäftsentwicklung", "Christoph Braun (BD)", "8", "M&A, Partnerschaften"],
        ]),
    ]
    save_pdf(FOLDER_HR, "HR_001_Organigramm_2024.pdf", "Organisationsstruktur", sections)


def gen_betriebsvereinbarung(topic, num):
    doc = Document()
    make_docx_header(doc, f"Betriebsvereinbarung – {topic}", confidential=False)
    doc.add_paragraph(
        f"zwischen\n\n{C['name']}, {C['full_address']} – nachfolgend \"Arbeitgeber\" –\n\n"
        "und\n\ndem Betriebsrat der Sentavia Precision GmbH, München – nachfolgend \"Betriebsrat\" –\n\n"
        f"wird folgende Betriebsvereinbarung geschlossen:\n\n"
        f"Präambel: Arbeitgeber und Betriebsrat sind sich einig, dass {topic} im gemeinsamen "
        "Interesse der Belegschaft und des Unternehmens liegt. Diese Betriebsvereinbarung dient "
        "der Regelung der damit verbundenen Rechte und Pflichten."
    )
    add_clause(doc, 1, "Geltungsbereich",
               ["Diese Betriebsvereinbarung gilt für alle Arbeitnehmer der Gesellschaft mit Ausnahme "
                "der leitenden Angestellten im Sinne des § 5 Abs. 3 BetrVG.",
                "Für Leiharbeitnehmer gilt diese Betriebsvereinbarung entsprechend, soweit dies "
                "gesetzlich vorgesehen ist."])
    add_clause(doc, 2, "Grundsätze",
               ["Die Parteien bekennen sich zu den Grundsätzen der partnerschaftlichen Zusammenarbeit, "
                "der gegenseitigen Rücksichtnahme sowie der Förderung einer positiven Unternehmenskultur.",
                "Alle Maßnahmen im Rahmen dieser Betriebsvereinbarung werden auf ihre Auswirkungen auf "
                "die Beschäftigten und das Unternehmen hin geprüft."])
    add_clause(doc, 3, "Regelungen",
               [f"Im Bereich {topic} gelten folgende Regelungen: Alle betroffenen Maßnahmen sind "
                "frühzeitig mit dem Betriebsrat abzustimmen. Dem Betriebsrat sind alle relevanten "
                "Informationen rechtzeitig zur Verfügung zu stellen.",
                "Der Betriebsrat ist vor der Einführung neuer Maßnahmen gemäß §§ 87 ff. BetrVG "
                "zu beteiligen. Die Zustimmung des Betriebsrats ist in den jeweils gesetzlich "
                "vorgesehenen Fällen einzuholen."])
    add_clause(doc, 4, "Laufzeit",
               ["Diese Betriebsvereinbarung tritt am 1. Januar 2024 in Kraft und gilt auf unbestimmte Zeit.",
                "Sie kann von jeder Seite mit einer Frist von 3 Monaten zum Quartalsende gekündigt werden. "
                "Nach Kündigung wirkt die Betriebsvereinbarung nach bis eine neue Regelung getroffen ist."])
    doc.add_paragraph(
        f"\nMünchen, den {ds(2024, 1, 1)}\n\n"
        f"_______________________\t\t_______________________\n"
        f"{C['ceo']}\t\t\tBetriebsrat\n"
        f"Geschäftsführerin\t\t\tVorsitzender"
    )
    save_docx(doc, FOLDER_HR, f"BV_{num:02d}_Betriebsvereinbarung_{sfn(topic[:25])}.docx")


def gen_personalstammdaten():
    rows = []
    for i, (name, title, dept, salary) in enumerate(EMPLOYEES, 1):
        rows.append((f"EMP-{2000+i}", name, title, dept,
                     ds(random.randint(2012, 2022), random.randint(1, 12), 1),
                     f"{salary:,}", "Unbefristet", "Vollzeit", "Aktiv"))
    sheets = [
        ("Personalstammdaten", ["Pers-Nr.", "Name", "Position", "Abteilung",
                                 "Eintritt", "Jahresgehalt EUR", "Vertragstyp", "Beschäftigungsart", "Status"],
         rows, [12, 22, 28, 18, 14, 16, 14, 16, 10]),
    ]
    make_xlsx_doc(FOLDER_HR, "HR_002_Personalstammdaten_2024.xlsx",
                  "Personalstammdaten Stand Januar 2024", sheets)


def gen_salary_structure():
    sheets = [
        ("Gehaltsstruktur", ["Level", "Bezeichnung", "Bandbreite von EUR", "Bandbreite bis EUR",
                              "Ø Ist-Gehalt EUR", "Anz. MA"],
         [
             ("L1", "Junior / Berufseinsteiger", "45.000", "58.000", "51.500", "42"),
             ("L2", "Professional", "58.000", "75.000", "66.800", "128"),
             ("L3", "Senior Professional", "75.000", "95.000", "84.200", "156"),
             ("L4", "Lead / Manager", "90.000", "115.000", "102.400", "89"),
             ("L5", "Director / Head of", "110.000", "145.000", "128.600", "38"),
             ("L6", "VP / Senior Director", "140.000", "185.000", "162.000", "14"),
             ("GF", "Geschäftsführung", "280.000", "400.000", "340.000", "3"),
         ], [6, 28, 18, 16, 18, 10]),
    ]
    make_xlsx_doc(FOLDER_HR, "HR_003_Gehaltsstruktur_Baender_2024.xlsx",
                  "Gehaltsstruktur und -bänder 2024", sheets)


def generate_03_personal_hr():
    print("  Generating 03_Personal_HR...")
    for i, (name, title, dept, salary) in enumerate(EMPLOYEES, 1):
        gen_employment_contract(name, title, dept, salary, i)
    gen_organigramm()
    gen_personalstammdaten()
    gen_salary_structure()
    bvs = ["Mobiles Arbeiten (Homeoffice)", "Arbeitszeit und Gleitzeitregelung",
           "Nutzung von IT-Systemen", "Gesundheitsförderung und BGM",
           "Weiterbildung und Schulungsmaßnahmen", "Reisekostenregelung"]
    for i, topic in enumerate(bvs, 1):
        gen_betriebsvereinbarung(topic, i)
    # Job descriptions
    for name, title, dept, _ in EMPLOYEES[:8]:
        sections = [
            ("Stellenprofil", f"Position: {title}\nAbteilung: {dept}\nStandort: {C['city']}\n"
             f"Berichtsweg: Abteilungsleitung {dept}"),
            ("Aufgaben und Verantwortlichkeiten",
             f"• Fachliche Verantwortung im Bereich {dept}\n"
             "• Entwicklung und Implementierung von Lösungen und Prozessen\n"
             "• Enge Zusammenarbeit mit internen Stakeholdern und externen Partnern\n"
             "• Sicherstellung der Einhaltung regulatorischer Anforderungen (MDR, ISO 13485)\n"
             "• Regelmäßige Berichterstattung an die Abteilungsleitung"),
            ("Anforderungen",
             "• Abgeschlossenes Studium der Ingenieurwissenschaften, Naturwissenschaften oder vergleichbar\n"
             "• Mindestens 3 Jahre Berufserfahrung in der Medizintechnikbranche\n"
             "• Kenntnisse in MDR 2017/745, ISO 13485:2016 und relevanten Normen\n"
             "• Fließende Deutsch- und Englischkenntnisse in Wort und Schrift"),
        ]
        save_pdf(FOLDER_HR, f"JD_{sfn(title[:20])}.pdf", f"Stellenbeschreibung – {title}", sections)
    # Training records
    sheets = [
        ("Schulungsnachweise 2023", ["Mitarbeiter", "Schulung", "Datum", "Trainer", "Bestanden"],
         [(name, random.choice(["GMP-Grundschulen", "MDR-Update 2023", "ISO 13485 Refresher",
                                 "Cybersecurity Awareness", "CAPA-Training"]),
           ds(2023, random.randint(1, 12), random.randint(1, 28)),
           random.choice(["intern", C["qra"], "extern"]),
           "JA") for name, _, _, _ in EMPLOYEES],
         [25, 30, 15, 20, 10]),
    ]
    make_xlsx_doc(FOLDER_HR, "HR_004_Schulungsnachweise_2023.xlsx",
                  "Schulungsnachweise und Trainingshistorie 2023", sheets)
    # Works council election protocol
    sections = [
        ("Betriebsratswahl 2023 – Wahlprotokoll",
         f"Die Betriebsratswahl bei der {C['name']} wurde am {ds(2023, 3, 15)} durchgeführt. "
         "Es wurden 5 Betriebsratsmitglieder gewählt."),
        ("Wahlergebnis", [
            ["Kandidat", "Stimmen", "Platz", "Status"],
            ["Klaus Schmidt", "189", "1", "Gewählt"],
            ["Petra Hartmann", "178", "2", "Gewählt"],
            ["Andreas Müller", "165", "3", "Gewählt"],
            ["Sandra Weiß", "154", "4", "Gewählt"],
            ["Dirk Becker", "142", "5", "Gewählt"],
        ]),
    ]
    save_pdf(FOLDER_HR, "HR_005_Betriebsratswahl_2023_Protokoll.pdf",
             "Betriebsratswahl 2023 – Wahlprotokoll", sections)


# ══════════════════════════════════════════════════════════════════════
# FOLDER 04 – VERTRIEB / DISTRIBUTOREN
# ══════════════════════════════════════════════════════════════════════
FOLDER_SALES = "04_Vertrieb_Distributoren"


def gen_distribution_agreement(dist_name, dist_short, territory, excl=True, lang="de"):
    if lang == "en":
        _gen_dist_agreement_en(dist_name, dist_short, territory)
        return
    doc = Document()
    make_docx_header(doc, f"Distributionsvertrag – {dist_name}",
                     f"{C['name']} und {dist_name}", confidential=True)
    doc.add_paragraph("DISTRIBUTIONSVERTRAG\n", style='Heading 2')
    doc.add_paragraph(
        f"zwischen\n\n{C['name']}, {C['full_address']}, {C['hrb']}\n"
        f"– nachfolgend \"Hersteller\" –\n\nund\n\n{dist_name}\n"
        f"– nachfolgend \"Distributor\" –\n\n"
        "Präambel:\n\nDer Hersteller entwickelt, produziert und vertreibt innovative Medizinprodukte "
        f"der Klassen IIa und IIb gemäß EU-MDR 2017/745 sowie In-vitro-Diagnostika nach IVDR 2017/746. "
        f"Der Distributor verfügt über umfangreiche Erfahrung im Vertrieb von Medizinprodukten im "
        f"Gebiet {territory} und hat die regulatorischen, logistischen und vertrieblichen Voraussetzungen "
        f"geschaffen, die für den Vertrieb der Produkte des Herstellers erforderlich sind. "
        f"Die Parteien beabsichtigen, ihre Zusammenarbeit im Bereich des Vertriebs von Medizinprodukten "
        "auf der Grundlage dieses Vertrags zu regeln."
    )
    add_clause(doc, 1, "Vertragsgegenstand", [
        f"Der Hersteller ernennt den Distributor zum {'exklusiven' if excl else 'nicht-exklusiven'} "
        f"Distributor für die in Anlage 1 aufgeführten Produkte ({C['prod1_name']}, {C['prod2_name']}, "
        f"{C['prod3_name']}) für das Vertragsgebiet {territory}.",
        "Der Distributor erwirbt die Produkte vom Hersteller auf eigene Rechnung und im eigenen Namen "
        "und vertreibt sie im Vertragsgebiet. Der Distributor ist kein Handelsvertreter im Sinne "
        "des HGB; er handelt nicht als Beauftragter oder Vertreter des Herstellers.",
        "Dieser Vertrag gilt ausschließlich für die in Anlage 1 aufgelisteten Produkte in den jeweils "
        "aktuellen CE-gekennzeichneten Versionen. Änderungen der Produktpalette bedürfen einer "
        "schriftlichen Ergänzungsvereinbarung.",
    ])
    add_clause(doc, 2, "Exklusivität und Gebietsschutz", [
        f"Der Hersteller verpflichtet sich, während der Laufzeit dieses Vertrags keine anderen "
        f"Distributoren für das Vertragsgebiet {territory} zu ernennen und das Vertragsgebiet "
        f"nicht selbst aktiv zu bearbeiten.",
        "Der Distributor verpflichtet sich, aktive Verkaufsbemühungen außerhalb des Vertragsgebiets "
        "zu unterlassen. Passive Verkäufe – d.h. unaufgeforderte Anfragen aus dem EU-Ausland – sind "
        "gestattet, sofern dies EU-rechtlich zulässig ist.",
        "Der Gebietsschutz erlischt, wenn der Distributor in zwei aufeinanderfolgenden Jahren die "
        "vereinbarten Mindestabnahmemengen (§ 3) unterschreitet. In diesem Fall kann der Hersteller "
        "nach 60-tägiger Abhilfefrist die Exklusivität schriftlich widerrufen.",
        "Bei Beendigung der Exklusivität hat der Distributor für eine Übergangszeit von 12 Monaten "
        "das Recht, bestehende Kundenbeziehungen zu Sonderkonditionen weiterzubedienen.",
    ])
    add_clause(doc, 3, "Mindestabnahmemengen und Forecasting", [
        f"Der Distributor verpflichtet sich, in den Jahren der Vertragslaufzeit folgende "
        f"Mindestabnahmemengen zu erwerben: Jahr 1: {eur(random.randint(3,6)*1_000_000)}; "
        f"Jahr 2: {eur(random.randint(4,7)*1_000_000)}; Jahr 3: {eur(random.randint(5,9)*1_000_000)}.",
        "Zum 1. Oktober jedes Jahres übermittelt der Distributor dem Hersteller eine Absatzplanung "
        "(Forecast) für das folgende Kalenderjahr, aufgeschlüsselt nach Produkten und Quartalen. "
        "Die Absatzplanung ist nicht verbindlich, dient aber als Grundlage für die Produktionsplanung "
        "des Herstellers.",
        "Zusätzlich übermittelt der Distributor monatlich bis zum 5. Werktag des Folgemonats einen "
        "Rolling-Forecast für die nächsten 3 Monate (\"Rolling 3-Month Forecast\").",
        "Unterschreitet der Distributor die Mindestabnahmemenge eines Jahres um mehr als 20 %, "
        "hat der Hersteller das Recht, die vereinbarte Exklusivität zu widerrufen oder den Vertrag "
        "außerordentlich zu kündigen.",
        "Der Distributor informiert den Hersteller unverzüglich über Umstände, die eine wesentliche "
        "Abweichung von der Absatzplanung erwarten lassen, insbesondere über verlorene Ausschreibungen "
        "oder Veränderungen in der Wettbewerbssituation.",
    ])
    add_clause(doc, 4, "Preise und Zahlungsbedingungen", [
        "Die Einkaufspreise des Distributors sind in Anlage 2 (Preisliste) festgelegt. Der Hersteller "
        "behält sich das Recht vor, die Preisliste einmal jährlich mit einer Vorankündigungsfrist "
        "von 90 Tagen anzupassen. Preissenkungen gelten sofort.",
        "Rechnungen sind innerhalb von 30 Tagen nach Rechnungsdatum ohne Abzug zu begleichen. "
        "Bei Zahlung innerhalb von 10 Tagen gewährt der Hersteller ein Skonto von 2 %. "
        "Bei Zahlungsverzug werden Verzugszinsen gemäß § 288 Abs. 2 BGB berechnet.",
        "Der Distributor ist nicht berechtigt, gegenüber Preisforderungen des Herstellers mit "
        "eigenen Gegenansprüchen aufzurechnen, es sei denn, die Gegenansprüche sind unbestritten "
        "oder rechtskräftig festgestellt.",
        "Alle Preise verstehen sich zuzüglich der jeweils geltenden gesetzlichen Umsatzsteuer. "
        "Einfuhrzölle und Importabgaben im Vertragsgebiet trägt der Distributor.",
    ])
    add_clause(doc, 5, "Lieferbedingungen und Abnahme", [
        f"Lieferungen erfolgen ab Werk {C['city']} (Incoterms® 2020: EXW). "
        "Gefahr und Kosten gehen mit Übergabe an den Frachtführer auf den Distributor über.",
        "Der Distributor ist verpflichtet, bei Anlieferung eine Eingangskontrolle durchzuführen "
        "(Qualitätseingangsprüfung gemäß ISO 13485). Mängel sind innerhalb von 5 Werktagen "
        "nach Anlieferung schriftlich zu rügen.",
        "Lieferzeiten sind unverbindliche Richtzeiten. Der Hersteller bemüht sich, Bestellungen "
        "innerhalb von 10 Werktagen zu bearbeiten. Bei Engpässen werden verfügbare Mengen "
        "anteilig auf alle Distributoren aufgeteilt.",
        "Der Distributor verpflichtet sich, eine Mindestbevorratung von 4 Wochen Umsatz je "
        "Produkt aufrechtzuerhalten. Hierüber ist dem Hersteller monatlich Bericht zu erstatten.",
        "Mindesthaltbarkeitsdaten: Der Distributor bestellt nur Produkte mit einer Restlaufzeit "
        "von mindestens 18 Monaten. Produkte mit kürzerer Restlaufzeit werden nur auf "
        "ausdrücklichen schriftlichen Wunsch des Distributors geliefert.",
    ])
    add_clause(doc, 6, "Regulatorische Pflichten des Distributors", [
        "Der Distributor bestätigt, dass er alle regulatorischen Anforderungen erfüllt, die im "
        "Vertragsgebiet für den Vertrieb von Medizinprodukten gelten, insbesondere die "
        "Anforderungen der EU-MDR 2017/745 und der IVDR 2017/746 an Distributoren "
        "(Art. 14 MDR / Art. 10 IVDR).",
        "Der Distributor verpflichtet sich, bei Verdacht auf oder Kenntnis von schwerwiegenden "
        "Vorkommnissen (Serious Incidents gemäß Art. 87 MDR) den Hersteller unverzüglich, "
        "spätestens jedoch innerhalb von 24 Stunden, zu informieren. Der Distributor meldet "
        "keine Vorkommnisse direkt an Behörden, ohne vorherige Absprache mit dem Hersteller.",
        "Feldsicherheitskorrekturmaßnahmen (FSCA): Im Falle einer vom Hersteller angeordneten FSCA "
        "(Field Safety Corrective Action) ist der Distributor verpflichtet, betroffene Produkte "
        "innerhalb der vom Hersteller gesetzten Frist zurückzurufen und alle betroffenen Kunden "
        "schriftlich zu informieren. Kosten der FSCA trägt der Hersteller.",
        "UDI-Compliance: Der Distributor stellt sicher, dass an den Produkten keine Veränderungen "
        "vorgenommen werden, die die Konformität mit den UDI-Anforderungen (Art. 27 MDR) "
        "beeinträchtigen. Insbesondere dürfen Etiketten oder UDI-Träger nicht entfernt, "
        "überklebt oder verändert werden.",
        "Der Distributor führt ein vollständiges Rückverfolgungssystem (Traceability System) für "
        "alle gelieferten Produkte, das die Identifizierung von Lot-Nummern, Seriennummern, "
        "Endanwendern und Versorgungszeitpunkt ermöglicht. Diese Informationen sind dem Hersteller "
        "auf Anfrage unverzüglich mitzuteilen und für mindestens 15 Jahre aufzubewahren.",
        "Der Distributor kooperiert vollumfänglich bei Post-Market-Surveillance-Aktivitäten "
        "(PMS) des Herstellers gemäß Art. 83 ff. MDR. Er stellt dem Hersteller regelmäßig "
        "Informationen über die Marktbeobachtung zur Verfügung, einschließlich Kundenfeedback, "
        "Beschwerden und Informationen über vergleichbare Wettbewerbsprodukte.",
    ])
    add_clause(doc, 7, "Produkthaftung und Rückrufe", [
        "Der Hersteller übernimmt die Produkthaftung für Schäden, die auf Konstruktions-, "
        "Herstellungs- oder Instruktionsfehlern beruhen. Der Distributor haftet für Schäden, "
        "die durch unsachgemäße Lagerung, Transport oder Handhabung entstanden sind.",
        "Bei Rückrufmaßnahmen (Recalls) koordiniert der Hersteller alle Maßnahmen. Der "
        "Distributor führt die Rückrufmaßnahmen im Vertragsgebiet durch und erstattet dem "
        "Hersteller vollständige Bericht über alle betroffenen Einheiten und kontaktierten Kunden.",
        "Der Distributor unterhält eine eigene Produkthaftpflichtversicherung mit einer "
        "Mindestdeckungssumme von EUR 10.000.000 je Schadensfall und stellt dem Hersteller "
        "auf Verlangen eine Kopie des Versicherungsscheins zur Verfügung.",
        "Der Distributor stellt den Hersteller von allen Ansprüchen Dritter frei, die auf "
        "schuldhaftem Verhalten des Distributors beruhen, insbesondere auf unsachgemäßer "
        "Lagerung, falschem Transport oder unzulässigen Produktveränderungen.",
    ])
    add_clause(doc, 8, "Qualitätssicherung und Audits", [
        "Der Distributor verpflichtet sich, ein Qualitätsmanagementsystem gemäß ISO 13485:2016 "
        "für die distributionsspezifischen Tätigkeiten zu unterhalten und auf Anfrage nachzuweisen.",
        "Der Hersteller ist berechtigt, beim Distributor einmal jährlich (und bei begründetem "
        "Anlass auch außerordentlich) Qualitätsaudits durchzuführen. Der Distributor ist "
        "verpflichtet, die Auditoren zu empfangen und alle relevanten Unterlagen und Informationen "
        "bereitzustellen.",
        "Festgestellte Abweichungen sind in einem CAPA-Plan (Corrective and Preventive Action) "
        "festzuhalten, der innerhalb von 30 Tagen nach Audit-Abschluss vorzulegen ist.",
        "Der Distributor verpflichtet sich, Schulungen des Herstellers für das eigene Vertriebspersonal "
        "zu ermöglichen und sicherzustellen, dass Mitarbeiter, die Medizinprodukte verkaufen, "
        "über angemessene medizinische und technische Kenntnisse verfügen (Art. 14 Abs. 2 MDR).",
    ])
    add_clause(doc, 9, "Geistiges Eigentum und Marken", [
        "Der Hersteller gewährt dem Distributor eine nicht-ausschließliche, nicht-übertragbare "
        "Lizenz zur Nutzung der Marken, Logos und sonstigen Kennzeichen des Herstellers "
        "ausschließlich für den Vertrieb der Produkte im Vertragsgebiet.",
        "Der Distributor ist nicht berechtigt, Änderungen an Produkten, Verpackungen, Etiketten "
        "oder Marketingmaterialien vorzunehmen, es sei denn, der Hersteller hat schriftlich "
        "zugestimmt. Das gilt insbesondere für Übersetzungen der Gebrauchsanweisung (IFU).",
        "Alle Rechte des geistigen Eigentums verbleiben beim Hersteller. Mit Beendigung des "
        "Vertrags erlöschen alle Lizenzrechte des Distributors.",
    ])
    add_clause(doc, 10, "Vertraulichkeit", [
        "Beide Parteien verpflichten sich, alle vertraulichen Informationen der anderen Partei "
        "streng vertraulich zu behandeln und Dritten gegenüber nicht zu offenbaren. Vertraulich "
        "sind insbesondere: Preislisten, Produktentwicklungen, klinische Daten, Kundenlisten "
        "und Geschäftsstrategien.",
        "Die Vertraulichkeitspflicht gilt für einen Zeitraum von 5 Jahren nach Beendigung "
        "dieses Vertrags fort.",
        "Von der Vertraulichkeitspflicht ausgenommen sind Informationen, die (a) allgemein "
        "bekannt sind, (b) ohne Verletzung dieses Vertrags bekannt geworden sind, oder "
        "(c) von einem Dritten ohne Vertraulichkeitsverpflichtung mitgeteilt wurden.",
    ])
    add_clause(doc, 11, "Laufzeit und Kündigung", [
        "Dieser Vertrag wird für eine Laufzeit von 3 Jahren ab Unterzeichnung geschlossen. "
        "Er verlängert sich automatisch um jeweils 1 Jahr, sofern er nicht von einer Partei "
        "mit einer Frist von 6 Monaten zum Vertragsende schriftlich gekündigt wird.",
        "Das Recht zur außerordentlichen Kündigung aus wichtigem Grund bleibt unberührt. "
        "Wichtige Gründe sind insbesondere: (a) Insolvenz einer Partei; (b) schwerwiegende "
        "Verletzung regulatorischer Pflichten; (c) wiederholte Verletzung von Zahlungspflichten; "
        "(d) Verletzung des Wettbewerbsverbots.",
        "Bei Beendigung des Vertrags ist der Distributor verpflichtet, auf Verlangen des "
        "Herstellers den Restbestand der Produkte zu Einkaufspreisen an den Hersteller "
        "zurückzuverkaufen oder an einen vom Hersteller benannten Dritten zu übertragen.",
        "Handelsvertreterausgleich: Da der Distributor als Eigenhändler tätig ist, besteht "
        "kein Handelsvertreterausgleichsanspruch gemäß § 89b HGB. Diese Regelung ist "
        "ausdrücklicher Vertragsbestandteil.",
    ])
    add_clause(doc, 12, "Haftungsbeschränkung", [
        "Die Haftung jeder Partei für mittelbare Schäden, entgangenen Gewinn und Folgeschäden "
        "ist auf den Betrag der im Vertragsjahr gezahlten Vergütung begrenzt, außer bei "
        "Vorsatz oder grober Fahrlässigkeit.",
        "Die Gesamthaftung des Herstellers gegenüber dem Distributor in einem Vertragsjahr "
        f"ist auf EUR 5.000.000 begrenzt.",
        "Diese Haftungsbeschränkung gilt nicht für Ansprüche aus dem Produkthaftungsgesetz "
        "(ProdHaftG) oder anderen zwingenden gesetzlichen Regelungen.",
    ])
    add_clause(doc, 13, "Schlussbestimmungen", [
        "Änderungen und Ergänzungen dieses Vertrags bedürfen der Schriftform.",
        "Dieser Vertrag unterliegt deutschem Recht unter Ausschluss des UN-Kaufrechts (CISG).",
        f"Gerichtsstand für alle Streitigkeiten aus diesem Vertrag ist {C['city']}.",
        "Sollten einzelne Bestimmungen dieses Vertrags unwirksam sein, bleibt die Wirksamkeit "
        "des übrigen Vertrags unberührt. Anlagen 1 (Produktliste) und 2 (Preisliste) sind "
        "integraler Bestandteil dieses Vertrags." + missing_exhibit_note(),
    ])
    doc.add_paragraph(
        f"\nMünchen, den {ds(2022, 1, 15)}\n\n"
        f"_______________________\t\t_______________________\n"
        f"{C['name']}\t\t\t{dist_name}\n"
        f"(Hersteller)\t\t\t(Distributor)"
    )
    save_docx(doc, FOLDER_SALES, f"DIST_{dist_short}_Distributionsvertrag.docx")


def _gen_dist_agreement_en(dist_name, dist_short, territory):
    """English version for international distributor"""
    doc = Document()
    make_docx_header(doc, f"Distribution Agreement – {dist_name}",
                     f"Sentavia Precision GmbH and {dist_name}", confidential=True)
    add_clause(doc, 1, "Appointment and Territory", [
        f"Sentavia Precision GmbH (\"Manufacturer\") hereby appoints {dist_name} (\"Distributor\") "
        f"as exclusive distributor for the Products listed in Schedule 1 in the territory of {territory}.",
        "Distributor accepts this appointment and agrees to use commercially reasonable efforts to "
        "promote and sell the Products throughout the Territory.",
    ])
    add_clause(doc, 2, "Regulatory Compliance", [
        "Distributor represents that it holds all licenses, permits, and registrations required "
        "under applicable law to import, store, and distribute the Products in the Territory.",
        "Distributor shall comply with all applicable local regulations regarding the distribution "
        "of medical devices, including but not limited to EU MDR 2017/745 distributor obligations.",
        "Distributor shall maintain a traceability system for all Products distributed and shall "
        "provide Manufacturer with full traceability records within 48 hours of request.",
    ])
    add_clause(doc, 3, "Minimum Purchase Obligations", [
        f"Distributor commits to minimum annual purchases of EUR {random.randint(2,5)*1_000_000:,} "
        "in Year 1, EUR {random.randint(3,6)*1_000_000:,} in Year 2, and "
        "EUR {random.randint(4,8)*1_000_000:,} in Year 3.",
    ])
    add_clause(doc, 4, "Governing Law", [
        "This Agreement shall be governed by and construed in accordance with German law. "
        "The United Nations Convention on Contracts for the International Sale of Goods (CISG) "
        "shall not apply.",
        "All disputes shall be submitted to binding arbitration under DIS rules in Munich.",
    ])
    save_docx(doc, FOLDER_SALES, f"DIST_{dist_short}_Distribution_Agreement_EN.docx")


def gen_hospital_agreement(hosp_name, hosp_short):
    doc = Document()
    make_docx_header(doc, f"Rahmenliefervertrag – {hosp_name}",
                     f"{C['name']} und {hosp_name}", confidential=True)
    add_clause(doc, 1, "Vertragsgegenstand", [
        f"Dieser Rahmenliefervertrag regelt die Lieferung von Medizinprodukten der {C['name']} "
        f"an die {hosp_name}. Er gilt für alle Bestellungen im Gültigkeitszeitraum und ersetzt "
        "alle vorherigen Einzelvereinbarungen.",
        f"Umfasste Produkte: {C['prod1_name']} ({C['prod1_class']}), {C['prod2_name']} "
        f"({C['prod2_class']}).",
    ])
    add_clause(doc, 2, "Preise und Konditionen", [
        "Die vereinbarten Rahmenpreise sind in Anlage 1 festgelegt und gelten für die Laufzeit "
        "des Vertrags. Bei Abnahme von mehr als 50 Einheiten pro Produkt pro Jahr gelten "
        "gesonderte Mengenrabattstaffeln.",
        "Lieferbedingungen: DDP (Delivered Duty Paid) Lieferanschrift des Kunden. "
        "Lieferzeit: 5–7 Werktage ab Auftragseingang.",
        "Zahlungsziel: 45 Tage netto. Bei öffentlichen Einrichtungen gilt das Zahlungsgesetz "
        "(Gesetz zur Bekämpfung von Zahlungsverzug im Geschäftsverkehr).",
    ])
    add_clause(doc, 3, "Schulung und Service", [
        "Die Gesellschaft stellt für alle gelieferten Geräte eine kostenlose Erstschulung des "
        "medizinischen und technischen Personals zur Verfügung (\"Einweisung gemäß MPBetreibV\").",
        "Wartungsverträge und erweiterte Serviceoptionen werden in einem separaten "
        "Service Level Agreement (SLA) geregelt.",
    ])
    add_clause(doc, 4, "Vigilanz und PMS", [
        "Der Kunde verpflichtet sich, beobachtete schwerwiegende Vorkommnisse oder "
        "unerwünschte Ereignisse im Zusammenhang mit den gelieferten Produkten unverzüglich "
        "an die Gesellschaft zu melden.",
        "Der Kunde wirkt auf Anfrage bei Post-Market-Surveillance-Aktivitäten (PMS) mit, "
        "insbesondere bei klinischen Bewertungen (PMCF) und der Erhebung von Anwenderdaten.",
    ])
    doc.add_paragraph(
        f"\nMünchen, den {ds(2022, 6, 1)}\n\n"
        f"_______________________\t\t_______________________\n"
        f"{C['name']}\t\t\t{hosp_name}"
    )
    save_docx(doc, FOLDER_SALES, f"HOSP_{hosp_short}_Rahmenliefervertrag.docx")


def gen_sales_reports():
    for q in ["Q1", "Q2", "Q3", "Q4"]:
        rev = C["revenue_2023"] // 4
        factor = {"Q1": 0.88, "Q2": 0.95, "Q3": 1.02, "Q4": 1.15}[q]
        qrev = int(rev * factor)
        rows = [
            (C["prod1_name"], str(random.randint(180, 260)), f"{int(qrev*0.52):,}",
             str(random.randint(12, 28)), f"{random.randint(75,95)}%"),
            (C["prod2_name"], str(random.randint(280, 400)), f"{int(qrev*0.33):,}",
             str(random.randint(8, 20)), f"{random.randint(68,88)}%"),
            (C["prod3_name"], str(random.randint(4000, 8000)), f"{int(qrev*0.15):,}",
             str(random.randint(5, 15)), f"{random.randint(55,75)}%"),
        ]
        sheets = [
            ("Umsatzübersicht", ["Produkt", "Einheiten", "Umsatz EUR", "Neue Kunden", "Kundenzufriedenheit"],
             rows, [22, 12, 15, 14, 20]),
        ]
        make_xlsx_doc(FOLDER_SALES, f"SALES_{q}_2023_Umsatzbericht.xlsx",
                      f"Vertriebsbericht {q} 2023", sheets)


def gen_customer_complaints_log():
    rows = []
    for i in range(1, 41):
        d = rdate(2023)
        rows.append((
            f"CC-2023-{i:04d}",
            ds(d.year, d.month, d.day),
            random.choice([C["prod1_name"], C["prod2_name"], C["prod3_name"]]),
            random.choice([C["hosp1"], C["hosp2"], C["hosp3"], C["hosp4"],
                           "Universitätsklinikum Frankfurt", "Helios Klinikum Berlin"]),
            random.choice(["Kalibrierungsproblem", "Softwarefehler", "Verbindungsproblem",
                           "Falsches Testergebnis", "Etikettierungsfehler", "Verpackungsschaden",
                           "Bedienungsfehler (Anwender)", "Produktversagen", "Lieferproblem"]),
            random.choice(["Offen", "In Bearbeitung", "Abgeschlossen", "Abgeschlossen", "Abgeschlossen"]),
            random.choice(["Ja – MDR Art. 87", "Nein", "Nein", "Nein", "In Prüfung"]),
        ))
    sheets = [
        ("Beschwerdeprotokoll 2023", ["Beschwerde-Nr.", "Datum", "Produkt", "Kunde",
                                       "Beschwerdeart", "Status", "Meldepflicht"],
         rows, [16, 14, 20, 32, 28, 16, 18]),
    ]
    make_xlsx_doc(FOLDER_SALES, "SALES_Kundenbeschwerden_2023.xlsx",
                  "Kundenbeschwerden-Log 2023", sheets)


def gen_pricing_schedule(dist_name, dist_short):
    rows = [
        (C["prod1_name"], "CS-PRO-001-BASE", "Basisgerät", "1", f"{random.randint(18,25)*1000:,}", "15%", f"{random.randint(15,21)*1000:,}"),
        (C["prod1_name"], "CS-PRO-001-PRB", "Probenentnahmeset (10er)", "10", f"{random.randint(80,120)*10:,}", "20%", f"{random.randint(65,95)*10:,}"),
        (C["prod2_name"], "OF-NAV-002-SYS", "Navigationssystem Komplett", "1", f"{random.randint(28,38)*1000:,}", "18%", f"{random.randint(23,31)*1000:,}"),
        (C["prod2_name"], "OF-NAV-002-REF", "Referenzmarker Set", "10", f"{random.randint(300,500):,}", "22%", f"{random.randint(240,390):,}"),
        (C["prod3_name"], "DK-SF-003-KIT", "Testkit 24er-Box", "24", f"{random.randint(180,280):,}", "25%", f"{random.randint(140,210):,}"),
    ]
    sheets = [
        ("Preisliste", ["Produkt", "Artikel-Nr.", "Bezeichnung", "VE", "Listenpreis EUR",
                         "Distributor-Rabatt", "Netto-Einkaufspreis EUR"],
         rows, [20, 18, 28, 6, 16, 16, 22]),
    ]
    make_xlsx_doc(FOLDER_SALES, f"PREISE_{dist_short}_Preisliste_2024.xlsx",
                  f"Preisliste {dist_name} – gültig ab 01.01.2024", sheets)


def generate_04_vertrieb():
    print("  Generating 04_Vertrieb_Distributoren...")
    gen_distribution_agreement(C["dist1"], "SH", "DACH (Deutschland, Österreich, Schweiz)")
    gen_distribution_agreement(C["dist2"], "FMC", "Benelux, Frankreich, Nordeuropa")
    gen_distribution_agreement(C["dist3"], "BB", "Südeuropa (Italien, Spanien, Portugal)", lang="en")
    for short, hosp in [("CHR", C["hosp1"]), ("UKE", C["hosp2"]),
                        ("LMU", C["hosp3"]), ("UKH", C["hosp4"])]:
        gen_hospital_agreement(hosp, short)
    gen_sales_reports()
    gen_customer_complaints_log()
    for dist, short in [(C["dist1"], "SH"), (C["dist2"], "FMC"), (C["dist3"], "BB")]:
        gen_pricing_schedule(dist, short)
    # Tender documents
    for i, title in enumerate([
        "Ausschreibung Kardiologie-Systeme Charité 2023",
        "Europäische Ausschreibung IVD-Systeme Bundesbeschaffungsamt",
        "Vergabeverfahren UKE Hamburg Navigationssysteme",
    ], 1):
        sections = [
            ("Angebotsunterlagen", f"Ausschreibungsreferenz: TENDER-BTP-2023-{i:03d}\nTitel: {title}"),
            ("Produktangebot", [
                ["Produkt", "Menge", "Einheitspreis EUR", "Gesamtpreis EUR"],
                [C["prod1_name"], str(random.randint(5,15)), f"{random.randint(18,22)*1000:,}",
                 f"{random.randint(90,330)*1000:,}"],
            ]),
            ("Serviceangebot", "5 Jahre Vollwartung inklusive. Schulungen kostenlos. SLA: 4h Response-Zeit."),
        ]
        save_pdf(FOLDER_SALES, f"AUSSCH_{i:02d}_Angebot_{sfn(title[:25])}.pdf", title, sections)
    # Product training certificates
    for hosp_short, hosp in [("CHR", C["hosp1"]), ("UKE", C["hosp2"])]:
        for prod in [C["prod1_name"], C["prod2_name"]]:
            sections = [
                ("Schulungszertifikat",
                 f"Hiermit wird bestätigt, dass das medizinische Personal der {hosp} am "
                 f"{ds(2023, random.randint(1,12), random.randint(1,28))} erfolgreich an der "
                 f"Produktschulung für {prod} teilgenommen hat. Die Schulung umfasste "
                 "Inbetriebnahme, Bedienung, Reinigung und Fehlerdiagnose."),
                ("Teilnehmer", [
                    ["Name", "Funktion", "Unterschrift"],
                    [f"Dr. med. {random.choice(['Weber', 'Müller', 'Schmidt'])}", "Oberärztin/-arzt", "___"],
                    [f"MTRA {random.choice(['Bauer', 'Hoffmann', 'Klein'])}", "Med.-techn. Radiologieassistenz", "___"],
                ]),
            ]
            save_pdf(FOLDER_SALES, f"TRAINING_{hosp_short}_{sfn(prod[:12])}_Schulungszertifikat.pdf",
                     f"Schulungszertifikat – {prod}", sections)


# ══════════════════════════════════════════════════════════════════════
# FOLDER 05 – EINKAUF / LIEFERANTEN
# ══════════════════════════════════════════════════════════════════════
FOLDER_PURCH = "05_Einkauf_Lieferanten"


def gen_supply_agreement(lief_name, lief_short, component, critical=True):
    doc = Document()
    make_docx_header(doc, f"Liefervertrag und Qualitätssicherungsvereinbarung",
                     f"{C['name']} und {lief_name}", confidential=True)
    doc.add_paragraph(
        f"LIEFERVERTRAG UND QUALITÄTSSICHERUNGSVEREINBARUNG (QSV)\n\nzwischen\n\n"
        f"{C['name']}, {C['full_address']}, {C['hrb']}\n– nachfolgend \"Auftraggeber\" (AG) –\n\n"
        f"und\n\n{lief_name}\n– nachfolgend \"Lieferant\" (LI) –\n\n"
        f"Gegenstand dieser Vereinbarung ist die Lieferung von {component} sowie die "
        "Regelung der damit verbundenen Qualitätssicherungsanforderungen.\n\n"
        "Präambel: Der AG betreibt ein zertifiziertes Qualitätsmanagementsystem nach "
        f"{C['iso_cert']} und ist nach EU-MDR 2017/745 sowie IVDR 2017/746 als Hersteller "
        "von Medizinprodukten zertifiziert. Der LI ist ein Lieferant, dessen Produkte "
        "direkt oder indirekt in Medizinprodukte einfließen. Beide Parteien erkennen die "
        "besondere Bedeutung von Qualität, Sicherheit und Rückverfolgbarkeit in der "
        "Medizintechnik an."
    )
    add_clause(doc, 1, "Vertragsgegenstand und Lieferbedingungen", [
        f"Der LI liefert dem AG die in Anlage 1 spezifizierten {component} in der dort "
        "beschriebenen Qualität und Menge. Die Lieferungen erfolgen gemäß den Abrufbestellungen "
        "des AG zu den vereinbarten Konditionen.",
        "Lieferbedingungen: DDP (Incoterms® 2020) an die Fertigungsadresse des AG. "
        f"Lieferzeit: {random.randint(8,15)} Werktage nach Bestelleingang.",
        "Der LI ist verpflichtet, den AG unverzüglich zu informieren, wenn er Lieferengpässe, "
        "Kapazitätsprobleme oder Qualitätsprobleme erkennt, die die vereinbarte Lieferung "
        "beeinträchtigen könnten. Eine Vorwarnzeit von mindestens 8 Wochen ist einzuhalten.",
    ])
    add_clause(doc, 2, "Qualitätsanforderungen", [
        f"Der LI muss ein QMS gemäß {C['iso_cert']} unterhalten und dies durch eine gültige "
        "Zertifizierung nachweisen. Zertifikate sind dem AG vor Beginn der Liefertätigkeit "
        "und bei jeder Erneuerung unaufgefordert zu übersenden.",
        "Die Produkte müssen den in Anlage 1 festgelegten Spezifikationen, Zeichnungen und "
        "Normen entsprechen. Abweichungen bedürfen einer schriftlichen Genehmigung des AG "
        "(Concession Request).",
        "Der LI führt für jede Lieferung eine Produktionsprüfung durch und stellt dem AG "
        "Prüfzertifikate (CoC – Certificate of Conformance) bei. Bei kritischen Komponenten "
        "werden zusätzliche Materialzertifikate (CoA – Certificate of Analysis) verlangt.",
        "Change Notification: Der LI verpflichtet sich, den AG mindestens 6 Monate vor "
        "jeder geplanten Änderung (Produkt, Prozess, Standort, Material) schriftlich zu "
        "informieren (Supplier Change Notification). Änderungen ohne Zustimmung des AG "
        "können zur Kündigung dieses Vertrags führen.",
    ])
    add_clause(doc, 3, "Rückverfolgbarkeit", [
        "Der LI stellt sicher, dass alle gelieferten Produkte eindeutig mit Lot-Nummern oder "
        "Seriennummern gekennzeichnet sind, die eine vollständige Rückverfolgbarkeit bis auf "
        "Rohmaterialien und Vorlieferantenebene ermöglichen.",
        "Rückverfolgbarkeitsdaten und Produktionsdokumentationen sind vom LI für mindestens "
        "15 Jahre nach dem letzten Inverkehrbringen des zugehörigen Medizinprodukts aufzubewahren. "
        "Bei Einstellung der Geschäftstätigkeit sind diese Daten an den AG zu übertragen.",
    ])
    add_clause(doc, 4, "Audits und Inspektionen", [
        "Der AG ist berechtigt, beim LI Lieferantenaudits gemäß ISO 19011 durchzuführen, "
        "einmal jährlich angekündigt und bei begründetem Verdacht auch unangekündigt. "
        "Der LI kooperiert vollständig und stellt alle erforderlichen Unterlagen bereit.",
        "Auditfeststellungen sind in einem gemeinsamen CAPA-Plan zu dokumentieren. "
        "Der LI ist verpflichtet, Maßnahmen zu kritischen Findings innerhalb von 30 Tagen "
        "und zu anderen Findings innerhalb von 90 Tagen umzusetzen.",
    ])
    add_clause(doc, 5, "Geheimhaltung und Exklusivität", [
        "Der LI verpflichtet sich, alle technischen Informationen und Spezifikationen des "
        "AG streng vertraulich zu behandeln und ausschließlich für die Erfüllung dieses "
        "Vertrags zu verwenden.",
        "Der LI entwickelt auf Basis der technischen Informationen des AG keine eigenen "
        "Produkte und beliefert keine Wettbewerber des AG mit vergleichbaren Produkten, "
        "ohne vorherige schriftliche Zustimmung des AG.",
    ])
    add_clause(doc, 6, "Haftung und Produkthaftung", [
        "Der LI haftet für alle Schäden, die durch mangelhafte Lieferungen entstehen, "
        "einschließlich der Kosten von Produktrückrufen, soweit der Mangel beim LI liegt. "
        "Die Haftung umfasst auch Folgekosten aus der EU-Produkthaftungsrichtlinie.",
        "Der LI unterhält eine Produkthaftpflichtversicherung mit einer Deckungssumme von "
        "mindestens EUR 5.000.000 je Schadensfall und legt dem AG jährlich den aktuellen "
        "Versicherungsnachweis vor.",
    ])
    doc.add_paragraph(
        f"\nMünchen, den {ds(2022, 6, 1)}\n\n"
        f"_______________________\t\t_______________________\n"
        f"{C['name']}\t\t\t{lief_name}"
    )
    save_docx(doc, FOLDER_PURCH, f"LIEF_{lief_short}_Liefervertrag_QSV_{sfn(component[:20])}.docx")


def gen_supplier_qualification(lief_name, lief_short):
    sections = [
        ("Lieferantenqualifikation – Zusammenfassung",
         f"Lieferant: {lief_name}\nQualifikationsdatum: {ds(2023, random.randint(1,12), random.randint(1,28))}\n"
         f"Qualifiziert durch: {C['qra']}, Sentavia Precision GmbH\nQualifikationsstatus: FREIGEGEBEN"),
        ("Qualifikationskriterien", [
            ["Kriterium", "Anforderung", "Ergebnis", "Status"],
            ["ISO 13485-Zertifikat", "Gültig, akkreditierte Stelle", "Vorliegend, gültig bis 2025", "OK"],
            ["Qualitätsmanagementsystem", "Dokumentiert, wirksam", "Audit bestanden", "OK"],
            ["Finanzielle Stabilität", "Keine Insolvenz, positive BWA", "Geprüft", "OK"],
            ["Technische Kompetenz", "Spezifikationserfüllung nachgewiesen", "Bemusterung bestanden", "OK"],
            ["Rückverfolgbarkeit", "Lot-/Seriennummernverfolgung", "System vorhanden", "OK"],
            ["Liefertreue (3 Jahre)", "> 95 %", f"{random.randint(95,99)} %", "OK"],
            ["Zuverlässigkeit (ppm)", "< 500 ppm Fehler", f"{random.randint(50,400)} ppm", "OK"],
        ]),
        ("Auditbericht Kurzfassung",
         f"Das Lieferantenaudit bei {lief_name} wurde am {ds(2023, 4, 18)} durch "
         f"{C['qra']} und Sarah Krause durchgeführt. Der Lieferant erfüllt alle "
         "wesentlichen Anforderungen der ISO 13485:2016 und der MDR-Lieferantenanforderungen. "
         "Kleinere Verbesserungspotenziale wurden im CAPA-Plan festgehalten."),
    ]
    save_pdf(FOLDER_PURCH, f"QUAL_{lief_short}_Lieferantenqualifikation.pdf",
             f"Lieferantenqualifikation – {lief_name}", sections)


def generate_05_einkauf():
    print("  Generating 05_Einkauf_Lieferanten...")
    lief_data = [
        (C["lief1"], "TI", "Mikroprozessoren und Signalverarbeitungs-ICs (Typ TMS320 Serie)", True),
        (C["lief2"], "SE", "Feuchte- und Temperatursensoren (SHT-Serie)", True),
        (C["lief3"], "MT", "Medizinische Verbrauchsmaterialien und Sterilverpackungen", True),
        (C["lief4"], "SM", "Etiketten, Kennzeichnungsmaterial und UDI-Etiketten", False),
    ]
    for lief_name, short, component, crit in lief_data:
        gen_supply_agreement(lief_name, short, component, crit)
        gen_supplier_qualification(lief_name, short)
    # Purchase orders
    for i in range(1, 16):
        lief = random.choice(lief_data)
        d = rdate(2023)
        sections = [
            ("Bestelldaten",
             f"Bestellnummer: PO-2023-{i:04d}\nDatum: {ds(d.year, d.month, d.day)}\n"
             f"Lieferant: {lief[0]}\nLiefertermin: {ds(d.year, d.month, min(d.day+14, 28))}\n"
             f"Zahlungsbedingungen: 30 Tage netto"),
            ("Bestellpositionen", [
                ["Pos.", "Artikel-Nr.", "Bezeichnung", "Menge", "Einheit", "Einzelpreis EUR", "Gesamtpreis EUR"],
                ["1", f"ART-{random.randint(1000,9999)}", lief[2][:30],
                 str(random.randint(100, 1000)), "Stk",
                 f"{random.randint(10,500):,}", f"{random.randint(5000,250000):,}"],
            ]),
        ]
        save_pdf(FOLDER_PURCH, f"PO_2023_{i:04d}_Bestellung_{lief[1]}.pdf",
                 f"Bestellung PO-2023-{i:04d}", sections)
    # Critical component audit
    for lief_name, short, component, crit in lief_data[:2]:
        sections = [
            ("Critical Component Audit Report",
             f"Component: {component}\nSupplier: {lief_name}\n"
             f"Audit Date: {ds(2023, 9, random.randint(10,25))}\n"
             f"Auditor: {C['qra']}"),
            ("Findings", [
                ["Finding", "Severity", "Corrective Action", "Due Date"],
                ["Calibration records incomplete", "Minor", "Update SOP, re-train staff",
                 ds(2023, 11, 1)],
                ["Supplier change notification delay", "Observation",
                 "Revise supplier agreement", ds(2023, 12, 1)],
            ]),
            ("Conclusion",
             f"The supplier {lief_name} demonstrates adequate quality management for critical "
             "medical device components. Two minor findings require corrective action. "
             "Overall qualification status: APPROVED WITH CONDITIONS."),
        ]
        save_pdf(FOLDER_PURCH, f"AUDIT_{short}_Critical_Component_Audit_2023.pdf",
                 f"Critical Component Audit – {lief_name}", sections)


# ══════════════════════════════════════════════════════════════════════
# FOLDER 06 – REGULATORISCH MDR (THE BIG ONE)
# ══════════════════════════════════════════════════════════════════════
FOLDER_REG = "06_Regulatorisch_MDR"


def gen_technical_file_section(prod_name, prod_class, prod_ref, section_name, section_num):
    sections = [
        ("Dokumentenhistorie", [
            ["Version", "Datum", "Autor", "Änderungsbeschreibung"],
            ["1.0", ds(2020, 6, 15), C["qra"], "Erstfassung"],
            ["1.1", ds(2021, 3, 10), "Markus Bauer", "Aktualisierung nach MDR-Artikel 61 Review"],
            ["2.0", ds(2022, 8, 22), C["qra"], "Vollständige Überarbeitung nach MDR-Transition"],
            ["2.1", ds(2023, 5, 18), "Markus Bauer", "Aktualisierung PMCF-Daten"],
        ]),
        ("Geltungsbereich",
         f"Diese technische Dokumentation gemäß Anhang II der EU-MDR 2017/745 beschreibt "
         f"{section_name} des Medizinprodukts {prod_name}.\n\n"
         f"Produkt: {prod_name}\nProduktklassifizierung: {prod_class}\n"
         f"UDI-DI: {prod_ref}\nCE-Zertifikat: {C.get('prod1_ce', 'CE 0123-BTP-Ostevo-2021')}\n"
         f"Benannte Stelle: {C['nb']} (Kennnummer {C['nb_id']})\n"
         f"Hersteller: {C['name']}, {C['full_address']}"),
        ("Normative Referenzen", [
            ["Norm/Richtlinie", "Titel", "Anwendbar seit"],
            ["MDR 2017/745", "EU-Medizinprodukteverordnung", "26. Mai 2021"],
            ["ISO 13485:2016", "Qualitätsmanagementsysteme für Medizinprodukte", "Fortlaufend"],
            ["ISO 14971:2019", "Risikomanagement für Medizinprodukte", "Fortlaufend"],
            ["IEC 62304:2006+A1:2015", "Software-Lebenszyklusprozesse für Medizinprodukte", "Fortlaufend"],
            ["IEC 60601-1:2005+A1:2012", "Elektrische Sicherheit in der Medizin", "Fortlaufend"],
            ["ISO 10993-1:2018", "Biologische Bewertung von Medizinprodukten", "Fortlaufend"],
            ["EN ISO 15223-1:2021", "Symbole für Medizinprodukte", "Fortlaufend"],
        ]),
    ]
    fname = f"TD_{sfn(prod_name[:12])}_{section_num:02d}_{sfn(section_name[:20])}.pdf"
    save_pdf(FOLDER_REG, fname, f"Technische Dokumentation – {prod_name} – {section_name}", sections)


def gen_declaration_of_conformity(prod_name, prod_class, prod_ref):
    sections = [
        ("EU-KONFORMITÄTSERKLÄRUNG gemäß Artikel 19 MDR 2017/745", None),
        ("Hersteller",
         f"{C['name']}\n{C['full_address']}\n"
         f"Telefon: {C['phone']}\nE-Mail: {C['email']}\n"
         f"EUDAMED-Registrierungsnummer: {C['eudamed_reg']}"),
        ("Produkt",
         f"Produktbezeichnung: {prod_name}\n"
         f"Produktklassifizierung: {prod_class}\n"
         f"UDI-DI (Basic): {prod_ref}\n"
         f"Modell-/Typenreihe: {prod_name} Produktfamilie, alle Softwareversionen ≥ 3.0"),
        ("Erklärung",
         f"Hiermit erklärt {C['name']} in alleiniger Verantwortung, dass das oben beschriebene "
         "Medizinprodukt bei bestimmungsgemäßem Einsatz mit den Anforderungen der Verordnung "
         "(EU) 2017/745 über Medizinprodukte (MDR) übereinstimmt. "
         "Das Produkt entspricht den grundlegenden Sicherheits- und Leistungsanforderungen "
         "gemäß Anhang I der MDR. Die technische Dokumentation wurde gemäß Anhang II und III "
         "der MDR erstellt und wird aktuell gehalten.\n\n"
         f"Das Konformitätsbewertungsverfahren wurde gemäß Anhang IX der MDR durchgeführt "
         f"unter Einbeziehung der Benannten Stelle {C['nb']} (Kennnummer {C['nb_id']}).\n\n"
         f"Zertifikat-Nr.: {C['nb_cert1']}"),
        ("Angewandte harmonisierte Normen", [
            ["Norm", "Ausgabe", "Bereich"],
            ["ISO 13485", "2016", "QMS"],
            ["ISO 14971", "2019", "Risikomanagement"],
            ["IEC 60601-1", "2005+A1", "Elektrische Sicherheit"],
            ["IEC 62304", "2006+A1", "Software"],
            ["ISO 10993-1", "2018", "Biokompatibilität"],
        ]),
        ("Unterzeichnung",
         f"Ausgestellt in: {C['city']}\n"
         f"Datum: {ds(2022, 9, 15)}\n\n"
         f"_______________________\n"
         f"{C['qra']}\n"
         "Head of Quality & Regulatory Affairs\n"
         f"{C['name']}"),
    ]
    fname = f"DOC_{sfn(prod_name[:12])}_EU_Konformitaetserklaerung.pdf"
    save_pdf(FOLDER_REG, fname, f"EU-Konformitätserklärung – {prod_name}", sections)


def gen_pms_plan(prod_name, prod_short):
    sections = [
        ("Zusammenfassung",
         f"Dieses Post-Market Surveillance (PMS) Plan-Dokument wurde für das Medizinprodukt "
         f"{prod_name} gemäß Artikel 84 der MDR 2017/745 und dem begleitenden Leitfaden MDCG 2020-7 "
         f"erstellt. Der Plan beschreibt die systematischen Verfahren zur Erhebung, Verarbeitung "
         "und Analyse von Post-Market-Daten."),
        ("PMS-Quellen und Methoden", [
            ["Datenquelle", "Frequenz", "Verantwortlich", "Instrument"],
            ["Kundenbeschwerden (CMS)", "Kontinuierlich", "QA-Team", C["qms_tool"]],
            ["Vigilanz-Datenbank BfArM", "Monatlich", "Regulatory Affairs", "Manuelle Recherche"],
            ["EUDAMED-Datenbank", "Quartalsweise", C["qra"], "EUDAMED Webportal"],
            ["Wissenschaftliche Literatur", "Halbjährlich", "Medical Affairs", "PubMed, Embase"],
            ["Post-Market Clinical Follow-Up", "Jährlich", C["cmo"], C["clinical_db"]],
            ["Klinischer Nutzen (Clinical Benefit)", "Jährlich", "Clinical Affairs", "Interne Studie"],
            ["Soziale Medien / Nutzerfeedback", "Monatlich", "Regulatory Affairs", "Monitoring-Tool"],
        ]),
        ("Auswertungsschwellen (Thresholds)",
         "Folgende Signale lösen eine vertiefte Analyse und ggf. eine Aktualisierung der "
         "Risikobewertung aus:\n"
         "• Anstieg der Complaint-Rate um > 20 % gegenüber Vorjahr\n"
         "• ≥ 3 gleichartige schwerwiegende Ereignisse innerhalb von 6 Monaten\n"
         "• Neue wissenschaftliche Erkenntnisse, die die Sicherheitsbewertung in Frage stellen\n"
         "• Regulatorische Maßnahmen von Behörden anderer Länder (FDA, TGA, Health Canada)"),
        ("PMS-Bericht und PSUR",
         f"Der PMS-Bericht (PMSR gemäß Art. 85 MDR) für {prod_name} ist jährlich zu erstellen. "
         "Er umfasst eine Auswertung aller PMS-Daten, eine aktualisierte Nutzen-Risiko-Bewertung "
         "und eine Zusammenfassung aller Vigilanzberichte.\n\n"
         "Der Periodic Safety Update Report (PSUR gemäß Art. 86 MDR) ist bei Klasse IIb-Produkten "
         "mindestens alle zwei Jahre zu erstellen und der Benannten Stelle vorzulegen."),
    ]
    save_pdf(FOLDER_REG, f"PMS_{prod_short}_Plan.pdf",
             f"Post-Market Surveillance Plan – {prod_name}", sections)


def gen_pms_report(prod_name, prod_short, year):
    sections = [
        ("Executive Summary",
         f"Dieser Post-Market Surveillance Report (PMSR) für {prod_name} umfasst den Zeitraum "
         f"1. Januar {year} bis 31. Dezember {year}. Im Berichtszeitraum wurden "
         f"{random.randint(2,8)} Kundenbeschwerden registriert, davon {random.randint(0,2)} "
         "als schwerwiegende Ereignisse (Serious Incidents) klassifiziert. Die Nutzen-Risiko-Bilanz "
         "bleibt positiv. Keine sicherheitskritischen Signale identifiziert."),
        ("Beschwerdedaten", [
            ["Kategorie", "Anzahl", "Schweregrad", "Regulatorisch relevant"],
            ["Kalibrierungsfehler", str(random.randint(0,3)), "Niedrig", "Nein"],
            ["Software-Bug", str(random.randint(0,2)), "Mittel", "In Prüfung"],
            ["Mechanischer Defekt", str(random.randint(0,2)), "Mittel", "Nein"],
            ["Fehlalarme", str(random.randint(0,5)), "Niedrig", "Nein"],
            ["Gesamt", str(random.randint(3,10)), "—", "—"],
        ]),
        ("Vigilanzberichte",
         f"Im Berichtszeitraum {year} wurden {random.randint(0,1)} Vigilanzberichte an die "
         "Zuständige Behörde (BfArM) erstattet. Die gemeldeten Ereignisse wurden analysiert; "
         "es wurden keine systemischen Probleme identifiziert."),
        ("Nutzen-Risiko-Bewertung",
         f"Die Gesamtbewertung des Nutzen-Risiko-Verhältnisses für {prod_name} bleibt unverändert "
         "positiv. Der klinische Nutzen (präzise Diagnose, Behandlungsunterstützung) überwiegt "
         "die bekannten Risiken deutlich. Keine neuen, nicht akzeptablen Risiken identifiziert."),
        ("Maßnahmen",
         "Aus den PMS-Aktivitäten wurden folgende Maßnahmen abgeleitet:\n"
         f"• Aktualisierung der IFU (Instructions for Use) hinsichtlich Kalibrierungsintervalle\n"
         "• Einleitung einer PMCF-Befragung bei 50 Anwendern im Frühjahr des Folgejahres\n"
         "• Initiierung eines Software-Updates zur Verbesserung der Alarmfunktion"),
    ]
    save_pdf(FOLDER_REG, f"PMSR_{prod_short}_{year}_Bericht.pdf",
             f"Post-Market Surveillance Report – {prod_name} – {year}", sections)


def gen_pmcf_plan(prod_name, prod_short):
    sections = [
        ("Zusammenfassung",
         f"Dieser Post-Market Clinical Follow-Up (PMCF) Plan wurde für {prod_name} gemäß "
         "Artikel 83 und Anhang XIV Teil B der MDR 2017/745 sowie dem Leitfaden MDCG 2020-8 erstellt. "
         "Ziel ist die kontinuierliche Erhebung und Auswertung klinischer Daten aus dem "
         "Markt, um den klinischen Nutzen und die Sicherheit des Produkts zu bestätigen."),
        ("PMCF-Aktivitäten", [
            ["Aktivität", "Methode", "Stichprobengröße", "Zeitraum", "Verantwortlich"],
            ["Anwenderregistry", "Fragebogen / Online-Portal", "n ≥ 200", "Laufend", C["cmo"]],
            ["Klinische Studie PMCF-1", "Prospektive Kohortenstudie", "n = 120", "2023–2025", C["cmo"]],
            ["Literaturrecherche", "Systematische Suche PubMed/Embase", "Alle relevanten Artikel", "Jährlich", "Medical Affairs"],
            ["Expertenregister", "PMCF-Workshop mit KOLs", "10–15 Experten", "Jährlich", C["cmo"]],
        ]),
        ("Statistische Methodik",
         "Die Auswertung der PMCF-Daten erfolgt durch statistische Methoden gemäß ISO 14155:2020. "
         "Die Stichprobengröße wurde auf Basis einer power-kalkulierten Analyse festgelegt "
         "(Power: 80 %, Signifikanzniveau: 5 %). Konfidenzintervalle werden berichtet."),
        ("PMCF-Evaluation Report",
         f"Der PMCF Evaluation Report wird jährlich erstellt und fließt in den PMS-Bericht "
         "sowie den Periodic Safety Update Report (PSUR) ein. Der Bericht ist der Benannten "
         f"Stelle {C['nb']} im Rahmen der Überwachungsaudits vorzulegen."),
    ]
    save_pdf(FOLDER_REG, f"PMCF_{prod_short}_Plan.pdf",
             f"Post-Market Clinical Follow-Up Plan – {prod_name}", sections)


def gen_pmcf_evaluation(prod_name, prod_short, year):
    sections = [
        ("Zusammenfassung",
         f"Dieser PMCF Evaluation Report für {prod_name} wertet die PMCF-Aktivitäten des Jahres "
         f"{year} aus. Es wurden {random.randint(120,250)} Patientenfälle in die PMCF-Studie "
         "eingeschlossen. Die Ergebnisse bestätigen die klinische Leistung und Sicherheit "
         "des Produkts wie im aktuellen Clinical Evaluation Report (CER) beschrieben."),
        ("Klinische Ergebnisse", [
            ["Endpunkt", "Zielwert", "Ergebnis", "Bewertung"],
            ["Diagnostische Sensitivität", "≥ 92 %", f"{random.randint(93,98)} %", "ERFÜLLT"],
            ["Diagnostische Spezifizität", "≥ 88 %", f"{random.randint(89,96)} %", "ERFÜLLT"],
            ["Geräteversagen (Device Failure)", "< 1 %", f"{random.uniform(0.1,0.8):.1f} %", "ERFÜLLT"],
            ["Anwenderzufriedenheit", "≥ 80 % positiv", f"{random.randint(82,95)} %", "ERFÜLLT"],
        ]),
        ("Schlussfolgerung",
         f"Die PMCF-Daten bestätigen das positive Nutzen-Risiko-Verhältnis von {prod_name}. "
         "Der CER wird auf Basis der neuen Daten aktualisiert. Keine sicherheitskritischen "
         "Erkenntnisse. PMCF-Plan für das Folgejahr bleibt unverändert gültig."),
    ]
    save_pdf(FOLDER_REG, f"PMCF_{prod_short}_{year}_Evaluation_Report.pdf",
             f"PMCF Evaluation Report – {prod_name} – {year}", sections)


def gen_sscp(prod_name, prod_class, prod_ref, prod_short):
    sections = [
        ("Zusammenfassung für Sicherheit und klinische Leistung (SSCP)",
         f"Dieses Dokument ist eine Zusammenfassung der Sicherheits- und Leistungsdaten für "
         f"{prod_name} gemäß Artikel 32 der MDR 2017/745."),
        ("Produktidentifikation",
         f"Produktbezeichnung: {prod_name}\nProduktklasse: {prod_class}\n"
         f"UDI-DI: {prod_ref}\nHersteller: {C['name']}, {C['full_address']}"),
        ("Klinische Evidenz",
         f"Die klinische Evidenz für {prod_name} basiert auf folgenden Quellen: "
         f"(1) Systematische Literaturrecherche in PubMed und Embase (≥ 150 relevante Publikationen); "
         f"(2) Klinische Prüfung gemäß ISO 14155 mit n = {random.randint(80,200)} Probanden; "
         f"(3) PMCF-Daten aus {random.randint(1,3)} Jahren Post-Market-Beobachtung. "
         f"Die Studien belegen eine diagnostische Genauigkeit von ≥ 93 % und ein "
         "exzellentes Sicherheitsprofil ohne schwerwiegende gerätebedingte Ereignisse."),
        ("Sicherheits- und Leistungsinformationen",
         f"Zweckbestimmung: {prod_name} ist für die nicht-invasive Anwendung durch geschultes "
         "medizinisches Fachpersonal bestimmt. Das Produkt dient der Unterstützung klinischer "
         "Entscheidungen und ersetzt keine ärztliche Diagnose oder Behandlung.\n\n"
         "Restrisiken: Das Restrisiko nach Implementierung aller Risikominderungsmaßnahmen "
         "ist akzeptabel im Sinne der ISO 14971:2019. Bekannte Restrisiken: "
         "(a) Möglichkeit falsch-negativer Ergebnisse bei atypischen Patientenphysiologien; "
         "(b) Interferenz durch starke elektromagnetische Felder gemäß IEC 60601-1-2."),
        ("Kontraindikationen und Warnhinweise",
         "Kontraindikationen: Patienten mit aktiven implantierten Geräten (Herzschrittmacher). "
         "Warnhinweise: Nicht in MRT-Umgebung verwenden. Regelmäßige Kalibrierung erforderlich. "
         "Nur für die in der IFU beschriebenen Indikationen geeignet."),
    ]
    save_pdf(FOLDER_REG, f"SSCP_{prod_short}_Summary_Safety_Clinical_Performance.pdf",
             f"Summary of Safety and Clinical Performance (SSCP) – {prod_name}", sections)


def gen_vigilance_report(prod_name, prod_short, num):
    d = rdate(2023)
    sections = [
        ("Vigilanzbericht an das BfArM gemäß §§ 29, 30 MPG / Art. 87 MDR",
         f"Meldende Stelle: {C['name']}, {C['full_address']}\n"
         f"EUDAMED-SRN: {C['eudamed_reg']}\n"
         f"Vorgangsnummer: VIG-{C['short']}-2023-{num:04d}\n"
         f"Datum der Erstmeldung: {ds(d.year, d.month, d.day)}\n"
         f"Meldungstyp: {'Erstmeldung' if num % 3 != 0 else 'Folgemeldung / Abschlussbericht'}"),
        ("Produktidentifikation",
         f"Produkt: {prod_name}\nUDI-DI: {C.get(f'prod{1}_ref', prod_short)}\n"
         f"Lot-Nr.: LOT-{random.randint(2023001, 2023999):07d}\n"
         f"Seriennummer: SN-{random.randint(100000,999999)}"),
        ("Ereignisbeschreibung",
         f"Art des Ereignisses: {random.choice(['Fehlmessung', 'Softwarefehler', 'Geräteverlust', 'Mechanischer Defekt'])}\n"
         f"Schweregrad: {random.choice(['Klasse I (leicht)', 'Klasse II (moderat)'])}\n"
         f"Ort des Ereignisses: {random.choice([C['hosp1'], C['hosp2'], C['hosp3']])}\n\n"
         "Beschreibung: Im Rahmen der klinischen Anwendung wurde eine Abweichung beobachtet, "
         "die zu einer verzögerten Diagnosestellung führte. Der Patient wurde adäquat medizinisch "
         "versorgt; es wurde kein dauerhafter Schaden festgestellt. Das Gerät wurde aus dem Betrieb "
         "genommen und zur Analyse eingeschickt."),
        ("Maßnahmen",
         "Sofortmaßnahmen: Gerät gesperrt, Kunden informiert, interne Ursachenanalyse eingeleitet.\n"
         "Geplante Maßnahmen: CAPA-Eröffnung, Produktprüfung, Ggf. Field Safety Corrective Action.\n"
         f"Status: {'Offen' if num % 2 == 0 else 'Abgeschlossen'}"),
    ]
    save_pdf(FOLDER_REG, f"VIG_{prod_short}_{num:04d}_Vigilanzbericht.pdf",
             f"Vigilanzbericht – {prod_name} – VIG-{C['short']}-2023-{num:04d}", sections, scan_prefix=True)


def gen_fsca(prod_name, prod_short):
    sections = [
        ("Field Safety Corrective Action (FSCA) – Sicherheitskorrekturmaßnahme",
         f"FSCA-Referenz: FSCA-{C['short']}-{prod_short}-2023-001\n"
         f"Produkt: {prod_name}\nHersteller: {C['name']}\n"
         f"Datum: {ds(2023, 8, 14)}\nBetroffene Lot-Nummern: LOT-2023001 bis LOT-2023089"),
        ("Hintergrund",
         f"Im Rahmen der PMS-Aktivitäten für {prod_name} wurden mehrere gleichartige Vorkommnisse "
         "identifiziert, die auf einen systematischen Fehler in der Softwarekomponente Version 3.1.2 "
         "hinweisen. Unter bestimmten Bedingungen (Einsatz bei Patienten mit Körpergewicht < 50 kg) "
         "kann es zu einer Fehlmessung kommen, die zu falschen Diagnoseergebnissen führt."),
        ("Maßnahme",
         "Art der Maßnahme: Software-Update (Firmware-Update v3.1.3) + Schulungshinweis an Anwender\n"
         "Keine Geräterückgabe erforderlich. Fernaktualisierung per verschlüsseltem Update-Server.\n\n"
         "Handlungsanweisung an Kunden: Bitte aktualisieren Sie die Firmware auf Version 3.1.3. "
         "Bei Patienten mit Körpergewicht < 50 kg ist bis zur Aktualisierung manuell zu kalibrieren."),
        ("Behördenbenachrichtigung",
         "Diese FSCA wurde dem BfArM (Bundesinstitut für Arzneimittel und Medizinprodukte) "
         f"am {ds(2023, 8, 14)} gemäß Art. 88 MDR gemeldet. "
         "Alle zuständigen Behörden der betroffenen EU-Mitgliedstaaten wurden informiert."),
        ("Kundenkommunikation",
         f"Alle {random.randint(180,320)} bekannten Betreiber des Produkts wurden schriftlich "
         "über die FSCA informiert. Die Reaktionsrate auf das Software-Update beträgt nach "
         f"4 Wochen: {random.randint(72,91)} %."),
    ]
    save_pdf(FOLDER_REG, f"FSCA_{prod_short}_001_Sicherheitskorrekturmassnahme.pdf",
             f"FSCA – {prod_name} – Feldsicherheitskorrekturmaßnahme", sections)


def gen_udi_documentation(prod_name, prod_ref, prod_short):
    sections = [
        ("UDI-Dokumentation gemäß Artikel 27 MDR 2017/745", None),
        ("UDI-System",
         f"Hersteller: {C['name']}\nEUDAMED-SRN: {C['eudamed_reg']}\n"
         f"Ausgebende Organisation: GS1 (Global Standards One)\n"
         f"UDI-DI: {prod_ref}"),
        ("UDI-Träger", [
            ["Verpackungsebene", "UDI-Format", "Datenträger", "Pflichtinhalt"],
            ["Primärpackmittel", "GS1-128", "Linearer Barcode", "UDI-DI + UDI-PI (Lot, Verfallsdatum)"],
            ["Sekundärpackmittel", "GS1-128 + QR-Code", "Linear + 2D", "Vollständige UDI"],
            ["Gerätekennzeichnung", "GS1 DataMatrix", "2D-Code", "UDI-DI + Seriennummer"],
        ]),
        ("EUDAMED-Registrierung",
         f"Der Hersteller hat das Produkt {prod_name} in EUDAMED unter der SRN "
         f"{C['eudamed_reg']} registriert. Alle Datenfelder gemäß Anhang VI Teil B MDR "
         "sind vollständig ausgefüllt und aktuell. Jährliche Überprüfung der Daten: Jänner."),
        ("Labelling Review",
         f"Review-Datum: {ds(2023, 3, 20)}\nReviewer: {C['qra']}\n"
         "Ergebnis: Alle Etiketten erfüllen die Anforderungen gemäß Anhang I Abschnitt 23 MDR "
         "und EN ISO 15223-1:2021. Symbole: ISO-konform. Sprachversionen: DE, EN, FR, IT, ES."),
    ]
    save_pdf(FOLDER_REG, f"UDI_{prod_short}_Dokumentation.pdf",
             f"UDI-Dokumentation – {prod_name}", sections)


def gen_nb_correspondence(prod_name, prod_short, doc_type, num):
    type_titles = {
        "application": "Antrag auf Konformitätsbewertung",
        "questions": "Fragen der Benannten Stelle (Major Deficiencies)",
        "response": "Antwort auf Fragen der Benannten Stelle",
        "certificate": "Konformitätszertifikat",
    }
    sections = [
        ("Dokumenteninformation",
         f"Von: {C['nb'] if doc_type == 'questions' else C['name']}\n"
         f"An: {C['name'] if doc_type == 'questions' else C['nb']}\n"
         f"Produkt: {prod_name}\nReferenz: NB-{C['short']}-{prod_short}-2022-{num:03d}\n"
         f"Datum: {ds(2022, random.randint(1,12), random.randint(1,28))}"),
        ("Inhalt",
         f"Betreff: {type_titles.get(doc_type, doc_type)} für {prod_name}\n\n"
         f"Sehr geehrte Damen und Herren,\n\n"
         "im Rahmen des Konformitätsbewertungsverfahrens gemäß Anhang IX der MDR 2017/745 "
         f"übersenden wir {'nachfolgende Fragen' if doc_type == 'questions' else 'die angeforderten Unterlagen'} "
         f"zu dem oben genannten Medizinprodukt {prod_name}.\n\n"
         + ("Die folgenden Punkte bedürfen der Klärung:\n"
            "1. Technische Dokumentation Abschnitt 3.2: Bitte ergänzen Sie die Risikoanalyse um "
            "die Bewertung von Cybersecurity-Bedrohungen gemäß MDCG 2019-16.\n"
            "2. Klinische Bewertung: Die verwendeten Äquivalenzstudien erfüllen nicht vollständig "
            "die Anforderungen von MDCG 2020-5. Bitte ergänzen Sie die Äquivalenzbegründung.\n"
            "3. Post-Market Surveillance: Der PMCF-Plan enthält keine spezifischen Erfolgskriterien "
            "(KPIs) gemäß MDCG 2020-8. Bitte ergänzen."
            if doc_type == "questions" else "Die entsprechenden Unterlagen finden Sie in den beigefügten Anlagen.")),
    ]
    save_pdf(FOLDER_REG, f"NB_{prod_short}_{num:03d}_{doc_type[:6].upper()}.pdf",
             f"Benannte Stelle Korrespondenz – {prod_name} – {type_titles.get(doc_type, doc_type)}", sections)


def gen_eudamed_registration():
    sections = [
        ("EUDAMED-Registrierungsbestätigung",
         f"Hersteller: {C['name']}\nSingle Registration Number (SRN): {C['eudamed_reg']}\n"
         f"Registrierungsdatum: {ds(2021, 5, 26)}\nStatus: AKTIV"),
        ("Registrierte Produkte", [
            ["Produkt", "UDI-DI", "Klasse", "Status"],
            [C["prod1_name"], C["prod1_ref"], "IIb", "Zertifiziert"],
            [C["prod2_name"], C["prod2_ref"], "IIa", "Zertifiziert"],
            [C["prod3_name"], C["prod3_ref"], "IVD Klasse C", "Zertifiziert"],
        ]),
        ("Hinweis",
         "Diese Bestätigung wurde aus dem EUDAMED-Portal der Europäischen Kommission abgerufen. "
         "Die in EUDAMED gespeicherten Daten sind öffentlich zugänglich und werden kontinuierlich "
         "vom Hersteller aktualisiert."),
    ]
    save_pdf(FOLDER_REG, "EUDAMED_Registrierungsbestaetigung.pdf",
             "EUDAMED-Registrierungsbestätigung", sections)


def gen_mdr_compliance_checklist():
    items = [
        ("Art. 10 Abs. 1 – QMS", "ISO 13485 Zertifikat vorliegend", "Erfüllt", C["iso_cert_nr"]),
        ("Art. 10 Abs. 2 – Technische Dokumentation", "Anhang II/III-Dokumentation vollständig", "Erfüllt", "Jährlich überprüft"),
        ("Art. 10 Abs. 3 – Konformitätsbewertung", "Anhang IX Verfahren abgeschlossen", "Erfüllt", C["nb_cert1"]),
        ("Art. 10 Abs. 4 – EU-Konformitätserklärung", "DoC vorhanden", "Erfüllt", "Alle 3 Produkte"),
        ("Art. 10 Abs. 5 – CE-Kennzeichnung", "CE + NB-Nummer auf Produkten", "Erfüllt", "Geprüft"),
        ("Art. 10 Abs. 7 – PMS-System", "PMS-Pläne und -Berichte vorliegend", "Erfüllt", "Jährlich"),
        ("Art. 10 Abs. 9 – SSCP", "SSCP für alle Klasse IIb-Produkte", "Erfüllt", "In EUDAMED"),
        ("Art. 10 Abs. 11 – Verantwortliche Person", "QRA benannt", "Erfüllt", C["qra"]),
        ("Art. 14 – Distributor-Pflichten", "Distributor-Vereinbarungen vorliegend", "Erfüllt", "3 Verträge"),
        ("Art. 27 – UDI", "UDI-System implementiert, EUDAMED-Daten gepflegt", "Erfüllt", C["eudamed_reg"]),
        ("Art. 29 – Registrierung", "EUDAMED-Registrierung vollständig", "Erfüllt", C["eudamed_reg"]),
        ("Art. 32 – SSCP", "SSCP publiziert", "Erfüllt", "Alle IIb-Produkte"),
        ("Art. 83 – PMS-System", "PMS-Pläne gemäß MDCG 2020-7", "Erfüllt", "Jährlich"),
        ("Art. 84 – PMS-Bericht", "PMSR für alle Produkte vorliegend", "Erfüllt", "Jährlich"),
        ("Art. 86 – PSUR", "PSUR für Klasse IIb erstellt", "Erfüllt", "Alle 2 Jahre"),
        ("Art. 87 – Meldepflicht", "Vigilanzprozess implementiert", "Erfüllt", "SOP vorhanden"),
    ]
    sheets = [
        ("MDR Compliance Checklist", ["Artikel", "Anforderung", "Status", "Nachweis/Anmerkung"],
         items, [25, 40, 12, 30]),
    ]
    make_xlsx_doc(FOLDER_REG, "MDR_Compliance_Checklist_2024.xlsx",
                  "MDR 2017/745 Compliance-Checkliste 2024", sheets)


def gen_fda_510k_excerpt(prod_name, prod_short):
    sections = [
        ("FDA 510(k) Premarket Notification – Executive Summary", None),
        ("Device Description",
         f"The {prod_name} (Model: {prod_short}-US-V3) is a non-invasive medical device intended "
         "for use in clinical settings by trained healthcare professionals. "
         "The device uses proprietary sensor technology to provide real-time diagnostic support.\n\n"
         f"Manufacturer: {C['name']}, {C['full_address']}, Germany\n"
         f"U.S. Agent: BTP MedTech Inc., 100 Market Street, San Francisco, CA 94105"),
        ("Predicate Device",
         "The 510(k) submission relies on the following predicate device: "
         "[PREDICATE DEVICE NAME], K201234, cleared DATE. "
         "The proposed device and the predicate are substantially equivalent with respect to "
         "intended use, technological characteristics, and safety and effectiveness.\n\n"
         "Note: Full 510(k) submission to be filed Q3 2024. This excerpt is for internal review."),
        ("Substantial Equivalence",
         "The device has the same intended use as the predicate. "
         "Technological differences exist (new sensor type) but these differences do not "
         "raise new types of safety and effectiveness questions. "
         "Performance testing demonstrates the device is at least as safe and effective."),
    ]
    save_pdf(FOLDER_REG, f"FDA_{prod_short}_510k_Excerpt.pdf",
             f"FDA 510(k) Excerpt – {prod_name}", sections)


def gen_ukca_submission(prod_name, prod_short):
    sections = [
        ("UKCA Marking Documentation",
         f"Product: {prod_name}\nManufacturer: {C['name']}, {C['full_address']}, Germany\n"
         f"UK Responsible Person: BTP UK Ltd., 10 Downing Business Park, London EC1A 1BB\n"
         "Applicable Regulation: UK Medical Devices Regulations 2002 (as amended)\n"
         f"UKCA Registration: MHRA Reference UK-CA-BTP-{prod_short}-2022"),
        ("Registration Status",
         f"The {prod_name} has been registered with the MHRA (Medicines and Healthcare products "
         "Regulatory Agency) under the UKCA marking scheme. Registration confirmed "
         f"{ds(2022, 10, 1)}. Annual renewal due October 2024."),
        ("Conformity Assessment",
         "UKCA conformity assessment performed via the route equivalent to EU MDR Annex IX. "
         "UK Approved Body: BSI (British Standards Institution), AB Ref. 2797. "
         "Certificate No.: UKCA-BSI-BTP-2022-001."),
    ]
    save_pdf(FOLDER_REG, f"UKCA_{prod_short}_Zulassung_UK.pdf",
             f"UKCA Marking – {prod_name}", sections)


def gen_ifu_excerpt(prod_name, prod_short):
    sections = [
        ("Gebrauchsanweisung (IFU) – Kurzfassung",
         f"Produkt: {prod_name} | Artikelnummer: {prod_short}-DE-V3\n"
         "Hersteller: Sentavia Precision GmbH, Freimannstraße 45, 80939 München\n"
         "Diese Gebrauchsanweisung (Instructions for Use, IFU) enthält wichtige Informationen "
         "zur sicheren Anwendung des Produkts. Bitte lesen Sie diese Anweisung vollständig "
         "durch, bevor Sie das Gerät in Betrieb nehmen."),
        ("Zweckbestimmung",
         f"{prod_name} ist für die nicht-invasive Messung und Auswertung klinischer Parameter "
         "bei erwachsenen Patienten in medizinischen Einrichtungen bestimmt. Das Gerät darf "
         "ausschließlich von geschultem medizinischem Fachpersonal verwendet werden."),
        ("Sicherheitshinweise",
         "WARNUNG: Nicht in MRT-Umgebung verwenden.\n"
         "WARNUNG: Nicht bei Patienten mit aktiven implantierten Herzschrittmachern verwenden.\n"
         "ACHTUNG: Gerät regelmäßig kalibrieren gemäß Wartungsplan (Abschnitt 7).\n"
         "ACHTUNG: Gerät nicht in Flüssigkeiten tauchen."),
        ("Inbetriebnahme",
         "1. Gerät aus der Verpackung entnehmen und auf Unversehrtheit prüfen.\n"
         "2. Netzteil anschließen und Gerät einschalten (Taste: Power On).\n"
         "3. Selbsttest abwarten (ca. 90 Sekunden).\n"
         "4. Kalibrierung gemäß Abschnitt 4.2 durchführen.\n"
         "5. Patientendaten eingeben und Messung starten."),
        ("Reinigung und Desinfektion",
         "Das Gerät ist nach jeder Patientenanwendung zu reinigen. Zugelassene Desinfektionsmittel: "
         "Ethanol 70 %, Isopropanol 70 %, alkoholische Flächendesinfektionsmittel (VAH-Liste). "
         "Nicht autoklavierbar. Keine UV-Bestrahlung."),
    ]
    save_pdf(FOLDER_REG, f"IFU_{prod_short}_Gebrauchsanweisung_DE_Auszug.pdf",
             f"IFU Auszug – {prod_name}", sections)


def gen_risk_management_file(prod_name, prod_short):
    sections = [
        ("Risikomanagementakte gemäß ISO 14971:2019", None),
        ("Risikomanagementplan",
         f"Produkt: {prod_name}\nGeltungsbereich: Gesamter Produktlebenszyklus von Entwicklung bis Entsorgung\n"
         f"Risikomanagementbeauftragter: {C['qra']}\nGültig ab: {ds(2022, 1, 1)}"),
        ("Gefährdungsanalyse", [
            ["Gefährdungs-ID", "Gefährdung", "Gefährdungssituation", "Schaden", "W", "S", "Risiko"],
            ["GEF-001", "Falschmessung", "Sensor-Drift bei extremen Temperaturen", "Fehldiagnose", "2", "4", "8 (mittel)"],
            ["GEF-002", "Elektrischer Schlag", "Isolation beschädigt", "Herzrhythmusstörung", "1", "5", "5 (mittel)"],
            ["GEF-003", "Software-Fehler", "Fehler in Auswertungsalgorithmus", "Fehlentscheidung", "2", "3", "6 (mittel)"],
            ["GEF-004", "Infektionsübertragung", "Kontamination bei Reinigungsversagen", "Infektion", "2", "4", "8 (mittel)"],
            ["GEF-005", "Elektromagnetische Störung", "Betrieb neben Hochfrequenzgeräten", "Fehlmessung", "3", "3", "9 (mittel)"],
        ]),
        ("Risikoakzeptanzkriterien",
         "Als ALARP (As Low As Reasonably Practicable) Grenzwerte gelten:\n"
         "• Inakzeptables Risiko: ≥ 15 (W × S)\n"
         "• Akzeptables Risiko: < 6 (W × S)\n"
         "• ALARP-Bereich: 6–14 (erfordert weitere Risikominderung oder Begründung)\n\n"
         "Alle Risiken nach Implementierung der Risikominderungsmaßnahmen liegen im "
         "akzeptablen oder ALARP-Bereich."),
        ("Gesamtbewertung",
         f"Das Gesamtrisiko für {prod_name} ist akzeptabel im Sinne der ISO 14971:2019. "
         "Die Nutzen-Risiko-Bilanz ist positiv: Der klinische Nutzen (präzise Diagnose) "
         "überwiegt die verbleibenden Restrisiken deutlich."),
    ]
    save_pdf(FOLDER_REG, f"RM_{prod_short}_Risikomanagementakte_ISO14971.pdf",
             f"Risikomanagementakte ISO 14971:2019 – {prod_name}", sections)


def gen_design_control(prod_name, prod_short):
    sections = [
        ("Design History File (DHF) – Zusammenfassung", None),
        ("Design Inputs", [
            ["Anforderung", "Quelle", "Version", "Verifikationsmethode"],
            ["Messgenauigkeit ± 2 %", "Marktanalyse / KOL-Befragung", "v1.0", "Labortests"],
            ["Batterielaufzeit ≥ 8 Std.", "Benutzeranforderung", "v1.0", "Dauertest"],
            ["IP44-Schutzklasse", "Regulatorische Anforderung IEC 60529", "v1.0", "Typprüfung"],
            ["Softwareklasse IEC 62304 Klasse B", "Risikobewertung", "v1.0", "Code-Review + Tests"],
            ["Betriebstemperatur 10–40 °C", "Norm IEC 60601-1-11", "v1.0", "Klimatest"],
        ]),
        ("Design Verification Summary",
         f"Alle Design-Outputs wurden gegen die Design-Inputs verifiziert (V&V-Bericht: "
         f"VV-{prod_short}-2022-001). Ergebnis: Alle Anforderungen erfüllt. "
         f"Verifikationsprotokoll vom {ds(2022, 7, 15)} liegt vor."),
        ("Design Validation Summary",
         f"Die klinische Validierung von {prod_name} wurde in einer Studie mit n = {random.randint(80,150)} "
         f"Probanden an {random.randint(3,5)} Prüfzentren durchgeführt (Klinische Prüfung gemäß ISO 14155). "
         "Alle primären Endpunkte wurden erreicht. Design Transfer abgeschlossen."),
    ]
    save_pdf(FOLDER_REG, f"DHF_{prod_short}_Design_History_File_Summary.pdf",
             f"Design History File (DHF) – {prod_name}", sections)


def generate_06_regulatorisch():
    print("  Generating 06_Regulatorisch_MDR (big section)...")
    products = [
        (C["prod1_name"], C["prod1_class"], C["prod1_ref"], "CSP"),
        (C["prod2_name"], C["prod2_class"], C["prod2_ref"], "OFN"),
        (C["prod3_name"], C["prod3_class"], C["prod3_ref"], "DKS"),
    ]
    # Technical documentation sections (many per product)
    td_sections = [
        "Produktbeschreibung und Zweckbestimmung",
        "UDI und Produktkennzeichnung",
        "Konstruktions- und Fertigungsinformationen",
        "Grundlegende Sicherheits- und Leistungsanforderungen",
        "Nutzen-Risiko-Analyse",
        "Risikomanagement (ISO 14971)",
        "Produktverifizierung und -validierung",
        "Klinische Bewertung (MDR Art. 61)",
        "Post-Market Surveillance Plan",
        "Gebrauchsanweisung (IFU)",
        "Konformitätserklärung",
        "Prüfberichte – Elektrische Sicherheit (IEC 60601-1)",
        "Prüfberichte – Elektromagnetische Verträglichkeit (IEC 60601-1-2)",
        "Biokompatibilitätsbewertung (ISO 10993)",
        "Software-Dokumentation (IEC 62304)",
    ]
    for prod_name, prod_class, prod_ref, prod_short in products:
        for i, section in enumerate(td_sections, 1):
            gen_technical_file_section(prod_name, prod_class, prod_ref, section, i)
        gen_declaration_of_conformity(prod_name, prod_class, prod_ref)
        gen_pms_plan(prod_name, prod_short)
        gen_pms_report(prod_name, prod_short, 2022)
        gen_pms_report(prod_name, prod_short, 2023)
        gen_pmcf_plan(prod_name, prod_short)
        gen_pmcf_evaluation(prod_name, prod_short, 2023)
        gen_sscp(prod_name, prod_class, prod_ref, prod_short)
        gen_udi_documentation(prod_name, prod_ref, prod_short)
        gen_risk_management_file(prod_name, prod_short)
        gen_design_control(prod_name, prod_short)
        gen_ifu_excerpt(prod_name, prod_short)
        gen_fsca(prod_name, prod_short)
        gen_fda_510k_excerpt(prod_name, prod_short)
        gen_ukca_submission(prod_name, prod_short)
        # NB correspondence
        for doc_type, num in [("application", 1), ("questions", 2), ("response", 3), ("certificate", 4)]:
            gen_nb_correspondence(prod_name, prod_short, doc_type, num)
        # Vigilance reports
        for v_num in range(1, 4):
            gen_vigilance_report(prod_name, prod_short, v_num)
    gen_eudamed_registration()
    gen_mdr_compliance_checklist()
    # MDR Article 10(11) Responsible Person
    sections = [
        ("Verantwortliche Person gemäß Art. 15 MDR 2017/745",
         f"Benannte verantwortliche Person: {C['qra']}\n"
         f"Qualifikation: Dipl.-Ing. Medizintechnik, 12 Jahre Erfahrung in Regulatory Affairs\n"
         f"Bestellung: {ds(2021, 5, 26)}\nZuständig für: Alle Medizinprodukte der {C['name']}"),
        ("Aufgaben",
         "Die verantwortliche Person stellt sicher, dass:\n"
         "(a) die Konformität der Produkte angemessen überprüft wird\n"
         "(b) die technische Dokumentation und die EU-Konformitätserklärungen erstellt werden\n"
         "(c) PMS-Pflichten erfüllt werden\n"
         "(d) Meldepflichten gegenüber Behörden eingehalten werden\n"
         "(e) Bei Untersuchungen kooperiert wird"),
    ]
    save_pdf(FOLDER_REG, "MDR_Art15_Verantwortliche_Person.pdf",
             "Verantwortliche Person gemäß Art. 15 MDR 2017/745", sections)
    # Additional regulatory items
    for i, topic in enumerate([
        "MDR Transition Plan 2021", "IVDR Transition Plan 2022", "Regulatory Strategy 2024–2026",
        "Notified Body Selection Rationale", "MDR Gap Analysis 2020",
        "Regulatory Submission Tracker 2023",
    ], 1):
        sheets = [("Übersicht", ["Thema", "Status", "Verantwortlich", "Fälligkeit"],
                   [(f"Aufgabe {j}", random.choice(["Erledigt", "In Arbeit", "Geplant"]),
                     random.choice([C["qra"], "Markus Bauer", C["cmo"]]),
                     ds(2024, random.randint(1, 12), random.randint(1, 28)))
                    for j in range(1, 10)], [30, 15, 20, 18])]
        make_xlsx_doc(FOLDER_REG, f"REG_{i:02d}_{sfn(topic[:25])}.xlsx", topic, sheets)


# ══════════════════════════════════════════════════════════════════════
# FOLDER 07 – QUALITÄT / QMS
# ══════════════════════════════════════════════════════════════════════
FOLDER_QMS = "07_Qualitaet_QMS"

SOPS = [
    ("SOP-001", "Lenkung von Dokumenten und Aufzeichnungen"),
    ("SOP-002", "Behandlung von nichtkonformen Produkten"),
    ("SOP-003", "Korrektur- und Vorbeugungsmaßnahmen (CAPA)"),
    ("SOP-004", "Interne Audits"),
    ("SOP-005", "Lieferantenqualifikation und -überwachung"),
    ("SOP-006", "Sterilisationsvalidierung"),
    ("SOP-007", "Softwarevalidierung"),
    ("SOP-008", "Schulung und Kompetenzmanagement"),
    ("SOP-009", "Risikomanagement"),
    ("SOP-010", "Umgang mit Kundenbeschwerden"),
    ("SOP-011", "Vigilanz und Marktbeobachtung"),
    ("SOP-012", "Änderungskontrolle (Change Management)"),
    ("SOP-013", "Post-Market Surveillance"),
    ("SOP-014", "Kalibrierung und Messtechnik"),
    ("SOP-015", "Design und Entwicklung"),
    ("SOP-016", "Eingangs- und Ausgangsprüfung"),
    ("SOP-017", "Produktkennzeichnung und UDI"),
    ("SOP-018", "Rückverfolgbarkeit und Chargendokumentation"),
    ("SOP-019", "Reinraumbetrieb und Umweltüberwachung"),
    ("SOP-020", "Umgang mit biologischen Proben und IVD"),
    ("SOP-021", "Statistische Methoden"),
    ("SOP-022", "Management Review"),
    ("SOP-023", "Wareneingang und Lager"),
    ("SOP-024", "Produktfreigabe"),
    ("SOP-025", "Externe Audits und Behördeninspektionen"),
    ("SOP-026", "IT-Sicherheit und Datenschutz"),
    ("SOP-027", "Reklamationsbearbeitung Exportmärkte"),
    ("SOP-028", "Biokompatibilitätsbewertung"),
]


def gen_sop(sop_id, sop_title):
    doc = Document()
    make_docx_header(doc, f"{sop_id} – {sop_title}", f"Version 3.0 | Gültig ab {ds(2023, 1, 1)}")
    # SOP header table
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    table_data = [
        ("Dokument-Nr.:", sop_id, "Version:", "3.0"),
        ("Titel:", sop_title, "Gültig ab:", ds(2023, 1, 1)),
        ("Verantwortlich:", C["qra"], "Erstellt von:", "Markus Bauer"),
        ("Genehmigt von:", C["ceo"], "Genehmigt am:", ds(2023, 1, 1)),
        ("Änderungsstand:", "3.0 – Vollständige Überarbeitung nach MDR-Transition", "", ""),
        ("Geltungsbereich:", f"Alle Mitarbeiter der {C['name']}", "", ""),
    ]
    for i, (k1, v1, k2, v2) in enumerate(table_data):
        row = table.rows[i].cells
        row[0].text = k1
        row[0].paragraphs[0].runs[0].bold = True
        row[1].text = v1
        row[2].text = k2
        row[2].paragraphs[0].runs[0].bold = True
        row[3].text = v2
    doc.add_paragraph()
    add_section(doc, "1. Zweck und Anwendungsbereich", [
        f"Diese SOP beschreibt die Anforderungen, Verantwortlichkeiten und Vorgehensweise für "
        f"{sop_title.lower()} bei der {C['name']}.",
        "Diese SOP gilt für alle Mitarbeiter und Bereiche der Gesellschaft, soweit sie in die "
        "beschriebenen Tätigkeiten eingebunden sind. Sie ist verbindlich und ersetzt alle "
        "vorherigen Versionen.",
        f"Normative Referenzen: ISO 13485:2016 Abschnitte 4.2, 8.1 ff.; MDR 2017/745 Art. 10; "
        "ISO 14971:2019 (soweit anwendbar).",
    ])
    add_section(doc, "2. Verantwortlichkeiten", [
        f"Prozessverantwortlicher: {C['qra']}, Head of Quality & Regulatory Affairs\n"
        "Dieser ist für die Implementierung, Aufrechterhaltung und kontinuierliche Verbesserung "
        "des in dieser SOP beschriebenen Prozesses verantwortlich.",
        "Mitarbeiter: Alle Mitarbeiter, die in die beschriebenen Tätigkeiten eingebunden sind, "
        "sind für die Einhaltung dieser SOP verantwortlich. Schulungsnachweis ist zu erbringen.",
        f"Management: {C['ceo']} genehmigt diese SOP und alle wesentlichen Änderungen daran. "
        "Das Management stellt die notwendigen Ressourcen zur Umsetzung bereit.",
    ])
    add_section(doc, "3. Definitionen und Abkürzungen", [
        "QMS: Qualitätsmanagementsystem\n"
        "CAPA: Corrective and Preventive Action (Korrektiv- und Präventivmaßnahme)\n"
        "MDR: Medical Device Regulation (EU-Medizinprodukteverordnung 2017/745)\n"
        "IVD: In-vitro-Diagnostikum\n"
        f"NB: Benannte Stelle ({C['nb']})\n"
        "OOS: Out of Specification\n"
        "NC: Non-Conformance (Nichtkonformität)",
    ])
    add_section(doc, "4. Prozessbeschreibung", [
        f"Der Prozess der {sop_title} umfasst die folgenden Hauptschritte:\n\n"
        "Schritt 1: Planung – Vor Beginn jeder Aktivität ist eine schriftliche Planung zu "
        "erstellen. Der Plan muss Ziele, Methoden, Zeitrahmen und Verantwortliche definieren.\n\n"
        "Schritt 2: Durchführung – Die Tätigkeit wird gemäß Plan und nach einschlägigen "
        "Normen durchgeführt. Alle Schritte sind zu dokumentieren.\n\n"
        "Schritt 3: Überprüfung – Ergebnisse werden gegen Akzeptanzkriterien geprüft. "
        "Abweichungen werden als NC erfasst und einer CAPA zugeführt.\n\n"
        "Schritt 4: Verbesserung – Erkenntnisse fließen in die jährliche Management-Review "
        "und in die kontinuierliche Verbesserung des QMS ein.",
        "Bei Abweichungen vom beschriebenen Prozess ist dies zu dokumentieren und zu begründen. "
        "Wesentliche Abweichungen sind mit dem QRA abzustimmen.",
    ])
    add_section(doc, "5. Aufzeichnungen und Dokumentenpflege", [
        "Alle Aufzeichnungen im Rahmen dieser SOP sind im QMS-Tool (MasterControl) abzulegen "
        "und für mindestens 10 Jahre (bei Medizinprodukten: gemäß MDR Anhang IX Abschnitt 3.9) "
        "aufzubewahren.",
        "Dokumente werden gemäß SOP-001 (Dokumentenlenkung) versioniert und genehmigt. "
        "Veraltete Versionen werden archiviert und als 'obsolet' gekennzeichnet.",
    ])
    add_section(doc, "6. Referenzdokumente", [
        "• ISO 13485:2016 (Qualitätsmanagementsystem für Medizinprodukte)\n"
        "• MDR 2017/745 (EU-Medizinprodukteverordnung)\n"
        "• MDCG-Leitfäden (aktuelle Fassungen)\n"
        f"• Qualitätshandbuch {C['name']}\n"
        "• Relevante Produktnormen (IEC 60601-1, ISO 14971 etc.)",
    ])
    save_docx(doc, FOLDER_QMS, f"{sop_id}_{sfn(sop_title[:25])}_v3.0.docx")


def gen_quality_manual():
    doc = Document()
    make_docx_header(doc, "Qualitätshandbuch",
                     f"Version 5.0 | Zertifikat-Nr.: {C['iso_cert_nr']}", confidential=False)
    doc.add_paragraph(
        f"Das vorliegende Qualitätshandbuch (QHB) beschreibt das Qualitätsmanagementsystem (QMS) "
        f"der {C['name']}. Es gilt für alle Standorte und Prozesse der Gesellschaft.\n\n"
        f"Zertifizierungsstatus: {C['iso_cert']}, Zertifikat-Nr. {C['iso_cert_nr']}, "
        f"ausgestellt durch {C['nb']}."
    )
    for title, paragraphs in [
        ("1. Unternehmensprofil und Qualitätspolitik", [
            f"{C['name']} entwickelt, produziert und vertreibt Medizinprodukte der Klassen IIa und IIb "
            f"sowie In-vitro-Diagnostika der Klasse C. Das Stammkapital beträgt {eur(C['stammkapital'])}. "
            f"Sitz der Gesellschaft: {C['full_address']}.",
            "Qualitätspolitik: Wir entwickeln und fertigen Medizinprodukte höchster Qualität, Sicherheit "
            "und Wirksamkeit. Unsere Qualitätspolitik verpflichtet uns zur kontinuierlichen Verbesserung "
            "aller Prozesse sowie zur vollständigen Einhaltung aller regulatorischen Anforderungen.",
        ]),
        ("2. Geltungsbereich des QMS", [
            f"Das QMS der {C['name']} umfasst: Design und Entwicklung von Medizinprodukten; "
            "Fertigung und Prüfung; Vertrieb und After-Sales-Service; Regulatorische Tätigkeiten "
            "einschließlich klinischer Bewertung und Post-Market Surveillance.",
            "Ausschlüsse: Keine wesentlichen Ausschlüsse vom Anwendungsbereich der ISO 13485:2016.",
        ]),
        ("3. Normative Grundlagen", [
            f"Das QMS entspricht den Anforderungen der {C['iso_cert']}, der EU-MDR 2017/745, "
            "der EU-IVDR 2017/746 sowie allen einschlägigen harmonisierten Normen (IEC 60601-1, "
            "ISO 14971:2019, IEC 62304:2006+A1, ISO 10993-1:2018).",
        ]),
        ("4. Prozesslandschaft", [
            "Das QMS ist prozessorientiert aufgebaut. Hauptprozesse: Entwicklung (Design Control), "
            "Produktion und Prüfung, PMS und Vigilanz, Lieferantenmanagement, CAPA, Interne Audits, "
            "Management Review. Unterstützungsprozesse: HR, IT, Compliance, Finanzen.",
        ]),
        ("5. Verantwortung der Leitung", [
            f"Die Geschäftsführung ({C['ceo']}, {C['cto']}, {C['cfo']}) trägt die Gesamtverantwortung "
            "für die Aufrechterhaltung und Weiterentwicklung des QMS. Die Verantwortliche Person "
            f"gemäß Art. 15 MDR ist {C['qra']}.",
        ]),
    ]:
        add_section(doc, title, paragraphs)
    save_docx(doc, FOLDER_QMS, "QHB_001_Qualitaetshandbuch_v5.0.docx")


def gen_internal_audit_report(year, quarter):
    d = rdate(year, (quarter-1)*3+1, quarter*3)
    sections = [
        ("Interner Auditbericht gemäß ISO 13485:2016",
         f"Audit-Nr.: IA-{year}-{quarter:02d}\nDatum: {ds(d.year, d.month, d.day)}\n"
         f"Auditor: Sarah Krause (QA Lead)\nAuditierter Bereich: "
         f"{random.choice(['Produktion und Prüfung', 'Regulatory Affairs', 'Design & Entwicklung', 'Vertrieb und After-Sales'])}"),
        ("Auditumfang",
         f"Dieser interne Audit überprüft die Einhaltung der ISO 13485:2016 und der relevanten "
         "SOPs im auditierten Bereich. Normative Referenzen: ISO 13485:2016, MDR 2017/745, "
         "interne SOPs (SOP-001 bis SOP-025)."),
        ("Feststellungen", [
            ["Nr.", "Feststellung", "Norm-Abschnitt", "Klassifizierung", "Status"],
            ["F1", "Kalibrierungsnachweis für Messgerät XY fehlt", "ISO 13485 7.6", "Minor", "CAPA eröffnet"],
            ["F2", "SOP-008 Version veraltet (v1.2 statt v3.0)", "ISO 13485 4.2.3", "Minor", "Korrigiert"],
            ["F3", "Lieferantenaudit-Dokumentation unvollständig", "ISO 13485 7.4.1", "Minor", "In Bearbeitung"],
            ["F4", "Best Practice: Elektronische Unterschriften korrekt implementiert", "—", "Positive Finding", "—"],
        ]),
        ("Abschlussbewertung",
         f"Der auditierte Bereich entspricht insgesamt den Anforderungen der ISO 13485:2016. "
         "Drei Minor-Findings wurden identifiziert. Keine Major Nonconformances. "
         "CAPA-Maßnahmen werden innerhalb von 60 Tagen abgeschlossen. "
         "Folgeterminal: Quartalsaudit."),
    ]
    save_pdf(FOLDER_QMS, f"IA_{year}_{quarter:02d}_Interner_Auditbericht.pdf",
             f"Interner Auditbericht {year} Q{quarter}", sections)


def gen_management_review(year):
    rev = C.get(f"revenue_{year}", 87_200_000)
    sections = [
        ("Management Review gemäß ISO 13485:2016 Abschnitt 5.6",
         f"Datum: {ds(year, 12, 15)}\nOrt: {C['full_address']}, Konferenzraum\n"
         f"Teilnehmer: {C['ceo']}, {C['cto']}, {C['cfo']}, {C['cmo']}, {C['qra']}, Sarah Krause, Markus Bauer"),
        ("Eingaben zum Management Review", [
            ["Thema", "Kennzahl", "Wert", "Vorjahr", "Trend"],
            ["Kundenbeschwerden", "Anzahl", str(random.randint(15,40)), str(random.randint(20,45)), "↓ Verbessert"],
            ["CAPA-Abschlussrate", "%", f"{random.randint(88,98)}%", f"{random.randint(80,92)}%", "↑ Verbessert"],
            ["Interne Audit-Findings (Minor)", "Anzahl", str(random.randint(5,15)), str(random.randint(8,20)), "↓ Verbessert"],
            ["Liefertreue Lieferanten", "%", f"{random.randint(94,99)}%", f"{random.randint(90,96)}%", "↑ Verbessert"],
            ["Schulungsquote", "%", f"{random.randint(96,100)}%", f"{random.randint(90,97)}%", "↑ Verbessert"],
            ["Umsatz", "EUR", f"{rev:,}", f"{int(rev*0.82):,}", "↑ Wachstum"],
        ]),
        ("Ausgaben: Beschlüsse", [
            ["Maßnahme", "Verantwortlich", "Fälligkeit", "Ressourcen"],
            ["Erweiterung des PMCF-Programms um Anwenderregistry", C["cmo"], ds(year+1, 3, 31), "EUR 120.000"],
            ["Einführung neuer Kalibrierungssoftware", C["qra"], ds(year+1, 6, 30), "EUR 45.000"],
            ["Verstärkung des Regulatory-Teams um 2 FTE", C["ceo"], ds(year+1, 3, 1), "Personalkosten"],
        ]),
        ("Qualitätsziele für das Folgejahr",
         f"1. Complaint-Rate: ≤ {random.randint(12,20)} Beschwerden gesamt\n"
         "2. CAPA-Abschlussrate: ≥ 95 % fristgemäß\n"
         "3. Keine Major-Nonconformances bei internen Audits\n"
         "4. ISO 13485 Re-Zertifizierungsaudit ohne Major-Findings bestehen"),
    ]
    save_pdf(FOLDER_QMS, f"MR_{year}_Management_Review_Protokoll.pdf",
             f"Management Review {year} – Protokoll", sections)


def gen_capa_record(num):
    d = rdate(2023)
    prod = random.choice([C["prod1_name"], C["prod2_name"], C["prod3_name"]])
    sections = [
        ("CAPA-Bericht",
         f"CAPA-Nr.: CAPA-2023-{num:04d}\nDatum: {ds(d.year, d.month, d.day)}\n"
         f"Eröffnet durch: {random.choice([C['qra'], 'Sarah Krause', 'Markus Bauer'])}\n"
         f"Produkt/Prozess: {prod}\n"
         f"Klassifizierung: {'Korrektivmaßnahme' if num % 2 == 0 else 'Präventivmaßnahme'}"),
        ("Problembeschreibung",
         f"Nichtkonformität Nr. NC-2023-{num:04d}: Im Rahmen der {'internen Auditierung' if num % 3 == 0 else 'Beschwerdebearbeitung'} "
         "wurde eine Abweichung von den geltenden Qualitätsanforderungen festgestellt. "
         f"Betroffener Bereich: {random.choice(['Produktion', 'Lager', 'Dokumentation', 'Software'])}. "
         "Eine unmittelbare Gefährdung von Patienten oder Benutzern wurde ausgeschlossen."),
        ("Ursachenanalyse",
         f"Methode: {random.choice(['Fishbone (Ishikawa)', '5-Why-Analyse', 'Fault Tree Analysis'])}\n"
         "Ermittelte Ursache: Mangelnde Schulung der betroffenen Mitarbeiter auf aktuellem SOP-Stand. "
         "Beitragende Faktoren: Schnelles Mitarbeiterwachstum, unzureichende Einarbeitungszeit."),
        ("Maßnahmenplan", [
            ["Maßnahme", "Verantwortlich", "Fälligkeit", "Status"],
            ["Nachschulung aller betroffenen Mitarbeiter", "HR + QA", ds(d.year, min(d.month+1,12), 15), "Abgeschlossen"],
            ["Aktualisierung SOP-008", C["qra"], ds(d.year, min(d.month+1,12), 1), "Abgeschlossen"],
            ["Wirksamkeitsprüfung nach 60 Tagen", "Sarah Krause", ds(d.year, min(d.month+2,12), d.day), "Ausstehend"],
        ]),
    ]
    save_pdf(FOLDER_QMS, f"CAPA_2023_{num:04d}_Bericht.pdf",
             f"CAPA-Bericht CAPA-2023-{num:04d}", sections)


def gen_calibration_records():
    rows = []
    instruments = [
        ("KAL-001", "Digitalmultimeter Fluke 87V", "Produktion", "IEC 61010-1"),
        ("KAL-002", "Kraftmessgerät Kistler 9101A", "F&E-Labor", "ISO 376"),
        ("KAL-003", "Drucksensor PCB 106B", "Reinraum", "DIN EN 837-1"),
        ("KAL-004", "Thermohygrometer Testo 610", "Lager", "ISO 9001"),
        ("KAL-005", "Pipette Eppendorf Research Plus", "IVD-Labor", "ISO 8655"),
        ("KAL-006", "Waage Mettler Toledo XSR", "Qualitätsprüfung", "OIML R111"),
        ("KAL-007", "Oszilloskop Keysight DSOX1204G", "Elektronikentwicklung", "IEC 61010"),
        ("KAL-008", "pH-Meter Mettler Toledo S470", "Biolab", "DIN 19268"),
        ("KAL-009", "Partikelzähler Met One 3400+", "Reinraum", "ISO 14644-3"),
        ("KAL-010", "Klimakammer Memmert HPP110", "Validierung", "IEC 60068-3-5"),
    ]
    for kal_id, instrument, location, norm in instruments:
        last_cal = rdate(2023, 1, 6)
        next_cal = date(last_cal.year + 1, last_cal.month, last_cal.day)
        rows.append((kal_id, instrument, location, norm,
                     ds(last_cal.year, last_cal.month, last_cal.day),
                     ds(next_cal.year, next_cal.month, min(next_cal.day, 28)),
                     "Kalibrierungslabor München GmbH",
                     random.choice(["OK – Konform", "OK – Konform", "OK – Konform",
                                    "OOS – Außer Betrieb genommen"])))
    sheets = [
        ("Kalibriernachweise", ["Kal.-Nr.", "Gerät/Instrument", "Standort", "Norm",
                                 "Letzte Kalibrierung", "Nächste Fälligkeit", "Kalibrierungsdienstleister", "Ergebnis"],
         rows, [10, 30, 20, 14, 18, 18, 30, 22]),
    ]
    make_xlsx_doc(FOLDER_QMS, "QMS_Kalibriernachweise_2023.xlsx",
                  "Kalibriernachweise Messmittel 2023", sheets)


def gen_nc_report(num):
    d = rdate(2023)
    sections = [
        ("Nichtkonformitätsbericht",
         f"NC-Nr.: NC-2023-{num:04d}\nDatum: {ds(d.year, d.month, d.day)}\n"
         f"Bereich: {random.choice(['Produktion', 'Qualitätsprüfung', 'Lager', 'Einkauf'])}\n"
         f"Produkt: {random.choice([C['prod1_name'], C['prod2_name'], C['prod3_name']])}\n"
         f"Klassifizierung: {random.choice(['Klasse A (kritisch)', 'Klasse B (wesentlich)', 'Klasse C (geringfügig)'])}"),
        ("Beschreibung der Nichtkonformität",
         f"Art der Abweichung: {random.choice(['Dimensionale Abweichung', 'Kennzeichnungsfehler', 'Dokumentationslücke', 'Prozessabweichung'])}\n"
         "Die festgestellte Nichtkonformität erfordert sofortige Einleitung von Maßnahmen. "
         "Das betroffene Produkt wurde gesperrt und unter Quarantäne gestellt bis zur Entscheidung."),
        ("Dispositionsentscheidung",
         f"Entscheidung: {random.choice(['Freigabe nach Nacharbeit', 'Ausschuss (Vernichtung)', 'Sonderfreigabe', 'Rücksendung an Lieferant'])}\n"
         f"Genehmigt von: {C['qra']}\nDatum: {ds(d.year, d.month, min(d.day+5, 28))}"),
    ]
    save_pdf(FOLDER_QMS, f"NC_2023_{num:04d}_Nichtkonformitaetsbericht.pdf",
             f"Nichtkonformitätsbericht NC-2023-{num:04d}", sections)


def generate_07_qualitaet():
    print("  Generating 07_Qualitaet_QMS...")
    gen_quality_manual()
    for sop_id, sop_title in SOPS:
        gen_sop(sop_id, sop_title)
    for yr in [2022, 2023]:
        for q in range(1, 5):
            gen_internal_audit_report(yr, q)
    for yr in [2021, 2022, 2023]:
        gen_management_review(yr)
    for i in range(1, 16):
        gen_capa_record(i)
    for i in range(1, 11):
        gen_nc_report(i)
    gen_calibration_records()


# ══════════════════════════════════════════════════════════════════════
# FOLDER 08 – KLINISCHE BEWERTUNG
# ══════════════════════════════════════════════════════════════════════
FOLDER_CLIN = "08_Klinische_Bewertung"


def gen_cer(prod_name, prod_class, prod_short):
    sections = [
        ("Zusammenfassung",
         f"Dieser Clinical Evaluation Report (CER) bewertet die klinische Evidenz für {prod_name} "
         f"({prod_class}) gemäß Artikel 61 und Anhang XIV Teil A der MDR 2017/745, "
         "dem Leitfaden MDCG 2020-5 sowie der Norm ISO/TR 20416:2020.\n\n"
         f"Fazit: Die klinische Evidenz für {prod_name} ist ausreichend und belegt einen klaren "
         "klinischen Nutzen bei akzeptablem Restrisiko. Das Produkt erfüllt die grundlegenden "
         "Sicherheits- und Leistungsanforderungen gemäß Anhang I der MDR."),
        ("Zweckbestimmung und Indikationen",
         f"{prod_name} ist für die Anwendung durch geschultes medizinisches Fachpersonal in "
         "klinischen Einrichtungen bestimmt. Das Produkt unterstützt die klinische Diagnose und "
         "Entscheidungsfindung. Es ist nicht für den Hausgebrauch oder Patientenselbstanwendung vorgesehen."),
        ("Suchstrategie für Literaturdaten", [
            ["Datenbank", "Suchterme (vereinfacht)", "Treffer gesamt", "Nach Screening", "Eingeschlossen"],
            ["PubMed/MEDLINE", f'"{prod_name}" OR Äquivalenzprodukte', str(random.randint(1200,3500)), str(random.randint(200,600)), str(random.randint(45,120))],
            ["Embase", f'"{prod_name}" AND clinical evidence', str(random.randint(800,2000)), str(random.randint(150,400)), str(random.randint(30,90))],
            ["Cochrane Library", "Systematische Reviews", str(random.randint(50,200)), str(random.randint(15,60)), str(random.randint(5,20))],
            ["Herstellerinterne Daten", "Klinische Studien, PMCF", "n/a", "n/a", str(random.randint(3,8))],
        ]),
        ("Klinische Evidenzübersicht",
         f"Insgesamt {random.randint(60,180)} Publikationen wurden in die Bewertung eingeschlossen. "
         f"Davon {random.randint(5,15)} randomisierte kontrollierte Studien (RCTs), "
         f"{random.randint(10,30)} prospektive Kohortenstudien und {random.randint(20,80)} "
         "retrospektive Analysen und Fallserien.\n\n"
         "Diagnostische Genauigkeit: Über alle Studien gepoolte Sensitivität von "
         f"{random.randint(92,97)} % und Spezifizität von {random.randint(88,95)} %.\n"
         "Klinischer Nutzen: Signifikante Verbesserung der diagnostischen Sicherheit und "
         f"Reduzierung der Zeit bis zur Diagnose um durchschnittlich {random.randint(18,35)} %."),
        ("Äquivalenzbewertung",
         "Für den klinischen Nachweis wurden Äquivalenzprodukte herangezogen gemäß MDCG 2020-5 "
         "Kriterien (klinische, technische und biologische Äquivalenz). "
         "Die Äquivalenz wurde mit dem Produkt [Äquivalenzprodukt] nachgewiesen; "
         "ein Zugangsvertrag nach MDR Art. 61(5) liegt vor."),
        ("Nutzenbewertung",
         f"Klinischer Nutzen: {prod_name} trägt zur Verbesserung der Patientenversorgung bei "
         "durch präzisere Diagnose, Reduktion unnötiger Eingriffe und schnellere Behandlungsentscheidungen.\n\n"
         "Restrisiken: Alle nach Implementierung der Risikominderungsmaßnahmen verbleibenden "
         "Restrisiken sind akzeptabel. Die Nutzen-Risiko-Bilanz ist eindeutig positiv."),
        ("Schlussfolgerung",
         f"Die klinische Bewertung kommt zu dem Ergebnis, dass {prod_name} die grundlegenden "
         "Sicherheits- und Leistungsanforderungen gemäß Anhang I der MDR 2017/745 erfüllt. "
         "Der klinische Nutzen ist nachgewiesen und die Sicherheit ist belegt. "
         f"Nächste Überprüfung des CER: {date.today().year + 1}. Verantwortlich: {C['cmo']}."),
    ]
    save_pdf(FOLDER_CLIN, f"CER_{prod_short}_Clinical_Evaluation_Report_v3.0.pdf",
             f"Clinical Evaluation Report – {prod_name}", sections)


def gen_clinical_investigation_plan(prod_name, prod_short):
    sections = [
        ("Klinischer Prüfplan (CIP) – Synopsis",
         f"Titel: Klinische Prüfung zur Bestätigung der Sicherheit und Leistung von {prod_name}\n"
         f"Sponsor: {C['name']}, {C['full_address']}\n"
         f"Prüfarzt: Prof. Dr. Klaus Meier, {C['hosp1']}\n"
         f"Studientyp: Prospektive, multizentrische Kohortenstudie\n"
         f"Studienpopulation: Erwachsene Patienten (≥ 18 Jahre) mit relevanter Indikation\n"
         f"Stichprobengröße: n = {random.randint(80, 200)} (berechnet auf Basis Power-Analyse)\n"
         f"Studienorte: {random.randint(3, 6)} Prüfzentren in Deutschland und Österreich\n"
         f"Geplante Laufzeit: 24 Monate"),
        ("Primärer Endpunkt",
         f"Diagnostische Sensitivität und Spezifizität von {prod_name} im Vergleich zur "
         "Goldstandard-Methode, gemessen über den gesamten Beobachtungszeitraum. "
         "Erfolgscriterium: Sensitivität ≥ 92 %, Spezifizität ≥ 88 % (jeweils mit 95 %-KI)."),
        ("Statistische Methodik",
         "Auswertung nach ITT- (Intention to Treat) und PP-Prinzip (Per Protocol). "
         "Berechnung der diagnostischen Güte mit 95 %-Konfidenzintervallen nach Wilson. "
         "Subgruppenanalysen nach Altersgruppen und Indikationsuntergruppen."),
        ("Ethik und Regulatorik",
         f"Das Studienprotokoll wurde der Ethikkommission der Medizinischen Hochschule Hannover "
         f"vorgelegt. Votum Nr. EK-{random.randint(2000,2999)}/2023. Positives Votum erteilt am "
         f"{ds(2023, 3, random.randint(10,25))}."),
    ]
    save_pdf(FOLDER_CLIN, f"CIP_{prod_short}_Klinischer_Pruefplan.pdf",
             f"Klinischer Prüfplan – {prod_name}", sections)


def gen_csr(prod_name, prod_short):
    sections = [
        ("Klinischer Studienbericht – Executive Summary",
         f"Studie: Klinische Prüfung {prod_name} (Studie BTP-2021-{random.randint(1,5):02d})\n"
         f"Sponsor: {C['name']}\nAbschluss: {ds(2022, random.randint(6,12), random.randint(1,28))}"),
        ("Ergebnisse", [
            ["Endpunkt", "Ergebnis", "Zielwert", "Erfolg"],
            ["Sensitivität", f"{random.randint(93,98)} %", "≥ 92 %", "JA"],
            ["Spezifizität", f"{random.randint(89,96)} %", "≥ 88 %", "JA"],
            ["Gerätebezogene SAEs", f"{random.randint(0,2)}", "≤ 3", "JA"],
            ["Abbruchrate", f"{random.uniform(1,5):.1f} %", "< 10 %", "JA"],
        ]),
        ("Sicherheit",
         f"Im Beobachtungszeitraum wurden {random.randint(0,2)} gerätebedingte Serious Adverse Events "
         "(SAEs) berichtet, alle vollständig aufgeklärt und ohne Langzeitfolgen für die Probanden. "
         f"{random.randint(3,12)} nicht-gerätebedingte Adverse Events (AEs) wurden dokumentiert."),
        ("Schlussfolgerung",
         f"Die Studie belegt die Sicherheit und Leistungsfähigkeit von {prod_name}. "
         "Alle primären und sekundären Endpunkte wurden erreicht. "
         "Die Ergebnisse fließen in den Clinical Evaluation Report (CER) ein."),
    ]
    save_pdf(FOLDER_CLIN, f"CSR_{prod_short}_Klinischer_Studienbericht.pdf",
             f"Klinischer Studienbericht – {prod_name}", sections)


def gen_investigator_agreement(prod_short, hosp_name, hosp_short):
    doc = Document()
    make_docx_header(doc, "Prüfarztvertrag / Investigator Agreement",
                     f"{C['name']} und {hosp_name}", confidential=True)
    add_clause(doc, 1, "Gegenstand", [
        f"Dieser Vertrag regelt die Durchführung der klinischen Prüfung für {C['prod1_name']} "
        f"(Studie BTP-2023-{prod_short}) durch {hosp_name} als Prüfzentrum.",
        f"Die klinische Prüfung wird gemäß ISO 14155:2020 und MDR Anhang XV durchgeführt. "
        f"Der Sponsor ist {C['name']}.",
    ])
    add_clause(doc, 2, "Vergütung", [
        f"Für jeden vollständig dokumentierten Studienteilnehmer erhält das Prüfzentrum "
        f"eine Prüfungsvergütung von EUR {random.randint(600, 1200)} pro Patient.",
        "Reisekosten und Sachaufwendungen werden gemäß Aufstellung in Anlage 2 erstattet.",
    ])
    add_clause(doc, 3, "Datenschutz", [
        "Alle Patientendaten werden pseudonymisiert in REDCap gespeichert. "
        "Der Prüfarzt ist verantwortlich für die ordnungsgemäße Einholung der Patienteneinwilligung.",
    ])
    save_docx(doc, FOLDER_CLIN,
              f"IA_{prod_short}_{hosp_short}_Pruefarztvtrag.docx")


def generate_08_klinisch():
    print("  Generating 08_Klinische_Bewertung...")
    for pname, pclass, pref, pshort in [
        (C["prod1_name"], C["prod1_class"], C["prod1_ref"], "CSP"),
        (C["prod2_name"], C["prod2_class"], C["prod2_ref"], "OFN"),
        (C["prod3_name"], C["prod3_class"], C["prod3_ref"], "DKS"),
    ]:
        gen_cer(pname, pclass, pshort)
        gen_clinical_investigation_plan(pname, pshort)
        gen_csr(pname, pshort)
    for pshort, hosp, hshort in [
        ("CSP", C["hosp1"], "CHR"), ("CSP", C["hosp2"], "UKE"),
        ("OFN", C["hosp3"], "LMU"), ("OFN", C["hosp4"], "UKH"),
    ]:
        gen_investigator_agreement(pshort, hosp, hshort)
    # Literature search results
    for pname, pshort in [(C["prod1_name"], "CSP"), (C["prod2_name"], "OFN"), (C["prod3_name"], "DKS")]:
        rows = []
        for i in range(1, 26):
            rows.append((
                f"LIT-{pshort}-{i:03d}",
                f"Author{i} et al.",
                str(random.randint(2015, 2023)),
                random.choice(["Lancet", "NEJM", "BMJ", "JAMA", "Eur Heart J", "Med Devices"]),
                f"Clinical validation of {pname[:20]} in {random.randint(50,500)} patients",
                random.choice(["Level 1a", "Level 1b", "Level 2a", "Level 2b", "Level 3"]),
                random.choice(["Eingeschlossen", "Eingeschlossen", "Ausgeschlossen"]),
            ))
        sheets = [
            ("Literatursuche", ["Lit-ID", "Autoren", "Jahr", "Zeitschrift", "Titel (gekürzt)",
                                 "Evidence Level", "Status"],
             rows, [12, 20, 8, 18, 45, 12, 16]),
        ]
        make_xlsx_doc(FOLDER_CLIN, f"LIT_{pshort}_Literaturrecherche_Ergebnisse.xlsx",
                      f"Literaturrecherche Ergebnistabelle – {pname}", sheets)
    # Ethics committee approvals and patient consent forms
    for hosp_short, hosp in [("CHR", C["hosp1"]), ("UKE", C["hosp2"])]:
        sections = [
            ("Votum der Ethikkommission",
             f"Die Ethikkommission der {hosp} hat in ihrer Sitzung vom "
             f"{ds(2023, 2, random.randint(10,28))} das vorgelegte Studienprotokoll geprüft und "
             "erteilt hiermit ihr zustimmendes Votum."),
            ("Bedingungen", "Die Studie ist gemäß den Bestimmungen der Deklaration von Helsinki "
             "und ISO 14155:2020 durchzuführen. Alle Patienten sind vor Studieneinschluss "
             "vollständig aufzuklären und schriftlich einzuwilligen."),
        ]
        save_pdf(FOLDER_CLIN, f"ETHI_{hosp_short}_Ethikkommissionsvotum.pdf",
                 f"Ethikkommissionsvotum – {hosp}", sections)
    # Patient information
    sections = [
        ("Patienteninformation und Einwilligungserklärung",
         "Sehr geehrte Patientin, sehr geehrter Patient,\n\nwir laden Sie ein, an einer klinischen "
         f"Studie mit dem Medizinprodukt {C['prod1_name']} teilzunehmen. Ihre Teilnahme ist "
         "freiwillig. Sie können jederzeit ohne Angabe von Gründen aus der Studie ausscheiden."),
        ("Was wird untersucht?",
         f"{C['prod1_name']} ist ein nicht-invasives Diagnostikgerät. In der Studie wird untersucht, "
         "wie gut das Gerät bei Patienten mit Ihrer Erkrankung arbeitet. Es werden keine zusätzlichen "
         "Eingriffe an Ihnen vorgenommen, die über die normale Routineversorgung hinausgehen."),
        ("Einwilligungserklärung",
         "Ich, Unterzeichnete(r), habe die Patienteninformation gelesen und verstanden. "
         "Ich willige freiwillig in meine Studienteilnahme ein.\n\n"
         "___________________\t\t___________________\n"
         "Datum\t\t\t\tUnterschrift Patient"),
    ]
    save_pdf(FOLDER_CLIN, "CONSENT_Patientenaufklaerung_Einwilligungserklaerung.pdf",
             "Patienteninformation und Einwilligungserklärung", sections)


# ══════════════════════════════════════════════════════════════════════
# FOLDER 09 – IP / PATENTE
# ══════════════════════════════════════════════════════════════════════
FOLDER_IP = "09_IP_Patente"


def gen_patent_application(num, title, filing_date_year, status):
    sections = [
        ("Patentanmeldung / Patenterteilung",
         f"Anmeldenummer: {random.choice(['EP', 'DE'])}{random.randint(10,23)}{random.randint(100000,999999)}.{random.randint(0,9)}\n"
         f"Titel: {title}\nAnmelder: {C['name']}, {C['full_address']}\n"
         f"Erfinder: {random.choice([C['cto'], 'Dr. Lena Fischer', 'Jan Hoffmann', 'Dr. Marcus Vogt'])}\n"
         f"Anmeldetag: {ds(filing_date_year, random.randint(1,12), random.randint(1,28))}\n"
         f"Status: {status}"),
        ("Schutzanspruch 1",
         f"Vorrichtung zur {title.lower()}, umfassend:\n"
         "a) einen Sensor zur Erfassung physiologischer Messgrößen;\n"
         "b) eine Verarbeitungseinheit zur Signalverarbeitung;\n"
         "c) eine Ausgabeeinheit zur Darstellung der Messergebnisse;\n"
         "dadurch gekennzeichnet, dass die Verarbeitungseinheit einen neuartigen Algorithmus "
         "zur Rauschunterdrückung und Mustererkennung umfasst."),
        ("Technischer Hintergrund",
         f"Die vorliegende Erfindung betrifft das Gebiet der medizinischen Diagnostik, "
         "insbesondere Methoden und Vorrichtungen zur nicht-invasiven Messung und Analyse "
         "biologischer Parameter. Der Stand der Technik wird durch folgende Nachveröffentlichungen "
         "repräsentiert: [Zitation 1], [Zitation 2], [Zitation 3]."),
    ]
    save_pdf(FOLDER_IP, f"PAT_{num:03d}_{sfn(title[:25])}_Status_{sfn(status[:8])}.pdf",
             f"Patent – {title}", sections)


def gen_patent_portfolio():
    rows = []
    patents = [
        ("Kardiologisches Echtzeitanalyse-Verfahren", 2015, "Erteilt EP"),
        ("Adaptiver Signalverarbeitungsalgorithmus für kardiovaskuläre Biosignale", 2017, "Erteilt EP+US"),
        ("Navigationssystem für orthopädische Chirurgie", 2018, "Erteilt EP"),
        ("Mikrofluidik-basierter IVD-Schnelltest", 2019, "Erteilt EP"),
        ("KI-gestützte Bildanalyse für medizinische Diagnostik", 2020, "Angemeldet PCT"),
        ("Biosensor-Array zur Multiparameter-Messung", 2021, "Angemeldet EP"),
        ("Drahtlose Datenübertragung für implantierbare Sensoren", 2021, "Angemeldet DE"),
        ("Verfahren zur Qualitätskontrolle von IVD-Reagenzien", 2022, "Angemeldet EP"),
        ("Verfahren zur patientenspezifischen Normierung von Messdaten", 2022, "Angemeldet PCT"),
        ("Smarter Algorithmus für false-positive Reduktion", 2023, "Provisional US"),
    ]
    for i, (title, year, status) in enumerate(patents, 1):
        rows.append((f"BTP-P-{i:03d}", title, str(year),
                     random.choice(["DE/EP", "EP", "EP/US", "PCT", "PCT/EP"]),
                     status, random.choice([C["cto"], "Dr. Lena Fischer", "Dr. Marcus Vogt"]),
                     f"Anhängig bis {year + random.randint(18,22)}" if "Erteilt" not in status else "Aktiv"))
    sheets = [
        ("Patentportfolio", ["Patent-ID", "Titel", "Anmeldejahr", "Schutzgebiet",
                              "Status", "Erfinder", "Laufzeit"],
         rows, [12, 50, 12, 15, 18, 25, 20]),
    ]
    make_xlsx_doc(FOLDER_IP, "IP_001_Patentportfolio_Uebersicht.xlsx",
                  f"Patentportfolio {C['name']} – Stand 2024", sheets)


def gen_fto_opinion(prod_name, prod_short):
    sections = [
        ("Freedom-to-Operate Analyse (FTO)",
         f"Produkt: {prod_name}\nGutachter: {C['patent_anwalt']}\n"
         f"Auftragsdatum: {ds(2022, random.randint(1,12), random.randint(1,28))}\n"
         f"Auftraggeber: {C['name']}"),
        ("Zusammenfassung",
         f"Das FTO-Gutachten für {prod_name} hat folgende Schutzrechte Dritter analysiert, "
         f"die potenziell relevant sein könnten. Analysierte Schutzrechte: {random.randint(25,80)} Patente "
         "und Gebrauchsmuster im relevanten technologischen Bereich.\n\n"
         f"Ergebnis: Keine eindeutigen Verletzungen identifiziert. {random.randint(2,5)} Patente "
         "wurden als 'Watch List'-Kandidaten eingestuft und werden laufend überwacht."),
        ("Risikoeinschätzung",
         "Gesamtrisikobewertung: MITTEL. Die identifizierten Schutzrechte könnten theoretisch "
         "geltend gemacht werden, werden aber als lizenzierbar oder umgehbar eingestuft. "
         "Empfehlung: Designänderung an Komponente X prüfen; alternativ Lizenzierung erwägen."),
    ]
    save_pdf(FOLDER_IP, f"FTO_{prod_short}_Freedom_to_Operate_Opinion.pdf",
             f"Freedom-to-Operate Opinion – {prod_name}", sections)


def gen_license_agreement():
    doc = Document()
    make_docx_header(doc, "Lizenzvertrag (In-Licensing)",
                     f"{C['name']} und Fraunhofer-Gesellschaft", confidential=True)
    add_clause(doc, 1, "Vertragsgegenstand", [
        f"Die Fraunhofer-Gesellschaft zur Förderung der angewandten Forschung e.V., München "
        f"(\"Lizenzgeber\") gewährt der {C['name']} (\"Lizenznehmer\") eine exklusive weltweite "
        "Lizenz zur kommerziellen Nutzung des Patents EP2344567B1 (\"Gegenstand der Lizenz\"), "
        "das ein neuartiges Verfahren zur kardiologischen Echtzeitmessung betrifft.",
        "Die Lizenz umfasst das Recht zur Herstellung, Vermarktung und zum Vertrieb von Produkten, "
        "die das lizenzierte Verfahren nutzen oder beinhalten.",
    ])
    add_clause(doc, 2, "Lizenzgebühren", [
        "Der Lizenznehmer zahlt an den Lizenzgeber eine einmalige Upfront-Lizenzgebühr von "
        "EUR 150.000 bei Vertragsabschluss.",
        "Zusätzlich zahlt der Lizenznehmer eine laufende Umsatzlizenzgebühr (Royalty) von "
        "2,5 % der mit Produkten erzielten Nettoumsätze, die das lizenzierte Verfahren nutzen, "
        "mindestens jedoch EUR 50.000 pro Jahr ab dem dritten Vertragsjahr.",
    ])
    add_clause(doc, 3, "Laufzeit", [
        "Der Vertrag gilt für die Schutzdauer des lizenzierten Patents, längstens bis "
        f"{ds(2035, 12, 31)}. Er kann von beiden Parteien mit einer Frist von 12 Monaten "
        "zum Jahresende gekündigt werden.",
    ])
    save_docx(doc, FOLDER_IP, "IP_002_Lizenzvertrag_Fraunhofer_Inlicensing.docx")


def gen_trademark_register():
    sections = [
        ("Markenregister",
         f"Inhaber: {C['name']}\nAnwalt: {C['patent_anwalt']}"),
        ("Eingetragene Marken", [
            ["Marken-Nr.", "Bezeichnung", "Klasse (Nizza)", "Eintragungsdatum", "Ablaufdatum", "Gebiet"],
            [f"DE{random.randint(30000000, 39999999)}", "Sentavia Precision", "10, 42", ds(2010, 3, 1), ds(2030, 3, 1), "Deutschland"],
            [f"EM{random.randint(10000000, 19999999)}", "Cardevio Pro", "10", ds(2018, 7, 15), ds(2028, 7, 15), "EU"],
            [f"EM{random.randint(10000000, 19999999)}", "Ostevo Navigator", "10, 44", ds(2019, 11, 10), ds(2029, 11, 10), "EU"],
            [f"EM{random.randint(10000000, 19999999)}", "Veridiq", "5, 10", ds(2021, 4, 22), ds(2031, 4, 22), "EU"],
            [f"WIPO{random.randint(1000000, 1999999)}", "BTP (Logo)", "10, 42, 44", ds(2022, 9, 1), ds(2032, 9, 1), "International (Madrid)"],
        ]),
    ]
    save_pdf(FOLDER_IP, "IP_003_Markenregister_Uebersicht.pdf",
             "Markenregister – Eingetragene Marken", sections)


def generate_09_ip():
    print("  Generating 09_IP_Patente...")
    gen_patent_portfolio()
    gen_license_agreement()
    gen_trademark_register()
    patent_data = [
        ("Kardiologisches Echtzeitanalyse-Verfahren", 2015, "Erteilt"),
        ("Adaptiver Biosignal-Verarbeitungsalgorithmus", 2017, "Erteilt"),
        ("Navigationssystem für orthopädische Chirurgie", 2018, "Erteilt"),
        ("Mikrofluidik-basierter IVD-Schnelltest", 2019, "Erteilt"),
        ("KI-gestützte Bildanalyse Diagnostik", 2020, "Angemeldet PCT"),
        ("Biosensor-Array Multiparameter", 2021, "Angemeldet EP"),
        ("Drahtlose Datenübertragung Sensor", 2021, "Angemeldet DE"),
        ("Qualitätskontrolle IVD-Reagenzien", 2022, "Angemeldet EP"),
    ]
    for i, (title, year, status) in enumerate(patent_data, 1):
        gen_patent_application(i, title, year, status)
    for pname, pshort in [(C["prod1_name"], "CSP"), (C["prod2_name"], "OFN")]:
        gen_fto_opinion(pname, pshort)
    # NDAs with research partners
    for partner in ["Fraunhofer IPA", "TU München – Medizintechnik", "Helmholtz Zentrum München"]:
        doc = Document()
        make_docx_header(doc, "Vertraulichkeitsvereinbarung (NDA)",
                         f"{C['name']} und {partner}", confidential=True)
        add_clause(doc, 1, "Gegenstand", [
            f"Dieser Vertrag regelt den Austausch vertraulicher Informationen zwischen {C['name']} "
            f"und {partner} im Rahmen einer Forschungskooperation.",
        ])
        add_clause(doc, 2, "Vertraulichkeitspflicht", [
            "Beide Parteien verpflichten sich, alle ausgetauschten vertraulichen Informationen "
            "streng vertraulich zu behandeln und nicht an Dritte weiterzugeben.",
            "Vertraulich sind insbesondere: technische Spezifikationen, Forschungsdaten, "
            "Prototypen, Geschäftsdaten und Produktentwicklungen.",
        ])
        add_clause(doc, 3, "Laufzeit", [
            "Diese Vereinbarung gilt für die Dauer der Zusammenarbeit und für einen Zeitraum "
            "von 5 Jahren nach deren Beendigung.",
        ])
        save_docx(doc, FOLDER_IP, f"NDA_{sfn(partner[:20])}_Vertraulichkeitsvereinbarung.docx")


# ══════════════════════════════════════════════════════════════════════
# FOLDER 10 – VERSICHERUNGEN
# ══════════════════════════════════════════════════════════════════════
FOLDER_INS = "10_Versicherungen"


def gen_insurance_policy(ins_type, ins_num, insurer, coverage, premium):
    sections = [
        ("Versicherungsschein",
         f"Versicherungsnehmer: {C['name']}, {C['full_address']}\n"
         f"Versicherungsart: {ins_type}\nVersicherungsnummer: {ins_num}\n"
         f"Versicherer: {insurer}\nVersicherungssumme: {eur(coverage)}\n"
         f"Jahresprämie (netto): {eur(premium)}\n"
         f"Versicherungsjahr: 1. Januar 2024 – 31. Dezember 2024"),
        ("Deckungsumfang",
         f"Die Versicherung deckt {ins_type.lower()} für alle Produkte der {C['name']}, "
         f"insbesondere {C['prod1_name']}, {C['prod2_name']} und {C['prod3_name']}. "
         "Geographischer Geltungsbereich: Weltweit (exkl. USA/Kanada für Produkthaftung: "
         "gesonderte US-Policy)."),
        ("Besondere Klauseln", [
            ["Klausel", "Inhalt"],
            ["Rückwärtsdeckung", f"Deckung für Ereignisse ab {ds(2018, 1, 1)}"],
            ["Nachhaftung", "3 Jahre nach Vertragsende"],
            ["Selbstbehalt", f"EUR {random.randint(10,50)*1000:,} je Schadensfall"],
            ["Serienschadenklausel", "Ja, mit Sublimit EUR 20.000.000"],
        ]),
    ]
    save_pdf(FOLDER_INS, f"VER_{ins_type[:10].replace(' ','_')}_{ins_num[:10]}_Police.pdf",
             f"Versicherungspolice – {ins_type}", sections)


def generate_10_versicherungen():
    print("  Generating 10_Versicherungen...")
    policies = [
        ("Produkthaftpflichtversicherung", "PL-DE-2024-00123", "Zurich Insurance plc", 25_000_000, 185_000),
        ("Klinische Studienversicherung", "CS-DE-2024-00456", "Allianz Global Corporate & Specialty", 10_000_000, 42_000),
        ("D&O Versicherung (Directors & Officers)", "DO-DE-2024-00789", "AXA XL", 20_000_000, 95_000),
        ("Cyber-Versicherung", "CY-DE-2024-00321", "Hiscox", 5_000_000, 38_000),
        ("Berufshaftpflicht (E&O)", "EO-DE-2024-00654", "HDI Global SE", 8_000_000, 55_000),
        ("Betriebsunterbrechungsversicherung", "BU-DE-2024-00987", "Munich Re", 15_000_000, 78_000),
        ("Gebäude- und Inhaltsversicherung", "GI-DE-2024-00111", "Generali Deutschland AG", 12_000_000, 32_000),
    ]
    for ins_type, ins_num, insurer, coverage, premium in policies:
        gen_insurance_policy(ins_type, ins_num, insurer, coverage, premium)
    # Insurance overview spreadsheet
    sheets = [
        ("Versicherungsübersicht", ["Versicherungsart", "Versicherer", "Police-Nr.", "Versicherungssumme EUR",
                                     "Jahresprämie EUR", "Gültigkeit"],
         [(p[0], p[2], p[1], f"{p[3]:,}", f"{p[4]:,}", "01.01.2024–31.12.2024") for p in policies],
         [35, 32, 20, 22, 18, 22]),
    ]
    make_xlsx_doc(FOLDER_INS, "VER_Uebersicht_Versicherungsportfolio_2024.xlsx",
                  "Versicherungsportfolio 2024 – Übersicht", sheets)


# ══════════════════════════════════════════════════════════════════════
# FOLDER 11 – IMMOBILIEN
# ══════════════════════════════════════════════════════════════════════
FOLDER_PROP = "11_Immobilien"


def gen_lease_agreement():
    doc = Document()
    make_docx_header(doc, "Gewerbemietvertrag – Freimannstraße 45, München",
                     f"{C['name']} als Mieter", confidential=True)
    add_clause(doc, 1, "Mietobjekt", [
        f"Gegenstand dieses Mietvertrags sind die Räumlichkeiten in der Freimannstraße 45, "
        f"80939 München, bestehend aus: Büroflächen (ca. 2.800 m²), Laborräumen (ca. 1.200 m²), "
        "Reinraumanlage (ca. 800 m², Klasse ISO 7), Lagerräumen (ca. 500 m²).",
        "Die Gesamtmietfläche beträgt ca. 5.300 m². Das Mietobjekt ist abgebildet in den "
        "Grundrissplänen gemäß Anlage 1.",
    ])
    add_clause(doc, 2, "Mietzins und Nebenkosten", [
        "Die monatliche Kaltmiete beträgt EUR 127.200,00 (EUR 24,00/m² × 5.300 m²) zzgl. "
        "der gesetzlichen Mehrwertsteuer. Die Miete ist jeweils zum 3. Werktag eines Monats "
        "im Voraus auf das Konto des Vermieters zu entrichten.",
        "Zusätzlich zur Kaltmiete trägt der Mieter die Nebenkosten gemäß Anlage 2 "
        "(Betriebskostenabrechnung). Hierzu leistet der Mieter monatliche Vorauszahlungen "
        "von EUR 18.000,00.",
        "Eine jährliche Mietanpassung gemäß VPI (Verbraucherpreisindex des Statistischen "
        "Bundesamts) ist vereinbart. Anpassungen erfolgen erstmals nach 2 Jahren Mietzeit.",
    ])
    add_clause(doc, 3, "Mietdauer und Optionen", [
        f"Das Mietverhältnis beginnt am {ds(2020, 1, 1)} und läuft bis zum {ds(2030, 12, 31)} "
        "(feste Laufzeit 10 Jahre).",
        "Der Mieter hat das Recht, das Mietverhältnis zweimal um je 5 Jahre zu verlängern "
        "(Verlängerungsoption), sofern die Option mindestens 12 Monate vor Ablauf der "
        "jeweiligen Laufzeit schriftlich ausgeübt wird.",
    ])
    add_clause(doc, 4, "Instandhaltung und Umbau", [
        "Der Vermieter ist für wesentliche Instandhaltungsmaßnahmen (Dach, Fassade, "
        "technische Gebäudeausrüstung) verantwortlich. Der Mieter übernimmt die Instandhaltung "
        "der Innenräume und der mieterseits eingebrachten Einrichtungen.",
        "Umbaumaßnahmen bedürfen der schriftlichen Zustimmung des Vermieters. Bei Vertragsende "
        "ist der ursprüngliche Zustand wiederherzustellen, sofern nichts anderes vereinbart.",
    ])
    save_docx(doc, FOLDER_PROP, "MIETE_001_Gewerbemietvertrag_Freimannstrasse_45.docx")


def gen_cleanroom_contract():
    sections = [
        ("Werkvertrag Reinraumausbau",
         f"Auftraggeber: {C['name']}, {C['full_address']}\n"
         "Auftragnehmer: Reinraum-Technik Süd GmbH, Oberschleißheim\n"
         f"Auftragsnummer: RR-2022-0047\nDatum: {ds(2022, 3, 1)}"),
        ("Leistungsumfang",
         "Planung, Lieferung und Installation einer Reinraumanlage gemäß ISO 14644-1 Klasse ISO 7 "
         "in den Räumlichkeiten Freimannstraße 45, München (Halle 3, Erdgeschoss).\n\n"
         "Fläche: 800 m²\nReinraumklasse: ISO 7 (≤ 352.000 Partikel/m³ ≥ 0,5 μm)\n"
         "Luftwechselrate: ≥ 40/h\nGFMPP-Filterung: H14\nDruck: +15 Pa gegenüber angrenzenden Bereichen"),
        ("Vergütung",
         "Pauschalpreis: EUR 1.850.000 zzgl. MwSt.\n"
         "Zahlungsplan: 30 % Anzahlung, 40 % nach Rohbaufertigstellung, 30 % nach Abnahme."),
        ("Abnahme und Qualifizierung",
         "Die Abnahme erfolgt nach erfolgreicher IQ/OQ/PQ (Installation/Operational/Performance "
         "Qualification). Das Protokoll der Qualifizierung wird als Anlage zum Abnahmeprotokoll "
         f"beigefügt. Abnahmedatum: {ds(2022, 11, 15)}."),
    ]
    save_pdf(FOLDER_PROP, "PROP_002_Reinraumausbau_Werkvertrag.pdf",
             "Werkvertrag Reinraumausbau ISO-Klasse 7", sections)


def gen_equipment_lease():
    doc = Document()
    make_docx_header(doc, "Leasingvertrag – Medizintechnische Prüfausrüstung",
                     f"{C['name']} und Siemens Financial Services GmbH")
    add_clause(doc, 1, "Leasinggegenstand", [
        "Gegenstand dieses Leasingvertrags sind die in Anlage 1 aufgeführten medizintechnischen "
        f"Prüfgeräte und Laborausrüstungen mit einem Gesamtanschaffungswert von EUR 1.240.000. "
        "Die Ausrüstung wird für F&E- und Qualitätsprüfungszwecke genutzt.",
    ])
    add_clause(doc, 2, "Leasinggebühr und Laufzeit", [
        "Die monatliche Leasinggebühr beträgt EUR 22.400 zzgl. MwSt. "
        "Die Leasinglaufzeit beträgt 60 Monate ab Übergabe der Ausrüstung.",
        "Nach Ablauf der Laufzeit besteht eine Kaufoption zum Restwert von EUR 65.000.",
    ])
    save_docx(doc, FOLDER_PROP, "PROP_003_Leasingvertrag_Pruefausruestung.docx")


def generate_11_immobilien():
    print("  Generating 11_Immobilien...")
    gen_lease_agreement()
    gen_cleanroom_contract()
    gen_equipment_lease()
    # Property overview
    sections = [
        ("Immobilienportfolio Übersicht",
         f"Stand: Januar 2024 | Unternehmen: {C['name']}"),
        ("Standorte", [
            ["Standort", "Nutzung", "Fläche m²", "Eigentum/Miete", "Jahreskosten EUR"],
            [C["full_address"], "Hauptsitz, Produktion, Labor, Reinraum", "5.300", "Miete", "1.745.000"],
            ["Freimannstraße 51, 80939 München", "Logistik und Versandlager", "1.200", "Miete", "180.000"],
            ["Maximilianstraße 35, 80539 München", "Repräsentanzbüro / Showroom", "350", "Miete", "168.000"],
        ]),
    ]
    save_pdf(FOLDER_PROP, "PROP_001_Immobilienportfolio_Uebersicht.pdf",
             "Immobilienportfolio Übersicht", sections)
    # Inspection reports
    for i, topic in enumerate(["Brandschutz-Inspektion 2023", "Reinraum-Requalifizierung ISO 14644 2023",
                                "Elektrische Anlage Prüfbericht DGUV V3 2023"], 1):
        sections = [
            ("Prüfbericht", f"Objekt: {C['full_address']}\nPrüfer: Zugelassener Sachverständiger\n"
             f"Datum: {ds(2023, random.randint(2,11), random.randint(1,28))}\nBefund: Keine wesentlichen Mängel"),
        ]
        save_pdf(FOLDER_PROP, f"PROP_{i+3:03d}_Pruefbericht_{sfn(topic[:25])}.pdf", topic, sections)


# ══════════════════════════════════════════════════════════════════════
# FOLDER 12 – IT / DIGITAL
# ══════════════════════════════════════════════════════════════════════
FOLDER_IT = "12_IT_Digital"


def gen_it_infrastructure():
    sections = [
        ("IT-Infrastruktur Übersicht",
         f"Stand: Januar 2024 | IT-Leitung: Benjamin Wolf\n"
         f"Mitarbeiter: {C['employees_2023']} FTE | IT-Standorte: 3"),
        ("Kernsysteme", [
            ["System", "Anbieter", "Hosting", "Nutzer", "MDR-Relevanz"],
            [C["erp"], "SAP SE", "SAP Cloud (Frankfurt)", "420", "Ja – SOX-relevant"],
            [C["qms_tool"], "MasterControl Inc.", "US-Cloud / EU-Replika", "185", "Ja – Kritisch (FDA 21 CFR Part 11)"],
            [C["clinical_db"], "Vanderbilt University", "On-Premise + Backup-Cloud", "45", "Ja – Klinische Daten"],
            ["EUDAMED", "EU-Kommission", "EU Government Cloud", "8", "Ja – Regulatorisch"],
            ["SAP Concur (Reisekostenabrechnung)", "SAP SE", "Cloud", "all", "Nein"],
            ["Microsoft 365 / Teams", "Microsoft", "EU Cloud", "all", "Nein"],
        ]),
        ("Netzwerkarchitektur",
         "Das Unternehmensnetz ist segmentiert in: (1) Produktionsnetz (OT – Operational Technology), "
         "(2) IT-Büronetz, (3) F&E-Netz (Entwicklungsumgebungen), (4) Gastnetz. "
         "Zwischen den Segmenten sind Firewalls der nächsten Generation (Next-Gen Firewall) "
         "von Palo Alto Networks implementiert. Remote-Zugriff über Zero-Trust-Architektur "
         "(Zscaler Private Access)."),
        ("Cybersecurity", "Implementierte Maßnahmen gemäß MDR Art. 5 Abs. 5 und IEC 62443:\n"
         "• Multi-Faktor-Authentifizierung (MFA) für alle Systeme\n"
         "• Endpoint Detection & Response (EDR) von CrowdStrike\n"
         "• SIEM-System (Splunk) für Security Monitoring\n"
         "• Regelmäßige Penetrationstests (jährlich extern, quartalsweise intern)\n"
         "• Business Continuity Plan und Disaster Recovery (RTO: 4h, RPO: 1h)"),
    ]
    save_pdf(FOLDER_IT, "IT_001_Infrastruktur_Uebersicht.pdf",
             "IT-Infrastruktur Übersicht", sections)


def gen_svmp():
    sections = [
        ("Software Validation Master Plan (SVMP)",
         f"Gültig ab: {ds(2023, 1, 1)}\nVersion: 2.0\nVerantwortlich: Benjamin Wolf, IT-Leitung\n"
         f"Genehmigt: {C['qra']}"),
        ("Geltungsbereich",
         "Dieser SVMP gilt für alle computerisierten Systeme und Software, die bei der "
         f"{C['name']} zur Unterstützung regulierter Tätigkeiten (GxP, MDR, ISO 13485) "
         "eingesetzt werden."),
        ("Validierungskategorien", [
            ["System", "Kategorie", "Risiko", "Validierungsumfang"],
            [C["qms_tool"], "Cat 4 (Konfigurierbar)", "Hoch", "IQ/OQ/PQ + Revalidierung jährlich"],
            [C["erp"], "Cat 4", "Mittel", "IQ/OQ + Change Control"],
            [C["clinical_db"], "Cat 5 (Custom Code)", "Hoch", "Vollvalidierung + GAMP 5"],
            ["Microsoft 365", "Cat 1 (Infrastruktur)", "Niedrig", "Vendor Assessment"],
        ]),
        ("Validierungsablauf",
         "1. Risikoklassifizierung und Kategorisierung des Systems\n"
         "2. Erstellen des Validierungsplans (VP) und der User Requirements Specification (URS)\n"
         "3. Erstellung und Ausführung der Testprotokolle (IQ/OQ/PQ)\n"
         "4. Bewertung der Testergebnisse und Erstellung des Validierungsberichts\n"
         "5. Freigabe des Systems durch QA und IT-Leitung\n"
         "6. Laufende Überwachung und periodische Revalidierung"),
    ]
    save_pdf(FOLDER_IT, "IT_002_Software_Validation_Master_Plan.pdf",
             "Software Validation Master Plan (SVMP)", sections)


def gen_cybersecurity_policy():
    doc = Document()
    make_docx_header(doc, "Cybersecurity Policy für Medizinprodukte",
                     f"Gemäß MDR Art. 5 Abs. 5 und MDCG 2019-16")
    add_section(doc, "1. Geltungsbereich", [
        f"Diese Richtlinie gilt für alle vernetzten Medizinprodukte der {C['name']}, "
        f"insbesondere {C['prod1_name']} und {C['prod2_name']}, sowie alle zugehörigen "
        "IT-Systeme und Software-Komponenten.",
        "Die Richtlinie setzt die MDCG 2019-16 (Cybersecurity for Medical Devices) und "
        "IEC 81001-5-1:2021 (Medical device software – IT-security) um.",
    ])
    add_section(doc, "2. Sicherheitsanforderungen", [
        "Alle vernetzten Medizinprodukte müssen folgende Sicherheitsanforderungen erfüllen:\n"
        "• Authentifizierung: Passwort-Schutz oder biometrische Authentifizierung\n"
        "• Verschlüsselung: TLS 1.3 für alle Netzwerkkommunikation\n"
        "• Audit Trail: Unveränderliche Protokollierung aller Zugriffe und Änderungen\n"
        "• Update-Mechanismus: Sicherer und überprüfbarer Software-Update-Prozess\n"
        "• Härtung: Deaktivierung aller nicht benötigten Dienste und Ports",
        "Software-Komponenten werden gemäß SBOM (Software Bill of Materials) katalogisiert "
        "und auf bekannte Schwachstellen (CVE-Datenbank) überwacht.",
    ])
    add_section(doc, "3. Vulnerability Management", [
        "Bekannte Schwachstellen werden nach CVSS-Score klassifiziert und innerhalb folgender "
        "Fristen behoben: Kritisch (CVSS ≥ 9.0): 30 Tage; Hoch (7.0–8.9): 90 Tage; "
        "Mittel (4.0–6.9): 180 Tage.",
    ])
    save_docx(doc, FOLDER_IT, "IT_003_Cybersecurity_Policy_Medizinprodukte.docx")


def gen_cloud_dpa():
    sections = [
        ("Auftragsverarbeitungsvertrag (AVV) nach Art. 28 DSGVO",
         f"Verantwortlicher: {C['name']}, {C['full_address']}\n"
         "Auftragsverarbeiter: SAP SE, Dietmar-Hopp-Allee 16, 69190 Walldorf\n"
         f"Vertragsgegenstand: {C['erp']} – Cloud-Betrieb und Datenspeicherung"),
        ("Gegenstand und Dauer", "Die Verarbeitung personenbezogener Daten durch den Auftragsverarbeiter "
         "erfolgt ausschließlich auf Weisung des Verantwortlichen. Der Vertrag gilt für die Dauer "
         "der Hauptleistungsvereinbarung (Lizenzvertrag SAP)."),
        ("Technische und organisatorische Maßnahmen (TOM)", [
            ["Maßnahme", "Beschreibung", "Implementiert"],
            ["Pseudonymisierung", "Trennungs- und Pseudonymisierungskonzept", "Ja"],
            ["Verschlüsselung", "AES-256 at rest, TLS 1.3 in transit", "Ja"],
            ["Vertraulichkeit", "Rollenbasiertes Zugriffskonzept (RBAC)", "Ja"],
            ["Verfügbarkeit", "99,9 % SLA, georedundantes Rechenzentrum", "Ja"],
            ["Auslagerung", "Keine ohne vorherige Zustimmung des Verantwortlichen", "Ja"],
        ]),
    ]
    save_pdf(FOLDER_IT, "IT_004_AVV_DSGVO_SAP_Cloud.pdf",
             "Auftragsverarbeitungsvertrag DSGVO – SAP Cloud", sections)


def generate_12_it():
    print("  Generating 12_IT_Digital...")
    gen_it_infrastructure()
    gen_svmp()
    gen_cybersecurity_policy()
    gen_cloud_dpa()
    for title, fname in [
        ("IT-Sicherheitskonzept 2024", "IT_005_Sicherheitskonzept_2024.docx"),
        ("Disaster Recovery Plan", "IT_006_Disaster_Recovery_Plan.docx"),
        ("Penetrationstest Abschlussbericht 2023", "IT_007_Penetrationstest_2023_Bericht.pdf"),
    ]:
        if fname.endswith(".docx"):
            d = Document()
            make_docx_header(d, title)
            d.add_paragraph(f"Dieses Dokument beschreibt {title.lower()} der {C['name']}. "
                            "Es ist vertraulich und für den internen Gebrauch bestimmt.")
            save_docx(d, FOLDER_IT, fname)
        else:
            save_pdf(FOLDER_IT, fname, title,
                     [("Zusammenfassung", f"Bericht über {title.lower()} der {C['name']}. "
                       f"Datum: {ds(2023, random.randint(9,11), random.randint(1,28))}.")])
    # AVV for MasterControl
    sections = [
        ("AVV MasterControl",
         f"Verantwortlicher: {C['name']}\nAuftragsverarbeiter: MasterControl Inc., Salt Lake City, Utah, USA\n"
         "Rechtsgrundlage: Art. 46 DSGVO – Standarddatenschutzklauseln der EU-Kommission"),
    ]
    save_pdf(FOLDER_IT, "IT_008_AVV_DSGVO_MasterControl.pdf",
             "AVV DSGVO – MasterControl Quality Excellence Suite", sections)


# ══════════════════════════════════════════════════════════════════════
# FOLDER 13 – COMPLIANCE / RECHT
# ══════════════════════════════════════════════════════════════════════
FOLDER_COMP = "13_Compliance_Recht"


def gen_dsgvo_verzeichnis():
    sheets = [
        ("Verarbeitungsverzeichnis", ["Nr.", "Verarbeitungstätigkeit", "Zweck", "Rechtsgrundlage",
                                       "Datenarten", "Empfänger", "Löschfrist"],
         [
             ("VVT-001", "Mitarbeiterverwaltung", "HR, Lohnabrechnung", "Art. 6(1)(b) DSGVO",
              "Stammdaten, Gehalt", "Steuerberater, Krankenkassen", "10 Jahre nach Ausscheiden"),
             ("VVT-002", "Kundenbeziehungsmanagement", "Vertrieb", "Art. 6(1)(b) DSGVO",
              "Kontaktdaten, Bestellhistorie", "CRM (Intern)", "5 Jahre nach letztem Kontakt"),
             ("VVT-003", "Klinische Studien", "Produktzulassung, PMS", "Art. 9 DSGVO i.V.m. § 22 BDSG",
              "Gesundheitsdaten, Einwilligung", "Prüfzentren, Behörden", "15 Jahre nach Studienende"),
             ("VVT-004", "Lieferantenmanagement", "Einkauf", "Art. 6(1)(f) DSGVO",
              "Kontaktdaten, Verträge", "Intern, Steuerberater", "10 Jahre nach Vertragsende"),
             ("VVT-005", "Videoüberwachung Produktion", "Sicherheit, Qualitätssicherung", "Art. 6(1)(f) DSGVO",
              "Videodaten", "Intern", "72 Stunden (FIFO)"),
             ("VVT-006", "EUDAMED-Meldungen", "Regulatorisch", "Art. 6(1)(c) DSGVO",
              "UDI, Produktdaten", "EU-Behörden, BfArM", "Permanent (regulatorisch)"),
         ], [10, 28, 22, 22, 28, 25, 28]),
    ]
    make_xlsx_doc(FOLDER_COMP, "DSGVO_001_Verarbeitungsverzeichnis.xlsx",
                  "Verzeichnis von Verarbeitungstätigkeiten gemäß Art. 30 DSGVO", sheets)


def gen_compliance_policy(title, num):
    doc = Document()
    make_docx_header(doc, title, f"Version 2.0 | Gültig ab {ds(2023, 1, 1)}")
    doc.add_paragraph(
        f"Diese Richtlinie der {C['name']} legt die Grundsätze und Anforderungen für "
        f"{title.lower()} fest. Sie gilt für alle Mitarbeiter, Geschäftsführer und "
        "externe Dienstleister, die im Auftrag der Gesellschaft tätig sind."
    )
    add_clause(doc, 1, "Grundsätze", [
        "Die Gesellschaft bekennt sich zu ethischem und rechtmäßigem Geschäftsverhalten. "
        "Verstöße gegen diese Richtlinie werden konsequent verfolgt.",
        "Alle Mitarbeiter sind verpflichtet, diese Richtlinie zu kennen und einzuhalten. "
        "Bei Unklarheiten ist die Rechtsabteilung zu konsultieren.",
    ])
    add_clause(doc, 2, "Verbotene Handlungen", [
        "Untersagt sind insbesondere: Bestechung und Korruption, Interessenkonflikte ohne "
        "Offenlegung, Weitergabe vertraulicher Informationen an nicht-autorisierte Personen, "
        "Marktmanipulation oder Insiderhandel.",
        "Zuwendungen an Angehörige der Heilberufe (Ärzte, Kliniken) sind nur im Rahmen "
        "des § 307d SGB V und der Berufsordnungen zulässig. Für Zuwendungen > EUR 50 ist "
        "eine schriftliche Genehmigung erforderlich.",
    ])
    add_clause(doc, 3, "Meldewesen", [
        "Verstöße oder Verdachtsfälle sind umgehend an die Rechtsabteilung oder über das "
        "anonyme Whistleblower-System zu melden. Es gilt ein absolutes Repressalienverbot.",
    ])
    save_docx(doc, FOLDER_COMP, f"COMP_{num:02d}_{sfn(title[:25])}.docx")


def gen_legal_opinion(topic, num):
    sections = [
        ("Rechtsgutachten",
         f"Auftraggeber: {C['name']}, {C['full_address']}\n"
         f"Auftragnehmer: {C['anwalt']}\nBetreff: {topic}\n"
         f"Datum: {ds(2023, random.randint(1,12), random.randint(1,28))}"),
        ("Sachverhalt",
         f"Das folgende Gutachten befasst sich mit {topic.lower()}. "
         f"Der Auftraggeber hat folgende Fragen zur rechtlichen Beurteilung vorgelegt: "
         "(1) Wie ist die aktuelle Rechtslage zu beurteilen? "
         "(2) Welche Handlungsoptionen bestehen? "
         "(3) Welche Risiken sind zu beachten?"),
        ("Gutachten",
         "Nach eingehender Prüfung der Sach- und Rechtslage kommt die Kanzlei zu folgendem "
         "Ergebnis: Die vorliegende Situation ist rechtlich komplex. Aus den relevanten "
         "Vorschriften (MDR 2017/745, HWG, MPG, UWG) ergibt sich folgende Einschätzung: "
         "Das beschriebene Vorgehen ist grundsätzlich zulässig, bedarf jedoch der "
         "Einhaltung bestimmter Voraussetzungen."),
        ("Empfehlung",
         "Es wird empfohlen: (1) Interne Richtlinie anpassen; "
         "(2) Compliance-Schulung für betroffene Mitarbeiter durchführen; "
         "(3) Laufendes Monitoring der Rechtsentwicklung."),
    ]
    save_pdf(FOLDER_COMP, f"LEGAL_{num:02d}_Rechtsgutachten_{sfn(topic[:25])}.pdf",
             f"Rechtsgutachten – {topic}", sections, confidential=True)


def gen_whistleblower_system():
    doc = Document()
    make_docx_header(doc, "Hinweisgebersystem gemäß HinSchG",
                     f"{C['name']} – Richtlinie und Verfahrensbeschreibung")
    add_section(doc, "Einleitung", [
        f"Die {C['name']} hat gemäß dem Hinweisgeberschutzgesetz (HinSchG) vom 2. Juli 2023 "
        "ein internes Hinweisgebersystem eingerichtet. Dieses System ermöglicht es Mitarbeitern "
        "und externen Dritten, auf potenzielle Verstöße gegen rechtliche Vorschriften aufmerksam "
        "zu machen, ohne Nachteile befürchten zu müssen.",
    ])
    add_section(doc, "Meldekanäle", [
        "Meldungen können auf folgenden Wegen erstattet werden:\n"
        "(1) Online-Portal: compliance.sentavia-precision.de (verschlüsselt, anonym möglich)\n"
        "(2) Telefon-Hotline: +49 89 32178-999 (vertraulich)\n"
        "(3) Schriftlich: Rechtsabteilung, persönlich adressiert (Umschlag 'Vertraulich')\n"
        "(4) Extern: Bundesamt der Justiz als externe Meldestelle",
    ])
    add_section(doc, "Schutz für Hinweisgeber", [
        "Hinweisgeber, die in gutem Glauben Meldung erstatten, genießen umfassenden Schutz. "
        "Repressalien (Kündigung, Degradierung, Schikane) sind verboten und werden strafrechtlich "
        "verfolgt. Die Identität des Hinweisgebers wird vertraulich behandelt.",
    ])
    save_docx(doc, FOLDER_COMP, "COMP_HinSchG_Hinweisgebersystem_Richtlinie.docx")


def generate_13_compliance():
    print("  Generating 13_Compliance_Recht...")
    gen_dsgvo_verzeichnis()
    gen_whistleblower_system()
    compliance_policies = [
        "Anti-Korruptions-Richtlinie",
        "HWG-Compliance – Medizinproduktewerbung",
        "MBO-Ä und Antikorruptionsgesetz Gesundheitswesen",
        "Code of Conduct",
        "Interessenkonflikt-Richtlinie",
        "Sanktionslisten-Screening",
        "Datenschutzrichtlinie (DSGVO)",
        "Social Media Policy",
    ]
    for i, policy in enumerate(compliance_policies, 1):
        gen_compliance_policy(policy, i)
    legal_opinions = [
        "MDR-Konformität der Werbeaussagen Cardevio Pro",
        "HWG-Zulässigkeit der Fachmessepräsentation",
        "Antikorruptionsrechtliche Bewertung Sprecherhonorare",
        "Datenschutzrechtliche Bewertung klinischer Datenübermittlung in die USA",
        "Wettbewerbsrechtliche Einschätzung Markteinführungsstrategie",
        "Steuerliche Einschätzung Forschungszulage FZulG",
    ]
    for i, topic in enumerate(legal_opinions, 1):
        gen_legal_opinion(topic, i)
    # DSGVO documentation
    for title, fname in [
        ("Datenschutz-Folgenabschätzung (DSFA) – Klinische Daten", "DSGVO_002_DSFA_Klinische_Daten.pdf"),
        ("Datenschutzerklärung Website sentavia-precision.de", "DSGVO_003_Datenschutzerklaerung_Website.pdf"),
        ("Verzeichnis Auftragsverarbeiter", "DSGVO_004_Auftragsverarbeiter_Verzeichnis.xlsx"),
    ]:
        if fname.endswith(".xlsx"):
            sheets = [("AVV-Übersicht", ["Auftragsverarbeiter", "Sitz", "Vertragsart", "Zweck", "Status"],
                       [
                           ("SAP SE", "Walldorf, Deutschland", "AVV Art. 28 DSGVO", "ERP", "Unterzeichnet"),
                           ("MasterControl Inc.", "Utah, USA", "SCC + AVV", "QMS", "Unterzeichnet"),
                           ("Microsoft Ireland", "Dublin, Irland", "AVV Art. 28 DSGVO", "M365", "Unterzeichnet"),
                       ], [28, 22, 20, 20, 15])]
            make_xlsx_doc(FOLDER_COMP, fname, title, sheets)
        else:
            save_pdf(FOLDER_COMP, fname, title,
                     [("Zusammenfassung", f"Dokumentation zur {title.lower()} der {C['name']}.")])


# ══════════════════════════════════════════════════════════════════════
# FOLDER 14 – STRATEGIE / INVESTOREN
# ══════════════════════════════════════════════════════════════════════
FOLDER_STRAT = "14_Strategie_Investoren"


def gen_business_plan():
    doc = Document()
    make_docx_header(doc, "Businessplan 2024–2028",
                     f"{C['name']} – Strategische Planung", confidential=True)
    add_section(doc, "1. Executive Summary", [
        f"{C['name']} ist ein führender Hersteller von Medizinprodukten der Klassen IIa und IIb "
        "sowie IVD-Klasse-C-Diagnostika mit Sitz in München. Das Unternehmen erzielte 2023 einen "
        f"Umsatz von {eur(C['revenue_2023'])} bei einer EBITDA-Marge von "
        f"{C['ebitda_2023']/C['revenue_2023']*100:.1f} %.",
        "Strategische Ziele bis 2028: (1) Wachstum des Umsatzes auf EUR 200 Mio.; "
        "(2) Ausweitung der geographischen Präsenz auf Nordamerika und Südostasien; "
        "(3) Einführung von 2 neuen Produkten (NextGen Cardevio und AI-Diagnostik); "
        "(4) Möglicher IPO-Readiness-Prozess bis 2027.",
    ])
    add_section(doc, "2. Markt und Wettbewerb", [
        "Der globale Markt für diagnostische Medizinprodukte wächst mit ca. 8 % p.a. "
        "(Quelle: GlobalData 2023). Treiber: Alternde Bevölkerung, zunehmende chronische "
        "Erkrankungen, Digitalisierung im Gesundheitswesen, Präzisionsmedizin.",
        f"Wettbewerber: Philips Healthcare, GE HealthCare, Siemens Healthineers (Hauptsegment), "
        "Roper Technologies, Drägerwerk AG. Die {C['name']} differenziert sich durch "
        "überlegene Sensortechnologie, MDR-optimierte Plattform und Service-Exzellenz.",
    ])
    add_section(doc, "3. Wachstumsstrategie", [
        "Geografische Expansion: Aufbau einer US-Tochtergesellschaft (BTP MedTech Inc.) im Jahr 2025. "
        "FDA 510(k)-Zulassung für Cardevio Pro bis Q2 2025 angestrebt. "
        "Markteintritt Südostasien über Partnerdistributoren ab 2026.",
        "Produktinnovation: Investition von EUR 15 Mio. p.a. in F&E. Kernprojekte: "
        "NextGen Cardevio mit KI-Integration, erweitertes Veridiq SARS-Flex 2.0, "
        "neue Produktklasse: miniaturisierte Wearable-Diagnostik.",
        "M&A-Strategie: Akquisition eines Softwareunternehmens im Bereich medical AI "
        "(Budget: EUR 20–40 Mio.) bis 2025; Aufbau eines eigenen Digital Health-Ökosystems.",
    ])
    save_docx(doc, FOLDER_STRAT, "STRAT_001_Businessplan_2024_2028.docx")


def gen_board_presentation(year, quarter):
    sections = [
        (f"Board Presentation {quarter} {year} – Agenda", [
            ["TOP", "Thema", "Dauer", "Präsentator"],
            ["1", "Review vorige Präsentation und offene Punkte", "10 min", C["ceo"]],
            ["2", f"Geschäftsergebnis {quarter} {year}", "20 min", C["cfo"]],
            ["3", "Produktentwicklung und Regulatory Update", "15 min", C["cto"]],
            ["4", "Vertrieb und Marktentwicklung", "15 min", "Felix Schäfer, VP Sales"],
            ["5", "HR und Organisation", "10 min", C["ceo"]],
            ["6", "Strategische Initiativen", "20 min", C["ceo"]],
            ["7", "Finanzausblick und Guidance", "10 min", C["cfo"]],
        ]),
        (f"Finanzzusammenfassung {quarter} {year}",
         f"Umsatz: {eur(C['revenue_2023']//4 * (1+0.05*['Q1','Q2','Q3','Q4'].index(quarter)))}\n"
         f"EBITDA: {eur(C['ebitda_2023']//4)}\n"
         f"Mitarbeiter: {C['employees_2023']} FTE"),
        ("Strategische Highlights",
         f"• FDA-Vorgespräche für Cardevio Pro 510(k) begonnen\n"
         f"• Neuer Distributionsvertrag Skandinavien unterzeichnet\n"
         f"• PMCF-Studie {C['prod2_name']} erfolgreich abgeschlossen\n"
         f"• {random.randint(3,8)} neue Patentanmeldungen eingereicht"),
    ]
    save_pdf(FOLDER_STRAT, f"BOARD_{year}_{quarter}_Praesentation.pdf",
             f"Board Presentation {quarter} {year}", sections, confidential=True)


def gen_market_analysis():
    sections = [
        ("Marktanalyse – Europäischer Markt für Kardiologische Diagnostik",
         "Auftraggeber: Sentavia Precision GmbH\nErstellt durch: Christoph Braun, BD\n"
         f"Datum: {ds(2024, 1, 15)}"),
        ("Marktgröße und Wachstum", [
            ["Markt", "2022 EUR Mrd.", "2023 EUR Mrd.", "2024E EUR Mrd.", "CAGR 2022–2027E"],
            ["Kardiologische Diagnostik EU", "4,2", "4,6", "5,0", "+8,5%"],
            ["Orthopädische Navigationssysteme EU", "1,8", "2,0", "2,2", "+7,2%"],
            ["IVD Point-of-Care EU", "3,1", "3,4", "3,8", "+9,1%"],
        ]),
        ("Wettbewerbslandschaft",
         "Die Hauptwettbewerber in den Zielsegmenten sind:\n"
         "Kardiologie: Philips Healthcare (Marktanteil ~28 %), GE HealthCare (~22 %), "
         "Sentavia Precision (~4 %), sonstige (~46 %).\n\n"
         "Sentavia Precision positioniert sich als technologischer Innovationsführer im "
         "mittleren Preissegment mit besonderem Fokus auf MDR-Compliance und klinische Evidenz."),
    ]
    save_pdf(FOLDER_STRAT, "STRAT_002_Marktanalyse_EU_Medizinprodukte.pdf",
             "Marktanalyse Europäischer Medizinproduktemarkt 2024", sections)


def generate_14_strategie():
    print("  Generating 14_Strategie_Investoren...")
    gen_business_plan()
    gen_market_analysis()
    for yr in [2022, 2023]:
        for q in ["Q1", "Q2", "Q3", "Q4"]:
            gen_board_presentation(yr, q)
    # Investor relations documents
    for i, topic in enumerate([
        "M&A Screening Medical AI 2023",
        "Exit-Strategie Analyse IPO vs Trade Sale 2024",
        "Due Diligence Zielobjekt MedSoft Analytics GmbH",
        "Investorenpräsentation Series D Fundraising 2024",
        "Roadshow-Unterlagen Investoren Q1 2024",
    ], 1):
        sections = [
            ("Übersicht", f"Dokument: {topic}\nErstellt: {C['name']}, Geschäftsentwicklung\n"
             f"Datum: {ds(2023+i//5, random.randint(1,12), random.randint(1,28))}\nVertraulich: JA"),
            ("Inhalt", f"Dieses Dokument enthält vertrauliche strategische Analysen zu {topic.lower()}. "
             "Die Informationen dienen ausschließlich dem internen Gebrauch und sind an "
             "autorisierte Empfänger beschränkt."),
        ]
        save_pdf(FOLDER_STRAT, f"STRAT_{i+2:03d}_{sfn(topic[:25])}.pdf", topic, sections, confidential=True)
    # Financial modeling
    sheets = [
        ("Bewertungsmodell", ["Bewertungsmethode", "Enterprise Value EUR", "Equity Value EUR", "Multiplikator"],
         [
             ("DCF-Analyse (Base Case)", f"{int(C['revenue_2023']*9.5):,}", f"{int(C['revenue_2023']*9.5-50_000_000):,}", "EV/EBITDA: 9,5x"),
             ("DCF-Analyse (Bull Case)", f"{int(C['revenue_2023']*12):,}", f"{int(C['revenue_2023']*12-50_000_000):,}", "EV/EBITDA: 12,0x"),
             ("Vergleichbare Transaktionen (M&A Comps)", f"{int(C['revenue_2023']*11):,}", f"{int(C['revenue_2023']*11-50_000_000):,}", "EV/EBITDA: 11,0x"),
             ("Börsennotierte Vergleichsunternehmen", f"{int(C['revenue_2023']*10.2):,}", f"{int(C['revenue_2023']*10.2-50_000_000):,}", "EV/EBITDA: 10,2x"),
         ], [30, 22, 22, 22]),
    ]
    make_xlsx_doc(FOLDER_STRAT, "STRAT_010_Unternehmensbewertung_DCF_Comps.xlsx",
                  "Unternehmensbewertung – DCF und Vergleichsmethoden", sheets)


# ══════════════════════════════════════════════════════════════════════
# FOLDER 15 – F&E / INNOVATION
# ══════════════════════════════════════════════════════════════════════
FOLDER_RD = "15_FuE_Innovation"


def gen_rd_roadmap():
    sheets = [
        ("F&E Roadmap 2024–2027", ["Projekt", "Start", "Ende", "Budget EUR", "Phase", "Produkt"],
         [
             (f"NextGen {C['prod1_name']} – KI-Integration", "Q1 2024", "Q4 2025", "4.200.000", "Konzeptphase", C["prod1_name"]),
             ("AI Diagnostics Engine v2.0", "Q2 2024", "Q3 2026", "2.800.000", "Anforderungen", "Neu"),
             (f"{C['prod2_name']} – 5G-Konnektivität", "Q3 2024", "Q2 2025", "1.200.000", "Entwicklung", C["prod2_name"]),
             (f"{C['prod3_name']} 2.0 – Erweiterte Analytes", "Q1 2024", "Q4 2025", "1.800.000", "Validierung", C["prod3_name"]),
             ("Wearable Biosensor Plattform", "Q2 2024", "Q4 2027", "6.500.000", "Forschung", "Neu"),
             ("Digitale Gesundheitsplattform", "Q3 2024", "Q2 2026", "3.200.000", "Konzeptphase", "Neu"),
             ("Biomarker-Discovery Programm", "Q1 2024", "Laufend", "2.400.000/Jahr", "Grundlagenforschung", "Plattform"),
         ], [35, 10, 10, 16, 20, 22]),
    ]
    make_xlsx_doc(FOLDER_RD, "RD_001_FuE_Roadmap_2024_2027.xlsx",
                  "F&E Roadmap 2024–2027", sheets)


def gen_vv_report(prod_name, prod_short, test_type):
    type_titles = {
        "electrical": "Prüfbericht Elektrische Sicherheit (IEC 60601-1)",
        "emi": "Prüfbericht Elektromagnetische Verträglichkeit (IEC 60601-1-2)",
        "biocompat": "Biokompatibilitätsbewertung (ISO 10993-1)",
        "software": "Software-Verifikation und -Validierung (IEC 62304)",
        "performance": "Leistungsverifikation und Performance Testing",
    }
    type_results = {
        "electrical": [
            ["Test", "Anforderung", "Gemessen", "Ergebnis"],
            ["Schutzleiterwiderstand", "≤ 200 mΩ", f"{random.randint(20,80)} mΩ", "BESTANDEN"],
            ["Ableitstrom Gehäuse", "≤ 100 μA", f"{random.randint(5,60)} μA", "BESTANDEN"],
            ["Spannungsfestigkeit 1.500 V", "Kein Überschlag", "Kein Überschlag", "BESTANDEN"],
            ["Einschaltstrom", "≤ 40 A (peak)", f"{random.randint(8,25)} A", "BESTANDEN"],
        ],
        "emi": [
            ["Test", "Standard", "Frequenzbereich", "Ergebnis"],
            ["Emission (conducted)", "EN 55011 Class B", "0,15–30 MHz", "BESTANDEN"],
            ["Emission (radiated)", "EN 55011 Class B", "30–1000 MHz", "BESTANDEN"],
            ["Immunität ESD", "IEC 61000-4-2", "±4 kV Kontakt", "BESTANDEN"],
            ["Immunität RS", "IEC 61000-4-3", "10 V/m", "BESTANDEN"],
        ],
    }
    sections = [
        ("Prüfbericht",
         f"Produkt: {prod_name}\nTestlabor: TÜV SÜD Product Service GmbH, München\n"
         f"Prüfdatum: {ds(2022, random.randint(4,8), random.randint(1,28))}\n"
         f"Prüfbericht-Nr.: TR-{C['short']}-{prod_short}-{random.randint(1000,9999)}\n"
         f"Prüfgrundlage: Siehe Prüfplan VVP-{prod_short}-001"),
    ]
    if test_type in type_results:
        sections.append(("Testergebnisse", type_results[test_type]))
    else:
        sections.append(("Ergebnisübersicht",
                         f"Alle Prüfungen nach {type_titles.get(test_type, test_type)} wurden bestanden. "
                         "Keine Nichtkonformitäten festgestellt."))
    sections.append(("Fazit",
                      f"Das Produkt {prod_name} erfüllt alle geprüften Anforderungen. "
                      f"Der Prüfbericht liegt vollständig vor und ist in der technischen "
                      "Dokumentation (Technisches Dossier) abgelegt."))
    fname = f"VV_{prod_short}_{sfn(test_type[:10])}_Pruefbericht.pdf"
    save_pdf(FOLDER_RD, fname, type_titles.get(test_type, f"Prüfbericht – {test_type}"), sections)


def gen_biocompat_report(prod_name, prod_short):
    sections = [
        ("Biokompatibilitätsbewertung gemäß ISO 10993-1:2018",
         f"Produkt: {prod_name}\nMaterialien: Polymere (ABS, PC), Edelstahl 316L, "
         "biokompatible Beschichtungen\nErstellt: {C['cmo']}"),
        ("Bewertung nach ISO 10993-1", [
            ["Endpunkt", "Norm", "Methode", "Ergebnis"],
            ["Zytotoxizität", "ISO 10993-5", "Elution Test L929", "BESTANDEN"],
            ["Sensibilisierung", "ISO 10993-10", "GPMT / Buehler", "BESTANDEN"],
            ["Irritation/intrakutane Reaktion", "ISO 10993-10", "Kaninchentest", "BESTANDEN"],
            ["Systemische Toxizität", "ISO 10993-11", "Mäusetest (akut)", "BESTANDEN"],
            ["Genotoxizität", "ISO 10993-3", "Ames Test, MNT", "BESTANDEN"],
            ["Implantation (bei Bedarf)", "ISO 10993-6", "n/a (kein Implantat)", "n/a"],
        ]),
        ("Schlussfolgerung",
         f"Alle untersuchten Materialien und das Gesamtprodukt {prod_name} sind biokompatibel "
         "im Sinne der ISO 10993-1:2018. Die Bewertung bestätigt ein akzeptables biologisches "
         "Risikoprofil für die bestimmungsgemäße Anwendung."),
    ]
    save_pdf(FOLDER_RD, f"BIO_{prod_short}_Biokompatibilitaetsbewertung_ISO10993.pdf",
             f"Biokompatibilitätsbewertung – {prod_name}", sections)


def generate_15_fuer():
    print("  Generating 15_FuE_Innovation...")
    gen_rd_roadmap()
    for pname, pshort in [(C["prod1_name"], "CSP"), (C["prod2_name"], "OFN"), (C["prod3_name"], "DKS")]:
        for test_type in ["electrical", "emi", "biocompat", "software", "performance"]:
            gen_vv_report(pname, pshort, test_type)
        gen_biocompat_report(pname, pshort)
    # Design History Files
    for pname, pshort in [(C["prod1_name"], "CSP"), (C["prod2_name"], "OFN")]:
        for section_name in ["Design Input Spezifikation", "Design Output Dokumente",
                              "Verifikationsplan und -bericht", "Validierungsplan und -bericht",
                              "Design Transfer Protokoll"]:
            sections = [
                ("DHF-Abschnitt", f"Produkt: {pname}\nAbschnitt: {section_name}\n"
                 f"Status: Freigegeben | Freigabe: {C['qra']}"),
                ("Inhalt", f"Dieser DHF-Abschnitt dokumentiert {section_name.lower()} für {pname} "
                 "gemäß den Anforderungen der MDR Anhang II und dem internen Design Control SOP-015."),
            ]
            save_pdf(FOLDER_RD, f"DHF_{pshort}_{sfn(section_name[:20])}.pdf",
                     f"DHF – {pname} – {section_name}", sections)
    # Innovation projects
    for i, project in enumerate([
        "Forschungskooperation TU München – Kardiologie-KI",
        "EU-Forschungsprojekt Horizon Europe – MEDAI-2024",
        "BMBF-Förderprojekt Digitale Gesundheitstechnologien",
        "Kooperation Charité – Real-World-Evidence Studie",
    ], 1):
        sections = [
            ("Projektbeschreibung", f"Projekt: {project}\nProjektnummer: RD-{2023+i//5}-{i:04d}\n"
             f"Laufzeit: 24–36 Monate\nBudget: EUR {random.randint(500,2500)*1000:,}\n"
             f"Projektleitung: {C['cto']}"),
            ("Meilensteine", [
                ["Meilenstein", "Fälligkeit", "Status"],
                ["M1 – Konzeptphase abgeschlossen", ds(2024, 3, 31), "Abgeschlossen"],
                ["M2 – Prototyp vorliegend", ds(2024, 9, 30), "In Arbeit"],
                ["M3 – Vorklinische Tests", ds(2025, 3, 31), "Geplant"],
            ]),
        ]
        save_pdf(FOLDER_RD, f"RD_{i:03d}_{sfn(project[:25])}.pdf", project, sections)


# ══════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════
def create_folder_structure():
    folders = [
        FOLDER_GES, FOLDER_FIN, FOLDER_HR, FOLDER_SALES, FOLDER_PURCH,
        FOLDER_REG, FOLDER_QMS, FOLDER_CLIN, FOLDER_IP, FOLDER_INS,
        FOLDER_PROP, FOLDER_IT, FOLDER_COMP, FOLDER_STRAT, FOLDER_RD,
    ]
    for folder in folders:
        (BASE / folder).mkdir(parents=True, exist_ok=True)
    print(f"Created {len(folders)} folder(s) under {BASE}")


def print_summary():
    print("\n" + "=" * 60)
    print("GENERATION COMPLETE")
    print("=" * 60)
    total = len(GENERATED_FILES)
    print(f"Total documents generated: {total}")
    print()
    folders = {}
    for fpath in GENERATED_FILES:
        p = Path(fpath)
        folder = p.parent.name
        folders[folder] = folders.get(folder, 0) + 1
    for folder, count in sorted(folders.items()):
        print(f"  {folder}: {count} docs")
    print(f"\nOutput directory: {BASE}")


# ══════════════════════════════════════════════════════════════════════
# VOLUME FILL — extra documents to reach ~1000 total
# ══════════════════════════════════════════════════════════════════════

def gen_volume_fill():
    """Generate additional documents across all folders to reach ~1000 total."""
    print("  Generating volume-fill documents...")

    # ── 01 Gesellschaftsrecht – extra board resolutions and UBOs ──────
    more_ges = [
        (2020, 4, "Genehmigung Unternehmenssatzung Klausel 9 Abs. 4 Änderung"),
        (2020, 5, "Beschluss über Erteilung einer Untervollmacht für Bankverkehr"),
        (2021, 6, "Bestellung Dr. Annika Schmidt als Chief Medical Officer"),
        (2021, 7, "Genehmigung Mietvertragsverlängerung Freimannstraße 45 bis 2030"),
        (2022, 8, "Zustimmung zur Teilnahme an EU-Forschungsprojekt MEDAI-2024"),
        (2023, 9, "Genehmigung Einführung virtueller Gesellschafterversammlungen"),
        (2024, 10, "Genehmigung Due-Diligence-Prozess für strategische Transaktion"),
        (2024, 11, "Ausschüttungsbeschluss Vorabdividende EUR 2.000.000 an Gesellschafter"),
    ]
    for yr, n, topic in more_ges:
        gen_board_resolution(yr, n, topic)
    # Annual general meetings for additional years
    for yr in [2019, 2024]:
        sections = [
            ("Gesellschafterbeschluss Umlaufverfahren",
             f"Protokoll der {yr}er Gesellschafterversammlung der {C['name']}.\n"
             f"Datum: {ds(yr, 5, random.randint(15,28))}\nOrt: {C['full_address']}"),
            ("Tagesordnung und Beschlüsse",
             f"TOP 1: Jahresabschluss {yr-1} – einstimmig genehmigt.\n"
             f"TOP 2: Entlastung Geschäftsführung {yr-1} – erteilt.\n"
             f"TOP 3: Wirtschaftsplan {yr} – genehmigt."),
        ]
        save_pdf(FOLDER_GES, f"GV_{yr}_Protokoll_Gesellschafterversammlung.pdf",
                 f"Gesellschafterversammlung {yr}", sections)

    # ── 02 Finanzen – more monthly reports, treasury ──────────────────
    # Additional monthly BWA 2021
    gen_bwa_monatlich(2021)
    # VAT returns (Umsatzsteuervoranmeldungen)
    for month in range(1, 13):
        rev_m = int(C["revenue_2023"] / 12 * (1 + random.uniform(-0.1, 0.12)))
        vat_amount = int(rev_m * 0.19 * 0.12)  # only partial VATable
        sections = [
            ("Umsatzsteuervoranmeldung",
             f"Unternehmen: {C['name']}\nSteuernummer: {C['steuernr']}\n"
             f"Finanzamt: {C['fa']}\nAnmeldezeitraum: {MONTHS_DE[month]} 2023\n"
             f"Abgabefrist: {ds(2023, min(month+1,12), 10)}"),
            ("Steuerberechnung", [
                ["Position", "Bemessungsgrundlage EUR", "Steuerbetrag EUR"],
                ["Umsätze 19 % USt", f"{rev_m:,}", f"{int(rev_m*0.19):,}"],
                ["Umsätze 7 % USt", f"{int(rev_m*0.08):,}", f"{int(rev_m*0.08*0.07):,}"],
                ["EU-Umsätze (0 %)", f"{int(rev_m*0.22):,}", "0"],
                ["Vorsteuer", "—", f"-{int(rev_m*0.19*0.65):,}"],
                ["Zahllast", "—", f"{vat_amount:,}"],
            ]),
        ]
        save_pdf(FOLDER_FIN, f"UST_2023_{month:02d}_Voranmeldung_{MONTHS_DE[month]}.pdf",
                 f"Umsatzsteuervoranmeldung {MONTHS_DE[month]} 2023", sections)
    # Additional financial statements
    for title, fname in [
        ("Offenlegung Jahresabschluss 2023 (Bundesanzeiger)", "JA_2023_Bundesanzeiger_Offenlegung.pdf"),
        ("Leasingverbindlichkeiten Übersicht IFRS 16", "FIN_IFRS16_Leasingverbindlichkeiten.xlsx"),
        ("Fremdwährungsrisiko-Analyse 2023", "FIN_Fremdwaehrungsrisiko_2023.pdf"),
        ("Zinsänderungsrisiko Sensitivitätsanalyse", "FIN_Zinsaenderungsrisiko_Sensitivitaet.pdf"),
        ("Umsatzabgrenzung und Revenue Recognition", "FIN_Revenue_Recognition_Policy.pdf"),
        ("Überleitung HGB auf IFRS 2023", "FIN_HGB_IFRS_Ueberleitung_2023.xlsx"),
        ("Cash Pooling Agreement – Sofinnova Partners", "FIN_Cash_Pooling_Agreement.docx"),
        ("Eigenkapitalentwicklung 2019–2023", "FIN_Eigenkapitalentwicklung_5_Jahre.xlsx"),
    ]:
        if fname.endswith(".pdf"):
            fin_rows = [["Position", "Betrag EUR", "Anmerkung"]] + [
                [f"Posten {i}", f"{random.randint(100,5000)*1000:,}", "Vgl. Jahresabschluss"]
                for i in range(1, 7)
            ]
            save_pdf(FOLDER_FIN, fname, title,
                     [("Inhalt", f"Dieses Dokument enthält {title.lower()} der {C['name']}. "
                       f"Erstellt: {C['wp']}, {ds(2024, random.randint(1,4), random.randint(1,28))}."),
                      ("Daten", fin_rows)])
        elif fname.endswith(".xlsx"):
            sheets = [("Übersicht", ["Position", "2023 EUR", "2022 EUR"],
                       [(f"Posten {i}", f"{random.randint(100,5000)*1000:,}",
                         f"{random.randint(90,4500)*1000:,}") for i in range(1, 10)],
                       [30, 16, 16])]
            make_xlsx_doc(FOLDER_FIN, fname, title, sheets)
        else:
            d = Document()
            make_docx_header(d, title)
            d.add_paragraph(f"Dieses Dokument regelt {title.lower()} der {C['name']}.")
            save_docx(d, FOLDER_FIN, fname)

    # ── 03 HR – more docs ─────────────────────────────────────────────
    more_hr_docs = [
        ("Gehaltserhöhungsrunde 2023 – Entscheidungsmatrix", "HR_Gehalterhoehung_2023.xlsx"),
        ("Mitarbeiterbefragung 2023 – Auswertung", "HR_Mitarbeiterbefragung_2023.pdf"),
        ("Onboarding Checkliste für neue Mitarbeiter", "HR_Onboarding_Checkliste.docx"),
        ("Stellenbeschreibung Head of Quality", "HR_JD_Head_of_Quality.pdf"),
        ("Stellenbeschreibung Regulatory Affairs Specialist", "HR_JD_Regulatory_Affairs.pdf"),
        ("Stellenbeschreibung Software Engineer Medical", "HR_JD_Software_Engineer.pdf"),
        ("Stellenbeschreibung Key Account Manager", "HR_JD_Key_Account_Manager.pdf"),
        ("Exit-Interview-Protokoll Q3 2023", "HR_Exit_Interviews_Q3_2023.pdf"),
        ("Krankheitsstatistik 2022–2023", "HR_Krankheitsstatistik.xlsx"),
        ("Prämienprogramm ESOP 2023 – Teilnehmerliste", "HR_ESOP_2023_Teilnehmerliste.xlsx"),
        ("Weiterbildungsplanung 2024", "HR_Weiterbildungsplanung_2024.xlsx"),
        ("Remote-Work-Richtlinie", "HR_Remote_Work_Richtlinie.docx"),
        ("Diversity & Inclusion Bericht 2023", "HR_Diversity_Inclusion_2023.pdf"),
        ("Arbeitsschutzunterweisung Jahresprotokoll 2023", "HR_Arbeitsschutz_2023.pdf"),
        ("Stellenausschreibungen Archiv 2023", "HR_Stellenausschreibungen_2023.pdf"),
    ]
    for title, fname in more_hr_docs:
        if fname.endswith(".pdf"):
            save_pdf(FOLDER_HR, fname, title,
                     [("Zusammenfassung", f"Dieses Dokument dokumentiert {title.lower()}.\n"
                       f"Datum: {ds(2023, random.randint(1,12), random.randint(1,28))}\n"
                       f"Verantwortlich: {C['ceo']}")])
        elif fname.endswith(".xlsx"):
            sheets = [("Daten", ["Kategorie", "Anzahl", "Anmerkung"],
                       [(f"Eintrag {i}", str(random.randint(1, 100)), "Gemäß Aufzeichnung")
                        for i in range(1, 12)], [25, 12, 30])]
            make_xlsx_doc(FOLDER_HR, fname, title, sheets)
        else:
            d = Document()
            make_docx_header(d, title)
            d.add_paragraph(f"Inhalt: {title}")
            save_docx(d, FOLDER_HR, fname)

    # ── 04 Vertrieb – extra pricing, tender docs, agreements ──────────
    for yr in [2021, 2022]:
        for q in ["Q1", "Q2", "Q3", "Q4"]:
            rev = C.get(f"revenue_{yr}", 71_400_000) // 4
            rows = [
                (C["prod1_name"], str(random.randint(150, 230)), f"{int(rev*0.52):,}",
                 f"{random.randint(65,90)}%"),
                (C["prod2_name"], str(random.randint(240, 380)), f"{int(rev*0.33):,}",
                 f"{random.randint(60,85)}%"),
                (C["prod3_name"], str(random.randint(3000, 7000)), f"{int(rev*0.15):,}",
                 f"{random.randint(50,70)}%"),
            ]
            sheets = [("Umsatz", ["Produkt", "Einheiten", "Umsatz EUR", "Kundenzufr."],
                       rows, [22, 12, 15, 16])]
            make_xlsx_doc(FOLDER_SALES, f"SALES_{q}_{yr}_Umsatzbericht.xlsx",
                          f"Vertriebsbericht {q} {yr}", sheets)
    # More distributor docs
    for dist, short in [(C["dist1"], "SH"), (C["dist2"], "FMC"), (C["dist3"], "BB")]:
        for yr in [2022, 2023]:
            sections = [
                ("Jahresabrechnung Distributor",
                 f"Distributor: {dist}\nJahr: {yr}\n"
                 f"Jahresumsatz: {eur(random.randint(3,9)*1_000_000)}\n"
                 f"Zielerreichung: {random.randint(88,115)} %"),
                ("Detailauswertung", [
                    ["Produkt", "Geplant EUR", "Ist EUR", "Abweichung"],
                    [C["prod1_name"], f"{random.randint(1500,3000)*1000:,}",
                     f"{random.randint(1400,3200)*1000:,}", f"+{random.randint(-5,15)} %"],
                    [C["prod2_name"], f"{random.randint(800,2000)*1000:,}",
                     f"{random.randint(750,2200)*1000:,}", f"+{random.randint(-8,12)} %"],
                ]),
            ]
            save_pdf(FOLDER_SALES, f"DIST_{short}_{yr}_Jahresabrechnung.pdf",
                     f"Jahresabrechnung Distributor {dist} {yr}", sections)
    # Hospital framework addenda
    for hosp_short, hosp in [("CHR", C["hosp1"]), ("UKH", C["hosp4"])]:
        sections = [
            ("Nachtrag Rahmenliefervertrag",
             f"Nachtrag Nr. 1 zum Rahmenliefervertrag vom 1. Juni 2022 zwischen {C['name']} "
             f"und {hosp}.\n\nGegenstand: Aufnahme von {C['prod3_name']} in den Vertragsumfang."),
        ]
        save_pdf(FOLDER_SALES, f"HOSP_{hosp_short}_Nachtrag_2023.pdf",
                 f"Nachtrag Rahmenliefervertrag – {hosp}", sections)
    # Customer satisfaction surveys
    sheets = [
        ("Kundenzufriedenheit 2023", ["Kunde", "Produkt", "Zufriedenheit (1–10)", "NPS", "Kommentar"],
         [(random.choice([C["hosp1"], C["hosp2"], C["hosp3"], C["hosp4"]]),
           random.choice([C["prod1_name"], C["prod2_name"]]),
           str(random.randint(7, 10)),
           str(random.randint(30, 80)),
           random.choice(["Sehr gutes Gerät", "Einfache Bedienung", "Support verbessern",
                          "Sehr zufrieden", "Kalibrierung aufwändig"]))
          for _ in range(35)], [32, 22, 20, 8, 35]),
    ]
    make_xlsx_doc(FOLDER_SALES, "SALES_Kundenzufriedenheit_2023.xlsx",
                  "Kundenzufriedenheitsumfrage 2023", sheets)
    # Pipeline CRM export
    sheets = [
        ("Sales Pipeline Q4 2023", ["Opportunity", "Kunde", "Produkt", "Wert EUR",
                                     "Phase", "Wahrscheinlichkeit %", "Expected Close"],
         [(f"OPP-{i:04d}", random.choice([C["hosp1"], C["hosp2"], C["hosp3"], C["hosp4"],
                                           "Klinikum Stuttgart", "Städtisches Klinikum München"]),
           random.choice([C["prod1_name"], C["prod2_name"], C["prod3_name"]]),
           f"{random.randint(50, 500)*1000:,}",
           random.choice(["Prospect", "Angebot", "Verhandlung", "Won", "Lost"]),
           str(random.choice([20, 40, 60, 80, 90, 100])),
           ds(2024, random.randint(1, 6), random.randint(1, 28)))
          for i in range(1, 30)], [12, 32, 22, 14, 14, 18, 16]),
    ]
    make_xlsx_doc(FOLDER_SALES, "SALES_Pipeline_Q4_2023.xlsx",
                  "Sales Pipeline Q4 2023", sheets)

    # ── 05 Einkauf – more supplier docs ──────────────────────────────
    # Supplier evaluation annual reports
    for lief_name, short in [(C["lief1"], "TI"), (C["lief2"], "SE"),
                              (C["lief3"], "MT"), (C["lief4"], "SM")]:
        for yr in [2022, 2023]:
            sections = [
                ("Lieferantenbewertung",
                 f"Lieferant: {lief_name}\nBewertungsjahr: {yr}\n"
                 f"Bewerter: Patrick Richter, Supply Chain Management"),
                ("Bewertungskennzahlen", [
                    ["Kriterium", "Gewichtung", "Bewertung (1–10)", "Gewichtetes Ergebnis"],
                    ["Liefertreue", "30%", str(random.randint(7, 10)), f"{random.uniform(2.1, 3.0):.1f}"],
                    ["Qualität (ppm)", "30%", str(random.randint(7, 10)), f"{random.uniform(2.1, 3.0):.1f}"],
                    ["Preis-Leistung", "20%", str(random.randint(6, 9)), f"{random.uniform(1.2, 1.8):.1f}"],
                    ["Flexibilität", "10%", str(random.randint(6, 10)), f"{random.uniform(0.6, 1.0):.1f}"],
                    ["Kommunikation", "10%", str(random.randint(7, 10)), f"{random.uniform(0.7, 1.0):.1f}"],
                    ["GESAMT", "100%", "—", f"{random.uniform(7.0, 9.5):.1f}"],
                ]),
            ]
            save_pdf(FOLDER_PURCH, f"LIEF_BWRT_{short}_{yr}_Jahresbewertung.pdf",
                     f"Lieferantenjahresbewertung – {lief_name} {yr}", sections)
    # Critical supplier scorecards
    sheets = [
        ("Lieferanten-Scorecard Q4 2023",
         ["Lieferant", "Liefertreue %", "Qualität ppm", "Rechnungskonformität %",
          "Audit-Status", "Freigabestatus"],
         [(C["lief1"], f"{random.randint(94,99)} %", str(random.randint(80,400)),
           f"{random.randint(98,100)} %", "Bestanden 2023", "Freigegeben"),
          (C["lief2"], f"{random.randint(93,98)} %", str(random.randint(100,500)),
           f"{random.randint(97,100)} %", "Bestanden 2023", "Freigegeben"),
          (C["lief3"], f"{random.randint(92,97)} %", str(random.randint(150,600)),
           f"{random.randint(96,100)} %", "Ausstehend Q1 2024", "Freigegeben (unter Auflagen)"),
          (C["lief4"], f"{random.randint(96,100)} %", str(random.randint(50,200)),
           f"{random.randint(99,100)} %", "Bestanden 2023", "Freigegeben")],
         [28, 14, 14, 18, 22, 22]),
    ]
    make_xlsx_doc(FOLDER_PURCH, "LIEF_Scorecard_Q4_2023.xlsx",
                  "Lieferanten-Scorecard Q4 2023", sheets)
    # SABS – Single Approved Buyer Specs
    for lief_name, short, component in [
        (C["lief1"], "TI", "Mikroprozessoren TMS320C6748"),
        (C["lief2"], "SE", "Feuchte-/Temperatursensoren SHT40"),
    ]:
        sections = [
            ("Komponentenspezifikation",
             f"Komponente: {component}\nLieferant: {lief_name}\n"
             f"Interne Teilenummer: BTP-COMP-{short}-{random.randint(1000,9999)}\n"
             f"Klassifizierung: Kritische Komponente (MDR-relevant)"),
            ("Technische Anforderungen", [
                ["Parameter", "Einheit", "Min", "Nom", "Max", "Prüfmethode"],
                ["Betriebstemperatur", "°C", "-20", "25", "70", "IEC 60068-2-2"],
                ["Versorgungsspannung", "V", "3,1", "3,3", "3,5", "IEC 61010-1"],
                ["Stromaufnahme (aktiv)", "mA", "—", "12", "18", "Intern"],
                ["Ausgangsgenauigkeit", "%", "—", "±0,5", "±2,0", "ISO 376"],
            ]),
        ]
        save_pdf(FOLDER_PURCH, f"SPEC_{short}_{sfn(component[:20])}_Komponentenspez.pdf",
                 f"Komponentenspezifikation – {component}", sections)
    # Additional purchase orders (bulk)
    for i in range(16, 36):
        lief = random.choice([(C["lief1"], "TI"), (C["lief2"], "SE"),
                               (C["lief3"], "MT"), (C["lief4"], "SM")])
        d = rdate(2023)
        sections = [
            ("Bestelldaten",
             f"Bestellnummer: PO-2023-{i:04d}\nLieferant: {lief[0]}\n"
             f"Datum: {ds(d.year, d.month, d.day)}\n"
             f"Betrag: {eur(random.randint(10,250)*1000)}"),
        ]
        save_pdf(FOLDER_PURCH, f"PO_2023_{i:04d}_Bestellung_{lief[1]}.pdf",
                 f"Bestellung PO-2023-{i:04d}", sections)

    # ── 06 Regulatorisch – PSUR + extra MDR artifacts ────────────────
    for pname, pshort in [(C["prod1_name"], "CSP"), (C["prod2_name"], "OFN"), (C["prod3_name"], "DKS")]:
        # PSUR
        sections = [
            ("Periodic Safety Update Report (PSUR) gemäß Art. 86 MDR",
             f"Produkt: {pname}\nBerichtszeitraum: 2022–2023\n"
             f"Version: 2.0 | Datum: {ds(2024, 1, 31)}\nErstellt durch: {C['qra']}"),
            ("Übersicht PMS-Aktivitäten",
             f"Im Berichtszeitraum 2022–2023 wurden für {pname} folgende PMS-Aktivitäten durchgeführt: "
             f"2 PMS-Berichte (PMSR), 1 PMCF Evaluation Report, laufende Vigilanzüberwachung, "
             f"systematische Literaturrecherche (halbjährlich), EUDAMED-Datenpflege."),
            ("Vigilanzdaten", [
                ["Jahr", "Gemeldete Ereignisse", "Davon SAE", "Vigilanzberichte BfArM"],
                ["2022", str(random.randint(1,4)), str(random.randint(0,2)), str(random.randint(0,1))],
                ["2023", str(random.randint(1,5)), str(random.randint(0,2)), str(random.randint(0,1))],
            ]),
            ("Schlussfolgerung",
             f"Das Nutzen-Risiko-Verhältnis für {pname} bleibt über den Berichtszeitraum positiv. "
             "Keine systemischen Sicherheitsprobleme identifiziert. Alle Risikominderungsmaßnahmen "
             "sind wirksam. CER wurde auf Basis der PSUR-Erkenntnisse aktualisiert."),
        ]
        save_pdf(FOLDER_REG, f"PSUR_{pshort}_2022_2023_Bericht.pdf",
                 f"Periodic Safety Update Report (PSUR) – {pname} 2022–2023", sections)
        # MDR Art. 10(13) – Serious Incident Report
        for si_num in range(1, 4):
            d = rdate(2023)
            sections = [
                ("Meldung nach Art. 87 MDR – Schwerwiegendes Vorkommnis",
                 f"SI-Referenz: SI-{C['short']}-{pshort}-2023-{si_num:03d}\n"
                 f"Datum: {ds(d.year, d.month, d.day)}\n"
                 f"Produkt: {pname} | Lot-Nr.: LOT-2023-{random.randint(1000,9999)}"),
                ("Sachverhalt",
                 f"Im klinischen Betrieb bei {random.choice([C['hosp1'], C['hosp2'], C['hosp3']])} "
                 f"wurde ein Ereignis mit dem Produkt {pname} beobachtet. "
                 "Nach eingehender Analyse wurde entschieden, eine Vigilanzmeldung zu erstatten."),
                ("Status",
                 f"Status: {random.choice(['Offen', 'In Untersuchung', 'Abgeschlossen'])}\n"
                 f"Behördliche Rückmeldung BfArM: {random.choice(['Ausstehend', 'Erhalten', 'Keine Rückfragen'])}"),
            ]
            save_pdf(FOLDER_REG,
                     f"SI_{pshort}_{si_num:03d}_Schwerwiegendes_Vorkommnis_2023.pdf",
                     f"Schwerwiegendes Vorkommnis – {pname}", sections)
        # Trend analysis
        sheets = [("Trend-Analyse PMS", ["Monat", "Complaints", "Near Miss", "SAE", "Trend"],
                   [(MONTHS_DE[m], str(random.randint(0,3)), str(random.randint(0,2)),
                     str(random.randint(0,1)), random.choice(["Stabil", "Stabil", "↓", "↑"]))
                    for m in range(1, 13)], [12, 12, 12, 8, 12])]
        make_xlsx_doc(FOLDER_REG, f"TREND_{pshort}_PMS_Trendanalyse_2023.xlsx",
                      f"PMS Trendanalyse 2023 – {pname}", sheets)
        # Design changes
        sections = [
            ("Änderungsmitteilung (Field Change)",
             f"Änderungs-ID: FC-{pshort}-2023-{random.randint(1,5):02d}\n"
             f"Produkt: {pname}\nArt der Änderung: Software-Update\n"
             f"Regulatorisch: Keine erneute Konformitätsbewertung erforderlich"),
        ]
        save_pdf(FOLDER_REG, f"FC_{pshort}_Aenderungsmitteilung_2023.pdf",
                 f"Änderungsmitteilung – {pname}", sections)
    # MDR Transition documentation
    for topic in ["MDR_Art_120_Uebergangsbescheinigung_2021",
                  "MDR_Anhang_IX_Bewertungsplan_CSP_2022",
                  "MDR_Anhang_IX_Bewertungsplan_OFN_2021",
                  "IVDR_Anhang_IX_Bewertungsplan_DKS_2022",
                  "MDR_Self_Assessment_Checklist_2021",
                  "NB_Annual_Surveillance_Audit_2023",
                  "MDR_Labelling_Review_All_Products_2023",
                  "MDR_IFU_Review_EN_Translation_CSP",
                  "MDR_IFU_Review_FR_Translation_CSP",
                  "EUDAMED_Jahresupdate_Produktdaten_2023"]:
        sections = [("Dokumentübersicht",
                     f"Thema: {topic.replace('_', ' ')}\nErstellt: {C['name']}, {C['qra']}\n"
                     f"Datum: {ds(2023, random.randint(1,12), random.randint(1,28))}")]
        save_pdf(FOLDER_REG, f"{topic[:50]}.pdf",
                 topic.replace("_", " "), sections)

    # ── 07 QMS – more SOPs and QMS docs ──────────────────────────────
    # More CAPAs
    for i in range(16, 31):
        gen_capa_record(i)
    # More NC reports
    for i in range(11, 26):
        gen_nc_report(i)
    # External audit reports
    for yr, cert_type in [(2021, "Überwachungsaudit"), (2022, "Überwachungsaudit"),
                          (2023, "Re-Zertifizierungsaudit")]:
        sections = [
            (f"Externen Auditbericht {cert_type} ISO 13485:2016",
             f"Auditiertes Unternehmen: {C['name']}\nAuditdatum: {ds(yr, 9, random.randint(10,25))}\n"
             f"Auditor: {C['nb']}, Auditor-ID: {random.randint(1000,9999)}\n"
             f"Zertifikat-Nr.: {C['iso_cert_nr']}\nErgebnis: BESTANDEN / Keine Major Findings"),
            ("Auditfeststellungen", [
                ["Nr.", "Feststellung", "Klassifizierung", "Status"],
                ["1", "Kalibrierungsplanung komplett implementiert", "Positiv", "—"],
                ["2", "CAPA-Prozess sehr gut dokumentiert", "Positiv", "—"],
                ["3", "Kleinere Abweichung SOP-013 Version", "Minor", "Korrigiert"],
            ]),
            ("Beschluss",
             f"Das Qualitätsmanagementsystem der {C['name']} entspricht den Anforderungen der "
             f"{C['iso_cert']}. Das Zertifikat {C['iso_cert_nr']} wird {'erneuert' if yr == 2023 else 'aufrechterhalten'}."),
        ]
        save_pdf(FOLDER_QMS, f"EXTAUDIT_{yr}_{cert_type[:10].replace(' ','_')}_ISO13485.pdf",
                 f"{cert_type} ISO 13485 {yr}", sections)
    # Product release records
    for pname, pshort in [(C["prod1_name"], "CSP"), (C["prod2_name"], "OFN"), (C["prod3_name"], "DKS")]:
        for yr in [2022, 2023]:
            sheets = [
                ("Chargenfreigaben", ["Charge-Nr.", "Produkt", "Menge", "Freigabedatum", "Freigeber", "Status"],
                 [(f"LOT-{yr}-{str(i).zfill(4)}", pname, str(random.randint(20,200)),
                   ds(yr, random.randint(1,12), random.randint(1,28)),
                   C["qra"], "FREIGEGEBEN")
                  for i in range(1, 16)], [16, 20, 10, 16, 22, 14]),
            ]
            make_xlsx_doc(FOLDER_QMS, f"QMS_Chargenfreigaben_{pshort}_{yr}.xlsx",
                          f"Chargenfreigabedokumentation {pname} {yr}", sheets)
    # Product complaint investigation reports
    for i in range(1, 9):
        prod = random.choice([C["prod1_name"], C["prod2_name"], C["prod3_name"]])
        sections = [
            ("Beschwerde-Untersuchungsbericht",
             f"Beschwerde-Nr.: CC-2023-{i:04d}\nProdukt: {prod}\n"
             f"Untersucht durch: Sarah Krause, QA\n"
             f"Abschluss: {ds(2023, random.randint(1,12), random.randint(1,28))}"),
            ("Ursache und Maßnahme",
             f"Ursache: {random.choice(['Anwenderfehler', 'Herstellungsabweichung', 'Transportschaden', 'Softwarefehler'])}\n"
             f"Maßnahme: {random.choice(['IFU aktualisiert', 'Schulung veranlasst', 'CAPA eröffnet', 'Reklamation akzeptiert'])}"),
        ]
        save_pdf(FOLDER_QMS, f"QMS_CC_Investigation_{i:04d}_2023.pdf",
                 f"Beschwerde-Untersuchung CC-2023-{i:04d}", sections)

    # ── 08 Klinisch – more study documents ───────────────────────────
    # More adverse event reports
    for i in range(1, 16):
        prod = random.choice([C["prod1_name"], C["prod2_name"], C["prod3_name"]])
        sections = [
            ("Adverse Event Report",
             f"AE-Referenz: AE-2023-{i:04d}\nProdukt: {prod}\n"
             f"Datum: {ds(2023, random.randint(1,12), random.randint(1,28))}\n"
             f"Schweregrad: {random.choice(['Mild', 'Moderat', 'Schwer', 'Nicht bewertbar'])}\n"
             f"Kausalität: {random.choice(['Möglicherweise', 'Unwahrscheinlich', 'Keine Beziehung'])}"),
            ("Beschreibung",
             f"Im klinischen Betrieb bei {random.choice([C['hosp1'], C['hosp2']])} "
             "wurde folgendes unerwünschtes Ereignis beobachtet: "
             f"{random.choice(['Verzögerung in der Diagnostik', 'Gerätealarm ohne Ursache', 'Messwert außerhalb Norm'])}."),
        ]
        save_pdf(FOLDER_CLIN, f"AE_2023_{i:04d}_Adverse_Event_Report.pdf",
                 f"Adverse Event Report AE-2023-{i:04d}", sections)
    # PMCF sub-studies
    for pname, pshort in [(C["prod1_name"], "CSP"), (C["prod2_name"], "OFN")]:
        sections = [
            ("PMCF Teilstudie – Geriatrie",
             f"Produkt: {pname}\nPopulation: Patienten ≥ 75 Jahre\nn = {random.randint(40,80)}\n"
             f"Prüfzentrum: {C['hosp1']}"),
            ("Ergebnisse",
             f"Diagnostische Genauigkeit in geriatrischer Kohorte: Sensitivität {random.randint(90,96)} %, "
             f"Spezifizität {random.randint(86,94)} %. Vergleichbar mit allgemeiner Population."),
        ]
        save_pdf(FOLDER_CLIN, f"PMCF_{pshort}_Substudie_Geriatrie.pdf",
                 f"PMCF Teilstudie Geriatrie – {pname}", sections)
    # Clinical evaluation summary tables
    for pname, pshort in [(C["prod1_name"], "CSP"), (C["prod2_name"], "OFN"), (C["prod3_name"], "DKS")]:
        sheets = [
            ("Evidenztabelle", ["Studie-ID", "Autor", "Jahr", "n", "Design", "Sensitivität %",
                                 "Spezifizität %", "Inclusion"],
             [(f"S{i:03d}", f"Autor{i} et al.", str(random.randint(2015,2023)),
               str(random.randint(50,800)),
               random.choice(["RCT", "Kohortenstudie", "Fallserie", "Register"]),
               f"{random.randint(88,98)}", f"{random.randint(84,96)}",
               random.choice(["JA", "JA", "JA", "NEIN"]))
              for i in range(1, 21)], [10, 18, 8, 8, 18, 14, 14, 10]),
        ]
        make_xlsx_doc(FOLDER_CLIN, f"CER_{pshort}_Evidenztabelle.xlsx",
                      f"CER Evidenztabelle – {pname}", sheets)

    # ── 09 IP – more patent docs ───────────────────────────────────────
    more_patents = [
        ("Verfahren zur Rauschunterdrückung in biosensorischen Signalen", 2020, "Angemeldet DE"),
        ("Verfahren zur patientenindividuellen Algorithmusanpassung", 2021, "Angemeldet PCT"),
        ("Vorrichtung zur Echtzeitkalibrierung medizinischer Messgeräte", 2022, "Angemeldet EP"),
        ("Miniaturisierter Biosensor für Wearable-Anwendungen", 2023, "Provisional US"),
    ]
    for i, (title, year, status) in enumerate(more_patents, len(more_patents)+1):
        gen_patent_application(i + 10, title, year, status)
    # Patent renewal records
    for i in range(1, 6):
        sections = [
            ("Jahresgebühr-Zahlung Patent",
             f"Patent-ID: BTP-P-{i:03d}\nAnnuität Jahr {10+i}\n"
             f"Betrag: EUR {random.randint(1500,8000)}\n"
             f"Gezahlt: {ds(2023, random.randint(1,12), random.randint(1,28))}\n"
             f"Kanzlei: {C['patent_anwalt']}"),
        ]
        save_pdf(FOLDER_IP, f"PAT_ANNUAL_{i:03d}_Jahresgebuehr.pdf",
                 f"Jahresgebühr Patent BTP-P-{i:03d}", sections)
    # License tracking
    sheets = [
        ("Lizenzübersicht", ["Lizenz-ID", "Lizenzgeber", "Lizenzthema", "Royalty %",
                              "Jahresbetrag EUR", "Status"],
         [
             ("LIC-001", "Fraunhofer IPA", "Kardiologisches Messverfahren EP2344567", "2,5%", "180.000", "Aktiv"),
             ("LIC-002", "TU München", "Algorithmus maschinelles Lernen", "1,5%", "45.000", "Aktiv"),
             ("LIC-003", "Helmholtz", "Biosensor-Grundpatent", "1,0%", "28.000", "Aktiv"),
         ], [10, 22, 40, 10, 16, 12]),
    ]
    make_xlsx_doc(FOLDER_IP, "IP_004_Lizenzuebersicht.xlsx",
                  "Lizenzportfolio – Übersicht", sheets)

    # ── 10 Versicherungen – policy renewals ───────────────────────────
    for yr in [2022, 2023]:
        sections = [
            ("Versicherungsportfolio – Jahresübersicht",
             f"Jahr: {yr}\nVersicherungsnehmer: {C['name']}\nGesamtprämienaufwand: "
             f"{eur(random.randint(480,540)*1000)}"),
            ("Änderungen gegenüber Vorjahr",
             f"• Erhöhung der Produkthaftpflicht-Deckungssumme von 20 Mio. auf 25 Mio. EUR\n"
             "• Neue Cyber-Police abgeschlossen\n"
             "• Anpassung Berufshaftpflicht an aktuellen Umsatz"),
        ]
        save_pdf(FOLDER_INS, f"VER_Jahresuebersicht_{yr}.pdf",
                 f"Versicherungsportfolio Jahresübersicht {yr}", sections)
    # Claims history
    sheets = [
        ("Schadenhistorie", ["Jahr", "Schadensfall", "Schadensart", "Betrag EUR", "Status"],
         [
             ("2021", "PL-2021-001", "Produkthaftung – Kundenbeschädigung", "45.000", "Abgewickelt"),
             ("2022", "DO-2022-001", "D&O – Managemententscheidung angefochten", "0 (abgewiesen)", "Abgewickelt"),
             ("2023", "CY-2023-001", "Cyber – Phishing-Angriff, kein Datenverlust", "8.500", "Abgewickelt"),
         ], [8, 16, 35, 18, 16]),
    ]
    make_xlsx_doc(FOLDER_INS, "VER_Schadenhistorie_2020_2023.xlsx",
                  "Schadenhistorie Versicherungen 2020–2023", sheets)

    # ── 11 Immobilien – more docs ─────────────────────────────────────
    for title, fname in [
        ("Nebenkostenabrechnung 2022", "PROP_004_Nebenkostenabrechnung_2022.pdf"),
        ("Nebenkostenabrechnung 2023", "PROP_005_Nebenkostenabrechnung_2023.pdf"),
        ("Grundrissplan Freimannstraße 45 EG", "PROP_006_Grundriss_EG.pdf"),
        ("Grundrissplan Freimannstraße 45 OG", "PROP_007_Grundriss_OG.pdf"),
        ("Schlussbericht Reinraumabnahme IQ OQ PQ", "PROP_008_Reinraum_IQOQPQ_Abnahme.pdf"),
        ("Vermieter Schriftverkehr 2023 – Mängelrüge Heizung", "PROP_009_Schriftverkehr_Heizung.pdf"),
        ("Umbaugenehmigung Stadt München 2022", "PROP_010_Baugenehmigung_2022.pdf"),
        ("Energieausweis Freimannstraße 45", "PROP_011_Energieausweis.pdf"),
    ]:
        sections = [("Dokument", f"{title}\nObjekt: {C['full_address']}\n"
                     f"Datum: {ds(2023, random.randint(1,12), random.randint(1,28))}")]
        save_pdf(FOLDER_PROP, fname, title, sections)

    # ── 12 IT – more docs ─────────────────────────────────────────────
    more_it = [
        ("ERP-Validierungsbericht SAP S4HANA v2023", "IT_009_SAP_Validierungsbericht.pdf"),
        ("QMS-Validierungsbericht MasterControl", "IT_010_MasterControl_Validierungsbericht.pdf"),
        ("IT-Notfallkonzept und BCP Test 2023", "IT_011_BCP_Test_2023.pdf"),
        ("Software-SBOM Cardevio Pro v3.2", "IT_012_SBOM_Cardevio.xlsx"),
        ("Vulnerability Scan Report Q3 2023", "IT_013_Vulnerability_Scan_Q3_2023.pdf"),
        ("IT Change Management Log 2023", "IT_014_Change_Management_Log.xlsx"),
        ("Backup und Recovery Testprotokoll 2023", "IT_015_Backup_Recovery_Test.pdf"),
        ("Access Management Review 2023", "IT_016_Access_Management_Review.pdf"),
        ("Cloud Security Assessment SAP", "IT_017_Cloud_Security_Assessment.pdf"),
        ("REDCap Validation Report", "IT_018_REDCap_Validation_Report.pdf"),
    ]
    for title, fname in more_it:
        if fname.endswith(".xlsx"):
            sheets = [("Daten", ["ID", "Beschreibung", "Status", "Datum"],
                       [(f"ID-{i:04d}", f"Eintrag {i}", random.choice(["OK", "Offen", "Erledigt"]),
                         ds(2023, random.randint(1,12), random.randint(1,28)))
                        for i in range(1, 20)], [12, 40, 12, 16])]
            make_xlsx_doc(FOLDER_IT, fname, title, sheets)
        else:
            save_pdf(FOLDER_IT, fname, title,
                     [("Zusammenfassung", f"Dokumentiert {title.lower()}.\n"
                       f"Datum: {ds(2023, random.randint(1,12), random.randint(1,28))}\n"
                       f"Verantwortlich: Benjamin Wolf, IT-Leitung"),
                      ("Ergebnis", "Keine kritischen Findings. Empfehlungen wurden dokumentiert.")])

    # ── 13 Compliance – more docs ─────────────────────────────────────
    more_comp = [
        ("Compliance Monitoring Bericht Q1 2023", "COMP_MON_Q1_2023.pdf"),
        ("Compliance Monitoring Bericht Q2 2023", "COMP_MON_Q2_2023.pdf"),
        ("Compliance Monitoring Bericht Q3 2023", "COMP_MON_Q3_2023.pdf"),
        ("Compliance Monitoring Bericht Q4 2023", "COMP_MON_Q4_2023.pdf"),
        ("Anti-Kickback Schulungsunterlagen 2023", "COMP_AK_Schulung_2023.pdf"),
        ("HWG-Prüfung Produktwerbung Cardevio", "COMP_HWG_Cardevio_Review.pdf"),
        ("Datenschutz Schulungsnachweis 2023", "COMP_DS_Schulungsnachweise_2023.xlsx"),
        ("Hinweisgebersystem – Fallprotokoll 2023", "COMP_HinSchG_Fallprotokoll_2023.pdf"),
        ("Sanktionslistenprüfung Lieferanten Q4 2023", "COMP_Sanktionsliste_Q4_2023.xlsx"),
        ("Legal Opinion EU AI Act Medizinprodukte", "COMP_Legal_Opinion_AI_Act.pdf"),
        ("Legal Opinion MDR Art 120 Übergangsfrist", "COMP_Legal_Opinion_MDR_Art120.pdf"),
        ("Transparenzregister Eintragung 2023", "COMP_Transparenzregister_2023.pdf"),
        ("Antikorruptions-Training Teilnehmerliste 2023", "COMP_AK_Training_Teilnehmer.xlsx"),
    ]
    for title, fname in more_comp:
        if fname.endswith(".xlsx"):
            sheets = [("Daten", ["Name", "Datum", "Ergebnis"],
                       [(f"Person {i}", ds(2023, random.randint(1,12), random.randint(1,28)), "Bestanden")
                        for i in range(1, 15)], [25, 16, 12])]
            make_xlsx_doc(FOLDER_COMP, fname, title, sheets)
        else:
            save_pdf(FOLDER_COMP, fname, title,
                     [("Zusammenfassung",
                       f"Compliance-Dokument: {title}\nDatum: {ds(2023, random.randint(1,12), random.randint(1,28))}\n"
                       f"Verantwortlich: Laura Neumann, Legal")])

    # ── 14 Strategie – more docs ──────────────────────────────────────
    for yr in [2021, 2022]:
        for q in ["Q1", "Q2", "Q3", "Q4"]:
            gen_board_presentation(yr, q)
    # M&A and strategy analyses
    for i, title in enumerate([
        "Strategiepapier Expansion USA 2025",
        "Markteintritt Analyse ASEAN 2026",
        "IPO-Readiness Assessment 2024",
        "Synergieanalyse Akquisition MedSoft Analytics",
        "Wettbewerberanalyse Philips Healthcare und GE HealthCare",
        "Partnerschaftsmodelle CRO-Kooperationen klinische Studien",
        "Digital Health Strategie 2025–2028",
        "Nachhaltigkeitsstrategie (ESG) BTP 2024",
        "Kostenstrukturanalyse und Effizienzprogramm 2024",
        "Quarterly Review Deck CFO Q3 2023 (Vorlage Beirat)",
    ], 1):
        sections = [
            ("Dokument", f"Titel: {title}\nKlassifikation: Streng vertraulich\n"
             f"Erstellt: {random.choice([C['ceo'], C['cfo'], 'Christoph Braun'])}\n"
             f"Datum: {ds(2023, random.randint(1,12), random.randint(1,28))}"),
            ("Zusammenfassung",
             f"Dieses Strategiepapier analysiert {title.lower()} aus Sicht der {C['name']}. "
             "Es enthält quantitative Analysen, Szenarien und Handlungsempfehlungen für die Geschäftsführung."),
        ]
        save_pdf(FOLDER_STRAT, f"STRAT_{15+i:03d}_{sfn(title[:25])}.pdf", title, sections, confidential=True)

    # ── 15 F&E – more innovation docs ────────────────────────────────
    # Software design documentation
    for pname, pshort in [(C["prod1_name"], "CSP"), (C["prod2_name"], "OFN"), (C["prod3_name"], "DKS")]:
        sections = [
            ("Software Requirements Specification (SRS)",
             f"Produkt: {pname} Software-Modul v3.x\n"
             f"Erstellt: Jan Hoffmann, Software Engineering\nVersion: 2.0"),
            ("Anforderungen", [
                ["REQ-ID", "Anforderung", "Priorität", "Klasse IEC 62304"],
                ["SRS-001", "System muss Messdaten in < 100 ms verarbeiten", "MUSS", "B"],
                ["SRS-002", "Alarm bei Out-of-Range-Werten innerhalb 3 Sekunden", "MUSS", "B"],
                ["SRS-003", "Audit Trail für alle Benutzerereignisse", "MUSS", "B"],
                ["SRS-004", "Datenverschlüsselung AES-256", "MUSS", "B"],
                ["SRS-005", "Mehrsprachige Benutzeroberfläche (DE, EN, FR)", "SOLL", "A"],
            ]),
        ]
        save_pdf(FOLDER_RD, f"SRS_{pshort}_Software_Requirements_Spec.pdf",
                 f"Software Requirements Specification – {pname}", sections)
        # Software architecture
        sections = [
            ("Software Architecture Document (SAD)",
             f"Produkt: {pname} Software-Modul\nArchitekturmuster: Layered Architecture\n"
             f"IEC 62304 Software-Klasse: B"),
            ("Architekturüberblick",
             "Das System ist in folgende Schichten unterteilt:\n"
             "1. Präsentationsschicht (UI) – React Native, plattformübergreifend\n"
             "2. Anwendungsschicht – Messalgorithmen, Datenverarbeitung\n"
             "3. Datenzugriffsschicht – lokale Speicherung, Cloud-Sync\n"
             "4. Hardwareabstraktionsschicht (HAL) – Sensortreiber, Kommunikation"),
        ]
        save_pdf(FOLDER_RD, f"SAD_{pshort}_Software_Architecture_Document.pdf",
                 f"Software Architecture Document – {pname}", sections)
        # Usability engineering
        sections = [
            ("Usability Engineering File (IEC 62366-1)",
             f"Produkt: {pname}\nNorm: IEC 62366-1:2015+A1:2020\n"
             f"Verantwortlich: {C['cmo']}"),
            ("Benutzergruppen und Anwendungskontext",
             f"Primäre Benutzer: Ärzte, MTRA, medizinisches Assistenzpersonal.\n"
             "Anwendungsumgebung: Stationäre und ambulante Einrichtungen.\n"
             "Potenzielle Nutzungsfehler: Falsche Patientenzuordnung, "
             "unsachgemäße Reinigung, fehlerhafte Kalibrierung."),
            ("Usability Testing Ergebnisse", [
                ["Test", "Teilnehmer", "Erfolgsquote", "Kritische Fehler", "Status"],
                ["Summativer Usability Test", "15 Ärzte, 10 MTRA", "97%", "0", "BESTANDEN"],
                ["Formative Evaluation Prototyp", "8 Experten", "92%", "1 (behoben)", "BESTANDEN"],
            ]),
        ]
        save_pdf(FOLDER_RD, f"UEF_{pshort}_Usability_Engineering_File.pdf",
                 f"Usability Engineering File – {pname}", sections)
    # Additional innovation projects
    for i, topic in enumerate([
        "Machbarkeitsstudie Artificial Intelligence Diagnostics",
        "Vorentwicklung Wearable Kardiologiesensor",
        "Technologiescouting Quantencomputing Biosignale",
        "Prototyp-Bericht NextGen Cardevio v4.0",
        "Innovationsmanagement Prozessbeschreibung",
        "F&E Budgetplanung 2025 nach Projekten",
        "Kooperationsvertrag Fraunhofer IPA F&E",
        "F&E Mitarbeiterstunden-Reporting 2023",
    ], 1):
        if i % 3 == 0:
            sheets = [("Daten", ["Projekt", "Stunden", "Budget EUR", "Status"],
                       [(f"Projekt {j}", str(random.randint(100,500)), f"{random.randint(50,500)*1000:,}",
                         random.choice(["Aktiv", "Abgeschlossen"])) for j in range(1, 10)],
                       [25, 12, 16, 14])]
            make_xlsx_doc(FOLDER_RD, f"RD_{100+i:03d}_{sfn(topic[:25])}.xlsx", topic, sheets)
        else:
            sections = [("Übersicht", f"F&E-Dokument: {topic}\n"
                         f"Datum: {ds(2023, random.randint(1,12), random.randint(1,28))}\n"
                         f"Verantwortlich: {C['cto']}")]
            save_pdf(FOLDER_RD, f"RD_{100+i:03d}_{sfn(topic[:25])}.pdf", topic, sections)

    print(f"  Volume fill complete. Running total: {len(GENERATED_FILES)} docs")


def gen_topup():
    """Top-up pass to reach ~1000 total documents."""
    print("  Generating top-up documents...")

    # 02 Finanzen – 12 more USt (for 2022)
    for month in range(1, 13):
        rev_m = int(C["revenue_2022"] / 12 * (1 + random.uniform(-0.1, 0.12)))
        vat_amount = int(rev_m * 0.19 * 0.12)
        rows = [["Position", "Basis EUR", "Steuer EUR"],
                ["Umsätze 19%", f"{rev_m:,}", f"{int(rev_m*0.19):,}"],
                ["Vorsteuer", "—", f"-{int(rev_m*0.19*0.65):,}"],
                ["Zahllast", "—", f"{vat_amount:,}"]]
        save_pdf(FOLDER_FIN, f"UST_2022_{month:02d}_Voranmeldung_{MONTHS_DE[month]}.pdf",
                 f"Umsatzsteuervoranmeldung {MONTHS_DE[month]} 2022",
                 [("Steuermeldung", f"Zeitraum: {MONTHS_DE[month]} 2022"), ("Berechnung", rows)])

    # 03 HR – 30 more employment contracts (junior staff)
    junior_staff = [
        ("Leon Wagner", "Junior Software Engineer", "IT/Software", 62_000),
        ("Nina Braun", "QA Specialist", "Qualität", 58_000),
        ("Tobias Schulz", "Regulatory Affairs Coordinator", "Regulatory", 55_000),
        ("Lea Kramer", "Supply Chain Coordinator", "Einkauf", 52_000),
        ("Jonas Fischer", "Marketing Specialist", "Marketing", 56_000),
        ("Mia Zimmermann", "Clinical Data Manager", "Klinik", 68_000),
        ("Max Huber", "Production Engineer", "Produktion", 72_000),
        ("Sophie Hartmann", "Controlling Analyst", "Finanzen", 65_000),
        ("Tim Schwarz", "IT Support", "IT", 48_000),
        ("Eva Weiß", "HR Coordinator", "HR", 50_000),
        ("Daniel Becker", "Technical Sales", "Vertrieb", 70_000),
        ("Hannah Klein", "Legal Intern → Junior Legal", "Legal", 46_000),
        ("Niklas Groß", "F&E Engineer", "F&E", 74_000),
        ("Lara Koch", "Medical Writer", "Medical", 72_000),
        ("Florian Müller", "Business Intelligence Analyst", "IT/BI", 68_000),
        ("Miriam Lange", "Project Manager F&E", "F&E", 82_000),
        ("Andreas Richter", "Validation Engineer", "Qualität", 76_000),
        ("Laura Hoffmann", "Regulatory Specialist IVD", "Regulatory", 78_000),
        ("Sebastian Wolf", "Embedded Software Developer", "IT/Software", 82_000),
        ("Katharina Bauer", "Clinical Study Coordinator", "Klinik", 62_000),
        ("Lukas Schmidt", "Logistics Manager", "Produktion", 68_000),
        ("Jana Weber", "Patent Paralegal", "Legal", 54_000),
        ("Philipp Neumann", "Sales Manager Benelux", "Vertrieb", 88_000),
        ("Carina Maier", "Quality Engineer", "Qualität", 72_000),
        ("Stefan Jung", "Data Scientist Medical AI", "F&E", 95_000),
        ("Johanna Werner", "Medical Affairs Manager", "Medical", 88_000),
        ("Christian Braun", "Senior Controller", "Finanzen", 85_000),
        ("Franziska Koch", "Compliance Officer", "Legal", 78_000),
        ("Matthias Lange", "Head of After Sales Service", "Vertrieb", 98_000),
        ("Sandra Richter", "Training & Development Lead", "HR", 72_000),
    ]
    for i, (name, title, dept, salary) in enumerate(junior_staff, 16):
        gen_employment_contract(name, title, dept, salary, i)

    # 04 Vertrieb – monthly sales tracking 2023
    for month in range(1, 13):
        rev_m = int(C["revenue_2023"] / 12 * (1 + random.uniform(-0.08, 0.12)))
        rows = [
            (C["prod1_name"], str(random.randint(50, 80)), f"{int(rev_m*0.52):,}"),
            (C["prod2_name"], str(random.randint(70, 120)), f"{int(rev_m*0.33):,}"),
            (C["prod3_name"], str(random.randint(1200, 2800)), f"{int(rev_m*0.15):,}"),
        ]
        sheets = [("Monatsumsatz", ["Produkt", "Einheiten", "Umsatz EUR"], rows, [22, 12, 15])]
        make_xlsx_doc(FOLDER_SALES, f"SALES_MONTHLY_2023_{month:02d}_{MONTHS_DE[month]}.xlsx",
                      f"Monatsumsatz {MONTHS_DE[month]} 2023", sheets)

    # 06 Regulatorisch – labelling control per product per year
    for pname, pshort in [(C["prod1_name"], "CSP"), (C["prod2_name"], "OFN"), (C["prod3_name"], "DKS")]:
        for yr in [2021, 2022, 2023]:
            sections = [
                ("Etikettierungskontrolle gemäß MDR Anhang I Abschnitt 23",
                 f"Produkt: {pname}\nÜberprüfungsjahr: {yr}\nPrüfer: {C['qra']}\n"
                 f"Datum: {ds(yr, 10, random.randint(1,28))}"),
                ("Prüfergebnis", [
                    ["Anforderung", "Erfüllt", "Anmerkung"],
                    ["Name/Bezeichnung gemäß Art. 10 MDR", "JA", ""],
                    ["UDI-DI auf Etiketten", "JA", "GS1-128 und DataMatrix"],
                    ["LOT/Seriennummer", "JA", ""],
                    ["Herstellungsdatum oder Verfallsdatum", "JA", ""],
                    ["Hinweis Sterilisation (falls zutreffend)", "JA/N.A.", ""],
                    ["Symbol EN ISO 15223-1", "JA", "Alle Symbole konform"],
                    ["IFU-Referenz", "JA", "Beipackzettel und online"],
                    ["Mehrsprachige Versionen", "JA", "DE, EN, FR, IT, ES, NL"],
                ]),
            ]
            save_pdf(FOLDER_REG, f"LABEL_{pshort}_{yr}_Etikettierungskontrolle.pdf",
                     f"Etikettierungskontrolle – {pname} {yr}", sections)
        # IFU translation versions
        for lang in ["EN", "FR", "IT", "ES"]:
            sections = [
                (f"IFU Translation Control – {lang}",
                 f"Product: {pname}\nLanguage: {lang}\nTranslation completed: "
                 f"{ds(2022, random.randint(4,9), random.randint(1,28))}\n"
                 "Back-translation verified: YES"),
            ]
            save_pdf(FOLDER_REG, f"IFU_{pshort}_{lang}_Translation_Control.pdf",
                     f"IFU Translation Control {lang} – {pname}", sections)

    # 07 QMS – SOP revision history for key SOPs
    for sop_id, sop_title in SOPS[:10]:
        sheets = [
            ("Revisionshistorie", ["Version", "Datum", "Autor", "Änderungsgrund", "Genehmigt von"],
             [(f"{v}.0", ds(2023 - (4 - v), random.randint(1,12), random.randint(1,28)),
               random.choice([C["qra"], "Sarah Krause", "Markus Bauer"]),
               random.choice(["MDR-Anpassung", "Prozessoptimierung", "Normenupdate", "Erstfassung"]),
               C["ceo"]) for v in range(1, 5)],
             [10, 16, 22, 30, 22]),
        ]
        make_xlsx_doc(FOLDER_QMS, f"{sop_id}_Revisionshistorie.xlsx",
                      f"Revisionshistorie {sop_id} – {sop_title}", sheets)

    # 08 Klinisch – more investigator agreements and sub-study protocols
    for pshort, hosp, hshort in [
        ("DKS", C["hosp3"], "LMU2"),
        ("CSP", C["hosp4"], "UKH2"),
        ("OFN", C["hosp1"], "CHR2"),
    ]:
        gen_investigator_agreement(pshort, hosp, hshort)
    # Registry study protocols
    for pname, pshort in [(C["prod1_name"], "CSP"), (C["prod3_name"], "DKS")]:
        sections = [
            ("Registry-Studienprotokoll (PMCF-Registry)",
             f"Produkt: {pname}\nRegistry-Titel: BTP-{pshort}-Registry-Europe\n"
             f"Sponsor: {C['name']}\nLaufzeit: 5 Jahre (2023–2028)\n"
             f"Zielgröße: n ≥ 500 Patienten in 15 Zentren"),
            ("Einschlusskriterien",
             f"• Erwachsene Patienten ≥ 18 Jahre mit Indikation für {pname}\n"
             "• Einverständniserklärung vorliegend\n"
             "• Anwendung in Routineklinik (Real-World-Evidence)"),
            ("Endpunkte",
             "Primär: Sicherheit (SAE-Rate) und klinischer Nutzen (qualitative Befragung).\n"
             "Sekundär: Anwenderakzeptanz, Trainingsaufwand, Messzuverlässigkeit."),
        ]
        save_pdf(FOLDER_CLIN, f"REGISTRY_{pshort}_Studienprotokoll.pdf",
                 f"Registry-Studienprotokoll – {pname}", sections)

    # 09 IP – IP assignment agreements with employees
    for i, (name, _, _, _) in enumerate(EMPLOYEES[:6], 1):
        doc = Document()
        make_docx_header(doc, "Erfindungsabtretungsvereinbarung",
                         f"{C['name']} und {name}", confidential=True)
        add_clause(doc, 1, "Verpflichtung zur Meldung und Abtretung", [
            f"Der Arbeitnehmer {name} verpflichtet sich, alle im Rahmen seiner Tätigkeit "
            f"für {C['name']} gemachten Erfindungen und entwickelten schutzfähigen Leistungen "
            "unverzüglich schriftlich zu melden gemäß § 5 ArbnErfG.",
            "Die Gesellschaft kann die Diensterfindung in Anspruch nehmen. Mit Inanspruchnahme "
            "gehen alle Rechte daran auf die Gesellschaft über.",
        ])
        add_clause(doc, 2, "Vergütung", [
            "Der Arbeitnehmer erhält für in Anspruch genommene Diensterfindungen eine "
            "angemessene Vergütung gemäß Vergütungsrichtlinien Nr. 12 des DPMA.",
        ])
        save_docx(doc, FOLDER_IP, f"IP_ASSIGN_{i:02d}_{sfn(name.replace(' ','_'))}.docx")

    # 10 Versicherungen – supplementary docs
    for topic, fname in [
        ("D&O Versicherung – Selbstbehaltnachweis Geschäftsführung", "VER_DO_Selbstbehalt_Nachweis.pdf"),
        ("Produkthaftpflicht USA – gesonderte Policy", "VER_PL_USA_Policy.pdf"),
        ("Versicherungsbroker Mandate Letter 2024", "VER_Broker_Mandate_2024.pdf"),
        ("Claims Management Protokoll 2022–2023", "VER_Claims_Management.pdf"),
    ]:
        save_pdf(FOLDER_INS, fname, topic,
                 [("Inhalt", f"Versicherungsdokument: {topic}\n"
                   f"Datum: {ds(2024, 1, random.randint(1,28))}")])

    # 11 Immobilien – more docs
    for i, topic in enumerate([
        "Mietminderungsprotokoll 2022 – Klimaanlage defekt",
        "Wartungsvertrag Aufzug Freimannstraße 45",
        "Sicherheitsbegehungsprotokoll Q2 2023",
        "Abwasserentsorgungsvertrag München",
        "Reinigungsdienstleistungsvertrag 2023",
    ], 12):
        save_pdf(FOLDER_PROP, f"PROP_{i:03d}_{sfn(topic[:25])}.pdf", topic,
                 [("Dokument", f"{topic}\nObjekt: {C['full_address']}\n"
                   f"Datum: {ds(2023, random.randint(1,12), random.randint(1,28))}")])

    # 12 IT – cloud and system docs
    for i, (title, fname) in enumerate([
        ("IT Asset Inventar 2023", "IT_019_Asset_Inventar.xlsx"),
        ("Lizenzmanagement Software 2024", "IT_020_Lizenzmanagement.xlsx"),
        ("ERP Nutzerrollen und Berechtigungen", "IT_021_ERP_Nutzerrollen.xlsx"),
        ("IT Risikoanalyse nach BSI 200-2", "IT_022_IT_Risikoanalyse.pdf"),
        ("Netzwerkdiagramm und Segmentierungsplan", "IT_023_Netzwerkdiagramm.pdf"),
        ("Mobile Device Management Policy", "IT_024_MDM_Policy.docx"),
    ], 1):
        if fname.endswith(".xlsx"):
            sheets = [("Daten", ["ID", "Bezeichnung", "Status", "Verantwortlich"],
                       [(f"{i+j:03d}", f"Eintrag {j}", "Aktiv",
                         random.choice(["IT-Team", "Benjamin Wolf", "extern"])) for j in range(1, 18)],
                       [10, 35, 12, 20])]
            make_xlsx_doc(FOLDER_IT, fname, title, sheets)
        elif fname.endswith(".pdf"):
            save_pdf(FOLDER_IT, fname, title,
                     [("Inhalt", f"{title}. Erstellt: IT-Abteilung, {C['name']}.\n"
                       f"Datum: {ds(2023, random.randint(1,12), random.randint(1,28))}")])
        else:
            d = Document()
            make_docx_header(d, title)
            d.add_paragraph(f"Policy: {title}")
            save_docx(d, FOLDER_IT, fname)

    # 13 Compliance – more privacy and AML docs
    for i, (title, fname) in enumerate([
        ("Geldwäscheprävention Risikoanalyse", "COMP_AML_Risikoanalyse.pdf"),
        ("GwG Sorgfaltspflichten Lieferanten", "COMP_GwG_Sorgfaltspflichten.pdf"),
        ("DSGVO Einwilligungsverwaltung – Prozess", "COMP_DSGVO_Einwilligung_Prozess.docx"),
        ("Datenschutzbeauftragter Bestellungsurkunde", "COMP_DSB_Bestellung.pdf"),
        ("Verarbeitungsverzeichnis Update Q3 2023", "COMP_VVT_Update_Q3_2023.xlsx"),
        ("IT-Sicherheits-Compliance BSI C5", "COMP_BSI_C5_Self_Assessment.xlsx"),
        ("Exportkontrolle Dual-Use Güterprüfung", "COMP_Exportkontrolle_Dualuse.pdf"),
        ("Menschenrechtssorgfaltspflicht LkSG 2023", "COMP_LkSG_Sorgfalt_2023.pdf"),
    ], 1):
        if fname.endswith(".xlsx"):
            sheets = [("Daten", ["Kategorie", "Bewertung", "Maßnahme"],
                       [(f"Punkt {j}", random.choice(["Konform", "Bedingt konform", "Maßnahme erforderlich"]),
                         random.choice(["Keine", "Überprüfung", "Anpassung notwendig"])) for j in range(1, 12)],
                       [25, 22, 30])]
            make_xlsx_doc(FOLDER_COMP, fname, title, sheets)
        elif fname.endswith(".pdf"):
            save_pdf(FOLDER_COMP, fname, title,
                     [("Zusammenfassung", f"{title}.\nErstellt: Laura Neumann, Legal/Compliance\n"
                       f"Datum: {ds(2023, random.randint(1,12), random.randint(1,28))}")])
        else:
            d = Document()
            make_docx_header(d, title)
            d.add_paragraph(f"Compliance-Dokument: {title}")
            save_docx(d, FOLDER_COMP, fname)

    # 14 Strategie – more investor relations docs
    for i, title in enumerate([
        "Pitch Deck Series D – Version 4.2",
        "Due Diligence Data Room Index 2024",
        "Investor FAQ – Häufige Fragen zum Produkt",
        "Wettbewerbsvergleich Produktmatrix 2024",
        "SWOT-Analyse Sentavia Precision 2024",
        "KPI-Dashboard Investoren 2024 (Excel-Modell)",
        "Finanzprognose konservativ vs. optimistisch 2024–2028",
        "Liquidationspräferenz Berechnung Exit-Szenarien",
        "Term Sheet Entwurf Series D (Vorlage)",
    ], 1):
        if "Excel" in title:
            sheets = [("KPIs", ["KPI", "Q1 2024", "Q2 2024E", "Q3 2024E", "Q4 2024E"],
                       [(f"KPI {j}", f"{random.randint(70,110)}%", f"{random.randint(75,115)}%",
                         f"{random.randint(80,120)}%", f"{random.randint(85,125)}%") for j in range(1, 10)],
                       [30, 12, 12, 12, 12])]
            make_xlsx_doc(FOLDER_STRAT, f"STRAT_{25+i:03d}_{sfn(title[:25])}.xlsx", title, sheets)
        else:
            sections = [("Dokument", f"Titel: {title}\nDatum: {ds(2024, random.randint(1,4), random.randint(1,28))}\n"
                         "Vertraulich: JA – Nur für authorisierte Empfänger")]
            save_pdf(FOLDER_STRAT, f"STRAT_{25+i:03d}_{sfn(title[:25])}.pdf", title, sections, confidential=True)

    # 15 F&E – more innovation and study docs
    for pname, pshort in [(C["prod1_name"], "CSP"), (C["prod2_name"], "OFN"), (C["prod3_name"], "DKS")]:
        # Software test reports
        sw_test_rows = [["Testfall-ID", "Beschreibung", "Erwartet", "Ergebnis", "Status"]] + [
            [f"TC-{pshort}-{i:03d}", f"Testfall {i}", "Wie spezifiziert",
             "Wie spezifiziert", "BESTANDEN"] for i in range(1, 12)
        ]
        sections = [
            ("Software-Testbericht",
             f"Produkt: {pname} Software v3.x\nTest-Art: Systemtest / Regressionstest\n"
             f"Testdatum: {ds(2023, random.randint(1,12), random.randint(1,28))}\n"
             f"Tester: Jan Hoffmann, Software Engineering"),
            ("Testergebnisse", sw_test_rows),
        ]
        save_pdf(FOLDER_RD, f"TEST_{pshort}_Softwaretest_Systemtest.pdf",
                 f"Software-Testbericht Systemtest – {pname}", sections)
        # Sterilization/cleaning validation
        sections = [
            ("Reinigungsvalidierung",
             f"Produkt: {pname}\nValidierungsart: Reinigungsvalidierung\n"
             f"Norm: ISO 11135 / ISO 17664\nDatum: {ds(2022, random.randint(4,8), random.randint(1,28))}"),
            ("Ergebnisse",
             "Alle getesteten Reinigungsmethoden erfüllen die Akzeptanzkriterien. "
             "Endotoxinwerte nach Reinigung < 0,25 EU/mL. Bioburden-Reduktion > 3 log."),
        ]
        save_pdf(FOLDER_RD, f"CLEAN_{pshort}_Reinigungsvalidierung.pdf",
                 f"Reinigungsvalidierung – {pname}", sections)
        # Shelf life / stability testing
        sections = [
            ("Lagerstabilitätsstudie",
             f"Produkt: {pname}\nNorm: ISO 11607 / ASTM F1980\n"
             f"Laufzeit: 36 Monate (Echtzeit) / 6 Monate (Accelerated)\n"
             f"Temperatur: 23 ± 2 °C, 50 ± 5 % rF"),
            ("Zwischenergebnisse", [
                ["Zeitpunkt", "Dichtigkeit", "Visuell", "Funktionstest", "Ergebnis"],
                ["T=0 (Ausgangsprüfung)", "OK", "OK", "OK", "BESTANDEN"],
                ["T=6 Monate", "OK", "OK", "OK", "BESTANDEN"],
                ["T=12 Monate", "OK", "OK", "OK", "BESTANDEN"],
                ["T=24 Monate (laufend)", "OK bisher", "OK", "OK", "LÄUFT"],
            ]),
        ]
        save_pdf(FOLDER_RD, f"STAB_{pshort}_Lagerstabilitaet.pdf",
                 f"Lagerstabilitätsstudie – {pname}", sections)

    print(f"  Top-up complete. Running total: {len(GENERATED_FILES)} docs")


def gen_final_push():
    """Final pass to push total to ~1000."""
    print("  Generating final-push documents...")

    # ── 02 Finanzen – monthly BWA 2021 already done; add 2020 + konsolidiert ─
    for yr in [2020]:
        gen_bwa_monatlich(yr)
    for yr in [2021, 2022, 2023]:
        sections = [
            ("Konsolidierter Monatsbericht",
             f"Unternehmen: {C['name']}\nZeitraum: Geschäftsjahr {yr}\n"
             f"Konsolidiert: Alle Standorte\nErstellt: {C['cfo']}"),
            ("Schlüsselkennzahlen", [
                ["KPI", "Ist EUR", "Plan EUR", "Abweichung"],
                ["Umsatz", f"{C.get(f'revenue_{yr}', 87_200_000):,}", f"{int(C.get(f'revenue_{yr}', 87_200_000)*0.97):,}", "+3,1%"],
                ["EBITDA", f"{C.get(f'ebitda_{yr}', 12_400_000):,}", f"{int(C.get(f'ebitda_{yr}', 12_400_000)*0.96):,}", "+4,2%"],
                ["Mitarbeiter", str(C.get(f"employees_{yr}", 612)), str(int(C.get(f"employees_{yr}", 612)*0.98)), f"+{random.randint(5,20)}"],
            ]),
        ]
        save_pdf(FOLDER_FIN, f"KONSOL_{yr}_Jahresuebersicht.pdf",
                 f"Konsolidierter Jahresüberblick {yr}", sections)

    # ── 03 HR – personal development plans ───────────────────────────
    for name, title, dept, _ in EMPLOYEES:
        sections = [
            ("Entwicklungsplan",
             f"Mitarbeiter: {name}\nPosition: {title}\nAbteilung: {dept}\n"
             f"Beurteilungsjahr: 2023\nBeurteiler: Abteilungsleitung"),
            ("Entwicklungsziele", [
                ["Ziel", "Maßnahme", "Zeitraum", "Status"],
                ["Führungskompetenz stärken", "Führungsseminar extern", "Q2 2024", "Geplant"],
                ["MDR-Kenntnisse vertiefen", "Online-Training MDCG", "Q1 2024", "Abgeschlossen"],
                ["Englischkenntnisse", "Businesskurs (B2+C1)", "Laufend", "In Arbeit"],
            ]),
        ]
        save_pdf(FOLDER_HR, f"EP_{sfn(name.replace(' ','_'))}_Entwicklungsplan_2023.pdf",
                 f"Entwicklungsplan – {name} 2023", sections)

    # ── 04 Vertrieb – 2021 monthly sales ─────────────────────────────
    for month in range(1, 13):
        rev_m = int(C["revenue_2021"] / 12 * (1 + random.uniform(-0.08, 0.12)))
        rows = [(C["prod1_name"], str(random.randint(40,70)), f"{int(rev_m*0.52):,}"),
                (C["prod2_name"], str(random.randint(60,100)), f"{int(rev_m*0.33):,}"),
                (C["prod3_name"], str(random.randint(800,2000)), f"{int(rev_m*0.15):,}")]
        sheets = [("Monatsumsatz", ["Produkt", "Einheiten", "Umsatz EUR"], rows, [22, 12, 15])]
        make_xlsx_doc(FOLDER_SALES, f"SALES_MONTHLY_2021_{month:02d}_{MONTHS_DE[month]}.xlsx",
                      f"Monatsumsatz {MONTHS_DE[month]} 2021", sheets)

    # ── 06 Regulatorisch – annual QA agreement with NB ────────────────
    for yr in [2021, 2022, 2023]:
        sections = [
            ("Jahresbericht an die Benannte Stelle",
             f"Berichtsjahr: {yr}\nHersteller: {C['name']}\n"
             f"Benannte Stelle: {C['nb']} (Kennnummer {C['nb_id']})\n"
             f"Datum: {ds(yr+1, 1, 31)}"),
            ("Produktübersicht", [
                ["Produkt", "Zertifikat", "Gültig bis", "Änderungen", "Vigilanz"],
                [C["prod1_name"], C["nb_cert1"], ds(yr+3, 12, 31), "Keine wesentlichen", f"{random.randint(0,2)} Meldungen"],
                [C["prod2_name"], "G1 20 12 87024 021", ds(yr+3, 6, 30), "SW-Update v3.1", f"{random.randint(0,1)} Meldungen"],
                [C["prod3_name"], "G1 20 12 87024 034", ds(yr+2, 3, 31), "Keine", "0 Meldungen"],
            ]),
        ]
        save_pdf(FOLDER_REG, f"NB_ANNUAL_{yr}_Jahresbericht_BenannteStelle.pdf",
                 f"Jahresbericht an die Benannte Stelle {yr}", sections)
    # QMS surveillance audit prep checklists
    for yr in [2021, 2022, 2023]:
        sheets = [
            (f"Auditvorbereitungscheckliste {yr}",
             ["Nr.", "Anforderung MDR / ISO 13485", "Nachweis", "Status", "Anmerkung"],
             [(str(i), f"Anforderung {i}", f"Dokument-Ref. {i}", random.choice(["OK", "In Prüfung"]), "")
              for i in range(1, 25)],
             [6, 40, 22, 14, 20]),
        ]
        make_xlsx_doc(FOLDER_REG, f"NB_AUDIT_PREP_{yr}_Checkliste.xlsx",
                      f"Auditvorbereitungscheckliste Benannte Stelle {yr}", sheets)
    # Post-Market data summaries per year per product
    for pname, pshort in [(C["prod1_name"], "CSP"), (C["prod2_name"], "OFN")]:
        for yr in [2021, 2022]:
            gen_pms_report(pname, pshort, yr)

    # ── 07 QMS – SOP training records per SOP ────────────────────────
    for sop_id, sop_title in SOPS[10:20]:
        sheets = [
            ("Schulungsnachweis", ["Mitarbeiter", "Datum", "Trainer", "Bestanden", "Gültig bis"],
             [(name, ds(2023, random.randint(1,12), random.randint(1,28)),
               random.choice([C["qra"], "Sarah Krause"]),
               "JA",
               ds(2025, random.randint(1,12), 1))
              for name, _, _, _ in EMPLOYEES], [25, 16, 22, 10, 14]),
        ]
        make_xlsx_doc(FOLDER_QMS, f"{sop_id}_Schulungsnachweise_2023.xlsx",
                      f"Schulungsnachweise {sop_id}", sheets)

    # ── 08 Klinisch – protocol amendments ────────────────────────────
    for pname, pshort in [(C["prod1_name"], "CSP"), (C["prod2_name"], "OFN")]:
        for amend_num in range(1, 4):
            sections = [
                (f"Protokollamendment Nr. {amend_num}",
                 f"Studie: BTP-2021-{pshort}\nProdukt: {pname}\n"
                 f"Amendment: {amend_num}\nDatum: {ds(2022, random.randint(1,12), random.randint(1,28))}\n"
                 f"Genehmigt: Ethikkommission, {ds(2022, random.randint(1,12), random.randint(1,28))}"),
                ("Änderungsinhalt",
                 f"Das vorliegende Amendment Nr. {amend_num} modifiziert das Studienprotokoll "
                 f"für {pname} wie folgt: Anpassung der Einschlusskriterien (Erweiterung auf "
                 f"Patienten mit moderater Einschränkung), Ergänzung eines sekundären Endpunkts "
                 "(Behandlungszeit bis Diagnose), administrative Korrekturen."),
            ]
            save_pdf(FOLDER_CLIN, f"AMEND_{pshort}_{amend_num:02d}_Protokollamendment.pdf",
                     f"Protokollamendment Nr. {amend_num} – {pname}", sections)

    # ── 09 IP – additional FTO and watch notices ──────────────────────
    for i in range(1, 7):
        sections = [
            ("IP Watch Notice",
             f"Referenz: WATCH-{C['short']}-2023-{i:03d}\n"
             f"Erstellt: {C['patent_anwalt']}\nDatum: {ds(2023, random.randint(1,12), random.randint(1,28))}"),
            ("Beobachtetes Patent",
             f"Patentinhaber: Wettbewerber {i}\nPatent-Nr.: EP{random.randint(1000000,3999999)}\n"
             f"Titel: Verfahren zur {random.choice(['kardiologischen Diagnostik', 'orthopädischen Navigation', 'IVD-Analyse'])}\n"
             f"Status: {random.choice(['Angemeldet', 'Erteilt', 'Widerspruchsphase'])}\n"
             f"Handlungsbedarf: {random.choice(['Beobachten', 'Abgrenzung prüfen', 'Kein Handlungsbedarf'])}"),
        ]
        save_pdf(FOLDER_IP, f"IP_WATCH_{i:03d}_Patent_Beobachtung.pdf",
                 f"IP Watch Notice – Patent-Beobachtung {i}", sections)

    # ── 13 Compliance – training attendance records ────────────────────
    for training in ["MDR-Update Schulung 2023", "DSGVO-Grundschulung 2023",
                     "Anti-Korruptions-Training 2023", "HWG-Compliance Training 2023"]:
        rows = [(name, title, dept, ds(2023, random.randint(1,12), random.randint(1,28)), "Bestanden")
                for name, title, dept, _ in EMPLOYEES]
        sheets = [("Teilnahmenliste", ["Name", "Position", "Abteilung", "Datum", "Ergebnis"],
                   rows, [22, 28, 18, 16, 12])]
        make_xlsx_doc(FOLDER_COMP, f"COMP_TRAINING_{sfn(training[:25])}.xlsx", training, sheets)

    # ── 14 Strategie – pipeline analysis ──────────────────────────────
    for yr in [2021, 2022]:
        sheets = [("Jahresplanung", ["Initiative", "Status", "Budget EUR", "Verantwortlich"],
                   [(f"Initiative {i}", random.choice(["Abgeschlossen", "In Arbeit", "Geplant"]),
                     f"{random.randint(100,2000)*1000:,}",
                     random.choice([C["ceo"], C["cto"], C["cfo"], "Christoph Braun"]))
                    for i in range(1, 15)], [35, 16, 15, 22])]
        make_xlsx_doc(FOLDER_STRAT, f"STRAT_JAHRESPLAN_{yr}.xlsx",
                      f"Strategische Jahresplanung {yr}", sheets)

    # ── 15 F&E – design reviews ────────────────────────────────────────
    for pname, pshort in [(C["prod1_name"], "CSP"), (C["prod2_name"], "OFN"), (C["prod3_name"], "DKS")]:
        for dr_phase in ["PDR", "CDR", "Final"]:  # Preliminary/Critical/Final Design Review
            sections = [
                (f"Design Review – {dr_phase} ({pname})",
                 f"Produkt: {pname}\nPhase: {dr_phase} Design Review\n"
                 f"Datum: {ds(2021, random.randint(1,12), random.randint(1,28))}\n"
                 f"Geleitet von: {C['cto']}\nTeilnehmer: {C['ceo']}, {C['cmo']}, {C['qra']}, F&E-Team"),
                ("Befunde und Maßnahmen", [
                    ["Nr.", "Befund", "Kritikalität", "Maßnahme", "Status"],
                    ["DR-1", "Sensor-Kalibrierungskonzept überarbeiten", "Mittel", "Neue Kalibrierungs-SOP", "Erledigt"],
                    ["DR-2", "UI-Feedback zu Alarmton verbessern", "Gering", "Usability-Test geplant", "Erledigt"],
                    ["DR-3", "Dokumentation IEC 62304 vervollständigen", "Hoch", "Sofortmaßnahme", "Erledigt"],
                ]),
                ("Freigabe",
                 f"Der {dr_phase} Design Review wurde erfolgreich abgeschlossen. "
                 f"Alle Action Items wurden termingerecht umgesetzt. "
                 f"Freigabe erteilt von: {C['cto']}, {ds(2021, random.randint(8,12), random.randint(1,28))}."),
            ]
            save_pdf(FOLDER_RD, f"DR_{pshort}_{dr_phase}_Design_Review.pdf",
                     f"{dr_phase} Design Review – {pname}", sections)

    print(f"  Final push complete. Running total: {len(GENERATED_FILES)} docs")


if __name__ == "__main__":
    print("Sentavia Precision GmbH – Due Diligence Document Generator")
    print("=" * 60)
    create_folder_structure()

    generate_01_gesellschaftsrecht()
    generate_02_finanzen()
    generate_03_personal_hr()
    generate_04_vertrieb()
    generate_05_einkauf()
    generate_06_regulatorisch()
    generate_07_qualitaet()
    generate_08_klinisch()
    generate_09_ip()
    generate_10_versicherungen()
    generate_11_immobilien()
    generate_12_it()
    generate_13_compliance()
    generate_14_strategie()
    generate_15_fuer()
    gen_volume_fill()
    gen_topup()
    gen_final_push()

    print_summary()
