#!/usr/bin/env python3
"""
Generate 300 mock due-diligence documents for Halbreiter Maschinenbau GmbH.
All documents are cross-consistent (same company facts, financials, people).
"""
# --- portable-paths-prelude --- (do not edit) ---
import sys
from pathlib import Path as _PathlibPath
_RP = _PathlibPath(__file__).resolve().parent
while not (_RP / "enhance_lib.py").exists() and _RP.parent != _RP:
    _RP = _RP.parent
_ROOT = _RP
sys.path.insert(0, str(_ROOT))
# --- end prelude ---

import os
import random
from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.platypus import HRFlowable
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

# ─── CANONICAL COMPANY DATA ───────────────────────────────────────────────────

C = {
    "name":           "Halbreiter Maschinenbau GmbH",
    "short":          "MMB",
    "hrb":            "HRB 47312",
    "amtsgericht":    "Amtsgericht Köln",
    "founded":        "14. März 1985",
    "founded_year":   1985,
    "address":        "Industriestraße 12",
    "plz":            "50829",
    "city":           "Köln",
    "full_address":   "Industriestraße 12, 50829 Köln",
    "phone":          "+49 221 47832-0",
    "fax":            "+49 221 47832-99",
    "email":          "info@halbreiter-maschinenbau.de",
    "web":            "www.halbreiter-maschinenbau.de",
    "ust_id":         "DE 198 765 432",
    "steuernr":       "215/5765/9876",
    "fa":             "Finanzamt Köln-Nord",
    "iban":           "DE89 3007 0010 0123 4567 89",
    "bic":            "DEUTDEDB",
    "bank":           "Deutsche Bank AG, Filiale Köln",
    "geschaeftszweck": "Entwicklung, Herstellung und Vertrieb von Sondermaschinen und Anlagen für die metallverarbeitende Industrie",
    # Managing directors
    "gf1_name":       "Klaus Müller",
    "gf1_title":      "Geschäftsführer (CEO)",
    "gf1_born":       "12. Februar 1963",
    "gf2_name":       "Sandra Becker",
    "gf2_title":      "Geschäftsführerin (CFO)",
    "gf2_born":       "7. September 1977",
    # Shareholders
    "gesellschafter1": "Klaus Müller",
    "anteil1":         "60 %",
    "gesellschafter2": "Müller Familien-GbR",
    "anteil2":         "40 %",
    # Financials (consistent across all docs)
    "employees_2022":  231,
    "employees_2023":  247,
    "employees_2024e": 251,
    "revenue_2022":    43_250_000,
    "revenue_2023":    48_630_000,
    "revenue_2024e":   52_100_000,
    "ebit_2022":       3_215_000,
    "ebit_2023":       3_890_000,
    "ebit_2024e":      4_200_000,
    "ebitda_2022":     5_100_000,
    "ebitda_2023":     5_980_000,
    "ebitda_2024e":    6_400_000,
    "bilanzsumme_2022": 28_400_000,
    "bilanzsumme_2023": 31_200_000,
    "eigenkapital_2022": 14_100_000,
    "eigenkapital_2023": 16_900_000,
    "stammkapital":    250_000,
    # Key customers (referenced in contracts)
    "kunde1": "ThyssenKrupp Steel Europe AG",
    "kunde2": "Bosch Rexroth AG",
    "kunde3": "Hella GmbH & Co. KGaA",
    "kunde4": "Viessmann Climate Solutions SE",
    "kunde5": "Dürr AG",
    # Key suppliers
    "lieferant1": "Schunk GmbH & Co. KG",
    "lieferant2": "Igus GmbH",
    "lieferant3": "Siemens AG – Antriebstechnik",
    "lieferant4": "Trumpf SE + Co. KG",
    # Insurance
    "versicherer1": "Allianz SE",
    "versicherer2": "HDI Global SE",
    # Law firm / auditor
    "wp":       "BDO AG Wirtschaftsprüfungsgesellschaft",
    "steuerber": "KPMG AG Wirtschaftsprüfungsgesellschaft",
    "anwalt":    "Heuking Kühn Lüer Wojtek Partnerschaft mbB",
    # IT
    "erp":       "SAP S/4HANA (On-Premise), Release 2023",
    "crm":       "Salesforce Sales Cloud (Enterprise)",
    # Products
    "produkt1":  "Pressenlinie PL-500 (hydraulische Stanzpresse)",
    "produkt2":  "Förderbandanlage FB-200 (modulares Transportsystem)",
    "produkt3":  "Laserschneidanlage LS-800 (5-Achs-CNC)",
    "produkt4":  "Montageroboter MR-150 (Kollaborationsroboter-Integration)",
}

def fmt_eur(n):
    return f"{n:,.0f} €".replace(",", ".")

def fmt_num(n):
    return f"{n:,}".replace(",", ".")

# Output directory structure
BASE = Path(f"{_ROOT}/mueller_small")

FOLDERS = {
    "01_Gesellschaftsrecht":    "docx",
    "02_Finanzen":              "mixed",
    "03_Personal_HR":           "mixed",
    "04_Vertraege_Kunden":      "docx",
    "05_Vertraege_Lieferanten": "docx",
    "06_Immobilien":            "docx",
    "07_IP_Lizenzen":           "docx",
    "08_Versicherungen":        "pdf",
    "09_Compliance":            "docx",
    "10_IT_Infrastruktur":      "mixed",
    "11_Strategie_Planung":     "mixed",
}

for folder in FOLDERS:
    (BASE / folder).mkdir(parents=True, exist_ok=True)

# ─── HELPERS ──────────────────────────────────────────────────────────────────

def date_str(y, m, d):
    months_de = ["", "Januar", "Februar", "März", "April", "Mai", "Juni",
                 "Juli", "August", "September", "Oktober", "November", "Dezember"]
    return f"{d}. {months_de[m]} {y}"

def make_docx_header(doc, title, subtitle=None):
    doc.add_heading(C["name"], level=1)
    p = doc.add_paragraph()
    p.add_run(C["full_address"] + " | " + C["email"]).font.size = Pt(9)
    p.add_run("  |  " + C["hrb"] + ", " + C["amtsgericht"]).font.size = Pt(9)
    doc.add_paragraph("─" * 80)
    h = doc.add_heading(title, level=2)
    if subtitle:
        sub = doc.add_paragraph(subtitle)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def save_docx(doc, folder, filename):
    path = BASE / folder / filename
    doc.save(path)
    print(f"  ✓ {folder}/{filename}")
    return path

def make_pdf(folder, filename, title, sections):
    """sections: list of (heading, body_text) tuples"""
    path = BASE / folder / filename
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            leftMargin=2.5*cm, rightMargin=2.5*cm,
                            topMargin=2.5*cm, bottomMargin=2.5*cm)
    styles = getSampleStyleSheet()
    normal = ParagraphStyle('body', parent=styles['Normal'], fontSize=10, leading=14, spaceAfter=6)
    heading1 = ParagraphStyle('h1', parent=styles['Heading1'], fontSize=14, spaceBefore=12, spaceAfter=6)
    heading2 = ParagraphStyle('h2', parent=styles['Heading2'], fontSize=11, spaceBefore=10, spaceAfter=4)
    small = ParagraphStyle('small', parent=styles['Normal'], fontSize=8, textColor=colors.grey)

    story = []
    # Header
    story.append(Paragraph(C["name"], heading1))
    story.append(Paragraph(C["full_address"] + " | " + C["email"] + " | " + C["hrb"] + ", " + C["amtsgericht"], small))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#2A2E4B")))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(title, heading1))
    story.append(Spacer(1, 0.4*cm))

    for sec_title, sec_body in sections:
        if sec_title:
            story.append(Paragraph(sec_title, heading2))
        if isinstance(sec_body, str):
            for para in sec_body.split("\n\n"):
                if para.strip():
                    story.append(Paragraph(para.strip(), normal))
        elif isinstance(sec_body, list):
            # Table data
            t = Table(sec_body, hAlign='LEFT')
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#2A2E4B")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F5F5")]),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ]))
            story.append(t)
        story.append(Spacer(1, 0.3*cm))

    doc.build(story)
    print(f"  ✓ {folder}/{filename}")
    return path

def make_xlsx(folder, filename, title, sheets_data):
    """sheets_data: list of (sheet_name, headers, rows, formats)"""
    path = BASE / folder / filename
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    header_fill = PatternFill("solid", fgColor="2A2E4B")
    header_font = Font(bold=True, color="FFFFFF", size=10)
    title_font = Font(bold=True, size=12, color="2A2E4B")
    alt_fill = PatternFill("solid", fgColor="F0F4FA")
    thin = Side(style='thin', color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for sheet_name, headers, rows, col_widths in sheets_data:
        ws = wb.create_sheet(sheet_name)
        ws.append([title, "", "", ""])
        ws["A1"].font = title_font
        ws.append([C["name"], "", "", ""])
        ws["A2"].font = Font(italic=True, size=9, color="666666")
        ws.append([])

        header_row = 4
        ws.append(headers)
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=header_row, column=col_idx)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', wrap_text=True)
            cell.border = border

        for row_idx, row in enumerate(rows, header_row + 1):
            ws.append(row)
            fill = alt_fill if row_idx % 2 == 0 else None
            for col_idx in range(1, len(headers) + 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.border = border
                if fill:
                    cell.fill = fill

        for col_idx, width in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(col_idx)].width = width

    wb.save(path)
    print(f"  ✓ {folder}/{filename}")
    return path

# ─── DOCUMENT GENERATORS ──────────────────────────────────────────────────────

generated = []

print("\n📁 01_Gesellschaftsrecht")

# 1. Gesellschaftsvertrag
def gen_gesellschaftsvertrag():
    doc = Document()
    make_docx_header(doc, "Gesellschaftsvertrag", "in der Fassung vom 15. Januar 2019")
    doc.add_heading("§ 1 Firma und Sitz", level=3)
    doc.add_paragraph(
        f"(1) Die Gesellschaft führt die Firma: {C['name']}.\n"
        f"(2) Sitz der Gesellschaft ist {C['city']}.\n"
        f"(3) Das Geschäftsjahr ist das Kalenderjahr."
    )
    doc.add_heading("§ 2 Gegenstand des Unternehmens", level=3)
    doc.add_paragraph(C["geschaeftszweck"] + ". Die Gesellschaft ist berechtigt, alle Geschäfte vorzunehmen und Maßnahmen zu ergreifen, die dem Gesellschaftszweck zu dienen geeignet sind.")
    doc.add_heading("§ 3 Stammkapital und Geschäftsanteile", level=3)
    doc.add_paragraph(
        f"(1) Das Stammkapital der Gesellschaft beträgt {fmt_eur(C['stammkapital'])} (in Worten: Zweihundertfünfzigtausend Euro).\n"
        f"(2) Hiervon übernimmt {C['gesellschafter1']} einen Geschäftsanteil im Nennwert von "
        f"{fmt_eur(C['stammkapital'] * 0.6)} ({C['anteil1']}) und "
        f"{C['gesellschafter2']} einen Geschäftsanteil im Nennwert von "
        f"{fmt_eur(C['stammkapital'] * 0.4)} ({C['anteil2']})."
    )
    doc.add_heading("§ 4 Geschäftsführung und Vertretung", level=3)
    doc.add_paragraph(
        f"(1) Die Gesellschaft hat einen oder mehrere Geschäftsführer.\n"
        f"(2) Ist nur ein Geschäftsführer bestellt, so vertritt er die Gesellschaft allein. Sind mehrere Geschäftsführer bestellt, so wird die Gesellschaft durch zwei Geschäftsführer gemeinsam oder durch einen Geschäftsführer in Gemeinschaft mit einem Prokuristen vertreten.\n"
        f"(3) Die Gesellschafterversammlung kann einzelnen Geschäftsführern Einzelvertretungsbefugnis und die Befreiung von den Beschränkungen des § 181 BGB erteilen."
    )
    doc.add_heading("§ 5 Gesellschafterversammlung", level=3)
    doc.add_paragraph(
        "(1) Die ordentliche Gesellschafterversammlung findet jährlich innerhalb von acht Monaten nach Ablauf des Geschäftsjahres statt.\n"
        "(2) Die Einberufung erfolgt durch den Geschäftsführer mittels eingeschriebenen Briefes unter Einhaltung einer Frist von zwei Wochen.\n"
        "(3) Beschlüsse der Gesellschafterversammlung bedürfen, soweit nicht das Gesetz oder dieser Vertrag etwas anderes bestimmen, der einfachen Mehrheit der abgegebenen Stimmen."
    )
    doc.add_heading("§ 6 Jahresabschluss und Gewinnverwendung", level=3)
    doc.add_paragraph(
        "(1) Binnen drei Monaten nach Ablauf des Geschäftsjahres ist ein Jahresabschluss aufzustellen.\n"
        "(2) Über die Verwendung des Jahresüberschusses beschließt die Gesellschafterversammlung."
    )
    doc.add_paragraph()
    doc.add_paragraph(f"Köln, den 15. Januar 2019")
    doc.add_paragraph()
    doc.add_paragraph(f"_________________________        _________________________")
    doc.add_paragraph(f"{C['gf1_name']}                   {C['gesellschafter2']}, vertr. d. d. Gesellschafter")
    return save_docx(doc, "01_Gesellschaftsrecht", "GV_001_Gesellschaftsvertrag_2019.docx")

# 2. Handelsregisterauszug
def gen_hrb_auszug(year):
    sections = [
        ("Registereintrag", f"""
Registergericht: {C["amtsgericht"]}
Registernummer: {C["hrb"]}
Rechtsform: Gesellschaft mit beschränkter Haftung (GmbH)
Firma: {C["name"]}
Sitz: {C["city"]}
Geschäftsanschrift: {C["full_address"]}
Gegenstand: {C["geschaeftszweck"]}
Stammkapital: {fmt_eur(C["stammkapital"])}
"""),
        ("Geschäftsführer", f"""
1. {C["gf1_name"]}, geboren am {C["gf1_born"]}, {C["city"]}; einzelvertretungsberechtigt; befugt, im Namen der Gesellschaft mit sich im eigenen Namen oder als Vertreter eines Dritten Rechtsgeschäfte abzuschließen.

2. {C["gf2_name"]}, geboren am {C["gf2_born"]}, {C["city"]}; einzelvertretungsberechtigt.
"""),
        ("Gesellschafter", f"""
1. {C["gesellschafter1"]}, {C["city"]}: Geschäftsanteil {fmt_eur(C["stammkapital"] * 0.6)} ({C["anteil1"]})
2. {C["gesellschafter2"]}, {C["city"]}: Geschäftsanteil {fmt_eur(C["stammkapital"] * 0.4)} ({C["anteil2"]})
"""),
        ("Datum", f"Auszug erstellt am: {date_str(year, 3, 15)}\nBeglaubigt durch: {C['amtsgericht']}"),
    ]
    return make_pdf("01_Gesellschaftsrecht", f"GR_{year}_Handelsregisterauszug.pdf", f"Handelsregisterauszug {year}", sections)

# 3. Gesellschafterliste
def gen_gesellschafterliste():
    doc = Document()
    make_docx_header(doc, "Gesellschafterliste", f"gemäß § 40 GmbHG, Stand: 1. Januar 2024")
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(["Nr.", "Name / Firma", "Anschrift", "Nennbetrag", "Anteil"]):
        hdr[i].text = h
    row1 = table.add_row().cells
    row1[0].text = "1"
    row1[1].text = C["gesellschafter1"]
    row1[2].text = "Rosenweg 4, 50823 Köln"
    row1[3].text = fmt_eur(C["stammkapital"] * 0.6)
    row1[4].text = C["anteil1"]
    row2 = table.add_row().cells
    row2[0].text = "2"
    row2[1].text = C["gesellschafter2"]
    row2[2].text = "Eichenallee 12, 50737 Köln"
    row2[3].text = fmt_eur(C["stammkapital"] * 0.4)
    row2[4].text = C["anteil2"]
    doc.add_paragraph()
    doc.add_paragraph(f"Köln, den 1. Januar 2024\n\n{C['gf1_name']}\nGeschäftsführer")
    return save_docx(doc, "01_Gesellschaftsrecht", "GR_002_Gesellschafterliste_2024.docx")

# 4. GF-Anstellungsvertrag
def gen_gf_anstellungsvertrag(gf_name, gf_title, salary, date_str_val, num):
    doc = Document()
    make_docx_header(doc, f"Geschäftsführer-Anstellungsvertrag", f"zwischen {C['name']} und {gf_name}")
    doc.add_heading("§ 1 Tätigkeit", level=3)
    doc.add_paragraph(f"{gf_name} wird als {gf_title} der {C['name']} bestellt. Die Bestellung erfolgte durch Gesellschafterbeschluss vom {date_str_val}.")
    doc.add_heading("§ 2 Vergütung", level=3)
    doc.add_paragraph(f"(1) Die Gesellschaft zahlt dem Geschäftsführer ein festes Jahresgehalt von {fmt_eur(salary)} brutto, zahlbar in 12 gleichen Monatsraten.\n(2) Zusätzlich erhält der Geschäftsführer eine ergebnisabhängige Tantieme in Höhe von bis zu 20 % des festen Jahresgehalts, sofern die vereinbarten Ziele erreicht werden.")
    doc.add_heading("§ 3 Dienstwagen", level=3)
    doc.add_paragraph("Der Geschäftsführer erhält einen Dienstwagen der Mittelklasse (Bruttolistenpreis bis 65.000 €), der auch privat genutzt werden darf.")
    doc.add_heading("§ 4 Urlaub", level=3)
    doc.add_paragraph("Der Geschäftsführer hat Anspruch auf 30 Werktage Urlaub pro Jahr.")
    doc.add_heading("§ 5 Wettbewerbsverbot", level=3)
    doc.add_paragraph("Während der Dauer des Anstellungsverhältnisses ist dem Geschäftsführer jede Tätigkeit für Konkurrenzunternehmen untersagt.")
    doc.add_heading("§ 6 Vertragsdauer", level=3)
    doc.add_paragraph(f"Der Vertrag wird auf unbestimmte Zeit geschlossen. Er kann mit einer Frist von 12 Monaten zum Jahresende gekündigt werden.")
    doc.add_paragraph(f"\nKöln, den {date_str_val}\n\n_________________________        _________________________\n{C['name']}                         {gf_name}")
    return save_docx(doc, "01_Gesellschaftsrecht", f"GR_{num}_GF_Anstellungsvertrag_{gf_name.replace(' ', '_')}.docx")

# 5. Gesellschafterbeschlüsse
def gen_gesellschafterbeschluss(year, topic, content, num):
    doc = Document()
    make_docx_header(doc, f"Gesellschafterbeschluss", f"{date_str(year, random.randint(3,11), random.randint(1,28))}")
    doc.add_paragraph(f"Die Gesellschafter der {C['name']} haben folgenden Beschluss gefasst:")
    doc.add_heading(topic, level=3)
    doc.add_paragraph(content)
    doc.add_paragraph(f"\nAbstimmungsergebnis: einstimmig angenommen\nStimmen gesamt: 250.000\nJa-Stimmen: 250.000\nNein-Stimmen: 0\nEnthaltungen: 0")
    doc.add_paragraph(f"\nKöln, den {date_str(year, 5, 20)}\n\n_________________________\n{C['gf1_name']}, Versammlungsleiter")
    return save_docx(doc, "01_Gesellschaftsrecht", f"GR_{num}_Gesellschafterbeschluss_{year}_{topic[:20].replace(' ','_').replace('/', '-').replace('(', '').replace(')', '').replace('&', 'und').replace('+', 'plus').replace(',', '')}.docx")

# 6. Prokura-Erteilung
def gen_prokura():
    doc = Document()
    make_docx_header(doc, "Prokura-Erteilung", "Eintragung ins Handelsregister")
    prokuristen = [
        ("Thomas Schneider", "Prokurist", "Leiter Vertrieb"),
        ("Andrea Hoffmann", "Prokuristin", "Leiterin Finanzbuchhaltung"),
    ]
    for p_name, p_title, p_role in prokuristen:
        doc.add_heading(p_name, level=3)
        doc.add_paragraph(
            f"Herr/Frau {p_name} ({p_role}) ist mit Einzelprokura ausgestattet. "
            f"Die Prokura berechtigt zur Vertretung der Gesellschaft in allen Angelegenheiten, "
            f"die der Betrieb eines Handelsgewerbes mit sich bringt, mit Ausnahme der Veräußerung und Belastung von Grundstücken.\n"
            f"Eingetragen im Handelsregister: {C['hrb']}, {C['amtsgericht']}."
        )
    return save_docx(doc, "01_Gesellschaftsrecht", "GR_007_Prokura_Erteilung.docx")

generated.append(gen_gesellschaftsvertrag())
generated.append(gen_hrb_auszug(2022))
generated.append(gen_hrb_auszug(2023))
generated.append(gen_hrb_auszug(2024))
generated.append(gen_gesellschafterliste())
generated.append(gen_gf_anstellungsvertrag(C["gf1_name"], C["gf1_title"], 280_000, "1. April 2010", "004"))
generated.append(gen_gf_anstellungsvertrag(C["gf2_name"], C["gf2_title"], 220_000, "1. September 2015", "005"))
generated.append(gen_gesellschafterbeschluss(2022, "Jahresabschluss 2021", f"Der Jahresabschluss 2021 wird festgestellt. Der Jahresüberschuss in Höhe von {fmt_eur(2_850_000)} wird wie folgt verwendet: Einstellung in die Gewinnrücklage {fmt_eur(1_000_000)}, Ausschüttung an die Gesellschafter {fmt_eur(1_850_000)}.", "008"))
generated.append(gen_gesellschafterbeschluss(2023, "Jahresabschluss 2022", f"Der Jahresabschluss 2022 wird festgestellt. Der Jahresüberschuss in Höhe von {fmt_eur(3_215_000)} wird wie folgt verwendet: Einstellung in die Gewinnrücklage {fmt_eur(1_200_000)}, Ausschüttung an die Gesellschafter {fmt_eur(2_015_000)}.", "009"))
generated.append(gen_gesellschafterbeschluss(2024, "Jahresabschluss 2023", f"Der Jahresabschluss 2023 wird festgestellt. Der Jahresüberschuss in Höhe von {fmt_eur(3_890_000)} wird wie folgt verwendet: Einstellung in die Gewinnrücklage {fmt_eur(1_500_000)}, Ausschüttung an die Gesellschafter {fmt_eur(2_390_000)}.", "010"))
generated.append(gen_prokura())

print(f"  → {len(generated)} docs generated so far")

# ─── 02 FINANZEN ──────────────────────────────────────────────────────────────
print("\n📁 02_Finanzen")

def gen_jahresabschluss_xlsx(year):
    rev = C[f"revenue_{year}"]
    ebit = C[f"ebit_{year}"]
    ebitda = C[f"ebitda_{year}"]
    bs = C[f"bilanzsumme_{year}"]
    ek = C[f"eigenkapital_{year}"]
    prev_year = year - 1
    prev_rev = C.get(f"revenue_{prev_year}", int(rev * 0.9))
    prev_ebit = C.get(f"ebit_{prev_year}", int(ebit * 0.85))

    mat_cost = int(rev * 0.42)
    personal = int(rev * 0.28)
    abschr = ebitda - ebit
    sonstige = rev - mat_cost - personal - abschr - ebit
    umsatz_prev = prev_rev

    guv_rows = [
        ["1.", "Umsatzerlöse", fmt_eur(rev), fmt_eur(umsatz_prev)],
        ["2.", "Materialaufwand", f"-{fmt_eur(mat_cost)}", f"-{fmt_eur(int(umsatz_prev*0.42))}"],
        ["3.", "Personalaufwand", f"-{fmt_eur(personal)}", f"-{fmt_eur(int(umsatz_prev*0.28))}"],
        ["4.", "Abschreibungen", f"-{fmt_eur(abschr)}", f"-{fmt_eur(int(abschr*0.95))}"],
        ["5.", "Sonstige betriebliche Aufwendungen", f"-{fmt_eur(sonstige)}", f"-{fmt_eur(int(sonstige*0.97))}"],
        ["", "EBIT", fmt_eur(ebit), fmt_eur(prev_ebit)],
        ["6.", "Finanzergebnis", f"-{fmt_eur(int(ebit*0.08))}", f"-{fmt_eur(int(prev_ebit*0.09))}"],
        ["", "EBT", fmt_eur(int(ebit*0.92)), fmt_eur(int(prev_ebit*0.91))],
        ["7.", "Steuern vom Einkommen", f"-{fmt_eur(int(ebit*0.27))}", f"-{fmt_eur(int(prev_ebit*0.27))}"],
        ["", "Jahresüberschuss", fmt_eur(int(ebit*0.65)), fmt_eur(int(prev_ebit*0.64))],
    ]

    anl = int(bs * 0.52)
    uml = bs - anl
    fka = int(ek * 0.65)
    bank_forderung = int(uml * 0.35)
    vorr = int(uml * 0.30)

    bilanz_rows = [
        ["AKTIVA", "", "", ""],
        ["A. Anlagevermögen", "", fmt_eur(anl), fmt_eur(int(anl*0.94))],
        ["  Sachanlagen", "", fmt_eur(int(anl*0.72)), ""],
        ["  Immaterielle Vermögenswerte", "", fmt_eur(int(anl*0.08)), ""],
        ["  Finanzanlagen", "", fmt_eur(int(anl*0.20)), ""],
        ["B. Umlaufvermögen", "", fmt_eur(uml), fmt_eur(int(uml*0.92))],
        ["  Vorräte", "", fmt_eur(vorr), ""],
        ["  Forderungen LuL", "", fmt_eur(int(uml*0.28)), ""],
        ["  Bankguthaben", "", fmt_eur(bank_forderung), ""],
        ["  Sonstige", "", fmt_eur(int(uml*0.07)), ""],
        ["Bilanzsumme", "", fmt_eur(bs), fmt_eur(int(bs*0.91))],
        ["", "", "", ""],
        ["PASSIVA", "", "", ""],
        ["A. Eigenkapital", "", fmt_eur(ek), fmt_eur(int(ek*0.84))],
        ["  Stammkapital", "", fmt_eur(C["stammkapital"]), ""],
        ["  Rücklagen", "", fmt_eur(int(ek*0.60)), ""],
        ["  Jahresüberschuss", "", fmt_eur(int(ebit*0.65)), ""],
        ["B. Fremdkapital", "", fmt_eur(bs - ek), fmt_eur(int((bs - ek)*1.1))],
        ["  Bankverbindlichkeiten", "", fmt_eur(fka), ""],
        ["  Verbindlichkeiten LuL", "", fmt_eur(int((bs-ek)*0.35)), ""],
        ["  Rückstellungen", "", fmt_eur(int((bs-ek)*0.20)), ""],
        ["Bilanzsumme", "", fmt_eur(bs), fmt_eur(int(bs*0.91))],
    ]

    return make_xlsx("02_Finanzen", f"FIN_{year}_Jahresabschluss.xlsx",
                     f"Jahresabschluss {year} – {C['name']}",
                     [
                         ("GuV", ["Pos.", "Position", f"IST {year}", f"IST {prev_year}"], guv_rows,
                          [6, 40, 18, 18]),
                         ("Bilanz", ["Position", "", f"IST {year}", f"IST {prev_year}"], bilanz_rows,
                          [30, 6, 18, 18]),
                     ])

def gen_bwa_xlsx(year, month):
    import calendar
    m_names = ["", "Jan", "Feb", "Mrz", "Apr", "Mai", "Jun", "Jul", "Aug", "Sep", "Okt", "Nov", "Dez"]
    rev_m = C[f"revenue_{year}"] // 12
    pers_m = int(rev_m * 0.28)
    mat_m = int(rev_m * 0.42)
    ebit_m = C[f"ebit_{year}"] // 12
    rows = [
        ["Umsatzerlöse", fmt_eur(rev_m), fmt_eur(int(rev_m * (1 + random.uniform(-0.08, 0.08))))],
        ["Materialaufwand", f"-{fmt_eur(mat_m)}", f"-{fmt_eur(int(mat_m * (1+random.uniform(-0.05,0.05))))}"],
        ["Personalaufwand", f"-{fmt_eur(pers_m)}", f"-{fmt_eur(int(pers_m * (1+random.uniform(-0.03,0.03))))}"],
        ["Abschreibungen", f"-{fmt_eur(int(rev_m*0.048))}", f"-{fmt_eur(int(rev_m*0.046))}"],
        ["Sonstige Aufwendungen", f"-{fmt_eur(int(rev_m*0.09))}", f"-{fmt_eur(int(rev_m*0.091))}"],
        ["EBIT", fmt_eur(ebit_m), fmt_eur(int(ebit_m * (1+random.uniform(-0.1, 0.1))))],
    ]
    return make_xlsx("02_Finanzen", f"FIN_BWA_{year}_{month:02d}_{m_names[month]}.xlsx",
                     f"Betriebswirtschaftliche Auswertung {m_names[month]} {year}",
                     [("BWA", ["Position", f"IST {m_names[month]} {year}", f"VJ {m_names[month]} {year-1}"],
                       rows, [35, 18, 18])])

def gen_liquiditaetsplanung_xlsx(year):
    months = ["Jan", "Feb", "Mrz", "Apr", "Mai", "Jun", "Jul", "Aug", "Sep", "Okt", "Nov", "Dez"]
    headers = ["Position"] + months + ["Gesamt"]
    rev_m = C.get(f"revenue_{year}e", C[f"revenue_2023"]) // 12
    rows = []
    liq = 2_400_000
    einnahmen_total = 0
    for m in months:
        r = int(rev_m * (1 + random.uniform(-0.1, 0.12)))
        einnahmen_total += r
    rows.append(["Einzahlungen aus Umsatz"] + [fmt_eur(int(rev_m*(1+random.uniform(-0.1,0.12)))) for _ in months] + [fmt_eur(einnahmen_total)])
    pers_m = int(rev_m * 0.28)
    rows.append(["Personalkosten"] + [f"-{fmt_eur(pers_m)}" for _ in months] + [f"-{fmt_eur(pers_m*12)}"])
    mat_m = int(rev_m * 0.42)
    rows.append(["Materialeinkauf"] + [f"-{fmt_eur(mat_m)}" for _ in months] + [f"-{fmt_eur(mat_m*12)}"])
    rows.append(["Investitionen"] + ([f"-{fmt_eur(int(rev_m*0.06))}" if m in ['Mrz','Jun','Sep','Dez'] else "-" for m in months]) + [f"-{fmt_eur(int(rev_m*0.06*4))}"])
    rows.append(["Liquidität Monatsende"] + [fmt_eur(int(liq * (1+random.uniform(-0.05, 0.08)))) for _ in months] + [""])
    return make_xlsx("02_Finanzen", f"FIN_{year}_Liquiditaetsplanung.xlsx",
                     f"Liquiditätsplanung {year}",
                     [("Liquidität", headers, rows, [28] + [12]*12 + [14])])

def gen_steuerbescheid_pdf(year):
    ebit_y = C.get(f"ebit_{year}", int(C["ebit_2022"] * 0.88))
    tax = int(ebit_y * 0.27)
    kst = int(tax * 0.55)
    gewerbe = tax - kst
    sections = [
        ("Steuerbescheid", f"""
Steuernummer: {C['steuernr']}
USt-IdNr.: {C['ust_id']}
Finanzamt: {C['fa']}
Veranlagungszeitraum: {year}
"""),
        ("Festgesetzte Steuern", [
            ["Steuerart", "Bemessungsgrundlage", "Steuersatz", "Festgesetzte Steuer"],
            ["Körperschaftsteuer", fmt_eur(ebit_y), "15 %", fmt_eur(kst)],
            ["Solidaritätszuschlag", fmt_eur(kst), "5,5 %", fmt_eur(int(kst*0.055))],
            ["Gewerbesteuer", fmt_eur(ebit_y), "ca. 13,2 %", fmt_eur(gewerbe)],
            ["Summe", "", "", fmt_eur(tax)],
        ]),
        ("Zahlungstermine", f"Fällig am 30. September {year+1}.\nBitte überweisen Sie unter Angabe der Steuernummer auf das Konto des {C['fa']}.\n\nBescheid erteilt am: {date_str(year+1, 7, 15)}\n\nRechtsbehelf: Einspruch binnen eines Monats möglich."),
    ]
    return make_pdf("02_Finanzen", f"FIN_{year}_Steuerbescheid.pdf", f"Steuerbescheid {year}", sections)

def gen_kreditvertrag_pdf(lender, amount, rate, term_years, date_y):
    sections = [
        ("Darlehensvertrag", f"""
Darlehensgeber: {lender}
Darlehensnehmer: {C['name']}, {C['full_address']}
Vertretungsberechtigte: {C['gf1_name']}, {C['gf2_name']}
"""),
        ("Konditionen", [
            ["Parameter", "Wert"],
            ["Darlehensbetrag", fmt_eur(amount)],
            ["Nominalzins", f"{rate:.2f} % p.a."],
            ["Laufzeit", f"{term_years} Jahre"],
            ["Tilgung", "Annuitätentilgung, monatlich"],
            ["Monatliche Rate", fmt_eur(int(amount / (term_years * 12) * 1.04))],
            ["Auszahlungsdatum", date_str(date_y, 3, 1)],
            ["Endfälligkeit", date_str(date_y + term_years, 2, 28)],
            ["Sicherheiten", "Grundschuld Industriestraße 12, Köln"],
        ]),
        ("Verwendungszweck", "Das Darlehen dient der Finanzierung von Investitionen in Produktionsanlagen und Fertigungsinfrastruktur."),
        ("Sondertilgung", "Sondertilgungen bis zu 10 % der ursprünglichen Darlehenssumme p.a. sind zinslos möglich."),
    ]
    fname = f"FIN_Kreditvertrag_{lender[:10].replace(' ','_').replace('/', '-').replace('(', '').replace(')', '').replace('&', 'und').replace('+', 'plus').replace(',', '')}_{date_y}.pdf"
    return make_pdf("02_Finanzen", fname, f"Kreditvertrag {lender} {date_y}", sections)

def gen_kontoauszug_xlsx(year, month):
    m_names = ["", "Jan", "Feb", "Mrz", "Apr", "Mai", "Jun", "Jul", "Aug", "Sep", "Okt", "Nov", "Dez"]
    import random
    rows = []
    balance = 1_800_000 + random.randint(0, 600_000)
    transactions = [
        (f"Lohnzahlung {m_names[month]} {year}", -int(C[f"revenue_{year}"]*0.028)),
        (f"Eingang {C['kunde1']}", int(C[f"revenue_{year}"]*0.04)),
        (f"Eingang {C['kunde2']}", int(C[f"revenue_{year}"]*0.03)),
        (f"Material {C['lieferant1']}", -int(C[f"revenue_{year}"]*0.025)),
        (f"Material {C['lieferant2']}", -int(C[f"revenue_{year}"]*0.018)),
        (f"Miete Industriestr. 12", -28_500),
        (f"Versicherungsprämie {C['versicherer1']}", -14_200),
        (f"Tilgung Kredit Deutsche Bank", -int(C.get("revenue_2022", 43_250_000)*0.004)),
        (f"Eingang {C['kunde3']}", int(C[f"revenue_{year}"]*0.022)),
        (f"Sonstige Betriebsausgaben", -int(C[f"revenue_{year}"]*0.008)),
    ]
    day = 1
    for desc, amount in transactions:
        balance += amount
        rows.append([f"{day:02d}.{month:02d}.{year}", desc, fmt_eur(amount) if amount > 0 else "",
                     "" if amount > 0 else fmt_eur(abs(amount)), fmt_eur(balance)])
        day += 3
    return make_xlsx("02_Finanzen", f"FIN_Kontoauszug_{year}_{month:02d}.xlsx",
                     f"Kontoauszug {m_names[month]} {year} – IBAN {C['iban']}",
                     [("Kontoauszug", ["Datum", "Buchungstext", "Haben", "Soll", "Saldo"],
                       rows, [14, 45, 16, 16, 16])])

for y in [2022, 2023]:
    generated.append(gen_jahresabschluss_xlsx(y))
for m in range(1, 13):
    generated.append(gen_bwa_xlsx(2023, m))
generated.append(gen_liquiditaetsplanung_xlsx(2024))
for y in [2021, 2022, 2023]:
    generated.append(gen_steuerbescheid_pdf(y))
generated.append(gen_kreditvertrag_pdf("Deutsche Bank AG", 4_500_000, 3.85, 10, 2019))
generated.append(gen_kreditvertrag_pdf("KfW Bankengruppe", 2_000_000, 1.95, 15, 2021))
for m in [3, 6, 9, 12]:
    generated.append(gen_kontoauszug_xlsx(2023, m))

print(f"  → {len(generated)} docs generated so far")

# ─── 03 PERSONAL / HR ─────────────────────────────────────────────────────────
print("\n📁 03_Personal_HR")

EMPLOYEES = [
    ("Thomas Schneider", "Prokurist / Vertriebsleiter", 95_000, "01.03.2008"),
    ("Andrea Hoffmann", "Prokuristin / Leiterin Finanzbuchhaltung", 85_000, "15.09.2010"),
    ("Michael Weber", "Leiter Produktion", 88_000, "01.05.2007"),
    ("Julia Krause", "Leiterin Qualitätssicherung", 76_000, "01.10.2012"),
    ("Stefan Braun", "Leiter Einkauf", 72_000, "01.07.2014"),
    ("Petra Zimmermann", "HR-Managerin", 65_000, "01.02.2016"),
    ("Markus Fischer", "Key Account Manager", 68_000, "01.04.2015"),
    ("Sabine Koch", "Leiterin Entwicklung / FuE", 92_000, "01.08.2011"),
    ("Jan Müller", "Produktionsleiter Schicht A", 58_000, "01.03.2018"),
    ("Lisa Schulz", "Controllerin", 62_000, "01.11.2019"),
]

def gen_arbeitsvertrag(emp_name, emp_title, salary, start_date, num):
    doc = Document()
    make_docx_header(doc, "Arbeitsvertrag", f"zwischen {C['name']} und {emp_name}")
    doc.add_heading("§ 1 Beginn und Art der Beschäftigung", level=3)
    doc.add_paragraph(f"Das Arbeitsverhältnis beginnt am {start_date}. {emp_name} wird als {emp_title} eingestellt.")
    doc.add_heading("§ 2 Vergütung", level=3)
    doc.add_paragraph(f"Das monatliche Bruttogehalt beträgt {fmt_eur(salary//12)}. Das Jahresbruttogehalt beläuft sich auf {fmt_eur(salary)}.")
    doc.add_heading("§ 3 Arbeitszeit", level=3)
    doc.add_paragraph("Die regelmäßige wöchentliche Arbeitszeit beträgt 40 Stunden. Beginn und Ende der täglichen Arbeitszeit richten sich nach den betrieblichen Erfordernissen.")
    doc.add_heading("§ 4 Urlaub", level=3)
    doc.add_paragraph("Der Arbeitnehmer hat Anspruch auf 30 Arbeitstage Erholungsurlaub im Kalenderjahr.")
    doc.add_heading("§ 5 Kündigung", level=3)
    doc.add_paragraph("Das Arbeitsverhältnis kann von beiden Parteien mit einer Frist von drei Monaten zum Quartalsende gekündigt werden.")
    doc.add_paragraph(f"\nKöln, den {start_date}\n\n_________________________        _________________________\n{C['name']}                    {emp_name}")
    return save_docx(doc, "03_Personal_HR", f"HR_{num:03d}_Arbeitsvertrag_{emp_name.replace(' ','_').replace('/', '-').replace('(', '').replace(')', '').replace('&', 'und').replace('+', 'plus').replace(',', '')}.docx")

def gen_personalstammdaten_xlsx():
    headers = ["Nr.", "Name", "Position", "Abt.", "Eintrittsdatum", "Jahresgehalt (€)", "Beschäftigungsart", "Standort"]
    rows = []
    departments = ["Vertrieb", "Finanzen", "Produktion", "QS", "Einkauf", "HR", "F&E", "IT", "Logistik", "Verwaltung"]
    for i, (name, title, sal, start) in enumerate(EMPLOYEES):
        rows.append([str(i+1), name, title, departments[i % len(departments)], start, fmt_eur(sal), "Vollzeit", C["city"]])
    # Add more rows to represent broader workforce
    for i in range(20):
        rows.append([str(len(EMPLOYEES)+i+1), f"Mitarbeiter {i+1:03d}", "Facharbeiter Produktion",
                     "Produktion", f"0{(i%9)+1}.0{(i%12)+1}.20{10+(i%12)}", fmt_eur(45_000 + i*1_200),
                     "Vollzeit", C["city"]])
    return make_xlsx("03_Personal_HR", "HR_Personalstammdaten_2024.xlsx",
                     "Personalstammdaten (Auswahl) – Stand: Januar 2024",
                     [("Personal", headers, rows, [5, 28, 30, 18, 16, 16, 14, 12])])

def gen_organigramm_pdf():
    sections = [
        ("Organisationsstruktur", f"""
{C['name']}
Stand: Januar 2024

GESCHÄFTSFÜHRUNG
├── {C['gf1_name']} (CEO) – Verantwortlich für: Strategie, Vertrieb, Produktion
└── {C['gf2_name']} (CFO) – Verantwortlich für: Finanzen, HR, Controlling, IT

ZWEITE FÜHRUNGSEBENE (Prokuristen)
├── Thomas Schneider – Vertrieb & Key Account Management
└── Andrea Hoffmann – Finanzbuchhaltung & Controlling

ABTEILUNGEN
├── Vertrieb (12 MA) – Leitung: Thomas Schneider
├── Produktion (148 MA) – Leitung: Michael Weber
│   ├── Schicht A (Leitung: Jan Müller)
│   ├── Schicht B
│   └── Instandhaltung
├── Forschung & Entwicklung (22 MA) – Leitung: Sabine Koch
├── Qualitätssicherung (18 MA) – Leitung: Julia Krause
├── Einkauf/Logistik (14 MA) – Leitung: Stefan Braun
├── Finanzen/Controlling (9 MA) – Leitung: Andrea Hoffmann
├── Human Resources (7 MA) – Leitung: Petra Zimmermann
└── IT/Infrastruktur (5 MA)

Gesamtbelegschaft: {C['employees_2023']} Mitarbeiter (Stand: 31.12.2023)
"""),
    ]
    return make_pdf("03_Personal_HR", "HR_Organigramm_2024.pdf", "Organigramm", sections)

def gen_lohnliste_xlsx(year, month):
    m_names = ["","Jan","Feb","Mrz","Apr","Mai","Jun","Jul","Aug","Sep","Okt","Nov","Dez"]
    headers = ["Nr.", "Name", "Abteilung", "Brutto (€)", "SV-AG-Anteil", "Netto ca. (€)"]
    rows = []
    total_brutto = 0
    for i, (name, title, sal, _) in enumerate(EMPLOYEES):
        brutto_m = sal // 12
        sv = int(brutto_m * 0.195)
        netto = int(brutto_m * 0.63)
        rows.append([str(i+1), name, title[:25], fmt_eur(brutto_m), fmt_eur(sv), fmt_eur(netto)])
        total_brutto += brutto_m
    rows.append(["", "Weitere Mitarbeiter (237 MA)", "Div.", fmt_eur(C[f"revenue_{year}"] // 12 * 28 // 100 - total_brutto), "", ""])
    return make_xlsx("03_Personal_HR", f"HR_Lohnliste_{year}_{month:02d}.xlsx",
                     f"Lohnliste (anonymisiert) {m_names[month]} {year}",
                     [("Löhne", headers, rows, [5, 30, 26, 16, 16, 16])])

def gen_betriebsvereinbarung(topic, content, num):
    doc = Document()
    make_docx_header(doc, f"Betriebsvereinbarung", topic)
    doc.add_heading("Präambel", level=3)
    doc.add_paragraph(f"Die {C['name']} und der Betriebsrat schließen folgende Betriebsvereinbarung ab, um geordnete und faire Arbeitsbedingungen zu gewährleisten.")
    doc.add_heading("Regelungsinhalt", level=3)
    doc.add_paragraph(content)
    doc.add_heading("Geltungsbereich", level=3)
    doc.add_paragraph(f"Diese Betriebsvereinbarung gilt für alle Arbeitnehmer der {C['name']} am Standort {C['city']}.")
    doc.add_heading("Inkrafttreten und Kündigung", level=3)
    doc.add_paragraph("Diese Betriebsvereinbarung tritt mit Unterzeichnung in Kraft und kann mit einer Frist von drei Monaten zum Jahresende von jeder Seite gekündigt werden.")
    doc.add_paragraph(f"\nKöln, den 1. Januar 2023\n\n_________________________        _________________________\n{C['gf1_name']}                   Betriebsratsvorsitzender")
    return save_docx(doc, "03_Personal_HR", f"HR_{num}_BV_{topic[:20].replace(' ','_').replace('/', '-').replace('(', '').replace(')', '').replace('&', 'und').replace('+', 'plus').replace(',', '')}.docx")

for i, (name, title, sal, start) in enumerate(EMPLOYEES):
    generated.append(gen_arbeitsvertrag(name, title, sal, start, i+1))
generated.append(gen_personalstammdaten_xlsx())
generated.append(gen_organigramm_pdf())
generated.append(gen_lohnliste_xlsx(2023, 6))
generated.append(gen_lohnliste_xlsx(2023, 12))
generated.append(gen_betriebsvereinbarung("Arbeitszeit und Schichtmodell",
    "Die Regelarbeitszeit beträgt 40 h/Woche. Im Produktionsbereich wird ein 2-Schicht-Modell gefahren: Frühschicht 06:00–14:30, Spätschicht 14:30–23:00. Überstunden sind zu vergüten oder in Freizeit auszugleichen.", "011"))
generated.append(gen_betriebsvereinbarung("Mobiles Arbeiten (Homeoffice)",
    "Mitarbeitern in Verwaltungs- und Bürotätigkeiten steht nach Abstimmung mit dem Vorgesetzten bis zu 2 Tage Homeoffice pro Woche zu. Für Produktions- und Servicetätigkeiten ist Homeoffice betriebsbedingt nicht möglich.", "012"))
generated.append(gen_betriebsvereinbarung("Nutzung von IT-Systemen",
    f"Alle IT-Systeme des Unternehmens ({C['erp']}, {C['crm']}, Produktionsnetzwerk) sind ausschließlich für dienstliche Zwecke zu nutzen. Private Nutzung ist untersagt. Der Betriebsrat wurde über Monitoring-Maßnahmen informiert.", "013"))

print(f"  → {len(generated)} docs generated so far")

# ─── 04 VERTRÄGE KUNDEN ───────────────────────────────────────────────────────
print("\n📁 04_Vertraege_Kunden")

def gen_kundenvertrag(kunde, value, product, start_y, num):
    doc = Document()
    make_docx_header(doc, "Liefervertrag / Rahmenvertrag",
                     f"zwischen {C['name']} und {kunde}")
    doc.add_heading("§ 1 Vertragsgegenstand", level=3)
    doc.add_paragraph(f"Die {C['name']} (nachfolgend \"Lieferant\") liefert dem Kunden {kunde} (nachfolgend \"Auftraggeber\") die folgende Anlage bzw. das folgende System: {product}. Lieferumfang, technische Spezifikation und Preis sind im Angebot Nr. A-{start_y}-{num:04d} geregelt, das Bestandteil dieses Vertrages ist.")
    doc.add_heading("§ 2 Auftragswert und Zahlungsbedingungen", level=3)
    doc.add_paragraph(f"Der vereinbarte Auftragswert beträgt {fmt_eur(value)} zzgl. der gesetzlichen Umsatzsteuer.\n"
                      f"Zahlungsbedingungen: 30 % Anzahlung bei Auftragserteilung, 40 % nach Lieferung, 30 % nach Abnahme.\n"
                      f"Zahlungsziel: 14 Tage netto.")
    doc.add_heading("§ 3 Liefertermin", level=3)
    doc.add_paragraph(f"Die Lieferung erfolgt innerhalb von 24 Wochen ab Auftragsbestätigung, frühestens jedoch zum {date_str(start_y, 9, 1)}.")
    doc.add_heading("§ 4 Gewährleistung", level=3)
    doc.add_paragraph("Die Gewährleistungsfrist beträgt 24 Monate ab Abnahme. Der Lieferant verpflichtet sich, Mängel binnen 5 Werktagen zu beheben.")
    doc.add_heading("§ 5 Eigentumsvorbehalt", level=3)
    doc.add_paragraph("Die gelieferte Anlage bleibt bis zur vollständigen Bezahlung Eigentum des Lieferanten.")
    doc.add_paragraph(f"\nKöln, den {date_str(start_y, 3, 15)}\n\n_________________________        _________________________\n{C['name']}                    {kunde}")
    return save_docx(doc, "04_Vertraege_Kunden", f"KD_{num:03d}_Vertrag_{kunde[:20].replace(' ','_').replace('/', '-').replace('(', '').replace(')', '').replace('&', 'und').replace('+', 'plus').replace(',', '')}.docx")

def gen_agb_pdf():
    sections = [
        ("§ 1 Geltungsbereich", "Diese Allgemeinen Geschäftsbedingungen (AGB) gelten für alle Lieferungen und Leistungen der " + C["name"] + " (nachfolgend \"Lieferant\"). Sie gelten ausschließlich; entgegenstehende Bedingungen des Kunden werden nicht anerkannt."),
        ("§ 2 Angebot und Vertragsschluss", "Angebote des Lieferanten sind freibleibend. Ein Vertrag kommt erst durch schriftliche Auftragsbestätigung zustande."),
        ("§ 3 Preise und Zahlungsbedingungen", "Alle Preise verstehen sich ab Werk Köln zzgl. Verpackung und gesetzlicher Umsatzsteuer. Rechnungen sind innerhalb von 14 Tagen ohne Abzug zu zahlen."),
        ("§ 4 Lieferung und Gefahrübergang", "Lieferzeitangaben sind unverbindlich. Die Gefahr geht auf den Käufer über, sobald die Ware das Werk verlässt."),
        ("§ 5 Gewährleistung", "Die Gewährleistungsfrist beträgt 24 Monate. Bei Mängeln hat der Lieferant das Recht zur Nacherfüllung."),
        ("§ 6 Haftungsbeschränkung", "Die Haftung des Lieferanten ist auf Vorsatz und grobe Fahrlässigkeit beschränkt, soweit nicht Personenschäden oder zwingend gesetzliche Regelungen betroffen sind."),
        ("§ 7 Gerichtsstand und anwendbares Recht", f"Gerichtsstand ist {C['city']}. Es gilt deutsches Recht unter Ausschluss des UN-Kaufrechts."),
    ]
    return make_pdf("04_Vertraege_Kunden", "KD_000_AGB_Lieferbedingungen.pdf", f"Allgemeine Geschäftsbedingungen (AGB) – {C['name']}", sections)

generated.append(gen_agb_pdf())
generated.append(gen_kundenvertrag(C["kunde1"], 1_850_000, C["produkt1"], 2022, 1))
generated.append(gen_kundenvertrag(C["kunde2"], 980_000, C["produkt3"], 2022, 2))
generated.append(gen_kundenvertrag(C["kunde3"], 2_300_000, C["produkt1"], 2023, 3))
generated.append(gen_kundenvertrag(C["kunde4"], 750_000, C["produkt2"], 2023, 4))
generated.append(gen_kundenvertrag(C["kunde5"], 1_450_000, C["produkt4"], 2023, 5))
generated.append(gen_kundenvertrag(C["kunde1"], 2_100_000, C["produkt3"], 2024, 6))
generated.append(gen_kundenvertrag(C["kunde2"], 1_200_000, C["produkt2"], 2024, 7))

# Angebote
for i, (kunde, val, prod) in enumerate([
    (C["kunde3"], 890_000, C["produkt2"]),
    (C["kunde4"], 3_200_000, C["produkt1"]),
    (C["kunde5"], 1_650_000, C["produkt4"]),
], 1):
    doc = Document()
    make_docx_header(doc, f"Angebot Nr. A-2024-{i:04d}", f"an {kunde}")
    doc.add_heading("Leistungsbeschreibung", level=3)
    doc.add_paragraph(f"Hiermit unterbreiten wir Ihnen unser Angebot für: {prod}\n\nPreis: {fmt_eur(val)} zzgl. 19 % MwSt.\nLieferzeit: 20-24 Wochen ab Auftragserteilung\nGültigkeit: 60 Tage ab Angebotsdatum")
    doc.add_paragraph(f"\n{C['city']}, den {date_str(2024, 2, i*5)}\n\nMit freundlichen Grüßen\n{C['name']}\n{C['gf1_name']}, Geschäftsführer")
    generated.append(save_docx(doc, "04_Vertraege_Kunden", f"KD_A{i:03d}_Angebot_{kunde[:20].replace(' ','_').replace('/', '-').replace('(', '').replace(')', '').replace('&', 'und').replace('+', 'plus').replace(',', '')}.docx"))

print(f"  → {len(generated)} docs generated so far")

# ─── 05 VERTRÄGE LIEFERANTEN ──────────────────────────────────────────────────
print("\n📁 05_Vertraege_Lieferanten")

def gen_lieferantenvertrag(lieferant, product, annual_value, num):
    doc = Document()
    make_docx_header(doc, "Lieferantenrahmenvertrag", f"zwischen {C['name']} und {lieferant}")
    doc.add_heading("§ 1 Gegenstand", level=3)
    doc.add_paragraph(f"Der Lieferant {lieferant} verpflichtet sich, dem Auftraggeber {C['name']} folgende Waren/Leistungen zu liefern: {product}.\nJahresvolumen (Richtwert): {fmt_eur(annual_value)}.")
    doc.add_heading("§ 2 Preise und Preisanpassung", level=3)
    doc.add_paragraph("Die vereinbarten Preise sind in der jeweils gültigen Preisliste festgelegt. Preisanpassungen sind mit einer Frist von 90 Tagen anzukündigen und bedürfen der schriftlichen Zustimmung.")
    doc.add_heading("§ 3 Lieferbedingungen", level=3)
    doc.add_paragraph(f"Lieferung erfolgt DDP {C['full_address']}. Lieferzeiten sind verbindlich. Bei Überschreitung gilt eine Vertragsstrafe von 0,5 % des Auftragswertes pro Woche, max. 5 %.")
    doc.add_heading("§ 4 Qualitätssicherung", level=3)
    doc.add_paragraph("Der Lieferant ist zertifiziert nach ISO 9001:2015. Wareneingangsprüfungen werden stichprobenartig durchgeführt.")
    doc.add_paragraph(f"\nKöln, den {date_str(2022, 1, 15)}\n\n_________________________        _________________________\n{C['name']}                    {lieferant}")
    return save_docx(doc, "05_Vertraege_Lieferanten", f"LF_{num:03d}_Rahmenvertrag_{lieferant[:20].replace(' ','_').replace('/', '-').replace('(', '').replace(')', '').replace('&', 'und').replace('+', 'plus').replace(',', '')}.docx")

generated.append(gen_lieferantenvertrag(C["lieferant1"], "Spanntechnik, Werkzeughalter, CNC-Komponenten", 1_200_000, 1))
generated.append(gen_lieferantenvertrag(C["lieferant2"], "Energieketten, Leitungen, Gleitlager", 480_000, 2))
generated.append(gen_lieferantenvertrag(C["lieferant3"], "Elektromotoren, Frequenzumrichter, Schaltschränke", 2_100_000, 3))
generated.append(gen_lieferantenvertrag(C["lieferant4"], "Laserköpfe, Optiken, Serviceleistungen", 850_000, 4))

# Einkaufsbestellungen
for i, (lief, item, val) in enumerate([
    (C["lieferant1"], "Spanndorne Typ SD-200 (50 Stk.)", 24_500),
    (C["lieferant2"], "Energiekette E4/H (200 m)", 12_800),
    (C["lieferant3"], "Servomotor 1FK7 (10 Stk.)", 48_000),
    (C["lieferant4"], "Laserkopf BIMO CP (2 Stk.)", 32_000),
    (C["lieferant1"], "Werkzeughalter HSK-A63 (100 Stk.)", 18_500),
], 1):
    doc = Document()
    make_docx_header(doc, f"Bestellung Nr. B-2023-{i:04d}", f"an {lief}")
    doc.add_paragraph(f"Wir bestellen hiermit:\n\n{item}\nMenge: siehe Positionsliste\nGesamtpreis: {fmt_eur(val)} zzgl. MwSt.\nLiefertermin: {date_str(2023, i*2, 15)}\nLieferanschrift: {C['full_address']}")
    generated.append(save_docx(doc, "05_Vertraege_Lieferanten", f"LF_B{i:03d}_Bestellung_{lief[:15].replace(' ','_').replace('/', '-').replace('(', '').replace(')', '').replace('&', 'und').replace('+', 'plus').replace(',', '')}.docx"))

print(f"  → {len(generated)} docs generated so far")

# ─── 06 IMMOBILIEN ────────────────────────────────────────────────────────────
print("\n📁 06_Immobilien")

def gen_mietvertrag(objekt, qm, miete_m, mieter_ist_mmb, num):
    doc = Document()
    if mieter_ist_mmb:
        make_docx_header(doc, "Mietvertrag (Mieter)", f"Objekt: {objekt}")
        doc.add_heading("§ 1 Mietsache", level=3)
        doc.add_paragraph(f"Vermieter: Müller Immobilien GbR, Rosenweg 4, 50823 Köln.\nMieter: {C['name']}, {C['full_address']}.\nObjekt: {objekt}, {qm} m² Nutzfläche.")
    else:
        make_docx_header(doc, "Mietvertrag (Vermieter)", f"Objekt: {objekt}")
        doc.add_heading("§ 1 Mietsache", level=3)
        doc.add_paragraph(f"Vermieter: {C['name']}, {C['full_address']}.\nMieter: MV Handels GmbH, Köln.\nObjekt: {objekt}, {qm} m² Nutzfläche.")
    doc.add_heading("§ 2 Mietzins", level=3)
    doc.add_paragraph(f"Die monatliche Miete beträgt {fmt_eur(miete_m)} zzgl. Betriebskosten und gesetzlicher Umsatzsteuer.\nJährliche Miete: {fmt_eur(miete_m*12)}.\nIndexanpassung: gemäß Verbraucherpreisindex, jährlich zum 01.01.")
    doc.add_heading("§ 3 Mietdauer", level=3)
    doc.add_paragraph(f"Der Mietvertrag läuft bis zum 31. Dezember 2030 (Festlaufzeit). Verlängerungsoption: zweimal um je 5 Jahre.")
    doc.add_heading("§ 4 Kaution", level=3)
    doc.add_paragraph(f"Es wurde eine Mietkaution in Höhe von {fmt_eur(miete_m*3)} hinterlegt.")
    return save_docx(doc, "06_Immobilien", f"IMM_{num:03d}_Mietvertrag_{objekt[:20].replace(' ','_').replace('/', '-').replace('(', '').replace(')', '').replace('&', 'und').replace('+', 'plus').replace(',', '')}.docx")

def gen_grundbuchauszug_pdf():
    sections = [
        ("Grundbuchauszug", f"""
Grundbuchamt: Amtsgericht Köln
Grundbuch von: Köln-Ossendorf
Blatt: 4892
"""),
        ("Bestandsverzeichnis", f"""
Lfd. Nr.: 1
Gemarkung: Köln-Ossendorf
Flur: 12, Flurstück: 874
Fläche: 18.500 m²
Lage: Industriestraße 12, 50829 Köln
Nutzungsart: Industrie- und Gewerbefläche
"""),
        ("Abteilung I – Eigentümer", f"""
Eingetragen: Müller Immobilien GbR bestehend aus {C['gf1_name']} und Brigitte Müller, Köln
Erwerbsgrund: Kaufvertrag vom 15.03.1998
"""),
        ("Abteilung II – Lasten und Beschränkungen", """
Keine eintragungspflichtigen Lasten oder Beschränkungen.
"""),
        ("Abteilung III – Grundpfandrechte", f"""
1. Grundschuld: {fmt_eur(3_000_000)} zugunsten Deutsche Bank AG, eingetragen am 01.04.2019.
"""),
        ("Auszug erstellt", f"Datum: {date_str(2024, 1, 10)}\nBeglaubigt: Rechtspflegerin beim Amtsgericht Köln"),
    ]
    return make_pdf("06_Immobilien", "IMM_001_Grundbuchauszug_Industriestrasse_12.pdf", "Grundbuchauszug", sections)

generated.append(gen_mietvertrag("Produktionshalle + Verwaltungsgebäude, Industriestr. 12, 50829 Köln", 12_800, 28_500, True, 1))
generated.append(gen_mietvertrag("Außenlager, Niehler Str. 200, 50733 Köln", 3_200, 6_400, True, 2))
generated.append(gen_grundbuchauszug_pdf())

# Nebenkostenabrechnung
nb_sheet = [
    ("Betriebskosten", ["Position", "2022 (€)", "2023 (€)"],
     [["Grundsteuer", "18.400", "19.200"],
      ["Müllabfuhr", "3.200", "3.400"],
      ["Strom (Gemeinschaftsflächen)", "4.100", "4.600"],
      ["Wasser/Abwasser", "6.800", "7.100"],
      ["Heizung", "22.000", "28.500"],
      ["Hausmeister/Reinigung", "14.400", "15.200"],
      ["Summe", "68.900", "78.000"]],
     [35, 18, 18])
]
generated.append(make_xlsx("06_Immobilien", "IMM_003_Nebenkostenabrechnung_2023.xlsx",
                            "Nebenkostenabrechnung 2023 – Industriestr. 12", nb_sheet))

print(f"  → {len(generated)} docs generated so far")

# ─── 07 IP & LIZENZEN ─────────────────────────────────────────────────────────
print("\n📁 07_IP_Lizenzen")

def gen_patent_pdf(patent_nr, title, filing_year, num):
    sections = [
        ("Patentinformation", f"""
Patentamt: Deutsches Patent- und Markenamt (DPMA)
Patentnummer: DE {patent_nr}
Anmeldetag: {date_str(filing_year, 4, 12)}
Erteilungstag: {date_str(filing_year+2, 8, 5)}
Anmelder/Inhaber: {C['name']}, {C['full_address']}
Erfinder: Sabine Koch, Michael Weber (Arbeitnehmererfindung)
"""),
        ("Titel der Erfindung", title),
        ("Patentanspruch 1 (Hauptanspruch)", f"""
Vorrichtung und Verfahren zur {title.lower()}, gekennzeichnet durch eine neuartige Kombination aus CNC-gesteuerter Präzisionsmechanik und adaptiver Sensorik zur Echtzeitregelung des Fertigungsprozesses, wobei die Steuereinheit mittels maschinellen Lernens kontinuierlich optimiert wird.
"""),
        ("Laufzeit und Jahresgebühren", f"""
Schutzrechtsablauf: {date_str(filing_year+20, 4, 11)}
Status: In Kraft. Jahresgebühren zuletzt entrichtet am {date_str(2024, 3, 1)}.
"""),
    ]
    return make_pdf("07_IP_Lizenzen", f"IP_{num:03d}_Patent_DE{patent_nr}.pdf", f"Patent DE {patent_nr}", sections)

def gen_lizenzvertrag(lizenzgeber, software, jahreslizenz, num):
    doc = Document()
    make_docx_header(doc, "Lizenzvertrag", f"zwischen {C['name']} und {lizenzgeber}")
    doc.add_heading("§ 1 Lizenzgegenstand", level=3)
    doc.add_paragraph(f"Der Lizenzgeber {lizenzgeber} räumt dem Lizenznehmer {C['name']} ein nicht-exklusives, nicht übertragbares Recht zur Nutzung der Software/Lizenz: {software} ein.")
    doc.add_heading("§ 2 Lizenzgebühr", level=3)
    doc.add_paragraph(f"Die jährliche Lizenzgebühr beträgt {fmt_eur(jahreslizenz)} zzgl. MwSt., zahlbar bis zum 31. Januar eines jeden Jahres.")
    doc.add_heading("§ 3 Laufzeit", level=3)
    doc.add_paragraph("Der Vertrag läuft auf unbestimmte Zeit. Die Kündigung ist mit 6 Monaten Frist zum Jahresende möglich.")
    doc.add_paragraph(f"\nKöln, den {date_str(2022, 1, 1)}\n\n_________________________        _________________________\n{C['name']}                    {lizenzgeber}")
    safe = software[:20].replace(' ', '_').replace('/', '-').replace('(', '').replace(')', '')
    return save_docx(doc, "07_IP_Lizenzen", f"IP_{num:03d}_Lizenz_{safe}.docx")

generated.append(gen_patent_pdf("102019123456", "Adaptive Presssteuerung für Hochgeschwindigkeitsstanzprozesse", 2019, 1))
generated.append(gen_patent_pdf("102021098765", "Modulares Transportsystem mit KI-basierter Lastoptimierung", 2021, 2))
generated.append(gen_patent_pdf("102022054321", "Laserscanner-Integrationsmodul für 5-Achs-Bearbeitungszentren", 2022, 3))
generated.append(gen_lizenzvertrag("SAP SE", C["erp"], 185_000, 4))
generated.append(gen_lizenzvertrag("Salesforce.com Inc.", C["crm"], 48_000, 5))
generated.append(gen_lizenzvertrag("Siemens AG", "SINUMERIK ONE CNC-Software (50 Lizenzen)", 72_000, 6))
generated.append(gen_lizenzvertrag("Autodesk Inc.", "AutoCAD Mechanical 2024 (15 Lizenzen)", 28_500, 7))
generated.append(gen_lizenzvertrag("PTC Inc.", "SOLIDWORKS Premium (10 Lizenzen)", 32_000, 8))

print(f"  → {len(generated)} docs generated so far")

# ─── 08 VERSICHERUNGEN ────────────────────────────────────────────────────────
print("\n📁 08_Versicherungen")

def gen_versicherungspolice_pdf(versicherer, art, versumme, praemie_jahr, num):
    sections = [
        ("Versicherungsschein", f"""
Versicherungsnehmer: {C['name']}, {C['full_address']}
Versicherer: {versicherer}
Versicherungsart: {art}
Vertragsnummer: VS-{versicherer[:3].upper()}-2022-{num:04d}
"""),
        ("Versicherungsschutz", [
            ["Parameter", "Wert"],
            ["Versicherungssumme", fmt_eur(versumme)],
            ["Jahresprämie", fmt_eur(praemie_jahr)],
            ["Zahlungsweise", "Jährlich zum 01.01."],
            ["Versicherungsbeginn", date_str(2022, 1, 1)],
            ["Laufzeit", "Bis auf Widerruf"],
            ["Selbstbehalt", fmt_eur(int(versumme * 0.001))],
        ]),
        ("Besondere Bedingungen", f"Alle Risiken am Standort {C['full_address']} sowie bei Lieferung und Montage beim Kunden sind mitversichert. Im Schadensfall bitte unverzüglich unter Angabe der Vertragsnummer melden."),
    ]
    return make_pdf("08_Versicherungen", f"VS_{num:03d}_{art[:20].replace(' ','_').replace('/', '-').replace('(', '').replace(')', '').replace('&', 'und').replace('+', 'plus').replace(',', '')}.pdf", f"Versicherungspolice – {art}", sections)

generated.append(gen_versicherungspolice_pdf(C["versicherer1"], "Betriebshaftpflicht", 10_000_000, 38_500, 1))
generated.append(gen_versicherungspolice_pdf(C["versicherer1"], "Produkthaftpflicht", 15_000_000, 52_000, 2))
generated.append(gen_versicherungspolice_pdf(C["versicherer2"], "Gebäudeversicherung (Industrie)", 25_000_000, 44_200, 3))
generated.append(gen_versicherungspolice_pdf(C["versicherer2"], "Maschinenversicherung", 8_000_000, 28_500, 4))
generated.append(gen_versicherungspolice_pdf(C["versicherer1"], "D&O-Versicherung (Management-Haftpflicht)", 5_000_000, 24_000, 5))
generated.append(gen_versicherungspolice_pdf(C["versicherer2"], "Cyberversicherung", 3_000_000, 18_500, 6))
generated.append(gen_versicherungspolice_pdf(C["versicherer1"], "Transportversicherung (Warenversicherung)", 2_000_000, 12_800, 7))

print(f"  → {len(generated)} docs generated so far")

# ─── 09 COMPLIANCE ────────────────────────────────────────────────────────────
print("\n📁 09_Compliance")

def gen_iso_zertifikat_pdf(norm, bereich, zert_datum, num):
    sections = [
        ("Zertifizierungsstelle", "TÜV Rheinland Cert GmbH\nAm Grauen Stein, 51105 Köln\nAkkreditiert durch DAkkS nach DIN EN ISO/IEC 17021-1"),
        ("Zertifikat", f"""
Hiermit wird bestätigt, dass

{C['name']}
{C['full_address']}

das Managementsystem für den Bereich "{bereich}" die Anforderungen der

{norm}

erfüllt.

Zertifikat-Nummer: TÜV-{num:05d}-{zert_datum.year}
Erstzertifizierung: {date_str(2018, 4, 15)}
Rezertifizierung: {date_str(zert_datum.year, zert_datum.month, zert_datum.day)}
Gültig bis: {date_str(zert_datum.year + 3, zert_datum.month, zert_datum.day - 1)}
"""),
        ("Auditor", f"Leitender Auditor: Dr. Klaus Hartmann\nAuditteam: 3 Auditoren\nAudittage: 4"),
    ]
    return make_pdf("09_Compliance", f"COMP_{num:03d}_Zertifikat_{norm[:10].replace(' ','_').replace('/', '-').replace('(', '').replace(')', '').replace('&', 'und').replace('+', 'plus').replace(',', '')}.pdf",
                    f"Zertifikat {norm}", sections)

def gen_compliance_richtlinie(thema, content, num):
    doc = Document()
    make_docx_header(doc, f"Compliance-Richtlinie", thema)
    doc.add_heading("1. Zweck und Geltungsbereich", level=3)
    doc.add_paragraph(f"Diese Richtlinie gilt für alle Mitarbeiter, Geschäftsführer und beauftragten Dritten der {C['name']}.")
    doc.add_heading("2. Regelungsinhalt", level=3)
    doc.add_paragraph(content)
    doc.add_heading("3. Verantwortlichkeit", level=3)
    doc.add_paragraph(f"Compliance-Verantwortliche: {C['gf2_name']} (CFO)\nAnsprechpartner für Fragen: compliance@halbreiter-maschinenbau.de")
    doc.add_heading("4. Sanktionen", level=3)
    doc.add_paragraph("Verstöße gegen diese Richtlinie können arbeitsrechtliche Konsequenzen bis hin zur fristlosen Kündigung nach sich ziehen.")
    doc.add_paragraph(f"\nKöln, den {date_str(2023, 1, 1)}\n\n{C['gf1_name']}            {C['gf2_name']}\nGeschäftsführer            Geschäftsführerin")
    return save_docx(doc, "09_Compliance", f"COMP_{num:03d}_Richtlinie_{thema[:20].replace(' ','_').replace('/', '-').replace('(', '').replace(')', '').replace('&', 'und').replace('+', 'plus').replace(',', '')}.docx")

generated.append(gen_iso_zertifikat_pdf("DIN EN ISO 9001:2015", "Entwicklung, Herstellung und Vertrieb von Sondermaschinen", date(2023, 5, 15), 1))
generated.append(gen_iso_zertifikat_pdf("DIN EN ISO 14001:2015", "Umweltmanagementsystem, Standort Köln", date(2023, 5, 15), 2))
generated.append(gen_iso_zertifikat_pdf("DIN EN ISO 45001:2018", "Arbeitsschutzmanagementsystem", date(2023, 5, 15), 3))
generated.append(gen_compliance_richtlinie("Antikorruption und Zuwendungen",
    "Mitarbeiter dürfen keine Zuwendungen von Geschäftspartnern annehmen, die über einen Wert von 50 € hinausgehen. Einladungen zu Geschäftsessen sind zulässig, wenn sie dem üblichen Geschäftverkehr entsprechen. Alle Sponsoring-Zahlungen ab 1.000 € bedürfen der Genehmigung durch die Geschäftsführung.", 4))
generated.append(gen_compliance_richtlinie("Datenschutz (DSGVO)",
    f"Die {C['name']} verarbeitet personenbezogene Daten ausschließlich zur Erfüllung von Vertragspflichten und gesetzlichen Anforderungen. Der Datenschutzbeauftragte ist extern bestellt bei: DataCo GmbH, Köln. Alle Mitarbeiter mit Datenzugang werden jährlich geschult.", 5))
generated.append(gen_compliance_richtlinie("Exportkontrolle",
    "Als Hersteller von Dual-Use-Gütern unterliegt das Unternehmen der EU-Exportkontrollverordnung (EU 2021/821). Exporte in Embargländer sind untersagt. Exportkontrollverantwortlicher: Stefan Braun (Einkauf/Logistik).", 6))
generated.append(gen_compliance_richtlinie("Informationssicherheit (ISMS)",
    f"Das ISMS der {C['name']} basiert auf den Anforderungen der ISO 27001. Passwörter müssen mindestens 12 Zeichen umfassen. 2-Faktor-Authentifizierung ist für alle Cloud-Dienste verpflichtend. Datensicherungen erfolgen täglich auf redundanten Systemen.", 7))

# Genehmigungen BImSchG
sections_bimsch = [
    ("Genehmigungsbescheid", f"""
Behörde: Bezirksregierung Köln, Dezernat 53 (Immissionsschutz)
Antragsteller: {C['name']}, {C['full_address']}
Datum der Genehmigung: {date_str(2020, 6, 30)}
"""),
    ("Genehmigungsgegenstand", f"""
Betrieb einer Anlage zur Be- und Verarbeitung von Metallen nach § 4 BImSchG.
Anlage: Hydraulische Stanzpressen, Laserschneider, Lackieranlage
Standort: {C['full_address']}
Genehmigungsnummer: BImSchG-2020-KN-4892
"""),
    ("Auflagen", """
1. Der Schallleistungspegel am Anlagenrand darf 65 dB(A) tagsüber und 50 dB(A) nachts nicht überschreiten.
2. Entstaubungsanlagen sind ordnungsgemäß zu betreiben und jährlich zu warten.
3. Ein Störfallbeauftragter ist zu benennen.
4. Jährlicher Emissionsbericht bis 31. März beim zuständigen Landesamt einzureichen.
"""),
]
generated.append(make_pdf("09_Compliance", "COMP_008_Genehmigung_BImSchG.pdf", "Immissionsschutzrechtliche Genehmigung", sections_bimsch))

print(f"  → {len(generated)} docs generated so far")

# ─── 10 IT & INFRASTRUKTUR ────────────────────────────────────────────────────
print("\n📁 10_IT_Infrastruktur")

def gen_it_uebersicht_pdf():
    sections = [
        ("IT-Infrastruktur Übersicht", f"""
Stand: Januar 2024
IT-Verantwortlicher: Tobias Lange (IT-Leiter)
Mitarbeiter in der IT-Abteilung: 5
"""),
        ("ERP-System", f"""
System: {C['erp']}
Module: FI/CO, MM, PP, SD, QM, PM
Anzahl User: 185 Named Users
Hosting: On-Premise, Rechenzentrum Köln (eigenes Serverrack)
Support-Vertrag: SAP-Wartungsvertrag, {fmt_eur(185_000)}/Jahr
Systemintegrator: Itelligence AG
"""),
        ("CRM-System", f"""
System: {C['crm']}
Module: Opportunities, Accounts, Cases, Reports
Anzahl User: 28 Lizenzen
Hosting: Salesforce Cloud
Jahresgebühr: {fmt_eur(48_000)}
"""),
        ("Produktionsnetzwerk", """
Maschinenanbindung: OPC-UA über Industrial Ethernet (PROFINET)
SPS-Systeme: Siemens SIMATIC S7-1500 (18 Stück)
SCADA: WinCC OA (AVEVA)
Netzwerktrennung: DMZ zwischen OT- und IT-Netzwerk
Backup-System: Veeam, täglich inkrementell, wöchentlich voll
"""),
        ("Hardware", [
            ["Kategorie", "Anzahl", "Hersteller/Modell"],
            ["Server (physisch)", "4", "HPE ProLiant DL380 Gen10"],
            ["Virtualisierung (VMs)", "32", "VMware vSphere 8.0"],
            ["Workstations", "85", "HP EliteDesk / Dell OptiPlex"],
            ["Laptops", "62", "HP EliteBook 840 G9"],
            ["Industrieterminals (HMI)", "24", "Siemens SIMATIC IPC"],
            ["Switches (managed)", "18", "Cisco Catalyst 2960"],
            ["Firewalls", "2", "Fortinet FortiGate 200F"],
        ]),
    ]
    return make_pdf("10_IT_Infrastruktur", "IT_001_Infrastruktur_Uebersicht.pdf", "IT-Infrastruktur Übersicht", sections)

def gen_it_vertraege_xlsx():
    headers = ["Nr.", "Vertragspartner", "Leistung", "Jahresbetrag (€)", "Vertragsbeginn", "Laufzeit bis"]
    rows = [
        ["1", "SAP SE", C["erp"] + " – Wartung & Support", fmt_eur(185_000), "01.01.2022", "31.12.2026"],
        ["2", "Salesforce.com Inc.", C["crm"] + " – Enterprise", fmt_eur(48_000), "01.04.2022", "31.03.2025"],
        ["3", "Microsoft", "Microsoft 365 E3 (250 Lizenzen)", fmt_eur(68_500), "01.01.2023", "31.12.2025"],
        ["4", "Telekom Deutschland", "Glasfaser-Internet 10 Gbit/s, SD-WAN", fmt_eur(42_000), "01.06.2021", "31.05.2026"],
        ["5", "Veeam Software", "Backup & Replication Enterprise", fmt_eur(12_800), "01.01.2023", "31.12.2025"],
        ["6", "Siemens AG", "WinCC OA SCADA-Lizenz", fmt_eur(28_500), "01.01.2022", "31.12.2024"],
        ["7", "Autodesk Inc.", "AutoCAD Mechanical (15 Seats)", fmt_eur(28_500), "01.03.2023", "28.02.2026"],
        ["8", "PTC Inc.", "SOLIDWORKS Premium (10 Seats)", fmt_eur(32_000), "01.03.2022", "28.02.2025"],
        ["9", "Zscaler Inc.", "Zero Trust Network Access", fmt_eur(22_000), "01.07.2023", "30.06.2026"],
        ["10", "CrowdStrike", "Falcon Complete MDR (Endpoint)", fmt_eur(38_500), "01.01.2024", "31.12.2026"],
    ]
    return make_xlsx("10_IT_Infrastruktur", "IT_002_IT_Vertraege_Lizenzen.xlsx",
                     "IT-Verträge und Lizenzen – Stand: Januar 2024",
                     [("IT-Verträge", headers, rows, [5, 28, 40, 18, 16, 16])])

generated.append(gen_it_uebersicht_pdf())
generated.append(gen_it_vertraege_xlsx())

print(f"  → {len(generated)} docs generated so far")

# ─── 11 STRATEGIE & PLANUNG ───────────────────────────────────────────────────
print("\n📁 11_Strategie_Planung")

def gen_businessplan_pdf():
    sections = [
        ("Executive Summary", f"""
{C['name']} ist ein führender Hersteller von Sondermaschinen und Anlagen für die metallverarbeitende Industrie. Das Unternehmen wurde {C['founded']} gegründet und beschäftigt aktuell {C['employees_2023']} Mitarbeiter.

Umsatz 2023: {fmt_eur(C['revenue_2023'])} | EBITDA-Marge: {C['ebitda_2023']/C['revenue_2023']*100:.1f} % | Eigenkapitalquote: {C['eigenkapital_2023']/C['bilanzsumme_2023']*100:.1f} %
"""),
        ("Markt und Wettbewerb", f"""
Der Markt für Sondermaschinen in Deutschland hat ein Volumen von ca. 28 Mrd. €/Jahr (VDMA 2023).
MMB hält einen geschätzten Marktanteil von ca. 0,17 % im adressierten Kundensegment.

Hauptwettbewerber: KUKA AG, Grob-Werke GmbH, Schuler AG

Wettbewerbsvorteile MMB:
• Tiefe Kundenintegration und Individualisierung (>70 % der Projekte sind Unikate)
• Kurze Lieferzeiten (20–24 Wochen vs. Marktdurchschnitt 32 Wochen)
• Hohe Ingenieurskompetenz (F&E-Quote: 4,2 % des Umsatzes)
"""),
        ("Strategische Ziele 2024–2028", f"""
1. Umsatzwachstum auf {fmt_eur(C['revenue_2024e'])} bis Ende 2024, langfristig {fmt_eur(70_000_000)} bis 2028.
2. Internationalisierung: Aufbau Vertriebspräsenz in Polen, Tschechien, Österreich.
3. Digitalisierung: Einführung digitaler Zwillinge für alle Serienprodukte bis 2026.
4. Fachkräftesicherung: Erhöhung der Ausbildungsquote auf 8 % der Belegschaft.
5. Nachhaltigkeit: CO₂-Neutralität am Standort Köln bis 2027 (Photovoltaik + Wärmepumpe).
"""),
        ("Finanzplanung", [
            ["", "2023 IST", "2024e", "2025p", "2026p", "2027p", "2028p"],
            ["Umsatz (Mio. €)", "48,6", "52,1", "56,0", "61,5", "66,0", "70,0"],
            ["EBITDA (Mio. €)", "6,0", "6,4", "7,2", "8,1", "9,0", "9,8"],
            ["EBITDA-Marge", "12,3 %", "12,3 %", "12,9 %", "13,2 %", "13,6 %", "14,0 %"],
            ["Mitarbeiter", "247", "251", "265", "280", "295", "310"],
            ["Investitionen (Mio. €)", "3,2", "4,5", "5,0", "4,0", "3,5", "3,5"],
        ]),
    ]
    return make_pdf("11_Strategie_Planung", "STRAT_001_Businessplan_2024_2028.pdf", "Businessplan 2024–2028", sections)

def gen_investitionsplanung_xlsx():
    headers = ["Projekt", "Kategorie", "Investition (€)", "Zeitraum", "Finanzierung", "Erwarteter ROI"]
    rows = [
        ["Neue Laserschneideanlage LS-800 (3 Stk.)", "Produktionsanlagen", fmt_eur(1_800_000), "Q1-Q2 2024", "Eigenfinanzierung", "18 Monate"],
        ["Erweiterung Halle B (800 m²)", "Immobilien/Bau", fmt_eur(950_000), "Q2-Q4 2024", "KfW-Darlehen", "36 Monate"],
        ["ERP-Upgrade SAP S/4HANA 2024", "IT", fmt_eur(420_000), "Q3 2024", "Eigenfinanzierung", "24 Monate"],
        ["Robotermontagezelle (2 Einheiten)", "Produktion/Automatisierung", fmt_eur(680_000), "Q4 2024", "Leasing", "24 Monate"],
        ["Photovoltaikanlage (250 kWp)", "Nachhaltigkeit/Energie", fmt_eur(320_000), "Q2-Q3 2024", "Eigenfinanzierung/Förderung", "7 Jahre"],
        ["Vertriebsbüro Warschau (Ausbau)", "Internationalisierung", fmt_eur(180_000), "Q3 2024", "Eigenfinanzierung", "24 Monate"],
        ["FuE-Labor Digitale Zwillinge", "Innovation/F&E", fmt_eur(150_000), "Q1 2024", "Eigenfinanzierung", "Strategisch"],
    ]
    return make_xlsx("11_Strategie_Planung", "STRAT_002_Investitionsplanung_2024.xlsx",
                     "Investitionsplanung 2024", [("Investitionen", headers, rows, [40, 25, 16, 18, 22, 14])])

def gen_marktanalyse_pdf():
    sections = [
        ("Marktanalyse – Sondermaschinen Deutschland", f"""
Erstellt: {date_str(2024, 1, 15)}
Verantwortlich: Sabine Koch (F&E), Thomas Schneider (Vertrieb)
"""),
        ("Marktgröße und Wachstum", f"""
Gesamtmarkt Sondermaschinenbau Deutschland 2023: ca. 28,4 Mrd. €
Wachstum 2023 vs. 2022: +3,2 % (VDMA-Hochrechnung)
Prognose 2024: +4,1 % (Wachstumstreiber: Automobilindustrie, Energiewende)

Relevanter Teilmarkt für MMB (hydraulische Pressen, Lasertechnik, Robotik): ca. 6,8 Mrd. €
Marktanteil MMB: ~0,71 %
"""),
        ("Kundenstruktur (Top-5 nach Umsatz 2023)", [
            ["Kunde", "Umsatz MMB 2023", "Anteil am Gesamtumsatz", "Sektor"],
            [C["kunde1"], fmt_eur(8_200_000), "16,9 %", "Stahl/Metall"],
            [C["kunde2"], fmt_eur(6_100_000), "12,5 %", "Automotive Supplier"],
            [C["kunde3"], fmt_eur(5_400_000), "11,1 %", "Automotive Supplier"],
            [C["kunde4"], fmt_eur(4_800_000), "9,9 %", "Heizung/Energie"],
            [C["kunde5"], fmt_eur(3_900_000), "8,0 %", "Automotive"],
            ["Sonstige Kunden (63)", fmt_eur(20_230_000), "41,6 %", "Div."],
            ["Gesamt", fmt_eur(C["revenue_2023"]), "100 %", ""],
        ]),
        ("Technologietrends", """
1. Digitaler Zwilling / Digital Twin: Simulation und Optimierung vor dem Bau
2. Kollaborative Roboter (Cobots): Mensch-Maschine-Zusammenarbeit in der Montage
3. Predictive Maintenance: KI-gestützte Wartungsvorhersage über IoT-Sensoren
4. Grüne Produktion: Energieeffiziente Antriebe, Rückgewinnungssysteme
5. Nearshoring: Rückverlagerung von Produktion aus Asien → Nachfrage nach Investitionsgütern steigt
"""),
    ]
    return make_pdf("11_Strategie_Planung", "STRAT_003_Marktanalyse_2024.pdf", "Marktanalyse Sondermaschinen 2024", sections)

generated.append(gen_businessplan_pdf())
generated.append(gen_investitionsplanung_xlsx())
generated.append(gen_marktanalyse_pdf())

# Management Report Q4 2023
sections_mr = [
    ("Management Report Q4 2023", f"""
Datum: {date_str(2024, 1, 15)}
Erstellt von: {C['gf2_name']} (CFO)
"""),
    ("Finanzkennzahlen Q4 2023", [
        ["KPI", "Q4 2023 IST", "Q4 2022 IST", "Abweichung"],
        ["Umsatz", fmt_eur(C["revenue_2023"]//4 + 1_200_000), fmt_eur(C["revenue_2022"]//4 + 800_000), "+4,8 %"],
        ["EBITDA", fmt_eur(C["ebitda_2023"]//4 + 200_000), fmt_eur(C["ebitda_2022"]//4 + 100_000), "+6,2 %"],
        ["Auftragseingang", fmt_eur(15_800_000), fmt_eur(12_400_000), "+27,4 %"],
        ["Auftragsbestand", fmt_eur(38_200_000), fmt_eur(31_500_000), "+21,3 %"],
        ["Mitarbeiter", str(C["employees_2023"]), str(C["employees_2022"]), f"+{C['employees_2023']-C['employees_2022']}"],
    ]),
    ("Highlights und Risiken", f"""
HIGHLIGHTS:
• Großauftrag {C['kunde3']}: {fmt_eur(2_300_000)} für Pressenlinie PL-500 (Lieferung Q2 2024)
• ISO 9001 Re-Zertifizierung erfolgreich abgeschlossen
• Aufbau Vertriebsbüro Warschau initiiert

RISIKEN:
• Lieferkettenstörungen bei Elektronikkomponenten (Leadtime Siemens SPS: +8 Wochen)
• Fachkräftemangel: 12 offene Stellen im Produktionsbereich
• Energiekosten 2023: +23 % vs. Vorjahr, teilweise über Preisgleitklauseln weitergegeben
"""),
]
generated.append(make_pdf("11_Strategie_Planung", "STRAT_004_Management_Report_Q4_2023.pdf", "Management Report Q4 2023", sections_mr))

print(f"  → {len(generated)} docs generated so far")

# ─── ADDITIONAL FILL-UP DOCS TO REACH 300 ─────────────────────────────────────
print("\n📁 Filling up to 300 documents...")

# Additional BWA months for 2022
for m in range(1, 13):
    generated.append(gen_bwa_xlsx(2022, m))

# More Gesellschafterbeschlüsse
topics_extra = [
    (2021, "Erweiterung Geschäftsführung", f"Der Gesellschafterversammlung wird vorgelegt: Die Bestellung von {C['gf2_name']} zur weiteren Geschäftsführerin mit Einzelvertretungsbefugnis ab 1. September 2015 wird bestätigt."),
    (2020, "Investitionsfreigabe Erweiterungsbau", f"Die Gesellschafterversammlung genehmigt die Investition in den Erweiterungsbau an der Industriestraße 12 mit einem Gesamtvolumen von bis zu {fmt_eur(1_800_000)}."),
    (2019, "Änderung Gesellschaftsvertrag", "Der Gesellschaftsvertrag wird in § 4 dahingehend geändert, dass beide Geschäftsführer Einzelvertretungsbefugnis erhalten."),
    (2018, "Ausschüttungsbeschluss", f"Es wird eine Vorabausschüttung in Höhe von {fmt_eur(1_500_000)} auf die Gesellschafter beschlossen, zahlbar bis 30. November 2018."),
    (2023, "Freigabe Investitionsplanung 2024", f"Die Investitionsplanung 2024 mit einem Gesamtvolumen von {fmt_eur(4_500_000)} wird genehmigt."),
]
for y, topic, content in topics_extra:
    generated.append(gen_gesellschafterbeschluss(y, topic, content, f"{len(generated):03d}"))

# Additional Kontoauszüge
for m in [1, 2, 4, 5, 7, 8, 10, 11]:
    generated.append(gen_kontoauszug_xlsx(2023, m))

# Additional customer quotes / correspondence
for i, (kunde, product, val, yr) in enumerate([
    (C["kunde1"], C["produkt2"], 450_000, 2023),
    (C["kunde2"], C["produkt4"], 680_000, 2023),
    (C["kunde3"], C["produkt1"], 1_900_000, 2024),
    (C["kunde4"], C["produkt3"], 1_100_000, 2024),
    (C["kunde5"], C["produkt2"], 520_000, 2024),
    ("Haver & Boecker OHG", C["produkt1"], 2_200_000, 2024),
    ("Gebr. Heyl GmbH & Co. KG", C["produkt3"], 890_000, 2023),
    ("Maschinenfabrik Niehoff GmbH", C["produkt2"], 740_000, 2023),
], 10):
    doc = Document()
    make_docx_header(doc, f"Auftragsbestätigung Nr. AB-{yr}-{i:04d}", f"an {kunde}")
    doc.add_paragraph(f"Sehr geehrte Damen und Herren,\n\nwir bestätigen Ihren Auftrag vom {date_str(yr, random.randint(1,6), random.randint(1,28))}:\n\nProdukt: {product}\nAuftragswert: {fmt_eur(val)} zzgl. MwSt.\nLiefertermin: {date_str(yr, random.randint(7,12), 1)}\nLieferung DDP an Ihre Anschrift.\n\nMit freundlichen Grüßen\n{C['name']}\n{C['gf1_name']}, Geschäftsführer")
    generated.append(save_docx(doc, "04_Vertraege_Kunden", f"KD_AB{i:03d}_Auftragsbestätigung_{kunde[:20].replace(' ','_').replace('/', '-').replace('(', '').replace(')', '').replace('&', 'und').replace('+', 'plus').replace(',', '')}_{yr}.docx"))

# More HR docs: Abmahnungen, Zeugnisse, Stellenausschreibungen
hr_extras = [
    ("Arbeitszeugnis_Thomas_Schneider", "Thomas Schneider war vom 01.03.2008 bis zum 31.12.2023 als Prokurist und Vertriebsleiter bei uns beschäftigt. Herr Schneider hat seine Aufgaben stets zu unserer vollsten Zufriedenheit erledigt. Wir wünschen ihm für seine Zukunft alles Gute."),
    ("Stellenausschreibung_Produktionsleiter", "Wir suchen zum nächstmöglichen Zeitpunkt einen Produktionsleiter (m/w/d). Anforderungen: abgeschlossenes Studium Maschinenbau, mind. 5 Jahre Führungserfahrung in der Serienfertigung. Wir bieten: attraktives Gehaltspaket, 30 Tage Urlaub, Dienstwagen."),
    ("Stellenausschreibung_FuE_Ingenieur", "Zur Verstärkung unseres F&E-Teams suchen wir einen Entwicklungsingenieur (m/w/d) mit Schwerpunkt Automatisierungstechnik / Robotik. Erfahrung mit SPS-Programmierung (Siemens SIMATIC) erforderlich."),
    ("Betriebsratswahl_2023_Protokoll", f"Protokoll der Betriebsratswahl 2023. Wahltermin: 15. März 2023. Gewählt: 5 Mitglieder. Vorsitzender: Hans-Peter Dietrich. Stellvertreterin: Monika Sauer. Wahlbeteiligung: 84,7 % der {C['employees_2022']} wahlberechtigten Arbeitnehmer."),
    ("Pensionszusage_Klaus_Mueller", f"Herrn {C['gf1_name']} wird eine betriebliche Altersversorgung in Form einer Direktzusage gewährt. Renteneintritt mit 65 Jahren. Jährliche Rentenleistung: {fmt_eur(72_000)} brutto. Witwenrente: 60 % der Hauptrente."),
]
for fname, content in hr_extras:
    doc = Document()
    make_docx_header(doc, fname.replace("_", " "))
    doc.add_paragraph(content)
    generated.append(save_docx(doc, "03_Personal_HR", f"HR_{fname}.docx"))

# FuE / R&D documents
fuer_docs = [
    ("FuE_Bericht_2023", "Forschungs- und Entwicklungsbericht 2023",
     f"F&E-Aufwendungen 2023: {fmt_eur(C['revenue_2023']*0.042)} (4,2 % des Umsatzes)\n\nProjekte:\n1. Digitaler Zwilling für Pressenlinie PL-500: Abschluss Q4 2023, Rollout 2024\n2. KI-gestützte Qualitätsprüfung (Bildverarbeitung): Phase II läuft\n3. Kollaborationsroboter-Integration MR-150: Produktreife Q2 2024\n4. Energieoptimierung hydraulische Systeme: Einsparpotenzial identifiziert: 18 %"),
    ("FuE_Antrag_Foerderung_BMWi", "Förderantrag BMWK – ZIM-Kooperationsprojekt",
     "Antragsteller: Halbreiter Maschinenbau GmbH\nFörderprogramm: Zentrales Innovationsprogramm Mittelstand (ZIM)\nProjekttitel: SmartPress 4.0 – KI-gestützte Adaptivregelung für Präzisionsstanzprozesse\nBeantragtes Fördervolumen: 285.000 €\nGesamtprojektkosten: 570.000 €\nLaufzeit: 24 Monate\nKooperationspartner: RWTH Aachen, Institut für Fertigungstechnik"),
]
for fname, title, content in fuer_docs:
    doc = Document()
    make_docx_header(doc, title)
    doc.add_paragraph(content)
    generated.append(save_docx(doc, "09_Compliance", f"COMP_{fname}.docx"))

# Additional financial docs: planning models
def gen_5year_plan_xlsx():
    years = ["2024e", "2025p", "2026p", "2027p", "2028p"]
    rev_plan = [52_100_000, 56_000_000, 61_500_000, 66_000_000, 70_000_000]
    ebitda_plan = [6_400_000, 7_200_000, 8_100_000, 9_000_000, 9_800_000]
    ebit_plan = [4_200_000, 4_900_000, 5_600_000, 6_300_000, 7_000_000]
    headers = ["KPI"] + years
    rows = [
        ["Umsatz"] + [fmt_eur(r) for r in rev_plan],
        ["Umsatzwachstum"] + ["+7,1 %", "+7,5 %", "+9,8 %", "+7,3 %", "+6,1 %"],
        ["EBITDA"] + [fmt_eur(e) for e in ebitda_plan],
        ["EBITDA-Marge"] + [f"{e/r*100:.1f} %" for e, r in zip(ebitda_plan, rev_plan)],
        ["EBIT"] + [fmt_eur(e) for e in ebit_plan],
        ["Mitarbeiter"] + ["251", "265", "280", "295", "310"],
        ["Investitionen"] + [fmt_eur(v) for v in [4_500_000, 5_000_000, 4_000_000, 3_500_000, 3_500_000]],
        ["Cashflow aus Betrieb"] + [fmt_eur(v) for v in [5_800_000, 6_500_000, 7_400_000, 8_200_000, 8_900_000]],
    ]
    return make_xlsx("11_Strategie_Planung", "STRAT_005_5Jahresplan_Finanzmodell.xlsx",
                     "Mittelfristiger Finanzplan 2024–2028",
                     [("Finanzplan", headers, rows, [30] + [16]*5)])

generated.append(gen_5year_plan_xlsx())

# Kreditlinie / Kontokorrent
sections_kk = [
    ("Kontokorrentkreditvertrag", f"""
Kreditgeber: {C['bank']}
Kreditnehmer: {C['name']}, {C['full_address']}
IBAN: {C['iban']}, BIC: {C['bic']}
"""),
    ("Konditionen", [
        ["Merkmal", "Vereinbarung"],
        ["Kreditlinie", fmt_eur(3_000_000)],
        ["Zinssatz", "EURIBOR 3M + 1,85 % p.a."],
        ["Bereitstellungsprovision", "0,25 % p.a. auf nicht in Anspruch genommenen Betrag"],
        ["Überziehungsrahmen", fmt_eur(500_000)],
        ["Überziehungszins", "5,5 % p.a."],
        ["Gültigkeit", f"Bis zum {date_str(2026, 12, 31)}, automatische Verlängerung um 12 Monate"],
    ]),
]
generated.append(make_pdf("02_Finanzen", "FIN_Kontokorrentvertrag_DeutscheBank.pdf", "Kontokorrentkreditvertrag", sections_kk))

# Additional strategic memos
strat_memos = [
    ("SWOT-Analyse 2024", f"""
STÄRKEN:
• Hohe technische Kompetenz und langjährige Kundenbeziehungen (Ø >8 Jahre)
• Flexibilität bei Sonderlösungen, kurze Reaktionszeiten
• Solide Finanzbasis (EK-Quote {C['eigenkapital_2023']/C['bilanzsumme_2023']*100:.1f} %, schuldenfrei im operativen Geschäft)
• Patentportfolio mit 3 aktiven Schutzrechten

SCHWÄCHEN:
• Hohe Abhängigkeit von Top-5-Kunden ({C['kunde1']}, {C['kunde2']}, {C['kunde3']} = ca. 40 % Umsatz)
• Begrenzte Internationalisierung (>95 % DACH-Markt)
• Fachkräftemangel in der Produktion

CHANCEN:
• Nearshoring-Trend: Europäische Produktionsverlagerungen
• Digitalisierungsförderung (KfW, BMWK-Förderprogramme)
• Wachstum Elektromobilität → neue Maschinennachfrage

RISIKEN:
• Konjunkturabhängigkeit im Automobilsektor (ca. 35 % des Umsatzes)
• Steigender Wettbewerbsdruck aus Osteuropa und China
• Energiekostensteigerungen (operativer Hebel: ca. 3,8 %)
"""),
]
for title, content in strat_memos:
    doc = Document()
    make_docx_header(doc, title)
    doc.add_paragraph(content)
    generated.append(save_docx(doc, "11_Strategie_Planung", f"STRAT_{title[:30].replace(' ','_').replace('/', '-').replace('(', '').replace(')', '').replace('&', 'und').replace('+', 'plus').replace(',', '')}.docx"))

# Fill remaining to 300 with realistic operational docs
remaining = 300 - len(generated)
print(f"\n  Need {remaining} more documents to reach 300...")

# Generate more operational docs
op_docs = []

# Monthly management reports 2022 and 2023
for yr in [2022, 2023]:
    for q in [1, 2, 3]:
        sections = [
            (f"Management Report Q{q} {yr}", f"Datum: {date_str(yr, q*3, 15)}\nErstellt: {C['gf2_name']} (CFO)"),
            ("Finanzkennzahlen", [
                ["KPI", f"Q{q} {yr} IST", f"Q{q} {yr-1} IST", "Δ"],
                ["Umsatz", fmt_eur(C[f"revenue_{yr}"]//4), fmt_eur(C.get(f"revenue_{yr-1}", int(C[f"revenue_{yr}"]*0.89))//4), "+5,2 %"],
                ["EBITDA", fmt_eur(C[f"ebitda_{yr}"]//4), fmt_eur(C.get(f"ebitda_{yr-1}", int(C[f"ebitda_{yr}"]*0.86))//4), "+7,1 %"],
                ["Auftragseingang", fmt_eur(int(C[f"revenue_{yr}"]*0.28)), fmt_eur(int(C.get(f"revenue_{yr-1}", int(C[f"revenue_{yr}"]*0.89))*0.26)), "+8,4 %"],
            ]),
        ]
        generated.append(make_pdf("11_Strategie_Planung", f"STRAT_MgmtReport_Q{q}_{yr}.pdf", f"Management Report Q{q} {yr}", sections))

# More HR: salary review, training records
training_docs = [
    ("Schulungsplan 2024", f"Geplante Schulungen 2024 für {C['employees_2024e']} Mitarbeiter:\n\n• ISO 9001 Awareness (alle MA): März 2024\n• Arbeitssicherheit BG Holz+Metall (Produktion): April 2024\n• SAP S/4HANA Key User (FI/CO): Mai 2024\n• Exportkontrolle (Vertrieb, Einkauf): Juni 2024\n• Führungskräfteseminar (10 TN): September 2024\n• Datenschutz DSGVO (alle MA): Oktober 2024\nGesamtbudget: {fmt_eur(85_000)}"),
    ("Gehaltsrunde 2024 Protokoll", f"Die Gehaltsrunde 2024 wurde zwischen der Geschäftsführung ({C['gf2_name']}) und dem Betriebsrat verhandelt.\nErgebnis: Generelle Gehaltserhöhung von 3,5 % für alle Tarifmitarbeiter ab 01.04.2024.\nLeistungsorientierte Boni für außertarifliche Mitarbeiter: Auszahlung im März 2024.\nGesamte Mehrbelastung Personalkosten: {fmt_eur(C['revenue_2023']*0.28*0.035)}/Jahr."),
    ("Fluktuationsanalyse 2023", f"Fluktuationsquote 2023: 6,8 % (Marktdurchschnitt Maschinenbau: 8,2 %)\nEintritte 2023: 34\nAustritte 2023: 18 (davon 4 Renteneintritte, 9 eigene Kündigung, 5 betriebsbedingt)\nDurchschnittliche Betriebszugehörigkeit: 9,4 Jahre\nOffene Stellen (31.12.2023): 12"),
]
for title, content in training_docs:
    doc = Document()
    make_docx_header(doc, title)
    doc.add_paragraph(content)
    generated.append(save_docx(doc, "03_Personal_HR", f"HR_{title[:25].replace(' ','_').replace('/', '-').replace('(', '').replace(')', '').replace('&', 'und').replace('+', 'plus').replace(',', '')}.docx"))

# More compliance: DSGVO Verarbeitungsverzeichnis
sections_vvt = [
    ("Verzeichnis von Verarbeitungstätigkeiten (Art. 30 DSGVO)", f"Verantwortlicher: {C['name']}, {C['full_address']}\nDSB: DataCo GmbH (extern)\nStand: {date_str(2024, 1, 1)}"),
    ("Verarbeitungstätigkeiten (Auswahl)", [
        ["Nr.", "Bezeichnung", "Zweck", "Betroffene", "Speicherdauer"],
        ["1", "Personalverwaltung", "Durchführung Arbeitsverhältnis", "Mitarbeiter", "10 Jahre nach Austritt"],
        ["2", "Lohnbuchhaltung", "Entgeltabrechnung", "Mitarbeiter", "10 Jahre (§ 257 HGB)"],
        ["3", "Kundenverwaltung (CRM)", "Vertrieb, Vertragsabwicklung", "Kunden", "3 Jahre nach Vertragsende"],
        ["4", "Lieferantenverwaltung", "Einkauf, Vertragsabwicklung", "Lieferanten", "6 Jahre nach Vertragsende"],
        ["5", "Videoüberwachung", "Werksschutz", "Besucher, MA", "7 Tage rollierende Speicherung"],
        ["6", "Website-Nutzerdaten", "Webanalyse (Matomo)", "Websitebesucher", "90 Tage"],
    ]),
]
generated.append(make_pdf("09_Compliance", "COMP_DSGVO_Verarbeitungsverzeichnis.pdf", "DSGVO Verarbeitungsverzeichnis", sections_vvt))

# Additional contracts: maintenance contracts
for i, (partner, scope, val) in enumerate([
    ("Siemens Industry Services GmbH", "Wartungsvertrag SPS und Automatisierungstechnik, alle Siemens-Komponenten am Standort Köln", 85_000),
    ("SAP SE Maintenance", "SAP S/4HANA Enterprise Support, Release 2023", 185_000),
    ("Trumpf Service GmbH", "Full-Service-Wartungsvertrag Laserschneider LS-800 (3 Anlagen)", 72_000),
    ("Schindler Deutschland AG", "Wartungsvertrag Aufzüge und Hebebühnen, Industriestr. 12", 12_400),
    ("Bruni Elektroanlagen GmbH", "Elektroprüfung (BGV A3) und Wartungsvertrag Schaltanlagen", 28_500),
], 1):
    doc = Document()
    make_docx_header(doc, f"Wartungsvertrag / Service-Level-Agreement", f"mit {partner}")
    doc.add_heading("§ 1 Leistungsumfang", level=3)
    doc.add_paragraph(f"Der Auftragnehmer {partner} erbringt folgende Wartungs- und Serviceleistungen für {C['name']}: {scope}.")
    doc.add_heading("§ 2 Vergütung", level=3)
    doc.add_paragraph(f"Jahrespauschale: {fmt_eur(val)} zzgl. MwSt., zahlbar in vier gleichen Quartalsraten.")
    doc.add_heading("§ 3 Reaktionszeiten", level=3)
    doc.add_paragraph("Bei Störungen: Reaktion innerhalb 4 Stunden, Behebung innerhalb 48 Stunden (Werktage).\nPräventive Wartung: mindestens 2× jährlich nach Wartungsplan.")
    generated.append(save_docx(doc, "05_Vertraege_Lieferanten", f"LF_SLA{i:03d}_Wartung_{partner[:20].replace(' ','_').replace('/', '-').replace('(', '').replace(')', '').replace('&', 'und').replace('+', 'plus').replace(',', '')}.docx"))

# Monthly P&L summary 2022
for m in range(1, 13):
    if len(generated) < 300:
        generated.append(gen_bwa_xlsx(2022, m))

# Fill to exactly 300 if needed
counter = len(generated)
extra_topics = [
    ("Datenschutzerklärung Website", "09_Compliance", "COMP"),
    ("Notfallplan Produktionsausfall", "09_Compliance", "COMP"),
    ("Lieferantenbewertung 2023", "05_Vertraege_Lieferanten", "LF"),
    ("Qualitätsprüfbericht Q4 2023", "09_Compliance", "COMP"),
    ("Energieverbrauchsanalyse 2023", "11_Strategie_Planung", "STRAT"),
    ("CO2-Bilanz 2023 (Corporate Carbon Footprint)", "09_Compliance", "COMP"),
    ("Lieferantenaudit Schunk 2023", "05_Vertraege_Lieferanten", "LF"),
    ("Kundenreklamationsstatistik 2023", "04_Vertraege_Kunden", "KD"),
    ("Inventurliste Vorräte 31.12.2023", "02_Finanzen", "FIN"),
    ("Anlagenverzeichnis 2023", "02_Finanzen", "FIN"),
    ("Rückstellungsübersicht 2023", "02_Finanzen", "FIN"),
    ("Bürgschaft Deutsche Bank", "02_Finanzen", "FIN"),
    ("Rahmenvertrag Leasing Trumpf", "05_Vertraege_Lieferanten", "LF"),
    ("Abfallentsorgungsnachweis 2023", "09_Compliance", "COMP"),
    ("Betriebsbegehungsprotokoll BG 2023", "09_Compliance", "COMP"),
    ("Einheitliche Europäische Eigenerklärung (EEE)", "09_Compliance", "COMP"),
    ("Exportkontrollbericht 2023", "09_Compliance", "COMP"),
    ("Maßnahmenplan ISO 9001 Audit 2023", "09_Compliance", "COMP"),
    ("IT-Notfallhandbuch", "10_IT_Infrastruktur", "IT"),
    ("Penetrationstest-Bericht 2023", "10_IT_Infrastruktur", "IT"),
]

for topic, folder, prefix in extra_topics:
    if len(generated) >= 300:
        break
    doc = Document()
    make_docx_header(doc, topic)
    doc.add_paragraph(f"Dokument: {topic}\nUnternehmen: {C['name']}\nDatum: {date_str(2024, 1, 15)}\n\nDieses Dokument enthält die relevanten Informationen zu: {topic}.\n\nVerantwortlich: {C['gf2_name']}, {C['name']}, {C['full_address']}.\n\nBei Fragen wenden Sie sich bitte an: {C['email']}")
    generated.append(save_docx(doc, folder, f"{prefix}_{topic[:30].replace(' ','_').replace('/', '-').replace('(', '').replace(')', '').replace('&', 'und').replace('+', 'plus').replace(',', '')}.docx"))

# ─── WAVE 2: Additional docs to reach 300 ────────────────────────────────────
print(f"\n📁 Wave 2: Adding more docs (currently {len(generated)})...")

# More customer correspondence: Rechnungen
for i, (kunde, val, inv_yr, inv_m) in enumerate([
    (C["kunde1"], 1_850_000, 2023, 3), (C["kunde2"], 980_000, 2023, 4),
    (C["kunde3"], 2_300_000, 2023, 8), (C["kunde4"], 750_000, 2023, 10),
    (C["kunde5"], 1_450_000, 2023, 11), (C["kunde1"], 2_100_000, 2024, 2),
    (C["kunde2"], 1_200_000, 2024, 3), ("Haver & Boecker OHG", 2_200_000, 2024, 4),
    ("Gebr. Heyl GmbH & Co. KG", 890_000, 2023, 7),
    ("Maschinenfabrik Niehoff GmbH", 740_000, 2023, 9),
], 1):
    if len(generated) >= 300: break
    doc = Document()
    make_docx_header(doc, f"Rechnung Nr. RE-{inv_yr}-{i:05d}", f"an {kunde}")
    doc.add_heading("Rechnungsposition", level=3)
    doc.add_paragraph(
        f"Auftraggeber: {kunde}\nAuftragsnummer: A-{inv_yr}-{i:04d}\n\n"
        f"Pos. 1: Sondermaschine gemäß Spezifikation Angebot Nr. A-{inv_yr}-{i:04d}\n"
        f"Nettobetrag: {fmt_eur(val)}\nUmsatzsteuer (19 %): {fmt_eur(int(val*0.19))}\n"
        f"Rechnungsbetrag: {fmt_eur(int(val*1.19))}\n\n"
        f"Zahlungsziel: 14 Tage netto\nIBAN: {C['iban']}, BIC: {C['bic']}\n"
        f"Bitte angeben: RE-{inv_yr}-{i:05d}"
    )
    generated.append(save_docx(doc, "04_Vertraege_Kunden", f"KD_RE{i:03d}_Rechnung_{kunde[:18].replace(' ','_').replace('/', '-').replace('&','und')}.docx"))

# Lieferantenseitige Rechnungen / Eingangsrechnungen
lieferant_rechnungen = [
    (C["lieferant1"], 185_000, 2023, 2), (C["lieferant2"], 96_000, 2023, 3),
    (C["lieferant3"], 420_000, 2023, 4), (C["lieferant4"], 72_000, 2023, 6),
    (C["lieferant1"], 210_000, 2023, 8), (C["lieferant3"], 380_000, 2023, 10),
    (C["lieferant2"], 88_000, 2024, 1), (C["lieferant4"], 65_000, 2024, 2),
]
for i, (lief, val, yr, m) in enumerate(lieferant_rechnungen, 1):
    if len(generated) >= 300: break
    doc = Document()
    make_docx_header(doc, f"Eingangsrechnung / Buchungsbeleg", f"von {lief}")
    doc.add_paragraph(
        f"Lieferant: {lief}\nRechnungsdatum: {date_str(yr, m, random.randint(5,25))}\n"
        f"Bestellbezug: B-{yr}-{i:04d}\n\n"
        f"Lieferung gemäß Bestellung\nNettobetrag: {fmt_eur(val)}\n"
        f"MwSt. 19 %: {fmt_eur(int(val*0.19))}\nBruttobetrag: {fmt_eur(int(val*1.19))}\n\n"
        f"Gebucht auf Kostenstelle: 5100 (Materialeinkauf)"
    )
    generated.append(save_docx(doc, "05_Vertraege_Lieferanten", f"LF_ER{i:03d}_Eingangsrechnung_{lief[:18].replace(' ','_').replace('/', '-').replace('&','und').replace('+','plus')}.docx"))

# Immobilien: Gutachten, Grundrisse, Standortanalyse
immo_extra = [
    ("Verkehrswertgutachten_Industriestr_12", f"Gutachten über den Verkehrswert des Grundstücks Industriestraße 12, 50829 Köln.\n\nGutachter: Dipl.-Ing. Martin Kessler, Öffentlich bestellter und vereidigter Sachverständiger für Grundstücksbewertung, Köln.\nAuftraggeberin: {C['name']}\nBewertungsstichtag: 31. Dezember 2023\n\nVerkehrswert (gerundet): 8.200.000 €\nMethode: Sachwertverfahren in Kombination mit Ertragswertverfahren\nGrundstücksfläche: 18.500 m²\nGebäudefläche (BGF): 12.800 m²\nBaujahr Hauptgebäude: 1987 (grundsaniert 2018)"),
    ("Grundriss_Produktionshalle_A", "Grundrissplan Produktionshalle A, Industriestraße 12, 50829 Köln.\nMaßstab 1:200. Fläche: 5.400 m². Hallenhöhe: 8,5 m. Kranenanlage: 2 × 12,5 t Brückenkran.\nBaujahr: 1987. Letzte Renovierung: 2019."),
    ("Grundriss_Verwaltungsgebäude", "Grundrissplan Verwaltungsgebäude, Industriestraße 12, 50829 Köln.\nMaßstab 1:100. Fläche gesamt: 1.850 m² (3 Stockwerke). Büroflächen, Besprechungsräume, Kantine.\nBaujahr: 2002."),
    ("Mietvertrag_Kundenparkplatz_2022", f"Mietvertrag über 15 Außenparkplätze an der Niehler Straße, 50733 Köln.\nVermieter: Stadtwerke Köln GmbH.\nMieter: {C['name']}.\nMiete: {fmt_eur(3_600)}/Monat.\nLaufzeit: 01.01.2022 – 31.12.2026."),
]
for title, content in immo_extra:
    if len(generated) >= 300: break
    doc = Document()
    make_docx_header(doc, title.replace("_", " "))
    doc.add_paragraph(content)
    generated.append(save_docx(doc, "06_Immobilien", f"IMM_{title[:35].replace(' ','_')}.docx"))

# IT: Sicherheitsberichte, ISMS-Dokumente
it_extra = [
    ("IT-Sicherheitskonzept 2024", f"Informationssicherheitskonzept gemäß ISO/IEC 27001:2022.\n{C['name']}, {C['full_address']}.\nErstellt: {date_str(2024, 1, 15)}\nGültig bis: {date_str(2025, 1, 14)}\n\nSchutzklasse: INTERN\nVerantwortlich: Tobias Lange (IT-Leiter)\n\nRisikoanalyse: 3 kritische, 8 mittlere, 14 niedrige Risiken identifiziert.\nMaßnahmenplan: 18 Maßnahmen definiert, davon 12 bereits umgesetzt."),
    ("Backup-Restore-Konzept", f"Backup-Strategie {C['name']}\n\nVoll-Backup: Sonntags 22:00 Uhr (Veeam)\nInkrementelles Backup: Mo–Sa 22:00 Uhr\nAufbewahrung: 30 Tage lokal, 90 Tage im gesicherten RZ-Tresor (Iron Mountain, Köln)\nRPO: 24 Stunden | RTO: 8 Stunden\nLetzter Restore-Test: {date_str(2023, 11, 14)} – erfolgreich"),
    ("Netzwerkdokumentation 2024", f"Netzwerkinfrastruktur {C['name']}, Stand: Januar 2024.\n\nIP-Segmentierung:\n• 10.0.1.0/24 – Server-VLAN (AD, ERP, File)\n• 10.0.2.0/24 – Client-VLAN (Büro)\n• 10.0.10.0/24 – OT-Netzwerk (Maschinen)\n• 192.168.50.0/24 – Management-VLAN\n\nFirewall-Regeln: 142 aktive Regeln (Fortinet FortiGate 200F)\nInternet-Uplink: Deutsche Telekom, 10 Gbit/s"),
    ("ISMS-Richtlinie Passwortmanagement", f"Passwort-Policy {C['name']}, Version 2.1, gültig ab {date_str(2023, 7, 1)}.\n\nAnforderungen:\n• Mindestlänge: 12 Zeichen\n• Komplexität: Groß-, Kleinbuchstaben, Zahlen, Sonderzeichen\n• Ablauf: 90 Tage für Standard-Accounts, 180 Tage für Service-Accounts\n• 2FA: Pflicht für alle Cloud-Dienste und Remote Access (Cisco AnyConnect)\n• Passwort-Manager: Bitwarden Enterprise für alle Mitarbeiter"),
    ("Software-Inventar 2024", f"Software-Inventar {C['name']}, Stand: 01.01.2024.\n\nBetriebssysteme: Windows 11 Pro (147 Lizenzen), Windows 10 Pro (0 – Migration abgeschlossen)\nOffice-Produktivität: Microsoft 365 E3 (250 Lizenzen)\nERP: {C['erp']} (185 Named User)\nCRM: {C['crm']} (28 Lizenzen)\nAntivirus/EDR: CrowdStrike Falcon (247 Agenten)\nGesamte IT-Kosten p.a.: {fmt_eur(680_000)} (Lizenzen) + {fmt_eur(320_000)} (Hardware-Abschreibungen)"),
]
for title, content in it_extra:
    if len(generated) >= 300: break
    doc = Document()
    make_docx_header(doc, title)
    doc.add_paragraph(content)
    generated.append(save_docx(doc, "10_IT_Infrastruktur", f"IT_{title[:30].replace(' ','_').replace('/', '-')}.docx"))

# More financial: Quartalsbericht 2024, Kostenstellenrechnung
fin_extra = [
    ("FIN_Kostenstellenrechnung_2023", "02_Finanzen",
     "Kostenstellenrechnung 2023",
     [["Kostenstelle", "Bezeichnung", "Kosten 2023", "Budget 2023", "Abweichung"],
      ["1000", "Geschäftsführung", fmt_eur(850_000), fmt_eur(900_000), "-5,6 %"],
      ["2000", "Vertrieb", fmt_eur(2_800_000), fmt_eur(2_600_000), "+7,7 %"],
      ["3000", "Produktion", fmt_eur(18_400_000), fmt_eur(18_000_000), "+2,2 %"],
      ["4000", "Forschung & Entwicklung", fmt_eur(2_040_000), fmt_eur(2_100_000), "-2,9 %"],
      ["5000", "Qualitätssicherung", fmt_eur(1_200_000), fmt_eur(1_150_000), "+4,3 %"],
      ["6000", "Einkauf/Logistik", fmt_eur(980_000), fmt_eur(950_000), "+3,2 %"],
      ["7000", "IT", fmt_eur(780_000), fmt_eur(800_000), "-2,5 %"],
      ["8000", "HR/Verwaltung", fmt_eur(620_000), fmt_eur(600_000), "+3,3 %"],
      ["9000", "Finanzen/Controlling", fmt_eur(480_000), fmt_eur(500_000), "-4,0 %"],
      ["Gesamt", "", fmt_eur(C["revenue_2023"]*0.28), fmt_eur(int(C["revenue_2023"]*0.272)), "+2,9 %"]]),
]
for fname, folder, title, table_data in fin_extra:
    if len(generated) >= 300: break
    generated.append(make_xlsx(folder, f"{fname}.xlsx", title,
                                [("KST", table_data[0], table_data[1:], [10, 30, 18, 18, 14])]))

# Bewerbungsunterlagen / Einstellungsunterlagen (anonymisiert)
bewerber = [
    ("Einstellungsprotokoll_Ingenieur_FuE_2023", "03_Personal_HR",
     "Einstellungsprotokoll: Entwicklungsingenieur F&E",
     f"Gesprächsteilnehmer: Sabine Koch (Leiterin F&E), Petra Zimmermann (HR)\nBewerber: anonymisiert\nDatum: {date_str(2023, 8, 22)}\n\nBewertung: Sehr gute Fachkenntnisse SPS-Programmierung (Siemens S7), gute Englischkenntnisse, 5 Jahre Berufserfahrung.\nEmpfehlung: Einstellung zum 01.10.2023\nGehaltsvorstellung: 68.000 € – angeboten: 65.000 € + Zielvereinbarung"),
    ("Zeugnis_Michael_Weber_Senior", "03_Personal_HR",
     "Zwischenzeugnis Michael Weber",
     f"Herr Michael Weber ist seit dem 01.05.2007 als Leiter Produktion bei uns beschäftigt. Er koordiniert eine Belegschaft von 148 Mitarbeitern in der Serienfertigung und in Sonderprojekten. Herr Weber hat in allen Bereichen hervorragende Leistungen erbracht. Wir schätzen sein Engagement, seine Zuverlässigkeit und sein unternehmerisches Denken sehr."),
    ("Betriebsrat_Sitzungsprotokoll_Q1_2024", "03_Personal_HR",
     "Betriebsratssitzung Protokoll Q1 2024",
     f"Sitzung vom {date_str(2024, 2, 8)}\nTeilnehmer: Hans-Peter Dietrich (Vorsitzender), Monika Sauer, 3 weitere BR-Mitglieder, {C['gf2_name']} (Arbeitgeberseite)\n\nTagesordnung:\n1. Gehaltsrunde 2024 → Einigung auf +3,5 % ab 01.04.2024\n2. Homeoffice-Regelung → BV Mobiles Arbeiten verlängert bis 31.12.2025\n3. Stellenplan 2024 → 12 neue Stellen genehmigt\n4. Verschiedenes"),
]
for fname, folder, title, content in bewerber:
    if len(generated) >= 300: break
    doc = Document()
    make_docx_header(doc, title)
    doc.add_paragraph(content)
    generated.append(save_docx(doc, folder, f"{fname}.docx"))

# Compliance: Auditberichte
audit_extras = [
    ("COMP_Internes_Audit_ISO_9001_2023", "Interner Audit-Bericht ISO 9001:2015 – November 2023",
     f"Auditleitung: Julia Krause (QS)\nAuditiert: Produktion, Einkauf, Vertrieb\nDatum: {date_str(2023, 11, 7)}\n\nFeststellungen:\n• Konformitäten: 42\n• Geringfügige Abweichungen (Klasse B): 3 → Korrektivmaßnahmen bis 31.01.2024\n• Kritische Abweichungen (Klasse A): 0\n\nFazit: Das QM-System erfüllt die Anforderungen der ISO 9001:2015. Empfehlung zur Re-Zertifizierung ausgesprochen."),
    ("COMP_Lieferantenbewertung_Matrix_2023", "Lieferantenbewertung 2023 – Gesamtmatrix",
     f"Bewertungsperiode: Geschäftsjahr 2023\nVerantwortlich: Stefan Braun (Einkauf)\n\nBewertungskriterien: Qualität (40 %), Liefertreue (30 %), Preis/Leistung (20 %), Nachhaltigkeit (10 %)\n\nErgebnisse:\n• {C['lieferant1']}: 87/100 Punkte (A-Lieferant)\n• {C['lieferant2']}: 91/100 Punkte (A-Lieferant)\n• {C['lieferant3']}: 83/100 Punkte (A-Lieferant)\n• {C['lieferant4']}: 78/100 Punkte (B-Lieferant → Maßnahmenplan)"),
    ("COMP_Risikoregister_2024", "Unternehmensrisikoregister 2024",
     f"Erstellt: {C['gf2_name']} | Genehmigt: {C['gf1_name']} | Datum: {date_str(2024, 1, 10)}\n\nTop-Risiken:\n1. Klumpenrisiko Kundenbasis (ThyssenKrupp ~17 %) – Maßnahme: Diversifikation Neukunden\n2. Lieferkettenunterbrechungen (Elektronik) – Maßnahme: Dual Sourcing, Lagerhaltung +30 %\n3. Fachkräftemangel – Maßnahme: Kooperation RWTH Aachen, Ausbildungsoffensive\n4. Cyberangriff – Maßnahme: CrowdStrike MDR, Incident Response Plan\n5. Konjunkturabschwung Automobil – Maßnahme: Diversifikation in Energie-/Medizintechnik"),
    ("COMP_Whistleblower_Richtlinie", "Hinweisgebersystem (Whistleblower-Richtlinie)",
     f"Gemäß HinSchG (Hinweisgeberschutzgesetz) vom 2. Juli 2023 hat {C['name']} ein internes Hinweisgebersystem eingerichtet.\n\nMeldekanal: compliance@halbreiter-maschinenbau.de (verschlüsselt) oder per Post an den externen Ombudsmann.\nOmbudsmann: Dr. Hans-Georg Stern, Rechtsanwalt, Köln\nSchutz: Hinweisgeber genießen vollständigen Schutz vor Repressalien."),
    ("COMP_Nachhaltigkeitsbericht_2023", "Nachhaltigkeitsbericht 2023 (ESG)",
     f"Erstellt gemäß GRI Standards und CSRD-Vorbereitung.\n\nUMWELT:\n• CO₂-Ausstoß Scope 1+2: 1.840 t CO₂eq (Vorjahr: 2.100 t, -12,4 %)\n• Energieverbrauch: 8.400 MWh (-8 % ggü. Vj.)\n• Ziel 2027: CO₂-neutral am Standort Köln\n\nSOZIALES:\n• Ausbildungsquote: 5,8 % der Belegschaft\n• Mitarbeiterzufriedenheit (Umfrage 2023): 78 % zufrieden/sehr zufrieden\n\nGOVERNANCE:\n• Compliance-Schulungen: 100 % der Mitarbeiter absolviert\n• Keine wesentlichen Compliance-Verstöße in 2023"),
]
for fname, title, content in audit_extras:
    if len(generated) >= 300: break
    doc = Document()
    make_docx_header(doc, title)
    doc.add_paragraph(content)
    generated.append(save_docx(doc, "09_Compliance", f"{fname}.docx"))

# IP: Markenanmeldungen, Gebrauchsmuster
ip_extra = [
    ("IP_Marke_MMB_Logo_DPMA", "Markenschutzregister – MMB Wort-/Bildmarke",
     f"Markeninhaber: {C['name']}\nDPMA Aktenzeichen: 30 2019 012 345\nAnmeldetag: {date_str(2019, 3, 10)}\nEintragungstag: {date_str(2019, 9, 5)}\nMarkenkategorie: Kombinierte Marke (Wort/Bild)\nWarenklassen: Klasse 7 (Maschinen und Werkzeugmaschinen), Klasse 40 (Bearbeitung von Materialien)\nSchutzende: {date_str(2029, 3, 9)}"),
    ("IP_Gebrauchsmuster_DE202021", "Gebrauchsmuster DE 20 2021 005 678",
     f"Inhaber: {C['name']}\nBezeichnung: Einstellbarer Präzisionsanschlag für hydraulische Stanzwerkzeuge\nAnmeldetag: {date_str(2021, 6, 14)}\nEintragungstag: {date_str(2021, 8, 25)}\nSchutzende: {date_str(2029, 6, 13)} (max. Schutzdauer 8 Jahre)"),
    ("IP_Technologie_Transfer_RWTH", "Technologietransfer-Vereinbarung RWTH Aachen",
     f"Vereinbarung zwischen {C['name']} und RWTH Aachen, Institut für Produktionstechnik.\nGegenstand: Gemeinsame F&E im Bereich KI-gestützte Fertigungsoptimierung.\nLaufzeit: 2023–2025\nFinanzierungsbeitrag MMB: {fmt_eur(85_000)}/Jahr\nIP-Regelung: Gemeinsame Schutzrechte; MMB erhält exklusive Nutzungslizenz für 5 Jahre."),
]
for fname, title, content in ip_extra:
    if len(generated) >= 300: break
    doc = Document()
    make_docx_header(doc, title)
    doc.add_paragraph(content)
    generated.append(save_docx(doc, "07_IP_Lizenzen", f"{fname}.docx"))

# Versicherungen: Schadensmeldungen, Schadenshistorie
vs_extra = [
    ("VS_Schadensmeldung_2022_Maschinenschaden", "08_Versicherungen",
     "Schadensmeldung 2022 – Maschinenschaden Fräszentrum",
     f"Versicherungsnehmer: {C['name']}\nVersicherung: {C['versicherer2']}, Maschinenversicherung VS-HDI-2022-0004\nSchadensdatum: {date_str(2022, 7, 12)}\nSchadensbeschreibung: Defekt der Spindellagerung am 5-Achs-Fräszentrum durch Werkzeugbruch.\nSchadenssumme: {fmt_eur(48_500)}\nStatus: Reguliert am {date_str(2022, 9, 30)}. Selbstbehalt: {fmt_eur(5_000)}."),
    ("VS_Schadensmeldung_2023_Wasserschaden", "08_Versicherungen",
     "Schadensmeldung 2023 – Wasserschaden Verwaltungsgebäude",
     f"Versicherungsnehmer: {C['name']}\nVersicherung: {C['versicherer2']}, Gebäudeversicherung VS-HDI-2022-0003\nSchadensdatum: {date_str(2023, 1, 18)}\nSchadensbeschreibung: Rohrbruch Heizungsleitung 1. OG, Wasserschaden ca. 120 m² Bürofläche.\nSchadenssumme: {fmt_eur(32_000)}\nStatus: Vollständig reguliert am {date_str(2023, 4, 5)}."),
]
for fname, folder, title, content in vs_extra:
    if len(generated) >= 300: break
    doc = Document()
    make_docx_header(doc, title)
    doc.add_paragraph(content)
    generated.append(save_docx(doc, folder, f"{fname}.docx"))

# Fill remainder with financial planning worksheets
planning_sheets = [
    ("FIN_Umsatzplanung_2024_nach_Kunden", "Umsatzplanung 2024 nach Kunden",
     [["Kunde", "Auftragsbestand", "Prognose H1", "Prognose H2", "Gesamt 2024e"],
      [C["kunde1"], fmt_eur(3_200_000), fmt_eur(4_100_000), fmt_eur(4_200_000), fmt_eur(8_300_000)],
      [C["kunde2"], fmt_eur(1_800_000), fmt_eur(3_000_000), fmt_eur(2_900_000), fmt_eur(5_900_000)],
      [C["kunde3"], fmt_eur(2_400_000), fmt_eur(2_800_000), fmt_eur(2_700_000), fmt_eur(5_500_000)],
      [C["kunde4"], fmt_eur(950_000), fmt_eur(1_800_000), fmt_eur(1_900_000), fmt_eur(3_700_000)],
      [C["kunde5"], fmt_eur(1_100_000), fmt_eur(1_900_000), fmt_eur(2_000_000), fmt_eur(3_900_000)],
      ["Haver & Boecker OHG", fmt_eur(2_200_000), fmt_eur(2_200_000), fmt_eur(1_800_000), fmt_eur(4_000_000)],
      ["Sonstige (ca. 55 Kunden)", fmt_eur(6_500_000), fmt_eur(10_400_000), fmt_eur(10_400_000), fmt_eur(20_800_000)],
      ["GESAMT", fmt_eur(18_150_000), fmt_eur(26_200_000), fmt_eur(25_900_000), fmt_eur(C["revenue_2024e"])]]),
    ("FIN_Investitionscontrolling_Q3_2024", "Investitionscontrolling Q3 2024",
     [["Projekt", "Budget", "Verbraucht Q1-Q3", "Prognose Jahresende", "Abweichung"],
      ["Laserschneider LS-800 (3 Stk.)", fmt_eur(1_800_000), fmt_eur(1_820_000), fmt_eur(1_820_000), "+1,1 %"],
      ["Erweiterung Halle B", fmt_eur(950_000), fmt_eur(680_000), fmt_eur(970_000), "+2,1 %"],
      ["ERP-Upgrade SAP S-4HANA 2024", fmt_eur(420_000), fmt_eur(395_000), fmt_eur(430_000), "+2,4 %"],
      ["Robotermontagezelle", fmt_eur(680_000), fmt_eur(320_000), fmt_eur(690_000), "+1,5 %"],
      ["Photovoltaikanlage 250 kWp", fmt_eur(320_000), fmt_eur(315_000), fmt_eur(320_000), "0,0 %"],
      ["Gesamt", fmt_eur(4_170_000), fmt_eur(3_530_000), fmt_eur(4_230_000), "+1,4 %"]]),
    ("FIN_Forderungsmanagement_2023", "Offene Forderungen und Mahnwesen 2023",
     [["Debitor", "Offener Betrag", "Fällig seit", "Mahnstufe", "Status"],
      [C["kunde1"], fmt_eur(285_000), "< 30 Tage", "-", "In Zahlung"],
      [C["kunde2"], fmt_eur(148_000), "< 30 Tage", "-", "In Zahlung"],
      [C["kunde4"], fmt_eur(92_000), "31–60 Tage", "Mahnstufe 1", "Gemahnt"],
      ["Gebr. Heyl GmbH", fmt_eur(67_000), "61–90 Tage", "Mahnstufe 2", "RA eingeschaltet"],
      ["Diverse", fmt_eur(214_000), "< 30 Tage", "-", "In Zahlung"],
      ["Wertberichtigung", f"-{fmt_eur(35_000)}", "", "", "EWB gebildet"]]),
]
for fname, title, table_data in planning_sheets:
    if len(generated) >= 300: break
    generated.append(make_xlsx("02_Finanzen", f"{fname}.xlsx", title,
                                [(title[:15], table_data[0], table_data[1:], [40] + [18]*(len(table_data[0])-1))]))

# Strategy extras
for yr in [2021, 2022]:
    if len(generated) >= 300: break
    sections = [
        (f"Jahresbericht {yr} – Zusammenfassung", f"""
Umsatz {yr}: {fmt_eur(C.get(f'revenue_{yr}', int(C['revenue_2022'] * 0.88)))}
EBITDA {yr}: {fmt_eur(C.get(f'ebitda_{yr}', int(C['ebitda_2022'] * 0.88)))}
Mitarbeiter: {C.get(f'employees_{yr}', 225)}

Highlights:
• {C['name']} verzeichnete solides Wachstum trotz herausfordernden Marktbedingungen
• Großauftrag {C['kunde1']}: Erweiterung der bestehenden Maschinenlinie
• ISO 9001:2015 Zertifizierung erfolgreich erneuert
"""),
    ]
    generated.append(make_pdf("11_Strategie_Planung", f"STRAT_Jahresbericht_{yr}.pdf", f"Jahresbericht {yr}", sections))

# Final generic docs if still short
generic_docs = [
    ("Korrespondenz_Bank_Jahresgespräch_2024", "02_Finanzen",
     f"Protokoll Bankgespräch Deutsche Bank, {date_str(2024, 2, 6)}\n\nTeilnehmer: {C['gf1_name']}, {C['gf2_name']} (MMB); Frau Dr. Susanne Vogt, Herr Patrick Hülsmann (Deutsche Bank, Firmenkundenbetreuung Köln)\n\nThemen:\n1. Jahresabschluss 2023 – positiv bewertet, Rating unverändert 'BB+'\n2. Prolongation Kontokorrentlinie 3 Mio. € bis 31.12.2026 – mündlich zugesagt\n3. Investitionsdarlehen 2024 – Deutsche Bank bietet 2,0 Mio. € zu 4,2 % p.a. an\n4. Factoring-Angebot für Forderungsbestand prüfen\n\nNächstes Gespräch: Vorlage geprüfter Jahresabschluss 2023 nach Prüfungsbericht {C['wp']}."),
    ("Korrespondenz_WP_Prüfungsauftrag_2023", "02_Finanzen",
     f"Auftragserteilung Jahresabschlussprüfung 2023\n\nAuftraggeberin: {C['name']}\nWirtschaftsprüfer: {C['wp']}\nGegenstand: Prüfung Jahresabschluss zum 31. Dezember 2023\nPrüfungsstandard: IDW PS 200 (ISA)\nHonorar: {fmt_eur(95_000)} zzgl. MwSt.\nBerichtstermin: spätestens {date_str(2024, 4, 30)}\n\nBilanzeid bestätigt am {date_str(2024, 3, 15)}."),
    ("STRAT_Wettbewerbsanalyse_2024", "11_Strategie_Planung",
     f"Wettbewerbsanalyse Sondermaschinen 2024\n\nHauptwettbewerber:\n1. KUKA AG – Schwerpunkt Automotive Robotik, Umsatz >3 Mrd. €, Marktführer in Schweißrobotern\n2. Grob-Werke GmbH & Co. KG – Bearbeitungszentren, Umsatz ~1,8 Mrd. €, Fokus Automobilindustrie\n3. Schuler AG – Umformtechnik, Pressen, Umsatz ~1,2 Mrd. €\n4. WEILER GmbH – CNC-Drehmaschinen, Umsatz ~350 Mio. €\n\nPositionierung MMB: Mittelstands-Nische, höchste Individualisierung, schnellere Lieferzeiten als Großanbieter, Preis 10–15 % unter KUKA/Schuler bei vergleichbaren Projekten."),
    ("LF_Zollpräferenznachweis_2023", "05_Vertraege_Lieferanten",
     f"Zollpräferenznachweis / EUR.1 Warenverkehrsbescheinigung\n\nAussteller: {C['lieferant3']}\nEmpfänger: {C['name']}\nDatum: {date_str(2023, 6, 15)}\nWarenbezeichnung: Elektrische Antriebssysteme, Serienummer 2023-ST-4892\nWarenwert: {fmt_eur(380_000)}\nUrsprungsland: Deutschland (EU)\nZollposition: 8504.40 (Statische Umformer)"),
    ("KD_Kundenzufriedenheitsumfrage_2023", "04_Vertraege_Kunden",
     f"Kundenzufriedenheitsumfrage 2023 – Auswertung\n\nRücklaufquote: 68 Fragebögen von 82 angeschriebenen Kunden (83 %)\n\nBewertung (1=sehr schlecht, 5=sehr gut):\n• Produktqualität: 4,6\n• Liefertreue: 4,2\n• Kundenservice: 4,4\n• Preis-Leistungs-Verhältnis: 4,1\n• Weiterempfehlungsbereitschaft (NPS): +52\n\nMaßnahmen: Verbesserung Liefertreue durch Kapazitätserweiterung Q2 2024."),
    ("HR_Entgeltrahmentarifvertrag_ERA", "03_Personal_HR",
     f"Anwendung Entgeltrahmentarifvertrag (ERA) Metallindustrie NRW\n\nDer ERA Metallindustrie NRW wird bei {C['name']} seit 01.01.2010 vollständig angewendet.\nEntgeltgruppen: E4–E13 für gewerbliche Mitarbeiter; AT-Verträge für Führungskräfte ab Abteilungsleiterebene.\nTarif-Gehaltstabelle: gemäß aktuellem Tarifabschluss Gesamtmetall/IG Metall (letzter Abschluss: +3,5 % ab Oktober 2023).\nBetriebsrat informiert und einbezogen."),
    ("COMP_Aufsichtsbehörden_Kommunikation_2023", "09_Compliance",
     f"Schriftverkehr mit Aufsichtsbehörden 2023 – Übersicht\n\nArbeitssicherheit (BG Metall Nord Süd):\n• Jahresmeldung Betriebsrevision: {date_str(2023, 3, 31)} – keine Beanstandungen\n• Betriebsbegehung: {date_str(2023, 9, 12)} – 2 geringe Mängel, beseitigt bis {date_str(2023, 10, 15)}\n\nUmweltschutz (Bezirksregierung Köln):\n• Emissionsbericht 2022 eingereicht: {date_str(2023, 3, 28)} – alle Grenzwerte eingehalten\n• Abfallwirtschaftskonzept aktualisiert\n\nDatensschutz (LDI NRW):\n• Keine Meldepflicht-Vorfälle in 2023"),
    ("GR_Vollmacht_Steuerberater", "01_Gesellschaftsrecht",
     f"Steuerberatervertrag und Vollmacht\n\nAuftraggeber: {C['name']}\nSteuerberater: {C['steuerber']}\nGegenstand: Steuerliche Beratung, Erstellung Steuererklärungen, Vertretung gegenüber Finanzbehörden\nVollmacht erteilt am: {date_str(2018, 1, 15)}\nHonorar: Gemäß StBVV, Jahresumsatz Steuerberatung: ca. {fmt_eur(85_000)}\nAnwalt: {C['anwalt']} für gesellschaftsrechtliche Fragen"),
    ("GR_Notarielle_Urkundenliste", "01_Gesellschaftsrecht",
     f"Notarielle Urkunden – Verzeichnis\n\n1. Gesellschaftsvertrag, Notar Dr. Bernhard Klein, Köln, {C['founded']}, UR-Nr. 234/1985\n2. Änderung Gesellschaftsvertrag, Notar Dr. Klein, Köln, 15.01.2019, UR-Nr. 89/2019\n3. Bestellung GF {C['gf2_name']}, Notar Dr. Klein, Köln, 01.09.2015, UR-Nr. 412/2015\n4. Abtretungsvertrag Geschäftsanteil (Erbfall), Notar Dr. Thomas Maier, Köln, 12.03.2020, UR-Nr. 156/2020\n5. Grundschuldbestellung Deutsche Bank, Notarin Dr. Eva Roth, Köln, 01.04.2019, UR-Nr. 890/2019"),
]
for fname, folder, content in generic_docs:
    if len(generated) >= 300: break
    doc = Document()
    make_docx_header(doc, fname.replace("_", " "))
    doc.add_paragraph(content)
    generated.append(save_docx(doc, folder, f"{fname}.docx"))

# Absolute fill-up: remaining docs as realistic financial/operational records
fill_templates = [
    ("Reisekostenabrechnung_GF1_Q1_2024", "02_Finanzen", f"{C['gf1_name']}: Reisekosten Q1 2024\nGeschäftsreisen: Frankfurt (Kunde {C['kunde5']}), München (Messe bauma 2024), Warschau (Marktrecherche)\nSumme: {fmt_eur(8_420)}\nGenehmigt von: {C['gf2_name']}"),
    ("Reisekostenabrechnung_Vertrieb_Q2_2024", "02_Finanzen", f"Vertriebsreisen Q2 2024 – Thomas Schneider\nKundenbesuche: {C['kunde1']} (Duisburg), {C['kunde2']} (Stuttgart), {C['kunde3']} (Lippstadt), {C['kunde4']} (Allendorf)\nSumme: {fmt_eur(5_840)}\nGenehmigt."),
    ("Bestandsbewertung_Roh_Hilf_Betriebsstoffe_2023", "02_Finanzen", f"Inventur Roh-, Hilfs- und Betriebsstoffe per 31.12.2023\nRohstoffe (Stahl, Aluminium): {fmt_eur(1_840_000)}\nHilfsstoffe (Werkzeuge, Schmierstoffe): {fmt_eur(420_000)}\nBetriebsstoffe (Verbrauchsmaterial): {fmt_eur(180_000)}\nUnfertige Erzeugnisse: {fmt_eur(2_640_000)}\nFertige Erzeugnisse/Handelsware: {fmt_eur(380_000)}\nGesamt Vorräte: {fmt_eur(5_460_000)}"),
    ("Bankbürgschaft_Exportgarantie_2023", "02_Finanzen", f"Bürgschaftsurkunde\nBürgschaftsgeber: {C['bank']}\nBegünstigte: {C['kunde1']}\nBetrag: {fmt_eur(500_000)}\nZweck: Liefergarantie Pressenlinie PL-500\nGültig bis: {date_str(2025, 12, 31)}\nBürgschaftsgebühr: 0,8 % p.a."),
    ("Factoring_Rahmenvertrag_2023", "02_Finanzen", f"Factoringvertrag\nFactor: Targo Commercial Finance AG\nKlient: {C['name']}\nFactoringvolumen: bis {fmt_eur(8_000_000)} offene Forderungen\nFactoringgebühr: 0,65 % des Forderungsbetrags\nZinssatz auf Vorfinanzierung: EURIBOR 1M + 1,2 %\nInkrafttretung: 01.04.2024"),
    ("Jahresabschluss_Kennzahlenblatt_2023", "02_Finanzen", f"Kennzahlenübersicht Jahresabschluss 2023\nUmsatz: {fmt_eur(C['revenue_2023'])}\nEBITDA-Marge: {C['ebitda_2023']/C['revenue_2023']*100:.1f} %\nEBIT-Marge: {C['ebit_2023']/C['revenue_2023']*100:.1f} %\nEigenkapitalquote: {C['eigenkapital_2023']/C['bilanzsumme_2023']*100:.1f} %\nVerschuldungsgrad (Net Debt/EBITDA): 1,2×\nROCE: 21,4 %\nWorking Capital: {fmt_eur(int(C['bilanzsumme_2023']*0.18))}"),
    ("Preisliste_2024_Standardprodukte", "04_Vertraege_Kunden", f"Preisliste {C['name']} – gültig ab 01.01.2024\n\n{C['produkt1']}: ab {fmt_eur(580_000)} (Basisversion)\n{C['produkt2']}: ab {fmt_eur(180_000)} (Standardlänge 20 m)\n{C['produkt3']}: ab {fmt_eur(320_000)} (Standardkonfiguration)\n{C['produkt4']}: ab {fmt_eur(420_000)} (inkl. Sicherheitszertifizierung)\n\nAlle Preise ab Werk Köln, zzgl. MwSt. und Transportkosten.\nPreisanpassungsklausel: Stahlpreisindexgebunden (ifo Rohstoffindex)."),
    ("LF_Preferred_Supplier_Siemens_2023", "05_Vertraege_Lieferanten", f"Preferred Supplier Agreement – Siemens AG Antriebstechnik\n\nKunde: {C['name']}\nLieferant: {C['lieferant3']}\nVolumen 2024: {fmt_eur(2_100_000)} (Forecast)\nRabattstaffel: -3 % ab 1,5 Mio. €, -5 % ab 2,0 Mio. €\nLieferzeit-Garantie: max. 8 Wochen für Standardantriebe\nTechnischer Support: dedizierter Ansprechpartner, max. 4 h Reaktionszeit"),
    ("GR_Gesellschafterbeschluss_Digitalisierung", "01_Gesellschaftsrecht", f"Gesellschafterbeschluss – Freigabe Digitalisierungsstrategie 2024–2026\n\n{date_str(2024, 1, 12)} – Die Gesellschafter genehmigen die Digitalisierungsstrategie mit einem Investitionsrahmen von {fmt_eur(1_800_000)} über 3 Jahre.\nSchwerpunkte: ERP-Upgrade, Digitaler Zwilling, IoT-Integration Produktion, KI-gestützte Qualitätskontrolle."),
    ("COMP_Datenschutzfolgeabschaetzung_CRM", "09_Compliance", f"Datenschutz-Folgenabschätzung (Art. 35 DSGVO) – CRM-System\n\nVerarbeitungssystem: {C['crm']}\nVerantwortlicher: {C['name']}, {C['gf2_name']}\nDatenschutzbeauftragter: DataCo GmbH\nDatum: {date_str(2023, 4, 1)}\n\nRisikobewertung: Mittleres Risiko (Score: 3/5)\nMaßnahmen: Verschlüsselung at rest und in transit, Zugriffsprotokollierung, jährliche Löschung inaktiver Datensätze, Auftragsverarbeitung nach Art. 28 DSGVO mit Salesforce vereinbart."),
    ("IT_Cloud_Strategie_2024_2026", "10_IT_Infrastruktur", f"Cloud-Strategie {C['name']} 2024–2026\n\nAktuelle Cloud-Readiness: 35 % der Anwendungen cloud-fähig\nZiel 2026: 60 % cloud-basiert (Hybrid-Cloud-Ansatz)\n\nGeplante Migrationen:\n• Microsoft 365 (abgeschlossen)\n• Zscaler ZTNA (abgeschlossen 2023)\n• Salesforce CRM (bereits Cloud)\n• SAP S/4HANA → SAP RISE (geplant 2025, Entscheidung offen)\n• ERP-Backup in Azure: geplant Q3 2024\n\nSicherheitsanforderungen: BSI C5-Anforderungskatalog für alle Cloud-Dienste."),
]
for fname, folder, content in fill_templates:
    if len(generated) >= 300: break
    doc = Document()
    make_docx_header(doc, fname.replace("_", " "))
    doc.add_paragraph(content)
    generated.append(save_docx(doc, folder, f"{fname}.docx"))

# Last resort: numbered technical records
tr_topics = [
    ("Prüfprotokoll_PL500_SN2023001", "09_Compliance", "Abnahmeprotokoll Pressenlinie PL-500, Seriennummer 2023-001"),
    ("Prüfprotokoll_PL500_SN2023002", "09_Compliance", "Abnahmeprotokoll Pressenlinie PL-500, Seriennummer 2023-002"),
    ("Prüfprotokoll_LS800_SN2023001", "09_Compliance", "Abnahmeprotokoll Laserschneider LS-800, Seriennummer 2023-001"),
    ("CE_Konformitaetserklaerung_PL500", "09_Compliance", f"EG-Konformitätserklärung gemäß Maschinenrichtlinie 2006/42/EG für Pressenlinie PL-500.\nHersteller: {C['name']}. Notifizierte Stelle: TÜV Nord, Nr. 0045."),
    ("CE_Konformitaetserklaerung_LS800", "09_Compliance", f"EG-Konformitätserklärung für Laserschneideanlage LS-800 gemäß Maschinenrichtlinie und Niederspannungsrichtlinie.\nHersteller: {C['name']}. Laserklasse: 4."),
    ("Technische_Dokumentation_PL500_Rev3", "07_IP_Lizenzen", f"Technische Dokumentation Pressenlinie PL-500, Revision 3.1.\nGeheimhaltungsstufe: VERTRAULICH\nErstellt von: Sabine Koch (F&E)\nDokumentiert: konstruktive Merkmale, Schaltpläne, Hydraulikschema."),
    ("Wartungshandbuch_FB200", "07_IP_Lizenzen", f"Wartungs- und Betriebshandbuch Förderbandanlage FB-200.\nWartungsintervalle: 500 h (Inspektion), 2.000 h (Vollwartung), 8.000 h (Überholung).\nErsatzteilnummern: gemäß Ersatzteilkatalog Rev. 4 (2024)."),
    ("Sicherheitsdatenblatt_Kuehlschmierstoff", "09_Compliance", "Sicherheitsdatenblatt Kühlschmierstoff KSS-200 (wassermischbar).\nGefahrpiktogramm: GHS07 (Ausrufezeichen).\nHandhabungsvorschriften: Schutzbrille, Chemikalienschutzhandschuhe.\nEntsorgung: gemäß Abfallverzeichnis-Verordnung, AVV 12 01 09."),
    ("Anlagenakte_Maschine_FZ4891", "06_Immobilien", f"Anlagenakte Maschine FZ-4891 (5-Achs-Fräszentrum DMG MORI DMU 85 monoBLOCK)\nAnschaffungsdatum: {date_str(2020, 5, 12)}\nAnschaffungspreis: {fmt_eur(680_000)}\nNutzungsdauer: 10 Jahre\nJährliche Abschreibung: {fmt_eur(68_000)}\nBuchwert 31.12.2023: {fmt_eur(296_000)}\nStandort: Halle A, Position 12"),
    ("Fuhrparkverzeichnis_2024", "06_Immobilien", f"Fuhrpark {C['name']}, Stand: Januar 2024\n\n1. BMW 5er (GF {C['gf1_name']}): KFZ-KE 4821 H, Baujahr 2023, Leasingrate {fmt_eur(890)}/Monat\n2. Mercedes C-Klasse (GF {C['gf2_name']}): KFZ-KE 4822 H, Baujahr 2022, Leasingrate {fmt_eur(720)}/Monat\n3. VW Passat (Thomas Schneider): KFZ-KE 4823 H, Baujahr 2022, Leasingrate {fmt_eur(580)}/Monat\n4. 5× VW Transporter (Montage/Service): diverse KFZ\nGesamte Flottenkosten 2023: {fmt_eur(128_000)}/Jahr"),
]
for fname, folder, content in tr_topics:
    if len(generated) >= 300: break
    doc = Document()
    make_docx_header(doc, fname.replace("_", " "))
    doc.add_paragraph(content)
    generated.append(save_docx(doc, folder, f"{fname}.docx"))

print(f"\n  Wave 2 complete. Total: {len(generated)} documents")

# ─── WAVE 3: Final push to 300 ───────────────────────────────────────────────
print(f"\n📁 Wave 3: Pushing to 300 (currently {len(generated)})...")

wave3_docs = [
    # Gesellschaftsrecht
    ("GR_Handlungsvollmacht_Einkauf_Stefan_Braun", "01_Gesellschaftsrecht",
     f"Handlungsvollmacht\n\n{C['name']} erteilt Herrn Stefan Braun, Leiter Einkauf, Handlungsvollmacht für den Abschluss von Einkaufsverträgen bis zu einem Wert von {fmt_eur(250_000)} pro Einzelauftrag ohne weitere Genehmigung der Geschäftsführung.\n\nKöln, den {date_str(2022, 1, 1)}\n\n{C['gf1_name']}, Geschäftsführer"),
    ("GR_Erbfolgeklausel_Testament_Hinweis", "01_Gesellschaftsrecht",
     f"Vermerk zur Nachfolgeplanung / Erbfolgeklausel\n\nGemäß Gesellschaftsvertrag § 12 (Erbfolge) gilt: Im Todesfall eines Gesellschafters wird der Geschäftsanteil an die Erben übertragen, sofern diese durch Gesellschafterbeschluss zugelassen werden. Andernfalls sind die Erben zur Einziehung oder Abtretung verpflichtet.\n\nAktueller Sachstand (vertraulich): Nachfolgeplanung {C['gf1_name']} nicht abgeschlossen. Testamentarische Regelung liegt beim Notariat Dr. Klein vor, ist aber nicht Bestandteil dieser Due-Diligence-Unterlage."),
    # More Finance
    ("FIN_Abschlussprüfungsbericht_WP_2022", "02_Finanzen",
     f"Bestätigungsvermerk des unabhängigen Abschlussprüfers\n\nWir haben den Jahresabschluss der {C['name']}, Köln, bestehend aus Bilanz zum 31. Dezember 2022, Gewinn- und Verlustrechnung, Anhang und Lagebericht für das Geschäftsjahr 2022 geprüft.\n\nPrüfungsurteil: Nach unserer Beurteilung vermittelt der Jahresabschluss ein den tatsächlichen Verhältnissen entsprechendes Bild der Vermögens-, Finanz- und Ertragslage des Unternehmens.\n\n{C['wp']}, Köln, {date_str(2023, 4, 28)}\nppa. Dr. Sabine Richter, Wirtschaftsprüferin"),
    ("FIN_Abschlussprüfungsbericht_WP_2023", "02_Finanzen",
     f"Bestätigungsvermerk des unabhängigen Abschlussprüfers\n\nWir haben den Jahresabschluss der {C['name']}, Köln, zum 31. Dezember 2023 geprüft.\n\nPrüfungsurteil: Nach unserer Beurteilung vermittelt der Jahresabschluss ein den tatsächlichen Verhältnissen entsprechendes Bild.\n\nUmsatz 2023: {fmt_eur(C['revenue_2023'])} | Jahresüberschuss: {fmt_eur(int(C['ebit_2023']*0.65))}\n\n{C['wp']}, Köln, {date_str(2024, 4, 25)}\nppa. Dr. Sabine Richter, Wirtschaftsprüferin"),
    ("FIN_Leasingvertrag_Fuhrpark_2023", "02_Finanzen",
     f"Leasingrahmenvertrag Fuhrpark\n\nLeasinggeber: BMW Financial Services Deutschland GmbH\nLeasingnehmer: {C['name']}\nLeasingvolumen: 12 Fahrzeuge, Gesamtrate {fmt_eur(8_400)}/Monat\nLaufzeit: 36 Monate (rollierend)\nRestwertgarantie: keine"),
    ("FIN_Devisensicherung_2024", "02_Finanzen",
     f"Devisensicherungsvertrag\n\nVertragspartner: {C['bank']}\n{C['name']} sichert Exportforderungen in USD gegen EUR-Schwankungen ab.\nVolumen 2024: USD 1.800.000 (Lieferungen nach USA, hauptsächlich {C['kunde1']}-Kontrakte)\nDevisentermingeschäft: USD/EUR Kurs gesichert bei 1,08, Fälligkeit quartalsweise\nPrämie: {fmt_eur(12_400)} einmalig"),
    # More HR
    ("HR_Arbeitnehmererfindungsgesetz_Meldung_2021", "03_Personal_HR",
     f"Meldung einer Arbeitnehmererfindung gemäß § 5 ArbnErfG\n\nErfinder: Sabine Koch, Michael Weber\nDatum der Meldung: {date_str(2021, 3, 15)}\nBezeichnung: Modulares Transportsystem mit KI-basierter Lastoptimierung\nArbeitgeber: {C['name']}\nInanspruchnahme durch Arbeitgeber: Ja, gemäß § 6 ArbnErfG\nVergütung: Erfindervergütung pauschal {fmt_eur(15_000)} an jeden Erfinder\nPatent daraus: DE 10 2021 098 765 (s. IP-Ordner)"),
    ("HR_Gehaltsstruktur_AT_Mitarbeiter_2024", "03_Personal_HR",
     f"Gehaltsstruktur außertarifliche Mitarbeiter (AT) 2024 – vertraulich\n\nGf-Ebene: {fmt_eur(220_000)} – {fmt_eur(280_000)}\nAbteilungsleiter: {fmt_eur(80_000)} – {fmt_eur(105_000)}\nTeamleiter / Senior Specialist: {fmt_eur(65_000)} – {fmt_eur(82_000)}\nSpezialist: {fmt_eur(52_000)} – {fmt_eur(68_000)}\n\nBonus-Pool 2024: {fmt_eur(int(C['revenue_2024e'] * 0.008))} (ca. 1,5 % des Umsatzes)"),
    # More Compliance
    ("COMP_Revisionsplan_2024", "09_Compliance",
     f"Interner Revisionsplan 2024\n\nRevisions-Verantwortliche: {C['gf2_name']} (CFO)\n\nQ1 2024: Prüfung Einkaufsprozesse und Lieferantenverträge\nQ2 2024: Prüfung Produktion (KST-Abweichungen, Ausschussquoten)\nQ3 2024: IT-Sicherheitsaudit (extern: CrowdStrike-Bericht + interner Review)\nQ4 2024: Jahresabschluss-vorbereitung, Steuer-Compliance\n\nExterneRevision durch {C['wp']}: April 2025 (Jahresabschluss 2024)"),
    ("COMP_REACH_Verordnung_Dokumentation", "09_Compliance",
     f"REACH-Konformitätsdokumentation\n\nGemäß EU-Verordnung 1907/2006 (REACH) dokumentiert {C['name']} die verwendeten Chemikalien.\n\nRegistrierte Stoffe: 12 (Kühlschmierstoffe, Lacke, Reiniger)\nSVHC-Kandidatenliste: Keine SVHC-Stoffe in Mengen >0,1 % verwendet\nLieferantenabfragen: Jährlich per SCIP-Datenbank\nEinkaufsverantwortlicher REACH: Stefan Braun"),
    # More IP/Verträge
    ("IP_NDA_Entwicklungspartner_RWTH_2023", "07_IP_Lizenzen",
     f"Geheimhaltungsvereinbarung (NDA)\n\nZwischen: {C['name']} und RWTH Aachen, Institut für Produktionstechnik\nGegenstand: Gemeinsames F&E-Projekt 'SmartPress 4.0'\nGeheimhaltungsfrist: 5 Jahre ab Vertragsunterzeichnung\nSchutzumfang: Alle technischen Informationen, Know-how, Prototypen\nDatum: {date_str(2023, 2, 15)}"),
    ("KD_Rahmenvertrag_ThyssenKrupp_2023", "04_Vertraege_Kunden",
     f"Rahmenliefervertrag 2023–2025\n\nAuftraggeber: {C['kunde1']}\nLieferant: {C['name']}\nGeltungszeitraum: 01.01.2023 – 31.12.2025\nJahresvolumen (Mindestabnahme): {fmt_eur(7_000_000)}\nProdukte: {C['produkt1']}, {C['produkt3']}\nPreisanpassung: Halbjährlich gemäß vereinbartem Stahl-Preisindex\nExklusivität: Nicht vereinbart"),
    ("LF_Rahmenvertrag_Einkauf_Allgemein_2023", "05_Vertraege_Lieferanten",
     f"Allgemeine Einkaufsbedingungen (AEB) {C['name']}, Stand: {date_str(2023, 1, 1)}\n\nGelten für alle Bestellungen der {C['name']} gegenüber Lieferanten.\n§ 1 Geltungsbereich: Diese AEB gelten exklusiv; Lieferantenbedingungen werden nicht anerkannt.\n§ 2 Angebot/Auftrag: Bestellungen bedürfen der Schriftform.\n§ 3 Lieferfristen: Verbindlich; Vertragsstrafe 0,5 % pro Woche, max. 5 %.\n§ 4 Qualitätssicherung: Wareneingangskontrolle, Rüge-Pflicht 14 Tage.\n§ 5 Gewährleistung: 24 Monate Gewährleistungsfrist."),
    ("IMM_Versicherungskataster_Immobilien_2024", "06_Immobilien",
     f"Versicherungskataster Immobilien {C['name']}, Stand: Januar 2024\n\nGrundstück/Gebäude Industriestr. 12: Versicherungswert {fmt_eur(25_000_000)} (Neuwert), Versicherer: {C['versicherer2']}\nAußenlager Niehler Str.: Versicherungswert {fmt_eur(3_200_000)}, Versicherer: {C['versicherer2']}\nMaschinen gesamt: Versicherungswert {fmt_eur(8_000_000)}, Versicherer: {C['versicherer2']}\nGesamtversicherungswert Sachanlagen: {fmt_eur(36_200_000)}"),
    # Final bulk: more financial records
    ("FIN_Umsatzsteuervoranmeldung_Q4_2023", "02_Finanzen",
     f"Umsatzsteuervoranmeldung Q4 2023\n\nUnternehmen: {C['name']}, USt-IdNr.: {C['ust_id']}\nZeitraum: Oktober – Dezember 2023\nUmsätze 19 % MwSt.: {fmt_eur(C['revenue_2023']//4)}\nVorsteuer: {fmt_eur(int(C['revenue_2023']//4 * 0.42 * 0.19))}\nZahllast: {fmt_eur(int(C['revenue_2023']//4 * 0.19 * 0.58))}\nFälligkeit: 10. Januar 2024"),
    ("FIN_Gewinnverteilungsplan_2023", "02_Finanzen",
     f"Gewinnverteilungsplan Geschäftsjahr 2023\n\nJahresüberschuss 2023: {fmt_eur(int(C['ebit_2023']*0.65))}\nEinstellung Gewinnrücklage: {fmt_eur(1_500_000)}\nAusschüttung an Gesellschafter: {fmt_eur(int(C['ebit_2023']*0.65) - 1_500_000)}\n  davon {C['gesellschafter1']} ({C['anteil1']}): {fmt_eur(int((int(C['ebit_2023']*0.65) - 1_500_000) * 0.6))}\n  davon {C['gesellschafter2']} ({C['anteil2']}): {fmt_eur(int((int(C['ebit_2023']*0.65) - 1_500_000) * 0.4))}\nAusschüttungstermin: 30. Juni 2024"),
    ("FIN_Wertpapieranlage_Festgeld_2023", "02_Finanzen",
     f"Festgeldanlage – Bankbeleg\n\nAnlageinstitut: {C['bank']}\nAnleger: {C['name']}\nBetrag: {fmt_eur(2_000_000)}\nLaufzeit: 12 Monate ab {date_str(2023, 7, 1)}\nZinssatz: 3,2 % p.a.\nFälligkeit: {date_str(2024, 6, 30)}\nZinsen: {fmt_eur(int(2_000_000 * 0.032))}"),
    # More strategic
    ("STRAT_Pressemitteilung_Auftrag_ThyssenKrupp_2023", "11_Strategie_Planung",
     f"Pressemitteilung - {date_str(2023, 8, 22)}\n\n{C['name']} erhaelt Grossauftrag von {C['kunde1']}\n\n{C['city']} - Die {C['name']} hat einen Auftrag ueber die Lieferung einer vollautomatischen Pressenlinie PL-500 von {C['kunde1']} erhalten. Das Auftragsvolumen belaeuft sich auf {fmt_eur(2_300_000)}. Die Anlage soll im zweiten Quartal 2024 in Betrieb genommen werden.\n\nDieser Auftrag bestaetigt unsere Staerke als Systemlieferant fuer die Stahlindustrie, sagt {C['gf1_name']}, Geschaeftsfuehrer."),
    ("STRAT_Roadmap_Internationalisierung_2025", "11_Strategie_Planung",
     f"Internationalisierungsroadmap 2025–2027\n\nIST-Situation: >95 % DACH-Markt\nZiel 2027: 15 % Umsatz international\n\nPhase 1 (2025): Warschau – Vertriebsbüro eröffnet Q3 2024\nPhase 2 (2025/26): Prag, Wien – Handelsagenten kontraktiert\nPhase 3 (2026/27): Mailand – Kooperation mit lokalem Maschinenbauunternehmen\n\nBudget Internationalisierung: {fmt_eur(350_000)} über 3 Jahre\nZielkunden international: Europäische Automotive-Tier-1-Zulieferer"),
    ("STRAT_Digitalisierungsbericht_Q2_2024", "11_Strategie_Planung",
     f"Digitalisierungsbericht Q2 2024\n\nProjektleiter: Tobias Lange (IT) / Sabine Koch (F&E)\n\nFortschritt laufende Projekte:\n1. Digitaler Zwilling PL-500: 70 % abgeschlossen, Fertigstellung Q4 2024\n2. IoT-Sensorik Produktion: 24 von 32 Maschinen angebunden\n3. SAP S/4HANA 2024 Upgrade: Go-Live verschoben auf Q1 2025 wegen Ressourcenknappheit\n4. KI-Qualitätskontrolle (Vision AI): Pilotphase mit {C['kunde2']} läuft\n\nBudget verbraucht: {fmt_eur(820_000)} von {fmt_eur(1_800_000)} genehmigt"),
    # Extra compliance bulk
    ("COMP_Energiemanagementsystem_DIN_EN_ISO_50001", "09_Compliance",
     f"Energiemanagementsystem nach DIN EN ISO 50001\n\nZertifizierung angestrebt bis Q2 2025 im Rahmen der Nachhaltigkeitsstrategie.\nEnergieverbrauch 2023: 8.400 MWh (Strom: 6.200 MWh, Gas: 2.200 MWh)\nEnergie-Einsparziel 2025 vs. 2022: -15 %\nMaßnahmen: LED-Umrüstung Hallen (abgeschlossen), Wärmerückgewinnung Kompressoren (in Planung), Photovoltaik 250 kWp (Genehmigung erteilt, Bau Q3 2024)"),
    ("COMP_Arbeitsschutz_Gefaehrdungsbeurteilung_2023", "09_Compliance",
     f"Gefährdungsbeurteilung gemäß § 5 ArbSchG – Produktion 2023\n\nErstellt: Michael Weber (Leiter Produktion) / Fachkraft für Arbeitssicherheit: Dipl.-Ing. Klaus Richter (extern)\nDatum: {date_str(2023, 10, 5)}\n\nIdentifizierte Gefährdungen (Auswahl):\n• Lärm > 85 dB(A) an 6 Arbeitsplätzen → Lärmschutzhelme, PSA\n• Heben > 15 kg: Hebehilfen installiert\n• Elektromagnetische Felder Laserschneider: Sicherheitsabstände markiert\n\nKeine schwerwiegenden Gefährdungslagen identifiziert."),
    # Remaining extras to exactly 300
    ("HR_Altersstrukturanalyse_2024", "03_Personal_HR",
     f"Altersstrukturanalyse Belegschaft 2024\n\nGesamt: {C['employees_2023']} Mitarbeiter\nDurchschnittsalter: 41,3 Jahre\n\nAltersgruppen:\n< 25 Jahre: 18 (7,3 %) – Auszubildende und Berufseinsteiger\n25–35 Jahre: 52 (21,1 %) – Berufseinsteiger, Spezialisten\n36–45 Jahre: 87 (35,2 %) – erfahrene Mitarbeiter, Kernbelegschaft\n46–55 Jahre: 68 (27,5 %) – Senior-Experten\n> 55 Jahre: 22 (8,9 %) – davon 8 in den nächsten 5 Jahren im Rentenalter\n\nNachfolgeplanung: Für 3 Schlüsselpositionen (Produktionsleiter, Leiter QS, FuE) laufen Nachfolgevorbereitungen."),
    ("LF_Rahmenvertrag_Logistik_DB_Schenker", "05_Vertraege_Lieferanten",
     f"Logistikvertrag\n\nAuftragnehmer: DB Schenker Deutschland AG\nAuftraggeberin: {C['name']}\nLeistungsumfang: Nationale und internationale Transportlogistik für Maschinen und Anlagen\nVersandvolumen p.a.: ca. 180 Sendungen, davon 12 % Speditionsverkehr international\nVergütung: Einzelpreise nach Tariftabelle, Jahresrabatt 8 % ab 150 Sendungen\nHaftung: CMR, max. 8,33 SDR/kg Schadensfall"),
    ("KD_SLA_After_Sales_Service_Bosch_2023", "04_Vertraege_Kunden",
     f"After-Sales-Servicevertrag\n\nKunde: {C['kunde2']}\nDienstleister: {C['name']}\nMaschinenpark: 3 Presslinien PL-500 (Lieferjahre 2019, 2021, 2022)\nServiceumfang: Präventive Wartung 2×/Jahr, Telefonhotline 24/7, Reaktionszeit Notfall: 6 h\nJahrespauschale: {fmt_eur(88_000)}\nLaufzeit: 01.01.2023 – 31.12.2025"),
    ("GR_Vollmacht_Rechtsanwalt_Heuking", "01_Gesellschaftsrecht",
     f"Prozessvollmacht\n\nDie {C['name']}, vertreten durch die Geschäftsführer {C['gf1_name']} und {C['gf2_name']}, erteilt hiermit {C['anwalt']}, die Vollmacht zur Vertretung in allen rechtlichen Angelegenheiten einschließlich Prozessführung, Vergleichsabschluss und außergerichtlicher Verhandlung.\n\nKöln, den {date_str(2020, 3, 1)}\n{C['name']}\n{C['gf1_name']}"),
    ("VS_Versicherungsübersicht_Gesamt_2024", "08_Versicherungen",
     f"Versicherungsübersicht {C['name']} – Stand: Januar 2024\n\nVersicherungsart | Versicherer | Versicherungssumme | Jahresprämie\nBetriebshaftpflicht | {C['versicherer1']} | {fmt_eur(10_000_000)} | {fmt_eur(38_500)}\nProduktehaftpflicht | {C['versicherer1']} | {fmt_eur(15_000_000)} | {fmt_eur(52_000)}\nGebäude | {C['versicherer2']} | {fmt_eur(25_000_000)} | {fmt_eur(44_200)}\nMaschinen | {C['versicherer2']} | {fmt_eur(8_000_000)} | {fmt_eur(28_500)}\nD&O | {C['versicherer1']} | {fmt_eur(5_000_000)} | {fmt_eur(24_000)}\nCyber | {C['versicherer1']} | {fmt_eur(3_000_000)} | {fmt_eur(18_500)}\nTransport | {C['versicherer1']} | {fmt_eur(2_000_000)} | {fmt_eur(12_800)}\nGesamt-Jahresprämien: {fmt_eur(218_500)}"),
    ("COMP_UVV_Prüfprotokoll_Elektro_2023", "09_Compliance",
     f"UVV-Prüfprotokoll elektrische Betriebsmittel gemäß BGV A3\n\nPrüfer: Bruni Elektroanlagen GmbH\nPrüfdatum: {date_str(2023, 11, 20)}\nPrüfumfang: 847 elektrische Betriebsmittel (Maschinen, Werkzeuge, Verlängerungskabel)\nErgebnis: 839 i.O., 8 defekt → gesperrt und zur Entsorgung freigegeben\nNächste Prüfung: {date_str(2024, 11, 15)}"),
]

for fname, folder, content in wave3_docs:
    if len(generated) >= 300: break
    doc = Document()
    make_docx_header(doc, fname.replace("_", " "))
    doc.add_paragraph(content)
    generated.append(save_docx(doc, folder, f"{fname}.docx"))

print(f"  Wave 3 complete. Total: {len(generated)} documents")

# Final count
generated = [g for g in generated if g is not None]
print(f"\n✅ DONE — {len(generated)} documents generated in {BASE}")
print("\nDocument breakdown by folder:")
for folder in sorted(BASE.iterdir()):
    if folder.is_dir():
        files = list(folder.iterdir())
        docx_count = sum(1 for f in files if f.suffix == '.docx')
        pdf_count = sum(1 for f in files if f.suffix == '.pdf')
        xlsx_count = sum(1 for f in files if f.suffix == '.xlsx')
        print(f"  {folder.name}: {len(files)} files (docx:{docx_count}, pdf:{pdf_count}, xlsx:{xlsx_count})")
