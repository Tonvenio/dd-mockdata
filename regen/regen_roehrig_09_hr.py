"""
regen_roehrig_09_hr.py
Enrich thin .docx files in roehrig_large/09_Personal_HR.

Strategy:
  - For each thin (<200w) docx, read its existing entity-header (line 1+2)
    and main title (level-2 heading) so we keep consistent context.
  - Route by filename prefix to a type-specific generator.
  - Re-write with realistic body content.
"""
# --- portable-paths-prelude --- (do not edit) ---
import sys
from pathlib import Path as _PathlibPath
_RP = _PathlibPath(__file__).resolve().parent
while not (_RP / "enhance_lib.py").exists() and _RP.parent != _RP:
    _RP = _RP.parent
_ROOT = _RP
sys.path.insert(0, str(_ROOT))
# --- end prelude ---

import sys, re
from pathlib import Path

sys.path.insert(0, f"{_ROOT}")
from enhance_lib import ROEHRIG_AG as R, write_doc, signatures, ds

from docx import Document

BASE = Path(f"{_ROOT}/roehrig_large/09_Personal_HR")

# ---- entity header lookup (name -> (addr, hrb) split) -----------------------
HEADERS = {
    "Brennhagen Elektronik AG":      ("Vaihinger Strasse 120, 70567 Stuttgart", "HRB 726451, AG Stuttgart"),
    "Brennhagen Elektronik AG":       ("Vaihinger Straße 120, 70567 Stuttgart", "HRB 726451, AG Stuttgart"),
    "Brennhagen Elektronik GmbH":     ("Heilbronner Strasse 88, 74072 Heilbronn", "HRB 221456, AG Heilbronn"),
    "Brennhagen Software GmbH":       ("Riesstrasse 12, 80992 Muenchen", "HRB 319872, AG Muenchen"),
    "Brennhagen Polska sp. z o.o.":   ("ul. Roosevelta 18, 40-008 Katowice (PL)", "KRS 0000412876"),
    "Brennhagen Czech s.r.o.":        ("Slatina Industrial Park, 627 00 Brno (CZ)", "IČO: 28418762"),
    "Brennhagen Hungary Kft.":        ("Ipari park 4, 9027 Győr (HU)", "Cg. 08-09-025671"),
    "Brennhagen Electronics (China) Ltd.": ("No. 188, Songze Road, Qingpu, Shanghai (CN)",
                                        "Unified Social Credit: 91310115MA1FL42Q38"),
    "Brennhagen Holding GmbH":        ("Vaihinger Strasse 120, 70567 Stuttgart", "HRB 726450, AG Stuttgart"),
    "Freshfields Bruckhaus Deringer LLP": ("Bockenheimer Anlage 44, 60322 Frankfurt am Main",
                                           "Partnerschaftsregister PR 1234"),
}

# Map to canonical short codes (for signatures/Geschaeftsfuehrer)
GF = {
    "Brennhagen Elektronik AG":  ("Anna Mueller", "Vorstandsvorsitzende"),
    "Brennhagen Elektronik GmbH":("Andreas Maier", "Geschaeftsfuehrer Werk Heilbronn"),
    "Brennhagen Software GmbH":  ("Dr. Klaus Kessler", "Geschaeftsfuehrer RSG"),
    "Brennhagen Polska sp. z o.o.": ("Marek Wojciechowski", "Prezes Zarzadu"),
    "Brennhagen Czech s.r.o.":   ("Petr Novák", "Jednatel"),
    "Brennhagen Hungary Kft.":   ("László Kovács", "Ügyvezető"),
    "Brennhagen Electronics (China) Ltd.": ("Zhang Hao", "General Manager Shanghai"),
    "Brennhagen Holding GmbH":   ("Laura Bauer", "CFO / Geschaeftsfuehrerin RHO"),
}

CITY = {
    "Brennhagen Elektronik AG":  "Stuttgart",
    "Brennhagen Elektronik GmbH":"Heilbronn",
    "Brennhagen Software GmbH":  "Muenchen",
    "Brennhagen Polska sp. z o.o.": "Katowice",
    "Brennhagen Czech s.r.o.":   "Brno",
    "Brennhagen Hungary Kft.":   "Györ",
    "Brennhagen Electronics (China) Ltd.": "Shanghai",
    "Brennhagen Holding GmbH":   "Stuttgart",
}


def word_count(p):
    d = Document(p)
    t = " ".join(par.text for par in d.paragraphs)
    for tbl in d.tables:
        for r in tbl.rows:
            for c in r.cells:
                t += " " + c.text
    return len(t.split())


def read_meta(p):
    """Extract (entity_name, addr, hrb, title, confidential, draft) from existing thin doc."""
    d = Document(p)
    paras = [par.text.strip() for par in d.paragraphs if par.text.strip()]
    confidential = any("VERTRAULICH" in s for s in paras[:3])
    draft = any("ENTWURF" in s or "DRAFT" in s for s in paras[:3])
    # Strip flag lines
    body = [s for s in paras
            if "VERTRAULICH" not in s and "ENTWURF" not in s and "DRAFT" not in s]
    name = body[0] if body else "Brennhagen Elektronik AG"
    addr_hrb = body[1] if len(body) > 1 else ""
    title = ""
    # Find the heading-2 = title (skip the separator line of dashes)
    for par in d.paragraphs:
        if par.style.name.startswith("Heading 2") and par.text.strip():
            title = par.text.strip()
            break
    if not title and len(body) > 3:
        title = body[3] if "─" in body[2] else body[2]
    # Split addr / hrb
    if "  |  " in addr_hrb:
        addr, hrb = addr_hrb.split("  |  ", 1)
    else:
        addr, hrb = addr_hrb, ""
    return name, addr, hrb, title, confidential, draft


def header_for(name, addr, hrb):
    if name in HEADERS:
        a, h = HEADERS[name]
        return {"name": name, "addr": a, "hrb": h}
    return {"name": name, "addr": addr or "—", "hrb": hrb or "—"}


# ---------- AV (Arbeitsvertrag) ---------------------------------------------

ROLE_PROFILE = {
    "Vorstandsvorsitzende": ("E1", 350_000, 60, 30, "Vorstand"),
    "Vorstandsvorsitzen":   ("E1", 340_000, 60, 30, "Vorstand"),
    "Chief Financial Officer":  ("E1", 295_000, 50, 30, "Vorstand"),
    "Chief Financial Offi":     ("E1", 295_000, 50, 30, "Vorstand"),
    "Chief Operating Officer":  ("E1", 285_000, 50, 30, "Vorstand"),
    "Chief Operating Off":      ("E1", 285_000, 50, 30, "Vorstand"),
    "Chief Operating Of":       ("E1", 285_000, 50, 30, "Vorstand"),
    "Chief Technology Officer": ("E1", 280_000, 50, 30, "Vorstand"),
    "Chief Technology Off":     ("E1", 280_000, 50, 30, "Vorstand"),
    "Chief Technology O":       ("E1", 280_000, 50, 30, "Vorstand"),
    "VP Sales Automotive":      ("E2", 185_000, 40, 30, "Sales"),
    "VP Sales Automotiv":       ("E2", 185_000, 40, 30, "Sales"),
    "VP Engineering":           ("E2", 178_000, 40, 30, "Engineering"),
    "Head of Quality IATF":     ("E3", 142_000, 25, 30, "Quality"),
    "Head of Quality IA":       ("E3", 142_000, 25, 30, "Quality"),
    "Senior Project Manager":   ("AT1", 108_000, 20, 30, "PMO"),
    "Senior Project Man":       ("AT1", 108_000, 20, 30, "PMO"),
    "Lead Software Engineer":   ("AT1", 102_000, 20, 30, "Software"),
    "Lead Software Engi":       ("AT1", 102_000, 20, 30, "Software"),
    "Principal Hardware Engineer": ("AT1", 105_000, 20, 30, "Hardware"),
    "Principal Hardware E":     ("AT1", 105_000, 20, 30, "Hardware"),
    "Principal Hardware":       ("AT1", 105_000, 20, 30, "Hardware"),
    "Senior Controller":        ("AT2",  92_000, 18, 30, "Finance / Controlling"),
    "HR Business Partner":      ("AT2",  88_000, 18, 30, "Human Resources"),
    "HR Business Partne":       ("AT2",  88_000, 18, 30, "Human Resources"),
    "Supply Chain Manager":     ("AT2",  94_000, 18, 30, "Supply Chain"),
    "Supply Chain Manag":       ("AT2",  94_000, 18, 30, "Supply Chain"),
    "Key Account Manager":      ("AT2",  98_000, 18, 30, "Sales"),
    "Key Account Manage":       ("AT2",  98_000, 18, 30, "Sales"),
    "Production Manager":       ("AT2",  86_000, 15, 30, "Produktion"),
    "Quality Engineer":         ("AT3",  72_000, 12, 30, "Quality"),
    "Software Engineer":        ("AT3",  78_000, 12, 30, "Software"),
    "Embedded Software Developer": ("AT3", 82_000, 12, 30, "Embedded SW"),
    "Embedded Software De":     ("AT3",  82_000, 12, 30, "Embedded SW"),
    "Embedded Software":        ("AT3",  82_000, 12, 30, "Embedded SW"),
    "ADAS Systems Engineer":    ("AT3",  85_000, 12, 30, "ADAS"),
    "ADAS Systems Enginee":     ("AT3",  85_000, 12, 30, "ADAS"),
    "ADAS Systems Engin":       ("AT3",  85_000, 12, 30, "ADAS"),
    "Test Engineer":            ("AT3",  68_000, 12, 30, "Testing"),
    "Purchasing Manager":       ("AT2",  84_000, 15, 30, "Einkauf"),
    "Financial Analyst":        ("AT3",  72_000, 12, 30, "Finance"),
    "Legal Counsel":            ("AT2",  98_000, 15, 30, "Legal"),
    "IP Manager":               ("AT2",  92_000, 15, 30, "Legal / IP"),
    "Compliance Officer":       ("AT2",  92_000, 15, 30, "Compliance"),
    "IT Security Analyst":      ("AT3",  78_000, 12, 30, "IT Security"),
    "IT Security Analys":       ("AT3",  78_000, 12, 30, "IT Security"),
    "Customer Quality Manager": ("AT2",  88_000, 15, 30, "Quality"),
    "Customer Quality Man":     ("AT2",  88_000, 15, 30, "Quality"),
    "Customer Quality M":       ("AT2",  88_000, 15, 30, "Quality"),
    "Process Engineer":         ("AT3",  74_000, 12, 30, "Manufacturing"),
    "Logistics Manager":        ("AT2",  82_000, 15, 30, "Logistik"),
}


def parse_av_filename(fname):
    """Return (employee_name, role_string)."""
    stem = fname[:-5]  # strip .docx
    # remove suffixes
    stem = re.sub(r"_(ENTWURF|DRAFT|WIP|FINAL|ALT|v\d+|rev_[A-Za-z]+|\d{4}-\d{2}-\d{2}|\d{8})$", "", stem)
    # AV_<num>_<name parts>_<role parts>
    m = re.match(r"AV_(\d+)_(.+)", stem)
    if not m:
        return "Mitarbeiter/in", "Mitarbeiter/in"
    rest = m.group(2)
    # try matching a role at end (longest match)
    for role in sorted(ROLE_PROFILE.keys(), key=len, reverse=True):
        rkey = role.replace(" ", "_")
        if rest.endswith("_" + rkey):
            name = rest[:-(len(rkey)+1)].replace("_", " ")
            # strip trailing digit suffix tag from name (e.g., "Ute Krämer 61")
            return name, role
    # fallback: split at last 2 words
    parts = rest.split("_")
    if len(parts) >= 3:
        return " ".join(parts[:2]), " ".join(parts[2:])
    return rest.replace("_", " "), "Mitarbeiter/in"


def arbeitsvertrag(path):
    name_co, addr, hrb, title, conf, draft = read_meta(path)
    hdr = header_for(name_co, addr, hrb)
    emp_name, role = parse_av_filename(path.name)
    profile = ROLE_PROFILE.get(role, ("AT3", 72_000, 12, 30, "Operations"))
    band, salary, var_pct, urlaub, abteilung = profile

    sal_str = f"{salary:,} EUR".replace(",", ".")
    var_str = f"{var_pct} %"

    # number for the AV
    m = re.match(r"AV_(\d+)_", path.name)
    av_no = m.group(1) if m else "—"

    is_vorstand = role.startswith(("Vorstandsvorsitzen", "Chief"))
    is_at = band.startswith("AT") or band.startswith("E")

    start_year = 2022 if int(av_no) <= 30 else (2023 if int(av_no) <= 80 else 2024)
    start_date = ds(start_year, ((int(av_no) - 1) % 12) + 1, ((int(av_no) - 1) % 27) + 1)

    sections = [
        ("Vertragsparteien",
         f"Zwischen der {name_co}, {CITY.get(name_co,'Stuttgart')}, eingetragen im "
         f"{hdr['hrb']}, vertreten durch {GF.get(name_co,('Anna Mueller','Vorstand'))[0]}, "
         f"{GF.get(name_co,('Anna Mueller','Vorstand'))[1]} (nachfolgend »Gesellschaft«), "
         f"und Herrn/Frau {emp_name} (nachfolgend »Arbeitnehmer/in«) wird der folgende "
         f"Anstellungsvertrag geschlossen. Personalakte: AV-{av_no}."),
        ("§ 1 Beginn und Aufgabe",
         ("clauses", [
             ("§ 1 Beginn und Aufgabe", [
                 f"Das Anstellungsverhaeltnis beginnt am {start_date}. Es ist auf unbestimmte Zeit geschlossen.",
                 f"Der/Die Arbeitnehmer/in wird als {role} in der Abteilung {abteilung} eingestellt. "
                 f"Direkter Vorgesetzter: Leitung {abteilung}.",
                 "Die Probezeit betraegt sechs Monate. Waehrend der Probezeit gilt eine Kuendigungsfrist "
                 "von zwei Wochen zum Monatsende. Die Aufgaben koennen im Rahmen des Direktionsrechts "
                 "angepasst werden, soweit Qualifikation und vereinbarter Aufgabenbereich gewahrt bleiben.",
             ]),
         ])),
        ("§ 2 Verguetung",
         ("clauses", [
             ("§ 2 Verguetung", [
                 f"Das Jahresbruttofixgehalt betraegt {sal_str}, zahlbar in zwoelf gleichen Monatsraten "
                 "am Monatsende auf das vom/von der Arbeitnehmer/in angegebene Konto.",
                 f"Zusaetzlich wird eine variable Verguetung von bis zu {var_str} des Jahresfixgehalts "
                 "bei vollstaendiger Zielerfuellung gewaehrt. Naeheres regelt die Zielvereinbarung "
                 "(MBO) gemaess STIP 2023.",
                 ("Als Vorstands-/leitende Person gilt: Teilnahme am Long-Term Incentive Plan (LTIP 2022-2024) "
                  "der Brennhagen Elektronik AG (Performance Share Units, 4-Jahres-Vesting, Performance-Caps "
                  "ROCE und TSR vs. STOXX Europe 600 Auto & Parts). Ein gesonderter LTIP-Bestaetigungsbrief "
                  "wird ausgestellt." if is_vorstand else
                  "Eine Einbeziehung in das LTIP-Programm der Brennhagen Elektronik AG erfolgt nur auf "
                  "ausdruecklichen Einzelbeschluss des Aufsichtsratspraesidialausschusses. Im uebrigen "
                  "gelten die einschlaegigen Konzernrichtlinien zur variablen Verguetung."),
                 f"Eintaktung in Verguetungsband {band}. Die Verguetung wird jaehrlich im Rahmen der "
                 "Konzern-Gehaltsrunde (April) ueberprueft.",
             ]),
         ])),
        ("§ 3 Arbeitszeit und Arbeitsort",
         ("clauses", [
             ("§ 3 Arbeitszeit und Arbeitsort", [
                 f"Dienstort ist {CITY.get(name_co,'Stuttgart')}. Eine voruebergehende Versetzung an einen "
                 "anderen Standort des Konzerns (insbesondere Heilbronn, Stuttgart, Muenchen, Katowice, "
                 "Brno, Györ, Shanghai) ist im Rahmen der Vertretbarkeit moeglich.",
                 ("Es gilt Vertrauensarbeitszeit gemaess Konzern-Betriebsvereinbarung 2022. Eine feste "
                  "woechentliche Arbeitszeit ist nicht vereinbart; die Aufgabenerfuellung steht im "
                  "Vordergrund." if is_at else
                  "Die regelmaessige woechentliche Arbeitszeit betraegt 35 Stunden gemaess IGM-Tarifvertrag "
                  "Metall- und Elektroindustrie Baden-Wuerttemberg. Ueberstunden werden gemaess BV "
                  "Arbeitszeit erfasst und durch Freizeitausgleich oder Verguetung kompensiert."),
                 "Homeoffice ist nach Massgabe der Konzern-Betriebsvereinbarung »Homeoffice/Mobiles Arbeiten« "
                 "2022 grundsaetzlich an bis zu drei Tagen pro Woche moeglich, soweit betriebliche Belange "
                 "nicht entgegenstehen.",
             ]),
         ])),
        ("§ 4 Urlaub, Krankheit, Sonstiges",
         ("clauses", [
             ("§ 4 Urlaub, Krankheit, Sonstiges", [
                 f"Der Jahresurlaub betraegt {urlaub} Arbeitstage bei 5-Tage-Woche.",
                 "Im Krankheitsfall ist die Gesellschaft unverzueglich, spaetestens vor Schichtbeginn, "
                 "zu informieren. Eine Arbeitsunfaehigkeitsbescheinigung ist ab dem dritten Krankheitstag "
                 "elektronisch zu uebermitteln (eAU-Verfahren).",
                 "Es gelten die jeweils gueltigen Konzern-Betriebsvereinbarungen, insbesondere zu "
                 "betrieblicher Altersversorgung (Direktzusage), Jubilaeumsregelung, BEM (Betriebliches "
                 "Eingliederungsmanagement) und Mobilitaetszuschuss OePNV.",
             ]),
         ])),
        ("§ 5 Verschwiegenheit, Nebentaetigkeit, Wettbewerb",
         ("clauses", [
             ("§ 5 Verschwiegenheit, Nebentaetigkeit, Wettbewerb", [
                 "Der/Die Arbeitnehmer/in ist verpflichtet, ueber alle ihm/ihr im Rahmen der Taetigkeit "
                 "bekannt gewordenen Geschaefts- und Betriebsgeheimnisse Stillschweigen zu bewahren - "
                 "auch nach Beendigung des Arbeitsverhaeltnisses.",
                 ("Fuer die Dauer von 24 Monaten nach Beendigung des Vertrages besteht ein vertragliches "
                  "Wettbewerbsverbot fuer den Bereich Automotive-Elektronik / Embedded Systems / ADAS, "
                  "gegen Karenzentschaedigung in Hoehe von 50 % der zuletzt bezogenen "
                  "Festbezuege (§ 74 ff. HGB)." if is_vorstand else
                  "Die Aufnahme einer entgeltlichen Nebentaetigkeit bedarf der vorherigen schriftlichen "
                  "Zustimmung der Gesellschaft. Ein Wettbewerbsverbot kann fuer bis zu 12 Monate "
                  "gegen angemessene Karenzentschaedigung vereinbart werden."),
                 "Datenschutz: Die Verarbeitung personenbezogener Daten des/der Arbeitnehmer/in erfolgt "
                 "ausschliesslich gemaess Art. 88 DSGVO i. V. m. § 26 BDSG sowie der Konzern-BV »Datenschutz "
                 "und IT-Systeme« vom 12. Dezember 2022.",
             ]),
         ])),
        ("§ 6 Beendigung",
         ("clauses", [
             ("§ 6 Beendigung", [
                 "Das Anstellungsverhaeltnis kann nach Ablauf der Probezeit von beiden Seiten unter Einhaltung "
                 "der gesetzlichen Fristen gemaess § 622 BGB ordentlich gekuendigt werden.",
                 ("Das Anstellungsverhaeltnis endet automatisch mit Ablauf des Monats, in dem der/die "
                  "Arbeitnehmer/in die Regelaltersgrenze der gesetzlichen Rentenversicherung erreicht. "
                  "Ausserordentliche Kuendigungsrechte bleiben unberuehrt."),
                 "Bei Beendigung sind alle Unterlagen, Datentraeger, Pruefmittel, Werksausweise sowie "
                 "Konzern-Hardware (Laptop, Mobiltelefon) unverzueglich an HR/IT zurueckzugeben.",
             ]),
         ])),
        ("§ 7 Schlussbestimmungen",
         "Es gilt das Recht der Bundesrepublik Deutschland. Gerichtsstand ist - soweit zulaessig - "
         f"{CITY.get(name_co,'Stuttgart')}. Aenderungen und Ergaenzungen dieses Vertrages beduerfen der "
         "Schriftform; das gilt auch fuer die Aufhebung dieser Schriftformklausel. Sollte eine "
         "Bestimmung dieses Vertrages unwirksam sein, beruehrt dies die Wirksamkeit der uebrigen "
         "Bestimmungen nicht."),
        ("Unterschriften",
         signatures(GF.get(name_co, ("Anna Mueller","Vorstand"))[0],
                    GF.get(name_co, ("Anna Mueller","Vorstand"))[1],
                    name_co,
                    emp_name, role, "Arbeitnehmer/in",
                    place=CITY.get(name_co,"Stuttgart"), date_str_=start_date)),
    ]
    write_doc(path, hdr, f"Anstellungsvertrag - {emp_name} - {role}",
              sections, confidential=conf, draft=draft)


# ---------- JG (Jahresgespraech) --------------------------------------------

def parse_jg_filename(fname):
    stem = fname[:-5]
    stem = re.sub(r"_(FINAL|DRAFT|WIP|ALT|v\d+|rev_[A-Za-z]+|\d{4}-\d{2}-\d{2}|\d{8})$", "", stem)
    m = re.match(r"JG_(\d+)_(.+?)_\d+_(\d{4})$", stem) or re.match(r"JG_(\d+)_(.+?)_(\d{4})$", stem)
    if m:
        return m.group(1), m.group(2).replace("_", " "), m.group(3)
    return "—", "Mitarbeiter/in", "2023"


def jahresgespraech(path):
    name_co, addr, hrb, title, conf, draft = read_meta(path)
    hdr = header_for(name_co, addr, hrb)
    num, emp, year = parse_jg_filename(path.name)

    sections = [
        ("1. Allgemeine Angaben",
         [["Feld", "Angabe"],
          ["Mitarbeiter/in", emp],
          ["Personalnummer", f"P-{num}"],
          ["Gesellschaft", name_co],
          ["Standort", CITY.get(name_co, "Stuttgart")],
          ["Vorgesetzte/r", "Abteilungsleitung (gemaess Organigramm)"],
          ["Gespraechstermin", f"{ds(int(year), 11, 24)}"],
          ["Berichtszeitraum", f"01.01.{year} – 31.12.{year}"],
          ["HR Business Partner", "Prof. Renate Meyer / Bjoern Franke"]]),
        ("2. Rueckblick Vorjahr / Zielerreichung",
         f"Im Berichtszeitraum {year} hat der/die Mitarbeiter/in die vereinbarten Ziele ueberwiegend "
         "erreicht. Die quantitativen Ziele (Liefertreue, Qualitaetskennzahlen, Projektmeilensteine) "
         "wurden zu durchschnittlich 105 % erfuellt. Bei qualitativen Zielen (Teamentwicklung, "
         "Wissenstransfer, kontinuierliche Verbesserung) wurden die Erwartungen ebenfalls erfuellt. "
         "Eine besondere Leistung zeigte sich bei der Beteiligung an konzernweiten Projekten zur "
         "ADAS- und BMS-Plattform sowie bei der Umsetzung der IATF 16949-Anforderungen.\n\n"
         "Gesamtbewertung Zielerreichung: 108 % (entspricht Faktor 1,15 fuer STIP-Auszahlung)."),
        ("3. Kompetenzbeurteilung",
         [["Kompetenzfeld", "Bewertung 1-5", "Kommentar"],
          ["Fachliche Expertise", "4", "Sehr gutes Fachwissen, geht ueber Erwartungen hinaus"],
          ["Loesungsorientierung", "4", "Pragmatisch, Eigenstaendig im Problemloesen"],
          ["Kommunikation", "4", "Klar, adressatengerecht, auch in Englisch"],
          ["Team-/Stakeholder-Mgmt", "3", "Solide, leichter Entwicklungsbedarf"],
          ["Veraenderungsbereitschaft", "4", "Treibt Verbesserungen aktiv"],
          ["Compliance-Bewusstsein", "5", "Konsequent regelkonform, Vorbildwirkung"]]),
        ("4. Entwicklungsziele und Massnahmen "+(year if year else ""),
         ("list", [
             "Teilnahme an Konzern-Programm »Future Leaders« (Mercer-Curriculum, Q2-Q4 nachfolgendes Jahr).",
             "Vertiefung Funktionale Sicherheit ISO 26262 ASIL-D (interne Schulung RSG Muenchen).",
             "Mentoring eines Junior-Mitarbeiters (Q1 Onboarding-Begleitung).",
             "Englisch C1-Zertifikat (Telc) bis 30.06. nachfolgendes Jahr - extern bezuschusst.",
         ])),
        ("5. Verguetungs- und Karriereperspektive",
         f"Eine Anpassung des Grundgehalts wurde im Rahmen der Gehaltsrunde {int(year)+1} besprochen "
         "und dem HR Komitee vorgelegt (siehe Gehaltsaenderungsmitteilung). Die naechste regulaere "
         "Karriereperspektive eroeffnet sich nach Abschluss der Entwicklungsmassnahmen voraussichtlich "
         "in den naechsten 18-24 Monaten (Senior-/Lead-Position bzw. Aufnahme in Talent-Pool »Mitte«)."),
        ("6. Unterschriften",
         signatures(emp, "Mitarbeiter/in", name_co,
                    "Prof. Renate Meyer", "HR Business Partner", "Brennhagen Elektronik AG",
                    place=CITY.get(name_co,"Stuttgart"), date_str_=ds(int(year), 11, 24))),
    ]
    write_doc(path, hdr, f"Jahresgespraech {year} - {emp}", sections,
              confidential=conf, draft=draft)


# ---------- MBO (Zielvereinbarung) ------------------------------------------

def parse_mbo_filename(fname):
    stem = fname[:-5]
    m = re.match(r"MBO_Zielvereinbarung_(\d{4})_(\d+)_(.+)$", stem)
    if m:
        return m.group(1), m.group(2), m.group(3).replace("_", " ")
    return "2023", "—", "Mitarbeiter/in"


def mbo(path):
    name_co, addr, hrb, title, conf, draft = read_meta(path)
    hdr = header_for(name_co, addr, hrb)
    year, num, emp = parse_mbo_filename(path.name)

    sections = [
        ("1. Grundlage",
         f"Diese Zielvereinbarung wird zwischen der {name_co} und Herrn/Frau {emp} "
         f"(Personalnummer P-{num}) fuer das Geschaeftsjahr {year} geschlossen. Sie konkretisiert "
         "die variable Verguetung gemaess Anstellungsvertrag § 2 sowie gemaess Short-Term "
         "Incentive Plan (STIP) " + year + " der Brennhagen Elektronik AG. Zielmessung erfolgt zum "
         "31.12. " + year + " auf Basis testierter IST-Daten. Auszahlung im April " + str(int(year)+1) + "."),
        ("2. Ziele und Gewichtung " + year,
         [["#", "Ziel", "Messgroesse / KPI", "Gewichtung", "Zielerreichung 100 %"],
          ["1", "Konzern-EBITDA-Margin", "EBITDA / Umsatz IFRS",          "20 %", ">= 12,1 %"],
          ["2", "Bereichsziel (Funktional)", "Bereichs-KPI (s. Anlage)",  "30 %", "siehe Anlage"],
          ["3", "Projektziel (z.B. BMS-12 SOP)", "Meilenstein-Erreichung","20 %", "100 % on time"],
          ["4", "Kunden-PPM (extern)", "PPM bei OEM-Kunden",              "10 %", "<= 25 PPM"],
          ["5", "ESG / Diversitaet", "Frauenanteil Fuehrung, Energie",    "10 %", "+1,5 PP / -3 %"],
          ["6", "Individuelles Ziel", "Persoenliche Entwicklung",          "10 %", "Erfolg messbar"]]),
        ("3. Auszahlungskurve",
         ("list", [
             "0 % Zielerreichung -> 0 % Bonusauszahlung (Cliff bei < 70 %).",
             "70 % Zielerreichung -> 50 % Auszahlung des Zielbonus.",
             "100 % Zielerreichung -> 100 % Auszahlung des vereinbarten Zielbonus.",
             "150 % oder mehr -> 150 % Cap; daruber hinaus keine weitere Auszahlung.",
             "Auszahlung im April " + str(int(year)+1) + " mit dem Aprilgehalt.",
         ])),
        ("4. Anpassungsklauseln",
         "Bei wesentlichen Veraenderungen der Geschaeftsgrundlage (M&A, Restrukturierung, Force "
         "Majeure, regulatorische Eingriffe) koennen Ziele im gegenseitigen Einvernehmen angepasst "
         "werden. Die Anpassung erfolgt schriftlich und beduerfen der Zustimmung des HR Business "
         "Partners sowie - bei Vorstaenden - des Aufsichtsrats-Praesidialausschusses. "
         "Bei Eintritt einer Compliance-Verfehlung kann der Bonus ganz oder teilweise einbehalten "
         "werden (Malus-Regelung gemaess STIP/LTIP)."),
        ("5. Unterschriften",
         signatures(emp, "Mitarbeiter/in", name_co,
                    "Vorgesetzte/r", "Abteilungsleitung", name_co,
                    place=CITY.get(name_co,"Stuttgart"),
                    date_str_=ds(int(year), 2, 15))),
    ]
    write_doc(path, hdr, f"Zielvereinbarung {year} - {emp}", sections,
              confidential=conf, draft=draft)


# ---------- Gehaltsaenderung -----------------------------------------------

def parse_gehalt_filename(fname):
    stem = fname[:-5]
    stem = re.sub(r"-?(?:_v\d+|_rev_[A-Za-z]+)?$", "", stem)
    m = re.match(r"Gehaltsaenderung_(\d{4})_(\d+)_(.+?)-?$", stem)
    if m:
        return m.group(1), m.group(2), m.group(3).replace("_", " ")
    return "2023", "—", "Mitarbeiter/in"


def gehaltsaenderung(path):
    name_co, addr, hrb, title, conf, draft = read_meta(path)
    hdr = header_for(name_co, addr, hrb)
    year, num, emp = parse_gehalt_filename(path.name)

    # Plausible salaries by personalnummer band
    inum = int(num) if num.isdigit() else 0
    if inum <= 5:
        old, new = 280_000, 295_000
    elif inum <= 12:
        old, new = 155_000, 168_000
    elif inum <= 20:
        old, new = 92_000, 99_000
    else:
        old, new = 72_000, 78_000

    sections = [
        ("Empfaenger",
         f"Herr/Frau {emp}\nPersonalnummer P-{num}\nGesellschaft: {name_co}\n"
         f"Standort: {CITY.get(name_co, 'Stuttgart')}"),
        ("Betreff",
         f"Anpassung Ihrer Verguetung im Rahmen der Gehaltsrunde {year}"),
        ("Mitteilung",
         f"Sehr geehrte/r Herr/Frau {emp},\n\n"
         f"im Rahmen der jaehrlichen Gehaltsueberpruefung {year} freuen wir uns, Ihnen mitteilen "
         "zu koennen, dass Ihr Grundgehalt mit Wirkung zum 1. April " + year + " wie folgt "
         "angepasst wird:\n\n"
         f"Bisheriges Jahresfixgehalt: {old:,} EUR\n".replace(",", ".") +
         f"Neues Jahresfixgehalt: {new:,} EUR\n".replace(",", ".") +
         f"Erhoehung: {((new-old)/old*100):.1f} % entsprechend {new-old:,} EUR brutto p. a.".replace(",", ".")),
        ("Begruendung und Hintergrund",
         f"Die Erhoehung erfolgt im Rahmen der Konzern-Gehaltsrunde {year} und beruecksichtigt "
         "Ihre individuelle Leistung gemaess Jahresgespraech " + str(int(year)-1) + ", die Entwicklung "
         "der relevanten Mercer-Benchmark-Werte fuer Ihre Funktion in der Region D/A/CH sowie "
         "die uebergreifende Verguetungspolitik der Brennhagen Elektronik AG. Die individuelle "
         "Leistung wurde mit »uebertrifft Erwartungen« bewertet."),
        ("Auswirkungen auf weitere Komponenten",
         ("list", [
             "STIP-Zielbonus bleibt prozentual unveraendert (Basis: neues Fixgehalt).",
             "LTIP-Berechtigung unveraendert (sofern bereits vorhanden).",
             "Sozialversicherungsbeitraege werden automatisch angepasst (Beitragsbemessungsgrenze 2024 "
             "beachten).",
             "Direktzusage betriebliche Altersversorgung: ungefoerderte Komponente wird auf neuer Basis "
             "fortgefuehrt.",
         ])),
        ("Bestaetigung",
         f"Diese Mitteilung wird Bestandteil Ihrer Personalakte (P-{num}). Eine Aenderung des "
         "Anstellungsvertrages ist nicht erforderlich; alle anderen Bestimmungen bleiben unveraendert "
         "gueltig.\n\n"
         "Wir bedanken uns ausdruecklich fuer Ihr Engagement und Ihren Beitrag zum Konzern-Erfolg. "
         "Wir wuenschen Ihnen weiterhin Erfolg und Freude bei der Arbeit."),
        ("Unterschriften",
         signatures("Prof. Renate Meyer", "HR Business Partner", name_co,
                    GF.get(name_co, ("Anna Mueller","Vorstand"))[0],
                    GF.get(name_co, ("Anna Mueller","Vorstand"))[1], name_co,
                    place=CITY.get(name_co,"Stuttgart"), date_str_=ds(int(year), 3, 28))),
    ]
    write_doc(path, hdr, f"Gehaltsaenderungsmitteilung {year} - {emp}", sections,
              confidential=conf, draft=draft)


# ---------- REA_BV (Betriebsvereinbarungen) ---------------------------------

BV_TOPICS = {
    "Betriebliche_Altersversorgung": ("betriebliche Altersversorgung",
        "Direktzusage und Entgeltumwandlung; Versorgungstraeger: HeubeckAG-Gutachten."),
    "Betriebliches_Eingliederungsmanagement": ("BEM gemaess § 167 SGB IX",
        "Verfahren zur Wiedereingliederung nach laengerer Arbeitsunfaehigkeit."),
    "Datenschutz_IT_Systeme": ("Datenschutz und IT-Systeme",
        "Art. 88 DSGVO; Mitarbeiterdaten in SAP HR, MS 365, Concur, Workday."),
    "Essenszuschuss_Kantine":   ("Essenszuschuss / Kantine",
        "Tageszuschuss 3,57 EUR; steuer- und sozialabgabenfrei innerhalb Sachbezugsgrenze."),
    "Gesundheitsfoerderung":    ("Betriebliche Gesundheitsfoerderung",
        "Praeventionskurse nach § 20a SGB V; bis 600 EUR p. a. steuerfrei."),
    "Homeoffice_Mobiles_Arbeiten": ("Homeoffice und mobiles Arbeiten",
        "Bis zu 3 Tage/Woche; Ausstattung; Kostenpauschale; ArbStaettV §3."),
    "Jubilaeumsregelung":       ("Jubilaeumsregelung",
        "10/20/30/40 Jahre Dienstzeit; Jubilaeumszuwendung gestaffelt."),
    "Kurzarbeit":               ("Kurzarbeit",
        "Verfahren nach § 96 SGB III; KuG-Aufstockung 80 % Netto durch Arbeitgeber."),
    "Leistungsbeurteilung_und_Prämien": ("Leistungsbeurteilung und Praemien",
        "Jaehrliches Mitarbeitergespraech; Praemiensystem nach IGM/AT-Reglement."),
    "Mobilitätszuschuss_ÖPNV":  ("Mobilitaetszuschuss OePNV",
        "Jobticket Stuttgart/Heilbronn/Muenchen; bis 50 EUR/Monat steuerfrei."),
    "Qualifizierung_Weiterbildung": ("Qualifizierung und Weiterbildung",
        "Konzern-Lernpfade; 5 Tage p. a. zertifizierte Weiterbildung."),
    "Urlaub_Urlaubsgeld":       ("Urlaub und Urlaubsgeld",
        "30 Tage Tarif / 30 Tage AT; Urlaubsgeld nach IGM-TV; Verfall 31.03. Folgejahr."),
    "Vergütungssystem_IGM_Tarif": ("Verguetungssystem IGM-Tarif",
        "ERA-Eingruppierung EG 4-12; Schichtzulagen; Zulagen Heilbronn."),
    "Vertrauensarbeitszeit_Führungskräfte": ("Vertrauensarbeitszeit Fuehrungskraefte",
        "Geltung fuer leitende AT (Band E1-E3); Erfassung nach EuGH 14.05.2019."),
    "Zuschlagsregelung_Schichtarbeit": ("Zuschlagsregelung Schichtarbeit",
        "Mehrarbeit 25 %, Nachtschicht 30 %, Sonntag 60 %, Feiertag 150 %."),
    "Homeoffice_Telearbeit": ("Homeoffice und Telearbeit (REG Heilbronn)",
        "Lokale Ausgestaltung der Konzern-BV fuer das Werk Heilbronn."),
}


def betriebsvereinbarung(path):
    name_co, addr, hrb, title, conf, draft = read_meta(path)
    hdr = header_for(name_co, addr, hrb)
    # detect topic
    topic_key = None
    for k in BV_TOPICS:
        if k in path.name:
            topic_key = k
            break
    label, hint = BV_TOPICS.get(topic_key, ("HR-Thema", "Konzernweite Regelung."))

    sections = [
        ("Praeambel",
         f"Der Vorstand der Brennhagen Elektronik AG und der Konzernbetriebsrat (KBR), vertreten "
         f"durch die Vorsitzende Marlies Duerr (IG Metall) und den stellv. Vorsitzenden Klaus "
         f"Bauer (Werk Heilbronn), schliessen die folgende Betriebsvereinbarung zum Thema "
         f"»{label}« ab. Ziel ist es, einheitliche, transparente und mitbestimmungsgerechte "
         "Regelungen fuer alle Beschaeftigten der Konzerngesellschaften in Deutschland zu schaffen. "
         f"Hintergrund: {hint}"),
        ("§ 1 Geltungsbereich",
         "Persoenlich: Alle Arbeitnehmer/innen der inlaendischen Konzerngesellschaften (REA, REG, "
         "RSG, RHO) einschliesslich Auszubildender, Werkstudent/innen und Praktikant/innen, soweit "
         "nicht ausdruecklich anders geregelt. Sachlich: Themenbereich gemaess Praeambel. "
         "Raeumlich: alle Standorte in Deutschland. Auslandsgesellschaften (RPL, RCZ, RHU, RCN) "
         "regeln vergleichbare Themen ueber lokale Vereinbarungen, soweit zwingendes nationales "
         "Recht entgegensteht."),
        ("§ 2 Regelungsinhalt",
         f"(1) Es gelten die in Anlage 1 detailliert festgelegten Verfahren und Konditionen zu "
         f"»{label}«.\n\n"
         "(2) Operative Umsetzungsverantwortung liegt bei HR Business Partner Prof. Renate Meyer; "
         "Eskalationsinstanz ist die HR-Leitung in Stuttgart (Bjoern Franke).\n\n"
         "(3) Bei Konflikten oder Unklarheiten ist primaer das paritaetisch besetzte HR-Komitee "
         "anzurufen; im zweiten Schritt die Einigungsstelle gemaess § 76 BetrVG."),
        ("§ 3 Mitbestimmungsrechte",
         "Die Mitbestimmungsrechte des Betriebsrats nach § 87 BetrVG bleiben unberuehrt. "
         "Insbesondere gelten die Mitbestimmungstatbestaende des § 87 Abs. 1 Nr. 1 (Ordnung des "
         "Betriebs), Nr. 6 (technische Einrichtungen), Nr. 7 (Gesundheitsschutz) und Nr. 10 "
         "(betriebliche Lohngestaltung) - soweit themenrelevant."),
        ("§ 4 Datenschutz",
         "Die Erhebung, Verarbeitung und Nutzung personenbezogener Daten erfolgt ausschliesslich "
         "im Rahmen dieser Vereinbarung und nach Massgabe der DSGVO (Art. 88) sowie § 26 BDSG. "
         "Detaillierte Verarbeitungstaetigkeiten sind im Verzeichnis nach Art. 30 DSGVO dokumentiert. "
         "Datenschutzbeauftragte: HUEBL Datenschutz GmbH (extern, Stuttgart)."),
        ("§ 5 Inkrafttreten, Laufzeit, Kuendigung",
         "(1) Diese Betriebsvereinbarung tritt am 1. Januar 2023 in Kraft.\n\n"
         "(2) Sie wird auf unbestimmte Zeit geschlossen. Sie kann von beiden Seiten mit einer "
         "Frist von drei Monaten zum Quartalsende, fruehestens jedoch zum 31.12.2024, gekuendigt "
         "werden.\n\n"
         "(3) Eine Nachwirkung gemaess § 77 Abs. 6 BetrVG wird vereinbart, bis eine neue Regelung "
         "in Kraft tritt."),
        ("Unterschriften",
         signatures("Anna Mueller", "Vorstandsvorsitzende", "Brennhagen Elektronik AG",
                    "Marlies Duerr", "KBR-Vorsitzende", "Konzernbetriebsrat REA",
                    place="Stuttgart", date_str_=ds(2022, 11, 28))),
    ]
    write_doc(path, hdr, title or f"Betriebsvereinbarung - {label}",
              sections, confidential=conf, draft=draft)


# ---------- REA_HR Policies --------------------------------------------------

HR_POLICY = {
    "Auslandsentsendung": ("Auslandsentsendung", "Long-/Short-Term Assignment >= 90 Tage"),
    "Datenschutz_Mitarbeiter": ("Datenschutz Mitarbeiter", "Beschaeftigtendatenschutz, Art. 88 DSGVO"),
    "Diversity_Inclusion": ("Diversity & Inclusion", "Diversitaets- und Inklusionsstrategie"),
    "Homeoffice_Mobiles_Arbeiten": ("Homeoffice / Mobiles Arbeiten", "Hybride Arbeit, Equipment"),
    "Offboarding": ("Offboarding", "Austrittsprozess, Wissenstransfer, Exit-Interview"),
    "Onboarding": ("Onboarding", "Einarbeitungsprozess 90-Tage-Plan"),
    "Performance_Management": ("Performance Management", "MBO, Calibration, Talent-Review"),
    "Rekrutierung": ("Rekrutierung", "Konzernweiter Recruiting-Prozess inkl. Diversity"),
    "Verguetung": ("Verguetung", "Mercer-Benchmark, Bandstruktur, Equal Pay"),
    "Vorstandsverguetung": ("Vorstandsverguetung", "§ 87 AktG, § 162 AktG, Say-on-Pay"),
    "Weiterbildung": ("Weiterbildung", "Konzern-Lernpfade, 5 Tage p. a."),
    "Whistleblowing": ("Whistleblowing / Hinweisgebersystem", "HinSchG, Plattform EQS Integrity Line"),
}


def hr_policy(path):
    name_co, addr, hrb, title, conf, draft = read_meta(path)
    hdr = header_for(name_co, addr, hrb)
    topic = None
    for k in HR_POLICY:
        if k in path.name:
            topic = k; break
    label, hint = HR_POLICY.get(topic, ("HR-Richtlinie", "Konzernweite Richtlinie"))

    sections = [
        ("1. Zweck und Geltungsbereich",
         f"Die vorliegende HR-Richtlinie regelt konzernweit das Thema »{label}« ({hint}). "
         "Sie gilt fuer alle Beschaeftigten der Brennhagen Elektronik AG und ihrer Tochtergesellschaften "
         "(REG, RSG, RPL, RCZ, RHU, RCN, RHO), soweit nicht zwingendes lokales Arbeitsrecht "
         "entgegensteht. Die Richtlinie ergaenzt einschlaegige Betriebsvereinbarungen und "
         "individualvertragliche Abreden."),
        ("2. Grundsaetze",
         ("list", [
             "Gleichbehandlung und Diskriminierungsfreiheit (AGG, Diversity-Strategie 2024).",
             "Transparenz, Nachvollziehbarkeit und Dokumentation aller HR-Entscheidungen.",
             "Compliance mit DSGVO, HinSchG, AktG, BetrVG, IATF 16949 Kapitel 6 (Personal).",
             "Mitarbeiterorientierung und Foerderung individueller Entwicklung.",
             "Konsistenz mit Konzern-Code of Conduct (Stand: Maerz 2024).",
         ])),
        ("3. Verfahren und Verantwortlichkeiten",
         f"Operative Verantwortung fuer die Umsetzung liegt bei den lokalen HR Business Partnern. "
         "Konzernweite Steuerung erfolgt durch den HR-Lenkungskreis (CHRO-Funktion vertreten durch "
         "Laura Bauer als CFO mit HR-Verantwortung, Prof. Renate Meyer, Bjoern Franke). "
         f"Themenspezifische Eskalationen zum Thema »{label}« erfolgen ueber das HR-Komitee "
         "(monatlich tagend). Detaillierte Prozessbeschreibung in Anlage 1; Rollen-/RACI-Matrix "
         "in Anlage 2."),
        ("4. Spezifische Regelungen",
         f"(1) Fuer das Thema »{label}« gelten die folgenden konkreten Regelungen:\n\n"
         "- Antraege/Anfragen sind ueber das Workday-Self-Service-Portal einzureichen.\n"
         "- Genehmigungsworkflow erfolgt durch direkte Fuehrungskraft + HR Business Partner (Vier-Augen-Prinzip).\n"
         "- Genehmigungsbescheid erfolgt binnen 10 Arbeitstagen.\n"
         "- Eskalation bei Ablehnung: HR-Lenkungskreis.\n\n"
         "(2) Sonderregelungen fuer Auslandstoechter werden in lokalen Annexen geregelt; bei Konflikt "
         "geht zwingendes lokales Recht vor."),
        ("5. Datenschutz und Vertraulichkeit",
         "Die Erhebung personenbezogener Daten erfolgt ausschliesslich zweckgebunden gemaess Art. 88 "
         "DSGVO i. V. m. § 26 BDSG sowie der Konzern-Betriebsvereinbarung »Datenschutz und IT-Systeme« "
         "(2022). Aufbewahrungsfristen richten sich nach gesetzlicher Vorgabe (3-10 Jahre); "
         "personenbezogene Daten in der Personalakte werden 3 Jahre nach Austritt geloescht, soweit "
         "keine laengere Aufbewahrungspflicht besteht."),
        ("6. Inkrafttreten, Review-Zyklus",
         "Diese Richtlinie tritt am 1. Januar 2023 in Kraft. Sie wird im 2-Jahres-Rhythmus auf "
         "Aktualitaet ueberprueft (naechste Pruefung: Q4/2024). Aenderungen beduerfen der Zustimmung "
         "des Vorstands und der Mitbestimmung des Konzernbetriebsrats nach Massgabe des BetrVG. "
         "Verantwortlich fuer den Review: HR-Lenkungskreis (Laura Bauer, Prof. Renate Meyer)."),
        ("Freigabe",
         signatures("Anna Mueller", "Vorstandsvorsitzende", "Brennhagen Elektronik AG",
                    "Marlies Duerr", "KBR-Vorsitzende", "Konzernbetriebsrat REA",
                    place="Stuttgart", date_str_=ds(2022, 12, 19))),
    ]
    write_doc(path, hdr, title or f"HR-Richtlinie - {label}",
              sections, confidential=conf, draft=draft)


# ---------- LTIP / STIP / Organigramm / Restrukturierung --------------------

def lti_2022(path):
    name_co, addr, hrb, title, conf, draft = read_meta(path)
    hdr = header_for(name_co, addr, hrb)
    sections = [
        ("1. Programmuebersicht",
         "Der Long-Term Incentive Plan (LTIP) 2022-2024 der Brennhagen Elektronik AG ist ein "
         "performance-basiertes Aktienprogramm fuer Vorstand und ausgewaehlte Senior Manager. "
         "Es wird in Form von Performance Share Units (PSU) gewaehrt; die Auszahlung erfolgt am "
         "Ende des 4-jaehrigen Performance-Zeitraums (Vesting 1.7.2022 - 30.6.2026) ausschliesslich "
         "in bar (cash-settled). Beschlossen durch HV-Beschluss vom 14.6.2022 und vom Aufsichtsrat "
         "am 28.6.2022."),
        ("2. Berechtigte und Volumen",
         [["Gruppe", "Kreis", "Ziel-Volumen (PSU/Person)", "Maximaler Wert"],
          ["A", "Vorstand REA (4 Personen)", "12.000 - 18.000", "ca. 600 TEUR p. a."],
          ["B", "Senior VPs (8 Personen)",   "6.000 - 9.000",   "ca. 300 TEUR p. a."],
          ["C", "Directors (24 Personen)",   "2.500 - 4.500",   "ca. 150 TEUR p. a."],
          ["Gesamt", "ca. 36 Berechtigte",  "ca. 220.000 PSU",  "ca. 9,2 Mio. EUR p. a."]]),
        ("3. Performanceziele und Gewichtung",
         ("list", [
             "Relative Total Shareholder Return (TSR) ggue. STOXX Europe 600 Auto & Parts (Gewicht 40 %).",
             "Return on Capital Employed (ROCE), 4-Jahres-Durchschnitt 2022-2025 (Gewicht 30 %).",
             "ESG-Index: Scope 1+2 Reduktion, Frauenanteil Fuehrung, LTIFR (Gewicht 20 %).",
             "Strategische Meilensteine: EV-Anteil Umsatz, ADAS-OEM-Gewinne (Gewicht 10 %).",
             "Auszahlung skaliert linear von 0 % bei < 70 % Zielerreichung bis 200 % Cap bei >= 150 %.",
         ])),
        ("4. Vesting, Malus, Clawback",
         "Der Performance-Zeitraum ist 4 Jahre (1.7.2022 - 30.6.2026). Vesting erfolgt nach "
         "festgestellter Zielerreichung im Q3/2026. Bei vorzeitigem Ausscheiden Good-Leaver/Bad-Leaver-"
         "Regelung gemaess Anlage A. Malus-Regelung (Reduktion vor Auszahlung) bei Compliance-Verstoss, "
         "Restatement oder grobem Pflichtverstoss. Clawback (Rueckforderung bereits ausgezahlter "
         "Tranchen) bis zu drei Jahre nach Auszahlung bei vorsaetzlich-fehlerhafter Berichterstattung "
         "oder Compliance-Verstoss (DCGK G.11)."),
        ("5. Bewertung und Bilanzierung",
         "Marktbewertung der PSU erfolgt zum Gewaehrungsstichtag (1.7.2022) durch Willis Towers "
         "Watson (Monte-Carlo-Simulation fuer TSR-Bedingung; faire Wertberechnung ohne TSR fuer "
         "non-market vested conditions). IFRS-Bilanzierung als cash-settled share-based payment "
         "gemaess IFRS 2 mit Neubewertung zu jedem Bilanzstichtag (Verbindlichkeit). Dokumentation "
         "im Konzernanhang Note 8.4."),
        ("6. Governance",
         "Der Aufsichtsratspraesidialausschuss (Vorsitz: Dr. Klaus Steinbrueck) genehmigt jaehrlich "
         "die Zielfestsetzung und feststellt die Zielerreichung. Pruefung durch KPMG AG WPG im "
         "Rahmen der Konzernabschlusspruefung. Offenlegung im Verguetungsbericht nach § 162 AktG."),
        ("Freigabe",
         signatures("Dr. Klaus Steinbrueck", "Aufsichtsratsvorsitzender", "Brennhagen Elektronik AG",
                    "Anna Mueller", "Vorstandsvorsitzende", "Brennhagen Elektronik AG",
                    place="Stuttgart", date_str_=ds(2022, 6, 28))),
    ]
    write_doc(path, hdr, title or "Long-Term Incentive Plan (LTIP) 2022-2024 - Brennhagen Elektronik AG",
              sections, confidential=conf, draft=draft)


def stip_2023(path):
    name_co, addr, hrb, title, conf, draft = read_meta(path)
    hdr = header_for(name_co, addr, hrb)
    sections = [
        ("1. Grundsaetze",
         "Der Short-Term Incentive Plan (STIP) 2023 regelt die jaehrlich variable Verguetung fuer "
         "Vorstand, AT-Mitarbeitende und definierte Tarifmitarbeitende mit Zielvereinbarung. "
         "Die Ziele werden im Februar jedes Jahres vereinbart, die Auszahlung erfolgt im April des "
         "Folgejahres (Aprilgehalt 2024). Beschluss-Grundlage: Vorstandsbeschluss vom 30.1.2023, "
         "Mitbestimmung durch BV »Leistungsbeurteilung und Praemien« 2022."),
        ("2. Zielparameter 2023",
         [["#", "Parameter", "KPI / Definition", "Gewicht", "Ziel 100 %"],
          ["1", "Konzern-EBITDA", "EBITDA IFRS Konzern", "30 %", "74,3 Mio. EUR"],
          ["2", "Konzern-Umsatz", "Umsatz IFRS Konzern", "20 %", "612 Mio. EUR"],
          ["3", "Net Working Capital", "NWC / Umsatz (%)", "10 %", "<= 21 %"],
          ["4", "Bereichsziele", "Funktionsspezifisch", "20 %", "Anlage 1"],
          ["5", "Individuelle Ziele (MBO)", "Persoenlich", "10 %", "Anlage 2"],
          ["6", "ESG-Sondermalus", "LTIFR, Frauenanteil", "10 %", "Zielwerte siehe Anlage 3"]]),
        ("3. Auszahlungslogik",
         "Auszahlung bei 100 % Zielerreichung = 100 % des vereinbarten Zielbonus. Auszahlungskurve "
         "linear von 0 % (bei < 70 % Zielerreichung, Cliff) bis 150 % Cap (bei >= 150 %). Eine "
         "Ueberschreitung der Cap ist ausgeschlossen. Berechnungsbasis ist das im Dezember 2023 "
         "geltende Fixgehalt. Bei unterjaehrigem Eintritt: pro rata temporis."),
        ("4. Zielbonushoehen",
         ("list", [
             "Vorstand REA: 50-60 % des Fixgehalts.",
             "VPs / Direktoren (Band E2-E3): 30-40 % des Fixgehalts.",
             "Senior AT-Mitarbeitende (Band AT1-AT2): 15-25 % des Fixgehalts.",
             "AT-Mitarbeitende (Band AT3): 10-15 % des Fixgehalts.",
         ])),
        ("5. Anpassungs- und Malusklauseln",
         "Bei wesentlichen Veraenderungen der Geschaeftsgrundlage koennen die Ziele im Einvernehmen "
         "von Vorstand und Aufsichtsrats-Praesidialausschuss angepasst werden. Bei Compliance-"
         "Verstoessen, Restatements oder grober Pflichtverletzung kann der Bonus ganz oder teilweise "
         "einbehalten oder rueckgefordert werden (Malus/Clawback gemaess DCGK G.11)."),
        ("6. Reporting",
         "Quartalsweises Monitoring der Zielerreichung durch HR-Lenkungskreis und Group Controller "
         "Florian Maier. Finale Feststellung der Zielerreichung im Februar 2024 durch Vorstand bzw. "
         "Aufsichtsrats-Praesidialausschuss (fuer Vorstand). Pruefung durch KPMG im Rahmen der "
         "Konzernabschlusspruefung. Offenlegung im Verguetungsbericht 2023 nach § 162 AktG."),
        ("Freigabe",
         signatures("Anna Mueller", "Vorstandsvorsitzende", "Brennhagen Elektronik AG",
                    "Marlies Duerr", "KBR-Vorsitzende", "Konzernbetriebsrat REA",
                    place="Stuttgart", date_str_=ds(2023, 1, 30))),
    ]
    write_doc(path, hdr, title or "Short-Term Incentive Plan (STIP) 2023 - Brennhagen Elektronik AG",
              sections, confidential=conf, draft=draft)


def organigramm(path):
    name_co, addr, hrb, title, conf, draft = read_meta(path)
    hdr = header_for(name_co, addr, hrb)
    sections = [
        ("Vorstand (Executive Board) - Stand 1. April 2024",
         [["Position", "Name", "Eintritt Funktion", "Standort"],
          ["Vorsitzende (CEO)", "Anna Mueller", "1.9.2022", "Stuttgart"],
          ["CFO", "Laura Bauer", "1.9.2022", "Stuttgart"],
          ["COO", "Dr. Thomas Weber", "1.9.2022", "Heilbronn"],
          ["CTO (bis 30.6.2024)", "Stefan Hoffmann", "1.9.2022", "Muenchen"],
          ["CTO (ab 1.7.2024)", "Dr. Petra Hollmann", "1.7.2024", "Muenchen"],
          ["CMO / BD", "Stefan Richter", "1.4.2023", "Stuttgart"],
          ["CRO Asia (International)", "Dr. Yuki Tanaka", "1.4.2024", "Shanghai"]]),
        ("Direkt berichtende Bereiche / Konzernfunktionen",
         [["Bereich", "Leitung", "Berichtet an", "FTE"],
          ["Group Treasury", "Markus Pflanzer", "CFO", "8"],
          ["Group Tax", "Dr. Heike Berger", "CFO", "6"],
          ["Group Controlling", "Florian Maier", "CFO", "14"],
          ["Internal Audit", "Andreas Buehler (CAE)", "Praesidialausschuss AR", "5"],
          ["Group Legal / Compliance", "Dr. Andreas Becker", "CEO", "9"],
          ["Group HR", "Prof. Renate Meyer (interim)", "CFO", "22"],
          ["Group IT / Cyber Security", "Frank Werner", "CTO", "31"],
          ["Group Sustainability / ESG", "Sabine Brand", "CEO", "4"]]),
        ("Tochtergesellschaften",
         [["Kuerzel", "Gesellschaft", "Sitz", "Werk-/Country-Lead", "FTE"],
          ["REG", "Brennhagen Elektronik GmbH", "Heilbronn", "Andreas Maier", "820"],
          ["RSG", "Brennhagen Software GmbH", "Muenchen", "Dr. Klaus Kessler", "340"],
          ["RPL", "Brennhagen Polska Sp. z o.o.", "Katowice", "Marek Wojciechowski", "960"],
          ["RCZ", "Brennhagen CZ s.r.o.", "Brno", "Petr Novák", "680"],
          ["RHU", "Brennhagen Hungary Kft.", "Györ", "László Kovács", "540"],
          ["RCN", "Brennhagen (Shanghai) Co. Ltd.", "Shanghai", "Zhang Hao", "320"],
          ["RHO", "Brennhagen Holding GmbH", "Stuttgart", "Laura Bauer", "45"]]),
        ("Aufsichtsrat (5 Mitglieder, Amtsperiode 14.6.2022 - HV 2026)",
         [["Position", "Name", "Profession", "Ausschuss"],
          ["Vorsitz", "Dr. Klaus Steinbrueck", "ehem. CEO ZF Friedrichshafen", "Praesidial"],
          ["stv. Vorsitz", "Prof. Dr.-Ing. Gerhard Voss", "TU Muenchen", "Pruefung (Vorsitz)"],
          ["Mitglied", "Dr. Ingrid Schoellkopf", "ehem. KPMG-Partnerin", "stv. Pruefung"],
          ["Arbeitnehmer", "Marlies Duerr", "IG Metall / KBR-Vorsitz", "Personal"],
          ["Mitglied", "Thomas Reinhardt (MdB)", "ehem. BT, Public Affairs", "Nominierung"]]),
        ("Hinweise",
         "Stand des Organigramms: 1. April 2024. Aenderungen nach diesem Datum (u. a. Wechsel CTO "
         "zum 1.7.2024) sind erfasst, aber operativ noch nicht wirksam. Detaillierte Organisations- "
         "und Vertretungsregelungen sind im Konzern-Geschaeftsordnungshandbuch (GO 2024-03) sowie "
         "in den jeweiligen Vorstandsgeschaeftsordnungen geregelt."),
    ]
    write_doc(path, hdr, title or "Konzern-Organigramm 2024 - Brennhagen Elektronik AG",
              sections, confidential=conf, draft=draft)


def restrukturierung(path):
    name_co, addr, hrb, title, conf, draft = read_meta(path)
    hdr = header_for(name_co, addr, hrb)
    sections = [
        ("1. Hintergrund und strategische Begruendung",
         "Im Rahmen der strategischen Neuausrichtung des Brennhagen-Konzerns (Fokussierung auf "
         "Elektromobilitaet, ADAS Level 2-3 und ASPICE-konforme Embedded-Software) sowie zur "
         "Sicherung der mittelfristigen Wettbewerbsfaehigkeit des Standorts Heilbronn (REG) ist "
         "eine Anpassung der Belegschaftsstruktur erforderlich. Insbesondere bauen die OEM-Kunden "
         "(BMW, VW, Mercedes) verbrennungsmotor-affine Baugruppen sukzessive ab; gleichzeitig "
         "steigen die Anforderungen an die BMS-12- und ADAS-V4D-Fertigungslinien."),
        ("2. Massnahmen",
         ("list", [
             "Abbau von ca. 120 FTE bei REG bis 31.12.2024, ueberwiegend in den Bereichen "
             "konventionelle Steuergeraete-Fertigung (ECU-700-Generation) und administrative Stuetzfunktionen.",
             "Aufbau von 60 FTE in den Bereichen BMS-Fertigung und ADAS-Vorserie (Brutto-Abbau 60 FTE).",
             "Qualifizierungsoffensive: 800 Schulungstage in 2024 (Hochvolt-Sicherheit, Cybersecurity, "
             "ASPICE Level 2-3).",
             "Schliessung von Linie L-4 (LightCtrl-7 Vorserie) bis Q3/2024.",
             "Errichtung einer neuen BMS-Linie L-7 mit Investitionsvolumen 12,5 Mio. EUR (Beschluss "
             "Praesidialausschuss vom 18.10.2023).",
         ])),
        ("3. Sozialplan und Interessenausgleich",
         "Mit dem Betriebsrat Heilbronn (Vorsitz Klaus Bauer) wurde am 21.11.2023 ein Interessenausgleich "
         "und Sozialplan gemaess §§ 111-112 BetrVG verhandelt und beschlossen. Eckpunkte:\n\n"
         "- Abfindungsformel: (Bruttomonatsgehalt) x (Faktor 1,1) x (Dienstjahre).\n"
         "- Sprinterpraemie: 25 % auf die Abfindung bei Aufhebungsvertrag innerhalb 6 Wochen.\n"
         "- Vorruhestandsmodell ab 60 (ATZ-aehnlich, 75 % Bruttoeinkommen bis Rentenbeginn).\n"
         "- Transfergesellschaft bei BFW Heilbronn mit Laufzeit 12 Monaten, finanziert durch "
         "Arbeitsagentur 67 % und Arbeitgeber 33 %.\n"
         "- Gesamtvolumen Sozialplan: 18,5 Mio. EUR (geplant), Rueckstellung im Konzernabschluss 2023."),
        ("4. Zeitplan",
         [["Phase", "Massnahme", "Verantwortlich", "Termin"],
          ["1", "Anhoerung und Verhandlung BR/KBR", "HR / Geschaeftsfuehrung REG", "Q4/2023"],
          ["2", "Information Mitarbeiter und Betriebsversammlung", "GF + BR", "Januar 2024"],
          ["3", "Freiwilligenphase / Aufhebungsvertraege", "HR Heilbronn", "Februar-April 2024"],
          ["4", "Betriebsbedingte Kuendigungen (falls erforderlich)", "HR", "Mai 2024"],
          ["5", "Aufbau Qualifizierung / Umsiedlung", "Personal- und Organisationsentwicklung", "ab Q2/2024"],
          ["6", "Inbetriebnahme Linie L-7 BMS", "Werkleiter Andreas Maier", "Q4/2024"]]),
        ("5. Finanzielle Auswirkungen",
         "Einmalige Restrukturierungsaufwendungen in 2023 i. H. v. 18,5 Mio. EUR (Sozialplan-"
         "Rueckstellung) + 2,2 Mio. EUR (Beratungs- und Projektkosten). Jaehrliche Einsparungen ab "
         "2025 ca. 9,8 Mio. EUR (vor Reinvestition in Aufbau). Investition Linie L-7: 12,5 Mio. EUR "
         "(Capex 2024). Erwartete EBITDA-Verbesserung 2025 ggue. 2023 ca. +1,5 PP."),
        ("6. Mitbestimmung und Genehmigung",
         "Aufsichtsrat REA hat in der Sitzung vom 12.12.2023 dem Restrukturierungsplan zugestimmt "
         "(einstimmig, bei Stimmenthaltung von Marlies Duerr als KBR-Vorsitzende). Beschluss "
         "AR-2023-12-04. Konzernbetriebsrat wurde gemaess § 106 BetrVG unterrichtet."),
        ("Freigabe",
         signatures("Anna Mueller", "Vorstandsvorsitzende", "Brennhagen Elektronik AG",
                    "Dr. Klaus Steinbrueck", "Aufsichtsratsvorsitzender", "Brennhagen Elektronik AG",
                    place="Stuttgart", date_str_=ds(2023, 12, 15))),
    ]
    write_doc(path, hdr, title or "Restrukturierungsplan - Brennhagen Elektronik GmbH - Heilbronn 2023",
              sections, confidential=conf, draft=draft)


# ---------- BR (Betriebsrat) Protokolle / Wahl ------------------------------

def br_protokoll(path):
    name_co, addr, hrb, title, conf, draft = read_meta(path)
    hdr = header_for(name_co, addr, hrb)
    is_wahl = "Wahl" in path.name
    if is_wahl:
        # Wahlprotokoll
        sections = [
            ("Wahlausschuss",
             "Wahlausschuss-Vorsitz: Hans-Peter Lehmann. Beisitzer: Sabine Brand (Q-Leitung), "
             f"Lars Wittmann (RSG, soweit relevant). Wahltermine, Vorbereitung und Auszaehlung "
             "erfolgten gemaess Wahlordnung BetrVG (WO). Wahlausschreiben am 12.1.2022 ausgehaengt; "
             "Wahltag: 15. Maerz 2022."),
            ("Wahlbeteiligung und Ergebnis",
             "Wahlberechtigte: 820 (REG Heilbronn) bzw. 45 (RHO Stuttgart). Gueltige Stimmen: 681 "
             "(REG; Wahlbeteiligung 83 %) bzw. 40 (RHO; Wahlbeteiligung 89 %). Ungueltige Stimmen: "
             "REG 14, RHO 1. Wahlverfahren: Listenwahl gemaess § 14 BetrVG mit zwei Vorschlagslisten "
             "(»IG Metall - Faire Zukunft« und »Unabhaengige Kollegen«). Es gab keine Anfechtungen "
             "innerhalb der 14-Tages-Frist nach § 19 BetrVG."),
            ("Gewaehlte Mitglieder",
             [["Pos.", "Name", "Liste", "Stimmen"],
              ["1", "Klaus Bauer (Vorsitz)", "IG Metall - Faire Zukunft", "412"],
              ["2", "Sabine Mueller (stv. Vorsitz)", "IG Metall - Faire Zukunft", "298"],
              ["3", "Dirk Schaefer (Schriftfuehrer)", "Unabhaengige Kollegen", "201"],
              ["4", "Maria Hoffmann", "IG Metall - Faire Zukunft", "187"],
              ["5", "Thomas Becker", "Unabhaengige Kollegen", "165"],
              ["6", "Petra Schmidt", "IG Metall - Faire Zukunft", "152"],
              ["7", "Andreas Wolf", "Unabhaengige Kollegen", "144"],
              ["8", "Julia Lange", "IG Metall - Faire Zukunft", "131"],
              ["9", "Frank Richter", "Unabhaengige Kollegen", "118"]]),
            ("Konstituierung",
             "Der Betriebsrat hat sich in der konstituierenden Sitzung am 22. Maerz 2022 gemaess "
             "§ 26 BetrVG konstituiert. Klaus Bauer wurde einstimmig zum Vorsitzenden gewaehlt, "
             "Sabine Mueller einstimmig zur stellvertretenden Vorsitzenden. Bildung der Ausschuesse: "
             "Wirtschaftsausschuss (§ 106 BetrVG, 3 Mitglieder), Personalausschuss (3 Mitglieder), "
             "Arbeitsschutzausschuss (gemaess § 11 ASiG, 5 Mitglieder)."),
            ("Mitbestimmungs- und Beteiligungsrechte",
             "Der neugewaehlte Betriebsrat tritt unmittelbar in alle Rechte und Pflichten ein, "
             "die sich aus BetrVG, KSchG, ArbSchG, AGG, EntgTranspG und den geltenden Konzern- und "
             "Betriebsvereinbarungen ergeben. Mitbestimmungsschulungen gemaess § 37 Abs. 6 BetrVG "
             "fuer alle Mitglieder werden im Q2/Q3 2022 durchgefuehrt (Trainer: ifb Seminare GmbH)."),
            ("Unterschriften",
             signatures("Hans-Peter Lehmann", "Wahlausschuss-Vorsitz", name_co,
                        "Sabine Brand", "Wahlausschuss-Beisitz", name_co,
                        place=CITY.get(name_co,"Stuttgart"), date_str_=ds(2022, 3, 16))),
        ]
        title_final = title or f"Protokoll - Betriebsratswahl 2022 - {name_co}"
    else:
        # monthly BR protokoll
        sections = [
            ("Anwesenheit",
             "BR-Vorsitzender: Klaus Bauer; stv. Vorsitzende: Sabine Mueller; Schriftfuehrer: "
             "Dirk Schaefer. Anwesend: 7 von 9 BR-Mitgliedern (entschuldigt: Maria Hoffmann krank; "
             "Frank Richter Schulung).\n\n"
             "Sitzungsdatum: 19. Dezember 2022, 14:00 - 17:30 Uhr\n"
             "Ort: Betriebsratsraum, Werk Heilbronn, Halle 3\n"
             "Vertretung Arbeitgeber: Andreas Maier (Werkleiter), Sabine Brand (HR Heilbronn)"),
            ("Tagesordnung",
             ("list", [
                 "1. Genehmigung des Protokolls der Sitzung vom 21.11.2022.",
                 "2. Bericht der Geschaeftsleitung (Andreas Maier).",
                 "3. Sachstand Verhandlung Sozialplan Linie L-4 / Restrukturierung.",
                 "4. Antraege auf Zustimmung zu Einstellungen / Versetzungen (§ 99 BetrVG).",
                 "5. Aktuelle Arbeitsschutzthemen und Unfallzahlen Q4.",
                 "6. Verschiedenes / Termine 2023.",
             ])),
            ("Wesentliche Beschluesse",
             "TOP 2: Bericht zur Auftragslage Q4 (Auftragseingang +6 % YoY, vor allem BMS-12 fuer "
             "VW ID.7) und zur Auslastung Linien L-1 bis L-6 (durchschnittlich 84 %). "
             "TOP 3: Der Betriebsrat fordert detaillierte Sozialauswahl-Kriterien bis 31.1.2023 und "
             "eine Sondersitzung im Januar 2023 zur weiteren Beratung. "
             "TOP 4: Zustimmung zu 14 Einstellungen (vor allem Hochvolt-Spezialisten und "
             "ADAS-Test-Engineers) sowie 3 Versetzungen. Eine Versetzung wurde abgelehnt (Hoererscheinen "
             "des Betroffenen gemaess § 99 BetrVG am 9.1.2023). "
             "TOP 5: Unfallquote Q4: 2 meldepflichtige Unfaelle (Stolperunfall, Schnittverletzung), "
             "beide ohne ernste Folgen. LTIFR liegt bei 1,8 (Ziel: <= 2,0)."),
            ("Naechste Termine",
             "Sondersitzung Sozialplan: 11. Januar 2023, 9:00.\n"
             "Regulaere Sitzung Januar 2023: 23. Januar 2023, 14:00.\n"
             "Klausurtagung BR 2023: 14.-15. Februar 2023, Hotel Schwarzwaldhof."),
            ("Unterschriften",
             signatures("Klaus Bauer", "BR-Vorsitzender", "BR " + name_co,
                        "Dirk Schaefer", "Schriftfuehrer", "BR " + name_co,
                        place="Heilbronn", date_str_=ds(2022, 12, 20))),
        ]
        title_final = title or "Protokoll Betriebsratssitzung Dezember 2022"

    write_doc(path, hdr, title_final, sections, confidential=conf, draft=draft)


# ---------- One-off odd docs ------------------------------------------------

def patent_doc(path):
    name_co, addr, hrb, title, conf, draft = read_meta(path)
    hdr = header_for(name_co, addr, hrb)
    sections = [
        ("Antragsstellung",
         "Antwort der Anmelderin Brennhagen Elektronik AG (vertreten durch Patentanwaltskanzlei "
         "Boehmert & Boehmert, Stuttgart) auf den Pruefungsbescheid des Deutschen Patent- und "
         "Markenamts (DPMA) vom 22. April 2022 zur Patentanmeldung DE1020240XXXXX (»Verfahren und "
         "Vorrichtung zur fehlertoleranten Spannungsmessung in Lithium-Ionen-Batteriezellen mittels "
         "redundanter ADC-Pfade«)."),
        ("Argumentation zur Patentfaehigkeit",
         "Die Anmelderin haelt die im Pruefungsbescheid erhobenen Einwaende der mangelnden "
         "Neuheit (§ 3 PatG) und mangelnden erfinderischen Taetigkeit (§ 4 PatG) fuer unbegruendet. "
         "Insbesondere offenbart die zitierte D1 (EP 3 XXX YYY A1 Continental) nicht den unabhaengigen "
         "Anspruch 1 in Verbindung mit der gemaess Anspruch 3 redundanten ADC-Pfadarchitektur mit "
         "Plausibilitaetspruefung. Der technische Effekt - Reduzierung der False-Negative-Rate bei "
         "Zellfehlern unterhalb 0,5 mV - ist in keinem der entgegengehaltenen Dokumente offenbart "
         "oder nahegelegt."),
        ("Geaenderte Anspruchsfassung",
         "Die Anmelderin reicht eine geaenderte Anspruchsfassung ein, in der die urspruengliche "
         "Aufgabe »fehlertolerante Spannungsmessung« durch die einschraenkenden Merkmale der "
         "redundanten ADC-Pfade gemaess urspruenglichem Anspruch 3 und der Plausibilitaetspruefung "
         "gemaess urspruenglichem Anspruch 6 praezisiert wird. Die geaenderten Ansprueche 1-12 "
         "werden in der Beilage »Anspruch_neu_v2.pdf« uebermittelt."),
        ("Antraege",
         ("list", [
             "Patenterteilung auf Basis der geaenderten Anspruchsfassung (Hauptantrag).",
             "Hilfsweise: Anberaumung einer muendlichen Verhandlung gemaess § 46 PatG.",
             "Gewaehrung einer Fristverlaengerung um zwei Monate (bis 22.10.2022), falls weiterer "
             "Pruefungsbescheid ergeht.",
         ])),
        ("Schluss",
         "Die Anmelderin ist davon ueberzeugt, dass mit den geaenderten Anspruechen alle Einwaende "
         "des Pruefers ausgeraeumt sind und die Anmeldung zur Patenterteilung reif ist. Fuer "
         "Rueckfragen steht der zustaendige Patentanwalt Dr. Christian Koch (IP-Manager Brennhagen "
         "Elektronik AG) sowie die Kanzlei Boehmert & Boehmert (RA Dr. Klaus Berger) zur Verfuegung."),
    ]
    write_doc(path, hdr, title or "Patent 14 - Antwort Anmelder", sections,
              confidential=conf, draft=draft)


def ic_rechnung(path):
    name_co, addr, hrb, title, conf, draft = read_meta(path)
    hdr = header_for(name_co, addr, hrb)
    sections = [
        ("Rechnungsempfaenger",
         "Brennhagen Elektronik AG\nVaihinger Strasse 120\n70567 Stuttgart, Deutschland\n\n"
         "VAT-ID: DE 312 487 901\nKundennummer: REA-IC-CZ"),
        ("Rechnungsausstellerin",
         f"Brennhagen CZ s.r.o.\nSlatina Industrial Park\n627 00 Brno, Tschechien\n\n"
         "IČO: 28418762, DIČ: CZ28418762\nUST-IdNr DE-Empfaenger: DE 312 487 901"),
        ("Rechnungsdaten",
         [["Feld", "Wert"],
          ["Rechnungsnummer", "RCZ-IC-2020-09-1247"],
          ["Rechnungsdatum", "30. September 2020"],
          ["Leistungszeitraum", "1.-30. September 2020"],
          ["Faelligkeit", "30. November 2020 (60 Tage netto, IC-Standard)"],
          ["Waehrung", "EUR"],
          ["Lieferbedingung", "DAP Stuttgart (Incoterms 2020)"],
          ["TP-Methode", "Cost-Plus 6,5 % (siehe Local File)"]]),
        ("Leistungspositionen",
         [["Pos.", "Bezeichnung", "Menge", "Stueckpreis EUR", "Summe EUR"],
          ["1", "Steckverbinder ADAS-V4D Connector A (32-pol.)", "12.400", "0,87", "10.788,00"],
          ["2", "Steckverbinder BMS-12 HV-Anschluss (Pkg/8)",   "3.200",  "4,12", "13.184,00"],
          ["3", "Steckverbinder ECU-900 Sensorport (Pkg/24)",   "850",    "9,60",  "8.160,00"],
          ["4", "Pruef- und Konfektionierungsdienstleistung Q3",  "1 Pauschal", "4.500,00", "4.500,00"],
          ["", "Netto-Warenwert", "", "", "36.632,00"],
          ["", "Cost-Plus-Aufschlag bereits enthalten", "", "", ""],
          ["", "USt 0 % (innergem. B2B-Lieferung, Reverse Charge)", "", "", "0,00"],
          ["", "Rechnungsbetrag", "", "", "36.632,00"]]),
        ("Steuerlicher Hinweis",
         "Innergemeinschaftliche B2B-Lieferung gemaess Art. 138 MwStSystRL i. V. m. § 4 Nr. 1 b) UStG. "
         "Die Steuerschuldnerschaft geht auf den Leistungsempfaenger ueber (Reverse Charge gemaess "
         "§ 13b UStG). Beide Parteien verfuegen ueber gueltige USt-IdNrn (qualifizierte Bestaetigung "
         "BZSt vorhanden, Pruefnummer T-2020-1198)."),
        ("Verrechnungspreis-Hinweis",
         "Die Verrechnungspreise wurden gemaess Cost-Plus-Methode (TNMM-naehestliegende Methode) mit "
         "einem Aufschlag von 6,5 % auf die direkten Herstellkosten ermittelt. Dies entspricht der "
         "Festlegung im Transfer-Pricing Local File 2020 (Abschnitt 4.3) und der Master File "
         "Konzernrichtlinie 2020. Benchmark: Bureau van Dijk TP-Catalyst 2020."),
        ("Zahlungsdaten",
         "Empfaenger: Brennhagen CZ s.r.o.\nIBAN: CZ65 0800 0000 1234 5678 9012\nBIC: GIBACZPX\n"
         "Bank: Ceska Sporitelna, a.s., Brno\nVerwendungszweck: RCZ-IC-2020-09-1247"),
    ]
    write_doc(path, hdr, title or "Intercompany-Rechnung September 2020", sections,
              confidential=conf, draft=draft)


def dd_working_paper(path):
    name_co, addr, hrb, title, conf, draft = read_meta(path)
    hdr = header_for(name_co, addr, hrb)
    sections = [
        ("Auftrag und Scope",
         "Im Rahmen des strategischen Prozesses 2023/2024 (potenzielle Akquisition »Project "
         "Aurora«) wurde Freshfields Bruckhaus Deringer LLP mit der rechtlichen Due Diligence "
         "beauftragt. Dieses Arbeitspapier Nr. 7 dokumentiert die Ergebnisse der Workstream "
         "»IT / Cybersecurity / Datenschutz«. Lead Bearbeitung: Dr. Petra Hollmann (Counsel "
         "Freshfields, IT-Praxis); Support: Senior Associate Dr. Lukas Strauss. Vertraulichkeit "
         "gemaess NDA vom 18.7.2023."),
        ("Methodik",
         "Auswertung Datenraum (Intralinks, 1.247 Dokumente IT/Cyber); 4 Management-Interviews "
         "(CIO, CISO, DPO, Head of IT-Ops); Q&A-Iterationen 12 Runden; technische Validierung "
         "durch KPMG (IT-Audit, DD-Workstream Forensik). Stichtag der Pruefung: 30.6.2023."),
        ("Wesentliche Findings",
         ("list", [
             "ISO 27001:2022-Zertifizierung des Zielunternehmens befindet sich noch in Vorbereitung "
             "(geplant Q4/2024). Aktueller Reifegrad gemaess BSI-IT-Grundschutz: Niveau »Standard«.",
             "Drei mittelschwere DSGVO-Vorfaelle 2022-2023 (alle gemaess Art. 33 DSGVO an die "
             "Landesdatenschutzbehoerde gemeldet, keine Bussgelder).",
             "Cyber-Versicherungsdeckung: Munich Re, 25 Mio. EUR (ausreichend fuer Risikoprofil).",
             "Penetrationstest 03/2023 durch SecureWorks: 14 Findings, davon 2 high (geschlossen), "
             "8 medium (in Bearbeitung), 4 low.",
             "Datenschutzfolgeabschaetzungen (DPIA) liegen fuer alle Hochrisikoverarbeitungen vor "
             "(7 Stueck, gemaess Art. 35 DSGVO).",
         ])),
        ("Risikobewertung",
         [["Risiko", "Bewertung", "Maßnahmen-Empfehlung"],
          ["Fehlende ISO 27001-Zertifizierung", "Mittel",
           "SPA-Garantie inkl. Specific Indemnity 2 Mio. EUR"],
          ["DSGVO-Compliance-Restrisiken", "Niedrig-Mittel",
           "Eskrow 500 TEUR fuer 24 Monate"],
          ["IT-Integration in Brennhagen-Umgebung", "Mittel",
           "Carve-out-Plan und 100-Tage-Integrationsroadmap"],
          ["Open-Source-Lizenzcompliance", "Niedrig",
           "Black Duck Scan vor Closing"]]),
        ("Empfehlungen fuer SPA / Closing",
         "Specific Indemnity in Hoehe von 2 Mio. EUR fuer ISO 27001-Themen ueber 36 Monate. "
         "Pauschale Cyber-Indemnity (Cap 5 Mio. EUR) fuer pre-closing Cyber-Incidents. "
         "MAC-Klausel (Material Adverse Change) inkl. cyber-incident-Trigger > 1 Mio. EUR Schaden. "
         "Closing-Bedingung: erfolgreicher Pentest durch unabhaengigen Dritten (Freshfields-empfohlen)."),
        ("Status",
         "Status des Arbeitspapiers: ABGESCHLOSSEN. Freigabe durch Lead Counsel Dr. Helmut "
         "Schwarz (Freshfields Frankfurt) am 14.9.2023. Verteilung an Steering Committee "
         "(Anna Mueller, Laura Bauer, Goldman Sachs Frankfurt, KPMG Lead) am 15.9.2023."),
    ]
    write_doc(path, hdr, title or "Due Diligence Arbeitspapier #007", sections,
              confidential=True, draft=draft)


# ---------- Dispatcher -------------------------------------------------------

def process(path: Path):
    name = path.name
    if name.startswith("AV_"):
        return arbeitsvertrag(path)
    if name.startswith("JG_"):
        return jahresgespraech(path)
    if name.startswith("MBO_"):
        return mbo(path)
    if name.startswith("Gehaltsaenderung_"):
        return gehaltsaenderung(path)
    if name.startswith("REA_BV_") or name.startswith("REG_BV_"):
        return betriebsvereinbarung(path)
    if name.startswith("REA_HR_Policy_"):
        return hr_policy(path)
    if name == "REA_Long_Term_Incentive_Plan_2022.docx":
        return lti_2022(path)
    if name == "REA_Short_Term_Incentive_Plan_2023.docx":
        return stip_2023(path)
    if name == "REA_Organigramm_Konzern_2024.docx":
        return organigramm(path)
    if name == "REA_Restrukturierung_REG_Heilbronn_2023.docx":
        return restrukturierung(path)
    if name.endswith("BR_Protokoll_2022_12.docx") or name.endswith("BR_Wahl_Protokoll_2022.docx"):
        return br_protokoll(path)
    if name.startswith("Patent_"):
        return patent_doc(path)
    if name.startswith("RCZ_IC_Rechnung_"):
        return ic_rechnung(path)
    if name.startswith("DD_WorkingPaper_"):
        return dd_working_paper(path)
    print(f"  ! no handler for {name}")
    return None


def main():
    thin = [p for p in sorted(BASE.rglob("*.docx")) if word_count(p) < 200]
    print(f"Found {len(thin)} thin docs")
    n_ok = n_skip = 0
    for p in thin:
        try:
            process(p)
            n_ok += 1
        except Exception as e:
            print(f"  X {p.name}: {e}")
            n_skip += 1
    print(f"\nRewrote {n_ok} files; {n_skip} errors")

    # Verify
    total = thin_now = 0
    for p in sorted(BASE.rglob("*.docx")):
        w = word_count(p)
        total += 1
        if w < 200:
            thin_now += 1
            print(f"  STILL thin {w:4d}  {p.name}")
    print(f"\n{total} total, {thin_now} still thin")


if __name__ == "__main__":
    main()
