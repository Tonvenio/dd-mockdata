"""Brennhagen Elektronik GmbH (REG) – Werk Heilbronn / 03_Tochter_DE_REG_Heilbronn – 92 thin docs."""
# --- portable-paths-prelude --- (do not edit) ---
import sys
from pathlib import Path as _PathlibPath
_RP = _PathlibPath(__file__).resolve().parent
while not (_RP / "enhance_lib.py").exists() and _RP.parent != _RP:
    _RP = _RP.parent
_ROOT = _RP
sys.path.insert(0, str(_ROOT))
# --- end prelude ---
from enhance_lib import ROEHRIG_AG as R, ROEHRIG_SUBS as S, write_doc, signatures

BASE = f"{_ROOT}/roehrig_large/03_Tochter_DE_REG_Heilbronn"

# Subsidiary header (most docs are REG-issued; few exceptions handled inline)
REG_H = {
    "name": "Brennhagen Elektronik GmbH (REG)",
    "addr": "Industriestrasse 47, 74076 Heilbronn",
    "hrb":  "HRB 221456, Amtsgericht Heilbronn",
}
RCN_H = {
    "name": "Brennhagen (Shanghai) Co. Ltd. (RCN)",
    "addr": "Building 7, 888 Hongqiao Road, Changning District, 200336 Shanghai, P.R. China",
    "hrb":  "Unified Social Credit Code 91310115MA1FL42Q38",
}
REA_H = {"name": R["name"], "addr": R["addr"], "hrb": R["hrb"]}
KPMG_H = {
    "name": "KPMG AG Wirtschaftspruefungsgesellschaft",
    "addr": "Theodor-Heuss-Strasse 5, 70174 Stuttgart  |  Lead Partner Dr. Maximilian Brand",
    "hrb":  "HRB 14346, AG Berlin-Charlottenburg",
}

# ============================================================================
# 1) ARBEITSVERTRAEGE REG (5 docs)
# ============================================================================

def arbeitsvertrag(fname, position, name, monatsgehalt_eur, jahresboni_pct,
                   urlaub, kuendigung, taetigkeit, vorgesetzter, sonderregel=""):
    write_doc(f"{BASE}/{fname}", REG_H,
        f"Anstellungsvertrag – {position} – Brennhagen Elektronik GmbH",
        subtitle=f"Zwischen REG Heilbronn und {name}; gueltig ab 1. Januar 2022",
        sections=[
            ("Praeambel",
             f"Zwischen der Brennhagen Elektronik GmbH, Industriestrasse 47, 74076 Heilbronn, vertreten "
             f"durch den Geschaeftsfuehrer Herrn Andreas Maier (»Arbeitgeberin«), und Herrn/Frau "
             f"{name} (»Arbeitnehmer/in«) wird der nachfolgende Anstellungsvertrag geschlossen. "
             f"Die REG ist eine 100 %-Tochter der Brennhagen Holding GmbH (RHO Stuttgart) im Konzern der "
             f"boersennotierten Brennhagen Elektronik AG (Prime Standard, ISIN DE000RHGRP12). Die Stelle "
             f"ist eingebettet in das Hauptwerk Heilbronn mit ca. 820 Beschaeftigten und Produktion "
             f"der Plattformen ICP-3 (Infotainment), BMS-12 (Battery Management), ECU-900 (Powertrain) "
             f"und ADAS-V4D (Radar-Fusion-Steuergeraet)."),
            ("Vertragsklauseln",
             ("clauses", [
                ("§ 1 Taetigkeit und Dienstort", [
                    f"Der/Die Arbeitnehmer/in wird als {position} eingestellt. Dienstort ist das Werk "
                    f"Heilbronn (Industriestrasse 47, 74076 Heilbronn). Voruebergehende Versetzungen "
                    f"an andere Brennhagen-Standorte (Stuttgart, Muenchen, Katowice, Brno, Gyoer) sind "
                    f"im zumutbaren Rahmen moeglich, insbesondere fuer Audits, Trainings und Projekt-"
                    f"einsaetze.",
                    f"Aufgabenbereich: {taetigkeit}",
                    f"Disziplinarisch vorgesetzt: {vorgesetzter}. Fachlich berichtspflichtig im "
                    f"Rahmen der Konzern-Matrix-Organisation an die entsprechende Konzern-Funktion "
                    f"in Stuttgart (z. B. Group Controlling, Group HR, Group Compliance, Group "
                    f"Quality)."]),
                ("§ 2 Verguetung", [
                    f"Das monatliche Bruttogehalt betraegt {monatsgehalt_eur:,} EUR, zahlbar jeweils "
                    f"zum Monatsletzten auf das vom Arbeitnehmer benannte Konto. Das Jahresgehalt "
                    f"versteht sich als 13. Monatsgehalt (12 + 1 Weihnachtsgeld); Auszahlung des "
                    f"13. Gehaltes mit der Novemberabrechnung.".replace(",", "."),
                    f"Zusaetzlich wird eine variable Jahresvergutung in Hoehe von bis zu "
                    f"{jahresboni_pct} % des Jahresgrundgehaltes gewaehrt. Die Bonus-Bemessung "
                    f"erfolgt nach (a) Werk-KPI Heilbronn (OEE, PPM, Reklamationsquote, EBIT-Werk) "
                    f"zu 40 %, (b) individueller Zielerreichung zu 40 % sowie (c) Konzern-EBITDA "
                    f"der Brennhagen Elektronik AG zu 20 %. Auszahlung im Folgejahr nach Feststellung "
                    f"des Konzernabschlusses (regelmaessig Maerz).",
                    f"Sachbezuege: Dienstwagen der Mittelklasse (Audi A4 / BMW 3er / Mercedes C-Klasse, "
                    f"Listenpreis bis 55.000 EUR) zur dienstlichen und privaten Nutzung; 1 %-"
                    f"Regelung gemaess EStG. Mobiltelefon und Laptop werden gestellt; private Nutzung "
                    f"im angemessenen Umfang gestattet."]),
                ("§ 3 Arbeitszeit", [
                    f"Die regelmaessige woechentliche Arbeitszeit betraegt 40 Stunden, verteilt auf "
                    f"Montag bis Freitag. Bei der Position wird eine Vertrauensarbeitszeit vereinbart; "
                    f"Mehrarbeit ist mit dem Gehalt abgegolten.",
                    f"Im Rahmen der IG-Metall-Tarifregelungen (Bezirk Baden-Wuerttemberg) sind "
                    f"abweichende Arbeitszeitmodelle (mobile Arbeit bis 40 % der Wochenarbeitszeit) "
                    f"in Abstimmung mit dem Vorgesetzten und dem oertlichen Betriebsrat (Vorsitz "
                    f"Hans-Juergen Klemm) moeglich."]),
                ("§ 4 Urlaub", [
                    f"Der Jahresurlaubsanspruch betraegt {urlaub} Arbeitstage bei einer 5-Tage-Woche. "
                    f"Urlaubsgewaehrung erfolgt unter Beruecksichtigung der betrieblichen Belange "
                    f"(insbesondere Werksferien Heilbronn, regelmaessig KW 31-33 und KW 52)."]),
                ("§ 5 Verschwiegenheit, IP, Compliance", [
                    f"Der/Die Arbeitnehmer/in verpflichtet sich zur strengen Verschwiegenheit ueber "
                    f"alle Betriebs- und Geschaeftsgeheimnisse der Brennhagen-Gruppe, insbesondere "
                    f"OEM-Pricing, Margen, technische Spezifikationen der Produktplattformen "
                    f"(ICP-3, BMS-12, ECU-900, ADAS-V4D, LightCtrl-7) sowie ueber konzerninterne "
                    f"Verrechnungspreise und IFRS-Konzernzahlen vor Veroeffentlichung.",
                    f"Saemtliche im Rahmen der Taetigkeit entstehenden Arbeitsergebnisse, "
                    f"Diensterfindungen (ArbnErfG) und Schutzrechte gehen unentgeltlich auf die "
                    f"Arbeitgeberin ueber. Diensterfindungen sind unverzueglich der zentralen IP-"
                    f"Stelle (REA Stuttgart, Patent-Anwaltskanzlei Bardehle Pagenberg) zu melden.",
                    f"Es gelten die Brennhagen-Konzern-Compliance-Richtlinien (Code of Conduct 2023, "
                    f"Anti-Korruptions-Richtlinie, Kartellrechts-Richtlinie, Insider-Handelsrichtlinie). "
                    f"Verstoesse koennen zur fristlosen Kuendigung fuehren."]),
                ("§ 6 Wettbewerbsverbot", [
                    f"Fuer die Dauer von 12 Monaten nach Beendigung des Arbeitsverhaeltnisses "
                    f"verpflichtet sich der/die Arbeitnehmer/in, weder unmittelbar noch mittelbar "
                    f"fuer Wettbewerber der Brennhagen-Gruppe taetig zu werden. Wettbewerber sind "
                    f"insbesondere Automotive Tier-1 mit vergleichbarem Produktportfolio (Continental, "
                    f"Bosch, ZF, Aptiv, Marelli, LG Magna, Denso, Visteon, Harman).",
                    f"Als Karenzentschaedigung wird 50 % der zuletzt bezogenen festen Bezuege "
                    f"(monatliches Bruttogehalt zzgl. 1/12 des letztjaehrigen Bonus) gewaehrt; "
                    f"Auszahlung monatlich."]),
                ("§ 7 Kuendigung", [
                    f"Das Anstellungsverhaeltnis kann mit einer Frist von {kuendigung} jeweils zum "
                    f"Ende eines Kalendermonats ordentlich gekuendigt werden. Das Recht zur "
                    f"ausserordentlichen Kuendigung aus wichtigem Grund (§ 626 BGB) bleibt unberuehrt.",
                    f"Das Vertragsverhaeltnis endet automatisch mit Ablauf des Monats, in dem "
                    f"der/die Arbeitnehmer/in das gesetzliche Renteneintrittsalter erreicht."]),
                ("§ 8 Schlussbestimmungen", [
                    f"Sollten einzelne Bestimmungen dieses Vertrages unwirksam sein oder werden, "
                    f"so bleibt die Wirksamkeit der uebrigen Bestimmungen davon unberuehrt. "
                    f"Aenderungen und Ergaenzungen beduerfen der Schriftform. Gerichtsstand ist "
                    f"Heilbronn. Es gilt deutsches Recht.",
                    sonderregel or f"Keine weiteren Sonderregelungen."]),
             ])),
            ("Unterschriften",
             signatures("Andreas Maier", "Geschaeftsfuehrer / Werkleiter REG",
                        "Brennhagen Elektronik GmbH",
                        name, position, "Arbeitnehmer/in",
                        place="Heilbronn", date_str_="15. Dezember 2021")),
        ])


arbeitsvertrag("REG_Arbeitsvertrag_01_Geschäftsführer_in_R_2022.docx",
    "Geschaeftsfuehrer/in REG", "Andreas Maier", 12916, 50, 30,
    "9 Monaten",
    "Gesamtverantwortung fuer die Brennhagen Elektronik GmbH Heilbronn (820 MA, ca. 280 Mio. EUR "
    "Umsatz 2023). Strategische Steuerung des Werks, OEM-Beziehungen (BMW, VW, Mercedes, "
    "Stellantis), Investitionsplanung, Berichterstattung an Vorstand REA (CEO Anna Mueller, "
    "COO Dr. Thomas Weber).",
    "Vorstand der Brennhagen Elektronik AG (Konzernmutter)",
    "Berufung als Geschaeftsfuehrer erfolgt durch Gesellschafterbeschluss der Brennhagen Holding GmbH "
    "(Alleingesellschafterin). Eintragung im Handelsregister Heilbronn HRB 221456. D&O-Versicherung "
    "ueber Konzern-Police bei Allianz Global Corporate & Specialty SE (Versicherungssumme 50 Mio. EUR).")

arbeitsvertrag("REG_Arbeitsvertrag_02_Produktionsleiter_in_2022.docx",
    "Produktionsleiter/in", "Dipl.-Ing. Markus Henkel", 9450, 30, 30,
    "6 Monaten",
    "Operative Leitung der Produktionsbereiche SMD-Bestueckung (Linien 1-6), Endmontage "
    "(Linien A-D) und EOL-Test fuer die Plattformen ICP-3, BMS-12, ECU-900 und ADAS-V4D. "
    "Verantwortung fuer OEE (Zielwert 78 % gemaess KPI-Plan 2024), PPM-Quote (< 25 ppm) und "
    "Lieferpuenktlichkeit (OTD > 98 %). Personalverantwortung fuer ca. 480 Beschaeftigte in "
    "3-Schicht-Modell.",
    "Werkleiter Andreas Maier",
    "Vereinbarung zur Bereitschaft fuer Schicht-/Wochenend-Einsaetze bei Eskalationen (Linien-"
    "ausfall, Kunden-Reklamation, EOL-Faults). Hierfuer wird eine Rufbereitschaftspauschale von "
    "850 EUR/Monat zusaetzlich gezahlt.")

arbeitsvertrag("REG_Arbeitsvertrag_03_Qualitätsmanagerin_H_2022.docx",
    "Qualitaetsmanagerin / Q-Leitung Werk Heilbronn", "Sabine Brand", 9120, 25, 30,
    "6 Monaten",
    "Gesamtverantwortung Qualitaetsmanagement Werk Heilbronn nach IATF 16949 und ISO 9001:2015. "
    "Steuerung der Q-Prozesse APQP, PPAP, FMEA, MSA, SPC, 8D, Kundenreklamationen, Lieferanten-"
    "Audits (VDA 6.3) sowie OEM-Audits (BMW QM-CAQ, VW Formel-Q, Mercedes MBSP-Q, Stellantis "
    "AAQS). Lead-Auditor fuer interne Audits IATF 16949.",
    "Werkleiter Andreas Maier (disziplinarisch); Konzern-Q-Leiter Dr. Bernhard Klotz (fachlich)",
    "Verpflichtung zur Aufrechterhaltung der persoenlichen Zertifizierungen Lead-Auditor IATF "
    "16949, VDA 6.3, FMEA-Moderator. Externe Schulungskosten werden von REG getragen.")

arbeitsvertrag("REG_Arbeitsvertrag_04_Finanzcontroller_Hei_2022_rev_SRichter.docx",
    "Finanzcontroller/in Werk Heilbronn", "Stefan Richter", 7850, 20, 30,
    "3 Monaten",
    "Werkscontrolling REG Heilbronn: monatliches Werks-Reporting (MonatsPL, OEE, "
    "Material-Effizienz, Personal-PL), Forecasting (FC-Updates Q1-Q4), Plan-Ist-Abweichungs"
    "analysen, Investitionsbeurteilung (CapEx > 250 TEUR), Mitwirkung Jahresabschluss (HGB lokal, "
    "IFRS-Reporting an Group Controlling Florian Maier in Stuttgart).",
    "Geschaeftsfuehrer REG (disziplinarisch); Group Controller Florian Maier (fachlich)",
    "Revisionsstand 14.11.2022 nach Review durch CMO/BD Stefan Richter (Konzernvorstand) – "
    "Anpassung Bonus-Struktur (40 % Werk-KPI / 40 % individuell / 20 % Konzern-EBITDA).")

arbeitsvertrag("REG_Arbeitsvertrag_05_HR-Manager_Heilbronn_2022.docx",
    "HR-Manager/in Werk Heilbronn", "Petra Holzwarth", 7200, 18, 30,
    "3 Monaten",
    "Operative HR-Verantwortung fuer ca. 820 Mitarbeitende am Standort Heilbronn: Recruiting "
    "(direct hire + Zeitarbeit Manpower / Adecco), Personalentwicklung, Gehaltsadministration "
    "(SAP HCM, Payroll ueber DATEV), HR-Business-Partner-Funktion fuer Werkleitung, Schnittstelle "
    "zum Betriebsrat Heilbronn (Vorsitz Hans-Juergen Klemm, 13 Mitglieder).",
    "Werkleiter Andreas Maier",
    "Vertraulichkeitsverpflichtung bzgl. Personalakten, Gehaltsdaten und Verhandlungen mit dem "
    "Betriebsrat (insbesondere Betriebsvereinbarungen). Schnittstellenfunktion zur Konzern-HR "
    "in Stuttgart (Leiterin HR Birgit Voelkner).")


# ============================================================================
# 2) COMPLIANCE REPORT 2023
# ============================================================================

write_doc(f"{BASE}/REG_Compliance_Report_2023.docx", REG_H,
    "Compliance-Bericht 2023 – Brennhagen Elektronik GmbH",
    subtitle="Berichtsperiode: 1. Januar 2023 – 31. Dezember 2023; Stichtag 31.12.2023",
    sections=[
        ("Berichtszweck",
         "Der Compliance-Bericht 2023 dokumentiert den Stand des Compliance-Management-Systems "
         "(CMS) der Brennhagen Elektronik GmbH (REG) Werk Heilbronn. Der Bericht erfuellt die "
         "interne Berichtspflicht gegenueber dem Konzern-Compliance-Officer (Dr. Rainer Kunze, "
         "REA Stuttgart) sowie die Anforderungen des Pruefungsausschusses des Aufsichtsrates der "
         "Brennhagen Elektronik AG (Vorsitz Prof. Dr.-Ing. Gerhard Voss). Methodik orientiert sich "
         "am IDW PS 980 (Grundsaetze ordnungsmaessiger Pruefung von Compliance-Management-Systemen) "
         "sowie am Konzern-Code-of-Conduct 2023."),
        ("Compliance-Organisation REG",
         [["Funktion", "Verantwortlich", "Berichtsweg"],
          ["Local Compliance Officer", "Dr. Lena Engelhardt (REG Legal)", "Konzern-Compliance Stuttgart"],
          ["Geldwaeschebeauftragte", "Stefan Richter (Controlling)", "Geschaeftsfuehrer + KPMG"],
          ["Datenschutzbeauftragter", "extern: SecLaw GmbH, Heilbronn", "Konzern-Datenschutz"],
          ["Exportkontrolle", "Heike Stoerle (Logistik)", "Ausfuhrverantwortlicher REA"],
          ["Anti-Korruption", "Local Compliance Officer", "Konzern-Compliance + Vorstand"]]),
        ("Wesentliche Themen 2023",
         ("list", [
            "Aktualisierung des Code of Conduct (Version 3.2, Inkrafttreten 1.4.2023) – Verteilung "
            "in Deutsch/Englisch an alle 820 Beschaeftigten REG; Quittungsquote 99,2 %.",
            "Implementierung des Hinweisgebersystems gemaess HinSchG (Hinweisgeberschutzgesetz, "
            "wirksam 2.7.2023): Anbindung an Konzern-Whistleblower-Tool »Speak-Up« (Provider "
            "EQS Group, hosting in DE); im Berichtsjahr 4 Meldungen aus REG (alle geprueft, "
            "2 ohne Befund, 2 mit organisatorischen Massnahmen abgeschlossen).",
            "Auffrischungsschulung Kartellrecht (Schulungsmodul »Wettbewerbsrecht im Vertrieb«) "
            "– absolviert von 100 % der Vertriebs- und KAM-Mitarbeiter, 92 % Gesamtquote REG.",
            "Audit Exportkontrolle / Dual-Use (BAFA-relevant) durchgefuehrt durch SUEDDEUTSCHE "
            "TREUHAND – keine wesentlichen Feststellungen, 3 Hinweise (umgesetzt bis 31.10.2023).",
            "Datenschutz: 2 meldepflichtige Vorfaelle gemaess Art. 33 DSGVO (verloren gegangenes "
            "Notebook, irrtuemlicher E-Mail-Versand) – fristgerecht an LfD BW gemeldet, "
            "Folgemassnahmen umgesetzt."]) ),
        ("Schulungs-KPI",
         [["Modul", "Pflicht-Zielgruppe", "Abdeckung 2023", "Plan 2024"],
          ["Code of Conduct (E-Learning)", "alle 820 MA", "99 %", "100 %"],
          ["Anti-Korruption", "alle Fuehrungskraefte + Vertrieb", "97 %", "100 %"],
          ["Kartellrecht", "Vertrieb, KAM, Werkleitung", "100 %", "100 %"],
          ["Datenschutz / DSGVO", "alle MA mit Datenbezug", "94 %", "100 %"],
          ["Exportkontrolle", "Logistik, Vertrieb, Engineering", "92 %", "100 %"],
          ["Arbeitssicherheit (DGUV)", "alle 820 MA", "96 %", "100 %"]]),
        ("Vorfaelle / Whistleblower-Meldungen 2023",
         "Im Berichtsjahr 2023 wurden ueber das Speak-Up-System sowie ueber den Local Compliance "
         "Officer insgesamt 7 Meldungen erfasst (Vorjahr: 5). Davon: 4 anonym, 3 namentlich. "
         "Inhaltliche Schwerpunkte: 2 vermutete Mobbing-Faelle (1 bestaetigt, einvernehmliche "
         "Trennung), 2 mutmasslich falsche Spesenabrechnungen (1 bestaetigt, schriftliche Abmahnung), "
         "1 Hinweis Lieferanten-Bevorzugung (nicht bestaetigt), 2 Verfahrensfragen (Aufklaerung). "
         "Keine bestaetigten Korruptionsfaelle, keine bestaetigten Kartellrechts-Verstoesse."),
        ("Risikoeinschaetzung 2024",
         "Schwerpunkte 2024: (a) Aktualisierung der Lieferanten-Compliance-Erklaerungen (LkSG "
         "Lieferketten-Sorgfaltspflichtengesetz, Anwendung erweitert ab 2024 auf Unternehmen "
         ">1.000 MA Konzern), (b) Schulung CSRD-Berichtspflichten Werkebene, (c) Auffrischung "
         "Sanktionslisten-Screening (Russland-Sanktionen, China-Exportkontrollen 2024), "
         "(d) Verstaerkung IT-Security-Awareness (Phishing-Simulationen Q1/Q3)."),
        ("Bestaetigung",
         "Die Geschaeftsfuehrung der Brennhagen Elektronik GmbH bestaetigt, dass das Compliance-"
         "Management-System der REG im Berichtsjahr 2023 wirksam war und keine wesentlichen "
         "Verstoesse gegen geltendes Recht festgestellt wurden, die zu materiellen Bilanz- oder "
         "Reputationsrisiken fuehren wuerden. Bericht zur Kenntnis vorgelegt am 14.2.2024 an "
         "Konzern-Compliance (Dr. Rainer Kunze) und Pruefungsausschuss (Prof. Voss)."),
        ("Unterschriften",
         signatures("Andreas Maier", "Geschaeftsfuehrer REG", "Brennhagen Elektronik GmbH",
                    "Dr. Lena Engelhardt", "Local Compliance Officer", "Brennhagen Elektronik GmbH",
                    place="Heilbronn", date_str_="14. Februar 2024")),
    ])


# ============================================================================
# 3) CONTROL PLANS (4 Produktplattformen)
# ============================================================================

def control_plan(fname, produkt, full_name, oem, prozess_rows, weitere_specs):
    write_doc(f"{BASE}/{fname}", REG_H,
        f"Lenkungsplan (Control Plan) – {produkt} {full_name} – Werk Heilbronn",
        subtitle=f"IATF 16949 Anhang A.5; Revisionsstand 4.2; Freigabe Q-Leitung 14.2.2023",
        confidential=True,
        sections=[
            ("Geltungsbereich",
             f"Dieser Control Plan regelt die Prozesslenkung fuer die Serienproduktion der "
             f"Produktplattform {produkt} ({full_name}) im Werk Heilbronn der Brennhagen Elektronik "
             f"GmbH. Hauptkunde: {oem}. Erstellt nach IATF 16949 Anhang A.5 (Production Control "
             f"Plan); ergaenzend zu PFMEA Rev. 3.1 und APQP-Phase 5. Verteilung: Produktion, "
             f"QS, Engineering, Werkleitung, OEM-Kunde (auf Anforderung; vertraulich).",),
            ("Allgemeine Angaben",
             [["Feld", "Wert"],
              ["Werk", "Brennhagen Elektronik GmbH, Werk Heilbronn"],
              ["Werksleitung", "Andreas Maier (WL); Sabine Brand (Q-Leitung)"],
              ["Produktfamilie", f"{produkt} – {full_name}"],
              ["Kunde / OEM", oem],
              ["Plattform-Lebenszyklus", "SOP 03/2022 – EOP geplant 12/2028"],
              ["Erstellt durch", "Dr.-Ing. Kerstin Roggner (QS-Engineering)"],
              ["Genehmigt durch", "Sabine Brand (Q-Leitung), Markus Henkel (Produktionsleiter)"],
              ["Revisionsstand", "4.2 (vorhergehend 4.1 vom 22.8.2022)"],
              ["Naechstes Review", "Q3 2023 (oder bei Aenderungsanforderung)"]]),
            ("Prozess-Steuerungsplan (Auszug)",
             [["Prozessschritt", "Merkmal", "Spezifikation", "Pruefmethode", "Frequenz", "Reaktionsplan"]] + prozess_rows),
            ("Weitere Festlegungen", weitere_specs),
            ("Aenderungs-Historie",
             [["Rev.", "Datum", "Aenderung", "Autor"],
              ["4.0", "01.04.2022", "Initial-Freigabe nach PPAP zur SOP", "K. Roggner"],
              ["4.1", "22.08.2022", "Erweiterung Pruefumfang EOL (CISPR 25)", "S. Brand"],
              ["4.2", "14.02.2023", "Anpassung Reaktionsplan; Update Lieferanten", "K. Roggner"]]),
            ("Bemerkungen",
             "Aenderungen am Control Plan beduerfen Freigabe durch Q-Leitung Werk Heilbronn "
             "(Sabine Brand) sowie Information des OEM-Kunden (Q-Engineering Kontakt) gemaess "
             "Customer Specific Requirements (CSR). Bei Aenderungen am Produkt-Design ist "
             "ein PPAP-Re-Submission gemaess AIAG Level 3 erforderlich."),
        ])


control_plan("REG_ControlPlan_ADAS-V4D_2023.docx", "ADAS-V4D",
    "Radar Fusion Steuergeraet (Level-2/3)", "Mercedes-Benz Group AG; Stellantis N.V.",
    [["SMT-Bestueckung Linie 3", "Pad-Offset", "+/- 0,08 mm", "AOI Yamaha YSi-V", "100 %", "Linie stoppen; Eskalation Produktionsleiter"],
     ["Reflow-Loeten", "Peak-Temperatur", "245 +/- 5 °C", "KIC Profiler 24/7", "stuendlich", "Profil-Anpassung; Sperrung Charge"],
     ["AOI nach SMT", "Pin-Lift, Polaritaet", "0 Fehler", "MIRTEC MV-7", "100 %", "Re-Work-Station; Sperrung Panel"],
     ["ICT (Flying Probe)", "Funktion Klemmen", "100 % i.O.", "Takaya APT-1400F", "100 %", "Re-Test; Funktionsdiagnose Engineering"],
     ["FCT-EOL", "CAN-Bus, FlexRay, Diagnostik", "DTC-frei", "Brennhagen Test-Stand FCT-RM-04", "100 %", "Re-Test; falls erneut NIO -> Sperrung Endprodukt"],
     ["EMV-Vorpruefung (CISPR 25)", "Class 5 Limits", "alle Frequenzen i.O.", "Stichprobe Schaffner CDN, GTEM", "1 pro Charge", "Charge sperren; Engineering-Analyse"],
     ["Verguss / Sealing", "Dichtheit IP67", "kein Bubble bei 0,8 bar", "Helium-Lecksuche", "100 %", "Re-Sealing; max. 1 Iteration zulaessig"]],
    "Lieferant Radar-MMIC: Infineon BGT24LTR12 (Stuttgart-Backnang); Lieferant Mikrocontroller: "
    "NXP S32R294 (D2). Spezielle Charakteristiken: Radar-Antennen-Charakterisierung bei "
    "Stichprobe 1:100 ueber Anechoic Chamber Heilbronn (KW 14 erweitert wegen Lieferanten-Wechsel). "
    "CSR Mercedes: zusaetzlich Burn-In-Pruefung 168h bei 85 °C / 95 % rF; CSR Stellantis: "
    "verstaerkte Korrosionspruefung Salzspruehnebel 480h gemaess ASTM B117.")

control_plan("REG_ControlPlan_BMS-12_2023.docx", "BMS-12",
    "Battery Management System (Hochvolt EV)", "Volkswagen AG (ID.7); Hyundai Motor Company",
    [["SMT-Linie 1+2 (BMS-Master)", "Loetstellen-Qualitaet", "IPC-A-610 Class 3", "AOI Mirtec MV-7L", "100 %", "Re-Work; SMT-Profil-Review"],
     ["SMT-Linie 5 (Cell-Module)", "Pad-Position Sense-Pins", "+/- 0,06 mm", "AOI + 3D-SPI Koh Young", "100 %", "Sperrung Panel"],
     ["Isolations-Pruefung HV", "Iso-Widerstand", ">= 100 MOhm @ 1000 V DC", "MEGGER MIT525-T", "100 %", "Linien-Stop; Eskalation HV-Sicherheits-Officer"],
     ["Cell-Balancing Funktionstest", "Spannungsabgleich", "Delta-U < 5 mV", "Brennhagen BMS-Tester Type BT-12", "100 %", "Re-Programmierung; ggf. Sperrung"],
     ["CAN-FD-Diagnose", "DBC-Konformitaet", "alle Botschaften", "Vector CANalyzer V18", "100 %", "Re-Flash; Funktionsdiagnose"],
     ["EOL-Aging-Test", "Stress 30 min @ Tmax", "kein Drift > 2 %", "Brennhagen Klimakammer KK-04", "100 %", "Sperrung Endprodukt"],
     ["IP67-Dichtheit", "Druckabfall", "< 50 mbar / 60 s", "Sciemetric BlackBelt", "100 %", "Re-Sealing"]],
    "CSR VW (Formel-Q Capability 3.0): erweitertes Reporting Cpk pro KW; ESD-Massnahmen "
    "verschaerft (Class 0). CSR Hyundai (HMC SQ): zusaetzlich Vibrationspruefung (ISO 16750-3 "
    "Profil 5) bei Stichprobe 1:200. HV-Sicherheit: Bedienung der HV-Pruefplatze ausschliesslich "
    "durch geschulte Elektrofachkraefte (EUP-zertifiziert).")

control_plan("REG_ControlPlan_ECU-900_2023.docx", "ECU-900",
    "Powertrain-ECU Gen3", "Volkswagen AG (MEB+ / MQB-Evo); Stellantis N.V.",
    [["SMT-Linie 4", "Komponenten-Position", "IPC-A-610 Class 2", "AOI", "100 %", "Re-Work"],
     ["Reflow-Loeten", "Peak-Temperatur", "240 +/- 5 °C", "KIC Profiler", "stuendlich", "Profil-Anpassung"],
     ["ICT Inline", "Bauteil-Wert / Funktion", "100 % i.O.", "Genrad GR228x", "100 %", "Re-Work; Sperrung"],
     ["Flash-Programmierung", "Software-Stand", "SW V3.4.2 fuer MEB+", "Brennhagen Flash-Stand", "100 %", "Re-Flash; max. 3 Iterationen"],
     ["FCT-EOL", "Diagnose UDS", "0 DTC", "Vector CANoe", "100 %", "Diagnose-Analyse"],
     ["Klimatest Stichprobe", "Funktion @ -40/+105 °C", "100 % i.O.", "Brennhagen KK-02", "1:500", "Charge sperren"]],
    "Stellantis MEB+ Variante mit zusaetzlichem PSI5-Interface (Druck-Sensoren). Lieferant "
    "PowerSemi: Infineon TC397XT (TriCore Aurix 32-bit). Sicherheits-Anforderung ISO 26262 "
    "ASIL-D fuer Powertrain-Domain. CSR Stellantis: vollstaendige Traceability pro Charge ueber "
    "DataMatrix-Code auf Gehaeuse; CSR VW: WPK-Auswertung MSA 4 monatlich.")

control_plan("REG_ControlPlan_ICP-3_2023.docx", "ICP-3",
    "InfoConnect Pro – Infotainment-Hauptrechner", "BMW Group; Mercedes-Benz Group AG",
    [["SMT-Linie 6 (HPC-Board)", "BGA-Voiding < 25 %", "IPC-A-610 Class 3", "AXI Viscom XR8000", "Stichprobe 1:50", "Sperrung Charge"],
     ["Reflow-Loeten", "Peak 248 +/- 4 °C", "Profil P-7 ICP", "KIC", "stuendlich", "Anpassung Profil"],
     ["ICT", "Funktion", "100 %", "Takaya APT-1400F", "100 %", "Re-Work-Station"],
     ["Display-Connect-Test", "MIPI-DSI / LVDS", "i.O.", "Brennhagen DispTest DT-7", "100 %", "Re-Connect"],
     ["FCT-EOL Audio/Video", "HDMI, A2B, Ethernet", "alle Schnittstellen", "Brennhagen FCT-IC-08", "100 %", "Re-Test"],
     ["Akustik-Pruefung", "SNR > 95 dB", "@ 1 kHz, 1 V RMS", "Audio Precision APx555", "Stichprobe 1:100", "Akustik-Engineering"],
     ["OTA-Update-Test (FOTA)", "Image-Boot OK", "Boot < 12 s", "Brennhagen Boot-Logger", "Stichprobe 1:500", "Re-Flash"]],
    "Mercedes (MBSP-Q): Bestaetigung Cybersecurity-Konformitaet ISO/SAE 21434 fuer Boot-Loader. "
    "BMW (QM-CAQ): zusaetzliche Klimapruefung -40/+85 °C an Stichprobe 1:500. SoC: Qualcomm "
    "SA8295P (Cockpit-Plattform). Speicher: LPDDR5 12 GB Micron. Lebensdauer-Anforderung: "
    "15 Jahre / 250.000 km bzw. 25.000 Betriebsstunden (BMW VVT-Reliability-Spec).")


# ============================================================================
# 4) IC-QUARTALSBERICHTE (7 docs: 2019 Q1-Q3, 2020 Q1-Q4)
# ============================================================================

def ic_quartalsbericht(fname, jahr, quartal, lieferungen_eur, ic_marge_pct,
                       hauptprodukt, kommentare):
    write_doc(f"{BASE}/{fname}", REG_H,
        f"Intercompany-Quartalsbericht {quartal} {jahr} – Brennhagen Elektronik GmbH (REG)",
        subtitle=f"Verrechnungspreis-Reporting an Konzern-Steuerabteilung (Group Tax)",
        sections=[
            ("Berichtszweck",
             f"Quartalsweise Berichterstattung der REG Werk Heilbronn ueber die IC-Lieferungen "
             f"und -Leistungen an verbundene Unternehmen der Brennhagen-Gruppe. Vorlage gemaess "
             f"Konzern-Verrechnungspreis-Richtlinie 2018 (Update 2022) sowie OECD-VPL 2022. "
             f"Adressat: Group Tax Director Dr. Heike Berger (REA Stuttgart), Group Controller "
             f"Florian Maier, Local Tax Advisor Deloitte GmbH WPG (Heilbronn-Office)."),
            ("Berichtszeitraum",
             f"Quartal: {quartal} {jahr}. Berichtswaehrung: EUR. Daten aus SAP S/4HANA REG, "
             f"Buchungskreis 1000-REG; Konsolidierung Konzern-IFRS ueber OneStream."),
            ("Wesentliche IC-Transaktionen (Quartalswerte)",
             [["Empfaenger / Gegenpartei", "Leistung", "Volumen EUR", "Methode"],
              ["RSG Brennhagen Software GmbH (DE)", "Software-Integration ICP-3 / ECU-900", f"{int(lieferungen_eur*0.18):,}".replace(",", "."), "Cost+ 7 %"],
              ["RPL Brennhagen Polska Sp. z o.o.", "Komponenten-Bezug (SMD-Bestueckung Vorprodukte)", f"-{int(lieferungen_eur*0.22):,}".replace(",", "."), "Cost+ 5 % (Eingang)"],
              ["RCZ Brennhagen CZ s.r.o.", "Steckverbinder + Kabelkonfektion", f"-{int(lieferungen_eur*0.12):,}".replace(",", "."), "Cost+ 6 % (Eingang)"],
              ["RHU Brennhagen Hungary Kft.", "Sensorik-Module", f"-{int(lieferungen_eur*0.08):,}".replace(",", "."), "Cost+ 6 % (Eingang)"],
              ["RCN Brennhagen Shanghai", "Komponenten / Aftermarket Asien", f"{int(lieferungen_eur*0.06):,}".replace(",", "."), "Cost+ 7 %"],
              ["REA Stuttgart (Holding)", "Konzernumlage Shared Services", f"-{int(lieferungen_eur*0.04):,}".replace(",", "."), "Kostenumlage"],
              ["Konzern-Treasury (Cash Pool)", "Cash-Pool-Nettoposition Zinsen", f"{int(lieferungen_eur*0.001):,}".replace(",", "."), "Marktueblicher Zins"],
              ["Gesamt IC-Volumen Quartal (Netto)", "", f"{int(lieferungen_eur):,}".replace(",", "."), "—"]]),
            ("IC-Margen und Benchmark",
             f"Realisierte IC-Marge fuer Routine-Producer-Komponenten im {quartal} {jahr}: "
             f"{ic_marge_pct:.1f} % EBIT/Umsatz auf IC-Lieferungen REG -> Konzern. "
             f"Benchmark-Korridor gemaess PwC-Studie 2021 (Update 2023): 4,8 % – 7,2 % "
             f"(Interquartile Range); realisierter Wert liegt im Median. Keine Anpassungen "
             f"erforderlich. Lizenzgebuehren (Konzernpatente) wurden zu 3,8 % auf IC-Umsatz "
             f"abgefuehrt (CUP)."),
            ("Hauptprodukt im Berichtszeitraum",
             f"Im {quartal} {jahr} entfielen die wesentlichen IC-Lieferungen auf die Plattform "
             f"{hauptprodukt}. {kommentare}"),
            ("Compliance / Auffaelligkeiten",
             f"Im Berichtszeitraum {quartal} {jahr} keine wesentlichen Beanstandungen seitens "
             f"des Local Tax Advisors. Keine Anpassungen seitens Betriebspruefung Finanzamt "
             f"Heilbronn (Pruefungszeitraum 2017-2019 laeuft separat, Status: ohne wesentliche "
             f"Feststellungen). Verrechnungspreis-Studien fuer die wesentlichen Kategorien sind "
             f"vorhanden und werden jaehrlich aktualisiert."),
            ("Anlagen",
             "(1) Detail-Auszug IC-Buchungen SAP (Konten 70xxx / 71xxx / 79xxx); (2) Cost+ "
             "Kalkulation Stundensaetze REG-Engineering 2023 ueberarbeitet; (3) Bestaetigung "
             "Geschaeftsfuehrer REG zur Korrektheit; (4) Quartals-Bestaetigung Local Tax "
             "Advisor."),
            ("Unterschriften",
             signatures("Stefan Richter", "Finanzcontroller REG", "Brennhagen Elektronik GmbH",
                        "Andreas Maier", "Geschaeftsfuehrer REG", "Brennhagen Elektronik GmbH",
                        place="Heilbronn", date_str_=f"20. Tag nach Quartalsende {jahr}")),
        ])


ic_quartalsbericht("REG_IC_Quartalsbericht_2019_Q1.docx", 2019, "Q1", 58_333_583, 6.2,
    "ICP-3 (Vorgaenger-Generation, BMW F-Klasse)",
    "Anlauf der zweiten Linie ICP-3 fuer BMW G05; Sondereffekte aus Tooling-Abgrenzungen "
    "ueber RSG (Software-Integration) erhoehten IC-Volumen um ca. 2,4 Mio. EUR.")
ic_quartalsbericht("REG_IC_Quartalsbericht_2019_Q2.docx", 2019, "Q2", 61_240_120, 6.4,
    "ICP-3 / ECU-900 Erstgeneration",
    "Erste IC-Lieferungen ECU-900-Vorserie an VW-Werk Wolfsburg fuer MEB-Plattform "
    "Erprobung; Q2 mit erheblichen Engineering-Aufwendungen RSG.")
ic_quartalsbericht("REG_IC_Quartalsbericht_2019_Q3.docx", 2019, "Q3", 64_580_900, 6.5,
    "ECU-900 (Powertrain ID.3)",
    "Anlauf-Phase ECU-900 fuer VW ID.3 (SOP 11/2019); Material-Vorbezug RPL fuer SMD-"
    "Vorprodukte erhoehte Eingangs-IC um 12 %.")
ic_quartalsbericht("REG_IC_Quartalsbericht_2020_Q1.docx", 2020, "Q1", 52_120_400, 5.4,
    "ICP-3 / ECU-900",
    "Covid-bedingte Lieferunterbrechungen ab Mitte Maerz; IC-Umsatz Q1 ca. 9 % unter Plan. "
    "Schliessung VW-Werk Wolfsburg Mitte Maerz hatte direkte Auswirkungen auf ECU-900-"
    "Volumen.")
ic_quartalsbericht("REG_IC_Quartalsbericht_2020_Q2.docx", 2020, "Q2", 38_240_600, 4.1,
    "ICP-3 / ECU-900",
    "Tiefpunkt der Covid-Krise; IC-Volumen Q2 ca. 35 % unter Vorjahr. Kurzarbeit 100 % "
    "ueber 6 Wochen; Konzern-Cash-Pool zur Liquiditaetsstuetzung in Anspruch genommen.")
ic_quartalsbericht("REG_IC_Quartalsbericht_2020_Q3.docx", 2020, "Q3", 55_820_300, 5.8,
    "ICP-3 / BMS-12 Vorbereitung",
    "Erholung im Anschluss an Covid-Tief; erste Vorbereitungen Plattform BMS-12 (SOP 2022 "
    "geplant). Anlauf Investitionen ca. 8 Mio. EUR (Linie 5 SMD).")
ic_quartalsbericht("REG_IC_Quartalsbericht_2020_Q4.docx", 2020, "Q4", 68_410_500, 6.6,
    "ICP-3 / ECU-900",
    "Starkes Q4 mit Nachfrage-Boom aus BMW + VW; deutliche IC-Volumensteigerung. "
    "Sondereffekte: Materialpreissteigerung Halbleiter erhoehte Eingangs-IC-Volumen um "
    "ca. 4 %.")


# ============================================================================
# 5) IC-RECHNUNGEN REG (39 Stueck, 2020-2022) – einheitliches Layout
# ============================================================================

import random

def ic_rechnung_reg(fname, jahr, monat_idx, betrag_eur, draft=False, final=False,
                    leistungs_titel="Konzern-IC-Sammelrechnung"):
    monate = ["Januar","Februar","Maerz","April","Mai","Juni","Juli","August",
              "September","Oktober","November","Dezember"]
    monat = monate[monat_idx-1]
    netto = round(betrag_eur / 1.19)
    mwst = betrag_eur - netto
    rg_nr = f"REG-IC-{jahr}-{monat_idx:02d}"
    leistungszeitraum = f"01.{monat_idx:02d}.{jahr} – {min(31, [31,28,31,30,31,30,31,31,30,31,30,31][monat_idx-1])}.{monat_idx:02d}.{jahr}"

    write_doc(f"{BASE}/{fname}", REG_H,
        f"Intercompany-Rechnung {monat} {jahr} – REG -> RHO Stuttgart",
        subtitle=f"Rechnungs-Nr. {rg_nr}; Konzern-Sammelabrechnung Komponenten + Engineering",
        draft=draft,
        sections=[
            ("Rechnungsadressat",
             f"Brennhagen Holding GmbH (RHO) – Konzern-Verrechnung\n"
             f"Vaihinger Strasse 120, 70567 Stuttgart\n"
             f"HRB 726450, Amtsgericht Stuttgart\n"
             f"USt-ID: DE 312 487 901\n\n"
             f"Konzern-Buchungsstelle: Frau Andrea Pfaff, Group Accounts Payable, "
             f"REA Stuttgart (Tel. 0711/4892-3214; E-Mail: andrea.pfaff@roehrig.com)"),
            ("Rechnungssteller",
             f"Brennhagen Elektronik GmbH (REG)\n"
             f"Industriestrasse 47, 74076 Heilbronn\n"
             f"HRB 221456, Amtsgericht Heilbronn\n"
             f"USt-ID: DE 287 612 449\n\n"
             f"Bankverbindung: Volksbank Heilbronn eG\n"
             f"IBAN: DE48 6209 1800 0123 4567 89; BIC: GENODES1VHN\n"
             f"Verwendungszweck: {rg_nr}"),
            ("Rechnungsdaten",
             [["Feld", "Wert"],
              ["Rechnungsnummer", rg_nr],
              ["Rechnungsdatum", f"{min(15, [31,28,31,30,31,30,31,31,30,31,30,31][monat_idx-1])}.{(monat_idx % 12)+1:02d}.{jahr if monat_idx<12 else jahr+1}"],
              ["Leistungszeitraum", leistungszeitraum],
              ["Vertragsgrundlage", "Konzern-IC-Rahmenvertrag REA/REG vom 1.1.2018 (Update 1.7.2022)"],
              ["Verrechnungspreis-Methode", "Cost+ Routine-Marge 5-7 % (vgl. TP-Studie PwC 2022)"],
              ["Zahlungsziel", "30 Tage netto Konzern-Settlement"],
              ["Lieferbedingungen", "DAP Stuttgart (Incoterms 2020)"]]),
            ("Leistungsaufstellung",
             [["Position", "Beschreibung", "Betrag EUR"],
              ["1", f"{leistungs_titel} {monat} {jahr} – Plattformen ICP-3 / ECU-900 / BMS-12", f"{int(netto*0.62):,}".replace(",", ".")],
              ["2", "Engineering-Services (Cost+ 7 %, ca. 8.450 h)", f"{int(netto*0.18):,}".replace(",", ".")],
              ["3", "Materialweiterberechnung (Halbleiter, ICs, passive Komponenten)", f"{int(netto*0.12):,}".replace(",", ".")],
              ["4", "Tooling-Abgrenzungen (Sondervorrichtungen, Pruefadapter)", f"{int(netto*0.05):,}".replace(",", ".")],
              ["5", "Logistik / Verpackung (DAP Stuttgart)", f"{int(netto*0.03):,}".replace(",", ".")],
              ["Summe Netto", "", f"{netto:,}".replace(",", ".")],
              ["USt 19 %", "", f"{mwst:,}".replace(",", ".")],
              ["Gesamt brutto", "", f"{betrag_eur:,}".replace(",", ".")]]),
            ("Hinweise",
             f"Konzern-IC-Rechnung gemaess Verrechnungspreis-Rahmenvereinbarung REA/REG. "
             f"Die Leistungs- und Lieferspezifikation ist in der Anlage 1 zur Rechnung "
             f"detailliert ausgewiesen (separates Dokument REG_IC_Spezifikation_{jahr}_{monat_idx:02d}). "
             f"Die Rechnung enthaelt keine Mengen-Boni oder retroaktiven Preisanpassungen. "
             f"Cash-Pool-Settlement erfolgt ueber das Konzern-Treasury-System "
             f"(Sweep-Konto Deutsche Bank AG, Wertstellung T+1)."
             + (" Die finale Version dieser Rechnung wurde nach Abstimmung "
                "mit Konzern-Tax am 12. des Folgemonats freigegeben." if final else "")
             + (" *** Diese Version ist ein Entwurf und nicht zur Buchung freigegeben. "
                "Finale Freigabe nach Plausibilisierung Group Controller. ***" if draft else "")),
            ("Bestaetigungen",
             "Bestaetigung der sachlichen Richtigkeit: Stefan Richter (Finanzcontroller REG). "
             "Bestaetigung der Genehmigung: Andreas Maier (Geschaeftsfuehrer REG). "
             "Bestaetigung der konzerneinheitlichen Bilanzierung: Florian Maier (Group "
             "Controller, REA Stuttgart)."),
            ("Unterschrift",
             signatures("Stefan Richter", "Finanzcontroller REG", "Brennhagen Elektronik GmbH",
                        "Andreas Maier", "Geschaeftsfuehrer REG", "Brennhagen Elektronik GmbH",
                        place="Heilbronn", date_str_=f"{min(15, [31,28,31,30,31,30,31,31,30,31,30,31][monat_idx-1])}.{(monat_idx % 12)+1:02d}.{jahr if monat_idx<12 else jahr+1}")),
        ])


# REG IC-Rechnungen Liste – betraege variieren leicht
REG_IC_LIST = [
    ("REG_IC_Rechnung_2020_01.docx", 2020,  1, 22_409_836),
    ("REG_IC_Rechnung_2020_02.docx", 2020,  2, 21_180_500),
    ("REG_IC_Rechnung_2020_03.docx", 2020,  3, 19_840_200),
    ("REG_IC_Rechnung_2020_04.docx", 2020,  4, 14_220_400),
    ("REG_IC_Rechnung_2020_05.docx", 2020,  5, 12_840_600),
    ("REG_IC_Rechnung_2020_06.docx", 2020,  6, 16_410_500),
    ("REG_IC_Rechnung_2020_07_WIP.docx", 2020,  7, 18_580_700),
    ("REG_IC_Rechnung_2020_09.docx", 2020,  9, 19_240_800),
    ("REG_IC_Rechnung_2020_10.docx", 2020, 10, 21_120_400),
    ("REG_IC_Rechnung_2020_11.docx", 2020, 11, 22_410_900),
    ("REG_IC_Rechnung_2020_12.docx", 2020, 12, 24_180_500),
    ("REG_IC_Rechnung_2021_01.docx", 2021,  1, 23_240_700),
    ("REG_IC_Rechnung_2021_02_v2.docx", 2021,  2, 21_840_300),
    ("REG_IC_Rechnung_2021_03.docx", 2021,  3, 24_120_400),
    ("REG_IC_Rechnung_2021_04.docx", 2021,  4, 22_840_900),
    ("REG_IC_Rechnung_2021_05.docx", 2021,  5, 23_410_500),
    ("REG_IC_Rechnung_2021_09.docx", 2021,  9, 25_180_700),
    ("REG_IC_Rechnung_2021_10.docx", 2021, 10, 26_410_400),
    ("REG_IC_Rechnung_2021_11.docx", 2021, 11, 27_240_500),
    ("REG_IC_Rechnung_2021_12.docx", 2021, 12, 28_410_900),
    ("REG_IC_Rechnung_2022_01.docx", 2022,  1, 25_840_600),
    ("REG_IC_Rechnung_2022_02.docx", 2022,  2, 24_120_400),
    ("REG_IC_Rechnung_2022_03_DRAFT.docx", 2022,  3, 26_240_700),
    ("REG_IC_Rechnung_2022_04.docx", 2022,  4, 27_180_500),
    ("REG_IC_Rechnung_2022_05.docx", 2022,  5, 26_410_800),
    ("REG_IC_Rechnung_2022_06.docx", 2022,  6, 25_840_300),
    ("REG_IC_Rechnung_2022_08_FINAL.docx", 2022,  8, 27_240_500),
    ("REG_IC_Rechnung_2022_09.docx", 2022,  9, 28_120_900),
    ("REG_IC_Rechnung_2022_10.docx", 2022, 10, 29_410_600),
    ("REG_IC_Rechnung_2022_11.docx", 2022, 11, 30_180_400),
    ("REG_IC_Rechnung_2022_12.docx", 2022, 12, 32_120_700),
]

for fn, jahr, m, betrag in REG_IC_LIST:
    is_draft = "DRAFT" in fn or "WIP" in fn
    is_final = "FINAL" in fn
    ic_rechnung_reg(fn, jahr, m, betrag, draft=is_draft, final=is_final)


# ============================================================================
# 6) REG -> RSG IC-RECHNUNGEN (28 docs) – Vor-Lieferungen REG an RSG Muenchen
# ============================================================================

def ic_reg_to_rsg(fname, jahr, monat_idx, betrag_eur, draft=False, final=False):
    monate = ["Januar","Februar","Maerz","April","Mai","Juni","Juli","August",
              "September","Oktober","November","Dezember"]
    monat = monate[monat_idx-1]
    netto = round(betrag_eur / 1.19)
    mwst = betrag_eur - netto
    rg_nr = f"REG-RSG-IC-{jahr}-{monat_idx:02d}"

    write_doc(f"{BASE}/{fname}", REG_H,
        f"IC-Rechnung {monat} {jahr} – REG Heilbronn an RSG Muenchen",
        subtitle=f"Rechnungs-Nr. {rg_nr}; Hardware-Vorprodukte fuer Software-Integration",
        draft=draft,
        sections=[
            ("Rechnungsadressat",
             f"Brennhagen Software GmbH (RSG)\n"
             f"Bayerstrasse 33, 80335 Muenchen\n"
             f"HRB 319872, Amtsgericht Muenchen\n"
             f"USt-ID: DE 312 487 902\n\n"
             f"Kontakt Group Accounts Payable RSG: Sandra Kessler, "
             f"sandra.kessler@roehrig-software.com"),
            ("Rechnungssteller",
             f"Brennhagen Elektronik GmbH (REG)\n"
             f"Industriestrasse 47, 74076 Heilbronn\n"
             f"HRB 221456, Amtsgericht Heilbronn\n"
             f"USt-ID: DE 287 612 449\n\n"
             f"Ansprechpartner Forderung: Stefan Richter (Finanzcontroller), "
             f"stefan.richter@brennhagen-elektronik.com, Tel. 07131/4823-2104"),
            ("Rechnungsdaten",
             [["Feld", "Wert"],
              ["Rechnungsnummer", rg_nr],
              ["Rechnungsdatum", f"{min(10, [31,28,31,30,31,30,31,31,30,31,30,31][monat_idx-1])}.{(monat_idx % 12)+1:02d}.{jahr if monat_idx<12 else jahr+1}"],
              ["Leistungszeitraum", f"{monat} {jahr}"],
              ["Vertragsgrundlage", "IC-Liefervereinbarung REG-RSG vom 1.1.2018 (Update 14.3.2022)"],
              ["VP-Methode", "Cost+ 6 % auf Materialkosten + 8 % auf Engineering"],
              ["Zahlungsziel", "30 Tage netto"],
              ["Lieferbedingung", "FCA Heilbronn (Incoterms 2020)"]]),
            ("Leistungsaufstellung",
             [["Pos.", "Beschreibung", "Betrag EUR"],
              ["1", "Hardware-Vorprodukte ICP-3 Cockpit-Board (Snapdragon SA8295P-Baugruppen)", f"{int(netto*0.45):,}".replace(",", ".")],
              ["2", "ECU-900 Dev-Hardware (Aurix-Baugruppen, Adapter-Boards)", f"{int(netto*0.22):,}".replace(",", ".")],
              ["3", "ADAS-V4D Engineering-Samples (Stand 4.2, Q1)", f"{int(netto*0.15):,}".replace(",", ".")],
              ["4", "BMS-12 Test-Adapter und Mess-Equipment", f"{int(netto*0.10):,}".replace(",", ".")],
              ["5", "Logistik (DHL Industrial, FCA Heilbronn)", f"{int(netto*0.03):,}".replace(",", ".")],
              ["6", "Konzern-Umlage Sonstiges", f"{int(netto*0.05):,}".replace(",", ".")],
              ["Summe Netto", "", f"{netto:,}".replace(",", ".")],
              ["USt 19 %", "", f"{mwst:,}".replace(",", ".")],
              ["Gesamt brutto", "", f"{betrag_eur:,}".replace(",", ".")]]),
            ("Hinweise / Verrechnungspreise",
             f"IC-Rechnung gemaess der internen Liefervereinbarung REG-RSG. Die "
             f"Verrechnungspreis-Methodik (Cost+ 6 % bzw. 8 %) ist im Konzern-Local-File "
             f"REG sowie im Local-File RSG dokumentiert. Im Berichtsmonat {monat} {jahr} "
             f"keine ausserplanmaessigen Anpassungen, keine Retro-Boni. Die Rechnung wird im "
             f"Cash-Pool-Settlement T+1 zwischen REG (Volksbank Heilbronn) und RSG (HypoVereins"
             f"bank Muenchen) ueber die Konzern-Treasury (Markus Pflanzer) abgewickelt."
             + (" Diese Rechnung wurde nach Abstimmung mit RSG-Engineering "
                "und RSG-Tax am 12. des Folgemonats final freigegeben." if final else "")
             + (" *** Entwurf – noch nicht zur Buchung freigegeben. Pruefung "
                "durch Group Controller (F. Maier) ausstehend. ***" if draft else "")),
            ("Bestaetigungen",
             "Sachliche Richtigkeit: Stefan Richter (Finanzcontroller REG). Genehmigung: "
             "Andreas Maier (Geschaeftsfuehrer REG). Empfangs-Bestaetigung RSG: Dr. Klaus "
             "Kessler (Werkleiter RSG). Tax-Review: Group Tax Director Dr. Heike Berger."),
            ("Unterschrift",
             signatures("Stefan Richter", "Finanzcontroller REG", "Brennhagen Elektronik GmbH",
                        "Sandra Kessler", "Accounts Payable", "Brennhagen Software GmbH",
                        place="Heilbronn", date_str_=f"{min(10, [31,28,31,30,31,30,31,31,30,31,30,31][monat_idx-1])}.{(monat_idx % 12)+1:02d}.{jahr if monat_idx<12 else jahr+1}")),
        ])


RSG_IC_LIST = [
    ("REG_to_RSG_IC_2021_02.docx", 2021,  2,  2_131_559),
    ("REG_to_RSG_IC_2021_03.docx", 2021,  3,  2_348_120),
    ("REG_to_RSG_IC_2021_04.docx", 2021,  4,  2_410_900),
    ("REG_to_RSG_IC_2021_05.docx", 2021,  5,  2_540_300),
    ("REG_to_RSG_IC_2021_06.docx", 2021,  6,  2_618_400),
    ("REG_to_RSG_IC_2021_07.docx", 2021,  7,  2_842_700),
    ("REG_to_RSG_IC_2021_08.docx", 2021,  8,  2_710_500),
    ("REG_to_RSG_IC_2021_09.docx", 2021,  9,  2_980_400),
    ("REG_to_RSG_IC_2021_10.docx", 2021, 10,  3_120_800),
    ("REG_to_RSG_IC_2021_11.docx", 2021, 11,  3_240_500),
    ("REG_to_RSG_IC_2021_12.docx", 2021, 12,  3_410_900),
    ("REG_to_RSG_IC_2022_01.docx", 2022,  1,  2_980_400),
    ("REG_to_RSG_IC_2022_02.docx", 2022,  2,  2_840_600),
    ("REG_to_RSG_IC_2022_04.docx", 2022,  4,  3_180_200),
    ("REG_to_RSG_IC_2022_05_v1.docx", 2022,  5, 3_240_600),
    ("REG_to_RSG_IC_2022_06.docx", 2022,  6,  3_310_700),
    ("REG_to_RSG_IC_2022_07.docx", 2022,  7,  3_240_800),
    ("REG_to_RSG_IC_2022_08.docx", 2022,  8,  3_420_500),
    ("REG_to_RSG_IC_2022_09.docx", 2022,  9,  3_580_900),
    ("REG_to_RSG_IC_2022_10.docx", 2022, 10,  3_640_700),
    ("REG_to_RSG_IC_2022_11.docx", 2022, 11,  3_810_400),
    ("REG_to_RSG_IC_2022_12.docx", 2022, 12,  4_120_600),
    ("REG_to_RSG_IC_2023_01.docx", 2023,  1,  3_840_700),
    ("REG_to_RSG_IC_2023_02.docx", 2023,  2,  3_720_500),
    ("REG_to_RSG_IC_2023_03.docx", 2023,  3,  3_980_900),
    ("REG_to_RSG_IC_2023_04_ENTWURF.docx", 2023,  4, 4_120_400),
    ("REG_to_RSG_IC_2023_05.docx", 2023,  5,  4_080_700),
    ("REG_to_RSG_IC_2023_07.docx", 2023,  7,  4_240_500),
    ("REG_to_RSG_IC_2023_08.docx", 2023,  8,  4_180_600),
    ("REG_to_RSG_IC_2023_09_FINAL.docx", 2023,  9, 4_310_900),
    ("REG_to_RSG_IC_2023_10.docx", 2023, 10,  4_380_700),
    ("REG_to_RSG_IC_2023_12.docx", 2023, 12,  4_510_400),
]

for fn, jahr, m, betrag in RSG_IC_LIST:
    is_draft = "ENTWURF" in fn
    is_final = "FINAL" in fn
    ic_reg_to_rsg(fn, jahr, m, betrag, draft=is_draft, final=is_final)


# ============================================================================
# 7) MIETVERTRAG BETRIEBSGELAENDE 2020
# ============================================================================

write_doc(f"{BASE}/REG_Mietvertrag_Betriebsgelaende_2020.docx", REG_H,
    "Miet-/Pachtvertrag Betriebsgelaende Werk Heilbronn – Brennhagen Elektronik GmbH",
    subtitle="Vermieterin: Industriegrundbesitz Heilbronn KG; Mieterin: Brennhagen Elektronik GmbH",
    sections=[
        ("Vertragsparteien",
         "Vermieterin: Industriegrundbesitz Heilbronn GmbH & Co. KG, vertreten durch die persoenlich "
         "haftende Gesellschafterin Industriegrundbesitz Heilbronn Verwaltungs-GmbH, diese vertreten "
         "durch den Geschaeftsfuehrer Herrn Dr. Bernd Voelker; Sitz: Wilhelmstrasse 21, 74072 "
         "Heilbronn; HRA 6231, Amtsgericht Heilbronn.\n\n"
         "Mieterin: Brennhagen Elektronik GmbH, vertreten durch den Geschaeftsfuehrer Herrn Andreas Maier; "
         "Sitz: Industriestrasse 47, 74076 Heilbronn; HRB 221456, Amtsgericht Heilbronn.\n\n"
         "Konzernzugehoerigkeit Mieterin: 100 %-Tochter der Brennhagen Holding GmbH (RHO Stuttgart) im "
         "Konzern Brennhagen Elektronik AG (REA), boersennotiert Frankfurter Wertpapierboerse Prime "
         "Standard."),
        ("Vertragsklauseln",
         ("clauses", [
            ("§ 1 Mietobjekt", [
                "Vermietet wird das Betriebsgelaende Industriestrasse 47, 74076 Heilbronn, "
                "Flurstueck 8124/3 der Gemarkung Heilbronn, Gesamt-Grundflaeche 19.481 m². "
                "Das Gelaende umfasst: (a) Produktionshalle 1 (Bauteil A, ca. 8.420 m² Hallenflaeche, "
                "Baujahr 1994, refurbished 2018); (b) Produktionshalle 2 (Bauteil B, ca. 5.940 m², "
                "Baujahr 2012, gebaut speziell fuer SMD-Reinraum-Bereiche, ISO Klasse 8); "
                "(c) Buero- und Verwaltungsgebaeude (Bauteil C, 4 Stockwerke, ca. 2.840 m² BGF); "
                "(d) Logistikhalle inkl. Lagerregale (Bauteil D, ca. 1.420 m²); (e) Aussenanlagen "
                "(Parkplaetze 220 Stellplaetze, Anlieferzonen, Mitarbeiter-Pausenflaechen).",
                "Die uebergebene Sachsubstanz ist im Uebergabeprotokoll vom 14.12.2019 detailliert "
                "dokumentiert (separate Anlage 1). Zur Mietsache gehoeren ferner die fest verbauten "
                "haustechnischen Anlagen (Heizung, Lueftung, Sprinkler, Druckluft, Mittelspannungs"
                "trafo-Station 800 kVA).",
                "Die Mieterin nutzt das Mietobjekt fuer die Produktion und Montage elektronischer "
                "Komponenten der Plattformen ICP-3, BMS-12, ECU-900, ADAS-V4D und LightCtrl-7 fuer "
                "Automotive-OEM-Kunden."]),
            ("§ 2 Mietzins", [
                "Die monatliche Grundmiete betraegt 200.000 EUR (in Worten: zweihunderttausend Euro) "
                "zzgl. der jeweils gueltigen gesetzlichen Umsatzsteuer (derzeit 19 %, somit Bruttomiete "
                "238.000 EUR/Monat). Der Mietzins ist jeweils im Voraus bis zum 3. Werktag eines "
                "jeden Kalendermonats faellig auf das im § 19 bezeichnete Konto der Vermieterin.",
                "Die Miete wird jaehrlich, erstmals zum 1. Januar 2022, um 2,0 % erhoeht (fester "
                "Steigerungsindex). Alternative Indexklausel: VPI Baden-Wuerttemberg fuer Wirtschafts"
                "immobilien, Anpassung jeweils bei Indexaenderung > 5 Punkte; in jedem Fall der "
                "hoehere Wert. Die jaehrliche Mieterhoehung ist der Mieterin schriftlich anzuzeigen.",
                "Nebenkosten (Heizung, Lueftung, Strom Gemeinflaechen, Wasser, Abwasser, Muellabfuhr, "
                "Gebaeudereinigung, Versicherung Gebaeude, Wartung haustechnische Anlagen) werden "
                "anteilig nach Quadratmeter abgerechnet; monatliche Vorauszahlung 35.000 EUR; "
                "Endabrechnung jeweils zum 31. Mai des Folgejahres."]),
            ("§ 3 Mietzeit und Kuendigung", [
                "Das Mietverhaeltnis beginnt am 1. Januar 2020 und ist auf 15 Jahre festgeschlossen "
                "(Festlaufzeit bis 31. Dezember 2034). Eine ordentliche Kuendigung waehrend der "
                "Festlaufzeit ist ausgeschlossen.",
                "Die Mieterin erhaelt zwei Verlaengerungsoptionen von jeweils 5 Jahren (somit "
                "verlaengerbar bis 31.12.2044). Die Optionsausuebung muss schriftlich spaetestens "
                "12 Monate vor Ablauf der jeweiligen Laufzeit erfolgen.",
                "Das Recht zur ausserordentlichen Kuendigung aus wichtigem Grund bleibt unberuehrt. "
                "Insbesondere kann die Mieterin den Vertrag fristlos kuendigen, wenn die Mietsache "
                "infolge nicht durch die Mieterin zu vertretenden Ereignissen (Brand, Wasser-"
                "schaden, behoerdliche Anordnungen) ueber 90 Tage hinaus nicht nutzbar ist."]),
            ("§ 4 Erhaltungslast / Instandhaltung", [
                "Die Vermieterin traegt die Dach-und-Fach-Instandhaltung sowie die Erhaltung der "
                "Substanz der Gebaeudehuelle und der Tragkonstruktion. Saemtliche Instandhaltungs"
                "massnahmen im Inneren der Mietsache (Bodenbelaege, Trennwaende, Beleuchtung, "
                "Sanitaeranlagen, etc.) sowie die Wartung der haustechnischen Anlagen obliegen "
                "der Mieterin.",
                "Aenderungen am Mietobjekt (insbesondere Reinraum-Erweiterungen, neue SMD-Linien, "
                "neue Mediennetze) beduerfen der vorherigen schriftlichen Zustimmung der "
                "Vermieterin. Die Vermieterin wird ihre Zustimmung nicht ohne wichtigen Grund "
                "verweigern."]),
            ("§ 5 Kaution / Sicherheiten", [
                "Die Mieterin stellt eine Bankbuergschaft eines deutschen Kreditinstituts in Hoehe "
                "von 600.000 EUR (3 Monatsmieten netto). Die Buergschaft wird bei Vertragsbeginn "
                "gestellt; Verguetung der Buergschaft traegt die Mieterin.",
                "Die Mieterin verpflichtet sich, das Mietobjekt auf eigene Kosten gegen "
                "Feuer (Industrieversicherung), Leitungswasser und Sturm/Hagel zu versichern "
                "(Police HDI Global SE, VS 140 Mio. EUR, Mitversicherung Vermieterin)."]),
            ("§ 6 Rueckgabe", [
                "Bei Beendigung des Mietverhaeltnisses ist das Mietobjekt im Zustand zurueck"
                "zugeben, in dem es uebergeben wurde, abzueglich normaler Abnutzung. Mieter-"
                "spezifische Einbauten (insbesondere SMD-Reinraeume) koennen nach Wahl der "
                "Vermieterin gegen Restbuchwert-Ausgleich uebernommen oder zurueckgebaut werden.",
                "Ein Uebergabeprotokoll wird gemeinsam erstellt; etwaige Maengel sind binnen "
                "4 Wochen zu beseitigen."]),
            ("§ 7 Schlussbestimmungen", [
                "Aenderungen und Ergaenzungen dieses Vertrages beduerfen der Schriftform; dies "
                "gilt auch fuer die Aufhebung der Schriftformklausel. Sollten einzelne Bestimmungen "
                "unwirksam sein, bleibt die Wirksamkeit der uebrigen Bestimmungen unberuehrt. "
                "Gerichtsstand ist Heilbronn. Es gilt deutsches Recht."]),
         ])),
        ("Unterschriften",
         signatures("Dr. Bernd Voelker", "Geschaeftsfuehrer", "Industriegrundbesitz Heilbronn GmbH & Co. KG",
                    "Andreas Maier", "Geschaeftsfuehrer", "Brennhagen Elektronik GmbH",
                    place="Heilbronn", date_str_="14. Dezember 2019")),
    ])


# ============================================================================
# 8) STEUERBESCHEIDE (4 docs)
# ============================================================================

def steuerbescheid(fname, jahr, kst_eur, gewst_eur, anmerkung="", draft=False):
    BESCHEID_H = {
        "name": "Finanzamt Heilbronn",
        "addr": "Cae celiusstrasse 1, 74072 Heilbronn  |  Sachgebiet Koerperschaften",
        "hrb":  "Bundeszentralamt fuer Steuern - Behoerdenkennung 28840",
    }
    write_doc(f"{BASE}/{fname}", BESCHEID_H,
        f"Steuerbescheid Koerperschaftsteuer / Gewerbesteuer {jahr} – Brennhagen Elektronik GmbH",
        subtitle=f"Steuernummer 65085/12345; Veranlagungszeitraum 1.1.{jahr} – 31.12.{jahr}",
        draft=draft,
        sections=[
            ("Bescheidempfaengerin",
             f"Brennhagen Elektronik GmbH (»Steuerpflichtige«)\n"
             f"Industriestrasse 47, 74076 Heilbronn\n"
             f"HRB 221456, Amtsgericht Heilbronn\n"
             f"Steuernummer: 65085/12345\n"
             f"Steuer-Identifikationsnummer: DE 287 612 449\n\n"
             f"Steuerberatung: Deloitte GmbH WPG, Buero Stuttgart (verantwortlicher StB Herr "
             f"Dr. Michael Wenzel)"),
            ("Veranlagungsgrundlage",
             f"Auf Grundlage der eingereichten Koerperschaftsteuer-Erklaerung {jahr} (Eingang FA "
             f"Heilbronn am 31.7.{jahr+1}; mit Bewertung der Wirtschaftspruefer KPMG AG WPG "
             f"vom 12.2.{jahr+1}) sowie nach erfolgter steuerlicher Pruefung wird folgende "
             f"Veranlagung vorgenommen:"),
            ("Ermittlung des zu versteuernden Einkommens (EUR)",
             [["Position", "Betrag"],
              ["Jahresueberschuss laut HGB-Abschluss", f"{int(kst_eur * 6.5):,}".replace(",", ".")],
              ["+ nicht abziehbare Betriebsausgaben (§ 4 Abs. 5 EStG)", "182.400"],
              ["+ Hinzurechnungen Gewerbesteuer (§ 8 GewStG)", f"{int(gewst_eur * 0.85):,}".replace(",", ".")],
              ["- Steuerfreie Einnahmen / Schachtelprivilegien (§ 8b KStG)", "-450.200"],
              ["- Spenden / sonstige Abzuege", "-28.400"],
              ["Steuerpflichtiges Einkommen", f"{int(kst_eur / 0.15):,}".replace(",", ".")]]),
            ("Festsetzung",
             [["Steuerart", "Betrag EUR", "Hinweis"],
              ["Koerperschaftsteuer (15,0 %)", f"{int(kst_eur):,}".replace(",", "."), "festgesetzt"],
              ["Solidaritaetszuschlag (5,5 % auf KSt)", f"{int(kst_eur * 0.055):,}".replace(",", "."), "festgesetzt"],
              ["Gewerbesteuer (Messbetrag x Hebesatz 420 %)", f"{int(gewst_eur):,}".replace(",", "."), "Stadt Heilbronn"],
              ["Gesamtsteuerbelastung", f"{int(kst_eur + kst_eur*0.055 + gewst_eur):,}".replace(",", "."), "ca. {:.1f} %".format((kst_eur + kst_eur*0.055 + gewst_eur)/(kst_eur/0.15)*100)]]),
            ("Bereits geleistete Vorauszahlungen",
             f"Geleistete Vorauszahlungen im Veranlagungszeitraum {jahr}: "
             f"{int((kst_eur + kst_eur*0.055 + gewst_eur) * 0.92):,} EUR ".replace(",", ".") +
             f"(verteilt auf 10. 3. / 10. 6. / 10. 9. / 10. 12.). Abschlusszahlung bzw. Erstattung "
             f"siehe Bescheid-Anlage 1.\n\n"
             f"Vorauszahlungs-Festsetzung fuer {jahr+1}: ca. "
             f"{int((kst_eur + kst_eur*0.055 + gewst_eur) * 1.05):,} EUR (gleichmaessig auf 4 Termine).".replace(",", ".")),
            ("Anmerkungen / Pruefungsfeststellungen",
             anmerkung or f"Keine wesentlichen Anpassungen gegenueber der eingereichten Steuererklaerung. "
             f"Die Veranlagung erfolgt unter dem Vorbehalt der Nachpruefung (§ 164 AO). Eine Betriebs"
             f"pruefung gemaess § 193 AO ist fuer den Veranlagungszeitraum {jahr} bisher nicht "
             f"angeordnet."),
            ("Rechtsbehelfsbelehrung",
             f"Gegen diesen Bescheid kann innerhalb eines Monats nach Bekanntgabe Einspruch beim "
             f"Finanzamt Heilbronn (Anschrift wie Bescheidkopf) eingelegt werden. Der Einspruch "
             f"ist schriftlich oder zur Niederschrift einzulegen. Die Einlegung des Einspruchs hat "
             f"grundsaetzlich keine aufschiebende Wirkung; ggf. ist ein gesonderter Antrag auf "
             f"Aussetzung der Vollziehung (§ 361 AO) zu stellen."),
            ("Unterschrift Finanzamt",
             signatures("Frau Antje Weller", "Sachbearbeiterin Sachgebiet KSt", "Finanzamt Heilbronn",
                        "Herr Dr. Stefan Klotz", "Sachgebietsleiter", "Finanzamt Heilbronn",
                        place="Heilbronn", date_str_=f"15. November {jahr+1}")),
        ])


steuerbescheid("REG_Steuerbescheid_2020.docx", 2020, 1_842_500, 1_580_300,
    "Anpassung Hinzurechnungen Dauerschuldzinsen (§ 8 Nr. 1 GewStG) +120 TEUR gegenueber Erklaerung "
    "infolge Korrektur der Konsortialkredit-Zinsen Bereitstellungsprovision (vorher unter sonstigen "
    "Aufwendungen). Keine weitere materielle Aenderung.")
steuerbescheid("REG_Steuerbescheid_2021.docx", 2021, 2_120_400, 1_840_500,
    "Veranlagung erfolgt erklaerungsgemaess; geringfuegige formelle Anpassungen Konzern-Umlagen. "
    "Nachfrage bzgl. Verrechnungspreise REG -> RHO durch Beleg Local File 2021 (PwC) erledigt.")
steuerbescheid("REG_Steuerbescheid_2022_ENTWURF.docx", 2022, 2_410_900, 2_120_700,
    "*** Bescheidentwurf zur Kommentierung an Steuerberatung Deloitte – Frist 30.1.2024. *** "
    "Punkt zur Diskussion: Behandlung der IPO-Vorbereitungskosten (anteilig REG, ca. 480 TEUR), "
    "die in der KSt-Erklaerung als nicht abziehbar deklariert wurden; FA pruegt zusaetzlich "
    "Abgrenzung Eigenkapital-Beschaffungskosten.",
    draft=True)
steuerbescheid("REG_Steuerbescheid_KSt_2023.docx", 2023, 2_640_300, 2_240_900,
    "Veranlagung 2023 unter Beruecksichtigung der Reorganisation Konzernumlagen ab 1.7.2023 "
    "(neue Cost+ Methode 7 % anstelle 5 % auf Engineering-Leistungen, Update Local File 2023). "
    "Keine Beanstandungen seitens FA. Veranlagung gilt vorlaeufig unter dem Vorbehalt der "
    "Nachpruefung (Betriebspruefung 2022-2024 in Vorbereitung, voraussichtlich Q3/2025).")


# ============================================================================
# 9) WP MANAGEMENT LETTERS (3 docs, 2021-2023)
# ============================================================================

def wp_ml(fname, jahr, feststellungen):
    write_doc(f"{BASE}/{fname}", KPMG_H,
        f"Management Letter {jahr} – Brennhagen Elektronik GmbH (REG)",
        subtitle=f"Empfehlungen aus der Jahresabschlusspruefung {jahr} der REG Heilbronn",
        confidential=True,
        sections=[
            ("Adressat",
             f"Geschaeftsfuehrung der Brennhagen Elektronik GmbH\n"
             f"Herrn Andreas Maier (Geschaeftsfuehrer)\n"
             f"Industriestrasse 47, 74076 Heilbronn\n\n"
             f"In Kopie: Group CFO Frau Laura Bauer; Group Controller Herr Florian Maier; "
             f"Local Compliance Officer; Pruefungsausschuss-Vorsitz Prof. Dr.-Ing. Gerhard Voss "
             f"(REA-Aufsichtsrat)"),
            ("Vorbemerkungen",
             f"Im Rahmen der Pruefung des Jahresabschlusses {jahr} der Brennhagen Elektronik GmbH "
             f"(Werk Heilbronn) – durchgefuehrt nach IDW PS 200 und ueberschneidend mit der "
             f"Konzernabschluss-Pruefung der Brennhagen Elektronik AG nach IFRS – haben wir die "
             f"folgenden Feststellungen und Empfehlungen zur Optimierung des internen Kontroll"
             f"systems (IKS), der Bilanzierungsprozesse und der Risikolandschaft erarbeitet. "
             f"Die Pruefung erfolgte durch das KPMG-Pruefungsteam mit Lead Partner "
             f"Dr. Maximilian Brand sowie dem Lead Manager Herrn Felix Schwarzbauer. Pruefungs"
             f"zeitraum vor Ort 14.1. – 12.2. des Folgejahres."),
            ("Wesentliche Feststellungen",
             feststellungen),
            ("Folgewirkung Pruefung",
             "Die Feststellungen haben keine Auswirkung auf die uneingeschraenkte Erteilung des "
             "Bestaetigungsvermerks. Saemtliche Punkte sind als Verbesserungsempfehlungen zu "
             "verstehen. Wir empfehlen die Aufnahme in den IKS-Massnahmenplan REG Heilbronn "
             "sowie die Berichterstattung an den Pruefungsausschuss in der naechsten Sitzung."),
            ("Stellungnahme der Geschaeftsfuehrung",
             "Die Geschaeftsfuehrung der REG (Andreas Maier) sowie der Finanzcontroller Stefan "
             "Richter haben den Feststellungen und Empfehlungen zugestimmt. Massnahmenplan und "
             "Verantwortlichkeiten werden im Q1 des Folgejahres festgelegt und an KPMG zur "
             "Kenntnis vorgelegt. Umsetzungs-Review im Rahmen der naechsten Jahresabschluss"
             "pruefung."),
            ("Unterschriften",
             signatures("Dr. Maximilian Brand", "Wirtschaftspruefer / Partner", "KPMG AG WPG",
                        "Felix Schwarzbauer", "Manager / Pruefungsleiter", "KPMG AG WPG",
                        place="Stuttgart", date_str_=f"28. Februar {jahr+1}")),
        ])


wp_ml("REG_WP_Management_Letter_2021.docx", 2021,
    ("list", [
        "(1) Optimierung Reisekosten-Workflow: derzeit teils Papier-Belege; Empfehlung Roll-Out "
        "Concur Mobile mit Beleg-Foto und OCR (Umsetzung bis Q3 2022; verantwortlich HR-Manager "
        "Petra Holzwarth).",
        "(2) Vier-Augen-Prinzip bei Banking-Releases: in Stichproben einzelne Faelle mit nur "
        "einer Freigabe in SAP-FBL (Workflows). Empfehlung: Zwangs-Workflow ab 50 TEUR Konto-"
        "ausgehend (Konfiguration SAP S/4HANA bis Q2 2022).",
        "(3) Inventur-Differenzen Werks-Lager: Inventur 30.11.2021 zeigte ca. 0,8 % Differenz "
        "bei elektronischen Vorprodukten; Empfehlung Einfuehrung cycle counting (rollierende "
        "Inventur) und Verschaerfung der ZBV-Buchungen.",
        "(4) Dokumentation der Manuell-Buchungen Monatsabschluss: Empfehlung verstaerkte "
        "Pruefkommentare im SAP-Buchungsbeleg, insbesondere bei Periodenabgrenzungen > 50 TEUR.",
    ]))

wp_ml("REG_WP_Management_Letter_2022.docx", 2022,
    ("list", [
        "(1) Follow-up: Concur-Roll-Out abgeschlossen (Mai 2022); keine wesentlichen Belege "
        "in Papierform mehr.",
        "(2) Follow-up: Banking Vier-Augen-Prinzip in SAP konfiguriert (Juli 2022); KPMG-"
        "Stichproben keine Auffaelligkeiten.",
        "(3) Inventur 30.11.2022 mit 0,3 % Differenz – deutlich verbessert; cycle counting "
        "wirksam.",
        "(4) Neu: IFRS 15 Umsatzabgrenzungen Tooling-Erloese fuer BMW ICP-3 (Auftrag 4,8 Mio. "
        "EUR) zeigen Anpassungsbedarf in der Allokation auf Performance Obligations; Empfehlung "
        "Klaerung mit Group Controlling (F. Maier) und Anpassung im Q1 2023 Reporting.",
        "(5) Neu: SOX-aehnliche Kontrollen IT-Zugriffsmanagement Werks-MES (Manufacturing "
        "Execution System) – Empfehlung halbjaehrliche User-Access-Reviews (UAR).",
        "(6) Erinnerung: Aktualisierung des IKS-Handbuchs REG (zuletzt 2019) – ueberfaellig; "
        "Empfehlung Update bis 30.9.2023.",
    ]))

wp_ml("REG_WP_Management_Letter_2023.docx", 2023,
    ("list", [
        "(1) Follow-up: IFRS 15 Tooling-Allokation umgesetzt Q2 2023; keine weiteren "
        "Feststellungen.",
        "(2) Follow-up: UAR halbjaehrlich etabliert; KPMG bestaetigt Wirksamkeit. "
        "Verbesserungspunkt: zusaetzliche UAR fuer SAP S/4HANA Power-User (CFO Office) jaehrlich.",
        "(3) Neu: Erstmalige Pruefung CSRD-Bereitschaft (Vorbereitung Erstanwendung GJ 2024 fuer "
        "REA-Konzern) – Empfehlung Aufbau ESG-Datenbank in REG (Scope 1+2 Emissionen, Wasser, "
        "Abfall, Frauenanteil, Trainingsstunden). Verantwortlich: HR + EHS-Manager.",
        "(4) Neu: Quartalsweise Plausibilisierung Verrechnungspreis-Margen (Cost+ Engineering) "
        "ggu. Quartile der TP-Studie PwC 2023; Empfehlung Aufnahme als IKS-Kontrolle.",
        "(5) Neu: IT-Security – Penetrationstest Werks-OT (Operational Technology) durch Konzern-"
        "IT (Bayerische Security GmbH) zeigte 3 mittelkritische Findings auf der SMD-Linie 4 "
        "(Reflow-Ofen-Steuerung); Empfehlung Umsetzung bis 30.6.2024.",
        "(6) Hinweis: Bestaetigungsvermerk Jahresabschluss 2023 wurde am 12.2.2024 ohne "
        "Einschraenkung erteilt; uneingeschraenktes Pruefungsurteil.",
    ]))


# ============================================================================
# 10) VERSICHERUNGSNACHWEIS 2023
# ============================================================================

write_doc(f"{BASE}/REG_Versicherungsnachweis_2023.docx", REG_H,
    "Versicherungsnachweis 2023 – Brennhagen Elektronik GmbH (REG)",
    subtitle="Versicherungsuebersicht Werk Heilbronn; Stand 1.1.2023; verlaengert bis 31.12.2023",
    sections=[
        ("Zweck des Nachweises",
         "Dieser Versicherungsnachweis dokumentiert den Versicherungsschutz der Brennhagen Elektronik "
         "GmbH (REG) Werk Heilbronn fuer das Geschaeftsjahr 2023. Der Nachweis wird an die folgenden "
         "Empfaenger geliefert: (1) Group Risk Management der Brennhagen Elektronik AG (Frau Christine "
         "Meinhardt); (2) Pruefungsausschuss des Aufsichtsrates der REA (Prof. Voss); (3) Vermieterin "
         "des Betriebsgelaendes (Industriegrundbesitz Heilbronn KG) als Mitversicherte der Sachversicherung; "
         "(4) Hausbanken im Rahmen des Konsortialkredit-Reportings (auf Anforderung). Die Versicherungen "
         "sind ueberwiegend Bestandteil der Brennhagen-Konzern-Police, ergaenzt um lokale Spezial-"
         "Policen fuer das Werk Heilbronn."),
        ("Versicherungsuebersicht",
         [["Sparte", "Versicherer", "Police-Nr.", "Deckungssumme", "Praemie p.a.", "Selbstbehalt"],
          ["Betriebshaftpflicht (BHV)", "Allianz Versicherungs-AG", "BHV-RHG-2023-001", "50.000.000 EUR", "135.000 EUR", "25.000 EUR"],
          ["Produkthaftpflicht inkl. erw. Produkthaftpflicht (Recall)", "AXA XL (Lloyd's-Syndikat 1209)", "PHV-AXA-2023-RHG-04", "100.000.000 EUR", "312.000 EUR", "100.000 EUR"],
          ["Industrie-Feuerversicherung (Gebaeude + Inhalt)", "HDI Global SE", "FEU-HDI-2023-HBN-12", "140.000.000 EUR", "164.000 EUR", "50.000 EUR"],
          ["Maschinen-Bruchversicherung (SMD-Linien)", "HDI Global SE", "MB-HDI-2023-HBN-08", "28.000.000 EUR", "62.000 EUR", "10.000 EUR"],
          ["Betriebsunterbrechung (BU; 18 Monate Haftzeit)", "Allianz Global Corporate & Specialty SE", "BU-AGCS-2023-RHG", "85.000.000 EUR", "98.000 EUR", "5 Werktage"],
          ["Transportversicherung (offene Police)", "Generali Versicherung AG", "TRP-GEN-2023-OFF", "10.000.000 EUR / Reise", "61.000 EUR", "2.500 EUR"],
          ["Cyber-Versicherung (Konzern-Police)", "Munich Re / Hartford", "CYBER-RHG-2023-MAIN", "30.000.000 EUR", "82.000 EUR", "100.000 EUR"],
          ["D&O-Versicherung (Geschaeftsfuehrer-Police REG)", "Allianz Global Corporate & Specialty SE", "DO-AGCS-2023-RHG-SUB-REG", "10.000.000 EUR", "n/a (Konzern-Tragung)", "n/a"],
          ["Verkehrshaftpflicht Fuhrpark (35 Pkw + 8 LKW)", "HUK-COBURG", "KH-HUK-2023-FLOTTE-REG", "Pflichtdeckung + 100 Mio. EUR Excess", "48.000 EUR", "variabel"],
          ["Umwelthaftpflicht / -schadens-VS", "R+V Allgemeine Versicherung AG", "UHV-RV-2023-HBN-3", "20.000.000 EUR", "28.000 EUR", "25.000 EUR"]]),
        ("Praemien-Summe und Erlaeuterung",
         "Gesamtpraemien-Summe REG Werk Heilbronn 2023: ca. 990.000 EUR p.a. (ohne D&O, die ueber "
         "Konzern-Police getragen wird). Anstieg ggu. 2022 (+8,4 %) im Wesentlichen aufgrund "
         "Erhoehung der Sachversicherungs-Praemien (HDI Global) infolge Marktentwicklung sowie "
         "Anpassung der Cyber-Versicherung an erhoehte Risiko-Bewertung der Branche.\n\n"
         "Die Brennhagen-Konzern-Police wird zentral durch die Group Risk Management Funktion (REA "
         "Stuttgart, Christine Meinhardt) gesteuert; Broker Marsh GmbH (Frankfurt). Verlaengerung "
         "jaehrlich zum 1.1.; Mandat Marsh bis 31.12.2025 vereinbart."),
        ("Schadenstatistik 2023",
         "Im Berichtsjahr 2023 wurden 6 Schaeden in der REG gemeldet: (1) Wassereinbruch Halle B "
         "infolge Starkregen Juni 2023 (Schaden ca. 142 TEUR, FEU-HDI, abgewickelt); (2) Maschinen"
         "schaden Reflow-Ofen Linie 4 (Schaden 38 TEUR, MB-HDI, abgewickelt); (3) Kfz-Sachschaden "
         "Werks-Lkw (4,2 TEUR, KH-HUK, abgewickelt); (4) BHV-Personenschaden Besucher (1,8 TEUR, "
         "BHV-Allianz, abgewickelt); (5) BU-Vorfall Linie 4 (5 Werktage Stillstand, Selbstbehalt "
         "voll, keine Entschaedigung BU-AGCS); (6) Recall-Andeutung BMW LightCtrl-7 (in Pruefung, "
         "noch nicht entschieden). Schadenquote bezogen auf Praemien-Volumen: ca. 19 % – "
         "vergleichsweise guenstig."),
        ("Bestaetigung",
         "Die Geschaeftsfuehrung der REG bestaetigt die Vollstaendigkeit und Richtigkeit der "
         "Versicherungs-Uebersicht zum Stichtag 1. Januar 2023. Alle aufgefuehrten Policen sind "
         "in Kraft und die Praemien wurden fristgerecht entrichtet."),
        ("Unterschrift",
         signatures("Andreas Maier", "Geschaeftsfuehrer REG", "Brennhagen Elektronik GmbH",
                    "Christine Meinhardt", "Group Risk Management", "Brennhagen Elektronik AG",
                    place="Heilbronn", date_str_="12. Januar 2023")),
    ])


# ============================================================================
# 11) PRJ-Dokumente (2 docs – Testbericht ADAS-V4D + Gate G3 ICP-3)
# ============================================================================

write_doc(f"{BASE}/PRJ-2023-001_Testbericht_Funktionstest_EOL_ADAS-V4D_Radar_Fusio.docx", REA_H,
    "Pruefbericht Funktionstest EOL – ADAS-V4D Radar Fusion (PRJ-2023-001)",
    subtitle="Endpruefung der Serienproduktion Werk Heilbronn; Test-Lot W14-2023",
    sections=[
        ("Projekt-Kontext",
         "Projekt PRJ-2023-001 betrifft die Plattform ADAS-V4D (Radar Fusion Steuergeraet, Level-2/3 "
         "ADAS) der Brennhagen Elektronik GmbH Werk Heilbronn. Hauptkunden: Mercedes-Benz Group AG "
         "(MBSP-Programm GLA/GLB) und Stellantis N.V. (DS-Modell-Reihe). Plattform-SOP 03/2022; "
         "aktueller Revisionsstand des Produkts 4.2 (Mod-Index 14.2.2023). Testbericht erstellt "
         "fuer Test-Lot W14-2023 (KW 14, 3.4. – 9.4.2023; Losgroesse 4.820 Stueck)."),
        ("Pruefumfang",
         "Endpruefung End-of-Line (EOL) gemaess Control Plan REG_ControlPlan_ADAS-V4D_2023 "
         "(Revisionsstand 4.2). Pruefungen umfassen: (a) Flying-Probe ICT (Takaya APT-1400F); "
         "(b) FCT-EOL inkl. CAN/FlexRay/Ethernet-Diagnose; (c) EMV-Vorpruefung gemaess CISPR 25 "
         "Klasse 5 (Stichprobe 1:200); (d) Verguss-Dichtheit IP67 (Helium-Lecksuche); "
         "(e) Burn-In 168h @ 85 °C / 95 % rF (Stichprobe 1:500 fuer Mercedes-Auslieferung)."),
        ("Testergebnis",
         [["Prueftest", "Pruefnorm / Anforderung", "Ergebnis", "Status"],
          ["Flying-Probe ICT", "100 % Komponenten OK", "100 %", "Bestanden"],
          ["FCT-EOL CAN/FlexRay", "0 DTC, Diagnostic Profile P2.0", "0 DTC", "Bestanden"],
          ["Ethernet-Konnektivitaet 100BASE-T1", "Link-Up < 200 ms", "142 ms (Mittel)", "Bestanden"],
          ["EMV-Vorpruefung Stoerstrahlung", "CISPR 25 Kap. 8 Klasse 5", "alle Frequenzen i.O.", "Bestanden"],
          ["EMV-Vorpruefung Stoerfestigkeit", "ISO 11452-2/4", "alle Frequenzen i.O.", "Bestanden"],
          ["Verguss-Dichtheit", "IP67 (Helium < 1e-6 mbar*l/s)", "1,2e-7 mbar*l/s (Mittel)", "Bestanden"],
          ["Burn-In 168h @ 85 °C", "kein Funktions-Drift", "alle Stichproben i.O.", "Bestanden"],
          ["Korrosionspruefung Salzspruehnebel", "ASTM B117, 480h (Stellantis CSR)", "i.O.", "Bestanden"]]),
        ("Statistische Auswertung",
         "Yield-Rate Lot W14-2023: 99,82 % (8 NIO bei 4.820 Stueck Eingang). NIO-Verteilung: "
         "3 x FCT-EOL (CAN-Bus-Diagnose, Re-Test bestanden), 2 x ICT (Bauteil-Lift), 2 x EMV-"
         "Vorpruefung (Stichproben, re-test bestanden), 1 x Verguss-Bubble (Re-Sealing erfolgreich). "
         "Endgueltige NIO-Quote nach Re-Test: 0 Stueck; Yield 100 % nach Nacharbeit. "
         "PPM-Quote: 0 ppm im Lot."),
        ("Freigabe",
         "Auf Basis der Testergebnisse erfolgt die Freigabe des Lots W14-2023 zur Auslieferung an "
         "Mercedes-Benz (Werk Sindelfingen, Anlieferung 16.4.2023) und Stellantis (Werk Hordain/F, "
         "Anlieferung 17.4.2023). Freigabe durch QS-Engineering (Dr.-Ing. Kerstin Roggner) sowie "
         "Q-Leitung Werk (Sabine Brand) erteilt am 11.4.2023."),
        ("Anlagen",
         "(1) Test-Protokolle EOL pro Stueck (digital, abgelegt in Brennhagen-QS-Datenbank, Lot W14-2023); "
         "(2) Statistical Process Control Auswertung (Cpk-Werte je Pruefparameter); (3) EMV-"
         "Pruefprotokoll Schaffner; (4) Burn-In-Protokoll Stichprobe; (5) Freigabe-Liste mit "
         "Seriennummern."),
        ("Unterschrift",
         signatures("Dr.-Ing. Kerstin Roggner", "QS-Engineering", "Brennhagen Elektronik GmbH",
                    "Sabine Brand", "Q-Leitung Werk Heilbronn", "Brennhagen Elektronik GmbH",
                    place="Heilbronn", date_str_="11. April 2023")),
    ])


write_doc(f"{BASE}/PRJ-2023-003_Gate_G3_Detailentwicklung_ICP-3_OTA_Update_Platform.docx", REA_H,
    "Stage-Gate-Review G3 – Detailentwicklung ICP-3 OTA Update Platform (PRJ-2023-003)",
    subtitle="Gate G3 (Detailentwicklung -> Verifikation); Review-Sitzung am 24.5.2023, Heilbronn",
    sections=[
        ("Projektkontext",
         "Projekt PRJ-2023-003 erweitert die Plattform ICP-3 (InfoConnect Pro Infotainment) um eine "
         "Over-the-Air-Update-Faehigkeit (FOTA). Stakeholder-OEM: BMW Group (Lead-Kunde, Programm "
         "G05 LCI 2025), Mercedes-Benz (Sekundaer-Kunde, MMA-Plattform). Projekt-Budget: 12,4 Mio. "
         "EUR (genehmigt durch REA-Vorstand 14.11.2022). Projektleiter Werk Heilbronn: "
         "Dr.-Ing. Patrick Sandner; Konzern-Sponsoring: CTO Dr. Petra Hollmann (ab 1.7.2024); "
         "bis dahin uebergangsweise COO Dr. Thomas Weber."),
        ("Gate-Definition",
         "Gate G3 entspricht im Konzern-Stage-Gate-Modell dem Uebergang von der Detail-Entwicklungs"
         "phase (D) zur Verifikations- und Validierungsphase (V&V). Voraussetzungen fuer Freigabe: "
         "(a) abgeschlossene Spezifikation Software-Architektur (SDS Rev. 2.0); (b) abgeschlossene "
         "Hardware-Aenderungen ICP-3 fuer FOTA (zusaetzlicher 64 GB eUFS Speicher, Redundanz-Boot-"
         "Partition); (c) abgeschlossener Cybersecurity-Konzept-Review (ISO/SAE 21434); (d) "
         "Lieferanten-Freigaben fuer Schluesselkomponenten."),
        ("Status zum Gate-Datum 24.5.2023",
         [["Liefergegenstand", "Soll", "Ist", "Status"],
          ["SDS Software-Architektur Rev. 2.0", "abgeschlossen", "Rev. 2.0 final 12.5.2023", "OK"],
          ["Hardware-Aenderungspaket ICP-3-Mod-4 (eUFS-Boot)", "abgeschlossen", "EE-Freigabe 18.5.2023", "OK"],
          ["Cybersecurity-Konzept ISO/SAE 21434 (TARA)", "abgeschlossen", "Review TUEV Sued 22.5.2023", "OK"],
          ["Cryptographic Key Management Concept", "abgeschlossen", "Final 19.5.2023", "OK"],
          ["Lieferanten-Freigabe Micron LPDDR5 12 GB", "abgeschlossen", "PPAP Level 3 freigegeben", "OK"],
          ["Lieferanten-Freigabe Samsung eUFS 4.0 64 GB", "abgeschlossen", "PPAP Level 3 freigegeben", "OK"],
          ["Lieferanten-Freigabe Qualcomm SA8295P (Bestand)", "Bestaetigung", "weiterhin freigegeben", "OK"],
          ["Test-Konzept V&V (Sample-Plaene, Test-Cases)", "abgeschlossen", "Final 22.5.2023", "OK"],
          ["Reichweite-OTA Risk-Analyse (Bricking-Risiko)", "abgeschlossen", "Risiko-Score 4 (gering)", "OK"]]),
        ("Risiken / Open Issues",
         ("list", [
            "Risiko 1 (mittel): Verzoegerung Qualcomm SoC-Patch fuer FOTA-Bootloader (BL-3.2.1); "
            "Verfuegbarkeit aktuell 30.6.2023; mitigation: paralleles Test mit BL-3.1.4 (eingeschraenkte "
            "Funktion).",
            "Risiko 2 (niedrig): Mercedes-CSR-Update erwartet Q3 2023 fuer MMA-Plattform; ggf. "
            "Re-Design-Iteration noetig.",
            "Open Issue 1: Endgueltige Vereinbarung mit BMW ueber OTA-Update-Frequenz (vorgeschlagen "
            "alle 90 Tage; BMW Wunsch alle 60 Tage) – Verhandlung Q3 2023.",
            "Open Issue 2: Cybersecurity-Audit durch BMW-eigenes Red-Team noch nicht terminiert "
            "(erwartet Q3 2023).",
         ])),
        ("Gate-Entscheidung",
         "Auf Basis der vorliegenden Liefergegenstaende und der ueberwiegend gruenen Status-Ampel "
         "wird das Gate G3 mit Auflagen freigegeben (»Conditional Pass«). Auflagen: (1) "
         "Verfuegbarkeit Qualcomm BL-3.2.1 bis 30.6.2023 oder Eskalation Plan B; (2) Klaerung "
         "OTA-Update-Frequenz mit BMW bis 31.7.2023; (3) Cybersecurity-Audit BMW Red-Team bis "
         "30.9.2023 abschliessen. Naechstes Gate G4 (Verifikation abgeschlossen) geplant: 14.11.2023."),
        ("Teilnehmer Gate-Sitzung",
         "Dr. Patrick Sandner (Projektleiter Werk REG), Dr. Klaus Kessler (Werkleiter RSG, "
         "Software-Lead), Sabine Brand (Q-Leitung REG), Andreas Maier (Geschaeftsfuehrer REG), "
         "Dr. Thomas Weber (COO REA, Sponsor), Markus Henkel (Produktionsleiter REG), "
         "Dr.-Ing. Kerstin Roggner (QS-Engineering REG), Cybersecurity-Vertreter Hr. Patrick Voss "
         "(extern, TUEV Sued). Protokoll: Dr. Sandner."),
        ("Unterschrift",
         signatures("Dr.-Ing. Patrick Sandner", "Projektleiter PRJ-2023-003", "Brennhagen Elektronik GmbH",
                    "Dr. Thomas Weber", "COO / Gate-Sponsor", "Brennhagen Elektronik AG",
                    place="Heilbronn", date_str_="24. Mai 2023")),
    ])


# ============================================================================
# 12) RCN IC-Rechnung 2022_01 (wrong-folder, RCN-issued)
# ============================================================================

write_doc(f"{BASE}/RCN_IC_Rechnung_2022_01.docx", RCN_H,
    "Intercompany Invoice January 2022 – RCN Shanghai to REG Heilbronn",
    subtitle="Invoice No. RCN-IC-2022-01; Asien-Aftermarket-Komponenten Re-Routing",
    sections=[
        ("Hinweis Ablage",
         "Diese Rechnung ist in der Ablage 03_Tochter_DE_REG_Heilbronn versehentlich abgelegt (»wrong "
         "folder«); originaere Zustaendigkeit liegt bei der RCN-Buchhaltung Shanghai sowie der REG-"
         "Buchhaltung Heilbronn (Eingangsrechnung Konto 33xxx). Korrekturhinweis bereits an "
         "Konzern-Document-Management gemeldet (Ticket DMS-2022-2148)."),
        ("Invoice Header (English / German)",
         f"From / Rechnungssteller:\n"
         f"  Brennhagen (Shanghai) Co. Ltd.\n"
         f"  Building 7, 888 Hongqiao Road, Changning District, 200336 Shanghai, P.R. China\n"
         f"  Unified Social Credit Code (USCI): 91310115MA1FL42Q38\n"
         f"  Local Tax Bureau: Shanghai Changning District; Country Manager: Zhang Hao\n\n"
         f"To / Rechnungsempfaengerin:\n"
         f"  Brennhagen Elektronik GmbH (REG)\n"
         f"  Industriestrasse 47, 74076 Heilbronn, Germany\n"
         f"  HRB 221456, AG Heilbronn; VAT-ID DE 287 612 449\n"
         f"  Buchungsstelle: Stefan Richter (Finanzcontroller)"),
        ("Rechnungsdaten",
         [["Field", "Value"],
          ["Invoice No.", "RCN-IC-2022-01"],
          ["Invoice Date", "10. February 2022"],
          ["Service Period", "January 1-31, 2022"],
          ["Currency", "EUR (Settlement; Booking in CNY equivalent)"],
          ["Payment Terms", "30 days net via Konzern-Cash-Pool (Sweep)"],
          ["Transfer Pricing Method", "Cost+ 7 % on aftermarket components"],
          ["Incoterms", "DAP Heilbronn (Incoterms 2020)"],
          ["Vertragsgrundlage", "IC-Liefervereinbarung RCN-REG vom 1.4.2021"]]),
        ("Leistungsaufstellung",
         [["Pos.", "Beschreibung", "Betrag EUR"],
          ["1", "Aftermarket-Komponenten ICP-3 (Re-Routing China -> Heilbronn fuer Europa-Aftermarket)", "2.840.500"],
          ["2", "ADAS-V4D Sensor-Stueck-Komponenten Bestellung Werk Wuhan", "640.000"],
          ["3", "Logistik-Vorbereitung Containertransport Shanghai -> Hamburg (3x40ft HC)", "180.400"],
          ["4", "Zoll-/Verzollungs-Service (extern via DHL Industrial)", "62.300"],
          ["5", "Engineering-Quality-Inspection RCN Shanghai (Cost+ 8 %)", "98.800"],
          ["6", "Sonstige Sammelposten Material", "321.200"],
          ["Summe Netto", "", "4.143.281"],
          ["Innergemeinschaftliche Lieferung (steuerbefreit Export CN -> DE)", "", "0 EUR (umsatzsteuerfrei)"],
          ["Total Invoice", "", "4.143.281 EUR (Anzeige Gesamt brutto 4.361.349 EUR bei interner Verzollungs-Brutto-Buchung)"]]),
        ("Hinweise / Settlement",
         "Cross-Border-Settlement zwischen RCN (CNY) und REG (EUR) erfolgt ueber den Konzern-Cash-"
         "Pool unter Federfuehrung der Group Treasury (Markus Pflanzer, REA Stuttgart). FX-Hedging "
         "EUR/CNY erfolgt durch die Konzern-Treasury ueber die Hausbanken BNP Paribas und Deutsche "
         "Bank (siehe FX-Hedge-Berichte REA_FX_Hedge_EUR_CNY_2022_Q1). Die Buchung erfolgt zum "
         "ECB-Referenzkurs zum Rechnungsdatum (10.2.2022: EUR/CNY 7.1240). Verrechnungspreis-"
         "Methode dokumentiert im RCN-Local-File 2022 (geprueft durch KPMG China Shanghai)."),
        ("Bestaetigungen",
         "Sachliche Richtigkeit: Liang Wei (Finance Manager RCN). Genehmigung: Zhang Hao (Country "
         "Manager RCN). Empfangsbestaetigung: Stefan Richter (Finanzcontroller REG). Tax-Review "
         "(Konzern): Dr. Heike Berger (Group Tax Director)."),
        ("Signature",
         signatures("Liang Wei", "Finance Manager", "Brennhagen (Shanghai) Co. Ltd.",
                    "Zhang Hao", "Country Manager", "Brennhagen (Shanghai) Co. Ltd.",
                    place="Shanghai", date_str_="February 10, 2022")),
    ])


print("Script complete – 92 docs scheduled. Running verification.")


# ============================================================================
# VERIFICATION
# ============================================================================

from docx import Document
from pathlib import Path
base_path = Path(BASE)
total = thin = 0
thin_files = []
for p in base_path.rglob("*.docx"):
    d = Document(p); t = " ".join(par.text for par in d.paragraphs)
    for tbl in d.tables:
        for r in tbl.rows:
            for c in r.cells: t += " " + c.text
    w = len(t.split()); total += 1
    if w < 200:
        thin += 1
        thin_files.append((p.name, w))

print(f"{total} total .docx, {thin} still thin (<200w)")
if thin_files:
    print("Thin files:")
    for n, w in thin_files:
        print(f"  {w:4d}w  {n}")
