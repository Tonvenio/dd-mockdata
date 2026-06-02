"""
fix_real_persons.py — surgical replacement of real-person / real-entity references
identified by the PII scan with fictional equivalents. Idempotent.

Applies across docx + xlsx in all three datasets (real-person hits were roehrig-only,
but we scan all to be safe). Order matters: most specific patterns first.
"""
# --- portable-paths-prelude --- (do not edit) ---
import sys
from pathlib import Path as _PathlibPath
_RP = _PathlibPath(__file__).resolve().parent
while not (_RP / "enhance_lib.py").exists() and _RP.parent != _RP:
    _RP = _RP.parent
_ROOT = _RP
sys.path.insert(0, str(_ROOT))
# --- end prelude ---
import re
from pathlib import Path
from docx import Document
import openpyxl

BASE = Path(f"{_ROOT}")
ROOTS = ["mueller_small", "biotech_medium", "roehrig_large"]

# (search, replace) — applied in order to each text block.
# Outputs never contain a later search term, so sequential application is safe.
REPL = [
    # --- real famous automotive figures (HIGH) ---
    ("Dr. Carsten Breitfeld", "Dr. Carsten Brendel"),
    ("Carsten Breitfeld", "Carsten Brendel"),
    ("Breitfeld", "Brendel"),
    ("Prof. Dr.-Ing. Lutz Eckstein", "Prof. Dr.-Ing. Ludwig Eckhardt"),
    ("Prof. Dr. Lutz Eckstein", "Prof. Dr. Ludwig Eckhardt"),
    ("Lutz Eckstein", "Ludwig Eckhardt"),
    ("Eckstein", "Eckhardt"),
    ("Prof. Dr. Markus Lienkamp", "Prof. Dr. Martin Lindenmann"),
    ("Markus Lienkamp", "Martin Lindenmann"),
    ("Lienkamp", "Lindenmann"),
    ("Sebastian Schreiber", "Sebastian Strohmeier"),   # SySS founder (Sabine Schreiber in HR untouched)
    ("Frank Weber", "Frank Wendler"),                  # BMW exec (Dr. Thomas Weber COO untouched)
    ("Andreas Reschka", "Andreas Rehbein"),
    ("Reschka", "Rehbein"),
    # --- fictional-but-real-sounding OEM purchasing contacts (rename to be safe) ---
    ("Dr. Markus Heinz", "Dr. Markus Heller"),
    ("Markus Heinz", "Markus Heller"),
    ("Dr. Joachim Lessing", "Dr. Joachim Lenz"),
    ("Joachim Lessing", "Joachim Lenz"),
    ("Dr. Henning Brueggemann", "Dr. Henning Brungs"),
    ("Henning Brueggemann", "Henning Brungs"),
    ("Brueggemann", "Brungs"),
    ("Brüggemann", "Brungs"),
    ("Eva Brenner", "Eva Bredow"),
    # --- external auditors (verify-real → rename) ---
    ("Dr. Christian Welt", "Dr. Christian Walter"),
    ("Christian Welt", "Christian Walter"),
    ("Marlene Schönherr", "Marlene Schubert"),
    ("Marlene Schoenherr", "Marlene Schubert"),
    ("Schönherr", "Schubert"),
    ("Schoenherr", "Schubert"),
    # --- academic collision (WU Wien) ---
    ("Prof. Renate Meyer", "Prof. Renate Mahler"),
    ("Renate Meyer", "Renate Mahler"),
    # --- external legal counsel at a real firm (rename named partners) ---
    ("Hans-Jürgen Schramm", "Hans-Jürgen Stollberg"),
    ("Hans-Juergen Schramm", "Hans-Juergen Stollberg"),
    ("Dr. Wolfgang Sturm", "Dr. Wolfgang Stahl"),
    ("Wolfgang Sturm", "Wolfgang Stahl"),
    ("Sturm & Partner", "Stahl & Kollegen"),
    # --- real entities ---
    ("KUGLER MAAG CIE", "ProcessExcellence Consulting"),
    ("KUGLER MAAG", "ProcessExcellence Consulting"),
    ("SySS GmbH", "SecuVantage GmbH"),
    ("SySS", "SecuVantage"),
    ("Wilhelmstrasse 14", "Europaplatz 3"),
    ("Wilhelmstraße 14", "Europaplatz 3"),
    # --- Allianz HQ address used for fictional landlord ---
    ("ABG Allianz Immobilien GmbH", "ABG Bavaria Immobilien GmbH"),
    ("Koeniginstrasse 28, 80802 Muenchen", "Leopoldstrasse 244, 80807 Muenchen"),
    ("Königinstraße 28, 80802 München", "Leopoldstraße 244, 80807 München"),
    ("Koeniginstrasse 28", "Leopoldstrasse 244"),
    ("Königinstraße 28", "Leopoldstraße 244"),
    # --- internal ISIN inconsistency ---
    ("DE000REA1234", "DE000RHGRP12"),
]

counts = {old: 0 for old, _ in REPL}
files_changed = set()


def apply_repl(text):
    changed = False
    for old, new in REPL:
        if old in text:
            counts[old] += text.count(old)
            text = text.replace(old, new)
            changed = True
    return text, changed


def process_paragraph(par):
    """Replace within a paragraph; collapse runs only if a target is present."""
    full = par.text
    new, changed = apply_repl(full)
    if changed and new != full:
        # collapse: put all new text in first run, clear the rest (minor format loss,
        # acceptable for body prose where these names live)
        if par.runs:
            par.runs[0].text = new
            for r in par.runs[1:]:
                r.text = ""
        else:
            par.add_run(new)
        return True
    return False


def process_docx(path):
    try:
        d = Document(path)
    except Exception:
        return False
    changed = False
    for par in d.paragraphs:
        if process_paragraph(par):
            changed = True
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    if process_paragraph(par):
                        changed = True
    if changed:
        d.save(path)
    return changed


def process_xlsx(path):
    try:
        wb = openpyxl.load_workbook(path)
    except Exception:
        return False
    changed = False
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    new, ch = apply_repl(cell.value)
                    if ch and new != cell.value:
                        cell.value = new
                        changed = True
    if changed:
        wb.save(path)
    return changed


def main():
    for root in ROOTS:
        rp = BASE / root
        for p in rp.rglob("*.docx"):
            if process_docx(p):
                files_changed.add(str(p.relative_to(BASE)))
        for p in rp.rglob("*.xlsx"):
            if process_xlsx(p):
                files_changed.add(str(p.relative_to(BASE)))

    print("Replacements applied (occurrences):")
    for old, new in REPL:
        if counts[old]:
            print(f"  {counts[old]:>4}x  {old!r} -> {new!r}")
    print(f"\nFiles changed: {len(files_changed)}")


if __name__ == "__main__":
    main()
