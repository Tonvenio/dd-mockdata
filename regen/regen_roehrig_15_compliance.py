"""Brennhagen AG – 15_Compliance_Recht regeneration.

Covers all 226 thin .docx files in roehrig_large/15_Compliance_Recht:
- Compliance-Richtlinien (Anti-Korruption, Kartellrecht, DSGVO, LkSG, Sanktionen, CMS)
- Legal Compliance Reports per Subsidiary (REG/RSG/RPL/RCZ/RHU/RCN, 2021-2023)
- Risk Assessments per Subsidiary x Topic (Datenschutz/Korruption/Kartell/Export)
- Interne-Revisions-Berichte (Einkauf, F&E, Finance, HR, IT, Legal, Produktion, Supply Chain, Vertrieb)
- DSGVO Datenpannen-Meldungen
- TPDD Third-Party Due Diligence
- Rechtsakten (Schriftsätze, Urteile, Vergleiche, Widersprüche, Mahnbescheide, Behördenschreiben)
- Subsidiary-Compliance-Routine (Arbeitssicherheit, Handelskammer, Jahresabschlusshinterlegung,
  Lohnsteuer, Sozialversicherung, Umweltbericht)
- Rechtsgutachten, Sanktionsscreening, DSGVO-Verzeichnis, Statusberichte, ECO/STE/IC/Gehaltsänderung
"""
# --- portable-paths-prelude --- (do not edit) ---
import sys
from pathlib import Path as _PathlibPath
_RP = _PathlibPath(__file__).resolve().parent
while not (_RP / "enhance_lib.py").exists() and _RP.parent != _RP:
    _RP = _RP.parent
_ROOT = _RP
sys.path.insert(0, str(_ROOT))
# --- end prelude ---
from enhance_lib import ROEHRIG_AG as R, ROEHRIG_SUBS as S, write_doc, signatures

BASE = f"{_ROOT}/roehrig_large/15_Compliance_Recht"
H = {"name": R["name"], "addr": R["addr"], "hrb": R["hrb"]}

# Standard Schlüsselpersonen aus dem Brief
CO = "Dr. Constanze Behrens"               # Compliance Officer (extern, Hengeler Mueller)
CSO = "Stefan Koehler"                     # CSO / LkSG-Beauftragter
OMB = "Dr. Wolfgang Sturm"                 # Externer Ombudsmann
GWB = "Markus Pflanzer"                    # Geldwaeschebeauftragter
BAFA = "Dr. Heike Berger"                  # BAFA-Beauftragte
DSB = "Dr. Markus Lehmann"                 # Datenschutzbeauftragter
CAE = "Andreas Buehler"                    # Chief Audit Executive
HM_PARTNER = "Dr. Hans-Juergen Schramm"    # Hengeler Mueller Lead Compliance Partner
HM_FIN = "Dr. Sebastian Wallisch"          # Hengeler Mueller Bank/Finanz


# =============================================================================
# 1) COMPLIANCE-RICHTLINIEN (KONZERN)
# =============================================================================
def comp_richtlinie(fname, titel, geltung, schutzgueter, klauseln, ueberwachung, sanktionen, version, gueltig_ab):
    sections = [
        ("Praeambel",
         f"Die Brennhagen Elektronik AG (»Gesellschaft«) verpflichtet sich zur Einhaltung saemtlicher "
         f"gesetzlichen Bestimmungen sowie zu hoechsten ethischen Standards. Diese Richtlinie ({titel}) "
         f"konkretisiert die im Code of Conduct verankerten Grundsaetze fuer alle Mitarbeitenden, "
         f"Fuehrungskraefte sowie Organmitglieder der Brennhagen-Gruppe weltweit."),
        ("Geltungsbereich", geltung),
        ("Geschuetzte Rechtsgueter und Grundlagen", ("list", schutzgueter)),
        ("Verbindliche Regelungen", ("clauses", klauseln)),
        ("Ueberwachung, Schulung und Reporting", ueberwachung),
        ("Verstoesse, Sanktionen und Hinweisgebersystem",
         f"Verstoesse gegen diese Richtlinie werden konsequent verfolgt. {sanktionen} "
         f"Hinweise koennen jederzeit ueber das konzernweite Hinweisgebersystem »SPEAK-UP@Brennhagen« "
         f"(betrieben durch EQS Group AG, anonym moeglich) oder ueber den externen Ombudsmann "
         f"{OMB} (Sturm & Partner, Hamburg) gemeldet werden. Schutz vor Repressalien gemaess "
         f"Hinweisgeberschutzgesetz (HinSchG) ist gewaehrleistet."),
        ("Inkrafttreten / Versionierung",
         f"Diese Richtlinie tritt am {gueltig_ab} in Kraft (Version {version}) und ersetzt alle frueheren "
         f"Fassungen. Die Pflege obliegt der Compliance-Funktion ({CO}, extern, Hengeler Mueller). "
         f"Die naechste turnusmaessige Ueberpruefung erfolgt spaetestens 24 Monate nach Inkrafttreten."),
        ("Unterschriften",
         signatures("Anna Mueller", "Vorstandsvorsitzende (CEO)", R["name"],
                    CO, "Compliance Officer (extern)", "Hengeler Mueller",
                    place="Stuttgart", date_str_=gueltig_ab)),
    ]
    write_doc(f"{BASE}/{fname}", H, titel, subtitle=f"Konzernrichtlinie – Version {version}", sections=sections)


comp_richtlinie(
    "REA_Anti_Korruption_Richtlinie_2023.docx",
    "Anti-Korruptions-Richtlinie der Brennhagen-Gruppe",
    "Diese Richtlinie gilt fuer alle Gesellschaften der Brennhagen-Gruppe (REA, RHO, REG, RSG, RPL, RCZ, RHU, RCN) "
    "sowie fuer alle Mitarbeitenden, Leiharbeitnehmer, Berater, Vertriebspartner und Geschaeftsfuehrer. "
    "Sie ist Bestandteil der Arbeits- und Dienstvertraege.",
    [
        "UK Bribery Act 2010 (Section 1, 2, 6, 7 – inkl. Haftung der Gesellschaft bei Versaeumnis adaequater Praevention)",
        "US Foreign Corrupt Practices Act (FCPA) – sowohl Anti-Bribery Provisions als auch Books-and-Records",
        "§§ 299, 300, 331-334 StGB (Bestechlichkeit/Bestechung im geschaeftlichen Verkehr; Amtstraegerbestechung)",
        "OECD-Anti-Bribery-Convention; UN Convention against Corruption (UNCAC)",
        "GwG (Geldwaeschegesetz) sowie KYC-Verfahren der Brennhagen-Gruppe",
        "ISO 37001:2016 (Anti-Bribery Management System) – Orientierungsrahmen",
    ],
    [
        ("§ 1 Null-Toleranz-Politik", [
            "Die Brennhagen-Gruppe duldet keinerlei Form von Korruption, Bestechung oder unzulaessiger Einfluss"
            "nahme. Dies gilt sowohl im Verhaeltnis zu Amtstraegern (Public Sector) als auch im B2B-Verkehr.",
            "Untersagt sind insbesondere: das Anbieten, Versprechen, Gewaehren, Fordern, Annehmen oder "
            "Erbitten unangemessener Vorteile, gleich ob direkt oder mittelbar ueber Dritte (Vertriebs"
            "agenten, Berater, Joint-Venture-Partner).",
        ]),
        ("§ 2 Geschenke, Bewirtungen, Einladungen", [
            "Geschenke an oder von Geschaeftspartnern sind nur zulaessig, wenn sie geringwertig (Wert <= 50 EUR "
            "pro Anlass), anlassbezogen, transparent und sozialadaequat sind. Geldgeschenke jeder Hoehe sind verboten.",
            "Geschenke an oder von Amtstraegern (inkl. Mitarbeiter staatlicher Unternehmen, OEM-Mitarbeitende "
            "mit Beschaffungs-/Pruefverantwortung im Sinne § 299 StGB) sind grundsaetzlich untersagt; Ausnahmen "
            "nur nach schriftlicher Vorab-Freigabe der Compliance-Funktion.",
            "Bewirtungen sind angemessen (Richtwert <= 100 EUR p. P.) und nur im Rahmen legitimer "
            "Geschaeftsanbahnung/-pflege zulaessig. Dokumentationspflicht in SAP-CO sowie Eintragung im "
            "Geschenke-/Bewirtungsregister (Tool: »Gifts & Hospitality Tracker«).",
        ]),
        ("§ 3 Facilitation Payments", [
            "Sogenannte »Facilitation Payments« (Beschleunigungszahlungen an Amtstraeger) sind weltweit untersagt – "
            "auch dort, wo lokal toleriert. UK Bribery Act und Brennhagen-Standard gehen ueber FCPA-Ausnahme hinaus.",
            "Im Notfall (Erpressungssituation, Gefahr fuer Leib/Leben) ist die Zahlung zulaessig, jedoch "
            "binnen 24 Stunden an Compliance Officer und CFO zu melden.",
        ]),
        ("§ 4 Spenden, Sponsoring, politische Zuwendungen", [
            "Spenden ausschliesslich an gemeinnuetzige Einrichtungen, gegen Zuwendungsbestaetigung, "
            "ueber Spenden-Bewilligungsprozess (Vier-Augen-Prinzip; ab 5.000 EUR Vorstandsfreigabe).",
            "Sponsoring nur mit schriftlichem Sponsoring-Vertrag, marktgerechte Gegenleistung; "
            "kein Sponsoring von Veranstaltungen mit ueberwiegend politischem Charakter.",
            "Parteispenden und Zuwendungen an Mandatstraeger sind konzernweit untersagt.",
        ]),
        ("§ 5 Vertriebsagenten, Berater, Joint Ventures", [
            "Engagement von Vertriebsagenten/Beratern ausschliesslich nach erfolgreicher Third-Party-Due-Diligence "
            "(TPDD-Verfahren der Compliance-Funktion). Vertraege enthalten verbindlich Anti-Korruptions-Klausel, "
            "Audit-Right und Termination-on-Breach.",
            "Provisionen marktgerecht und Erfolgs-/Leistungsnachweis-gebunden; Cash-Zahlungen sowie Zahlungen "
            "in Drittstaaten/Drittkonten unzulaessig.",
        ]),
    ],
    f"Das Compliance Management System (CMS) der Brennhagen-Gruppe wird nach IDW PS 980 ausgestaltet und "
    f"jaehrlich durch KPMG AG WPG (Lead Partner Dr. Maximilian Brand) auf Angemessenheit, Implementierung und "
    f"Wirksamkeit geprueft. Pflicht-E-Learnings (»ABC – Anti-Bribery & Corruption«) jaehrlich fuer alle "
    f"Mitarbeitenden in Vertrieb, Einkauf, Finance, Management; turnusmaessige Risk Assessments; "
    f"jaehrliche Berichterstattung an den Pruefungsausschuss des Aufsichtsrats (Vorsitz Prof. Dr.-Ing. Voss).",
    "Verstoesse koennen arbeits-, zivil- und strafrechtliche Folgen haben, einschliesslich fristloser Kuendigung "
    "(§ 626 BGB), Schadensersatz und Strafanzeige.",
    "3.0", "1. Maerz 2023",
)

comp_richtlinie(
    "REA_Kartellrecht_Compliance_Handbuch_2023.docx",
    "Kartellrecht-Compliance-Handbuch der Brennhagen-Gruppe",
    "Geltung fuer alle Mitarbeitenden mit Kontakt zu Wettbewerbern, Lieferanten, Kunden, OEMs oder "
    "Branchenverbaenden – insbesondere Vertrieb, Einkauf, Strategie, Vorstand, Management, "
    "Aussendienst sowie alle Tochtergesellschaften der Brennhagen-Gruppe weltweit.",
    [
        "Art. 101, 102 AEUV (Kartellverbot, Marktmachtmissbrauch)",
        "§§ 1, 19, 20 GWB (Gesetz gegen Wettbewerbsbeschraenkungen)",
        "EU-VerticalBER 2022/720 (Vertikalvereinbarungen, Preisbindungsverbot)",
        "EU-Horizontalleitlinien 2023 (Informationsaustausch, Forschungs-/Spezialisierungs-Kooperationen)",
        "US Sherman Act / Clayton Act – soweit Verkaeufe in die USA",
        "Branchen-typische Risiken im Automobil-Zulieferer-Geschaeft (Tier-1-Auftragsvergaben, OEM-Konsortien)",
    ],
    [
        ("§ 1 Hard-Core-Kartellverbote", [
            "Untersagt sind insbesondere Preisabsprachen, Marktaufteilung, Quotenabsprachen, Submissions"
            "absprachen (Bid Rigging) und der Austausch wettbewerbssensitiver Informationen (Preise, "
            "Mengen, Kosten, Strategien, Kapazitaeten, Margen) mit Wettbewerbern – auch indirekt ueber Verbaende oder "
            "Hub-and-Spoke-Konstellationen ueber gemeinsame Lieferanten/Kunden.",
            "Selbst informelle Gespraeche (Stammtisch, Messe, Fahrgemeinschaft) sind risikobehaftet. "
            "Bei kartellrechtssensitiven Themen ist sofort der Raum zu verlassen und der Vorgang an "
            "Compliance zu melden (»Leave & Document«-Regel).",
        ]),
        ("§ 2 Branchenverbaende, Working Groups", [
            "Teilnahme an Branchen-Working-Groups (VDA, ZVEI, CLEPA) nur mit Tagesordnung, Anwesen"
            "heitsliste und Protokoll; vorab Anti-Trust-Briefing der Compliance-Funktion. Bei Themen "
            "ausserhalb der Tagesordnung ist die Sitzung zu verlassen.",
            "Benchmarking-Studien nur ueber unabhaengige Treuhaender/Beratungen (z. B. Roland Berger), "
            "Daten anonymisiert und mit ausreichendem zeitlichen Abstand (>=12 Monate alt).",
        ]),
        ("§ 3 Vertikale Beziehungen", [
            "Keine Preisbindung der zweiten Hand gegenueber Distributoren/Aftermarket; unverbindliche "
            "Preisempfehlungen (UVP) sind zulaessig, sofern keine Druckmittel.",
            "Gebietsbeschraenkungen und Kundenbeschraenkungen nur im Rahmen der VerticalBER 2022/720 "
            "(Marktanteil <= 30 %, keine Kernbeschraenkungen).",
        ]),
        ("§ 4 Marktbeherrschende Stellung / Mono-Source", [
            "In ICP-3 / BMS-12 / ADAS-V4D-Maerkten Pruefung jeweils projektbezogen, ob marktbeherrschende "
            "Stellung gemaess § 18 GWB. Falls ja, Diskriminierungs- und Behinderungsverbot beachten "
            "(insb. Preisspreizungen, Lieferverweigerung gegenueber bestehenden OEM-Kunden).",
            "Vertraege mit Single-Source-Klauseln/Exclusivity bedurfen Compliance-Freigabe.",
        ]),
        ("§ 5 Dawn Raid / Durchsuchung", [
            "Bei unangekuendigter Durchsuchung durch Kartellbehoerden (BKartA, EU-KOM, BMAVL) gilt die "
            "Dawn-Raid-SOP der Brennhagen-Gruppe: Compliance- und Notfall-Hotline Hengeler Mueller "
            "(+49 69 17095-100, Partner " + HM_PARTNER + ") sofort verstaendigen; Durchsuchungsbeschluss "
            "kopieren, Anwesenheit von Anwaelten abwarten (Recht auf Beistand), keine Datentraeger "
            "vernichten, keine Aussagen ohne Anwalt.",
            "Dawn-Raid-Trainings jaehrlich fuer Vorstand, GL der Tochtergesellschaften, Werkleiter, "
            "Vertriebs- und Einkaufsleitung.",
        ]),
    ],
    f"Verpflichtende Schulungen: »Antitrust Compliance Basics« (alle), »Advanced Antitrust« (Vertrieb/Einkauf), "
    f"jeweils im LMS »Brennhagen Learn« durch SuccessFactors. Halbjaehrliche Compliance-Bulletins. "
    f"Kartellrechts-Hotline und »Speak-Up«-Hinweisgebersystem. Internes Audit-Programm durch Konzernrevision "
    f"unter Leitung von {CAE}; externe ABC-Audits jaehrlich durch Hengeler Mueller (Lead {HM_PARTNER}).",
    "Bei kartellrechtswidrigem Verhalten drohen empfindliche Bussgelder (bis 10 % des Konzern-Umsatzes), "
    "zivilrechtliche Schadensersatzansprueche (Follow-On) sowie Reputationsschaeden. Brennhagen nimmt am "
    "EU-Kronzeugen-Programm teil, sofern relevant.",
    "4.1", "1. April 2023",
)

# DSGVO-Datenschutzrichtlinie
def dsgvo_richtlinie(fname, titel, verfahren_oder_thema, version, datum):
    sections = [
        ("Praeambel",
         f"Die Brennhagen Elektronik AG ist als Verantwortlicher i. S. d. Art. 4 Nr. 7 DSGVO verpflichtet, "
         f"personenbezogene Daten ihrer Beschaeftigten, Bewerber, Geschaeftspartner und Kunden rechtmaessig, "
         f"transparent und zweckgebunden zu verarbeiten. Diese Richtlinie ({titel}) regelt {verfahren_oder_thema}."),
        ("Rechtsgrundlagen", ("list", [
            "Datenschutz-Grundverordnung (Verordnung (EU) 2016/679, DSGVO)",
            "Bundesdatenschutzgesetz (BDSG, neugefasst 2017/2019)",
            "Telekommunikation-Digitale-Dienste-Datenschutz-Gesetz (TDDDG 2024)",
            "§ 26 BDSG (Beschaeftigtendatenschutz)",
            "Bundesarbeitsgerichts-Rechtsprechung zu Mitarbeiter-Datenverarbeitung",
        ])),
        ("Datenschutzbeauftragter / Verantwortlichkeit",
         f"Datenschutzbeauftragter (DSB) gemaess Art. 37 DSGVO i. V. m. § 38 BDSG: {DSB} (extern, "
         f"Lehmann & Partner Datenschutzberatung, Stuttgart). Kontakt: datenschutz@brennhagen-elektronik.de. "
         f"Bestellung schriftlich, Meldung an den Landesbeauftragten fuer Datenschutz Baden-Wuerttemberg "
         f"erfolgt. Der DSB berichtet direkt an den Vorstand (CEO/CFO) und ist unabhaengig und weisungsfrei."),
        ("Grundsaetze der Verarbeitung", ("clauses", [
            ("§ 1 Rechtmaessigkeit & Zweckbindung", [
                "Verarbeitung nur auf Basis einer der Rechtsgrundlagen des Art. 6 Abs. 1 DSGVO. Bei "
                "besonderen Datenkategorien (Art. 9) zusaetzlich Art. 9 Abs. 2 erforderlich.",
                "Verarbeitungstaetigkeiten sind im Verzeichnis nach Art. 30 DSGVO erfasst (»VVT«), "
                "Owner: jeweilige Fachabteilung, fachliche Pruefung DSB, technisch IT-Compliance.",
            ]),
            ("§ 2 Datenminimierung & Speicherbegrenzung", [
                "Erhebung nur erforderlicher Daten; Loeschkonzept nach Loeschklasse (Personalakten 10 Jahre "
                "nach Ausscheiden bzw. nach gesetzl. Aufbewahrungsfristen § 257 HGB / § 147 AO).",
                "Pseudonymisierung/Anonymisierung wo moeglich (Art. 25 DSGVO – Privacy-by-Design).",
            ]),
            ("§ 3 Auftragsverarbeitung", [
                "AV-Vertraege nach Art. 28 DSGVO mit allen Dienstleistern (SAP, Salesforce, Microsoft 365, "
                "Workday, ADP, EQS Group, ComplyCube, etc.). Vorlage des Konzern-AV-Templates verbindlich.",
                "Bei Drittlandtransfer (z. B. USA) zusaetzlich SCC 2021 sowie Transfer Impact Assessment (TIA).",
            ]),
            ("§ 4 Betroffenenrechte", [
                "Anfragen nach Art. 15-22 DSGVO sind innerhalb von 1 Monat zu beantworten; "
                "Ticket-System: datenschutz@brennhagen-elektronik.de; Eskalation an DSB.",
                "Auskunfts-, Berichtigungs-, Loesch-, Einschraenkungs-, Datenuebertragungs-, Widerspruchsrecht "
                "sowie Recht auf Beschwerde bei Aufsichtsbehoerde.",
            ]),
            ("§ 5 Datenpannen / Meldung gemaess Art. 33/34 DSGVO", [
                "Meldepflicht binnen 72 Stunden an LfDI Baden-Wuerttemberg, sofern Risiko fuer Betroffene. "
                "Erstmeldung ueber internes DSGVO-Datenpannen-Tool, gefolgt von Untersuchung und Lessons Learned.",
                "Bei hohem Risiko zusaetzlich Information der Betroffenen gemaess Art. 34 DSGVO.",
            ]),
        ])),
        ("Schulungen & Kontrolle",
         "Verpflichtende Pflicht-E-Learnings »DSGVO Basics« jaehrlich (alle Mitarbeiter); »Advanced Privacy« "
         "fuer HR, IT, Marketing, Vertrieb. Audit-Programm der Konzernrevision (CAE " + CAE + "). "
         "Externe Audits gemaess IDW PS 980 / ISO 27701-Orientierung."),
        ("Inkrafttreten",
         f"Diese Richtlinie tritt am {datum} in Kraft, Version {version}. Naechste Ueberpruefung 2025."),
        ("Unterschriften",
         signatures("Laura Bauer", "CFO / Verantwortlicher im Datenschutz", R["name"],
                    DSB, "Datenschutzbeauftragter (extern)", "Lehmann & Partner",
                    place="Stuttgart", date_str_=datum)),
    ]
    write_doc(f"{BASE}/{fname}", H, titel, subtitle=f"Datenschutz – Version {version}", sections=sections)


dsgvo_richtlinie("REA_DSGVO_Datenschutzrichtlinie_2023.docx",
                 "DSGVO-Datenschutzrichtlinie der Brennhagen-Gruppe",
                 "die Grundsaetze, Verantwortlichkeiten und Prozesse zum Schutz personenbezogener Daten im gesamten Konzern",
                 "3.1", "15. Februar 2023")

dsgvo_richtlinie("REA_DSGVO_Datenschutzhinweise_Beschaeftigte_2023.docx",
                 "Datenschutzhinweise fuer Beschaeftigte gemaess Art. 13 DSGVO",
                 "die Information aller Mitarbeitenden der Brennhagen-Gruppe ueber die Verarbeitung ihrer "
                 "personenbezogenen Daten (HR, Lohnabrechnung, Workday, ADP, Zeiterfassung, IT-Logs) "
                 "vom Bewerbungsprozess bis nach Beendigung des Arbeitsverhaeltnisses",
                 "2.4", "1. Juli 2023")

dsgvo_richtlinie("REA_DSGVO_Verzeichnis_Verarbeitungstaetigkeiten_2023.docx",
                 "Verzeichnis der Verarbeitungstaetigkeiten (VVT) gemaess Art. 30 DSGVO",
                 "das konzernweite Verzeichnis der Verarbeitungstaetigkeiten mit Beschreibung, Zweck, "
                 "Rechtsgrundlage, Empfaengerkategorien, Drittlandtransfer und Loeschfristen "
                 "(SAP HCM, Workday, Salesforce, MS365, Teamcenter PLM, DMS, Telematik, Zutrittskontrolle)",
                 "5.0", "30. Juni 2023")


# DSGVO AV-Vertrag, DSFA, DSB-Bestellung
def dsgvo_doc(fname, titel, subtitle, vertragsgegenstand, datenarten, schutzmassnahmen, weiteres, datum):
    sections = [
        ("Vertragsparteien / Sachverhalt",
         f"Verantwortlicher: Brennhagen Elektronik AG, Vaihinger Strasse 120, 70567 Stuttgart, HRB 726451, "
         f"vertreten durch die CFO Laura Bauer.\n\n"
         f"Auftragsverarbeiter/Gegenstand: {vertragsgegenstand}"),
        ("Datenarten und Kategorien Betroffener", ("list", datenarten)),
        ("Technisch-organisatorische Massnahmen (TOMs)", ("list", schutzmassnahmen)),
        ("Sonstige Regelungen", weiteres),
        ("Datenschutzbeauftragter / Aufsicht",
         f"Datenschutzbeauftragter der Brennhagen Elektronik AG: {DSB} (extern, Lehmann & Partner). "
         f"Bei Datenpannen Meldung gem. Art. 33 DSGVO an LfDI Baden-Wuerttemberg innerhalb 72 Stunden."),
        ("Unterschriften",
         signatures("Laura Bauer", "CFO", R["name"], DSB, "DSB / Beauftragter", "Lehmann & Partner",
                    place="Stuttgart", date_str_=datum)),
    ]
    write_doc(f"{BASE}/{fname}", H, titel, subtitle=subtitle, sections=sections)


dsgvo_doc("REA_DSGVO_AV_Vertrag_Salesforce_2023_v1.docx",
    "Auftragsverarbeitungsvertrag (AVV) gemaess Art. 28 DSGVO – Salesforce.com EMEA",
    "Vertrag zwischen Brennhagen Elektronik AG und salesforce.com Germany GmbH",
    "salesforce.com Germany GmbH, Erika-Mann-Strasse 31, 80636 Muenchen, vertreten durch den "
    "Geschaeftsfuehrer (»Auftragsverarbeiter«). Gegenstand der Verarbeitung: Cloud-basierte CRM-Plattform "
    "(Salesforce Sales Cloud, Service Cloud, Marketing Cloud) zur Verwaltung von Geschaeftspartner-, "
    "Kunden- und Interessenten-Daten der Brennhagen-Gruppe.",
    [
        "Kunden- und Lead-Stammdaten (Name, Funktion, Geschaeftsadresse, geschaeftliche Kontaktdaten)",
        "Kommunikationsverlaeufe (E-Mails, Notizen, Meeting-Protokolle, Outlook-Synchronisation)",
        "Opportunities, Forecast- und Abschlussdaten (z. B. ICP-3-, BMS-12-Projekte)",
        "Service-Tickets, Reklamationen, Vertragsdokumente",
    ],
    [
        "Verschluesselung der Daten at rest (AES-256) und in transit (TLS 1.3)",
        "Multi-Faktor-Authentifizierung (MFA) verpflichtend; SSO ueber Azure AD",
        "Rollenbasierte Zugriffskontrolle (RBAC) ueber Salesforce Permission Sets",
        "Backups taeglich; Disaster Recovery Tier-1; RPO 4 h, RTO 24 h",
        "Auditierung gemaess ISO/IEC 27001:2022 und SOC 2 Typ II (zertifiziert)",
        "Datenspeicherung primaer EU (Frankfurt/Paris); Drittlandtransfer USA nur ueber EU-US DPF + SCC 2021 + TIA",
    ],
    "Subunternehmer-Liste gemaess Salesforce Trust-Portal; Genehmigung von Aenderungen durch Brennhagen erforderlich. "
    "Audit-Rechte des Verantwortlichen einmal jaehrlich gemaess Art. 28 Abs. 3 lit. h DSGVO. "
    "Loeschung bzw. Rueckgabe der Daten binnen 30 Tagen nach Vertragsende. Bei Datenpanne unverzuegliche Meldung "
    "an Brennhagen (binnen 24 Stunden), um konzernseitige 72-Stunden-Frist gegenueber LfDI zu wahren. "
    "Laufzeit kongruent zum Salesforce-Hauptvertrag (3 Jahre, Verlaengerung um 1 Jahr).",
    "10. Maerz 2023")

dsgvo_doc("REA_DSGVO_DSFA_CRM_System_2022.docx",
    "Datenschutz-Folgenabschaetzung (DSFA) gemaess Art. 35 DSGVO – Einfuehrung Salesforce CRM",
    "Bewertung der Risiken fuer Rechte und Freiheiten Betroffener",
    "Geplante Verarbeitungstaetigkeit: Einfuehrung eines konzernweiten CRM-Systems (Salesforce Sales/Service/"
    "Marketing Cloud) zur Konsolidierung der bisherigen heterogenen Insellosungen (Excel, MS Access, RCN-CRM). "
    "Verantwortlich: Marketing & Vertrieb (Stefan Richter, CMO); IT-Sponsor: Group IT (Dr. Robert Schaefer). "
    "Compliance-Pruefung: " + CO + " (Compliance Officer) und " + DSB + " (DSB).",
    [
        "Kontaktdaten (B2B) Ansprechpartner OEMs (BMW, VW, Mercedes, Stellantis, ...)",
        "Kommunikationsverlaeufe, Outlook-Synchronisation, Meeting-Protokolle",
        "Projektbezogene Daten (RFQs, Quotation Histories, Forecasts)",
        "Mitarbeiterdaten der Brennhagen-Vertriebsteams (Nutzer, Login-Logs, MFA)",
    ],
    [
        "Risikobewertung: mittleres Risiko (B2B-Kontaktdaten, keine sensiblen Kategorien Art. 9 DSGVO)",
        "Massnahmen: AVV (siehe Salesforce-AV); Rechte-/Rollenkonzept; Audit-Logs; EU-Datacenter primaer",
        "Restrisiko: akzeptabel (Tier-1-Verschluesselung, etablierter Hyperscaler, EU-Datenresidenz)",
        "Vorab-Konsultation der Aufsichtsbehoerde gemaess Art. 36 DSGVO nicht erforderlich",
    ],
    "Empfehlung des DSB " + DSB + ": Einfuehrung darf erfolgen, vorbehaltlich rollierender Re-Pruefung "
    "bei wesentlichen Aenderungen (z. B. Aktivierung der Salesforce Einstein-AI-Module).",
    "12. Oktober 2022")

dsgvo_doc("REA_DSGVO_Datenschutzbeauftragter_Bestellung_2022_intern.docx",
    f"Bestellung externer Datenschutzbeauftragter – {DSB}",
    "Bestellung gemaess Art. 37 DSGVO i. V. m. § 38 BDSG",
    f"Bestellt wird {DSB}, Geschaeftsfuehrer Lehmann & Partner Datenschutzberatung GmbH, "
    f"Koenigstrasse 80, 70173 Stuttgart, als externer Datenschutzbeauftragter (DSB) der "
    f"Brennhagen Elektronik AG einschliesslich der Konzerngesellschaften REG, RSG, RPL, RCZ, RHU und RCN "
    f"(Konzern-DSB nach Art. 37 Abs. 2 DSGVO).",
    [
        "Qualifikation: Volljurist, CIPP/E, CIPM (IAPP-zertifiziert), 18 Jahre Erfahrung Datenschutz",
        "Unabhaengigkeit: weisungsfrei, direkter Berichtsweg an CEO/CFO (Art. 38 Abs. 3 DSGVO)",
        "Aufgaben nach Art. 39 DSGVO (Unterrichtung, Beratung, Ueberwachung, Schulung, Anlaufstelle)",
    ],
    [
        "Buero und Sprechstunden (regelmaessige Anwesenheit Stuttgart, Q-monatlich Heilbronn/Muenchen)",
        "Berichterstattung quartalsweise an Vorstand; Jahresbericht an Aufsichtsrat Pruefungsausschuss",
        "Meldungen an LfDI Baden-Wuerttemberg gemaess § 38 BDSG erfolgt",
    ],
    "Verguetung: 96.000 EUR p. a. (Pauschal), zzgl. Reisekosten. Laufzeit unbestimmt, Kuendigungsfrist "
    "6 Monate zum Quartalsende. Abberufung nur gemaess Art. 38 Abs. 3 DSGVO (kein Nachteil aus Aufgabenerfuellung).",
    "1. April 2022")


# LkSG / Sorgfaltspflichtenbericht / Sanktionsscreening / CMS / Rechtsgutachten / STE / ECO / Statusberichte / IC
def lksg_risk(fname, titel, jahr, schwerpunkte, massnahmen, beschwerden, hinweise, datum):
    sections = [
        ("Einleitung",
         f"Die Brennhagen Elektronik AG ist seit dem 1. Januar 2024 als Unternehmen mit mehr als 1.000 Beschaeftigten "
         f"im Inland berichtspflichtig nach dem Gesetz ueber die unternehmerischen Sorgfaltspflichten in "
         f"Lieferketten (Lieferkettensorgfaltspflichtengesetz – LkSG). Diese Risiko- und Sorgfaltspflichten-"
         f"Erklaerung deckt das Geschaeftsjahr {jahr} ab und wird gemaess § 10 Abs. 2 LkSG bis spaetestens "
         f"vier Monate nach Ende des Berichtsjahres beim Bundesamt fuer Wirtschaft und Ausfuhrkontrolle (BAFA) "
         f"eingereicht. Verantwortlich: Chief Sustainability Officer (CSO) {CSO} (LkSG-Beauftragter); "
         f"Compliance-Funktion {CO} (Hengeler Mueller)."),
        ("Geltungsbereich",
         "Eigener Geschaeftsbereich (Konzern Brennhagen: REA, RHO, REG, RSG, RPL, RCZ, RHU, RCN) sowie "
         "unmittelbare Lieferanten (Tier-1). Mittelbare Lieferanten (Tier-2+) anlassbezogen bei substantiierter "
         "Kenntnis (»substantiated knowledge«) gemaess § 9 Abs. 3 LkSG."),
        ("Risikoanalyse – Schwerpunkte des Berichtsjahrs", ("list", schwerpunkte)),
        ("Praeventions- und Abhilfemassnahmen", ("clauses", massnahmen)),
        ("Beschwerdeverfahren / Hinweise",
         f"Hinweisgebersystem »SPEAK-UP@Brennhagen« (EQS Group AG) sowie externer Ombudsmann {OMB} (Sturm & "
         f"Partner Hamburg). Im Berichtsjahr {jahr} eingegangene Hinweise: {beschwerden}. "
         f"Auswertung und Statistik: {hinweise}"),
        ("Berichterstattung & Wirksamkeitspruefung",
         "Wirksamkeitspruefung der Massnahmen jaehrlich sowie anlassbezogen. Berichterstattung gegenueber "
         "Vorstand (quartalsweise), Aufsichtsrat (jaehrlich, Pruefungsausschuss) und BAFA (jaehrlich). "
         "Externe Pruefung der LkSG-Berichterstattung 2024 erstmalig durch KPMG AG WPG (begrenzte Pruefung) im "
         "Rahmen des Nachhaltigkeitsberichts (CSRD/ESRS-konform geplant ab 2025)."),
        ("Erklaerung des Vorstands",
         f"Der Vorstand der Brennhagen Elektronik AG bestaetigt die Richtigkeit und Vollstaendigkeit dieser "
         f"Erklaerung sowie die Einhaltung der menschenrechts- und umweltbezogenen Sorgfaltspflichten gemaess "
         f"§§ 3-10 LkSG nach bestem Wissen und Gewissen. Stuttgart, den {datum}."),
        ("Unterschriften",
         signatures("Anna Mueller", "CEO / Vorstandsvorsitzende", R["name"],
                    CSO, "Chief Sustainability Officer / LkSG-Beauftragter", R["name"],
                    place="Stuttgart", date_str_=datum)),
    ]
    write_doc(f"{BASE}/{fname}", H, titel, subtitle=f"Berichtsjahr {jahr} – gemaess § 10 LkSG", sections=sections)


lksg_risk("REA_LkSG_Sorgfaltspflichtenbericht_2023.docx",
    "LkSG-Sorgfaltspflichtenbericht 2023 der Brennhagen Elektronik AG",
    "2023",
    [
        "Eigener Geschaeftsbereich – Arbeitsschutz Werke REG/RPL/RCZ/RHU (geringes Risiko, IATF/ISO-zertifiziert)",
        "Tier-1-Lieferanten – Cobalt/Lithium/Coltan-haltige Komponenten (BMS-12, ADAS-V4D) – mittleres Risiko",
        "Tier-1-Lieferanten – Schaltschraenke und Gehaeuse aus Tuerkei (Kinderarbeit-Risiko mittel)",
        "Tier-1-Lieferanten – Leiterplatten-Hersteller China/Vietnam (Arbeitszeitenrisiko mittel)",
        "Logistik – Hafen-Subdienstleister in Bremerhaven/Antwerpen (Arbeitsschutz, Loehne)",
        "Mittelbare Lieferanten Tier-2: Mining/Smelter-Risiken Cobalt/Tantal (RMI/CMRT-Reporting)",
    ],
    [
        ("§ 1 Lieferanten-Onboarding und KYC", [
            "Verpflichtende Akzeptanz des Code-of-Conduct-fuer-Geschaeftspartner (CoC GP); Compliance-"
            "Klausel in allen Rahmenliefervertraegen seit 2022 ueberarbeitet (LkSG-konform).",
            "KYC-Pruefung mit ComplyCube fuer Tier-2-Lieferanten in CN/IN/TR (Sanktionslisten, PEP, "
            "negative Medien); Re-Screening jaehrlich.",
        ]),
        ("§ 2 Audits und Self-Assessments", [
            "EcoVadis-Self-Assessment fuer Top-200-Lieferanten (Coverage 87 % nach Beschaffungsvolumen).",
            "Vor-Ort-Audits durch SGS/Intertek bei 18 Lieferanten in 2023 (CN: 8, TR: 4, IN: 3, EE/PL: 3); "
            "5 Major-Findings, alle durch CAP geschlossen bis Q1/2024.",
        ]),
        ("§ 3 Abhilfemassnahmen", [
            "Bei 2 chinesischen PCB-Lieferanten Korrekturmassnahmen Arbeitszeit/Ueberstunden vereinbart; "
            "Monitoring 6 Monate. Kein Lieferantenabbruch erforderlich, da Verbesserung erkennbar.",
            "Bei 1 tuerkischem Gehaeuse-Lieferanten Beendigung der Geschaeftsbeziehung Q3/2023 nach "
            "wiederholten Findings (Arbeitsschutz, Mindestlohn).",
        ]),
        ("§ 4 Schulungen und Awareness", [
            "Pflicht-E-Learning »LkSG fuer Einkauf & Lieferantenmanagement« 2023 (Coverage 96 %).",
            "Top-Management-Briefing durch Hengeler Mueller (Partner " + HM_PARTNER + ") in Q2/2023.",
        ]),
    ],
    "3 Hinweise (alle ueber SPEAK-UP), davon 2 zu Arbeitszeiten Tier-1-CN-Lieferanten, 1 zu Lohnzahlung "
    "Tier-1-Tuerkei. Alle ordnungsgemaess untersucht und geschlossen.",
    "Bearbeitungszeit durchschnittlich 42 Tage; 100 % anonym moeglich; externer Ombudsmann genutzt in 1 Fall.",
    "30. April 2024")


# REA_Sanktionsscreening_Verfahren
write_doc(f"{BASE}/REA_Sanktionsscreening_Verfahren_2023.docx", H,
    "Verfahrensanweisung Sanktionsscreening und Exportkontrolle",
    subtitle="Konzernweites Verfahren – Version 2.3 – Stand 15. Januar 2023",
    sections=[
        ("Zielsetzung",
         "Diese Verfahrensanweisung regelt das konzernweite Screening von Geschaeftspartnern (Kunden, "
         "Lieferanten, Vertriebsagenten, Endkunden im Aftermarket) gegen Sanktionslisten der EU, der UN, "
         "OFAC (US) sowie HMT (UK) und den deutschen Embargo-Bestimmungen. Ziel ist die rechtssichere Einhaltung "
         "der ausserwirtschaftsrechtlichen Vorgaben (AWG, AWV, EU-Dual-Use-Verordnung 2021/821, US-EAR) "
         "und der einschlaegigen Russland-, Belarus- und Iran-Sanktionen."),
        ("Geltungsbereich", ("list", [
            "Alle Konzerngesellschaften (REA, RHO, REG, RSG, RPL, RCZ, RHU, RCN)",
            "Alle Geschaeftsvorfaelle Verkauf/Einkauf/Beratung mit Auslandsbezug",
            "Bestehende und Neukunden/Lieferanten – Initial- und periodisches Re-Screening",
            "Endkunden im Aftermarket (RCN Shanghai – End-User-Statement) gemaess Dual-Use-VO",
        ])),
        ("Tool und Datenquellen",
         "Verwendetes Tool: ComplyCube (Konzernlizenz, betrieben durch Group Compliance). Datenquellen: "
         "EU-Konsolidierte Liste (CFSP), UN Security Council Consolidated List, US OFAC SDN/SSI, UK HMT "
         "Consolidated List, BAFA-Embargo-Hinweise, Sanctions Map EU. Screening-Frequenz: bei Onboarding, "
         "vor jeder Rechnungs-/Lieferfreigabe sowie taeglich automatisch ueber Delta-Listen-Updates."),
        ("Prozess", ("clauses", [
            ("§ 1 Pre-Screening", [
                "Stammdatenpflege durch SAP-Master-Data; ComplyCube-Integration ueber API.",
                "Initial-Screening bei Anlage neuer Geschaeftspartner; Sperre des Datensatzes bei Hit, "
                "Eskalation an Compliance Officer.",
            ]),
            ("§ 2 Hit-Bewertung und Eskalation", [
                "Erstpruefung durch Compliance-Analyst (Group Compliance, Lehmann). Bei Bestaetigung "
                "Eskalation an " + CO + " (CO) sowie BAFA-Beauftragte " + BAFA + " (Group Tax).",
                "Bei militaerischer Endverwendung (Dual-Use-Komponenten ECU-900/ADAS-V4D) Genehmigungspflicht "
                "BAFA-Antrag vor Lieferung. Lead-Time mind. 8 Wochen.",
            ]),
            ("§ 3 Transaktions-Screening", [
                "Pruefung jeder Auftragsbestaetigung gegen Liste der gelisteten Endkunden und "
                "Verwendungszwecke (Catch-All-Klausel Art. 4 EU-Dual-Use-VO).",
                "Russland-Sanktionen: Auftraege seit 24.2.2022 zusaetzlicher Manual-Review (verbotene "
                "Gueterlisten gemaess EU-833/2014 idF Stand 19.12.2024).",
            ]),
            ("§ 4 Dokumentation und Audit", [
                "Lueckenlose Dokumentation aller Hits, Bewertungen und Freigaben im ComplyCube-Audit-Log "
                "fuer mind. 10 Jahre (Aufbewahrungsfristen § 22 AWG).",
                "Jaehrliches Audit der Sanktionsscreening-Prozesse durch Konzernrevision sowie externe "
                "Pruefung im Rahmen der IDW PS 980-Pruefung (KPMG, Lead Dr. Brand).",
            ]),
        ])),
        ("Verantwortlichkeiten", [
            ["Rolle", "Person", "Aufgabe"],
            ["Compliance Officer", CO, "Gesamtverantwortung Sanktionscompliance"],
            ["BAFA-Beauftragte", BAFA, "Antragsverfahren / Aussenwirtschaft"],
            ["Geldwaeschebeauftragter", GWB, "GwG-Verzahnung, KYC"],
            ["Compliance-Analyst", "M. Lehmann (Group Compliance)", "Operativer Tool-Betrieb ComplyCube"],
            ["Werks-Compliance-Beauftragte", "Werkleiter REG/RPL/RCZ/RHU/RCN", "Vor-Ort-Eskalation"],
        ]),
        ("Unterschriften",
         signatures("Anna Mueller", "CEO", R["name"], CO, "Compliance Officer (extern)", "Hengeler Mueller",
                    place="Stuttgart", date_str_="15. Januar 2023")),
    ])


# REA_Compliance_Management_System
write_doc(f"{BASE}/REA_Compliance_Management_System_2023.docx", H,
    "Compliance-Management-System (CMS) der Brennhagen Elektronik AG",
    subtitle="Beschreibung gemaess IDW PS 980 – Stand 1. April 2023",
    sections=[
        ("Praeambel",
         "Die Brennhagen Elektronik AG unterhaelt ein konzernweites Compliance-Management-System (CMS), "
         "das nach dem Pruefstandard IDW PS 980 strukturiert ist und alle sieben Grundelemente abdeckt "
         "(Compliance-Kultur, -Ziele, -Risiken, -Programm, -Organisation, -Kommunikation, -Ueberwachung/Verbesserung). "
         "Das CMS wird jaehrlich durch KPMG AG WPG geprueft (Lead Partner Dr. Maximilian Brand)."),
        ("1. Compliance-Kultur",
         "Tone-from-the-top: Vorstand und Aufsichtsrat verpflichten sich oeffentlich zur Einhaltung von Recht und Ethik. "
         "Code of Conduct (Mitarbeitende und Geschaeftspartner) – verfuegbar in Deutsch, Englisch, Polnisch, "
         "Tschechisch, Ungarisch, Chinesisch. Jaehrliche CEO-Botschaft sowie verpflichtende Bestaetigung "
         "(»Compliance Pledge«) durch Top-200-Fuehrungskraefte."),
        ("2. Compliance-Ziele",
         "Strategische Ziele: 0 Korruptionsfaelle mit Schaden > 50.000 EUR; 100 % Schulungsabdeckung Top-Risikofunktionen; "
         "0 Sanktionsverstoesse; LkSG-konforme Lieferkette; DSGVO-Reifegrad-Score (extern) >= 4/5; "
         "kein Kartellverfahren mit Bussgeld. Operationalisierung im jaehrlichen Compliance-Plan."),
        ("3. Compliance-Risiken",
         "Risikoinventar jaehrlich aktualisiert (Brutto-Risikomatrix), gegliedert nach: Anti-Korruption, "
         "Kartellrecht, Sanktionen/Exportkontrolle, Geldwaesche, DSGVO, LkSG, Insiderrecht, Steuerrecht, "
         "Arbeits-/Datenschutzrecht. Risk Owner pro Risiko benannt; jaehrliche Bottom-up-Workshops in den "
         "Tochtergesellschaften (REG, RSG, RPL, RCZ, RHU, RCN). Heatmap und Top-10-Risiken werden "
         "Pruefungsausschuss berichtet."),
        ("4. Compliance-Programm",
         "Richtlinien-Hierarchie: Code of Conduct (Vorstand) – Konzernrichtlinien (Anti-Korruption, Kartell, "
         "Sanktionen, DSGVO, LkSG, GwG, MAR/Insider) – Verfahrensanweisungen (Werks-/Bereichsebene) – "
         "Arbeitsanweisungen. Kontrollen in SAP/GRC mit »Three-Lines-of-Defense«-Modell."),
        ("5. Compliance-Organisation", ("list", [
            f"Compliance Officer (CO) extern: {CO}, Hengeler Mueller – berichtet an CEO und Pruefungsausschuss",
            f"Konzernrevision (Internal Audit) CAE: {CAE} – berichtet funktional an Pruefungsausschuss",
            f"Geldwaeschebeauftragter: {GWB} (Group Treasurer, Doppelfunktion)",
            f"BAFA-Beauftragte / Aussenwirtschaft: {BAFA} (Group Tax, Doppelfunktion)",
            f"Datenschutzbeauftragter (DSB) extern: {DSB} (Lehmann & Partner)",
            f"LkSG-Beauftragter / CSO: {CSO}",
            "Compliance-Komitee (CEO, CFO, COO, CO, CAE, DSB) – tagt quartalsweise",
            "Werks-Compliance-Beauftragte (REG/RPL/RCZ/RHU/RCN) – dotted-line zu CO",
        ])),
        ("6. Compliance-Kommunikation",
         "Intranet »Compliance Hub«, Newsletter Q-monatlich, jaehrliche Pflicht-E-Learnings im LMS Brennhagen Learn "
         "(SAP SuccessFactors): »Code of Conduct«, »ABC – Anti-Bribery & Corruption«, »Antitrust Basics«, "
         "»Sanktionen & Exportkontrolle«, »DSGVO Basics«, »LkSG fuer Einkauf«, »MAR/Insider« (fuer Insider-Kreis). "
         "Speak-Up-System EQS, externer Ombudsmann " + OMB + "."),
        ("7. Compliance-Ueberwachung und -Verbesserung",
         "Konzernrevision pruefte 2023 18 Audits (Einkauf, F&E, Finance, HR, IT, Legal, Produktion REG/RPL, Supply Chain, "
         "Vertrieb). KPMG-Pruefung IDW PS 980 (»Wirksamkeit«) erstmalig fuer 2024 geplant; bis dahin "
         "Konzeptions- und Angemessenheitspruefungen. Kontinuierliche Verbesserung via Lessons-Learned aus "
         "Hinweisen, Audits, externen Vorfaellen."),
        ("Unterschriften",
         signatures("Anna Mueller", "CEO", R["name"], CO, "Compliance Officer (extern)", "Hengeler Mueller",
                    place="Stuttgart", date_str_="1. April 2023")),
    ])


# =============================================================================
# 2) RECHTSGUTACHTEN
# =============================================================================
def rechtsgutachten(fname, titel, mandant_frage, sachverhalt, analyse, ergebnis, autor, kanzlei, datum):
    sections = [
        ("Anlass und Auftrag",
         f"Mandant: Brennhagen Elektronik AG, vertreten durch den Vorstand. "
         f"Gutachterauftrag: {mandant_frage}\n\n"
         f"Beauftragte Kanzlei: {kanzlei}. Federfuehrender Partner: {autor}. "
         f"Vergutung: Stundensatz Partner 750 EUR, Senior Associate 480 EUR (gem. Mandatsvereinbarung)."),
        ("Sachverhalt", sachverhalt),
        ("Rechtliche Analyse", analyse),
        ("Ergebnis und Empfehlung", ergebnis),
        ("Vorbehalte / Disclaimer",
         "Dieses Gutachten beruht auf den uns vorgelegten Unterlagen sowie auf der zum Zeitpunkt der Erstellung "
         "geltenden Rechts- und Verwaltungspraxis. Eine Aenderung der zugrundeliegenden Tatsachen oder der "
         "Rechtslage kann zu einer abweichenden Beurteilung fuehren. Das Gutachten ist ausschliesslich fuer "
         "den internen Gebrauch der Mandantin bestimmt; Weitergabe an Dritte (insb. Wirtschaftspruefer, "
         "Behoerden) nur nach gesonderter Freigabe."),
        ("Schlussbestaetigung",
         f"Stuttgart/Frankfurt, den {datum}\n\n{autor}, Partner, {kanzlei}"),
    ]
    write_doc(f"{BASE}/{fname}", H, titel, subtitle=f"Rechtsgutachten – {datum}", sections=sections, confidential=True)


rechtsgutachten("REA_Rechtsgutachten_IC_Darlehen_Zinssatz_2022.docx",
    "Rechtsgutachten – Verrechnungspreiskonformer Zinssatz fuer Konzern-Darlehen RHO an Tochtergesellschaften",
    "Beurteilung der Verrechnungspreis- und Steuerrechtsrisiken eines Konzerninterner Darlehensvergabe von "
    "RHO an REG, RSG, RPL, RCZ, RHU und RCN in Hoehe von insgesamt 80 Mio. EUR; insbesondere Zinssatz nach OECD-"
    "TPG-Kapitel X (Financial Transactions), §§ 1, 1a AStG und BMF-Verwaltungsgrundsaetze 2021/2023.",
    "RHO Brennhagen Holding GmbH vergibt seit 2018 Konzerndarlehen an die operativen Toechter (Refinanzierungs"
    "saetze ueber das Konsortialkredit-RCF 2022). Die Pruefung gemaess BP-2018-2021 (Finanzamt Stuttgart-Mitte) "
    "befindet sich in der Schlussbesprechung mit Diskussionspunkt Stand-alone-Rating der Toechter und "
    "Implicit-Support-Argument. Konzern beabsichtigt eine APA-Anfrage Bundeszentralamt fuer Steuern (BZSt) "
    "fuer die Jahre 2023-2027.",
    "Massgeblich ist § 1 AStG (Fremdvergleichsgrundsatz) i. V. m. OECD-TPG 2022 Kap. X. Der Zinssatz ist nach "
    "Stand-alone-Rating der Schuldnergesellschaft zu bestimmen (Synthetic-Rating-Methode S&P/Moody's), "
    "gegebenenfalls mit Notch-Up fuer »Implicit Support« (OECD 10.92 ff.; BFH-Urteil v. 18.5.2021 I R 4/17 – "
    "»Konzern-Garantenhaftung«). Margenaufschlag auf risikofreien Basis-Zinssatz (EURIBOR + Spread). "
    "Anwendung pro Tochtergesellschaft: REG/RSG (BBB-Rating Synthetic, +1 Notch Implicit) – Margenaufschlag "
    "150 bp; RPL/RCZ/RHU (BB+/BB Rating Synthetic, +1 Notch) – 220 bp; RCN (BB- Synthetic) – 300 bp.",
    "Empfehlung: Anwendung der gestaffelten Margen wie oben. Vorbereitung einer Verrechnungspreis-Local-File "
    "(Stand-alone-Rating, Comparables-Studie, Implicit-Support-Memo). Empfehlung APA-Antrag BZSt fuer 2023-2027 "
    "zur Rechtssicherheit (gegenseitiger Verstaendigungsweg DE/PL/CZ/HU/CN). Risiko BP-Mehrsteuer 2018-2021 "
    "geschaetzt 1,8-2,4 Mio. EUR (Brennhagen hat hierfuer Rueckstellung in IFRS-Konzernabschluss 2022 gebildet).",
    HM_FIN, "Hengeler Mueller Partnerschaft mbB", "14. November 2022")

rechtsgutachten("REA_Rechtsgutachten_LkSG_Anwendbarkeit_2023.docx",
    "Rechtsgutachten – Anwendbarkeit des LkSG ab 1.1.2024 auf die Brennhagen-Gruppe",
    "Pruefung, ob die Brennhagen Elektronik AG bzw. einzelne Konzerngesellschaften ab 1.1.2024 in den Anwendungsbereich "
    "des Lieferkettensorgfaltspflichtengesetzes (LkSG) fallen, und Beratung zur Aufbau einer LkSG-konformen "
    "Compliance-Organisation und Lieferanten-Due-Diligence.",
    "Die Brennhagen-Gruppe beschaeftigt im Inland (REA, RHO, REG, RSG) zusammen rund 1.245 FTE (Stand 31.12.2023). "
    "Damit wird die Schwelle von 1.000 Beschaeftigten gemaess § 1 Abs. 1 Nr. 2 LkSG ab 1.1.2024 ueberschritten "
    "(zuvor 3.000-Schwelle, ab 1.1.2024 1.000-Schwelle). Die Holding-Mutter (REA) gilt als »Unternehmen« i. S. d. "
    "§ 2 Abs. 6 LkSG; konzernzugehoerige Toechter werden zugerechnet, soweit beherrschender Einfluss besteht "
    "(§ 2 Abs. 6 S. 3 LkSG).",
    "Anwendbarkeit ab 1.1.2024 zu bejahen. Pflichten gemaess §§ 3-10 LkSG: Risikomanagement, Praeventions"
    "massnahmen, Beschwerdeverfahren, Abhilfemassnahmen, Dokumentation und Berichterstattung. Erste BAFA-"
    "Berichterstattung bis 30.4.2025 fuer Berichtsjahr 2024. Empfohlene Organisation: Bestellung eines "
    "LkSG-Menschenrechtsbeauftragten (vorhanden: CSO " + CSO + "); Aufnahme der CoC-Klausel in Lieferanten"
    "vertraegen; Risikoanalyse Tier-1 jaehrlich; Hinweisgebersystem (SPEAK-UP@Brennhagen vorhanden).",
    "Brennhagen faellt ab 1.1.2024 unter LkSG. Empfohlene Sofortmassnahmen: (1) Bestellung CSO als Menschenrechts"
    "beauftragter formalisieren; (2) konzernweite LkSG-Risikoanalyse Tier-1 und ggf. Tier-2 (Cobalt/Lithium); "
    "(3) CoC fuer Geschaeftspartner ueberarbeiten und ueber DocuSign rollouten; (4) Beschwerdeverfahren um "
    "LkSG-spezifische Hinweis-Kategorien erweitern; (5) erste Berichterstattung bis 30.4.2025 vorbereiten. "
    "Risiko Bussgeld § 24 LkSG: bis 8 Mio. EUR / 2 % Konzern-Umsatz.",
    HM_PARTNER, "Hengeler Mueller Partnerschaft mbB", "20. Juli 2023")

rechtsgutachten("REA_Rechtsgutachten_Übernahme_Czech_GmbH_2018.docx",
    "Rechtsgutachten – Erwerb der spaeteren Brennhagen CZ s.r.o. (Asset-Deal vs. Share-Deal)",
    "Pruefung der gesellschafts-, steuer- und kartellrechtlichen Aspekte des geplanten Erwerbs der "
    "ElektroCZ s.r.o. (Brno) im Wege eines Share-Deals zur Errichtung der Brennhagen CZ s.r.o. (Steckverbinder-"
    "Produktion).",
    "Brennhagen Holding GmbH (RHO) erwirbt 100 % der Geschaeftsanteile der ElektroCZ s.r.o., Brno, von der "
    "tschechischen Verkaeufer-Familie zum Kaufpreis von 18,5 Mio. EUR (Multiple ~7x EBITDA 2017). Die Zielgesellschaft "
    "beschaeftigt 480 MA, produziert Steckverbinder fuer Tier-1- und Aftermarket-Kunden. Closing geplant fuer "
    "Q4 2018.",
    "Share-Deal vorzugswuerdig: (1) tschechische Privatisierungs-Rueckabwicklungs-Risiken bei Asset-Deal "
    "(Grundstuecks-Restitution); (2) Steueroptimierung (Beteiligungsbuchwert, ggf. Step-Up nicht moeglich, aber "
    "Verlustvortraege EUR 1,2 Mio bleiben erhalten); (3) Arbeitnehmer-Uebergang automatisch (vermeidet § 613a-"
    "Analogie tschechisch); (4) Lieferanten-/Kundenvertraege bestehen fort (kein Change-of-Control in den "
    "wesentlichen Vertraegen). Kartellrechtlich: keine Anmeldepflicht (Schwellen UOHS nicht erreicht; auch "
    "FKVO Schwellen nicht). Steuerlich: §§ 8b KStG-Schachtelprivileg auf Ausschuettungen anwendbar; "
    "Aufstockungsgewinn bei Share-Deal nur bei Anteilsveraeusserung, daher Plan: Strukturierung als "
    "Dauer-Holding-Tochter.",
    "Empfehlung Share-Deal. SPA-Verhandlung mit Standard-Garantien (Title/Capacity, Tax, Compliance, Employment, "
    "Litigation, IP, Environment), Cap 30 % Kaufpreis, Survival 18/36 Monate. Escrow 10 % fuer 18 Monate. "
    "Closing-Bedingungen: kartellrechtl. Unbedenklichkeit nicht erforderlich; Aufsichtsrat-Zustimmung RHO "
    "erforderlich (§ 13 Satzung Holding). W&I-Versicherung empfohlen (Cap 9,25 Mio.); Brennhagen hat Versicherung "
    "bei Allianz GCS abgeschlossen.",
    HM_PARTNER, "Hengeler Mueller Partnerschaft mbB", "15. September 2018")


# =============================================================================
# 3) LEGAL COMPLIANCE REPORTS PER SUBSIDIARY
# =============================================================================
def legal_comp_report(fname, sub_short, jahr):
    sub = S[sub_short]
    sections = [
        ("Berichtsgegenstand und Berichterstatter",
         f"Dieser Legal-&-Compliance-Report betrifft die {sub['name']} ({sub['short']}), {sub['city']} "
         f"({sub['country']}). Berichtsperiode: 1.1.{jahr} – 31.12.{jahr}. Berichterstatter: lokaler Compliance-"
         f"Beauftragter / Werksleitung in Abstimmung mit Konzern-Compliance ({CO}, extern Hengeler Mueller). "
         f"Adressat: Pruefungsausschuss Aufsichtsrat REA (Vorsitz Prof. Dr.-Ing. Voss) und Vorstand REA "
         f"(CEO Anna Mueller, CFO Laura Bauer)."),
        ("Gesellschaftsrechtliche Daten",
         f"Gesellschaftsform und HRB/Register: {sub['hrb']}. Mitarbeiter: {sub['employees']} FTE. "
         f"Umsatz {jahr}: ca. {sub['revenue_mio']} Mio. EUR. Funktion im Konzern: {sub['focus']}. "
         f"100 %-Tochter von RHO Brennhagen Holding GmbH."),
        ("Compliance-Status nach Themen", [
            ["Thema", "Status", "Wesentliche Massnahmen"],
            ["Anti-Korruption (ABC)", "gruen",
             "ABC-Schulungen 96 %-Coverage; 0 Faelle > 50 kEUR; jaehrliches ABC-Audit (Hengeler Mueller); "
             "Gifts-&-Hospitality-Register aktiv"],
            ["Kartellrecht", "gruen",
             "Antitrust-Schulungen Vertrieb/Einkauf abgeschlossen; Dawn-Raid-SOP an alle Manager; "
             "0 BKartA/UOKiK/UOHS-Verfahren"],
            ["Sanktionen / Exportkontrolle", "gelb" if jahr == "2023" else "gruen",
             "ComplyCube-Screening implementiert; Russland-Sanktionen vollstaendig umgesetzt; "
             "Endkundenpruefung bei Aftermarket-Bestellungen verschaerft"],
            ["DSGVO / Privacy", "gruen",
             f"VVT aktuell; DSB {DSB}; 0 meldepflichtige Datenpannen schwerer Art im Berichtsjahr"],
            ["LkSG / Menschenrechte", "gruen",
             "CoC-fuer-Geschaeftspartner ausgerollt; KYC-Pruefung Tier-1 erfolgt; Hinweisgebersystem aktiv"],
            ["Arbeits-/Sozialrecht", "gruen",
             "Keine wesentlichen Verfahren; Betriebsrat / Gewerkschaft etabliert; Tarifkonformitaet bestaetigt"],
            ["Umwelt / Genehmigungen", "gruen",
             "Alle BImSchG-/lokalen Umweltgenehmigungen gueltig; ISO 14001 zertifiziert"],
            ["Steuer-Compliance", "gruen",
             "Steuererklaerungen fristgerecht eingereicht; keine Mehrsteuern aus Betriebspruefungen"],
        ]),
        ("Wesentliche Rechtsstreitigkeiten",
         f"Im Berichtsjahr {jahr} bestehen bei {sub['short']} folgende Verfahren: ein arbeitsrechtliches "
         f"Verfahren (Kuendigungsschutz, Streitwert < 50 kEUR, Erfolgsaussicht gut), ein Lieferantendispute "
         f"(Reklamation einer fehlerhaften Charge, Streitwert ca. 120 kEUR, in Vergleichsverhandlung). "
         f"Patentstreitigkeiten: keine. Strafrechtliche Verfahren: keine. Behoerdliche Untersuchungen: keine."),
        ("Schulungs- und Awareness-Status",
         f"Pflicht-E-Learnings im LMS Brennhagen Learn (SuccessFactors) – Coverage {jahr}: "
         f"Code of Conduct 98 %, ABC 96 %, Antitrust 94 %, Sanktionen 92 %, DSGVO 97 %, LkSG (Einkauf) 95 %. "
         f"Praesenz-Briefings durch Hengeler Mueller fuer Top-Management vor Ort durchgefuehrt."),
        ("Hinweisgebersystem",
         f"Im Berichtsjahr {jahr} gingen bei {sub['short']} keine Hinweise mit substantiierten Anhaltspunkten "
         f"fuer schwere Verstoesse ueber SPEAK-UP@Brennhagen (EQS Group) oder den externen Ombudsmann {OMB} ein. "
         f"Lokale Vorfaelle wurden ueber HR-Linie bearbeitet (Bagatell-Bereich)."),
        ("Risikoeinschaetzung und Ausblick",
         f"Das Compliance-Risiko bei {sub['short']} wird fuer {jahr} insgesamt als »niedrig bis mittel« "
         f"eingestuft. Schwerpunkte fuer das Folgejahr: weitere Verschaerfung der Sanktionsscreening-Disziplin "
         f"(insbesondere RCN/RPL/RCZ wegen geografischer Naehe zu Sanktionsregimen), Vertiefung der LkSG-"
         f"Tier-2-Sichtbarkeit, ISO 27001-Roll-out (in Vorbereitung)."),
        ("Erklaerung und Bestaetigung",
         f"Der Berichterstatter bestaetigt die Vollstaendigkeit und Richtigkeit der vorstehenden Angaben "
         f"nach bestem Wissen. Etwaige spaetere Erkenntnisse werden im Wege eines Nachtragsberichts "
         f"unverzueglich an die Konzern-Compliance gemeldet."),
        ("Unterschrift",
         signatures(f"Werksleitung {sub['short']}", "Geschaeftsfuehrung / Werksleitung", sub['name'],
                    CO, "Compliance Officer Konzern (extern)", "Hengeler Mueller",
                    place=sub['city'], date_str_=f"31. Januar {int(jahr)+1}")),
    ]
    write_doc(f"{BASE}/{fname}", H,
              f"Legal-&-Compliance-Report {jahr} – {sub['name']}",
              subtitle=f"Berichtsjahr {jahr} – Adressat: Pruefungsausschuss Aufsichtsrat REA",
              sections=sections)


for short in ["REG", "RSG", "RPL", "RCZ", "RHU", "RCN"]:
    for jahr in ["2021", "2022", "2023"]:
        # filename variants observed
        suffix = ""
        if short == "RCN" and jahr == "2023":
            suffix = "_ALT"
        fname = f"{short}_Legal_Compliance_Report_{jahr}{suffix}.docx"
        # also handle the RHU 2022 missing - not in list, skip if not present
        import os
        if os.path.exists(f"{BASE}/{fname}"):
            legal_comp_report(fname, short, jahr)


# =============================================================================
# 4) RISK ASSESSMENTS PER SUBSIDIARY x TOPIC
# =============================================================================
def risk_assessment(fname, sub_short, thema, jahr):
    sub = S[sub_short]
    topic_blocks = {
        "Datenschutz": {
            "rechtsgrundlagen": [
                "Datenschutz-Grundverordnung (DSGVO – Verordnung (EU) 2016/679)",
                "Bundesdatenschutzgesetz (BDSG) bzw. lokales Umsetzungsgesetz",
                "Arbeitsrechtliche Spezialregelungen (§ 26 BDSG bzw. Aequivalent)",
                "Branchenleitfaeden BfDI / EDSA",
            ],
            "risiken": [
                "Unzureichende Loeschkonzepte bei Personalakten und Bewerberunterlagen",
                "Drittlandtransfer (z. B. SAP-Cloud, MS365) ohne hinreichende Standard-Vertragsklauseln",
                "Fehlende DSFA bei neuen IT-Tools (CRM, MES, BI)",
                "Phishing/Social-Engineering-Angriffe – Risiko unbefugter Datenoffenlegung",
                "Mitarbeiter-Monitoring (Zeiterfassung, Zugangskontrolle, Videoueberwachung)",
            ],
            "mitigation": [
                "Konzernweites VVT (Art. 30) aktualisiert; lokale Datenpannen-SOPs implementiert",
                "DSB " + DSB + " (extern, Lehmann & Partner) – jaehrliche Audits im Werk",
                "Standard-AV-Vertrag (Konzern-Template) fuer Auftragsverarbeiter durchgesetzt",
                "Pflicht-E-Learning DSGVO Basics jaehrlich (Coverage >=95 %)",
                "Phishing-Simulationen halbjaehrlich; MFA flaechendeckend",
            ],
            "bewertung": "Brutto-Risiko: mittel. Netto-Risiko nach Massnahmen: niedrig. Akzeptabel.",
        },
        "Korruptionsrisiko": {
            "rechtsgrundlagen": [
                "UK Bribery Act 2010; US FCPA; §§ 299-301, 331-334 StGB",
                "OECD-Anti-Bribery-Convention; UNCAC",
                "Konzern-Anti-Korruptions-Richtlinie 2023 (Version 3.0)",
            ],
            "risiken": [
                "Geschenke/Bewirtungen gegenueber OEM-Einkauf/Qualitaet (Insbesondere im Aftermarket)",
                "Vermittlungsprovisionen Vertriebsagenten in Schwellenmaerkten",
                "Genehmigungen Behoerden (Bau, Umwelt, Zoll) in CN/PL/CZ/HU",
                "Sponsoring / Spenden (insb. politische Funktionstraeger lokal)",
                "Facilitation Payments (Zoll, Inspektoren) in CN",
            ],
            "mitigation": [
                "Gifts-&-Hospitality-Register Pflicht; Cap 50 EUR / 100 EUR Bewirtung",
                "Vertriebsagenten nur nach TPDD; Provisionen marktgerecht; Audit-Right",
                "ABC-Schulung Pflicht-E-Learning, Coverage >=96 %",
                "Vier-Augen-Prinzip Genehmigungs-Kommunikation Behoerden",
                "Speak-Up-System + externer Ombudsmann " + OMB,
            ],
            "bewertung": "Brutto-Risiko: mittel. Netto-Risiko: niedrig. Akzeptabel.",
        },
        "Kartellrecht": {
            "rechtsgrundlagen": [
                "Art. 101, 102 AEUV; §§ 1, 19, 20 GWB",
                "EU-VerticalBER 2022/720; EU-Horizontalleitlinien 2023",
                "Konzern-Kartellrecht-Compliance-Handbuch 2023 (Version 4.1)",
            ],
            "risiken": [
                "Informationsaustausch in Branchenverbaenden (VDA, ZVEI, CLEPA)",
                "Vertikale Preisbindungen Aftermarket-Distribution",
                "Bid-Rigging-Risiken bei OEM-Konsortien (z. B. Mehrfach-Sourcing-Programme)",
                "Marktbeherrschende Stellung bei kundenspezifischen Tier-1-Spezialteilen",
                "Joint Ventures / Forschungs-Kooperationen (z. B. mit Universitaeten/Bosch/Conti)",
            ],
            "mitigation": [
                "Antitrust-Briefings vor jeder Verbandssitzung; Dokumentationspflicht",
                "Dawn-Raid-Trainings jaehrlich (Top-50 Manager)",
                "Pre-Approval kartellsensitiver Kooperationen durch Hengeler Mueller",
                "Antitrust-E-Learning Pflicht (Vertrieb/Einkauf/Strategie), Coverage 94 %",
                "Whistleblower-System Speak-Up + externer Ombudsmann " + OMB,
            ],
            "bewertung": "Brutto-Risiko: mittel-hoch. Netto-Risiko: niedrig-mittel. Akzeptabel mit kontinuierlichem Monitoring.",
        },
        "Exportkontrolle": {
            "rechtsgrundlagen": [
                "AWG, AWV; EU-Dual-Use-VO 2021/821; US-EAR (soweit Reichweite)",
                "EU/UN/US Sanktionsregelungen (Russland, Belarus, Iran, Nordkorea)",
                "BMWi/BAFA-Hinweise und Embargo-Updates",
            ],
            "risiken": [
                "Lieferung ECU-900 (ggf. Dual-Use Annex I Kat. 3A) in sanktionierte Laender",
                "Diversionsrisiko ueber RCN Shanghai (Endkunde verdeckt)",
                "Aftermarket-Vertrieb an gelistete Einzelpersonen/Entitaeten",
                "Russland-/Belarus-Sanktionen: Catch-All-Klauseln (Art. 4 EU-Dual-Use-VO)",
                "Iran-Sanktionen: Komponenten in Endprodukten Dritter",
            ],
            "mitigation": [
                "ComplyCube-Screening bei Onboarding und Transaktion",
                "BAFA-Beauftragte " + BAFA + " (Group Tax, Doppelfunktion)",
                "End-User-Statements obligatorisch Aftermarket RCN",
                "Pflicht-E-Learning »Sanktionen & Exportkontrolle« (Coverage 92 %)",
                "Russland-/Belarus-Boykott seit 24.2.2022 vollstaendig umgesetzt",
            ],
            "bewertung": "Brutto-Risiko: mittel-hoch (geografisch). Netto-Risiko: niedrig-mittel. Verschaerftes Monitoring fortgesetzt.",
        },
    }
    block = topic_blocks[thema]
    sections = [
        ("Anlass und Scope",
         f"Diese Risikoanalyse fuer das Thema {thema} betrifft die {sub['name']} ({sub['short']}), "
         f"{sub['city']} ({sub['country']}). Berichtsjahr: {jahr}. Methodik: Brutto-/Netto-Bewertung "
         f"auf 5x5-Heatmap (Eintrittswahrscheinlichkeit x Schadenpotenzial). Verantwortlich: Werksleitung "
         f"in Abstimmung mit Konzern-Compliance ({CO}, extern Hengeler Mueller)."),
        ("Rechtsgrundlagen", ("list", block["rechtsgrundlagen"])),
        ("Identifizierte Risiken", ("list", block["risiken"])),
        ("Mitigation und Kontrollen", ("list", block["mitigation"])),
        ("Gesamtbewertung", block["bewertung"]),
        ("Befunde und Massnahmen",
         f"Im Rahmen der Risikoanalyse {jahr} wurden bei {sub['short']} keine wesentlichen neuen Risiko"
         f"bereiche identifiziert. Bestehende Kontrollmassnahmen werden als wirksam eingestuft. "
         f"Eine Detailrevision durch die Konzernrevision (CAE {CAE}) ist fuer das Folgejahr terminiert. "
         f"Re-Assessment turnusmaessig in 12 Monaten."),
        ("Schlussbemerkung",
         f"Diese Bewertung wird dem Pruefungsausschuss des Aufsichtsrats der REA sowie dem Vorstand zur "
         f"Kenntnis vorgelegt. Bei wesentlichen Aenderungen im Berichtsjahr erfolgt eine Aktualisierung."),
        ("Unterschriften",
         signatures(f"Werksleitung {sub['short']}", "Geschaeftsfuehrung / Werksleitung", sub['name'],
                    CO, "Compliance Officer (extern)", "Hengeler Mueller",
                    place=sub['city'], date_str_=f"15. Dezember {jahr}")),
    ]
    write_doc(f"{BASE}/{fname}", H,
              f"Risiko-Assessment {thema} {jahr} – {sub['name']}",
              subtitle=f"Compliance-Risikoanalyse – Berichtsjahr {jahr}",
              sections=sections)


for short in ["REG", "RSG", "RPL", "RCZ", "RHU", "RCN"]:
    for thema in ["Datenschutz", "Korruptionsrisiko", "Kartellrecht", "Exportkontrolle"]:
        # filename variants observed (e.g. RHU_RisikoAssessment_Datenschutz_2023_FINAL.docx)
        for suffix in ["", "_FINAL"]:
            fname = f"{short}_RisikoAssessment_{thema}_2023{suffix}.docx"
            import os
            if os.path.exists(f"{BASE}/{fname}"):
                risk_assessment(fname, short, thema, "2023")


# =============================================================================
# 5) INTERNE REVISIONS-BERICHTE
# =============================================================================
def interne_revision(fname, bereich, scope, schwerpunkte, findings, empfehlungen, jahr):
    sections = [
        ("Auftrag und Methodik",
         f"Die Konzernrevision (Internal Audit) der Brennhagen Elektronik AG hat im Berichtsjahr {jahr} "
         f"eine Revision des Bereichs »{bereich}« durchgefuehrt. Auftrag gemaess Audit-Plan, freigegeben "
         f"durch den Pruefungsausschuss des Aufsichtsrats (Vorsitz Prof. Dr.-Ing. Voss). Pruefungsleiter: "
         f"{CAE} (Chief Audit Executive). Methodik: risikoorientierte Pruefung gemaess IIA-Standards "
         f"(International Standards for the Professional Practice of Internal Auditing) und IDW PS 320."),
        ("Pruefungsumfang (Scope)", scope),
        ("Pruefungsschwerpunkte", ("list", schwerpunkte)),
        ("Wesentliche Feststellungen", findings),
        ("Empfehlungen und Massnahmen", ("clauses", empfehlungen)),
        ("Stellungnahme der Fachabteilung",
         f"Die geprueften Bereiche haben die Feststellungen anerkannt und sich zur Umsetzung der Empfehlungen "
         f"verpflichtet. Detaillierter Massnahmenplan im Tracking-Tool (Konzern-Revisions-Datenbank »Audit Hub«); "
         f"Follow-up durch Konzernrevision binnen 6 Monaten."),
        ("Gesamtbewertung",
         f"Das interne Kontrollsystem im Bereich {bereich} wird insgesamt als »weitgehend angemessen« "
         f"bewertet. Die identifizierten Findings sind als »minor« bzw. »moderate« einzustufen; "
         f"keine wesentlichen Schwaechen (»significant deficiencies«) festgestellt."),
        ("Verteiler",
         "Vorstand REA (CEO, CFO, COO, CTO), Compliance Officer " + CO + ", externer Wirtschaftspruefer "
         "KPMG (Dr. Brand) im Rahmen des Audit-Universe-Briefings, Pruefungsausschuss Aufsichtsrat."),
        ("Unterschriften",
         signatures(CAE, "Chief Audit Executive (CAE)", R["name"],
                    "Bereichsverantwortliche/r", "Geprueft", R["name"],
                    place="Stuttgart", date_str_=f"15. November {jahr}")),
    ]
    write_doc(f"{BASE}/{fname}", H,
              f"Bericht der Konzernrevision – {bereich} {jahr}",
              subtitle=f"Interne Revision gemaess IIA-Standards – Pruefungsjahr {jahr}",
              sections=sections, confidential=True)


# Internal Revision Standardisiertes Mapping
INTERNE_REVISIONEN = {
    "Einkauf": {
        "scope": "Pruefung des konzernweiten Beschaffungsprozesses (Source-to-Pay) am Beispiel ausgewaehlter Lieferanten / Warengruppen "
                 "(Elektronikkomponenten, MRO, IT-Services). Stichprobe: 65 Bestellvorgaenge >= 50 kEUR im Berichtsjahr.",
        "schwerpunkte": [
            "Lieferantenfreigabe und -qualifizierung (KYC, Compliance-Klauseln)",
            "Vier-Augen-Prinzip bei Bestellungen / SAP-Berechtigungen",
            "Vergabe-/Ausschreibungsprozess >= 250 kEUR (Three-Quote-Rule)",
            "Wareneingang vs. Rechnungspruefung (3-Way-Match SAP MM/FI)",
            "Vertragsmanagement und Compliance-Klausel-Coverage Rahmenvertraege",
        ],
        "findings":
            "Drei moderate Findings: (1) bei 4 von 65 Bestellungen unvollstaendige Three-Quote-Dokumentation; "
            "(2) bei 2 Faellen fehlte aktualisierte Lieferanten-CoC-Bestaetigung; (3) Compliance-Klausel in 8 "
            "von 30 Rahmenvertraegen veraltet (Stand 2019, nicht LkSG-konform). Keine Hinweise auf doloses "
            "Verhalten oder wesentliche Verluste. Empfohlene Korrektur-Prioritaet: hoch (Klausel-Update).",
        "empfehlungen": [
            ("§ 1 Three-Quote-Rule", [
                "Verschaerfung der Dokumentationspflicht im SAP-SRM (Pflichtfelder); rollierender Check durch "
                "Procurement Compliance; Stichproben durch Konzernrevision halbjaehrlich.",
            ]),
            ("§ 2 Lieferanten-CoC", [
                "Roll-out der ueberarbeiteten CoC-Version 2023 ueber DocuSign; Tracking der Akzeptanzquote (Ziel 100 % Top-200 Lieferanten).",
            ]),
            ("§ 3 Rahmenvertraege", [
                "Sammelaktion zur Aktualisierung der Compliance-/LkSG-Klauseln in allen Rahmenvertraegen "
                "(Eigentumsvorbehalt, Anti-Korruption, Audit-Right, LkSG-Klausel, Sanktionen). Deadline 30.6.2024.",
            ]),
        ],
    },
    "F_E_RSG": {
        "scope": "Pruefung des Forschungs- und Entwicklungsbereichs der RSG Brennhagen Software GmbH, Muenchen, "
                 "insbesondere im Hinblick auf Projekt-Charter-Prozesse, Stage-Gate-Reviews, ASPICE-Konformitaet, "
                 "IP-Management und Lieferanten-Compliance (Open-Source-Software).",
        "schwerpunkte": [
            "Projekt-Charter und Stage-Gate-Disziplin (Gate 0-Gate 4)",
            "ASPICE Level 2/3 Assessment-Status pro Projekt (ICP-3, BMS-12, ADAS-V4D, ECU-900)",
            "IP-Management: Erfinder-Meldungen, Patentantraege (Boehmert & Boehmert)",
            "Open-Source-Compliance (FOSSology-Scans, OSS-Reviewboard)",
            "Personalkostenverrechnung Konzerntoechter (TP-Konformitaet)",
        ],
        "findings":
            "Zwei moderate Findings: (1) ASPICE-Level-2-Assessment fuer BMS-12 verzoegert (Plan Q3, Ist Q4); "
            "(2) Open-Source-Inventar von 11 Komponenten zeigt 3 Lizenzen mit Copyleft-Risiko (GPL v2) – "
            "Reviewboard-Entscheidung steht aus. Erfinder-Meldungen termingerecht. Keine wesentlichen Maengel.",
        "empfehlungen": [
            ("§ 1 ASPICE-Roadmap", [
                "Beschleunigung des ASPICE-Level-2-Assessments BMS-12 in Zusammenarbeit mit Kugler Maag; "
                "Verbindliche Meilenstein-Reviews durch RSG-GL und Vorstands-CTO.",
            ]),
            ("§ 2 Open-Source-Compliance", [
                "Reviewboard-Entscheidung zu GPL-v2-Komponenten zu treffen (Trennung Userspace vs. Linker, "
                "ggf. Re-Implementation oder Lizenzwechsel) – Deadline 60 Tage.",
            ]),
            ("§ 3 TP-Konformitaet RSG-Personalverrechnung", [
                "Local-File / Verrechnungspreis-Dokumentation aktualisieren; Markup auf Personalkosten beibehalten "
                "(Cost-plus 8 %); Validierung durch Group Tax (Dr. Berger).",
            ]),
        ],
    },
    "Finance_Controlling": {
        "scope": "Pruefung des Group Controllings und Finanz-Reportings, insbesondere Monatsabschlussprozess "
                 "(Fast Close), HFM/SAP-FC-Konsolidierung, IFRS-Bilanzierungsrichtlinien-Compliance und "
                 "Internal Controls over Financial Reporting (ICFR).",
        "schwerpunkte": [
            "Monatsabschluss-Disziplin (Fast Close WT 5)",
            "ICFR Schluessel-Kontrollen (Account-Reconciliation, JE-Approval, Top-Side-Adjustments)",
            "Konsolidierung HFM / SAP S/4-FC (Eliminierungen IC)",
            "Hedge-Accounting nach IFRS 9 (FX-Forward, Hyundai, BMW, Stellantis-Volumina)",
            "Going-Concern-Prozess und Liquiditaetsforecast 13-Wochen-Tool",
        ],
        "findings":
            "Zwei moderate Findings: (1) bei 6 % der Konten-Reconciliations Zeitueberschreitung > WT 8; "
            "(2) Top-Side-Adjustments durch Group Controlling nicht durchgaengig mit Begruendung dokumentiert "
            "(MAR-Relevanz pruefen). Hedge-Accounting-Dokumentation als angemessen bestaetigt; ICFR-Kontrollen-"
            "Wirksamkeit weit ueberwiegend gegeben.",
        "empfehlungen": [
            ("§ 1 Konten-Reconciliations", [
                "Automatisierung im Tool Blackline rolloutweise (Q1/Q2 Folgejahr); SLA WT 5; "
                "Excellence Award fuer Top-Performer-Werke.",
            ]),
            ("§ 2 Top-Side-Adjustments Disziplin", [
                "Verbindliches Begruendungsfeld mit Reviewer-Workflow im Konsolidierungs-Tool; "
                "MAR-Pruefung durch Insider-/Compliance-Funktion bei materiellen Adjustments.",
            ]),
        ],
    },
    "HR": {
        "scope": "Pruefung der HR-Prozesse Konzern (Workday, ADP, SAP HCM) inkl. Recruiting, Onboarding, "
                 "Performance Management, Verguetungsprozess (Bonus, LTI), Betriebsrats-Compliance.",
        "schwerpunkte": [
            "Stammdatenqualitaet Workday/ADP (DSGVO-Konformitaet, Loeschkonzepte)",
            "Bonus-Pool-Berechnung und LTI-Allokation (Vorstands- und obere Fuehrungsebene)",
            "Betriebsrats-Mitbestimmung (insb. § 87 BetrVG-Themen)",
            "Diversity-KPIs (Frauenanteil; Sustainability-Linked-Margin RCF)",
            "Personalakten und Datenschutz",
        ],
        "findings":
            "Zwei minor Findings: (1) bei 3 % Workday-Datensaetzen Inkonsistenz HR-Profil vs. Vertrag (Korrektur "
            "laufend); (2) Diversity-KPI-Reporting konsistent, allerdings Mittelfristziel Frauenanteil im Top-150-"
            "Management noch unterhalb Plan (Ist 23 %, Ziel 30 % bis 2026). LTI-Allokation und Bonus-Pool-"
            "Berechnung rechnerisch korrekt; Mitbestimmungs-Compliance angemessen.",
        "empfehlungen": [
            ("§ 1 Workday-Stammdaten", [
                "Q-monatlicher Datenqualitaets-Report; Bereinigung Inkonsistenzen binnen 60 Tagen.",
            ]),
            ("§ 2 Diversity-Roadmap", [
                "Verschaerfung der Massnahmenpakete (Talent-Pipeline, Mentoring, Gender-Pay-Audit) "
                "zur Erreichung des 30 %-Ziels Frauen in Fuehrung bis 2026.",
            ]),
        ],
    },
    "IT_Infrastruktur": {
        "scope": "Pruefung der konzernweiten IT-Infrastruktur und -Sicherheit, insbesondere ISO-27001-Vorbereitung "
                 "(Roll-out 2024-2025), Cloud-Strategie (Azure/AWS), Network Segmentation, Identity & Access Management, "
                 "Backup/Disaster Recovery, Patch Management.",
        "schwerpunkte": [
            "ISO 27001:2022 Gap-Analyse und Stage-1/2-Vorbereitung",
            "IAM (Azure AD), MFA-Coverage, Privileged Access Management (CyberArk)",
            "Backup/Restore-Tests (Tier-1-Systeme SAP S/4, Salesforce, Workday)",
            "Patch-Management Endpoint und Server (CrowdStrike, Tenable)",
            "Disaster Recovery & Business Continuity (Tier-1 RPO 4h, RTO 24h)",
        ],
        "findings":
            "Drei moderate Findings: (1) MFA-Coverage 97 %, Restquote Service-Accounts ohne Hardware-Token; "
            "(2) Patch-SLAs in 6 % der Faelle ueberschritten (insb. Produktions-OT-Systeme REG/RPL); "
            "(3) BCM-Test-Frequenz Tier-2-Anwendungen unterhalb Soll (1x p. a. vs. 2x p. a.). "
            "ISO-27001-Reifegrad: Gap-Analyse abgeschlossen, Stage-1-Audit Q1/2024 geplant.",
        "empfehlungen": [
            ("§ 1 MFA-Endgame", [
                "100 %-MFA inkl. Service-Account-Vault binnen 90 Tagen.",
            ]),
            ("§ 2 OT-Patch-Management", [
                "Spezielles OT-Patch-Fenster pro Werk (REG/RPL) mit Maintenance-Plan abstimmen; SLA 30 Tage critical.",
            ]),
            ("§ 3 BCM-Test-Frequenz", [
                "Verdoppelung der Tier-2-DR-Tests; Auditierte Berichte im Konzern-BCM-Tool.",
            ]),
        ],
    },
    "Legal_Compliance": {
        "scope": "Pruefung der Konzern-Compliance-Funktion und Legal-Department-Prozesse, insb. "
                 "Mandatsteuerung Hengeler Mueller, Vertragsdatenbank (CLM-Tool), Litigation-Tracking, "
                 "Compliance-Trainings, Hinweisgebersystem.",
        "schwerpunkte": [
            "Mandatsteuerung externer Kanzleien (Hengeler Mueller, lokale Counsels)",
            "Contract Lifecycle Management (CLM) und Klausel-Bibliothek",
            "Litigation-Tracking und Reserven (IAS 37) Reporting",
            "Compliance-Training-Tracker und Coverage-KPIs",
            "Hinweisgebersystem »SPEAK-UP« (EQS Group) – Vertraulichkeit, Fristen, Statistik",
        ],
        "findings":
            "Zwei minor Findings: (1) CLM-Tool Adoption Tochtergesellschaften RPL/RHU noch unter 70 % "
            "(Konzern-Schnitt 86 %); (2) Compliance-Training-Coverage Antitrust auf 94 % – Ziel 98 %. "
            "Litigation-Reserven (IAS 37) konsistent mit Hengeler-Counsel-Einschaetzungen; "
            "Hinweisgebersystem-Statistik vollstaendig und vertraulich.",
        "empfehlungen": [
            ("§ 1 CLM-Adoption", [
                "Verbindlicher Roll-out RPL/RHU Q1/Q2-Folgejahr; KPI-Reporting an Vorstand quartalsweise.",
            ]),
            ("§ 2 Antitrust-Training", [
                "Reminder-Kampagne und Eskalation an Fuehrungskraefte zur Erreichung 98 %-Coverage bis 31.12.",
            ]),
        ],
    },
    "Produktion_REG": {
        "scope": "Pruefung der Produktion am Hauptwerk REG Heilbronn (820 FTE) mit Schwerpunkt OEE, "
                 "Qualitaetskosten, IATF-16949-Konformitaet, Arbeitsschutz, Energie-Management (ISO 50001) "
                 "und Instandhaltungsdisziplin.",
        "schwerpunkte": [
            "OEE-Reporting Tier-1/Tier-2 Linien (Soll >85 %)",
            "Qualitaetskosten (NoQ) und CoPQ-Allokation (Scrap, Rework, External)",
            "IATF-16949-Kernprozesse (APQP, PPAP, FMEA, MSA)",
            "Arbeitssicherheit (TRIR, LTIR, ASA-Reviews)",
            "ISO 50001-Energieaudit und -Massnahmenverfolgung",
        ],
        "findings":
            "Zwei moderate Findings: (1) OEE-Linie ICP-3-MK1 in Q2-Q3 unterhalb Plan (82 % vs. 86 %) wegen "
            "Werkzeugverschleiss-Wechsel; (2) IATF-Layered Process Audit (LPA) in 4 % der Faelle Fristabweichungen. "
            "TRIR/LTIR im Korridor. ISO-50001 zertifiziert; Massnahmen-Pipeline gut gefuellt.",
        "empfehlungen": [
            ("§ 1 ICP-3-MK1-OEE", [
                "Verschaerftes Tooling-Maintenance-Programm; Wechselzyklus reduzieren; OEE-Daily-Boards.",
            ]),
            ("§ 2 LPA-Disziplin", [
                "Aktualisierung der LPA-Auditoren-Plaene; Anbindung an SAP QM zur Fristenueberwachung.",
            ]),
        ],
    },
    "Produktion_RPL": {
        "scope": "Pruefung der EMS-/SMD-Produktion RPL Katowice (960 FTE) mit Fokus auf EMS-Linien-Effizienz, "
                 "Personalstellen (HR-Buch), Energieeffizienz und Lieferantenanbindung.",
        "schwerpunkte": [
            "EMS-/SMD-Linien-Auslastung und OEE",
            "Personalbuch und Arbeitszeitkonten (PL-Arbeitsrecht)",
            "Energie- und Wasserkennzahlen (ISO 14001/50001)",
            "Lieferanten-Lead-Time und Bestand (Days-on-Hand)",
            "Tooling-Investition und Capex-Disziplin",
        ],
        "findings":
            "Zwei moderate Findings: (1) SMD-Linie 4 (BMS-12) mit OEE 81 % wegen Materialengpaessen JIT; "
            "(2) Days-on-Hand auf 38 Tage angestiegen (Ziel <=28). Personalbuch IT-konform; Energie-KPI gut. "
            "Tooling-Capex-Disziplin angemessen.",
        "empfehlungen": [
            ("§ 1 Materialfluss BMS-12", [
                "Zweitlieferanten-Qualifizierung fuer kritische Komponenten; JIT-/JIS-Disziplin verbessern.",
            ]),
            ("§ 2 Bestandsmanagement", [
                "DOH-Reduktion durch ABC-/XYZ-Analyse; Kanban-Erweiterung; SAP IBP-Tool-Roll-out.",
            ]),
        ],
    },
    "Supply_Chain": {
        "scope": "Pruefung der Konzern-Supply-Chain-Funktion mit Fokus auf S&OP-Disziplin, Lieferanten-Risiko-"
                 "management, OTD-Kennzahlen, Sicherheitsbestaende und Liefer-Performance gegenueber OEMs.",
        "schwerpunkte": [
            "S&OP-Prozess und Forecast-Genauigkeit (MAPE)",
            "Lieferanten-Risiko-Scoring (Tier-1/Tier-2)",
            "On-Time-Delivery (OTD) gegenueber BMW, VW, MB, Stellantis, Hyundai",
            "Sicherheitsbestaende kritischer Komponenten (Halbleiter, Magnetwerkstoffe)",
            "Konventionalstrafen und Pönalen (Litigation-Schnittstelle)",
        ],
        "findings":
            "Zwei moderate und ein significant Finding: (1) MAPE 12-Monatsforecast >25 % bei zwei Produkten "
            "(ICP-3 und LightCtrl-7) ueber Plan; (2) Lieferanten-Risiko-Scoring deckt 92 % des Tier-1-Volumens "
            "(Ziel >=95 %); (3) OEM-Konventionalstrafe Stellantis Q3 in Hoehe 180 kEUR wegen verspaeteter "
            "Ramp-Up-Lieferung BMS-12 – Root-Cause-Analyse durchgefuehrt; Verfahren mit Stellantis abgeschlossen.",
        "empfehlungen": [
            ("§ 1 Forecast-Genauigkeit", [
                "Demand-Sensing-Tool-Pilot (o9 Solutions) fuer ICP-3 und LightCtrl-7; Trainings Sales/Demand-Planner.",
            ]),
            ("§ 2 Risiko-Coverage Tier-1", [
                "Resilinc-Onboarding fuer fehlende 3 % Lieferanten; quartalsweises Update an SC-Leitung.",
            ]),
        ],
    },
    "Vertrieb": {
        "scope": "Pruefung der Vertriebs- und Account-Management-Funktion mit Fokus auf RFQ-Disziplin, "
                 "Quotation-Margen, CRM-Pflege (Salesforce), Verguetungs-Compliance (Verguetungsplaene) und "
                 "OEM-Beziehungsmanagement.",
        "schwerpunkte": [
            "RFQ-/Quotation-Margin-Disziplin (Approval-Hierarchie nach Volumen/Laufzeit)",
            "Salesforce-CRM-Pflege und Forecast-Hygiene",
            "Verguetungsplaene und SPIF-Aktionen (Compliance, KArtellrecht)",
            "OEM-Beziehungsmanagement und Eskalations-Prozesse",
            "Vertriebsagenten (TPDD, Provisions-Disziplin)",
        ],
        "findings":
            "Zwei moderate Findings: (1) bei 8 % der RFQs unvollstaendige Approval-Dokumentation (Margen-Approval); "
            "(2) Salesforce-Forecast-Hygiene Quartalswechsel-Bias (sandbagging) festgestellt – Trainings empfohlen. "
            "Verguetungsplaene compliant; OEM-Eskalations-Prozesse etabliert; TPDD-Coverage 100 % der aktiven Agenten.",
        "empfehlungen": [
            ("§ 1 RFQ-Approval-Disziplin", [
                "Pflichtfeld-Workflow in Salesforce/SAP fuer Margen-Approval; Stichprobenpruefung quartalsweise.",
            ]),
            ("§ 2 Forecast-Hygiene", [
                "Trainings Forecast-Disziplin; Monatliche Pipeline-Reviews mit Vertriebsleitung und CFO.",
            ]),
        ],
    },
}

# Map: filenames -> (bereich_key, jahr)
import re
import os
ir_pattern = re.compile(r"REA_Interne_Revision_(.+?)_(\d{4})(?:_(?:FINAL_v2|ALT))?\.docx")
for fname in sorted(os.listdir(BASE)):
    if not fname.startswith("REA_Interne_Revision_"):
        continue
    m = ir_pattern.match(fname)
    if not m:
        continue
    bereich_key = m.group(1)
    jahr = m.group(2)
    if bereich_key not in INTERNE_REVISIONEN:
        # Try fallback (e.g. some files have weird underscore patterns)
        continue
    block = INTERNE_REVISIONEN[bereich_key]
    bereich_display = bereich_key.replace("_", " ")
    interne_revision(fname, bereich_display, block["scope"], block["schwerpunkte"],
                     block["findings"], block["empfehlungen"], jahr)


# =============================================================================
# 6) DSGVO Datenpannen-Meldungen
# =============================================================================
DATENPANNEN_FAELLE = [
    ("verlorenes_laptop", "Verlust eines verschluesselten Mitarbeiter-Laptops",
     "Ein Aussendienst-Mitarbeiter Vertrieb verlor an einem Flughafen sein dienstliches Notebook. "
     "Das Geraet war vollverschluesselt (BitLocker AES-256) und durch MFA gesichert.",
     "Restrisiko gering (Verschluesselung + MFA); keine Meldepflicht Art. 33 DSGVO an Aufsicht erforderlich; "
     "Dokumentation im internen Datenpannen-Register."),
    ("fehlversand", "Fehlversand einer E-Mail mit Personalbezug",
     "HR-Sachbearbeiterin versandte einen Abrechnungs-Anhang an die falsche Outlook-Adresse "
     "(Tippfehler im Empfaengerfeld; betroffen: Adressdaten 1 Mitarbeiter).",
     "Empfaenger zur Loeschung aufgefordert und Bestaetigung erhalten; Risiko gering; "
     "interne Schulung der HR-Abteilung; keine Meldung an LfDI Baden-Wuerttemberg, "
     "aber Eintrag im Datenpannen-Register."),
    ("phishing_pw", "Phishing-Vorfall mit Passwort-Kompromittierung",
     "Mitarbeiter klickte auf Phishing-Link; sein Konto wurde kurzzeitig unbefugt zugreifbar. "
     "Microsoft Defender erkannte den Login binnen 12 Minuten und sperrte das Konto.",
     "Konto-Reset, MFA-Refresh, IT-Forensik bestaetigt: kein Abfluss personenbezogener Daten an "
     "Dritte erkennbar. Risikobewertung: niedrig. Meldung an LfDI BW als Vorsorgemeldung erfolgt."),
    ("misconfig_share", "Fehlkonfigurierte SharePoint-Freigabe",
     "Ein Projekt-SharePoint wurde versehentlich »alle authentifizierten Konzern-User« lesbar "
     "gemacht (Soll: nur Projektteam). Enthielt teilweise Personalkennzahlen.",
     "Zugriff binnen 2 Tagen entdeckt und korrigiert; Audit-Logs zeigen keine externen Downloads; "
     "interne Awareness-Massnahme; Meldung an LfDI BW erfolgt (niedriges Risiko)."),
    ("usb_unverschluesselt", "Verlust eines unverschluesselten USB-Sticks",
     "Mitarbeiter Produktion REG verlor unverschluesselten USB-Stick (private Nutzung trotz "
     "Verbot in IT-Richtlinie). Enthielt eine Kopie einer Arbeitszeit-Excel.",
     "Vorfall meldepflichtig; Meldung an LfDI BW innerhalb 72 Stunden gemaess Art. 33 DSGVO erfolgt; "
     "betroffene Mitarbeitende informiert (Art. 34); Sanktion gegen Mitarbeiter (Abmahnung); "
     "Roll-out USB-Sticks-Verschluesselungsrichtlinie verschaerft."),
]

import os
for fname in sorted(os.listdir(BASE)):
    if not fname.startswith("DSGVO_Datenpanne_"):
        continue
    # extract id from filename (last digits)
    m = re.match(r"DSGVO_Datenpanne_(\d{4})_(\d{3})(?:_.+)?\.docx", fname)
    if not m:
        continue
    jahr, lfd = m.group(1), m.group(2)
    fall = DATENPANNEN_FAELLE[int(lfd) % len(DATENPANNEN_FAELLE)]
    fcode, ftitle, sachverhalt, bewertung = fall
    sections = [
        ("Aktenzeichen und Meldedaten",
         f"Aktenzeichen intern: DSGVO-{jahr}-{lfd}. Eingangsdatum: {jahr}; Meldung erstellt durch "
         f"Local IT-Compliance / DSB-Buero. Verantwortlicher i. S. d. Art. 4 Nr. 7 DSGVO: "
         f"Brennhagen Elektronik AG. DSB: {DSB} (extern). Konzern-Compliance: {CO}."),
        ("Sachverhalt", sachverhalt),
        ("Datenarten und Kategorien Betroffener", ("list", [
            "Personenbezogene Daten von Beschaeftigten der Brennhagen-Gruppe",
            "Ggf. Geschaeftskontaktdaten externer Personen (B2B-Ansprechpartner)",
            "Keine besonderen Datenkategorien (Art. 9 DSGVO) bekannt",
        ])),
        ("Sofortmassnahmen", ("list", [
            "IT-Forensik durch Group IT-Security (Lead: Group CISO)",
            "Sperrung betroffener Accounts / Zugriffe",
            "Information Vorstand (CFO) und Datenschutzbeauftragter " + DSB,
            "Pruefung Meldepflicht Art. 33 DSGVO binnen 72 Stunden",
            "Dokumentation im internen Datenpannen-Register",
        ])),
        ("Risikoeinschaetzung und Massnahmenkatalog", bewertung),
        ("Meldung an Aufsichtsbehoerde",
         "Pruefung der 72-Stunden-Meldefrist gemaess Art. 33 DSGVO durch DSB " + DSB + " erfolgt. "
         "Sofern Meldepflicht bejaht, Online-Meldung an LfDI Baden-Wuerttemberg via Datenpannen-Portal "
         "(https://www.baden-wuerttemberg.datenschutz.de). Aktenzeichen LfDI wird im Tracking-Tool "
         "fortgefuehrt. Information der Betroffenen gemaess Art. 34 DSGVO bei hohem Risiko."),
        ("Lessons Learned und Praeventivmassnahmen",
         "Im Rahmen der Lessons-Learned-Analyse wurden folgende Praeventivmassnahmen identifiziert: "
         "(1) Verschaerfung der Awareness-Massnahmen (E-Learning Update); (2) ggf. Update der "
         "technisch-organisatorischen Massnahmen (Verschluesselung, MFA, Berechtigungskonzepte); "
         "(3) Aufnahme des Vorfallstyps in den jaehrlichen Datenpannen-Report an den Pruefungsausschuss; "
         "(4) Re-Test der entsprechenden Kontrollen durch Konzernrevision."),
        ("Status",
         f"Vorgang dokumentiert und im internen Datenpannen-Register ({jahr}-{lfd}) abgelegt. "
         f"Status: abgeschlossen / Massnahmen umgesetzt. Naechste Pruefung im Rahmen des Jahresberichts DSB an Vorstand."),
        ("Unterschriften",
         signatures(DSB, "Datenschutzbeauftragter (extern)", "Lehmann & Partner",
                    "Laura Bauer", "CFO / Verantwortlicher", R["name"],
                    place="Stuttgart", date_str_=f"15. Maerz {jahr}")),
    ]
    write_doc(f"{BASE}/{fname}", H,
              f"DSGVO-Datenpannen-Meldung – {ftitle}",
              subtitle=f"Aktenzeichen DSGVO-{jahr}-{lfd}",
              sections=sections, confidential=True)


# =============================================================================
# 7) TPDD Third-Party Due Diligence
# =============================================================================
TPDD_PARTNER = [
    ("Shanghai Aftermarket Dist. Co. Ltd.", "Shanghai (CN)", "Vertrieb Aftermarket Asien (RCN)",
     "Vertrieb",
     "Mittel (Geografie CN, Risiko Sanktionsdiversionsrisiko, ABC-Risiko bei Verzollung)"),
    ("Istanbul Trade Group A.S.", "Istanbul (TR)", "Vertriebs- und Service-Partner Tuerkei/MEA",
     "Vertrieb",
     "Mittel-Hoch (TR-Sanktionsexposure Russland, USD-Zahlungen)"),
    ("Mumbai EMS Services Pvt. Ltd.", "Mumbai (IN)", "Vertriebs- und Engineering-Partner Indien",
     "Vertrieb",
     "Mittel (Schwellenmarkt, lokale Genehmigungspraxis)"),
    ("Bratislava Industrial Brokers s.r.o.", "Bratislava (SK)", "Vermittler EU-Aftermarket",
     "Vermittler",
     "Niedrig (EU-Mitglied, geringe Geografie-Risiken)"),
    ("Detroit Aftermarket Consulting LLC", "Detroit (US)", "Beratungspartner US-OEM-Geschaeft",
     "Berater",
     "Mittel (US-Antitrust-/FCPA-Pflichten)"),
    ("Beijing Procurement Advisors Co.", "Beijing (CN)", "Lieferanten-Suche/-Qualifizierung China",
     "Berater",
     "Mittel-Hoch (Naehe zu Staatsunternehmen, FCPA-/UKBA-Risiko)"),
    ("Warsaw Logistics Brokers Sp. z o.o.", "Warschau (PL)", "Logistik-Brokerage Mittel-Osteuropa",
     "Logistik",
     "Niedrig-Mittel (PL/EU, Standardrisiken)"),
    ("Budapest IT Consulting Kft.", "Budapest (HU)", "IT-Beratung HU-Werk",
     "IT-Berater",
     "Niedrig (EU/HU, Standardrisiken)"),
    ("Prag Logistics Services s.r.o.", "Prag (CZ)", "Logistik-Dienstleister Werk RCZ",
     "Logistik",
     "Niedrig (EU/CZ, Standardrisiken)"),
    ("Singapore Trading Hub Pte. Ltd.", "Singapur", "Asien-Pazifik-Trading-Hub Aftermarket",
     "Vertrieb",
     "Mittel (SG ist niedriges Korruptionsrisikoland, jedoch Hub fuer Diversion)"),
    ("Sao Paulo OEM Liaison Ltda.", "Sao Paulo (BR)", "OEM-Liaison Stellantis Brasilien",
     "Vermittler",
     "Mittel (BR-Lavajato-Hintergrund, FCPA-Praezedenz)"),
    ("Seoul Engineering Partners Inc.", "Seoul (KR)", "Engineering-Partner Hyundai-Programm",
     "Engineering",
     "Niedrig-Mittel (KR, etablierte Compliance-Standards)"),
    ("Cape Town Distribution Pty.", "Kapstadt (ZA)", "Aftermarket Afrika",
     "Vertrieb",
     "Mittel (ZA, BBBEE-Vorgaben, FCPA-Praezedenz Nachbarmaerkte)"),
    ("Madrid Industrial Reps S.L.", "Madrid (ES)", "Vertrieb Iberische Halbinsel",
     "Vertrieb",
     "Niedrig (EU/ES, Standardrisiken)"),
    ("Brussels EU Government Affairs SPRL", "Bruessel (BE)", "EU-Politikberatung",
     "GovAffairs",
     "Mittel (Public-Sector-Naehe, Lobbying-Register-Pflicht)"),
    ("Tokyo Tier1 Liaison K.K.", "Tokio (JP)", "Liaison japanische OEMs/Tier-1",
     "Vermittler",
     "Niedrig-Mittel (JP, JFTC, Standardrisiken)"),
    ("Bucharest Sales Reps S.R.L.", "Bukarest (RO)", "Vertrieb Suedosteuropa",
     "Vertrieb",
     "Mittel (RO, GRECO-Korruptionsindex)"),
    ("Munich Patent Brokers GmbH", "Muenchen (DE)", "Patent-/IP-Vermittler",
     "IP-Berater",
     "Niedrig (DE, etablierte Standards)"),
    ("Krakow Trade Services Sp. z o.o.", "Krakau (PL)", "Trade-Services Mittel-Osteuropa",
     "Vertrieb",
     "Niedrig-Mittel (EU/PL, Standardrisiken)"),
    ("Shenzhen Components Trading Co.", "Shenzhen (CN)", "Halbleiter-/Komponenten-Trading",
     "Vertrieb",
     "Mittel-Hoch (CN-Halbleiter-Embargo-Risiken)"),
]

import os
for fname in sorted(os.listdir(BASE)):
    if not fname.startswith("TPDD_Third_Party_"):
        continue
    m = re.match(r"TPDD_Third_Party_(\d{4})_(\d{3})\.docx", fname)
    if not m:
        continue
    jahr, lfd = m.group(1), m.group(2)
    idx = int(lfd) % len(TPDD_PARTNER)
    partner_name, sitz, scope, kategorie, risiko = TPDD_PARTNER[idx]
    sections = [
        ("Aktenzeichen und Auftrag",
         f"Aktenzeichen: TPDD-{jahr}-{lfd}. Auftrag durch Konzern-Compliance ({CO}) im Rahmen des "
         f"Third-Party-Due-Diligence-Verfahrens (TPDD). Zweck: Risikobewertung vor Aufnahme einer "
         f"Geschaeftsbeziehung mit dem genannten externen Partner gemaess Konzern-Anti-Korruptions-"
         f"Richtlinie (Ziff. § 5 Vertriebsagenten, Berater, Joint Ventures) sowie LkSG-Anforderungen."),
        ("Geschaeftspartner",
         f"Name: {partner_name}\n\nSitz: {sitz}\n\nGeplanter Geschaeftszweck: {scope}\n\n"
         f"Kategorie: {kategorie}"),
        ("Pruefumfang und Datenquellen", ("list", [
            "Handelsregisterauszug / Konzernstrukturauswertung (Local Counsel / Dun & Bradstreet)",
            "ComplyCube-Sanktions-/PEP-/Adverse-Media-Screening",
            "Beneficial-Ownership-Auswertung (UBO, Ultimate Beneficial Owner)",
            "Reputations-Recherche (Negativberichterstattung, Strafverfahren)",
            "ABC-Questionnaire (selbstauskunft des Partners)",
            "Compliance-Audit-Berichte (sofern vom Partner verfuegbar)",
        ])),
        ("Befunde", ("list", [
            f"Identitaet und Konzernstruktur des Partners ({partner_name}) ueber Handelsregister bestaetigt",
            "Sanktionslisten-Screening EU/UN/OFAC/HMT: keine Treffer (zum Stichtag)",
            "Adverse Media: keine substantiierten negativen Meldungen identifiziert",
            "UBO offengelegt; Auswirkungen auf Sanktions-/PEP-Status keine",
            "ABC-Questionnaire vollstaendig ausgefuellt; keine kritischen Selbstangaben",
            "Branchenuebliche Reputations-Recherche unauffaellig",
        ])),
        ("Risikoeinstufung",
         f"Gesamtrisiko: {risiko}. Empfehlung: Aufnahme der Geschaeftsbeziehung freigegeben unter Auflagen "
         f"(Compliance-Klausel im Vertrag verbindlich; Provision/Honorar marktgerecht und Erfolgsnachweis-"
         f"gebunden; Audit-Right zugunsten Brennhagen; Termination-on-Breach-Klausel; Schulung des Partners "
         f"zu Anti-Korruption; Re-Screening jaehrlich)."),
        ("Vertragliche Auflagen", ("clauses", [
            ("§ 1 Compliance-Klausel", [
                "Im Vertrag mit dem Partner sind die Konzern-Standardklauseln zu Anti-Korruption, Sanktionen, "
                "LkSG-Sorgfaltspflichten und Datenschutz aufzunehmen.",
            ]),
            ("§ 2 Audit-Right", [
                "Recht von Brennhagen zur einseitigen Pruefung (Books-and-Records-Klausel) mind. einmal pro Jahr "
                "fuer die Vertragsdauer plus 3 Jahre.",
            ]),
            ("§ 3 Re-Screening und Monitoring", [
                "Jaehrliches Re-Screening durch ComplyCube; Anlassbezogenes Re-Screening bei wesentlichen "
                "Aenderungen (UBO, Adverse Media, Sanktionslistenaenderungen).",
            ]),
        ])),
        ("Freigabe",
         f"Empfehlung Compliance: Freigabe mit Auflagen. Eskalation an CO {CO}: erfolgt am {jahr}-XX-XX. "
         f"Final-Freigabe durch CFO Laura Bauer fuer Verguetungs-Cap >= 100 kEUR p. a. erforderlich. "
         f"Re-TPDD: 12 Monate."),
        ("Unterschriften",
         signatures(CO, "Compliance Officer (extern)", "Hengeler Mueller",
                    "Laura Bauer", "CFO", R["name"],
                    place="Stuttgart", date_str_=f"30. Juni {jahr}")),
    ]
    write_doc(f"{BASE}/{fname}", H,
              f"Third-Party-Due-Diligence – {partner_name}",
              subtitle=f"TPDD-Akte {jahr}-{lfd}",
              sections=sections, confidential=True)


# =============================================================================
# 8) RECHTSAKTEN (Schriftsaetze, Urteile, Vergleiche, Widersprueche, Mahnbescheide, Behoerdenschreiben)
# =============================================================================
RECHTSAKTEN_FAELLE = [
    ("Lieferanten-Reklamation Steckverbinder GmbH",
     "Lieferantenreklamation 8D-Fall, fehlerhafte Charge LL-2022-014 (Steckverbinder); "
     "Streitwert 187.500 EUR; Klaegerin Brennhagen Elektronik GmbH (REG); Beklagte Steckverbinder Mueller GmbH.",
     "LG Stuttgart, Aktenzeichen 12 O 145/22", "Hengeler Mueller (Frankfurt)"),
    ("Patentstreit ICP-3-Schnittstelle",
     "Patentverletzungsklage; Beklagte Brennhagen Elektronik AG; Klaegerin Tech-Patents Holdings LLC (USA); "
     "behauptete Verletzung EP 2 845 678 (»Audio-Bus-Synchronisationsverfahren«); Brennhagen macht "
     "Nichtigkeitseinrede geltend.",
     "OLG Duesseldorf, Aktenzeichen I-2 U 78/23", "Hengeler Mueller (Patent: Dr. Christoph Voelker)"),
    ("Arbeitsrechtliches Kuendigungsverfahren",
     "Kuendigungsschutzklage gegen REG Heilbronn; Klaegerin: ehem. Mitarbeiterin (Vertrieb, "
     "betriebsbedingte Kuendigung im Rahmen Restrukturierung 2022); Streitwert 32.500 EUR.",
     "ArbG Heilbronn, Aktenzeichen 8 Ca 215/22", "Lokale Kanzlei: Mueller & Partner Heilbronn"),
    ("Verwaltungsstreit BImSchG-Genehmigung",
     "Streit ueber Naebenbestimmungen einer BImSchG-Genehmigung Werk REG Heilbronn (Loesemittelemissionen); "
     "Klaeger Brennhagen Elektronik GmbH; Beklagte Regierungspraesidium Stuttgart.",
     "VG Stuttgart, Aktenzeichen 11 K 4520/23", "Hengeler Mueller (Umweltrecht: Dr. Stefan Wolfsberger)"),
    ("Vergleich Vertriebsagent Tuerkei",
     "Vergleich ueber ausstehende Provisionen; Vertriebsagent Istanbul Trade Group A.S.; "
     "Brennhagen zahlt EUR 95.000 final, ohne Anerkenntnis einer Rechtspflicht; Geheimhaltung; "
     "Generalquittung. Compliance-Pruefung erfolgt.",
     "Schiedsgericht ICC Paris, Fall Nr. 25.834/MM", "Hengeler Mueller"),
    ("Mahnbescheid Bestand-Lieferant",
     "Mahnbescheid gegen Lieferant Display-Tech GmbH wegen verspaeteter Lieferung und Konventionalstrafen; "
     "Hauptforderung 124.300 EUR zzgl. Zinsen; Mahnbescheid am 12.7.2022 erlassen.",
     "AG Stuttgart Mahngericht", "Inhouse Legal"),
    ("Behoerdenschreiben BAFA Dual-Use",
     "Schreiben BAFA Eschborn: Anfrage zu Endverwendung ECU-900 mit Lieferung nach Vietnam; "
     "Brennhagen hat Antrag auf Ausfuhrgenehmigung mit End-User-Statement und Endverbleibserklaerung "
     "eingereicht; Klassifizierungsfrage Annex I Kat. 3A.",
     "BAFA Aktenzeichen 81-DL-2022/4521", "Group Tax / BAFA-Beauftragte " + BAFA),
    ("Urteilsauszug Lieferantenstreit",
     "Urteil LG Stuttgart vom 14.3.2023, Az. 12 O 145/22: Beklagte Steckverbinder Mueller GmbH wird "
     "verurteilt, an Klaegerin REG EUR 145.700 nebst Zinsen 5 %-Punkte ueber BZ ab 15.5.2022 zu zahlen; "
     "Berufung der Beklagten anhaengig OLG Stuttgart.",
     "LG Stuttgart, Az. 12 O 145/22; Berufung OLG Stuttgart Az. 7 U 92/23", "Hengeler Mueller"),
    ("Widerspruch Steuerbescheid",
     "Einspruch gegen Koerperschaftsteuer-Bescheid 2020 (Finanzamt Stuttgart-Mitte); strittig: "
     "Verrechnungspreis-Korrektur IC-Darlehen RHO-REG, Mehrsteuer 850 kEUR; Einspruch eingelegt 8.9.2022.",
     "FA Stuttgart-Mitte, Steuernummer 99012/12345", "Hengeler Mueller (Steuer: Dr. Tina Kraft)"),
    ("Vergleich D&O-Schaden",
     "Vergleich ueber D&O-Versicherungsleistung Allianz GCS in Hoehe von 2,1 Mio. EUR zur Abgeltung "
     "etwaiger Pflichtverletzungs-Ansprueche gegen ein ehemaliges Vorstandsmitglied; "
     "Geheimhaltung; Verzicht der Gesellschaft auf weitergehende Ansprueche.",
     "Allianz Schadennummer DO-2022-451", "Hengeler Mueller (D&O: Dr. Schramm)"),
]

import os
for fname in sorted(os.listdir(BASE)):
    if not fname.startswith("Rechtsakte_"):
        continue
    m = re.match(r"Rechtsakte_(\d{4})_(\d{3})_(Schriftsatz|Urteil_Auszug|Vergleich|Widerspruch|Mahnbescheid|Behoerdenschreiben)(?:_.+)?\.docx", fname)
    if not m:
        continue
    jahr, lfd, typ = m.group(1), m.group(2), m.group(3)
    fall = RECHTSAKTEN_FAELLE[int(lfd) % len(RECHTSAKTEN_FAELLE)]
    titel_fall, sachverhalt, az, vertreter = fall
    typ_display = typ.replace("_", " ")
    sections = [
        ("Sache / Aktenzeichen",
         f"Fall: {titel_fall}\n\nAktenzeichen: {az}\n\nAktentyp: {typ_display} ({jahr})\n\n"
         f"Vertretung Brennhagen: {vertreter}\n\nKonzern-Compliance: {CO} (Hengeler Mueller)"),
        ("Sachverhalt", sachverhalt),
        ("Rechtliche Wuerdigung",
         f"Die rechtliche Wuerdigung dieses Falles erfolgt durch die mandatierte Kanzlei ({vertreter}) "
         f"im Einklang mit den einschlaegigen Vorschriften des materiellen und prozessualen Rechts. "
         f"Im Wesentlichen wird auf die Schluessigkeit der Anspruchsgrundlagen, Verjaehrungsfragen, "
         f"einschlaegige Rechtsprechung der Berufungsgerichte sowie die Beweissituation abgestellt. "
         f"Bei {typ_display}-Aktentypen gelten die jeweiligen formellen Voraussetzungen (Schriftform, "
         f"Fristen, Begruendungspflichten)."),
        ("Aktueller Verfahrensstand",
         f"Stand zum Aktendatum: {typ_display} vom {jahr}; weiteres Verfahren wie folgt: "
         f"Sofern Schriftsatz, wird auf den naechsten Termin / die Erwiderungsfrist verwiesen. "
         f"Sofern Urteil, wird die Rechtsmittelfrist und ggf. Berufungsentscheidung benannt. "
         f"Sofern Vergleich, gilt die Abgeltungswirkung und Geheimhaltung. Sofern Widerspruch, wird "
         f"auf die Bescheidung durch die Behoerde gewartet. Sofern Mahnbescheid, ist Widerspruchsfrist "
         f"zu beachten. Sofern Behoerdenschreiben, ist die Antwort fristgerecht erfolgt."),
        ("Risikobewertung / Rueckstellung (IAS 37)",
         f"Die Rueckstellungs- und Reserven-Hoehe ist im Einklang mit IAS 37 (Provisions, Contingent "
         f"Liabilities and Contingent Assets) zu beurteilen. Erfolgsaussicht und Schadens-Cap werden "
         f"jaehrlich (bzw. quartalsweise bei wesentlichen Aenderungen) im Litigation-Tracker "
         f"(»REA_Rechtsstreitigkeitsregister_2024.xlsx«) erfasst und mit der Kanzlei abgestimmt."),
        ("Empfehlung und naechste Schritte",
         "Empfehlungen der mandatierten Kanzlei werden in Abstimmung mit dem Konzern-Legal "
         "(General Counsel / Compliance Officer) umgesetzt. Eskalation an Vorstand (CFO) erfolgt bei "
         "Streitwerten / Ergebniswirkungen > 250 kEUR oder reputativer Relevanz."),
        ("Vertraulichkeit",
         "Diese Akte unterliegt der anwaltlichen Verschwiegenheit (§ 43a BRAO) sowie dem "
         "Mandatsgeheimnis. Weitergabe ausschliesslich an autorisierte interne Personen sowie an "
         "den externen Wirtschaftspruefer (KPMG) im Rahmen der Litigation-Letter-Anfrage."),
        ("Unterschrift",
         signatures(vertreter.split(" (")[0], "mandatierte Rechtsvertretung", "Kanzlei",
                    CO, "Compliance Officer (extern)", "Hengeler Mueller",
                    place="Frankfurt", date_str_=f"{jahr}")),
    ]
    write_doc(f"{BASE}/{fname}", H,
              f"Rechtsakte {typ_display} – {titel_fall}",
              subtitle=f"Aktenzeichen {az} – {jahr}",
              sections=sections, confidential=True)


# =============================================================================
# 9) SUBSIDIARY ROUTINE COMPLIANCE DOCS
# (Arbeitssicherheit, Handelskammer, Jahresabschlusshinterlegung, Lohnsteuer, Sozialversicherung, Umweltbericht)
# =============================================================================
SUB_ROUTINE_MAP = {
    "DE": ("REG", "Heilbronn", "Deutschland"),
    "PL": ("RPL", "Katowice", "Polen"),
    "CZ": ("RCZ", "Brno", "Tschechien"),
    "HU": ("RHU", "Gyoer", "Ungarn"),
    "CN": ("RCN", "Shanghai", "VR China"),
}


def sub_routine(fname, country_code, sub_short, sub_city, sub_country, jahr, typ):
    sub = S[sub_short]
    body_by_type = {
        "Arbeitssicherheitsbericht": (
            "Jahresbericht Arbeitssicherheit",
            f"Dieser Bericht der {sub['name']} stellt die Arbeitssicherheitsleistung am Standort {sub_city} "
            f"fuer das Geschaeftsjahr {jahr} dar. Rechtsgrundlagen: ArbSchG (DE) bzw. lokale Aequivalente "
            f"(BHP-Gesetz PL, Zakonik prace CZ, Munka Toervenykoenyve HU, Arbeitsschutzgesetz CN); "
            f"DGUV-Vorschriften (DE) sowie ISO 45001:2018-Orientierung.",
            [
                "Sicherheits-Organisation: Fachkraft fuer Arbeitssicherheit (SiFa) etabliert; Sicherheitsausschuss vierteljaehrlich",
                f"Mitarbeiterstand: {sub['employees']} FTE",
                "Pflichtunterweisungen: 100 %-Coverage (jaehrliche Sicherheitsunterweisung, Erste-Hilfe-Kurse)",
                "Gefaehrdungsbeurteilungen aktuell; Update bei Aenderungen Arbeitsmittel/Arbeitsplaetze",
                "Persoenliche Schutzausruestung (PSA) bedarfsgerecht zur Verfuegung gestellt",
                "Notfallplaene und Evakuierungsuebungen halbjaehrlich",
            ],
            [
                "TRIR (Total Recordable Incident Rate): unter Branchen-Benchmark VDA/CLEPA",
                "LTIR (Lost Time Incident Rate): rueckl. ggue. Vorjahr",
                "Meldepflichtige Unfaelle: <5 (geringfuegig, kein toedlicher Vorfall)",
                "Beinahe-Unfaelle erfasst und ausgewertet (Lessons-Learned-Loop)",
                "Audits durch Berufsgenossenschaft / lokale Aufsichtsbehoerde: keine wesentlichen Beanstandungen",
            ],
        ),
        "Handelskammer_Bericht": (
            "Bericht an die Handelskammer",
            f"Pflichtbericht der {sub['name']} an die zustaendige Handelskammer ({sub_country}) fuer das "
            f"Geschaeftsjahr {jahr}. Erstattet auf Grundlage der landesueblichen IHK-/Handelskammer-Mitglieds"
            f"pflichten (DE: IHK Heilbronn-Franken; PL: KIG; CZ: HK CR; HU: MKIK; CN: CCPIT Shanghai). "
            f"Zweck: Mitgliederbestandspflege, Statistik, Konjunkturberichterstattung.",
            [
                f"Gesellschaftsdaten: {sub['hrb']}",
                f"Mitarbeiter: {sub['employees']} FTE; Veraenderung ggue. Vorjahr",
                f"Umsatz {jahr}: ca. {sub['revenue_mio']} Mio. EUR",
                f"Investitionen {jahr}: Maschinen, IT, Werks-Modernisierung",
                "Marktentwicklung: Automotive-Tier-1-Geschaeft; Hauptmaerkte EU, NAFTA, APAC",
                "Ausbildung/Qualifizierung: Berufsausbildung (DE: IHK-Lehrberufe; PL: szkoly branzowe; CZ: ucni; HU: szakkepzes; CN: vocational)",
            ],
            [
                "Branchenkonjunktur: Stabilisierung nach Halbleiter-/Lieferketten-Engpaessen",
                "Investitionsklima: Brennhagen investiert kontinuierlich in Automatisierung und Energieeffizienz",
                "Personalentwicklung: Erweiterung F&E in Muenchen; stabiles Produktionsniveau in Heilbronn/Katowice/Brno/Gyoer",
                "Aussenhandel: stark EU-getrieben; China und USA als Wachstumsmaerkte",
                "Mitgliedsbeitrag: gemaess Beitragsordnung lokal entrichtet",
            ],
        ),
        "Jahresabschlusshinterlegung": (
            "Jahresabschluss-Hinterlegung (lokal)",
            f"Pflichtenmaessige Hinterlegung des Jahresabschlusses der {sub['name']} fuer das Geschaeftsjahr "
            f"{jahr} bei der zustaendigen Register-/Veroeffentlichungsstelle ({sub_country}): "
            f"DE: elektronischer Bundesanzeiger / Unternehmensregister; PL: KRS-Repozytorium Dokumentow Finansowych; "
            f"CZ: Sbirka listin Obchodniho rejstriku; HU: Igazsagugyi Miniszterium (e-beszamolo); "
            f"CN: National Enterprise Credit Information Publicity System.",
            [
                f"Jahresabschluss {jahr} bestehend aus: Bilanz, Gewinn-/Verlustrechnung, Anhang, Lagebericht",
                f"Aufstellung gemaess lokalen GAAP; Konsolidierung gemaess IFRS auf Konzernebene (REA)",
                "Wirtschaftspruefer (lokal): bei groesseren Toechtern PwC bzw. KPMG-Netzwerk; Konzern KPMG",
                "Hinterlegung fristgerecht innerhalb gesetzlicher Frist",
                "Aufbewahrungsfristen lokale Buchfuehrungspflichten (DE 10 J. § 257 HGB, PL 5 J., CZ 10 J., HU 8 J., CN 10 J.)",
            ],
            [
                f"Bilanzsumme {jahr}: marktueblich gegenueber Vorjahr",
                f"Umsatzerloese {jahr}: ca. {sub['revenue_mio']} Mio. EUR",
                "Ergebnis nach Steuern: positiv im Korridor der Konzernerwartung",
                "Eigenkapitalquote: solide ueber 30 %",
                "Konzern-Verflechtungen: ueber RHO Holding (100 %); IC-Saldenausweis im Anhang",
            ],
        ),
        "Lohnsteuer_Jahresabschluss": (
            "Lohnsteuer-Jahresabschluss",
            f"Lohnsteuer-Jahresabschluss der {sub['name']} ({sub_country}) fuer das Lohnsteuer-Jahr {jahr}. "
            f"Berechnung und Anmeldung nach den jeweiligen lokalen Bestimmungen "
            f"(DE: § 41a EStG / § 41b EStG / § 42b EStG; PL: PIT-4R, PIT-8AR; CZ: Zakon o danich z prijmu; "
            f"HU: SZJA-Gesetz; CN: IIT-Bestimmungen).",
            [
                f"Anzahl Beschaeftigte im Lohnsteuer-Jahr {jahr}: {sub['employees']} FTE",
                "Lohnsteuer/IIT-Abfuehrung monatlich termingerecht",
                "Sozialversicherungsbeitraege getrennt erfasst (siehe Sozialversicherung-Meldung)",
                "Jahresbescheinigung an Beschaeftigte ausgegeben",
                "Lohnsteuer-Anmeldungen elektronisch eingereicht",
                "Lokale Lohnsteuer-Beauftragte/r: HR/Payroll-Funktion am Standort",
            ],
            [
                "Lohnsumme brutto gesamt: konsistent mit HR-Buch",
                "Lohnsteuer/IIT: ordnungsgemaess einbehalten und abgefuehrt",
                "Spezielle Bezuege (LTI, Bonuspoolausschuettung): korrekt versteuert",
                "Auslandsentsendungen: Doppelbesteuerung-Abkommen DBA beachtet",
                "Keine Beanstandungen lokaler Steuerbehoerden festgestellt",
            ],
        ),
        "Sozialversicherung_Meldung": (
            "Sozialversicherung – Jahresmeldung",
            f"Jahresmeldung Sozialversicherung der {sub['name']} ({sub_country}) fuer das Beitragsjahr {jahr}. "
            f"Erstattet gemaess lokalen Sozialversicherungsgesetzen "
            f"(DE: DEUEV, SGB IV; PL: ZUS; CZ: CSSZ; HU: NEAK/NAV; CN: Social-Insurance-Law).",
            [
                "Renten-, Kranken-, Arbeitslosen-, Pflege-, Unfallversicherung (DE); bzw. ZUS-Trager (PL); CSSZ + VZP (CZ); NAV/NEAK (HU); Wujian Yi Jin (CN)",
                f"Anzahl gemeldeter Beschaeftigter: {sub['employees']} FTE",
                "An- und Abmeldungen elektronisch zeitgerecht erfolgt (DEUEV-elektronische Meldung; ZUS; CSSZ etc.)",
                "Beitragsabfuehrung monatlich; Saeumniszuschlaege keine",
                "Korrekturmeldungen unterjaehrig bei Statusaenderungen erfolgt",
            ],
            [
                "Beitragssumme gesamt: konsistent mit Lohnbuchhaltung",
                "Pruefungen Rentenversicherung / lokale SV-Behoerde: keine wesentlichen Beanstandungen",
                "Betriebliche Altersversorgung (bAV): separate Meldungen erfolgt; Heubeck-Aktuargutachten Konzernebene",
                "Auslandseinsaetze: A1-Bescheinigungen / E101 / lokale Aequivalente vorhanden",
                "Datensicherheit (DSGVO) bei Lohnabrechnungs-Dienstleister (ADP/Workday) gewaehrleistet",
            ],
        ),
        "Umweltbericht": (
            "Umweltbericht (lokal)",
            f"Umweltbericht der {sub['name']} fuer das Berichtsjahr {jahr}. Erstellt auf Basis der internen "
            f"ISO-14001-/50001-Managementsysteme sowie der lokalen umweltrechtlichen Berichtspflichten "
            f"(DE: BImSchG / KrWG / WHG; PL: PRTR; CZ: Zakon o ochrane ovzdusi; HU: koernyezetvedelmi torveny; "
            f"CN: HJ-Standards). Verantwortlich: Werks-EHS-Beauftragte/r in Abstimmung mit Konzern-CSO {CSO}.",
            [
                "ISO 14001:2015 (Umwelt) und ISO 50001:2018 (Energie) zertifiziert",
                "Scope-1- und Scope-2-Emissionen (CO2e) erfasst und an Konzern-CSRD-Reporting geliefert",
                "Wasserverbrauch / Abwasser / Abfall nach Fraktion getrennt erfasst",
                "Genehmigungen (BImSchG/aequivalent) gueltig; Auflagen eingehalten",
                "Energieeffizienz-Massnahmen (LED, PV-Eigenversorgung, Druckluft-Leckage-Management)",
                "Schulungen Umweltbewusstsein jaehrlich (Coverage 100 %)",
            ],
            [
                f"Energieverbrauch {jahr}: ggue. Vorjahr reduziert (Effizienzpipeline)",
                f"Scope-1-Emissionen: -3-5 %; Scope-2 (location-based) -5-8 %",
                "Erneuerbare Stromanteile: kontinuierlich ausgebaut (Sustainability-Linked-RCF-Margin)",
                "Abfallquoten: Wiederverwertung > 85 %",
                "Umweltvorfaelle: keine meldepflichtigen Ereignisse",
            ],
        ),
    }
    typ_data = body_by_type[typ]
    titel = typ_data[0]
    sections = [
        ("Berichtsgegenstand", typ_data[1]),
        ("Status / Pflichten", ("list", typ_data[2])),
        ("Kennzahlen und Befunde", ("list", typ_data[3])),
        ("Konzern-Bezug",
         f"Die {sub['name']} ist 100 %ige Tochter der RHO Brennhagen Holding GmbH (Stuttgart, Konzernmutter "
         f"Brennhagen Elektronik AG, Vaihinger Strasse 120, 70567 Stuttgart). Berichterstattung an "
         f"Konzernfunktionen (Group HR/Tax/CSO/Compliance/Group Audit) erfolgt im Rahmen der "
         f"konzernweiten Reporting-Prozesse. Externe Pruefung im Rahmen Konzernabschluss durch KPMG."),
        ("Erklaerung",
         f"Die Werkleitung der {sub['name']} bestaetigt nach bestem Wissen die Vollstaendigkeit und "
         f"Richtigkeit der vorstehenden Angaben. Alle gesetzlichen Berichtspflichten {jahr} wurden "
         f"fristgerecht und ordnungsgemaess erfuellt."),
        ("Unterschriften",
         signatures(f"Werksleitung {sub_short}", "Geschaeftsfuehrung", sub['name'],
                    CO if typ != "Umweltbericht" else CSO,
                    "Compliance / CSO Konzern", R["name"],
                    place=sub_city, date_str_=f"31. Maerz {int(jahr)+1}")),
    ]
    write_doc(f"{BASE}/{fname}", H,
              f"{titel} – {sub['name']} {jahr}",
              subtitle=f"Berichtsjahr {jahr} – {sub_country}",
              sections=sections)


import os
sub_pattern = re.compile(r"(DE|PL|CZ|HU|CN)_(REG|RPL|RCZ|RHU|RCN)_"
                         r"(Arbeitssicherheitsbericht|Handelskammer_Bericht|Jahresabschlusshinterlegung|"
                         r"Lohnsteuer_Jahresabschluss|Sozialversicherung_Meldung|Umweltbericht)_(\d{4})"
                         r"(?:_(?:WIP|ENTWURF|ALT|FINAL|DRAFT))?\.docx")
for fname in sorted(os.listdir(BASE)):
    m = sub_pattern.match(fname)
    if not m:
        continue
    cc, sub_short, typ, jahr = m.group(1), m.group(2), m.group(3), m.group(4)
    _, sub_city, sub_country = SUB_ROUTINE_MAP[cc]
    sub_routine(fname, cc, sub_short, sub_city, sub_country, jahr, typ)


# =============================================================================
# 10) MISCELLANEOUS DOCS
# =============================================================================
# REA_STE_BMS-12_Nomination_Letter_2022.docx
write_doc(f"{BASE}/REA_STE_BMS-12_Nomination_Letter_2022.docx", H,
    "Nomination Letter STE / BMS-12 – Volkswagen AG (Plattform ID.7)",
    subtitle="Single-Tier-Engagement / Long-Term Agreement (LTA)",
    sections=[
        ("Vorgang",
         "Mit Schreiben vom 12.10.2022 hat die Volkswagen AG, Wolfsburg, die Brennhagen Elektronik AG / Brennhagen "
         "Elektronik GmbH (REG) als nominierten Lieferanten fuer das Batteriemanagementsystem BMS-12 in der "
         "Plattform ID.7 ausgewaehlt (SOP geplant Q3/2023). Diese Single-Tier-Engagement-Bestaetigung (STE) "
         "ist Grundlage fuer den anschliessenden Long-Term Agreement (LTA) mit 5 Jahren Laufzeit "
         "(1.7.2023 – 30.6.2028, kuendbar gemaess VW-Konzern-AGB)."),
        ("Inhalt der Nomination", [
            ["Parameter", "Wert"],
            ["Produkt", "BMS-12 (BatteryMS-12)"],
            ["Plattform", "Volkswagen MEB+ / ID.7"],
            ["SOP", "Q3 2023"],
            ["Programm-Volumen", "ca. 350.000 Einheiten p. a."],
            ["Programm-Laufzeit", "5 Jahre + 2 Optionsjahre"],
            ["Standort Produktion", "REG Heilbronn (Hauptlinie) + RPL Katowice (SMD)"],
            ["LCC", "Lifetime ca. 60 Mio. Einheiten"],
            ["Pricing", "Annual Productivity Index 3,5 % p. a. (VDA-Standard)"],
        ]),
        ("Compliance-Erklaerungen",
         "Die Brennhagen-Gruppe bestaetigt die Einhaltung saemtlicher VW-Konzern-Lieferanten-Compliance-Anforderungen, "
         "insbesondere: VW-Konzern-Nachhaltigkeitsrating-Module (Sustainability Rating-S-Rating), VW-Code-of-"
         "Conduct fuer Geschaeftspartner, IATF-16949-Zertifizierung (alle relevanten Werke), ASPICE-Level-2-3-"
         "Reifegrad RSG, ISO-14001 und ISO-50001 (Werke), Anti-Korruption gemaess UK Bribery Act / FCPA / "
         "§§ 299, 333 StGB, Sanktionen-/Exportkontroll-Konformitaet und LkSG-Anforderungen (ab 1.1.2024 "
         "berichtspflichtig)."),
        ("Naechste Schritte",
         "Verhandlung und Abschluss des LTA-Hauptvertrags bis 30.6.2023; Vereinbarung der jaehrlichen "
         "Productivity-Index-Klausel, Rohstoff-Pass-Through-Klausel (insb. fuer Halbleiter/Magnetwerkstoffe) "
         "und Sustainability-Linked-Klauseln (Scope 1-3 Emissionen). Parallel laufende PPAP-Vorbereitung "
         "und Werkzeugbau in Heilbronn."),
        ("Unterschriften",
         signatures("Stefan Richter", "CMO / Business Development", R["name"],
                    "Andreas Maier", "Werkleiter REG Heilbronn", "Brennhagen Elektronik GmbH",
                    place="Stuttgart", date_str_="15. Oktober 2022")),
    ])


# RSG_to_RCZ_IC_2023_01
write_doc(f"{BASE}/RSG_to_RCZ_IC_2023_01.docx", H,
    "Konzern-interne Leistungsverrechnung (IC-Charge) – RSG an RCZ – Q1/2023",
    subtitle="Engineering-Support-Leistungen RSG Muenchen fuer RCZ Brno",
    sections=[
        ("Vorgang",
         "RSG Brennhagen Software GmbH (Muenchen) erbringt im Rahmen der ADAS-V4D- und ECU-900-Plattformen "
         "Engineering-Support-Leistungen (Embedded Software Adaptionen, Test-Bench-Support, EMV-Begleitung) "
         "an RCZ Brennhagen CZ s.r.o. (Brno). Die Leistungen werden auf Cost-plus-Basis (Markup 8 %) gemaess "
         "der konzerninternen TP-Richtlinie und § 1 AStG verrechnet."),
        ("Leistungsdetails Q1/2023", [
            ["Leistung", "Stunden", "Stundensatz EUR", "Betrag EUR"],
            ["Embedded Software Engineering (ADAS-V4D)", "412", "115", "47.380"],
            ["Test-Bench Bring-Up Support (ECU-900)", "286", "110", "31.460"],
            ["EMV-Pre-Compliance Begleitung", "84", "125", "10.500"],
            ["Projekt-Management Schnittstelle", "62", "140", "8.680"],
            ["Summe Kosten", "844", "—", "98.020"],
            ["Markup 8 % (Cost-plus)", "", "", "7.842"],
            ["Verrechnungsbetrag netto", "", "", "105.862"],
        ]),
        ("Verrechnungspreis-Konformitaet",
         "Die Verrechnung erfolgt im Einklang mit OECD-TPG 2022 (Kap. VII, Kap. I), § 1 AStG (Fremdvergleich) "
         "und der konzerninternen TP-Richtlinie. Cost-plus 8 % wurde durch Benchmark-Studie der KPMG bestaetigt "
         "(Vergleichbarkeitsstudie Engineering-Services Automotive, 2022). Local File DE / CZ aktualisiert."),
        ("Verrechnung",
         "Rechnungsstellung RSG -> RCZ in Hoehe von 105.862 EUR netto zzgl. USt nach Reverse-Charge-Verfahren "
         "(EU-B2B-Leistung gemaess Art. 44 MwStSyst-RL / § 3a Abs. 2 UStG). Faelligkeit 60 Tage. SAP-Buchung "
         "Konzern-Eliminierung in HFM unter IC-Kennzeichen DE-RSG/CZ-RCZ-EngSrv."),
        ("Genehmigung und Unterschriften",
         signatures("Dr. Klaus Kessler", "Werkleiter RSG Muenchen", "Brennhagen Software GmbH",
                    "Petr Novak", "Werkleiter RCZ Brno", "Brennhagen CZ s.r.o.",
                    place="Muenchen / Brno", date_str_="5. April 2023")),
    ])


# ECO_ECU-900_005_Engineering_Change_2023
write_doc(f"{BASE}/ECO_ECU-900_005_Engineering_Change_2023.docx", H,
    "Engineering Change Order (ECO) – ECU-900 005 (2023)",
    subtitle="Aenderungsantrag mit Compliance- und Lieferketten-Auswirkungen",
    sections=[
        ("ECO-Daten und Geltungsbereich",
         "ECO-Nummer: ECU-900-005-2023. Antragsteller: Lars Wittmann (Lead Developer RSG Muenchen). "
         "Genehmigung: ECO-Board (CTO, Q-Leitung, Einkauf, Produktion REG, RSG). Geltung: alle laufenden "
         "ECU-900-Programme (VW MEB+, MQB-Evo; Stellantis CMP)."),
        ("Aenderungsbeschreibung",
         "Wechsel des Mikrocontroller-Lieferanten von Infineon (TC397) auf NXP S32G3 zur Erhoehung der "
         "Lieferketten-Resilienz (Second-Source-Strategie) und Verbesserung der Funktional-Safety-Reifegrad "
         "(ASIL-D Native Lockstep). Software-Stack-Adaption Embedded Linux + AUTOSAR Adaptive."),
        ("Compliance-Pruefungen", ("list", [
            "Funktionale Sicherheit ISO 26262: HAZOP-Re-Analyse erforderlich (Lead RSG Munich)",
            "Cybersecurity ISO 21434: Threat-Re-Assessment erforderlich",
            "ASPICE: System- und Software-Reifegrad re-evaluation Level 2-3",
            "Exportkontrolle: NXP S32G3 Pruefung Dual-Use Annex I (BAFA-Beauftragte Dr. Berger)",
            "PPAP: Re-PPAP-Level-3 erforderlich gegenueber VW/Stellantis",
            "Kosten-Gating: BOM-Delta +1,80 EUR/Stk (Genehmigung CFO)",
        ])),
        ("Lieferanten-/Vertrags-Auswirkungen",
         "Neuer NXP-Rahmenvertrag mit Allocation-Garantie 2024-2027 verhandelt (Volumen 12 Mio Stk Lifetime). "
         "Bestehender Infineon-Vertrag wird als Backup gepflegt (Last-Time-Buy-Option). Lieferanten-Compliance-"
         "Klauseln im neuen NXP-Vertrag entsprechen Konzern-Standard 2023 (Anti-Korruption, LkSG, Sanktionen, AuditRight)."),
        ("Genehmigung",
         "ECO genehmigt am 18. Juli 2023 durch ECO-Board mit den Auflagen: (1) ASIL-D HAZOP-Re-Analyse bis "
         "31.10.; (2) Re-PPAP gegenueber VW bis 15.12.; (3) BAFA-Pre-Check positiv erforderlich. "
         "Risiko-Reserven im Programmbudget gebildet."),
        ("Unterschriften",
         signatures("Stefan Hoffmann", "CTO (bis 30.6.2024)", R["name"],
                    "Lars Wittmann", "Lead Developer RSG", "Brennhagen Software GmbH",
                    place="Stuttgart", date_str_="18. Juli 2023")),
    ])


# Gehaltsaenderung_2023_006_Julia_Hartma
write_doc(f"{BASE}/Gehaltsaenderung_2023_006_Julia_Hartma.docx", H,
    "Gehaltsanpassung – Julia Hartmann (Compliance-Funktion)",
    subtitle="Personal-/HR-Akte – vertraulich",
    sections=[
        ("Mitarbeiterin und Position",
         "Mitarbeiterin: Julia Hartmann, geb. 14.3.1989, eingestellt 1.9.2018, Personalnummer 10245. "
         "Position: Senior Compliance Analyst, Group Compliance (Bereich Anti-Korruption / Sanktionsscreening). "
         "Berichtsweg: an Compliance Officer (extern) " + CO + " mit dotted-line an Group General Counsel."),
        ("Anlass der Anpassung",
         "Erweiterung der Verantwortlichkeiten um den Bereich Sanktionsscreening (ComplyCube-Lead) und "
         "Implementierungssteuerung der ISO 37001-Anti-Bribery-Standards. Marktgehaltsbenchmark Korn Ferry "
         "Compliance-Vergleichsstudie 2023 ergab Anpassungsbedarf zur Markt-Median."),
        ("Vereinbarte Aenderung", [
            ["Komponente", "Alt (2023)", "Neu (ab 1.10.2023)"],
            ["Jahres-Grundgehalt brutto", "82.500 EUR", "92.000 EUR"],
            ["Variable Verguetung Ziel", "12 % (9.900 EUR)", "15 % (13.800 EUR)"],
            ["Dienstwagen / Mobilitaetspauschale", "550 EUR/Mo", "650 EUR/Mo"],
            ["bAV (Direktversicherung)", "200 EUR/Mo", "300 EUR/Mo"],
            ["Urlaub", "30 Tage", "30 Tage"],
        ]),
        ("Vertragsbasis",
         "Aenderung erfolgt als Aenderungsvertrag zum bestehenden Arbeitsvertrag vom 1.9.2018 (in der Fassung "
         "des Aenderungsvertrages 2021). Alle uebrigen Vertragsbestimmungen bleiben unveraendert. Insbesondere "
         "gelten unveraendert: Verschwiegenheitsklauseln (§ 12), Wettbewerbsverbot waehrend des Arbeitsverhaeltnisses, "
         "Erfindungs- und Schutzrechtsklauseln, Datenschutzklauseln gemaess Art. 13 DSGVO."),
        ("Genehmigung",
         "HR-Buero: Erstellung durch HR Business Partner Group Compliance. Genehmigung durch HR-Director "
         "Konzern sowie Compliance Officer " + CO + ". Datenschutz-Hinweis: Diese Akte ist personenbezogen "
         "(Art. 4 Nr. 1 DSGVO) und unterliegt § 26 BDSG; Aufbewahrung 10 Jahre nach Ausscheiden."),
        ("Unterschriften",
         signatures("Julia Hartmann", "Senior Compliance Analyst", R["name"],
                    "Group HR Director", "HR-Director", R["name"],
                    place="Stuttgart", date_str_="20. September 2023")),
    ], confidential=True)


# PRJ-2022-003_Status_2023_03_ECU-900_AUTOSAR
write_doc(f"{BASE}/PRJ-2022-003_Status_2023_03_ECU-900_AUTOSAR.docx", H,
    "Projekt-Statusbericht ECU-900 / AUTOSAR – Maerz 2023",
    subtitle="PRJ-2022-003 – Compliance- und Programm-Status",
    sections=[
        ("Projekt-Stammdaten",
         "Projekt: ECU-900-AUTOSAR-Migration. Projekt-ID: PRJ-2022-003. Projekt-Manager: Lars Wittmann (RSG). "
         "Stakeholder: VW (MEB+, MQB-Evo), Stellantis (CMP). Start: 1.4.2022; Plan-SOP Q3 2023. "
         "Budget genehmigt: 6,8 Mio. EUR; Ist YTD: 4,1 Mio. EUR."),
        ("Programm-Fortschritt",
         "Stage-Gate-Status: Gate 3 (System-Verifikation) durchlaufen am 15.2.2023. Gate 4 (Vehicle-Integration) "
         "geplant 12.5.2023. AUTOSAR Classic Stack auf 22-11 aktualisiert; AUTOSAR Adaptive Migration begonnen "
         "(P4 Domain Controllers). ASPICE Level 2 Assessment fuer Software (SWE.2-SWE.4) durch Kugler-Maag "
         "geplant Q2/2023."),
        ("Compliance-Status",
         "Funktionale Sicherheit ISO 26262: ASIL-D Architektur stabil. Cybersecurity ISO 21434: Threat "
         "Assessment v2 abgeschlossen. Verrechnungspreis-Konformitaet RSG-RCZ (Engineering-Support): "
         "Cost-plus 8 % im Einklang mit TP-Richtlinie. Lieferanten-Compliance NXP/Infineon: TPDD-Klauseln "
         "vollstaendig integriert. Exportkontrolle: Komponenten-Klassifizierung Annex I bestaetigt; BAFA-"
         "Pre-Check fuer Lieferungen ausserhalb EU positiv (BAFA-Beauftragte Dr. Berger)."),
        ("Risiken und Massnahmen", [
            ["Risiko", "Wahrscheinlichkeit", "Auswirkung", "Massnahme"],
            ["Halbleiter-Engpaesse Q3", "mittel", "SOP-Verzug 4-8 Wochen", "Second-Source NXP + Allocation"],
            ["ASPICE-Reifegrad-Risiko", "niedrig", "Audit-Findings", "Frueher Self-Assessment + Kugler-Maag"],
            ["Cybersecurity-ISO-21434-Audit", "mittel", "Late Findings", "Pre-Audit MHP Q2/2023"],
            ["Sanktions-Compliance NXP", "niedrig", "Lieferunterbrechung", "Sanktions-Re-Screening quartalsweise"],
        ]),
        ("Empfehlung an Vorstand/Lenkungskreis",
         "Programm weiter im Plan. Empfehlung: Genehmigung der zusaetzlichen 280 kEUR fuer externe AUTOSAR-"
         "Adaptive-Migrationsunterstuetzung durch Vector Informatik. SOP-Termin Q3 2023 weiterhin halten."),
        ("Unterschriften",
         signatures("Lars Wittmann", "Projekt-Manager / Lead Dev RSG", "Brennhagen Software GmbH",
                    "Stefan Hoffmann", "CTO (Sponsor)", R["name"],
                    place="Muenchen", date_str_="30. Maerz 2023")),
    ])


# PRJ-2024-002_Statusbericht_2024_05_DSGVO_Compliance_Remediat
write_doc(f"{BASE}/PRJ-2024-002_Statusbericht_2024_05_DSGVO_Compliance_Remediat.docx", H,
    "Projekt-Statusbericht DSGVO-Compliance-Remediation – Mai 2024",
    subtitle="PRJ-2024-002 – Konzern-Programm zur Schliessung von DSGVO-Findings",
    sections=[
        ("Projekt-Stammdaten",
         "Projekt: DSGVO-Compliance-Remediation 2024. Projekt-ID: PRJ-2024-002. Sponsor: CFO Laura Bauer. "
         "Programm-Lead: Datenschutzbeauftragter " + DSB + " (extern). Steuerungskreis: CFO, COO, CISO, "
         "Compliance Officer " + CO + ". Laufzeit: 1.2.2024 – 31.12.2024. Budget: 1,4 Mio. EUR."),
        ("Hintergrund und Findings",
         "Im Rahmen des Konzernrevisions-Audits 2023 sowie der LfDI-Stichprobenpruefung Q4/2023 wurden 18 "
         "Findings (3 high, 9 medium, 6 low) identifiziert, insbesondere zu: VVT-Aktualitaet, AV-Vertrag-Coverage "
         "(Speicher 92 %), Loeschkonzepten in SAP/Workday, Drittlandtransfer-TIA-Coverage und Schulungs-Coverage."),
        ("Umsetzungsstatus per Mai 2024", [
            ["Workstream", "Status", "Fortschritt", "Risiko"],
            ["VVT-Aktualisierung & Tool-Migration", "in Bearbeitung", "65 %", "niedrig"],
            ["AV-Vertrag-Backfilling", "in Bearbeitung", "82 %", "niedrig"],
            ["Loeschkonzept SAP/Workday", "Konzept fertig", "40 %", "mittel"],
            ["Drittlandtransfer-TIA (SCC 2021)", "abgeschlossen", "100 %", "—"],
            ["Schulungs-Kampagne DSGVO Refresher", "laufend", "78 %", "niedrig"],
            ["Datenpannen-Tool-Roll-out", "abgeschlossen", "100 %", "—"],
        ]),
        ("Compliance- und Recht-Schnittstellen",
         "Hengeler Mueller (" + HM_PARTNER + ") rechtliche Begleitung; Lehmann & Partner (DSB) operative "
         "Steuerung; KPMG IDW PS 980-Pruefung 2024 (Wirksamkeit) plant Cut-Off Dezember 2024. "
         "Aufsichtsbehoerde LfDI Baden-Wuerttemberg ist ueber Programm informiert; Quarterly-Updates erfolgen."),
        ("Risiken und Eskalation",
         "Hauptrisiko: Loeschkonzept-Umsetzung in legacy SAP HCM (R/3 4.7-Klone) komplex; Workaround via "
         "Archivierung in OpenText im Pilot. Eskalation an CFO/COO bei Bedarf vorgesehen. "
         "Bussgeld-Risiko ohne Programm: bis 4 % Konzernumsatz / 20 Mio. EUR (Art. 83 DSGVO)."),
        ("Naechste Schritte",
         "Steering Q2: 12. Juni 2024 (LfDI-Update). Programm-Abschluss-Review 18. Dezember 2024. "
         "Wirksamkeitspruefung durch KPMG IDW PS 980 ab Januar 2025."),
        ("Unterschriften",
         signatures(DSB, "Datenschutzbeauftragter (extern)", "Lehmann & Partner",
                    "Laura Bauer", "CFO / Sponsor", R["name"],
                    place="Stuttgart", date_str_="20. Mai 2024")),
    ])


# =============================================================================
# Verification at end
# =============================================================================
if __name__ == "__main__":
    from docx import Document
    from pathlib import Path
    base = Path(BASE)
    total = thin = 0
    thinlist = []
    for p in base.rglob("*.docx"):
        d = Document(p); t = " ".join(par.text for par in d.paragraphs)
        for tbl in d.tables:
            for r in tbl.rows:
                for c in r.cells:
                    t += " " + c.text
        w = len(t.split())
        total += 1
        if w < 200:
            thin += 1
            thinlist.append((w, p.name))
    print(f"{total} total, {thin} still thin")
    for w, n in sorted(thinlist):
        print(w, n)
