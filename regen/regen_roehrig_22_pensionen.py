"""Regenerate thin docs in roehrig_large/22_Pensionen_bAV/.

Idempotent: overwrites .docx files atomically.
"""
# --- portable-paths-prelude --- (do not edit) ---
import sys
from pathlib import Path as _PathlibPath
_RP = _PathlibPath(__file__).resolve().parent
while not (_RP / "enhance_lib.py").exists() and _RP.parent != _RP:
    _RP = _RP.parent
_ROOT = _RP
sys.path.insert(0, str(_ROOT))
# --- end prelude ---

import sys
from pathlib import Path

sys.path.insert(0, f"{_ROOT}")
from enhance_lib import ROEHRIG_AG as R, write_doc, signatures  # noqa: E402

BASE = Path(f"{_ROOT}/roehrig_large/22_Pensionen_bAV")
H = {"name": R["name"], "addr": R["addr"], "hrb": R["hrb"]}
H_HEUBECK = {
    "name": "Heubeck AG",
    "addr": "Friedrich-Ebert-Allee 27, 50931 Koeln",
    "hrb": "HRB 12476, Amtsgericht Koeln",
}

# ----------------------------------------------------------------------------
# Helpers
# ----------------------------------------------------------------------------

PLAN_INFO = {
    "REA": {
        "long": "Brennhagen Elektronik AG (Konzernobergesellschaft)",
        "kreis": "Vorstand und ausgewaehlte erste Fuehrungsebene",
        "anwartschaften": 12,
        "rentner": 4,
        "ausgeschiedene": 2,
        "dbo_2021": 13.2,
        "dbo_2022": 14.0,
        "dbo_2023": 14.8,
        "plan_assets_2023": 6.4,
        "typ": "Direktzusage (leistungsorientiert)",
    },
    "REG": {
        "long": "Brennhagen Elektronik GmbH, Werk Heilbronn",
        "kreis": "Belegschaft tariflich + AT (Altzusagen vor 2015)",
        "anwartschaften": 612,
        "rentner": 184,
        "ausgeschiedene": 95,
        "dbo_2021": 38.4,
        "dbo_2022": 40.6,
        "dbo_2023": 42.8,
        "plan_assets_2023": 18.9,
        "typ": "Unterstuetzungskasse (rueckgedeckt, Talanx HDI)",
    },
    "RSG": {
        "long": "Brennhagen Software GmbH, Muenchen",
        "kreis": "Fuehrungskraefte (AT) und IT-Spezialisten Neuzusage 2015+",
        "anwartschaften": 145,
        "rentner": 22,
        "ausgeschiedene": 18,
        "dbo_2021": 16.1,
        "dbo_2022": 17.2,
        "dbo_2023": 18.4,
        "plan_assets_2023": 7.3,
        "typ": "Direktzusage (beitragsorientiert, Verzinsung 2,75 %)",
    },
    "RPL": {
        "long": "Brennhagen Polska Sp. z o.o., Katowice",
        "kreis": "Auslandsplan - defined contribution (PPK)",
        "anwartschaften": 960,
        "rentner": 0,
        "ausgeschiedene": 12,
        "dbo_2021": 0.0,
        "dbo_2022": 0.0,
        "dbo_2023": 0.0,
        "plan_assets_2023": 0.0,
        "typ": "Beitragsorientiert (PPK Polen, kein DBO)",
    },
}


def _money(v: float) -> str:
    return f"{v:,.1f}".replace(",", "X").replace(".", ",").replace("X", ".")


def aktuargutachten(jahr: int, plan: str, path: Path) -> None:
    info = PLAN_INFO[plan]
    dbo = info[f"dbo_{jahr}"]
    diskontsatz = {2021: "1,10", 2022: "3,75", 2023: "3,40"}[jahr]
    vj_diskont = {2021: "0,95", 2022: "1,10", 2023: "3,75"}[jahr]
    gehaltstrend = "2,5"
    rententrend = "1,75"
    inflation = "2,0"
    dbo_vj = info[f"dbo_{jahr - 1}"] if jahr > 2021 else dbo * 0.94

    title = f"Versicherungsmathematisches Gutachten {jahr} – {plan} ({info['long']})"
    subtitle = (
        f"IFRS IAS 19 / HGB § 253 Abs. 2 / BilMoG  |  Stichtag 31.12.{jahr}  |  "
        f"Aktuar: Heubeck AG, Lead Dipl.-Math. Andreas Geyer"
    )

    sections = [
        (
            "1. Auftrag und Grundlagen",
            (
                f"Die Heubeck AG, Koeln, wurde von der {R['name']} mit der Erstellung "
                f"des versicherungsmathematischen Gutachtens zum Stichtag 31. Dezember {jahr} "
                f"fuer den Pensionsplan »{plan}« beauftragt. Auftraggeber ist die "
                f"Konzernzentrale (Group Treasury, Markus Pflanzer; IFRS-Review: "
                f"Dr. Heike Berger, Group Tax).\n\n"
                f"Das Gutachten erfolgt parallel nach IAS 19 (revised 2011) fuer den "
                f"IFRS-Konzernabschluss und nach § 253 Abs. 2 HGB (BilMoG, "
                f"10-Jahres-Durchschnittszinssatz) fuer den Einzelabschluss. "
                f"Plantyp: {info['typ']}.\n\n"
                f"Berechnungsverfahren: Projected Unit Credit Method (PUCM) gemaess "
                f"IAS 19.67 ff. Biometrische Grundlagen: Heubeck-Richttafeln 2018 G."
            ),
        ),
        (
            "2. Bewertungsannahmen",
            [
                ["Parameter", f"31.12.{jahr}", f"31.12.{jahr - 1}", "Bemerkung"],
                ["Diskontierungszinssatz IFRS", f"{diskontsatz} %", f"{vj_diskont} %", "Iboxx EUR AA 10+ Y"],
                ["Diskontierungszinssatz HGB", "1,82 %", "1,78 %", "§ 253 Abs. 2 HGB"],
                ["Gehaltstrend", f"{gehaltstrend} %", f"{gehaltstrend} %", "DE-Tarif AT-Mitarbeiter"],
                ["Rententrend", f"{rententrend} %", f"{rententrend} %", "§ 16 BetrAVG Annahme"],
                ["Inflation", f"{inflation} %", f"{inflation} %", "ECB Survey of Professional Forecasters"],
                ["Fluktuation", "3,2 %", "3,5 %", "alters-/dienstzeitabhaengig"],
                ["Pensionierungsalter", "67", "67", "Regelaltersgrenze SGB VI"],
                ["Sterbetafeln", "RT 2018 G", "RT 2018 G", "Heubeck-Richttafeln"],
            ],
        ),
        (
            "3. Bestand der Versorgungsberechtigten",
            [
                ["Kategorie", "Anzahl", "Ø Alter", "Ø Dienstjahre"],
                ["Aktive Anwartschaftsberechtigte", str(info["anwartschaften"]), "44", "12"],
                ["Ausgeschiedene mit unverfallbarer Anwartschaft", str(info["ausgeschiedene"]), "52", "—"],
                ["Laufende Rentner / Hinterbliebene", str(info["rentner"]), "71", "—"],
                ["Gesamt", str(info["anwartschaften"] + info["ausgeschiedene"] + info["rentner"]), "—", "—"],
            ],
        ),
        (
            "4. Defined Benefit Obligation (DBO) zum Stichtag",
            [
                ["Position", f"31.12.{jahr} (Mio. EUR)", f"31.12.{jahr - 1} (Mio. EUR)"],
                [f"DBO {plan} gesamt", _money(dbo), _money(dbo_vj)],
                ["davon Aktive", _money(dbo * 0.62), _money(dbo_vj * 0.62)],
                ["davon Ausgeschiedene", _money(dbo * 0.14), _money(dbo_vj * 0.13)],
                ["davon Rentner", _money(dbo * 0.24), _money(dbo_vj * 0.25)],
                ["Service Cost (laufender Dienstzeitaufwand)", _money(dbo * 0.041), _money(dbo_vj * 0.042)],
                ["Interest Cost", _money(dbo_vj * float(vj_diskont.replace(",", ".")) / 100), _money(dbo_vj * 0.011)],
            ],
        ),
        (
            "5. Versicherungsmathematische Gewinne und Verluste",
            (
                f"Die Neubewertungen aus Aenderungen finanzieller Annahmen "
                f"(insbesondere Diskontierungszinssatz von {vj_diskont} % auf "
                f"{diskontsatz} %) und demografischer Annahmen werden gemaess "
                f"IAS 19.127 (Remeasurements) erfolgsneutral im sonstigen "
                f"Ergebnis (OCI) erfasst.\n\n"
                f"Erfahrungsbedingte Abweichungen (Fluktuation, "
                f"Gehaltsentwicklung, Sterblichkeit) wurden ausgewertet und in "
                f"den Annahmen reflektiert. Eine wesentliche Aenderung der "
                f"demografischen Basis (RT 2018 G) wird nicht empfohlen."
            ),
        ),
        (
            "6. Sensitivitaetsanalyse (IAS 19.145)",
            [
                ["Parametervariation", "DBO-Wirkung (Mio. EUR)", "in %"],
                ["Diskontsatz +0,50 %-Pkt.", f"-{_money(dbo * 0.078)}", "-7,8 %"],
                ["Diskontsatz -0,50 %-Pkt.", f"+{_money(dbo * 0.084)}", "+8,4 %"],
                ["Gehaltstrend +0,50 %-Pkt.", f"+{_money(dbo * 0.027)}", "+2,7 %"],
                ["Rententrend +0,25 %-Pkt.", f"+{_money(dbo * 0.034)}", "+3,4 %"],
                ["Lebenserwartung +1 Jahr", f"+{_money(dbo * 0.038)}", "+3,8 %"],
            ],
        ),
        (
            "7. Bestaetigung des Aktuars",
            (
                f"Wir bestaetigen, dass das vorliegende Gutachten unter Beachtung "
                f"der berufsstaendischen Grundsaetze der Deutschen "
                f"Aktuarvereinigung e.V. (DAV) und nach den Vorgaben des IAS 19 "
                f"(revised 2011) sowie HGB § 253 i.V.m. RuekStuckV erstellt wurde. "
                f"Die im Gutachten verwendeten Daten wurden von der {R['name']} "
                f"zur Verfuegung gestellt und auf Plausibilitaet geprueft.\n\n"
                f"Heubeck AG, Koeln, im Februar {jahr + 1}\n"
                f"Dipl.-Math. Andreas Geyer (Lead-Aktuar, DAV-Mitglied Nr. 4271)\n"
                f"Aktuar Junior: Anna-Lena Becker, M.Sc. Math."
            ),
        ),
    ]

    write_doc(
        path,
        H_HEUBECK,
        title,
        sections,
        subtitle=subtitle,
        confidential=True,
    )


def jahresinfo(num: int, name: str, path: Path, draft: bool = False) -> None:
    """Persoenliche Jahresmitteilung Versorgungsberechtigte."""
    rentenbaustein_pa = round(120 + (num * 37.4) % 850, 2)
    anwartschaft_kapital = round(rentenbaustein_pa * 12 * (28 + num % 12), 2)
    eintritt = 2000 + (num * 7) % 24

    title = f"Persoenliche Jahresmitteilung zur betrieblichen Altersversorgung 2023"
    subtitle = (
        f"{name}  |  Personal-Nr. P-{1000 + num}  |  Stichwert 31.12.2023  |  "
        f"Versorgungswerk REG/RSG-Neuzusage 2015"
    )

    sections = [
        (
            "Sehr geehrte/r Versorgungsberechtigte/r,",
            (
                f"hiermit erhalten Sie - gemaess § 4a Betriebsrentengesetz "
                f"(BetrAVG) - Ihre persoenliche Jahresmitteilung zur "
                f"betrieblichen Altersversorgung (bAV) zum Stichtag "
                f"31. Dezember 2023.\n\n"
                f"Diese Mitteilung erfolgt unaufgefordert und jaehrlich. Bitte "
                f"pruefen Sie die persoenlichen Daten und melden sich bei "
                f"Abweichungen innerhalb von 6 Wochen bei der HR-Abteilung "
                f"(bav@brennhagen-elektronik.de)."
            ),
        ),
        (
            "1. Persoenliche Versorgungsdaten",
            [
                ["Position", "Wert"],
                ["Name", name],
                ["Personalnummer", f"P-{1000 + num}"],
                ["Eintrittsdatum", f"01.04.{eintritt}"],
                ["Anwartschaft unverfallbar seit", f"01.04.{eintritt + 3}"],
                ["Versorgungswerk", "Neuzusage 2015 (beitragsorientiert)"],
                ["Beitragssatz Arbeitgeber", "3,5 % Bemessungsentgelt"],
                ["Entgeltumwandlung Arbeitnehmer", f"{(num % 4) * 50 + 100} EUR/Monat"],
                ["Garantieverzinsung", "2,75 % p.a."],
            ],
        ),
        (
            "2. Anwartschaftsbarwerte zum 31.12.2023",
            [
                ["Komponente", "Wert (EUR)"],
                ["Bisher erworbener Rentenbaustein (p.a.)", _money(rentenbaustein_pa)],
                ["Gesamt-Anwartschaftskapital (Barwert)", _money(anwartschaft_kapital)],
                ["Prognostizierte Monatsrente ab Alter 67", _money(rentenbaustein_pa * 1.25)],
                ["Alternativ: einmalige Kapitalauszahlung", _money(anwartschaft_kapital * 1.04)],
                ["Hinterbliebenenrente (60 % der Altersrente)", _money(rentenbaustein_pa * 0.75)],
                ["Berufsunfaehigkeitsrente (50 %)", _money(rentenbaustein_pa * 0.625)],
            ],
        ),
        (
            "3. Erlaeuterungen",
            (
                "Die ausgewiesenen Werte basieren auf dem versicherungs"
                "mathematischen Gutachten der Heubeck AG, Koeln, vom Februar "
                "2024, mit folgenden Annahmen: Diskontierungssatz 3,40 % p.a., "
                "Gehaltstrend 2,5 % p.a., Rententrend 1,75 % p.a., Sterbetafeln "
                "Heubeck Richttafeln 2018 G.\n\n"
                "Der Wert »Gesamt-Anwartschaftskapital« stellt den Barwert Ihrer "
                "bisher erworbenen Versorgungsleistungen dar und ist nicht mit "
                "einem garantierten Auszahlungsbetrag gleichzusetzen. Die "
                "endgueltige Hoehe der Leistungen bemisst sich nach der zum "
                "Zeitpunkt des Versorgungsfalls geltenden Versorgungsordnung."
            ),
        ),
        (
            "4. Insolvenzsicherung",
            (
                "Ihre Anwartschaft ist gemaess § 7 BetrAVG durch den "
                "Pensions-Sicherungs-Verein a.G. (PSVaG), Koeln, gegen "
                "Insolvenz des Arbeitgebers gesichert. Die Beitraege werden "
                "vom Arbeitgeber getragen.\n\n"
                "Zusaetzlich besteht ein Contractual Trust Arrangement (CTA) "
                "zur insolvenzfesten Auslagerung des Pensionsvermoegens "
                "(Brennhagen Pension Trust e.V., gegruendet 2018, Treuhaender: "
                "RBS RoeverBroennerSusat Mazars)."
            ),
        ),
        (
            "5. Steuerliche Behandlung",
            (
                "Beitraege zur Entgeltumwandlung sind im Rahmen der § 3 Nr. 63 "
                "EStG-Grenze (8 % BBG-West, 2023 = 7.008 EUR p.a.) steuer- und "
                "sozialabgabenfrei. Bei spaeterer Auszahlung erfolgt nachgelagerte "
                "Besteuerung gemaess § 22 Nr. 5 EStG.\n\n"
                "Fragen zur individuellen steuerlichen Behandlung beantwortet "
                "Ihre HR-Ansprechperson oder die Steuerberatung."
            ),
        ),
        (
            "Ansprechpartner und Rueckfragen",
            (
                f"Bei Fragen wenden Sie sich bitte an:\n\n"
                f"HR Pensionen / bAV-Service Center\n"
                f"{R['name']}, {R['addr']}\n"
                f"E-Mail: bav@brennhagen-elektronik.de  |  Tel.: +49 711 4567-{200 + num % 50}\n"
                f"Sachbearbeiterin: Frau Susanne Loeffler (Mo-Do 9-15 Uhr)\n"
                f"Pensionsbeauftragter: Markus Pflanzer (Group Treasurer)\n\n"
                f"Mit freundlichen Gruessen\n"
                f"i.A. HR Compensation & Benefits"
            ),
        ),
    ]

    write_doc(
        path,
        H,
        title,
        sections,
        subtitle=subtitle,
        confidential=True,
        draft=draft,
    )


def direktzusage(name: str, position: str, vs_geburtsdatum: str, ek_jahr: int, monatsbeitrag_eur: int, path: Path) -> None:
    title = f"Einzelne Versorgungszusage (Direktzusage) – {name}"
    subtitle = f"{position}  |  Versorgungsvertrag {ek_jahr}  |  beitragsorientiert"

    sections = [
        (
            "Praeambel",
            (
                f"Zwischen der {R['name']}, {R['addr']} ({R['hrb']}), "
                f"vertreten durch den Vorstand (nachfolgend »Arbeitgeber«), und "
                f"Herrn/Frau {name}, geboren am {vs_geburtsdatum} (nachfolgend "
                f"»Versorgungsberechtigte/r«), wird auf Grundlage des "
                f"Vorstands-/Anstellungsvertrages und unter Beruecksichtigung "
                f"der konzernweiten Versorgungsordnung »Neuzusage 2015« die "
                f"nachfolgende Einzel-Direktzusage zur betrieblichen "
                f"Altersversorgung (bAV) erteilt."
            ),
        ),
        (
            "Vertragsklauseln",
            (
                "clauses",
                [
                    (
                        "§ 1 Gegenstand der Zusage",
                        [
                            f"Der Arbeitgeber erteilt {name} eine wertgleiche, "
                            f"unverfallbare Versorgungszusage in Form einer "
                            f"Direktzusage (§ 1 Abs. 1 Nr. 1 BetrAVG) zur "
                            f"betrieblichen Altersversorgung. Die Zusage umfasst "
                            f"Leistungen bei Erreichen der Regelaltersgrenze "
                            f"(67), bei Berufsunfaehigkeit sowie an "
                            f"Hinterbliebene.",
                        ],
                    ),
                    (
                        "§ 2 Beitragsaequivalent",
                        [
                            f"Der Arbeitgeber stellt monatlich einen "
                            f"Versorgungsbeitrag von {monatsbeitrag_eur:,} EUR "
                            f"(i.W. {monatsbeitrag_eur:,} Euro) zur Verfuegung. "
                            f"Dieser wird mit 2,75 % p.a. verzinst und bildet "
                            f"das Anwartschaftskapital. Zusaetzliche "
                            f"Entgeltumwandlung gemaess § 1a BetrAVG ist "
                            f"moeglich.".replace(",", "."),
                        ],
                    ),
                    (
                        "§ 3 Unverfallbarkeit",
                        [
                            "Die Zusage ist gemaess § 1b BetrAVG mit "
                            "Erteilung sofort unverfallbar (vertragliche "
                            "Verbesserung gegenueber gesetzlicher 3-Jahres-Frist).",
                        ],
                    ),
                    (
                        "§ 4 Leistungsformen",
                        [
                            "Bei Erreichen der Regelaltersgrenze wird auf Wunsch "
                            "des Versorgungsberechtigten eine lebenslange "
                            "monatliche Rente oder eine Kapitalauszahlung (auch "
                            "Ratenzahlung in bis zu 10 Jahresraten) gewaehrt. "
                            "Die Wahl ist 3 Jahre vor Versorgungsbeginn schriftlich "
                            "anzuzeigen.",
                            "Hinterbliebenenleistungen: 60 % an Ehegatten/eingetragene "
                            "Lebenspartner, 20 % je Waisenkind (max. 100 % gesamt).",
                        ],
                    ),
                    (
                        "§ 5 Rueckdeckung und Insolvenzschutz",
                        [
                            "Die Verpflichtung wird durch eine kongruente "
                            "Rueckdeckungsversicherung bei der Talanx HDI "
                            "Lebensversicherung AG bzw. Allianz "
                            "Lebensversicherung AG abgedeckt; die Anspruechen "
                            "aus der Versicherung sind dem "
                            "Versorgungsberechtigten verpfaendet.",
                            "Zusaetzlich besteht Insolvenzschutz durch den "
                            "Pensions-Sicherungs-Verein a.G., Koeln (PSVaG) "
                            "gemaess § 7 BetrAVG.",
                        ],
                    ),
                    (
                        "§ 6 Anpassungspruefung",
                        [
                            "Laufende Renten werden gemaess § 16 BetrAVG alle "
                            "3 Jahre auf Anpassungsbedarf geprueft; die "
                            "Verpflichtung gilt als erfuellt, wenn die Anpassung "
                            "mindestens dem Verbraucherpreisindex (VPI) "
                            "entspricht.",
                        ],
                    ),
                    (
                        "§ 7 Sonstiges",
                        [
                            f"Aenderungen beduerfen der Schriftform. "
                            f"Gerichtsstand: Stuttgart. Anwendbares Recht: "
                            f"Recht der Bundesrepublik Deutschland.",
                        ],
                    ),
                ],
            ),
        ),
        (
            "Unterschriften",
            signatures(
                "Anna Mueller", "CEO", R["name"],
                name, position, R["name"],
                place="Stuttgart", date_str_=f"15.03.{ek_jahr}",
            ),
        ),
        (
            "Gegenzeichnung HR",
            signatures(
                "Sandra Berger", "Head of HR Group", R["name"],
                "Markus Pflanzer", "Pensionsbeauftragter / Group Treasurer", R["name"],
                place="Stuttgart", date_str_=f"15.03.{ek_jahr}",
            ),
        ),
    ]

    write_doc(path, H, title, sections, subtitle=subtitle, confidential=True)


def ias19_disclosure(plan_name: str, kurzbeschreibung: str, dbo: float, plan_assets: float, path: Path) -> None:
    """IAS-19 Plan Disclosure / Versorgungsordnung Beschreibung."""
    title = f"Konzern-Versorgungsplan: {plan_name}"
    subtitle = (
        f"IAS 19 Plan Disclosure / Versorgungsordnung  |  Stichtag 31.12.2023  |  "
        f"{R['name']} IFRS-Konzernabschluss"
    )

    netto = dbo - plan_assets

    sections = [
        (
            "1. Planbeschreibung",
            kurzbeschreibung
            + "\n\nDer Plan wird gemaess IAS 19 (revised 2011) als "
            "leistungsorientierter Plan (defined benefit plan) klassifiziert. "
            "Die versicherungsmathematische Bewertung erfolgt jaehrlich durch "
            "die Heubeck AG, Koeln (Lead-Aktuar: Dipl.-Math. Andreas Geyer).",
        ),
        (
            "2. Regelungsinhalt",
            (
                "clauses",
                [
                    (
                        "§ 1 Persoenlicher Geltungsbereich",
                        [
                            "Versorgungsberechtigt sind alle Mitarbeiterinnen "
                            "und Mitarbeiter mit unbefristetem Arbeitsvertrag "
                            "nach Vollendung des 21. Lebensjahres und einer "
                            "Mindestbeschaeftigung von 12 Monaten.",
                        ],
                    ),
                    (
                        "§ 2 Leistungsarten",
                        [
                            "Altersrente (ab Regelaltersgrenze 67), vorzeitige "
                            "Altersrente (ab 63 mit Abschlag 0,3 %/Monat), "
                            "Erwerbsminderungsrente, Hinterbliebenenrente, "
                            "Waisenrente.",
                        ],
                    ),
                    (
                        "§ 3 Finanzierung",
                        [
                            "Arbeitgeberbeitrag 3,5 % des Bemessungsentgelts "
                            "(max. BBG-West). Entgeltumwandlung freiwillig im "
                            "Rahmen § 1a BetrAVG, mit AG-Zuschuss 15 %.",
                            "Die Verpflichtungen werden durch ein Contractual "
                            "Trust Arrangement (CTA, Brennhagen Pension Trust e.V., "
                            "Stuttgart, 2018) sowie Rueckdeckungsversicherungen "
                            "bei Talanx HDI und Allianz besichert.",
                        ],
                    ),
                    (
                        "§ 4 Unverfallbarkeit",
                        [
                            "Gesetzliche Unverfallbarkeit gemaess § 1b BetrAVG "
                            "(3 Jahre Zusagedauer, Vollendung 21. Lebensjahr). "
                            "Bei Neuzusage 2015 sofort unverfallbar.",
                        ],
                    ),
                    (
                        "§ 5 Anpassung",
                        [
                            "Laufende Renten werden gemaess § 16 BetrAVG alle "
                            "3 Jahre geprueft. Mindestanpassung in Hoehe des "
                            "Verbraucherpreisindex (VPI).",
                        ],
                    ),
                ],
            ),
        ),
        (
            "3. Beteiligte Konzerngesellschaften",
            [
                ["Gesellschaft", "Beteiligte MA (ca.)", "Plan-Status"],
                ["Brennhagen Elektronik AG (Holding/Vorstand)", "12", "Direktzusage"],
                ["Brennhagen Elektronik GmbH (REG Heilbronn)", "612", "U-Kasse"],
                ["Brennhagen Software GmbH (RSG Muenchen)", "145", "Direktzusage NZ-2015"],
                ["Brennhagen Holding GmbH (RHO Stuttgart)", "32", "Direktzusage"],
                ["Auslandstoechter (RPL/RCZ/RHU/RCN)", "n/a", "defined contribution"],
            ],
        ),
        (
            "4. Quantitative IAS-19-Angaben zum Stichtag",
            [
                ["Position", "31.12.2023 (Mio. EUR)", "31.12.2022 (Mio. EUR)"],
                ["Defined Benefit Obligation (DBO)", _money(dbo), _money(dbo * 0.94)],
                ["Beizulegender Zeitwert Planvermoegen", _money(plan_assets), _money(plan_assets * 0.92)],
                ["Nettoverpflichtung (Bilanz)", _money(netto), _money(dbo * 0.94 - plan_assets * 0.92)],
                ["Aktivische Hoechstwertbegrenzung (Asset Ceiling)", "0,0", "0,0"],
                ["Service Cost (laufender Aufwand)", _money(dbo * 0.041), _money(dbo * 0.040)],
                ["Net Interest", _money(netto * 0.034), _money(netto * 0.0375)],
            ],
        ),
        (
            "5. Asset Allocation Planvermoegen",
            [
                ["Anlageklasse", "Quote 2023", "Quote 2022", "Strategische Bandbreite"],
                ["Anleihen Investment Grade", "52 %", "55 %", "45-65 %"],
                ["Aktien (Industrielaender)", "28 %", "26 %", "20-35 %"],
                ["Aktien (Emerging Markets)", "6 %", "6 %", "0-10 %"],
                ["Immobilien (offene Fonds)", "8 %", "7 %", "5-12 %"],
                ["Liquide Mittel / Geldmarkt", "4 %", "4 %", "0-10 %"],
                ["Alternative Investments (Private Debt)", "2 %", "2 %", "0-5 %"],
            ],
        ),
        (
            "6. Verantwortlichkeiten",
            (
                "Pensionsbeauftragter: Markus Pflanzer (Group Treasurer).\n"
                "Aktuargutachten-Review IFRS: Dr. Heike Berger (Group Tax / IFRS).\n"
                "Asset-Allocation / ALM: Investment Committee (Vorsitz: Laura Bauer, CFO).\n"
                "Aktuar: Heubeck AG, Koeln. Wirtschaftspruefer: KPMG AG WPG."
            ),
        ),
    ]

    write_doc(path, H, title, sections, subtitle=subtitle, confidential=True)


def versorgungsordnung(werk: str, planname: str, dbo: float, path: Path, draft: bool = False, alt: bool = False) -> None:
    suffix = " (ALT - bis 2014)" if alt else ""
    ias19_disclosure(planname + suffix, f"Versorgungsordnung des Werks {werk}. Geltungsbereich: alle Mitarbeitenden gemaess Versorgungsplan {planname}.", dbo, dbo * 0.42, path)


# ----------------------------------------------------------------------------
# Aktuarberichte
# ----------------------------------------------------------------------------

AKTUAR_FILES = [
    ("Aktuarbericht_REA_2021.docx", 2021, "REA"),
    ("Aktuarbericht_REA_2022.docx", 2022, "REA"),
    ("Aktuarbericht_REA_2023.docx", 2023, "REA"),
    ("Aktuarbericht_REG_2021.docx", 2021, "REG"),
    ("Aktuarbericht_REG_2022.docx", 2022, "REG"),
    ("Aktuarbericht_REG_2023.docx", 2023, "REG"),
    ("Aktuarbericht_RPL_2021.docx", 2021, "RPL"),
    ("Aktuarbericht_RPL_2022.docx", 2022, "RPL"),
    ("Aktuarbericht_RPL_2023.docx", 2023, "RPL"),
    ("Aktuarbericht_RSG_2021.docx", 2021, "RSG"),
    ("Aktuarbericht_RSG_2022.docx", 2022, "RSG"),
    ("Aktuarbericht_RSG_2023_FINAL_v2.docx", 2023, "RSG"),
]

# Einzelzusagen
EINZELZUSAGEN = [
    ("bAV_Einzelzusage_Dr._Petra_Hollmann_2020.docx", "Dr. Petra Hollmann", "CTO (designiert)", "14.06.1969", 2020, 3200),
    ("bAV_Einzelzusage_Klaus-Peter_Zimmermann_2020.docx", "Klaus-Peter Zimmermann", "Bereichsleiter Operations", "22.11.1972", 2020, 1800),
    ("bAV_Einzelzusage_Stefan_Richter_2020.docx", "Stefan Richter", "CMO/Business Development", "08.03.1971", 2020, 2400),
]

# Versorgungsplaene
VERSORGUNGSPLAENE = [
    ("REA_bAV_BZML_Beitragsorientierter_Plan_2023.docx", "REA", "Beitragsorientierter Plan REA-BZML 2023", 14.8),
    ("REA_bAV_Direktzusage_Plan_REG_Heilbronn_2023_reviewed.docx", "REG Heilbronn", "Direktzusagenplan REG Heilbronn 2023", 42.8),
    ("REA_bAV_Entgeltumwandlung_Plan_2020_2023_ALT.docx", "REA", "Entgeltumwandlung Plan 2020-2023", 8.6),
    ("REA_bAV_Neuzusage_Plan_2015_2023.docx", "RSG", "Neuzusage 2015 (konzernweit)", 18.4),
    ("REA_bAV_Pensionskasse_Plan_REA_AG_2023.docx", "REA", "Pensionskasse REA AG", 6.2),
]


# ----------------------------------------------------------------------------
# Special docs
# ----------------------------------------------------------------------------

def write_it_change_request():
    path = BASE / "IT_Change_Request_2023_0025.docx"
    title = "IT Change Request 2023-0025 — bAV-Stammdaten Migration SAP HCM PA0021"
    subtitle = "Change Advisory Board (CAB)  |  System SAP HCM (Produktiv P01)  |  Antrag 14.03.2023"
    sections = [
        (
            "1. Antragsdaten",
            [
                ["Feld", "Wert"],
                ["CR-Nummer", "CR-2023-0025"],
                ["Antragsteller", "Susanne Loeffler (HR Comp&Ben)"],
                ["Verantwortlich IT", "Markus Hellweg (SAP HCM Team Lead)"],
                ["Genehmiger Fachbereich", "Sandra Berger (Head of HR Group)"],
                ["Genehmiger Pensionen", "Markus Pflanzer (Group Treasurer)"],
                ["Datum Antrag", "14.03.2023"],
                ["Datum CAB-Beschluss", "21.03.2023"],
                ["Geplante Umsetzung", "Wochenende 08./09.04.2023"],
                ["Prioritaet", "Mittel"],
                ["Risikoklasse", "B (Standard-CR mit Test)"],
            ],
        ),
        (
            "2. Beschreibung der Aenderung",
            (
                "Migration der bAV-Stammdaten aus dem Altsystem »BavTool 2.0« "
                "(Eigenentwicklung 2008, abgeloest) in das SAP HCM Infotyp 0021 "
                "(Familie/Angehoerige) und Infotyp 0167 (Versicherungsverhaeltnisse) "
                "fuer 1.247 aktive Versorgungsberechtigte sowie 184 Rentner.\n\n"
                "Zweck: Konsolidierung der Stammdaten in einem System, Wegfall "
                "des Schnittstellenrisikos zur Heubeck-Aktuarsoftware, "
                "Vereinfachung der jaehrlichen Datenuebernahme fuer das "
                "Aktuargutachten (heute 18 Personentage, kuenftig 4 PT)."
            ),
        ),
        (
            "3. Betroffene Systeme",
            [
                ["System", "Aenderung", "Verantwortlicher"],
                ["SAP HCM P01 (Produktiv)", "Neue Infotyp-Auspraegungen 0021/0167", "M. Hellweg"],
                ["BavTool 2.0 Legacy", "Read-only ab 10.04.2023, Stilllegung 31.12.2023", "T. Schaefer"],
                ["Heubeck-Schnittstelle", "Anpassung CSV-Export an neue Felder", "Heubeck AG Externe"],
                ["DWH/Reporting", "Neue Views pension_mart_v2", "Florian Maier (Group Controlling)"],
            ],
        ),
        (
            "4. Risikobewertung und Tests",
            (
                "Risiken: Datenmigrationsfehler (Mitigation: Dual-Run mit "
                "Abstimmung Heubeck-Aktuarwerte vor/nach Migration), "
                "Datenschutzbedenken (DSFA durch Datenschutzbeauftragten "
                "Dr. Stefan Roth durchgefuehrt 18.03.2023, keine kritischen "
                "Befunde).\n\n"
                "Testumfang: 4 Wochen UAT auf Q-System mit 50 reprodukten "
                "Testfaellen aus 8 Versorgungsplaenen. Abnahme HR: 04.04.2023, "
                "Abnahme Treasury: 05.04.2023.\n\n"
                "Rollback-Plan: Snapshot SAP HCM vor Migration, "
                "Wiederherstellung innerhalb 4h moeglich; Legacy-System bleibt "
                "bis 31.12.2023 als Read-only-Backup verfuegbar."
            ),
        ),
        (
            "5. CAB-Beschluss",
            (
                "Das Change Advisory Board (CAB) genehmigt den CR mit folgenden "
                "Auflagen: (a) Dual-Run Aktuarbericht 2023 mit beiden "
                "Datenstaenden, (b) Reconciliation-Report an Group Tax bis "
                "30.04.2023, (c) Datenschutz-Audit Q3-2023 durch Internal Audit "
                "(CAE Andreas Buehler).\n\n"
                "Stimmen: 7 dafuer, 0 dagegen, 1 Enthaltung (IT-Security wegen "
                "noch ausstehender Pen-Test-Bestaetigung)."
            ),
        ),
        (
            "6. Unterschriften",
            signatures(
                "Markus Hellweg", "SAP HCM Team Lead", R["name"],
                "Sandra Berger", "Head of HR Group", R["name"],
                place="Stuttgart", date_str_="21.03.2023",
            ),
        ),
    ]
    write_doc(path, H, title, sections, subtitle=subtitle, confidential=True)


def write_mr_rpl():
    path = BASE / "MR_RPL_Managementbewertung_2023.docx"
    title = "Management-Review: Auslandsplan RPL Katowice (PPK Polen) 2023"
    subtitle = "Bewertung defined-contribution-Plan / Brennhagen Polska Sp. z o.o."
    sections = [
        (
            "1. Auftrag und Zielsetzung",
            (
                "Auf Anweisung des Group Treasurer (Markus Pflanzer) wurde der "
                "polnische Pensionsplan PPK (Pracownicze Plany Kapitalowe) der "
                "Brennhagen Polska Sp. z o.o. (RPL), Katowice, im Hinblick auf "
                "IAS-19-Klassifizierung, Beitragsstruktur und Performance des "
                "externen Anbieters NN Investment Partners TFI S.A. einer "
                "Managementbewertung unterzogen."
            ),
        ),
        (
            "2. Eckdaten Plan",
            [
                ["Position", "Wert"],
                ["Anbieter / Verwalter", "NN Investment Partners TFI S.A., Warschau"],
                ["Plantyp", "Defined Contribution (PPK gesetzlich verpflichtend)"],
                ["IAS-19-Klassifikation", "Beitragsorientierter Plan, IAS 19.50"],
                ["Aktive Mitarbeiter (TN)", "748 von 960 (78 %)"],
                ["Arbeitgeberbeitrag (RPL)", "1,5 % Bruttogehalt (gesetzlicher Mindestbeitrag)"],
                ["Arbeitgeber-Zusatzbeitrag", "0,5 % freiwillig (Beschluss VL 2022)"],
                ["Arbeitnehmerbeitrag", "2,0 % (gesetzlich) + bis 2,0 % freiwillig"],
                ["Staatlicher Welcome-Bonus", "250 PLN/Eintritt"],
                ["Beitragsvolumen 2023 (PLN)", "8,4 Mio. PLN (RPL-Anteil)"],
                ["Beitragsvolumen 2023 (EUR)", "1,87 Mio. EUR (Kurs 4,48)"],
            ],
        ),
        (
            "3. Bilanzielle Behandlung Konzern",
            (
                "Da es sich um einen beitragsorientierten Plan handelt, werden "
                "die Beitraege als laufender Personalaufwand erfasst (IAS 19.51). "
                "Es entsteht keine DBO und kein Plan-Asset auf "
                "Konzernebene. Die Beitraege fliessen direkt an NN Investment "
                "Partners; das Anlagerisiko traegt der einzelne Arbeitnehmer.\n\n"
                "Eine Risikoeinschaetzung aus Sicht der Brennhagen-Gruppe ist daher "
                "auf das Beitragsvolumen (operatives Cost-Risk) beschraenkt."
            ),
        ),
        (
            "4. Fazit / Empfehlung",
            (
                "Der Plan wird als unverdaechtig beurteilt. Empfehlungen: "
                "(a) Beitragsstabilitaet sicherstellen, (b) jaehrliches Reporting "
                "an Group HR durch RPL HR (Anna Kowalska), (c) Pruefung "
                "Verschmelzung mit anderen EU-Auslandsplaenen im Rahmen einer "
                "EBA-Initiative 2025.\n\n"
                "Geprueft: Markus Pflanzer (Group Treasurer), 12.12.2023.\n"
                "Reviewed: Dr. Heike Berger (Group Tax / IFRS), 19.12.2023."
            ),
        ),
    ]
    write_doc(path, H, title, sections, subtitle=subtitle)


def write_bmw_ecr():
    path = BASE / "REA_BMW_ECU-900_ECR_2022_002.docx"
    title = "BMW Engineering Change Request ECR-2022-002 – ECU-900 Powertrain Gen3"
    subtitle = "Kunde BMW Group  |  Bauteil ECU-900-PT  |  ECR-Nr. BMW-2022-EN-7421"
    sections = [
        (
            "1. Aenderungsanlass",
            (
                "BMW Group fordert eine Aenderung am ECU-900-PT (Powertrain Gen3) "
                "zur Anpassung an MEB+/MQB-Evo-Plattform-Variante 2024 MJ. "
                "Konkret: Erhoehung der CAN-FD-Bandbreite von 2 Mbit/s auf "
                "5 Mbit/s sowie Einfuehrung eines optionalen Ethernet-Ports "
                "(100BASE-T1) fuer Diagnose."
            ),
        ),
        (
            "2. Technische Aenderungen",
            [
                ["Komponente", "Aenderung", "Stufe"],
                ["Mikrocontroller", "Tausch Infineon TC275 -> TC375", "Major"],
                ["CAN-Transceiver", "TJA1144 -> TJA1463 (CAN-FD 5 Mbit/s)", "Minor"],
                ["Ethernet-PHY", "Neu: NXP TJA1100 100BASE-T1", "Add-on"],
                ["Gehaeuse", "Identisch (kein Tooling-Change)", "—"],
                ["Software", "AUTOSAR 4.4 -> 4.7, Update Bootloader", "Major"],
                ["Pruefumfang", "Erweiterung EMV-Test nach LV 124 2021-08", "—"],
            ],
        ),
        (
            "3. Auswirkungen bAV (Pension-Folge)",
            (
                "Diese ECR ist in den Pensionsordner falsch abgelegt - sie betrifft "
                "produktseitig keinen bAV-Sachverhalt. Allerdings wird durch die "
                "Beschaeftigungsausweitung in RSG Muenchen (zusaetzliche 4 SW-Entwickler "
                "fuer AUTOSAR-Migration) ein zusaetzlicher bAV-Aufwand von ca. "
                "16 TEUR/Jahr fuer den Neuzusageplan 2015 erwartet. Dies ist im "
                "HR-Budget 2024 bereits enthalten."
            ),
        ),
        (
            "4. Zeitschiene und Kommerzielles",
            (
                "ECR-Approval BMW: erwartet bis 15.05.2022. PPAP-Run-at-Rate: "
                "Q4-2022. SOP-Variante 2024 MJ: 12.09.2022. "
                "Stueckpreisaenderung: +3,40 EUR/Stueck (Material + ENT-Anteil), "
                "Verhandlung mit BMW Einkauf (Frau Dr. Friederich) noch offen."
            ),
        ),
        (
            "5. Freigaben",
            signatures(
                "Dr. Klaus Kessler", "Werkleiter RSG Muenchen", "Brennhagen Software GmbH",
                "Markus Berthold", "Lead Engineer ECU", R["name"],
                place="Muenchen", date_str_="03.05.2022",
            ),
        ),
    ]
    write_doc(path, H, title, sections, subtitle=subtitle)


def write_cta_treuhand():
    path = BASE / "REA_CTA_Treuhandvertrag_2018.docx"
    title = "Contractual Trust Arrangement (CTA) – Treuhandvertrag Brennhagen Pension Trust e.V."
    subtitle = "Doppelseitige Treuhand (administrativ + Sicherungstreuhand)  |  Abschluss 14.06.2018"
    sections = [
        (
            "Praeambel",
            (
                f"Zwischen der {R['name']}, {R['addr']} (»Treugeberin«), "
                f"und dem Brennhagen Pension Trust e.V., Stuttgart (»Treuhaender«), "
                f"vertreten durch den Vorstand des Vereins (Vorsitz: "
                f"Markus Pflanzer; stv. Vorsitz: Dr. Heike Berger; "
                f"Schatzmeister: Florian Maier), wird zur Auslagerung des "
                f"Pensionsvermoegens ein Contractual Trust Arrangement (CTA) "
                f"in Form einer doppelseitigen Treuhand vereinbart."
            ),
        ),
        (
            "Vertragsklauseln",
            (
                "clauses",
                [
                    (
                        "§ 1 Zweck der Treuhand",
                        [
                            "Die Treuhand dient (a) der insolvenzfesten "
                            "Sicherung der Versorgungsanwartschaften zugunsten "
                            "der Versorgungsberechtigten der Brennhagen-Gruppe und "
                            "(b) der bilanziellen Saldierung des Plan-Asset mit "
                            "der DBO gemaess IAS 19.8.",
                            "Der Treuhaender haelt das Sicherungsvermoegen "
                            "ausschliesslich zugunsten der "
                            "Versorgungsberechtigten und nicht zugunsten der "
                            "Treugeberin.",
                        ],
                    ),
                    (
                        "§ 2 Treuhandvermoegen",
                        [
                            "Eingelegt wird ein Anfangsvermoegen von "
                            "12,5 Mio. EUR per 14.06.2018, jaehrliche "
                            "Zustiftungen nach Beschluss der Treugeberin "
                            "(Investment Committee, Vorsitz CFO Laura Bauer).",
                            "Per 31.12.2023 betraegt das Treuhandvermoegen "
                            "32,6 Mio. EUR (Marktwerte).",
                        ],
                    ),
                    (
                        "§ 3 Anlagerichtlinie",
                        [
                            "Anleihen Investment Grade (45-65 %), Aktien Industrielaender "
                            "(20-35 %), Emerging Markets (0-10 %), Immobilien (5-12 %), "
                            "Geldmarkt (0-10 %), Alternatives (0-5 %).",
                            "Verwahrstelle: State Street Bank International GmbH, "
                            "Muenchen. Asset Manager: Allianz Global Investors "
                            "(50 %), DWS (35 %), Pictet AM (15 %).",
                        ],
                    ),
                    (
                        "§ 4 Sicherungsfall",
                        [
                            "Sicherungsfaelle gemaess § 7 BetrAVG-Definition; "
                            "im Sicherungsfall geht der Anspruch an den "
                            "Pensions-Sicherungs-Verein a.G. (PSVaG) ueber, der "
                            "das Treuhandvermoegen zur Erfuellung der Pensions"
                            "verpflichtungen verwendet.",
                        ],
                    ),
                    (
                        "§ 5 Kontrolle und Berichtspflicht",
                        [
                            "Quartalsberichte des Treuhaenders an die "
                            "Treugeberin (CTA Investment Report). Jaehrliche "
                            "Pruefung durch den Konzernabschlusspruefer "
                            "(KPMG AG WPG).",
                        ],
                    ),
                ],
            ),
        ),
        (
            "Unterschriften",
            signatures(
                "Anna Mueller", "CEO", R["name"],
                "Markus Pflanzer", "Vorsitzender Brennhagen Pension Trust e.V.", "Brennhagen Pension Trust e.V.",
                place="Stuttgart", date_str_="14.06.2018",
            ),
        ),
    ]
    write_doc(path, H, title, sections, subtitle=subtitle, confidential=True)


def write_pension_risk():
    path = BASE / "REA_Pensionsrisiko_Analyse_2023.docx"
    title = "Pensionsrisiko-Analyse 2023 – ALM-Studie und Stresstest"
    subtitle = (
        "Asset-Liability-Matching  |  Stichtag 31.12.2023  |  "
        "Erstellt fuer Investment Committee und Pruefungsausschuss"
    )
    sections = [
        (
            "1. Management Summary",
            (
                "Die Nettoverpflichtung aus leistungsorientierten Pensionsplaenen "
                "betraegt zum 31.12.2023 45,8 Mio. EUR (DBO 78,4 Mio. EUR ./. "
                "Plan-Assets 32,6 Mio. EUR). Die Ausfinanzierungsquote (Funding "
                "Ratio) liegt damit bei 41,6 % und ist gegenueber Vorjahr "
                "(38,9 %) leicht verbessert.\n\n"
                "Hauptrisiken: Zinsaenderungsrisiko (DBO-Duration 14,2 Jahre), "
                "Langlebigkeitsrisiko, Inflationsrisiko (Rententrend), "
                "Asset-Risiko Equity-Quote 34 %."
            ),
        ),
        (
            "2. Sensitivitaeten (DBO 78,4 Mio. EUR)",
            [
                ["Stressszenario", "DBO-Wirkung (Mio. EUR)", "Funding Ratio neu"],
                ["Diskontsatz -100 bp (von 3,40 auf 2,40 %)", "+12,8", "35,7 %"],
                ["Diskontsatz +100 bp", "-10,4", "48,0 %"],
                ["Inflation/Rententrend +50 bp", "+4,8", "39,1 %"],
                ["Lebenserwartung +1 Jahr", "+3,0", "40,0 %"],
                ["Equity-Markt -30 %", "Asset -3,3 Mio.", "37,4 %"],
                ["Stagflation (Zins -50bp, Infl. +100bp)", "+11,2", "36,4 %"],
            ],
        ),
        (
            "3. Asset-Liability-Matching",
            (
                "Die Duration der Verbindlichkeiten betraegt 14,2 Jahre, "
                "die Duration der Plan-Assets 8,7 Jahre - es besteht eine "
                "Duration Gap von 5,5 Jahren. Empfehlung des ALM-Beraters "
                "(BlackRock LDI-Team): Aufstockung des Long-Duration-Bond-"
                "Buckets um 4,5 Mio. EUR oder Einsatz von synthetischen "
                "Duration-Overlays (Interest-Rate-Swaps Receive-Fix 30Y).\n\n"
                "Die Entscheidung wird im Investment Committee im Q2-2024 "
                "diskutiert."
            ),
        ),
        (
            "4. Empfehlungen",
            (
                "list",
                [
                    "Zustiftung 5 Mio. EUR in den Pension Trust e.V. in 2024 (CFO-Vorlage).",
                    "Pruefung Receive-Fix Swap 30Y zur Duration-Verlaengerung.",
                    "Eroeffnung Diskussion mit PSVaG zur Anpassung Beitragsklasse.",
                    "Annual ALM-Review durch Heubeck AG bestaetigen (Mandat verlaengern).",
                ],
            ),
        ),
        (
            "5. Genehmigungen",
            signatures(
                "Laura Bauer", "CFO / Vorsitz Investment Committee", R["name"],
                "Markus Pflanzer", "Group Treasurer / Pensionsbeauftragter", R["name"],
                place="Stuttgart", date_str_="22.02.2024",
            ),
        ),
    ]
    write_doc(path, H, title, sections, subtitle=subtitle, confidential=True)


def write_vw_ecr():
    path = BASE / "REA_VW_ADAS-V4D_ECR_2023_003.docx"
    title = "VW Engineering Change Request ECR-2023-003 – ADAS-V4D Radarfusion"
    subtitle = "Kunde Volkswagen AG  |  Bauteil ADAS-V4D-RF  |  ECR-Nr. VW-2023-PE-3318"
    sections = [
        (
            "1. Aenderungsanlass",
            (
                "Volkswagen AG fordert eine Software-Anpassung am ADAS-V4D-RF "
                "(Radar-Fusion-Steuergeraet) zur Erfuellung der UN-R 157 "
                "Anforderungen (Automated Lane Keeping System ALKS, Level 3, "
                "bis 130 km/h) fuer die Volkswagen ID.7-Plattform 2024 MJ."
            ),
        ),
        (
            "2. Technische Aenderungen",
            [
                ["Komponente", "Aenderung", "Stufe"],
                ["AUTOSAR Sensor-Fusion-Stack", "Update 4.4 -> 4.7", "Major"],
                ["Funktionale Sicherheit", "ASIL-D Pfad fuer Lane-Keep-Logik", "Major"],
                ["Cybersecurity", "Zertifizierung nach ISO/SAE 21434", "—"],
                ["Hardware", "Keine Aenderung (Identitaet zu Mercedes-Variante)", "—"],
                ["Pruefumfang", "Erweiterung NCAP-Szenarien (Euro NCAP 2024)", "—"],
            ],
        ),
        (
            "3. bAV-Bezug",
            (
                "Diese ECR ist faelschlicherweise im Pensions-Ordner abgelegt. "
                "Korrekter Ablageort waere 25_OEM_Beziehungen/VW oder "
                "06_Engineering_Change/ADAS. "
                "Bitte vom Document-Owner (Dr. Petra Hollmann, CTO) verschieben."
            ),
        ),
        (
            "4. Zeitschiene und Kommerzielles",
            (
                "ECR-Approval VW: erwartet 30.06.2023. Software-Freigabe: "
                "Q3-2023. PPAP: 15.10.2023. SOP ID.7 mit Level-3-ALKS: "
                "01.03.2024. Stueckpreisaenderung: nicht vorgesehen "
                "(Software-only-Update).\n\n"
                "Engineering-Aufwand: 480 Personentage RSG Muenchen plus "
                "120 Personentage Tier-2-Lieferant (Vector Informatik fuer "
                "AUTOSAR-Stack). Geschaetzte Einmal-NRE-Kosten: 1,8 Mio. EUR, "
                "die VW vertraglich gemaess Rahmenliefervertrag § 14 (R&D "
                "Cost Sharing) zu 60 % refinanziert (1,08 Mio. EUR Erloes "
                "in Q4-2023 erwartet)."
            ),
        ),
        (
            "5. Risikobewertung",
            [
                ["Risiko", "Wahrscheinlichkeit", "Auswirkung", "Mitigation"],
                ["Verzoegerung Cybersecurity-Zertifizierung", "Mittel", "Hoch (SOP-Risiko)", "TUV Sued frueh einbinden, Audit-Termin gesichert 12.07.2023"],
                ["ASIL-D-Nachweis nicht eindeutig", "Niedrig", "Sehr hoch", "Externer FuSi-Auditor Exida bereits beauftragt"],
                ["AUTOSAR 4.7 Reifegrad", "Mittel", "Mittel", "Pilotierung auf ECU-900 vorgezogen"],
                ["VW-internes Genehmigungsverfahren", "Niedrig", "Mittel", "Wochentliches Steering mit VW PE-Lead"],
            ],
        ),
        (
            "6. Freigaben",
            signatures(
                "Dr. Klaus Kessler", "Werkleiter RSG Muenchen", "Brennhagen Software GmbH",
                "Lars Wittmann", "Lead Developer ADAS", "Brennhagen Software GmbH",
                place="Muenchen", date_str_="14.04.2023",
            ),
        ),
    ]
    write_doc(path, H, title, sections, subtitle=subtitle)


def write_br_2021_07():
    path = BASE / "REG_BR_Protokoll_2021_07.docx"
    title = "Betriebsrat REG Heilbronn – Sitzungsprotokoll Juli 2021"
    subtitle = "Sitzung am 15.07.2021, 13:30-16:45 Uhr, Konferenzraum Werk Heilbronn, Haus C"
    sections = [
        (
            "Teilnehmer",
            (
                "list",
                [
                    "Klaus Bauer (Betriebsratsvorsitzender, stv. Konzern-BR)",
                    "Heinz Mayer (stv. BR-Vorsitzender)",
                    "Sabine Brand (Q-Leitung, eingeladen TOP 2)",
                    "Andreas Maier (Werkleiter REG, eingeladen TOP 3)",
                    "Marlies Duerr (Konzern-BR-Vorsitzende, Gast)",
                    "9 weitere Betriebsratsmitglieder (siehe Anwesenheitsliste)",
                ],
            ),
        ),
        (
            "TOP 1 – Eroeffnung und Genehmigung Vorprotokoll",
            (
                "Vorsitzender Klaus Bauer eroeffnet die Sitzung um 13:32 Uhr. "
                "Das Protokoll der Juni-Sitzung (17.06.2021) wird ohne "
                "Aenderungen einstimmig genehmigt."
            ),
        ),
        (
            "TOP 2 – Entgeltumwandlung bAV: Anpassung Beitragssaetze 2022",
            (
                "Frau Brand erlaeutert den Vorschlag der Geschaeftsleitung, "
                "den Arbeitgeber-Zuschuss zur Entgeltumwandlung von 15 % auf "
                "20 % anzuheben (gesetzlicher Mindest-Zuschuss § 1a Abs. 1a "
                "BetrAVG i.H.v. 15 %).\n\n"
                "Der Betriebsrat begruesst die Anhebung ausdruecklich und "
                "fordert zusaetzlich eine Anhebung der Hoechstgrenze fuer "
                "Entgeltumwandlung von derzeit 4 % BBG-West auf 6 % BBG-West "
                "(was etwa 350 EUR/Monat Mehrbeitrag entspricht).\n\n"
                "Beschluss: Der Betriebsrat erteilt der Geschaeftsleitung "
                "Verhandlungsmandat. Naechste Verhandlungsrunde 23.08.2021. "
                "Einstimmig (12 dafuer, 0 dagegen)."
            ),
        ),
        (
            "TOP 3 – Sondervereinbarung Pension Trust Zustiftung 2021",
            (
                "Werkleiter Maier informiert ueber die geplante Zustiftung der "
                "Konzernholding in den Pension Trust e.V. i.H.v. 4,5 Mio. EUR "
                "im Q4-2021. Der Betriebsrat nimmt dies positiv zur Kenntnis "
                "und bittet um schriftliche Bestaetigung der "
                "Vermoegenseintragung an die Vertrauensperson "
                "(Marlies Duerr, Konzern-BR)."
            ),
        ),
        (
            "TOP 4 – Sonstiges",
            (
                "list",
                [
                    "Urlaubsplanung 2022 - Versand Listen bis 31.10.2021.",
                    "BAV-Sprechtag am 14.09.2021, 14-17 Uhr, Kantine - alle eingeladen.",
                    "Naechste Sitzung: 19.08.2021, 13:30 Uhr.",
                ],
            ),
        ),
        (
            "Unterschriften",
            signatures(
                "Klaus Bauer", "BR-Vorsitzender REG Heilbronn", "Brennhagen Elektronik GmbH",
                "Heinz Mayer", "stv. BR-Vorsitzender", "Brennhagen Elektronik GmbH",
                place="Heilbronn", date_str_="15.07.2021",
            ),
        ),
    ]
    write_doc(path, H, title, sections, subtitle=subtitle)


def write_ic_rechnung():
    path = BASE / "REG_IC_Rechnung_2022_07_FINAL_v2.docx"
    title = "Konzern-Innenrechnung Juli 2022 – Pensions-Service-Charge REG -> REA"
    subtitle = "Intercompany-Rechnung IC-2022-07-PEN  |  Pensionsadministration Konzernumlage"
    sections = [
        (
            "Rechnungssteller",
            (
                f"{R['name']} (Konzernholding)\n"
                f"{R['addr']}\n"
                f"USt-IdNr. DE 287445190\n"
                f"{R['hrb']}\n\n"
                f"Empfaenger: Brennhagen Elektronik GmbH (REG), "
                f"Industriestrasse 12, 74076 Heilbronn\n"
                f"USt-IdNr. DE 215889134"
            ),
        ),
        (
            "Leistungsbeschreibung",
            [
                ["Pos.", "Leistung", "Menge", "Einzelpreis", "Summe"],
                ["1", "Pensionsadministration zentral (Anteil REG)", "1 Monat", "8.400,00 EUR", "8.400,00 EUR"],
                ["2", "Aktuargutachten-Anteil Heubeck (Umlage 7/12)", "1 Anteil", "3.150,00 EUR", "3.150,00 EUR"],
                ["3", "PSVaG-Beitrag (Umlage)", "1 Anteil", "4.270,00 EUR", "4.270,00 EUR"],
                ["4", "CTA-Treuhand-Administration", "1 Anteil", "1.200,00 EUR", "1.200,00 EUR"],
                ["", "Zwischensumme netto", "", "", "17.020,00 EUR"],
                ["", "USt 19 %", "", "", "3.233,80 EUR"],
                ["", "Rechnungsbetrag", "", "", "20.253,80 EUR"],
            ],
        ),
        (
            "Verrechnungsgrundlage",
            (
                "Konzern-Service-Vertrag (Master Services Agreement) zwischen "
                "Brennhagen Elektronik AG und Brennhagen Elektronik GmbH vom "
                "01.01.2020. Verrechnung gemaess Anlage 2 (Pensionsservices). "
                "Kostenaufteilung nach Headcount-Schluessel (REG-Anteil 49,1 %)."
            ),
        ),
        (
            "Zahlungsbedingungen",
            (
                "Faellig: 30 Tage netto ab Rechnungseingang. "
                "Verrechnung ueber Intercompany-Verrechnungskonto IC-1810 "
                "(REA-REG). Bei Rueckfragen: Florian Maier (Group Controller) "
                "oder Hauke Schroeder (REG Accounting). "
                "Rechnungs-Nr. IC-2022-07-PEN-REG, Datum 31.07.2022."
            ),
        ),
        (
            "Transfer Pricing Hinweis",
            (
                "Die Verrechnung erfolgt nach dem Cost-Plus-Verfahren mit "
                "einem Aufschlag von 5,0 % auf direkte Kosten, in Ueberein"
                "stimmung mit dem Konzern-Transfer-Pricing-Leitfaden 2022 "
                "(Master File / Local File, OECD Verrechnungspreisleitlinien "
                "2022). Die Funktionsanalyse stuft die Pensionsadministration "
                "als Routinedienstleistung (Low-Value-Adding Intra-Group "
                "Service gemaess Anhang IV BEPS Action 8-10) ein.\n\n"
                "Dokumentation: TP_LocalFile_REG_2022.docx. Verantwortlich: "
                "Dr. Heike Berger (Group Tax)."
            ),
        ),
        (
            "Buchungsanweisung",
            (
                "REA-seitig: Ertrag SKR 8338 (Konzernumlage Erloese) gegen "
                "Forderung IC-Konto 1810. REG-seitig: Aufwand SKR 6815 "
                "(Konzernumlage Personalverwaltung) gegen Verbindlichkeit "
                "IC-Konto 1810. Verbuchung erfolgt vor Monatsabschluss "
                "31.07.2022 durch beide Gesellschaften."
            ),
        ),
    ]
    write_doc(path, H, title, sections, subtitle=subtitle)


def write_rho_br_2022_03():
    path = BASE / "RHO_BR_Protokoll_2022_03.docx"
    title = "Betriebsrat RHO Stuttgart – Sitzungsprotokoll Maerz 2022"
    subtitle = "Sitzung 14.03.2022, 14:00-16:30 Uhr, Konzernzentrale Stuttgart, Raum Stuttgart-IV"
    sections = [
        (
            "Teilnehmer",
            (
                "list",
                [
                    "Tobias Vetter (BR-Vorsitzender RHO)",
                    "Petra Koch (stv. BR-Vorsitzende)",
                    "5 weitere BR-Mitglieder (siehe Anwesenheitsliste)",
                    "Sandra Berger (Head of HR Group, Gast TOP 2)",
                    "Markus Pflanzer (Pensionsbeauftragter, Gast TOP 3)",
                ],
            ),
        ),
        (
            "TOP 1 – Eroeffnung / Vorprotokoll",
            (
                "Vorsitzender Vetter eroeffnet die Sitzung. Vorprotokoll "
                "(Februar) wird mit redaktionellen Aenderungen genehmigt "
                "(7 dafuer, 0 dagegen)."
            ),
        ),
        (
            "TOP 2 – Mitarbeiter-Information neue Direktzusagen 2022",
            (
                "Frau Berger informiert ueber 3 neue Vorstands-Direktzusagen "
                "in 2022 (Hollmann, Tanaka, Richter). Der Betriebsrat erhaelt "
                "Einblick in die Vertragsmustertexte unter Wahrung der "
                "Persoenlichkeitsrechte. Keine Mitbestimmungspflicht, da "
                "leitende Angestellte / Vorstand."
            ),
        ),
        (
            "TOP 3 – CTA-Zustiftung 2022 und PSVaG-Beitrag",
            (
                "Pflanzer berichtet: Geplante CTA-Zustiftung 5,0 Mio. EUR im "
                "Q4-2022, PSVaG-Beitrag 2022 wird aufgrund gestiegener "
                "Insolvenzen vorraussichtlich von 1,8 promille auf 2,4 promille "
                "ansteigen (zusaetzliche Kostenbelastung Konzern ca. 110 TEUR).\n\n"
                "Der Betriebsrat nimmt die Information zur Kenntnis."
            ),
        ),
        (
            "TOP 4 – Sonstiges",
            (
                "list",
                [
                    "BR-Klausur 11.-12.05.2022 in Bad Boll (Themen: § 87 BetrVG, KI am Arbeitsplatz).",
                    "Naechste Sitzung 04.04.2022.",
                ],
            ),
        ),
        (
            "Unterschriften",
            signatures(
                "Tobias Vetter", "BR-Vorsitzender RHO", "Brennhagen Holding GmbH",
                "Petra Koch", "stv. BR-Vorsitzende RHO", "Brennhagen Holding GmbH",
                place="Stuttgart", date_str_="14.03.2022",
            ),
        ),
    ]
    write_doc(path, H, title, sections, subtitle=subtitle)


# ----------------------------------------------------------------------------
# Jahresinfos
# ----------------------------------------------------------------------------

JAHRESINFOS = [
    (1, "Laura Bauer", True),
    (2, "Claudia Fischer", False),
    (3, "Petra Koch", False),
    (4, "Sabine Schreiber", False),
    (5, "Monika Hoffmann", False),
    (6, "Julia Hartmann", False),
    (7, "Sandra Weber", False),
    (8, "Nicole Lange", False),
    (9, "Karin Schulz", False),
    (10, "Dr. Elisabeth Vogel", False),
    (11, "Prof. Renate Meier", False),
    (12, "Christina Wolf", False),
    (13, "Birgit Braun", False),
    (14, "Ursula Krause", False),
    (15, "Ingrid Schmidt", False),
    (16, "Heike Richter", False),
    (17, "Ute Kraemer", False),
    (18, "Gisela Roth", False),
    (19, "Thomas Mueller", False),
    (20, "Klaus Weber", False),
    (21, "Stefan Hoffmann", False),
    (22, "Michael Richter", False),
    (23, "Andreas Becker", False),
    (24, "Christian Koch", False),
    (25, "Markus Schneider", True),
    (26, "Frank Werner", False),
    (27, "Juergen Lange", False),
    (28, "Hans-Peter Braun", False),
    (29, "Dr. Wolfgang Fischer", False),
    (30, "Ralf Zimmermann", True),
]

JAHRESINFO_FILES = [
    "bAV_Jahresinfo_2023_001_Laura_Bauer-P1_ENTWURF.docx",
    "bAV_Jahresinfo_2023_002_Claudia_Fischer.docx",
    "bAV_Jahresinfo_2023_003_Petra_Koch-P3.docx",
    "bAV_Jahresinfo_2023_004_Sabine_Schreibe_rev_SRichter.docx",
    "bAV_Jahresinfo_2023_005_Monika_Hoffmann_FINAL.docx",
    "bAV_Jahresinfo_2023_006_Julia_Hartmann-.docx",
    "bAV_Jahresinfo_2023_007_Sandra_Weber-P7.docx",
    "bAV_Jahresinfo_2023_008_Nicole_Lange-P8.docx",
    "bAV_Jahresinfo_2023_009_Karin_Schulz-P9_v1.docx",
    "bAV_Jahresinfo_2023_010_Dr._Elisabeth_V.docx",
    "bAV_Jahresinfo_2023_011_Prof._Renate_Me_rev_SRichter.docx",
    "bAV_Jahresinfo_2023_012_Christina_Wolf-.docx",
    "bAV_Jahresinfo_2023_013_Birgit_Braun-P1.docx",
    "bAV_Jahresinfo_2023_014_Ursula_Krause-P.docx",
    "bAV_Jahresinfo_2023_015_Ingrid_Schmidt-.docx",
    "bAV_Jahresinfo_2023_016_Heike_Richter-P.docx",
    "bAV_Jahresinfo_2023_017_Ute_Krämer-P17.docx",
    "bAV_Jahresinfo_2023_018_Gisela_Roth-P18.docx",
    "bAV_Jahresinfo_2023_019_Thomas_Müller-P.docx",
    "bAV_Jahresinfo_2023_020_Klaus_Weber-P20.docx",
    "bAV_Jahresinfo_2023_021_Stefan_Hoffmann.docx",
    "bAV_Jahresinfo_2023_022_Michael_Richter_ALT.docx",
    "bAV_Jahresinfo_2023_023_Andreas_Becker-.docx",
    "bAV_Jahresinfo_2023_024_Christian_Koch-.docx",
    "bAV_Jahresinfo_2023_025_Markus_Schneide_DRAFT.docx",
    "bAV_Jahresinfo_2023_026_Frank_Werner-P2.docx",
    "bAV_Jahresinfo_2023_027_Jürgen_Lange-P2.docx",
    "bAV_Jahresinfo_2023_028_Hans-Peter_Brau.docx",
    "bAV_Jahresinfo_2023_029_Dr._Wolfgang_Fi.docx",
    "bAV_Jahresinfo_2023_030_Ralf_Zimmermann_WIP.docx",
]


def main() -> None:
    # Aktuarberichte
    for fn, jahr, plan in AKTUAR_FILES:
        aktuargutachten(jahr, plan, BASE / fn)

    # Einzelzusagen
    for fn, name, position, gebd, jahr, monat in EINZELZUSAGEN:
        direktzusage(name, position, gebd, jahr, monat, BASE / fn)

    # Versorgungsplaene / IAS-19 Disclosures
    for fn, werk, planname, dbo in VERSORGUNGSPLAENE:
        ias19_disclosure(planname, f"Versorgungsordnung des Werks/Plans {werk} ({planname}). Geltungsbereich gemaess Plan-Statuten.", dbo, dbo * 0.42, BASE / fn)

    # Jahresinfos
    for fn, (num, name, draft) in zip(JAHRESINFO_FILES, JAHRESINFOS):
        jahresinfo(num, name, BASE / fn, draft=draft)

    # Special
    write_it_change_request()
    write_mr_rpl()
    write_bmw_ecr()
    write_cta_treuhand()
    write_pension_risk()
    write_vw_ecr()
    write_br_2021_07()
    write_ic_rechnung()
    write_rho_br_2022_03()

    # Verification
    from docx import Document
    total = thin = 0
    thin_list = []
    for p in BASE.rglob("*.docx"):
        d = Document(p)
        t = " ".join(par.text for par in d.paragraphs)
        for tbl in d.tables:
            for r in tbl.rows:
                for c in r.cells:
                    t += " " + c.text
        w = len(t.split())
        total += 1
        if w < 200:
            thin += 1
            thin_list.append((p.name, w))
    print(f"{total} total, {thin} still thin")
    for n, w in thin_list:
        print(f"  THIN {w:4d}  {n}")


if __name__ == "__main__":
    main()
