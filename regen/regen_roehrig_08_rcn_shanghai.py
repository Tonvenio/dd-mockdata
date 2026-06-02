"""Brennhagen (Shanghai) Co. Ltd. – Folder 08_Tochter_CN_Shanghai
Re-generates all 57 thin .docx files in this folder with realistic content.

Country Manager: Mr. Zhang Hao. Finance Manager: Liang Wei.
Registration USCC: 91310000789123456X. FIE (Foreign-Invested Enterprise).
CNY currency, SAFE Devisenkontrolle, chinesische Sozialversicherung, FIE-Steuer.
"""
# --- portable-paths-prelude --- (do not edit) ---
import sys
from pathlib import Path as _PathlibPath
_RP = _PathlibPath(__file__).resolve().parent
while not (_RP / "enhance_lib.py").exists() and _RP.parent != _RP:
    _RP = _RP.parent
_ROOT = _RP
sys.path.insert(0, str(_ROOT))
# --- end prelude ---
from enhance_lib import ROEHRIG_AG as R, write_doc, signatures

BASE = f"{_ROOT}/roehrig_large/08_Tochter_CN_Shanghai"

# Headers
H_RCN = {
    "name": "Brennhagen (Shanghai) Co. Ltd. / 罗瑞格(上海)有限公司",
    "addr": "Building 7, No. 1188 Lianhang Road, Minhang District, 201112 Shanghai, PRC",
    "hrb":  "USCC 91310000789123456X  |  FIE-WFOE  |  Reg.-Kap. 32 Mio. CNY",
}
H_RCN_SHORT = {
    "name": "Brennhagen Electronics (China) Ltd.",
    "addr": "Shanghai",
    "hrb":  "USCC 91310115MA1FL42Q38  |  FIE-WFOE",
}
H_REA = {"name": R["name"], "addr": R["addr"], "hrb": R["hrb"]}
H_RHO = {
    "name": "Brennhagen Holding GmbH",
    "addr": "Vaihinger Strasse 120, 70567 Stuttgart",
    "hrb":  "HRB 726450, Amtsgericht Stuttgart",
}
H_RSG = {
    "name": "Brennhagen Software GmbH",
    "addr": "Landsberger Strasse 110, 80339 Muenchen",
    "hrb":  "HRB 319872, Amtsgericht Muenchen",
}
H_KPMG = {
    "name": "KPMG AG Wirtschaftspruefungsgesellschaft",
    "addr": "Klingelhoeferstrasse 18, 10785 Berlin",
    "hrb":  "WPK-Nr. 2063  |  Lead Partner Dr. Maximilian Brand",
}

CM = "Mr. Zhang Hao"
FIN = "Ms. Liang Wei"
HR_CN = "Ms. Chen Xiu"
PROD_CN = "Mr. Wang Tao"
QA_CN = "Ms. Sun Mei"

CONTEXT_PARA = (
    "Brennhagen (Shanghai) Co. Ltd. (»RCN«, 罗瑞格上海有限公司) ist eine 100%ige Tochter "
    "der Brennhagen Holding GmbH, Stuttgart, und Teil des Konzerns Brennhagen Elektronik AG "
    "(Prime Standard, WKN RHGRP1). RCN ist als Wholly Foreign-Owned Enterprise (WFOE) "
    "im chinesischen Handelsregister Shanghai eingetragen (USCC 91310000789123456X), "
    "registriertes Kapital 32 Mio. CNY, voll eingezahlt. Belegschaft per 31.12.2023: "
    "320 Mitarbeiter. Geschaeftsfelder: (i) Vertrieb und Application Engineering der "
    "Brennhagen-Produktfamilien ICP-3, BMS-12, ADAS-V4D, ECU-900, LightCtrl-7 fuer "
    "chinesische OEM-Werke (Mercedes-Benz Beijing, Stellantis Wuhan, BMW Brilliance "
    "Shenyang, FAW-VW, SAIC, NIO, Li Auto), (ii) Aftermarket China und Suedost-Asien, "
    "(iii) Local Sourcing/Lieferantenmanagement CNY-Region. Country Manager: "
    f"{CM}. Finance Manager: {FIN}."
)

# ─────────────────────────────────────────────────────────────────────────────
# 1) ARBEITSVERTRAEGE (4) – Mandarin + Deutsch/Englisch
# ─────────────────────────────────────────────────────────────────────────────
def arbeitsvertrag(fname, position_de, position_zh, mitarbeiter, monatsgehalt_cny,
                   bonus_pct, urlaubstage, start_date, sv_kasse):
    write_doc(f"{BASE}/{fname}", H_RCN,
        f"Arbeitsvertrag / Labor Contract / 劳动合同 – {position_de}",
        subtitle=f"Brennhagen (Shanghai) Co. Ltd. – Mr./Ms. {mitarbeiter} – Wirksam ab {start_date}",
        sections=[
            ("Vertragsparteien / Parties / 合同双方",
             f"Arbeitgeber: Brennhagen (Shanghai) Co. Ltd. (»Gesellschaft«), Building 7, "
             f"No. 1188 Lianhang Road, Minhang District, 201112 Shanghai, P.R. China, "
             f"USCC 91310000789123456X, vertreten durch den Country Manager (»Legal "
             f"Representative«) {CM}.\n\n"
             f"Arbeitnehmer/in: Mr./Ms. {mitarbeiter}, Wohnsitz Shanghai, "
             f"Hukou-Registrierung und Ausweis-Nr. liegen der Gesellschaft vor. "
             f"Position: {position_de} / {position_zh}.\n\n"
             "Der Vertrag wird in chinesischer Originalfassung (Mandarin) sowie in einer "
             "deutsch/englischen Uebersetzung geschlossen. Im Konfliktfall ist die "
             "Mandarin-Fassung verbindlich (gemaess Art. 11 PRC Labor Contract Law)."),
            ("Vertragsregelungen",
             ("clauses", [
                 ("§ 1 Arbeitsort & Aufgaben / Position & Duties", [
                     f"Arbeitsort: Hauptsitz Shanghai-Minhang, Building 7, No. 1188 "
                     f"Lianhang Road. Reisetaetigkeit innerhalb VR China sowie nach "
                     f"Deutschland (Stuttgart, Heilbronn) bis zu 40 Tage p. a.",
                     f"Funktionsbezeichnung: {position_de} (chinesisch: {position_zh}). "
                     f"Disziplinarisch und fachlich berichtet die Stelle an den Country "
                     f"Manager {CM}; fuer das Group-Reporting ggf. dotted-line an die "
                     f"jeweilige Konzernfunktion in Stuttgart.",
                     "Stellenbeschreibung gemaess Anlage 1 (»Job Description / 岗位说明书«), "
                     "die wesentlicher Vertragsbestandteil ist.",
                 ]),
                 ("§ 2 Vertragsdauer / Term / 合同期限", [
                     f"Befristung: 3 Jahre (Erstvertrag), beginnend am {start_date}. "
                     "Verlaengerung nach Massgabe Art. 14 PRC Labor Contract Law "
                     "(nach zwei befristeten Vertraegen ist ein unbefristeter Vertrag "
                     "abzuschliessen, sofern beide Parteien zustimmen).",
                     "Probezeit (Trial Period / 试用期): 6 Monate gemaess Art. 19 PRC "
                     "LCL (max. zulaessig fuer Vertraege ueber 3 Jahre).",
                 ]),
                 ("§ 3 Verguetung / Compensation / 薪酬", [
                     f"Monatsgrundgehalt: {monatsgehalt_cny:,} CNY brutto (vor "
                     f"Individual Income Tax IIT und Sozialversicherungs-Beitraegen), "
                     "zahlbar zum 10. Bankarbeitstag des Folgemonats per Bank-Transfer "
                     "auf das vom Arbeitnehmer benannte CNY-Konto bei der Bank of "
                     "Communications, Shanghai.",
                     f"13. Monatsgehalt (Annual Bonus / 年终奖金) in Hoehe von einem "
                     f"Monatsgrundgehalt, faellig im Februar des Folgejahres (vor "
                     "Spring Festival), pro-rata bei unterjaehrigem Ein-/Austritt.",
                     f"Variable Erfolgsbeteiligung (STI / 绩效奖金): Zielwert "
                     f"{bonus_pct}% des Jahresgrundgehalts, Auszahlung 0–150% in "
                     "Abhaengigkeit von (a) Konzern-EBIT, (b) RCN-Umsatz CNY, (c) "
                     "individuellen MbO-Zielen (Mercedes/Stellantis/BMW-Brilliance-"
                     "Programmziele). Bonus-Festsetzung durch Country Manager nach "
                     "Group-CHRO-Abstimmung.",
                     "Gehaltsanpassung: jaehrliches Review per 1. April; "
                     "Inflationsausgleich CPI Shanghai sowie individuelle Performance.",
                 ]),
                 ("§ 4 Sozialversicherung & Housing Fund / 社保和公积金", [
                     f"Arbeitgeber meldet den Arbeitnehmer fristgerecht zu allen "
                     f"fuenf Pflicht-Sozialversicherungen (五险) sowie zum Housing "
                     f"Provident Fund (HPF / 住房公积金) bei der zustaendigen Behoerde "
                     f"{sv_kasse} an. Beitraege werden anteilig gemaess Shanghai-"
                     "Saetzen 2023 abgefuehrt (Pension 16% AG / 8% AN; Krankenversicherung "
                     "9.5% / 2%; Arbeitslosen 0.5% / 0.5%; Arbeitsunfall 0.16–1.52%; "
                     "Mutterschaft 1%; HPF 7% / 7%).",
                     "Beitragsbemessungsgrenze gemaess SHHRSS-Bekanntmachung 2023 "
                     "(Untergrenze 7.310 CNY / Obergrenze 36.549 CNY pro Monat).",
                 ]),
                 ("§ 5 Arbeitszeit & Urlaub / Working Hours & Leave", [
                     "Regelarbeitszeit: 40 Std/Woche, Montag bis Freitag, "
                     "8:30 – 17:30 Uhr (1 Std. Mittagspause). Ueberstunden gemaess "
                     "Art. 41 PRC LCL nur mit Zustimmung und Vergutung 150 / 200 / "
                     "300% je nach Wochentag/Feiertag.",
                     f"Jahresurlaub: {urlaubstage} Arbeitstage (gesetzl. Mindesturlaub "
                     "5–15 Tage je nach Betriebszugehoerigkeit + freiwillige "
                     "RCN-Mehrleistung). Zusaetzlich alle gesetzlichen Feiertage "
                     "(Spring Festival, Qingming, Labor Day, Dragon Boat, "
                     "Mid-Autumn, National Day).",
                     "Krankheitsurlaub (Sick Leave / 病假) nach Shanghai Local Rule "
                     "(min. 60% Gehalt, Dauer abhaengig von Betriebszugehoerigkeit).",
                 ]),
                 ("§ 6 Vertraulichkeit & Non-Compete / 保密与竞业限制", [
                     "Vertraulichkeit ueber Geschaeftsgeheimnisse, Source-Code, "
                     "Kunden-/Lieferantenbeziehungen, Preise und alle nicht-"
                     "oeffentlich bekannten Informationen waehrend der Vertragsdauer "
                     "und 5 Jahre nach Beendigung.",
                     "Nachvertragliches Wettbewerbsverbot (Non-Compete) max. 24 Monate "
                     "gemaess Art. 24 PRC LCL; Kompensation 30% des letzten "
                     "Monatsgehalts pro Monat der Karenzzeit. Geltungsbereich: "
                     "Wettbewerber im Bereich Automotive-Elektronik in der VR China "
                     "(insb. Continental China, Bosch China, Aptiv Asia, Huawei "
                     "Automotive, Desay SV, Joyson Electronics).",
                 ]),
                 ("§ 7 Beendigung / Termination", [
                     "Beendigung gemaess Art. 36–44 PRC LCL. Ordentliche Kuendigung "
                     "mit 30 Tagen Schriftform-Frist; ausserordentliche Kuendigung "
                     "bei schwerwiegender Pflichtverletzung. Abfindung (Severance / "
                     "经济补偿金) gemaess Art. 47 (1 Monatsgehalt pro Beschaeftigungsjahr, "
                     "max. 3x lokaler Durchschnittslohn / max. 12 Jahre).",
                 ]),
                 ("§ 8 Anwendbares Recht / Applicable Law", [
                     "Es gilt das Recht der Volksrepublik China unter Ausschluss "
                     "kollisionsrechtlicher Verweisungsnormen. Gerichtsstand: "
                     "Shanghai Pudong Labor Dispute Arbitration Commission "
                     "(上海市浦东新区劳动人事争议仲裁委员会); fuer nachgelagerte "
                     "Gerichtsverfahren das Volksgericht des Bezirks Minhang.",
                     "Salvatorische Klausel und Aenderungen der Schriftform; "
                     "Anlagen Bestandteil dieses Vertrages: Job Description, "
                     "Compensation Sheet, IT/IP-Richtlinie, Code of Conduct.",
                 ]),
             ])),
            ("Unterschriften / Signatures / 签字",
             signatures(CM, "Country Manager / Legal Representative",
                        "Brennhagen (Shanghai) Co. Ltd.",
                        mitarbeiter, position_de, "Arbeitnehmer/in / Employee",
                        place="Shanghai", date_str_=start_date)),
        ])

arbeitsvertrag("RCN_Arbeitsvertrag_01_Geschäftsführer_in_R_2022.docx",
    "Country Manager / General Manager", "总经理",
    "Zhang Hao", 95000, 40, 25, "1. April 2022",
    "Shanghai Human Resources and Social Security Bureau (SHHRSS), Minhang Branch")

arbeitsvertrag("RCN_Arbeitsvertrag_02_Produktionsleiter_in_2022.docx",
    "Operations / Application Engineering Manager", "运营/应用工程经理",
    "Wang Tao", 38000, 20, 18, "1. Juni 2022",
    "SHHRSS Minhang Branch")

arbeitsvertrag("RCN_Arbeitsvertrag_03_Qualitätsmanagerin_S_2022.docx",
    "Quality Manager OEM-China", "OEM中国质量经理",
    "Sun Mei", 32000, 18, 15, "15. Juli 2022",
    "SHHRSS Minhang Branch")

arbeitsvertrag("RCN_Arbeitsvertrag_05_HR-Manager_Shanghai_2022.docx",
    "HR Manager Shanghai", "人力资源经理",
    "Chen Xiu", 28000, 15, 15, "1. September 2022",
    "SHHRSS Minhang Branch")


# ─────────────────────────────────────────────────────────────────────────────
# 2) INTERCOMPANY-RECHNUNGEN (31 Stueck) RSG ebenfalls 1
# ─────────────────────────────────────────────────────────────────────────────
def ic_rechnung(fname, year, month, header=H_RCN_SHORT, sender_label="RCN", betrag_cny=None):
    monate = ["", "Januar", "Februar", "Maerz", "April", "Mai", "Juni",
              "Juli", "August", "September", "Oktober", "November", "Dezember"]
    rg_nr = f"{sender_label}-IC-{year}-{month:02d}-001"
    if betrag_cny is None:
        # variable invoice amount around 4-9 Mio CNY
        base = 5_400_000 + ((year - 2019) * 320_000) + (month * 41_000)
        betrag_cny = base
    eur_kurs = {2019: 7.85, 2020: 7.92, 2021: 7.51, 2022: 7.07, 2023: 7.65}.get(year, 7.65)
    betrag_eur = betrag_cny / eur_kurs

    if sender_label == "RSG":
        leistung = ("Softwareentwicklung Embedded ADAS – Modul-Verifikation und "
                    "ASPICE-Level-3-konforme Testautomatisierung")
        sender_name = "Brennhagen Software GmbH, Muenchen (RSG)"
        empf_name = "Brennhagen Holding GmbH, Stuttgart (RHO)"
        currency = "EUR"; betrag_str = f"{betrag_eur:,.2f} EUR"
        bank_line = "Deutsche Bank AG Muenchen, IBAN DE21 7007 0010 0987 6543 21, BIC DEUTDEMMXXX"
    else:
        leistung = (f"Verkaufsfoerderung und Application Engineering (Service-Charge) "
                    f"sowie Margenabrechnung Lieferungen ICP-3 / BMS-12 / ADAS-V4D / "
                    f"ECU-900 an OEM-Werke China im Berichtsmonat {monate[month]} {year}")
        sender_name = "Brennhagen Elektronik GmbH, Heilbronn (REG)"
        empf_name = "Brennhagen (Shanghai) Co. Ltd. (RCN), 上海"
        currency = "CNY"; betrag_str = f"{betrag_cny:,.2f} CNY (umgerechnet ca. {betrag_eur:,.2f} EUR @ {eur_kurs})"
        bank_line = ("Bank of Communications, Shanghai Branch, "
                     "CNY-Konto 310 066 700 123 456 789  |  SWIFT COMMCNSHSHI  |  "
                     "SAFE-Genehmigung: Service-Trade Allgemeine FX-Lizenz")

    write_doc(f"{BASE}/{fname}", header,
        f"Intercompany-Rechnung {monate[month]} {year}",
        subtitle=f"Rechnungs-Nr. {rg_nr} – Service- & Vertriebs-Charge gemaess "
                 f"Konzern-Verrechnungspreis-Richtlinie 2018",
        sections=[
            ("Rechnungssteller / Invoice Issuer",
             f"{sender_name}\n\n"
             "Steuer-/Reg.-Nummer: siehe Briefkopf. Ust-/VAT-ID gemaess CN-Fapiao "
             "(special VAT invoice / 增值税专用发票) bzw. Reverse-Charge-Abrechnung "
             "im EU-Bereich. Konzern-Verrechnung gemaess OECD-Verrechnungspreis-"
             "Leitlinien (TPG 2022, Kap. VII) und Section 6 China Tax Administration "
             "Bulletin 2017/6 (Special Tax Investigation / SAT-Mitteilung)."),
            ("Rechnungsempfaenger / Bill-To",
             f"{empf_name}\n\n"
             "Verrechnung im Rahmen des gueltigen Intercompany-Service-Agreements "
             "vom 1. Januar 2019 (Master Service Agreement RCN/REA/REG/RSG). "
             "Methode: Cost-Plus 7,5% auf Kostenbasis (Distribution Function) bzw. "
             "TNMM (Berry-Ratio 1,08 fuer Application-Engineering-Anteil), "
             "dokumentiert im Local File 2022/2023 (China) sowie Master File "
             "Brennhagen-Konzern 2023."),
            ("Leistung / Service Description",
             leistung + ".\n\n"
             "Periode: " + monate[month] + " " + str(year) + ". Detail-Aufstellung "
             "(Verkaufsfoerderung-Volumen, Mannstunden Application Engineering, "
             "Anteilige Allgemeinkosten) gemaess beigefuegtem Annex A "
             "(Excel-Aufstellung intern abrufbar unter "
             "RCN_IC_Detail_" + str(year) + "_" + f"{month:02d}.xlsx)."),
            ("Rechnungsbetrag / Invoice Amount",
             [["Position", "Bemerkung", "Betrag"],
              ["Service-Charge (Cost+7,5%)", "Distribution Function RCN",
               f"{betrag_cny*0.62:,.2f} {currency}"],
              ["Application Engineering",
               "ASPICE Lvl 2 Test-Support, OEM Onsite",
               f"{betrag_cny*0.28:,.2f} {currency}"],
              ["Local Marketing China",
               "Messen Auto China Shanghai 2023, OEM-Roadshows",
               f"{betrag_cny*0.10:,.2f} {currency}"],
              ["GESAMT-NETTO", "ohne CN-VAT (Reverse-Charge / 跨境服务零税率)",
               betrag_str]]),
            ("Zahlungsbedingungen / Payment Terms",
             f"Zahlungsziel: 60 Tage netto ab Rechnungsdatum (gemaess "
             "Intercompany-Standardkonditionen).\n\n"
             f"Bankverbindung: {bank_line}.\n\n"
             "FX-Abwicklung CNY-Outbound: SAFE-Devisenkontroll-Pflicht "
             "(State Administration of Foreign Exchange). Fuer Beraege > 50.000 USD "
             "ist die jeweilige Bank zur Pruefung der Underlying-Dokumente "
             "(Service-Agreement, Tax-Filing, Bank-Form-Filing) verpflichtet. "
             "Steuerliche Quellensteuer (Withholding Tax, EIT) 10% sowie "
             "Surcharge 6,72% VAT auf grenzueberschreitende Dienstleistungen "
             "werden gemaess Cai-Shui 2016/36 von RCN einbehalten und abgefuehrt; "
             "der genehmigte Outbound-Remittance-Betrag wird entsprechend reduziert."),
            ("Hinweise",
             "Verrechnungspreis-Konformitaet wird vom Group Tax (Dr. Heike Berger) "
             "und KPMG Shanghai (Lead Partner Wendy Zhao) jaehrlich gepruegt. "
             "Bei Aenderungen der OECD- oder SAT-Regulierung wird die Rechnung "
             "ex-post angepasst. Reklamationen sind binnen 14 Tagen schriftlich "
             "(E-Mail liang.wei@rohrig.cn, cc tax@rohrig.de) anzuzeigen."),
        ])

# 2020 (RCN-Rechnungen)
for m in [1,2,3,5,6,7,8,9,10,12]:
    suffix = "_ALT" if m == 9 else ""
    ic_rechnung(f"RCN_IC_Rechnung_2020_{m:02d}{suffix}.docx", 2020, m)
# 2021
for m in range(1, 13):
    suffix = "_v2" if m == 2 else ("_FINAL_v2" if m == 8 else "")
    ic_rechnung(f"RCN_IC_Rechnung_2021_{m:02d}{suffix}.docx", 2021, m)
# 2022
for m in [2,3,4,5,6,7,8,10,11,12]:
    ic_rechnung(f"RCN_IC_Rechnung_2022_{m:02d}.docx", 2022, m)

# RSG May 2022
ic_rechnung("RSG_IC_Rechnung_2022_05.docx", 2022, 5, header=H_RSG, sender_label="RSG",
            betrag_cny=1_240_000)


# ─────────────────────────────────────────────────────────────────────────────
# 3) IC-QUARTALSBERICHTE (8)
# ─────────────────────────────────────────────────────────────────────────────
def ic_quartalsbericht(fname, year, q, draft=False):
    umsatz_cny = {2019: [128.4, 142.6, 138.2, 167.8],
                  2020: [98.3, 121.4, 145.6, 178.9]}[year][q-1]
    ebit_cny = umsatz_cny * 0.082
    fte = {2019: [165,172,178,185], 2020: [192,205,218,234]}[year][q-1]
    write_doc(f"{BASE}/{fname}", H_RCN_SHORT,
        f"Intercompany-Quartalsbericht Q{q} {year} – Brennhagen (Shanghai) Co. Ltd.",
        subtitle=f"Berichtszeitraum: Quartal {q}/{year} – Berichtswaehrung CNY (Mio.), "
                 f"Konzern-Reporting EUR (IFRS) per Tagesmittelkurs ECB",
        draft=draft,
        sections=[
            ("1. Management Summary",
             CONTEXT_PARA + "\n\n"
             f"Im Berichtsquartal Q{q}/{year} erzielte RCN einen Umsatz von "
             f"{umsatz_cny:.1f} Mio. CNY (Vorquartal Vergleichszahl im Anhang) "
             f"bei einem EBIT von {ebit_cny:.2f} Mio. CNY (EBIT-Marge ca. 8,2%). "
             f"Auftragseingang lag mit 1,12x ueber Umsatz (Book-to-Bill > 1), "
             f"getrieben durch Mercedes-Beijing-Programm (W206/EQE) sowie "
             f"BMW Brilliance Shenyang (i3-EV Plattform). FTE Stand "
             f"{fte} Mitarbeiter (Vergleich Vorjahresquartal +12%)."),
            ("2. Umsatz nach Kunden (CNY Mio.)",
             [["Kunde / Programm", "Q-Umsatz", "YTD", "Anteil"],
              ["Mercedes-Benz Beijing (BBAC) – W206/EQE",
               f"{umsatz_cny*0.34:.1f}", f"{umsatz_cny*0.34*q:.1f}", "34%"],
              ["Stellantis Wuhan (DPCA) – C-SUV",
               f"{umsatz_cny*0.18:.1f}", f"{umsatz_cny*0.18*q:.1f}", "18%"],
              ["BMW Brilliance Shenyang (BBA) – iX3/i3-EV",
               f"{umsatz_cny*0.21:.1f}", f"{umsatz_cny*0.21*q:.1f}", "21%"],
              ["FAW-VW Changchun – MEB-Programm",
               f"{umsatz_cny*0.12:.1f}", f"{umsatz_cny*0.12*q:.1f}", "12%"],
              ["Aftermarket China + ASEAN (Tier-1 Distributoren)",
               f"{umsatz_cny*0.10:.1f}", f"{umsatz_cny*0.10*q:.1f}", "10%"],
              ["Sonstige (NIO, Li Auto Engineering-Samples)",
               f"{umsatz_cny*0.05:.1f}", f"{umsatz_cny*0.05*q:.1f}", "5%"],
              ["SUMME", f"{umsatz_cny:.1f}", f"{umsatz_cny*q:.1f}", "100%"]]),
            ("3. Ergebnis & KPI",
             [["Kennzahl", "Q-Wert", "Plan", "Abw."],
              ["Umsatz (Mio. CNY)", f"{umsatz_cny:.1f}", f"{umsatz_cny*0.95:.1f}", "+5%"],
              ["EBIT (Mio. CNY)", f"{ebit_cny:.2f}", f"{ebit_cny*0.92:.2f}", "+9%"],
              ["EBIT-Marge", "8,2%", "8,0%", "+0,2pp"],
              ["DSO (Tage)", "78", "72", "+6"],
              ["FTE", str(fte), str(fte-3), "+3"],
              ["FX-Hedging-Anteil (CNY/EUR)", "65%", "60%", "+5pp"]]),
            ("4. Operative Highlights",
             ("list", [
                 "Nominierung BMW Brilliance Shenyang fuer LightCtrl-7 Matrix-LED "
                 "Steuermodul (Volumen ca. 240k Stk/Jahr, SOP Q1 2021).",
                 "Mercedes-Benz Beijing PPAP Level-3 Freigabe ICP-3 (Infotainment) "
                 "fuer W206 erteilt; SOP konfirmiert.",
                 "Stellantis Wuhan: Lieferprobleme Steckverbinder von RCZ Brno "
                 "wegen Hafen-Engpaessen Shanghai-Yangshan – Mitigation per "
                 "Luftfracht (Sonderkosten ca. 280k CNY).",
                 "Aftermarket-Joint-Venture mit Bosch-China-Aftermarket: "
                 "Vorgespraeche begonnen (NDA unterzeichnet).",
                 "Lokale Sozialversicherungs-Beitragssaetze Shanghai 2020: "
                 "leicht reduziert (COVID-19-Entlastung), positive EBIT-Auswirkung "
                 "ca. 0,4 Mio. CNY p.a.",
             ])),
            ("5. Risiken & FX",
             "FX-Risiko CNY/EUR: Im Quartal Schwankungsbreite 7,82–7,94 CNY/EUR. "
             "Hedging via Forward-Kontrakte (Bank of Communications) ueber "
             "Group Treasury (Markus Pflanzer) abgesichert.\n\n"
             "Regulatorische Risiken: SAFE-Verschaerfung bei Outbound-Remittance "
             "Service-Trade; KPMG Shanghai begleitet die Filings. FIE-Steuer"
             "verguenstigungen (15% reduzierter EIT-Satz fuer Hi-Tech-Status) "
             "laufen bis 31.12.2025; Verlaengerung wird 2024 beantragt.",
             ),
            ("6. Ausblick",
             f"Fuer das Folgequartal wird ein Umsatz im Korridor "
             f"{umsatz_cny*0.95:.0f}–{umsatz_cny*1.10:.0f} Mio. CNY erwartet. "
             "Hauptthemen: Anlauf-Unterstuetzung Mercedes-Beijing EQE, BMW "
             "Brilliance LightCtrl-7 Engineering-Phase, sowie SAFE-Audit "
             "Vorbereitung 2020/Q4."),
            ("Berichtsfreigabe",
             signatures(CM, "Country Manager", "Brennhagen (Shanghai) Co. Ltd.",
                        FIN, "Finance Manager", "Brennhagen (Shanghai) Co. Ltd.",
                        place="Shanghai",
                        date_str_=f"30. {['Maerz','Juni','September','Dezember'][q-1]} {year}")),
        ])

ic_quartalsbericht("RCN_IC_Quartalsbericht_2019_Q1_ENTWURF.docx", 2019, 1, draft=True)
ic_quartalsbericht("RCN_IC_Quartalsbericht_2019_Q2.docx", 2019, 2)
ic_quartalsbericht("RCN_IC_Quartalsbericht_2019_Q3.docx", 2019, 3)
ic_quartalsbericht("RCN_IC_Quartalsbericht_2019_Q4.docx", 2019, 4)
ic_quartalsbericht("RCN_IC_Quartalsbericht_2020_Q1.docx", 2020, 1)
ic_quartalsbericht("RCN_IC_Quartalsbericht_2020_Q2.docx", 2020, 2)
ic_quartalsbericht("RCN_IC_Quartalsbericht_2020_Q3_WIP.docx", 2020, 3, draft=True)
ic_quartalsbericht("RCN_IC_Quartalsbericht_2020_Q4.docx", 2020, 4)


# ─────────────────────────────────────────────────────────────────────────────
# 4) COMPLIANCE-BERICHT 2023
# ─────────────────────────────────────────────────────────────────────────────
write_doc(f"{BASE}/RCN_Compliance_Report_2023_reviewed.docx", H_RCN_SHORT,
    "Compliance-Bericht 2023 – Brennhagen (Shanghai) Co. Ltd.",
    subtitle="Jaehrlicher Compliance-Bericht zur Vorlage an Group Compliance "
             "Stuttgart und Group Internal Audit (CAE Andreas Buehler)",
    sections=[
        ("1. Ueberblick",
         CONTEXT_PARA + "\n\n"
         "Der vorliegende Bericht dokumentiert den Status der Compliance- und "
         "Governance-Funktion der RCN fuer das Geschaeftsjahr 2023, basierend "
         "auf den vom Group Compliance Officer (Frau Dr. Sabine Mertens, "
         "Stuttgart) vorgegebenen Reporting-Templates. Stichtag: 31.12.2023."),
        ("2. Regulatorisches Umfeld",
         ("list", [
             "Foreign-Invested Enterprise Law (FIL, in Kraft seit 1.1.2020) – RCN "
             "weiterhin als WFOE registriert; Conversion gemaess SAMR-Bekanntmachung "
             "2019/64 abgeschlossen.",
             "FIE-Steuerverguenstigungen (Hi-Tech Enterprise Status, HNTE) gelten "
             "bis 31.12.2025 (15% statt 25% EIT). Re-Zertifizierung Q3/2024 "
             "geplant; Vorbereitung durch KPMG Shanghai.",
             "PRC Personal Information Protection Law (PIPL, in Kraft 1.11.2021) – "
             "Datenuebertragungs-Assessment fuer Cross-Border-Datenfluesse zur "
             "Konzernzentrale Stuttgart durchgefuegt (CAC-Filing 2023/Q2 eingereicht).",
             "Anti-Corruption (PRC Criminal Law Sec. 163/164, AUCL) – jaehrliches "
             "ABC-Training (Anti-Bribery/Corruption) fuer alle 320 RCN-MA "
             "absolviert; Quote 100% (Vorjahr 97%).",
             "Export Control Law (ECL, 1.12.2020) – Dual-Use-Pruefung etabliert; "
             "keine kritischen Ausfuhren in 2023.",
         ])),
        ("3. SAFE-Devisenkontrolle / FX-Compliance",
         "Alle 87 Outbound-Remittances 2023 (kumuliert 42,3 Mio. CNY an Konzern: "
         "Service-Charges, Lizenzgebuehren, Dividende vorbereitet) wurden gemaess "
         "SAFE-Verfahren (Hui Fa 2014/19, Hui Fa 2020/8) abgewickelt. Bank of "
         "Communications hat sechs Transaktionen mit erweitertem Underlying-"
         "Dokumenten-Review gepruegt (Service-Vertrag, Tax-Filing-Receipt, "
         "Form-Filing), in allen Faellen freigegeben.\n\n"
         "Steuerliche Quellensteuer (10% EIT + 6% VAT Surcharge) wurde "
         "einbehalten und an die Shanghai Tax Authority (Minhang Branch) "
         "abgefuegt; Tax-Receipts vorliegend.\n\n"
         "Keine Beanstandungen durch SAFE im Audit-Zyklus 2023."),
        ("4. Anti-Korruption & Geschenke-Register",
         "Geschenke- und Bewirtungs-Register 2023: 142 Eintraege, davon 18 "
         "ueber dem Schwellenwert 1.500 CNY (Genehmigungs-pflichtig). Alle "
         "ueberschwelligen Faelle wurden vom Country Manager Zhang Hao "
         "genehmigt und im Group Tool »ComplianceNet« dokumentiert.\n\n"
         "Keine Whistleblower-Meldungen ueber die SpeakUp-Hotline (Concentric) "
         "im Berichtsjahr aus RCN."),
        ("5. Sozialversicherung / Arbeitsrecht",
         "Alle 320 MA sind ordnungsgemaess bei SHHRSS (Pension, Krankenversicherung, "
         "Arbeitslosen-, Arbeitsunfall-, Mutterschafts-Versicherung) sowie HPF "
         "(Housing Provident Fund) gemeldet und abgerechnet. Lohn-/SV-Audit durch "
         "KPMG Shanghai (Februar 2024) ohne Findings. Beitragsbemessungsgrundlage "
         "monatlich automatisch an SHHRSS-Tabelle angepasst (Untergrenze 7.310 CNY, "
         "Obergrenze 36.549 CNY in 2023)."),
        ("6. Datenschutz (PIPL)",
         "PIPL-Cross-Border-Transfer-Assessment fuer HR-Daten (Stuttgart-Konzern-"
         "Reporting) sowie Customer-Engineering-Daten (Mercedes-Beijing, BMW-"
         "Brilliance Cross-Border-Engineering) abgeschlossen. CAC-Filing am "
         "12.05.2023 eingereicht, Bestaetigung am 28.07.2023 erhalten."),
        ("7. Findings & Massnahmen",
         [["Finding", "Quelle", "Status", "Owner"],
          ["DSO Mercedes-Beijing > 90 Tage", "Group Treasury",
           "in Bearbeitung", FIN],
          ["FIE-Hi-Tech-Status Re-Zert. 2024",
           "Group Tax", "vorbereitet Q3 2024", "KPMG / Tax"],
          ["Lokales Code-Repository (PIPL)",
           "IT-Audit", "Migration abgeschlossen", "IT"],
          ["Geschenke-Register: 3 Late-Filings", "Internal",
           "Trainings durchgefuehrt", CM]]),
        ("8. Bestaetigung",
         signatures(CM, "Country Manager", "Brennhagen (Shanghai) Co. Ltd.",
                    "Dr. Sabine Mertens", "Group Compliance Officer",
                    "Brennhagen Elektronik AG",
                    place="Shanghai / Stuttgart", date_str_="15. Februar 2024")),
    ])


# ─────────────────────────────────────────────────────────────────────────────
# 5) STEUERBESCHEIDE
# ─────────────────────────────────────────────────────────────────────────────
def steuerbescheid(fname, year, kst=False):
    title = f"Steuerbescheid / Tax Assessment {year}" + (
        " – Koerperschaftsteuer (Enterprise Income Tax)" if kst else "")
    rev_cny = {2022: 612.4, 2023: 728.6}[year]
    write_doc(f"{BASE}/{fname}", H_RCN_SHORT,
        f"{title} – Brennhagen (Shanghai) Co. Ltd.",
        subtitle=f"Bescheid der Shanghai Municipal Tax Service, Minhang Branch – "
                 f"Veranlagungsjahr {year}",
        sections=[
            ("Ausstellende Behoerde",
             "Shanghai Municipal Tax Service, Minhang District Branch "
             "(上海市税务局闵行区税务局)\n"
             "Adresse: No. 555 Zhongchun Road, Minhang District, 201101 Shanghai\n"
             f"Aktenzeichen: SH-MTS-MH-EIT-{year}-RCN-00321\n"
             f"Ausstellungsdatum: 28. Mai {year+1}\n"
             "Verfasser: Mr. Liu Jianhua, Senior Tax Inspector (高级税务稽查员)"),
            ("Steuerpflichtige Gesellschaft",
             f"Brennhagen (Shanghai) Co. Ltd. / 罗瑞格(上海)有限公司\n"
             "USCC 91310000789123456X – Wholly Foreign-Owned Enterprise (WFOE)\n"
             "Hi-Tech Enterprise (HNTE) Status: Zertifikat-Nr. GR202031003845, "
             "gueltig bis 31.12.2025 (reduzierter EIT-Satz 15% gemaess Cai-Shui "
             "2008/1).\n"
             "Geschaeftsjahr: 1.1. bis 31.12."),
            ("Veranlagungsergebnis (EIT)",
             [["Position", "Betrag (CNY)", "Bemerkung"],
              ["Steuerpflichtiger Umsatz", f"{rev_cny:,.1f} Mio.", "IFRS-konform, GAAP-Ueberleitung"],
              ["Betriebsausgaben anerkannt", f"{rev_cny*0.913:,.1f} Mio.",
               "inkl. R&D-Super-Deduction 175%"],
              ["Zu versteuerndes Einkommen (Taxable Income)",
               f"{rev_cny*0.087:,.1f} Mio.", "Vor Abzug R&D-Bonus"],
              ["R&D-Super-Deduction (75% bonus)", f"{rev_cny*0.012:,.1f} Mio.",
               "gemaess Cai-Shui 2018/99"],
              ["Endgueltiges Taxable Income", f"{rev_cny*0.075:,.1f} Mio.", ""],
              ["EIT-Satz (HNTE-Status)", "15,00%", "statt regulaer 25%"],
              ["Festgesetzte EIT", f"{rev_cny*0.075*0.15:,.2f} Mio.", "endfaellig"],
              ["bereits geleistete Quartals-Vorauszahlungen",
               f"{rev_cny*0.075*0.15*0.92:,.2f} Mio.", "Q1–Q4"],
              ["Nachzahlung / (Erstattung)",
               f"{rev_cny*0.075*0.15*0.08:,.2f} Mio.", "faellig zum 30.06."]]),
            ("Quellensteuer auf Auslandszahlungen (Withholding EIT)",
             f"Im Veranlagungsjahr {year} wurden grenzueberschreitende Service-"
             "Charges und Lizenzgebuehren an die Konzernmuetter (REA, REG, RSG) "
             f"in Hoehe von {rev_cny*0.063:,.1f} Mio. CNY abgerechnet. Hierauf "
             "wurde Quellensteuer (Withholding EIT) in Hoehe von 10% einbehalten "
             "und gemaess Cai-Shui 2016/36 sowie DBA Deutschland-China (Art. 12) "
             "abgefuehrt. Ferner 6% VAT auf grenzueberschreitende Dienstleistungen "
             "(Reverse-Charge)."),
            ("Rechtsbehelf",
             "Gegen diesen Bescheid kann binnen 60 Tagen nach Zustellung "
             "(gemaess Art. 88 PRC Tax Collection Administration Law) Einspruch "
             "(Administrative Review / 行政复议) bei der Shanghai Municipal Tax "
             "Service eingelegt werden. Die festgesetzte Steuer ist gleichwohl "
             "fristgerecht zu entrichten; Aussetzung der Vollziehung nur in "
             "Ausnahmefaellen.\n\n"
             "Hinweis Wirtschaftspruefer: KPMG Shanghai hat den Bescheid am "
             "12.06. fuer plausibel befunden; keine Einspruchsempfehlung "
             "(Memo KPMG Shanghai/RCN/" + str(year) + "/Tax-04)."),
            ("Bestaetigung Erhalt",
             signatures(FIN, "Finance Manager", "Brennhagen (Shanghai) Co. Ltd.",
                        CM, "Country Manager / Legal Representative",
                        "Brennhagen (Shanghai) Co. Ltd.",
                        place="Shanghai", date_str_=f"5. Juni {year+1}")),
        ])

steuerbescheid("RCN_Steuerbescheid_2022.docx", 2022, kst=False)
steuerbescheid("RCN_Steuerbescheid_KSt_2023.docx", 2023, kst=True)


# ─────────────────────────────────────────────────────────────────────────────
# 6) VERSICHERUNGSNACHWEIS 2023
# ─────────────────────────────────────────────────────────────────────────────
write_doc(f"{BASE}/RCN_Versicherungsnachweis_2023.docx", H_RCN_SHORT,
    "Versicherungsnachweis – Brennhagen (Shanghai) Co. Ltd. – 2023",
    subtitle="Uebersicht aller in 2023 abgeschlossenen Versicherungen "
             "(lokal China + Group Master Cover)",
    sections=[
        ("Versicherungsnehmer",
         CONTEXT_PARA + "\n\n"
         "Die Versicherungs-Struktur kombiniert (i) lokale chinesische Pflicht-"
         "Versicherungen, (ii) freiwillige lokale Sachversicherungen (Property), "
         "(iii) globale Konzern-Master-Cover (Allianz Global Corporate & "
         "Specialty SE, Marsh Broker)."),
        ("Versicherungs-Uebersicht 2023",
         [["Sparte", "Versicherer", "Police-Nr.", "Versicherungssumme",
           "Praemie 2023", "Laufzeit"],
          ["Sach- / Inhalt (Property + BU)",
           "PICC P&C Shanghai Branch (人保财险)", "PICC-SH-PRP-2023-RCN-001",
           "120 Mio. CNY (Geb./Inhalt) + 18 Mio. BU",
           "287.500 CNY", "01.01.–31.12.2023"],
          ["Allgemeine Haftpflicht (Public Liability)",
           "Ping An P&C Insurance (平安产险)", "PA-SH-GL-2023-44782",
           "20 Mio. CNY", "62.000 CNY", "01.01.–31.12.2023"],
          ["Produkthaftpflicht (Product Liability) China",
           "AIG China (lokale Nfdg.)", "AIG-CN-PL-2023-RCN",
           "50 Mio. CNY", "180.000 CNY", "01.01.–31.12.2023"],
          ["Produkthaftpflicht – Group Master Cover Excess",
           "Allianz GCS SE (Group)", "AGCS-GBL-PL-2023-RHGRP-01",
           "USD 150 Mio. xs USD 5 Mio.", "anteilig (Konzernverrechnung)",
           "01.01.–31.12.2023"],
          ["D&O (Group Master)",
           "Allianz GCS SE", "AGCS-GBL-DO-2023-RHGRP-02",
           "EUR 50 Mio.", "anteilig (Group Allocation)",
           "01.01.–31.12.2023"],
          ["Marine / Cargo (Transport)",
           "PICC Marine, Shanghai", "PICC-MAR-2023-RCN-009",
           "Open Cover bis 5 Mio. USD/Sendung",
           "98.000 CNY (variabel)", "Open Cover Jahr"],
          ["Cyber Risk (Group Master)",
           "Beazley / Allianz (Konzern-Programm)",
           "BEAZ-GBL-CY-2023-RHGRP", "EUR 25 Mio.",
           "anteilig", "01.01.–31.12.2023"],
          ["Arbeitnehmer-Gruppen-Unfall (Group Accident)",
           "China Life P&C (中国人寿财险)", "CLPC-SH-GA-2023-RCN",
           "200.000 CNY pro MA", "210.000 CNY", "01.01.–31.12.2023"],
          ["Pflicht-Sozialversicherung (5 Lines) + HPF",
           "SHHRSS Minhang", "(gesetzlich)", "n/a",
           "ca. 18,2 Mio. CNY AG-Anteil", "fortlaufend"]]),
        ("Schadensquote 2023 (vorlaeufig)",
         "Im Berichtsjahr 2023 wurden zwei Schaeden gemeldet: (i) Wasserschaden "
         "Lager Building 7 vom 17.08.2023 nach Taifun »Doksuri« – Schaden ca. "
         "780.000 CNY, Regulierung durch PICC abgeschlossen; (ii) Frachtschaden "
         "ICP-3 Containerlieferung Mercedes Beijing vom 22.10.2023 – Schaden "
         "ca. 140.000 USD, Regulierung durch PICC Marine in Bearbeitung. "
         "D&O-Faelle: keine. Produkthaftpflicht-Faelle: keine."),
        ("Group Broker / Kontakte",
         "Globaler Versicherungsmakler: Marsh (Munich Office) – "
         "Senior Client Executive Frau Bettina Reiff. Lokaler China-Korrespondent: "
         "Marsh China (Shanghai) – Mr. Eric Liu.\n\n"
         "Group Insurance Manager (Stuttgart): Frau Heike Schroeder, "
         "h.schroeder@rohrig.de.\n\n"
         "Lokaler RCN-Ansprechpartner: " + FIN + " (Finance / Risk Management)."),
        ("Bestaetigung",
         signatures(FIN, "Finance Manager", "Brennhagen (Shanghai) Co. Ltd.",
                    CM, "Country Manager", "Brennhagen (Shanghai) Co. Ltd.",
                    place="Shanghai", date_str_="15. Januar 2024")),
    ])


# ─────────────────────────────────────────────────────────────────────────────
# 7) WP MANAGEMENT LETTERS 2021/2022/2023 (3 Stueck)
# ─────────────────────────────────────────────────────────────────────────────
def wp_letter(fname, year):
    write_doc(f"{BASE}/{fname}", H_KPMG,
        f"Management Letter {year} – Brennhagen (Shanghai) Co. Ltd.",
        subtitle=f"Pruefungsergaenzendes Schreiben zur Jahresabschlusspruefung "
                 f"zum 31.12.{year} – Konzern-Reporting Package",
        confidential=True,
        sections=[
            ("An",
             f"Vorstand der Brennhagen Elektronik AG, Stuttgart\n"
             f"z. Hd. Frau Laura Bauer (CFO) und Frau Dr. Heike Berger (Group Tax)\n"
             f"sowie an den Pruefungsausschuss des Aufsichtsrats "
             f"(Vorsitz Prof. Dr.-Ing. Gerhard Voss)\n\n"
             f"in Kopie: Country Manager RCN Shanghai – Herr Zhang Hao; "
             f"Finance Manager RCN – Frau Liang Wei; CAE Andreas Buehler."),
            ("Gegenstand",
             f"Im Rahmen unserer Pruefung des Jahresabschlusses {year} der "
             "Brennhagen (Shanghai) Co. Ltd. (»RCN«) als Reporting-Package fuer den "
             f"Konzernabschluss {year} der Brennhagen Elektronik AG haben wir die "
             "internen Kontrollen und Prozesse im Bereich Finance, Tax, "
             "Compliance und IT gepruegt. Die Pruefung erfolgte gemeinsam mit "
             "unserem chinesischen Pruefungsnetzwerk (KPMG Huazhen LLP, "
             "Shanghai Office; Lead Partner Frau Wendy Zhao, CICPA).\n\n"
             "Die Pruefung erfolgte nach ISA (International Standards on Auditing) "
             "sowie chinesischen CAS (China Auditing Standards) fuer das lokale "
             "Statutory-Filing. Der uneingeschraenkte Bestaetigungsvermerk "
             "wurde am Folgetag dieses Letters erteilt."),
            ("Wesentliche Feststellungen",
             ("clauses", [
                 (f"§ 1 Verrechnungspreise (Transfer Pricing)", [
                     "Die Service-Charge-Methode (Cost-Plus 7,5%) ist im Local "
                     f"File {year} (China) sowie im Master File {year} dokumentiert. "
                     "Wir empfehlen die jaehrliche Aktualisierung der "
                     "Benchmark-Studie (Distributoren-Vergleichbarkeitsanalyse) "
                     "mit aktuellen Datenbankvergleichen (Bureau van Dijk Oriana).",
                     "Berry-Ratio liegt im Korridor (Interquartil 1,03–1,12); "
                     "kein Anpassungsbedarf erkennbar.",
                 ]),
                 (f"§ 2 SAFE-Compliance / Outbound-Remittances", [
                     "Sample-Pruefung 12 Outbound-Remittances zeigt ordnungsgemaesse "
                     "SAFE-Filings (Hui Fa 2014/19). Empfehlung: "
                     "Centralised-Process-Documentation mit eindeutiger Vier-"
                     "Augen-Freigabe (Finance Manager + Country Manager) bei "
                     "Beraegen > USD 100k.",
                 ]),
                 (f"§ 3 R&D-Super-Deduction (Cai-Shui 2018/99)", [
                     "RCN beansprucht keine R&D-Super-Deduction; Empfehlung: "
                     "ab 2024 ggf. lokale Application-Engineering-Aufwendungen "
                     "(Anteil Mercedes-Beijing PPAP-Engineering) als R&D-"
                     "qualifiziert pruefen.",
                 ]),
                 (f"§ 4 Forderungs-Wertberichtigung (ECL nach IFRS 9)", [
                     "DSO-Spitze bei Mercedes-Beijing (95 Tage); EWB pauschal "
                     "1,8% gemaess Konzern-Policy. Wir halten dies fuer angemessen.",
                 ]),
                 (f"§ 5 IT/PIPL", [
                     "Cross-Border-Data-Transfer-Assessment (CAC-Filing) liegt vor. "
                     "Empfehlung: jaehrlicher Penetrationstest des Shanghai-RZ "
                     "(KPMG IT-Audit-Team verfuegbar).",
                 ]),
             ])),
            ("Zusammenfassung",
             f"Im Berichtsjahr {year} wurden keine schwerwiegenden Mangel in der "
             "internen Kontrolle festgestellt. Die identifizierten Verbesserungs"
             "punkte sind operativer Natur und werden mit dem RCN-Management "
             "(Zhang Hao, Liang Wei) in Q1 des Folgejahres adressiert."),
            ("Unterzeichnung",
             signatures("Dr. Maximilian Brand", "Lead Partner / WP",
                        "KPMG AG WPG, Berlin",
                        "Wendy Zhao, CICPA", "Lead Partner China",
                        "KPMG Huazhen LLP, Shanghai",
                        place="Berlin / Shanghai", date_str_=f"15. Maerz {year+1}")),
        ])

wp_letter("RCN_WP_Management_Letter_2021.docx", 2021)
wp_letter("RCN_WP_Management_Letter_2022.docx", 2022)
wp_letter("RCN_WP_Management_Letter_2023.docx", 2023)


# ─────────────────────────────────────────────────────────────────────────────
# 8) MIETVERTRAG BETRIEBSGELAENDE 2020
# ─────────────────────────────────────────────────────────────────────────────
write_doc(f"{BASE}/RCN_Mietvertrag_Betriebsgelaende_2020.docx", H_RCN_SHORT,
    "Miet- / Pachtvertrag Betriebsgelaende – Brennhagen (Shanghai) Co. Ltd.",
    subtitle="Anmietung Building 7, Caohejing Hi-Tech Park Minhang – Laufzeit 5 Jahre",
    sections=[
        ("Vertragsparteien",
         "Vermieter: Shanghai Caohejing Hi-Tech Park Development Co. Ltd. "
         "(上海漕河泾新兴技术开发区发展总公司), Adresse: No. 200 Tianlin Road, "
         "Xuhui District, 200233 Shanghai, USCC 91310000132456789X, "
         "vertreten durch den Direktor Mr. Chen Yulong.\n\n"
         "Mieter: Brennhagen (Shanghai) Co. Ltd. / 罗瑞格(上海)有限公司, "
         "USCC 91310000789123456X, vertreten durch den Country Manager und "
         "Legal Representative " + CM + ".\n\n"
         "Der Vertrag wird in zwei Originalfassungen geschlossen (Mandarin und "
         "deutsch/englisch); im Konfliktfall ist die Mandarin-Fassung verbindlich."),
        ("Vertragsregelungen",
         ("clauses", [
             ("§ 1 Mietobjekt", [
                 "Building 7, Caohejing Hi-Tech Park, No. 1188 Lianhang Road, "
                 "Minhang District, 201112 Shanghai. Nutzflaeche gesamt 8.450 m² "
                 "BGF, davon (i) 5.200 m² Buero/Application Engineering, "
                 "(ii) 1.800 m² Lager fuer ICP-3, BMS-12, ADAS-V4D, ECU-900, "
                 "LightCtrl-7 Aftermarket-Bestaende, (iii) 1.450 m² "
                 "Labor- und Werkstattflaechen.",
                 "Inkl. 65 Stellplaetze Tiefgarage, 4 Lkw-Verladestellen, "
                 "uneingeschraenkter Zugang zum Park-Gemeinschafts-Glasfasernetz.",
             ]),
             ("§ 2 Vertragsdauer", [
                 "5 Jahre fest: 01.01.2020 – 31.12.2024.",
                 "Option auf Verlaengerung um weitere 5 Jahre, auszuueben "
                 "spaetestens 12 Monate vor Vertragsablauf. Mietpreis-"
                 "Anpassung an Caohejing-Marktpreis (max. +12% gegenueber "
                 "Bestandsmiete).",
             ]),
             ("§ 3 Mietzins", [
                 "Kaltmiete: 6,80 CNY/m²/Tag (= ca. 207 CNY/m²/Monat) "
                 "fuer Buero/Engineering; 3,40 CNY/m²/Tag fuer Lager; "
                 "5,50 CNY/m²/Tag fuer Labor.",
                 "Gesamtmonatsmiete (Kaltmiete) ca. 1.612.000 CNY exkl. Steuern. "
                 "Zzgl. 9% VAT (Property Lease) und Nebenkostenpauschale 12,5 "
                 "CNY/m²/Monat. Jaehrliche Mietsteigerung 4% (Index-gebunden).",
                 "Kaution: 3 Monatsmieten (ca. 4,84 Mio. CNY) als Sicherheitsleistung; "
                 "Hinterlegung als Bank-Guarantee der Bank of Communications.",
             ]),
             ("§ 4 Nebenkosten", [
                 "Strom: gesonderte Subzaehlung, Abrechnung quartalsweise nach "
                 "Verbrauch zum jeweils gueltigen Industrie-Tarif Shanghai. "
                 "Wasser, Klimatisierung, Internet (Glasfaser), Wachschutz und "
                 "Reinigung der Gemeinschaftsflaechen sind Vermieterleistung.",
             ]),
             ("§ 5 Nutzungszweck", [
                 "Geneigt fuer Verwaltung, Application Engineering, Lager- und "
                 "Distribution-Center fuer Automotive-Elektronik-Produkte. "
                 "Aenderung des Nutzungszwecks bedarf schriftlicher Zustimmung "
                 "des Vermieters und ggf. Aenderung der FIE-Geschaeftslizenz.",
             ]),
             ("§ 6 Versicherung & Brandschutz", [
                 "Mieter verpflichtet sich zur Property- und Inhalt-Versicherung "
                 "(PICC P&C Shanghai, Police PICC-SH-PRP-2023-RCN-001). "
                 "Vermieter unterhaelt Gebaeude-Versicherung. Brandschutz-Konzept "
                 "gemaess Shanghai Municipal Fire Code 2018; jaehrliche "
                 "Sicherheits-Inspektion durch zustaendige Behoerde.",
             ]),
             ("§ 7 Kuendigung & Ausserordentliches", [
                 "Ordentliche Kuendigung waehrend der Festlaufzeit ausgeschlossen. "
                 "Ausserordentliche Kuendigung des Mieters bei Verlust der "
                 "FIE-Geschaeftslizenz oder wesentlicher Mietpflichtverletzung "
                 "des Vermieters mit 90 Tagen Frist.",
                 "Bei vorzeitiger Beendigung durch den Mieter Konventionalstrafe "
                 "in Hoehe von 3 Monatsmieten.",
             ]),
             ("§ 8 Anwendbares Recht / Gerichtsstand", [
                 "Recht der VR China. Streitigkeiten werden vor der Shanghai "
                 "International Arbitration Center (SHIAC) gemaess deren Regeln "
                 "verhandelt; Schiedsort Shanghai, Verfahrenssprache Chinesisch.",
             ]),
         ])),
        ("Unterschriften",
         signatures("Mr. Chen Yulong", "Direktor / Vermieter",
                    "Shanghai Caohejing Hi-Tech Park Dev. Co. Ltd.",
                    CM, "Country Manager / Legal Representative",
                    "Brennhagen (Shanghai) Co. Ltd.",
                    place="Shanghai", date_str_="18. Dezember 2019")),
    ])


# ─────────────────────────────────────────────────────────────────────────────
# 9) BR PROTOKOLL OKTOBER 2020 (RHO)
# ─────────────────────────────────────────────────────────────────────────────
write_doc(f"{BASE}/RHO_BR_Protokoll_2020_10.docx", H_RHO,
    "BR-Sitzungsprotokoll Oktober 2020 – Brennhagen Holding GmbH",
    subtitle="Protokoll der Sitzung des Betriebsrats am 14. Oktober 2020 – Stuttgart",
    sections=[
        ("Sitzungsdaten",
         "Datum: Mittwoch, 14. Oktober 2020, 09:30 – 12:15 Uhr\n"
         "Ort: Brennhagen Holding GmbH, Vaihinger Strasse 120, 70567 Stuttgart, "
         "Konferenzraum »Neckar« (Praesenz) sowie MS Teams (hybrid wg. COVID-19)\n"
         "Sitzungsleitung: Marlies Duerr (BR-Vorsitzende RHO und "
         "Konzernbetriebsrats-Vorsitzende, IG Metall)\n"
         "Protokollfuehrung: Stefan Roth (BR-Schriftfuehrer)\n\n"
         "Teilnehmer (BR): Marlies Duerr, Stefan Roth, Heike Brenner, Jens "
         "Lechner, Sabine Pohl (5 von 7 BR-Mitgliedern).\n"
         "Eingeladene Vertreter Arbeitgeber: Frau Anna Mueller (CEO), Frau Laura "
         "Bauer (CFO), Frau Dr. Anke Vogt (Head of Group HR)."),
        ("Tagesordnung",
         ("list", [
             "TOP 1: Genehmigung Protokoll September 2020",
             "TOP 2: COVID-19 – Mobile Arbeit Konzern-Rahmenvereinbarung (Status)",
             "TOP 3: Konzern-Reorganisation Tochter China (RCN Shanghai) – "
             "Information durch CFO",
             "TOP 4: Geplante Konsolidierung der HR-Systeme (Workday-Migration)",
             "TOP 5: Sozialplan-Vorgespraeche Werk Stuttgart-Vaihingen (40 Stellen)",
             "TOP 6: Verschiedenes",
         ])),
        ("TOP 1 – Protokollgenehmigung",
         "Das Protokoll der Sitzung vom 09.09.2020 wird ohne Aenderungen "
         "einstimmig angenommen (5 Ja, 0 Nein, 0 Enthaltung)."),
        ("TOP 2 – COVID-19 / Mobile Arbeit",
         "Die Konzern-Rahmenvereinbarung »Mobile Arbeit waehrend Pandemielage« "
         "vom 15.03.2020 wurde verlaengert bis 31.03.2021. Frau Bauer informiert "
         "ueber die Ausstattung mit Notebooks, Headsets sowie Erstattungspauschale "
         "fuer Stromkosten (50 EUR/Monat). BR begruesst die Verlaengerung und "
         "fordert Beruecksichtigung der Ergonomie-Anforderungen auch im Home-"
         "Office (Stuhl, Tisch). Vereinbart: HR prueft Ergonomie-Foerderung "
         "(bis 500 EUR pro MA, einmalig) bis November-Sitzung."),
        ("TOP 3 – Information RCN Shanghai",
         "Frau Bauer informiert ueber die Geschaeftslage der Tochter Brennhagen "
         "(Shanghai) Co. Ltd. (RCN): COVID-19-Erholung schneller als in DE; "
         "Mercedes-Beijing-PPAP Level-3 ICP-3 erteilt; FTE-Ausbau auf 234 (Q3) "
         "geplant. Keine direkte Berichts-Pflicht des chinesischen BR – die "
         "Konzernbetriebsrats-Vorsitzende Frau Duerr bittet um halbjaehrliche "
         "Informations-Aktualisierung zu Personalentwicklung RCN. CFO sagt zu."),
        ("TOP 4 – HR-System / Workday-Migration",
         "Geplant Q1/2021 Migration der HR-Systeme aller Konzerngesellschaften "
         "(REA, RHO, REG, RSG, RPL, RCZ, RHU, RCN) auf Workday. BR verlangt "
         "rechtzeitige Mitbestimmungsverhandlungen (§ 87 Abs. 1 Nr. 6 BetrVG). "
         "Frau Dr. Vogt sagt Start der Verhandlungen ab November 2020 zu."),
        ("TOP 5 – Sozialplan-Vorgespraeche Stuttgart-Vaihingen",
         "Geplanter Stellenabbau in Holding-Verwaltung (40 Stellen bis Mitte "
         "2021) durch Effizienzprogramm. BR fordert Sozialplan-Verhandlungen "
         "gemaess §§ 111, 112 BetrVG. Erstgespraech 28.10.2020. CFO bestaetigt."),
        ("TOP 6 – Verschiedenes",
         "Naechste BR-Sitzung am Mittwoch, 18.11.2020, 09:30 Uhr."),
        ("Unterschriften",
         signatures("Marlies Duerr", "BR-Vorsitzende", "Brennhagen Holding GmbH",
                    "Stefan Roth", "Schriftfuehrer", "Betriebsrat RHO",
                    place="Stuttgart", date_str_="14. Oktober 2020")),
    ])


# ─────────────────────────────────────────────────────────────────────────────
# 10) PRJ-2023-001 GATE G3 ADAS-V4D
# ─────────────────────────────────────────────────────────────────────────────
write_doc(f"{BASE}/PRJ-2023-001_Gate_G3_Detailentwicklung_ADAS-V4D_Radar_Fusion.docx",
    H_REA,
    "Gate Review G3 – Detailentwicklung – ADAS-V4D Radar Fusion",
    subtitle="Projekt PRJ-2023-001 – Stage-Gate G3 (Design-Freeze / Pre-A-Sample) – "
             "Konzern-Innovationsprozess KIP-2022 Rev. 4",
    sections=[
        ("Projektidentifikation",
         "Projekt-ID: PRJ-2023-001\n"
         "Projektname: ADAS-V4D Radar-Fusion Steuergeraet (Level-2/3)\n"
         "Hauptkunden / Lead-Customer: Mercedes-Benz Group AG (Plattform MMA), "
         "Stellantis N.V. (Plattform STLA Medium), als Long-Lead-Backup "
         "BMW Brilliance Shenyang (Programm Neue Klasse-CN).\n"
         "Verantwortlicher Bereich: RSG Muenchen (Embedded Software / "
         "ASPICE-Engineering), HW-Entwicklung REG Heilbronn, Asien-Programm "
         "Support RCN Shanghai.\n"
         "Projektleitung: Dr. Klaus Kessler (Werkleiter RSG); Co-Lead Stefan "
         "Hoffmann (CTO bis 30.06.2024) / ab 01.07.2024 Dr. Petra Hollmann (CTO).\n"
         "Gate-Datum: 18. Mai 2023, 09:00 Uhr, Werk Heilbronn / Hybrid Stuttgart"),
        ("Gate-Kriterien (G3-Checklist)",
         [["Krit.", "Beschreibung", "Soll", "Ist", "Status"],
          ["K1", "Lastenheft-Reife (SyRS)", ">= 95%", "97%", "OK"],
          ["K2", "ASPICE Assessment Level", "Level 2 (Pre)", "Level 2 erreicht",
           "OK"],
          ["K3", "Pre-A-Sample (HW) Verfuegbarkeit",
           "10 Stk", "12 Stk (REG/Vaihingen)", "OK"],
          ["K4", "EMV-Pre-Compliance (CISPR 25)",
           "in-spec", "in-spec mit 6 dB Marge", "OK"],
          ["K5", "Sicherheits-Konzept (FuSi ISO 26262)",
           "ASIL B(D)", "ASIL B(D) freigegeben", "OK"],
          ["K6", "Cyber-Security Konzept (ISO 21434)",
           "TARA abgeschlossen", "TARA + Concept Phase abg.", "OK"],
          ["K7", "Kunde Mercedes – Concept-Freeze",
           "schriftl. Zustimmung", "Freigabe vom 12.05.2023", "OK"],
          ["K8", "Kunde Stellantis – Lessons-Learned-Review",
           "geplant", "verschoben auf 30.05.2023", "GELB"],
          ["K9", "Make-or-Buy Radar-Transceiver",
           "Entscheidung", "Make (TI AWR2243 + IFX BGT24-eigene HF-Stage)",
           "OK"],
          ["K10", "Business Case NPV @ 9% WACC",
           ">= 22 Mio. EUR", "27,4 Mio. EUR", "OK"]]),
        ("Gate-Entscheidung",
         "Beschluss des Steering Committees: PROCEED (mit Auflage). Der "
         "Stage Gate G3 wird unter folgender Auflage freigegeben:\n\n"
         "(a) Nachholung des Lessons-Learned-Reviews mit Stellantis bis "
         "30.05.2023; Bericht an CTO-Office.\n"
         "(b) China-spezifische Cyber-Security-Anforderungen (GB/T 40861-2021) "
         "muessen bis G4 in der TARA ergaenzt sein; Verantwortung bei RCN "
         "(Country Manager " + CM + ") in Abstimmung mit RSG.\n"
         "(c) Local-Sourcing-Strategie (Anteil chinesischer Lieferanten 25–30% "
         "bei SOP) ist bis G4 Konzept-faehig."),
        ("Naechste Schritte (Plan)",
         ("list", [
             "G4 (Pilot-Serie / B-Sample): geplant Q1 2024 (Werk Heilbronn)",
             "G5 (PPAP Mercedes Beijing): Q3 2024",
             "SOP Mercedes MMA: Q2 2025 (Beijing-Werk + Sindelfingen)",
             "Stellantis STLA Medium SOP: Q4 2025",
             "China Cyber-Security Re-TARA: bis 31.10.2023 (Verant. RSG/RCN)",
         ])),
        ("Anwesende / Steering Committee",
         "Dr. Thomas Weber (COO, Vorsitz Gate-Board), Stefan Hoffmann (CTO), "
         "Dr. Klaus Kessler (Werkleiter RSG / Projektleitung), Andreas Maier "
         "(Werkleiter REG), Sabine Brand (Q-Leitung REG), " + CM + " (Country "
         "Manager RCN), Lars Wittmann (Lead Developer RSG), Stefan Richter "
         "(CMO/BD), Dr. Yuki Tanaka (CRO Asia ab 01.04.2024 als Gast bereits "
         "anwesend), Florian Maier (Group Controlling)."),
        ("Beschlussbestaetigung",
         signatures("Dr. Thomas Weber", "COO / Vorsitz Gate-Board",
                    "Brennhagen Elektronik AG",
                    "Dr. Klaus Kessler", "Projektleiter / Werkleiter RSG",
                    "Brennhagen Software GmbH",
                    place="Heilbronn / Stuttgart", date_str_="18. Mai 2023")),
    ])


# ─────────────────────────────────────────────────────────────────────────────
# 11) REA-BMW Kuendigungs-Androhung 2023
# ─────────────────────────────────────────────────────────────────────────────
write_doc(f"{BASE}/REA_BMW_Rahmenvertrag_Kuendigung_Androhung_2023.docx", H_REA,
    "Korrespondenz – Rahmenvertrag Kuendigungs-Androhung – BMW Group 2023",
    subtitle="Entwurf eines Mahn- und Eskalationsschreibens an BMW AG / "
             "BMW Brilliance – Programm LightCtrl-7 / iX3-CN",
    draft=True,
    sections=[
        ("Adressat",
         "BMW AG, Petuelring 130, 80809 Muenchen\n"
         "z. Hd. Herrn Dr. Joachim Stoltenberg (Senior Vice President Purchasing "
         "Electrics/Electronics)\n"
         "in Kopie: BMW Brilliance Automotive Ltd., Beijing-Daxing – "
         "Mr. Liu Chen (Director Strategic Sourcing E/E)\n\n"
         "Von: Brennhagen Elektronik AG, Stuttgart – Vorstand, vertreten durch "
         "Stefan Richter (CMO/BD) sowie Stefan Hoffmann (CTO).\n"
         "Aktenzeichen: REA-VTRG-BMW-2023-077-LightCtrl7.\n"
         "Datum (Entwurf): 11. November 2023 – noch NICHT VERSENDET, "
         "Freigabe-Workflow Anwaltskanzlei Hengeler Mueller offen."),
        ("Sachverhalt",
         "Mit dem Rahmenvertrag vom 14. Maerz 2021 (Vertrag-Nr. BMW-PR-21-04432) "
         "verpflichtet sich die Brennhagen Elektronik AG zur Lieferung des Matrix-LED "
         "Steuermoduls »LightCtrl-7« fuer die BMW iX3-CN-Plattform (Werk Shenyang, "
         "BMW Brilliance) mit Soll-Liefervolumen 240.000 Stueck p.a. und einer "
         "Festpreis-Bindung fuer den Zeitraum 2022–2025.\n\n"
         "Seit dem 14. Maerz 2023 sind Forderungen aus 23 Lieferungen "
         "(insgesamt 4,82 Mio. EUR, davon 1,67 Mio. EUR ueberfaellig > 90 Tage) "
         "trotz mehrfacher Mahnung (Schreiben vom 12.06., 22.08. und 18.10.2023) "
         "nicht beglichen. BMW-Werk Shenyang fuehrt die Verzoegerung auf eine "
         "interne SAP-Migration zurueck; diese Begruendung erscheint jedoch "
         "angesichts der Dauer nicht mehr ausreichend.\n\n"
         "Parallel hat BMW Brilliance einen Pre-Audit der Matrix-LED-Module mit "
         "der lokalen Wettbewerberin Desay SV gestartet – nach Einschaetzung "
         "des Country Managers RCN (" + CM + ") ein Indikator fuer eine "
         "drohende Re-Sourcing-Entscheidung."),
        ("Rechtliche Bewertung (Vorab Hengeler Mueller)",
         "Gemaess § 7 des Rahmenvertrages (Zahlungsmodalitaeten) sowie § 12 "
         "(Vertragslaufzeit / Beendigung) ist Brennhagen nach erfolglosem Ablauf "
         "einer angemessenen Nachfrist berechtigt, weitere Lieferungen zu "
         "verweigern und (i) Verzugszinsen 9 PP ueber Basiszinssatz zu fordern, "
         "(ii) den Vertrag aus wichtigem Grund zu kuendigen.\n\n"
         "Empfehlung Hengeler Mueller (Schreiben vom 03.11.2023): Eskalation in "
         "drei Stufen – (1) formelle Mahnung mit 30-Tage-Nachfrist, (2) "
         "Liefer-Stop-Androhung nach Fristablauf, (3) Vertragskuendigung. "
         "Konzern-Risikoabwaegung: Verlust BMW-iX3-CN-Volumen vs. Liquiditaets-"
         "und Reputationsrisiko."),
        ("Vorschlag Schreiben (Auszug)",
         "»Sehr geehrter Herr Dr. Stoltenberg, mit Bedauern muessen wir "
         "feststellen, dass trotz mehrfacher Erinnerungen Forderungen in Hoehe "
         "von 4,82 Mio. EUR aus dem Liefer-Rahmenvertrag vom 14.03.2021 weiterhin "
         "offen sind. Wir setzen Ihnen hiermit eine letzte Nachfrist von 30 "
         "Tagen, gerechnet ab Zugang dieses Schreibens. Sollte bis zum Ablauf "
         "der Frist kein vollstaendiger Zahlungseingang erfolgen, werden wir "
         "von unserem Recht Gebrauch machen, weitere Lieferungen einzustellen "
         "und ggf. den Rahmenvertrag aus wichtigem Grund zu kuendigen. (...)«"),
        ("Interne Freigaben",
         ("list", [
             "Stefan Richter (CMO/BD) – Vorab-OK",
             "Stefan Hoffmann (CTO) – Vorab-OK",
             "Laura Bauer (CFO) – Vorab-OK (Verzugszins-Forderung)",
             "Anna Mueller (CEO) – Final-Approval ausstehend",
             "Hengeler Mueller (Frankfurt) – Pruefung abgeschlossen 03.11.2023",
             "Aufsichtsrat-Information (vorsorglich) durch Prof. Voss",
         ])),
        ("Unterschrift / Status",
         signatures("Stefan Richter", "CMO / BD", R["name"],
                    "Stefan Hoffmann", "CTO", R["name"],
                    place="Stuttgart", date_str_="11. November 2023 (Entwurf)")),
    ])


# ─────────────────────────────────────────────────────────────────────────────
# 12) REA-VW Preisanpassungsschreiben 2023
# ─────────────────────────────────────────────────────────────────────────────
write_doc(f"{BASE}/REA_VW_Preisanpassungsschreiben_2023.docx", H_REA,
    "Korrespondenz – Preisanpassungsschreiben – Volkswagen AG 2023",
    subtitle="Schreiben an Volkswagen AG / Group Procurement Wolfsburg "
             "betr. Material- und Energiepreis-Eskalations-Anpassung",
    sections=[
        ("Adressat",
         "Volkswagen AG, Berliner Ring 2, 38440 Wolfsburg\n"
         "z. Hd. Frau Dr. Karin Stein (Group Senior Vice President Purchasing "
         "Electrics/Electronics & Powertrain)\n"
         "in Kopie: Herr Marco Pirelli (Lead Buyer BMS), Frau Stefanie Lang "
         "(Commodity Manager Powertrain ECUs)\n\n"
         "Von: Brennhagen Elektronik AG, Stuttgart – Vorstand, vertreten durch "
         "Frau Laura Bauer (CFO) und Herrn Stefan Richter (CMO/BD).\n"
         "Aktenzeichen: REA-VW-PA-2023-014\n"
         "Datum: 28. September 2023."),
        ("Anlass",
         "Bezug nehmend auf die Liefer-Rahmenvertraege Nr. VW-EE-21-00876 "
         "(BMS-12 Batteriemanagement, ID.7 / MEB+) sowie Nr. VW-PT-22-00321 "
         "(ECU-900 Powertrain Gen3, MQB-Evo) bitten wir um Aufnahme von "
         "Verhandlungen ueber eine Preisanpassung der Listenpreise mit Wirkung "
         "zum 01.01.2024. Hintergrund sind die seit Vertragsabschluss "
         "eingetretenen wesentlichen Kostensteigerungen, die die ueblichen "
         "Toleranzkorridore deutlich ueberschreiten."),
        ("Indexierung & Begruendung",
         [["Kostenfaktor", "Anteil Materialkosten",
           "Index-Veraenderung 2021–H1 2023",
           "Forderung Anpassung"],
          ["Halbleiter (MCU, Power-IC, DC/DC)",
           "32%", "+47% (SIA Semiconductor Index)",
           "+9,4 PP auf Listenpreis"],
          ["Kupfer (PCB, Stecker)",
           "12%", "+22% (LME Cash Settlement)",
           "+2,6 PP"],
          ["Epoxidharz / FR-4 Substrat",
           "8%", "+38% (lokaler Asien-Spotmarkt)",
           "+3,0 PP"],
          ["Energie (Werk Heilbronn, Bezugs-Strom)",
           "9%", "+72% (EEX Phelix DE/AT Future)",
           "+6,5 PP"],
          ["Logistik / Seefracht China-Europa",
           "5%", "+18% (Drewry WCI)",
           "+0,9 PP"],
          ["Loehne IGM-Tarif 2023",
           "16%", "+5,2% Tarif + Inflations-Ausgleich",
           "+0,8 PP"],
          ["SUMME (gewichtet)",
           "", "", "+23,2 PP auf Listenpreis"]]),
        ("Forderung",
         "Wir fordern eine Anpassung der Listenpreise der oben genannten "
         "Produktfamilien um durchschnittlich +9,8% (gewichtet) mit Wirkung "
         "zum 01.01.2024. Dieser Wert spiegelt einen substanziellen Eigenanteil "
         "der Brennhagen Elektronik AG (Effizienzprogramme, Mehrwerk-Allokation, "
         "Hedging-Erfolge im FX-Bereich) wider und liegt deutlich unter der "
         "nominalen Index-Veraenderung.\n\n"
         "Wir bitten um einen Verhandlungstermin in den Kalenderwochen 42–44, "
         "vorzugsweise bei Volkswagen Konzernzentrale Wolfsburg, hilfsweise "
         "Stuttgart. Auf Wunsch begleiten unsere Konzern-Verrechnungspreis-"
         "Experten (Dr. Heike Berger, Group Tax Director) die Sitzung."),
        ("Hinweis",
         "Wir betonen die Bedeutung der seit 2017 sehr vertrauensvollen "
         "Geschaeftsbeziehung zwischen Volkswagen AG und Brennhagen Elektronik AG. "
         "Sollte eine angemessene Anpassung nicht erfolgen, behalten wir uns "
         "vor, gemaess § 9 (Material-/Energiepreis-Klausel) der genannten "
         "Rahmenvertraege das Sonderkuendigungsrecht zu pruefen. Eine solche "
         "Eskalation moechten wir ausdruecklich vermeiden."),
        ("Unterschriften",
         signatures("Laura Bauer", "CFO", R["name"],
                    "Stefan Richter", "CMO / BD", R["name"],
                    place="Stuttgart", date_str_="28. September 2023")),
    ])


# ─────────────────────────────────────────────────────────────────────────────
# Run verification at the end
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    from docx import Document
    from pathlib import Path
    base = Path(BASE)
    total = thin = 0
    for p in sorted(base.rglob("*.docx")):
        d = Document(p); t = " ".join(par.text for par in d.paragraphs)
        for tbl in d.tables:
            for r in tbl.rows:
                for c in r.cells: t += " " + c.text
        w = len(t.split()); total += 1
        if w < 200:
            thin += 1
            print(f"  THIN({w}): {p.name}")
    print(f"\n{total} total .docx, {thin} still thin")
