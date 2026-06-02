"""Brennhagen AG / 06_Tochter_CZ_Brno – 101 thin docs.

RCZ = Brennhagen CZ s.r.o., Brno, 680 MA, Steckverbinder-Produktion.
Court: KS Brno C 87654. Werkleiter Petr Novak, Q-Leitung Eva Cerna.
Currency CZK (with EUR conversion ~24,5 CZK/EUR for IC).
"""
# --- portable-paths-prelude --- (do not edit) ---
import sys
from pathlib import Path as _PathlibPath
_RP = _PathlibPath(__file__).resolve().parent
while not (_RP / "enhance_lib.py").exists() and _RP.parent != _RP:
    _RP = _RP.parent
_ROOT = _RP
sys.path.insert(0, str(_ROOT))
# --- end prelude ---
import sys
from pathlib import Path
sys.path.insert(0, f"{_ROOT}")
from enhance_lib import ROEHRIG_AG as R, ROEHRIG_SUBS as S, write_doc, signatures

BASE = Path(f"{_ROOT}/roehrig_large/06_Tochter_CZ_Brno")

# ---- Local entity facts -----------------------------------------------------
RCZ_NAME = "Brennhagen CZ s.r.o."
RCZ_ADDR = "Tuzemska 47, 627 00 Brno-Slatina, Ceska republika"
RCZ_HRB  = "C 87654, KS Brno (Krajsky soud v Brne)"
RCZ_ICO  = "ICO 28765432, DIC CZ28765432"
RCZ_HDR  = {"name": RCZ_NAME,
            "addr": RCZ_ADDR + "  |  " + RCZ_ICO,
            "hrb":  RCZ_HRB}

REG_HDR  = {"name": R["name"], "addr": R["addr"], "hrb": R["hrb"]}

# Local plant management & roles
WL  = "Petr Novak"          # Werkleiter / Jednatel
QL  = "Eva Cerna"           # Q-Leitung
PL  = "Tomas Dvorak"        # Produktionsleiter
FC  = "Martina Svobodova"   # Finanzcontroller
HR  = "Lenka Prochazkova"   # HR Manager
GF2 = "Jana Horakova"       # 2. Jednatelka

# Trade union
OZ = "OZ KOVO – Zakladni organizace pri Brennhagen CZ s.r.o. (Vorsitzender Pavel Marek)"

# Advisors
WP_CZ  = "KPMG Ceska republika, s.r.o., Pobrezni 648/1a, 186 00 Praha 8"
ANW_CZ = "Havel & Partners s.r.o., advokatni kancelar, Na Florenci 2116/15, 110 00 Praha 1"

OEM_RECV = "Brennhagen Elektronik GmbH (REG), Heilbronner Strasse 88, 74072 Heilbronn"

# FX baseline (used for narrative; not hedged in this folder)
FX_CZK_EUR = 24.5


def num_cz(n):
    """Czech format with dot as thousands separator (we use space typically; dot OK)."""
    return f"{n:,}".replace(",", ".")


MONTHS_CZ = ["", "leden", "unor", "brezen", "duben", "kveten", "cerven",
             "cervenec", "srpen", "zari", "rijen", "listopad", "prosinec"]
MONTHS_DE = ["", "Januar", "Februar", "Maerz", "April", "Mai", "Juni",
             "Juli", "August", "September", "Oktober", "November", "Dezember"]


# ===========================================================================
# 1) RCZ_Arbeitsvertrag_*  (5 docs) – Czech Arbeitsrecht
# ===========================================================================

ARBEITSVERTRAEGE = [
    ("RCZ_Arbeitsvertrag_01_Geschäftsführer_in_R_2022.docx",
     "Petr Novak", "Jednatel / Werkleiter (Geschaeftsfuehrer)",
     185000, 22, "1. Februar 2022", "neurcito (unbefristet)",
     "Vollzeit 40 h/Woche", True,
     "Berichtsweg an COO Dr. Thomas Weber (REA Stuttgart). Gesamtverantwortung "
     "Werk Brno (680 MA), Budget 1,76 Mrd. CZK Umsatz p. a., Produktionslinien "
     "Steckverbinder fuer ICP-3 / ADAS-V4D / ECU-900 / BMS-12 / LightCtrl-7."),
    ("RCZ_Arbeitsvertrag_02_Produktionsleiter_in_2022.docx",
     "Tomas Dvorak", "Produktionsleiter (Vedouci vyroby)",
     128000, 21, "1. Maerz 2022", "neurcito (unbefristet)",
     "Vollzeit 40 h/Woche, Schichtbereitschaft 3-Schicht", False,
     "Berichtsweg an Werkleiter Petr Novak. Verantwortung 3-Schicht-Betrieb 5 "
     "Steckverbinder-Linien (Crimp-, Spritzguss-, Montagelinien), 540 direkte MA, "
     "Output ca. 18 Mio. Steckverbinder p. a."),
    ("RCZ_Arbeitsvertrag_03_Qualitätsmanagerin_B_2022.docx",
     "Eva Cerna", "Manazerka kvality / Quality Manager",
     115000, 20, "15. Maerz 2022", "neurcito (unbefristet)",
     "Vollzeit 40 h/Woche", False,
     "Berichtsweg dotted line an Group QA Sabine Brand (REG Heilbronn). "
     "Verantwortung IATF 16949 Zertifizierung Werk Brno, ISO 9001:2015, "
     "ISO 14001:2015, VDA 6.3 Lieferanten-Audits, 8D-Prozess."),
    ("RCZ_Arbeitsvertrag_04_Finanzcontroller_Brn_2022.docx",
     "Martina Svobodova", "Financni controller / Finanzcontroller",
     98000, 20, "1. April 2022", "neurcito (unbefristet)",
     "Vollzeit 40 h/Woche", False,
     "Berichtsweg dotted line an Group Controller Florian Maier (REA Stuttgart). "
     "Verantwortung Monats-/Quartalsabschluss CZ-GAAP + IFRS-Reporting-Paket "
     "(IFRS 15/16/9), Forecast-Update FC, MonatsPL, Transfer-Pricing-Dokumentation."),
    ("RCZ_Arbeitsvertrag_05_HR-Manager_Brno_2022.docx",
     "Lenka Prochazkova", "Personalni manazerka / HR Manager",
     92000, 20, "1. Mai 2022", "neurcito (unbefristet)",
     "Vollzeit 40 h/Woche", False,
     "Berichtsweg dotted line an Group HR (REA Stuttgart). Verantwortung HR "
     "fuer 680 MA, Lohnabrechnung (Lohnsteuer CZ via Externer Anbieter ADP CZ), "
     "Verhandlungen mit OZ KOVO, Krankenversicherung VZP/Zamestnanecka pojistovna."),
]


def write_arbeitsvertrag(fname, name, role, base_brutto_czk, urlaub_tage,
                         eintritt, befristung, arbeitszeit, leitend, scope):
    base_eur = round(base_brutto_czk * 12 / FX_CZK_EUR / 1000) * 1000
    sections = [
        ("Smluvni strany / Vertragsparteien",
         f"Zamestnavatel / Arbeitgeber:\n{RCZ_NAME}\n{RCZ_ADDR}\n{RCZ_ICO}\n"
         f"Zapsano: {RCZ_HRB}\nZastoupeno / vertreten durch: {WL}, jednatel\n\n"
         f"Zamestnanec / Arbeitnehmer:\n{name}\nBytem v Brne, CR\n"
         f"Rodne cislo: auf Anfrage / on file (HR-Akte vertraulich)\n\n"
         f"Tato pracovni smlouva (Arbeitsvertrag) se uzavira podle zakona "
         f"c. 262/2006 Sb., zakonik prace (Tschechisches Arbeitsgesetzbuch) "
         f"in der jeweils geltenden Fassung."),
        ("Art. 1 – Druh prace / Taetigkeit",
         f"Zamestnanec se zavazuje vykonavat pro zamestnavatele praci na pozici »{role}«. "
         f"Mistem vykonu prace je sidlo zamestnavatele v Brne ({RCZ_ADDR}). "
         f"Eintrittsdatum / Den nastupu: {eintritt}. "
         f"{scope}"),
        ("Art. 2 – Doba trvani / Vertragsdauer",
         f"Pracovni pomer se sjednava na dobu {befristung}. "
         f"Zkusebni doba / Probezeit: 3 mesice (Manager: 6 mesice) podle § 35 zakoniku prace. "
         f"Vypovedni doba / Kuendigungsfrist: 2 mesice na konci kalendarniho mesice "
         f"(podle § 51 zakoniku prace; bei leitenden Stellen vertraglich 3 Monate)."),
        ("Art. 3 – Pracovni doba / Arbeitszeit",
         f"{arbeitszeit}. Stanovena tydenni pracovni doba podle § 79 zakoniku prace "
         f"betraegt 40 Stunden. Bei Schichtarbeit gelten die Pausen- und Ruheregelungen "
         f"nach §§ 88 f. zakoniku prace. Ueberstunden (prescasova prace) werden mit "
         f"einem Zuschlag von mindestens 25 % auf den Durchschnittsverdienst verguetet "
         f"(§ 114), Nachtarbeit mit 10 % Zuschlag (§ 116), Wochenendarbeit mit 10 % "
         f"Zuschlag (§ 118)."),
        ("Art. 4 – Mzda / Verguetung",
         f"Hrubá mesicni mzda / monatliches Bruttogehalt: {num_cz(base_brutto_czk)} CZK "
         f"(ca. {num_cz(int(base_brutto_czk/FX_CZK_EUR))} EUR). "
         f"Vyplata 13. den nasledujiciho mesice na bankovni ucet. "
         f"Jahresbonus / rocni odmena: bis zu 15 % des Jahresbruttogehalts, "
         f"abhaengig von Ziel-Erreichung (50 % Werks-EBIT Brno, 30 % Quality KPI – "
         f"PPM Customer Complaints, 20 % persoenliche Ziele). "
         f"Bonus-Auszahlung erfolgt im April des Folgejahres. "
         f"Verguetungs-Review jaehrlich zum 1. April. "
         f"Indikative Jahresvergutung total target: ca. "
         f"{num_cz(int(base_brutto_czk*12*1.15))} CZK."),
        ("Art. 5 – Pojisteni / Sozialversicherung",
         f"Zdravotni pojisteni / Krankenversicherung: Zamestnanec je pojisten u "
         f"Vseobecna zdravotni pojistovna CR (VZP, kod 111) oder wahlweise "
         f"Zamestnanecka pojistovna Skoda (kod 209). Arbeitgeberanteil 9 % vom "
         f"Bruttogehalt, Arbeitnehmeranteil 4,5 %. "
         f"Sociální pojisteni / Sozialversicherung (penzijni + nemocenske + statni "
         f"politika zamestnanosti) gemaess Zakon c. 589/1992 Sb.: Arbeitgeber 24,8 %, "
         f"Arbeitnehmer 6,5 %. Die Anmeldung bei OSSZ Brno-mesto erfolgt durch HR."),
        ("Art. 6 – Dovolena / Urlaub",
         f"Zakonni narok na dovolenou / gesetzlicher Urlaub: 4 Wochen pro Kalenderjahr "
         f"(§ 213 zakoniku prace). Vertraglich gewaehrt der Arbeitgeber {urlaub_tage} "
         f"Werktage / {round(urlaub_tage/5,1)} Wochen Urlaub (Betriebsstandard 5 Wochen + 5 Sick-Days). "
         f"Urlaubsplanung erfolgt jaehrlich bis 31.3. (Rocni plan dovolene). "
         f"Resturlaub aus Vorjahr verfaellt zum 31.3. des Folgejahres, sofern "
         f"nicht aus betrieblichen Gruenden uebertragen."),
        ("Art. 7 – Mlcenlivost / Verschwiegenheit / NDA",
         f"Zamestnanec se zavazuje k zachovani mlcenlivosti o vsech obchodnich, "
         f"technickych a vyrobnich informacich. Geschuetzte Informationen umfassen "
         f"insbesondere: OEM-Kundendaten (BMW, VW, Mercedes, Stellantis), Stueckpreise, "
         f"Konstruktionsdaten Steckverbinder (ICP-3/ADAS-V4D/ECU-900/BMS-12/LightCtrl-7), "
         f"Lieferantendaten, Preise und Konditionen. "
         f"Konkurencni dolozka / Wettbewerbsverbot (§ 310 zakoniku prace): "
         f"12 Monate nach Vertragsende, ortlich CZ + SK + PL + DE, fuer alle "
         f"OEM-Steckverbinder-Hersteller (TE Connectivity, Yazaki, Aptiv, JST, Molex). "
         f"Kompenzace / Karenzentschaedigung: 1,0 monatliches Durchschnittsgehalt."),
        ("Art. 8 – Skoleni / Schulungen",
         f"Pflichtschulungen jaehrlich: ISO 9001 / IATF 16949 (Q-Bewusstsein), "
         f"ISO 14001 (Umwelt), Arbeitssicherheit / BOZP (§ 103 zakoniku prace), "
         f"Brandschutz (PO), GDPR/AVG. Kostenuebernahme zu 100 % durch Arbeitgeber. "
         f"Bei externen Weiterbildungen ueber 50.000 CZK Zurueckzahlungsklausel "
         f"(prorata 36 Monate) gemaess § 234 zakoniku prace."),
        ("Art. 9 – Zaverecna ustanoveni / Schlussbestimmungen",
         f"Diese Vertrag wird in zweifacher Ausfertigung in tschechischer und "
         f"deutscher Sprache ausgefertigt; bei Widerspruechen ist die tschechische "
         f"Fassung massgeblich. Aenderungen beduerfen der Schriftform. "
         f"Gerichtsstand / Soud: Mestsky soud v Brne. "
         f"Anwendbares Recht: tschechisches Arbeitsrecht (Zakon c. 262/2006 Sb. "
         f"v platnem zneni)."),
        ("Podpisy / Unterschriften",
         signatures(WL, "Jednatel", RCZ_NAME,
                    name, role, "(Zamestnanec)",
                    place="Brno", date_str_=eintritt)),
    ]
    write_doc(BASE / fname, RCZ_HDR,
              f"Pracovni smlouva / Arbeitsvertrag – {name}",
              sections,
              subtitle=f"Position: {role} | Eintritt: {eintritt} | "
                       f"Brutto-Grundgehalt {num_cz(base_brutto_czk)} CZK / Monat")


for args in ARBEITSVERTRAEGE:
    write_arbeitsvertrag(*args)


# ===========================================================================
# 2) RCZ_to_REG_IC_*  (35 docs) – Intercompany Rechnungen RCZ -> REG
# 3) RCZ_IC_Rechnung_*  (~25 docx Rechnungen, plus 12 Quartalsberichte 2019-2020)
# ===========================================================================

def ic_rechnung(fname, header, sender, sender_addr, sender_id,
                recv, recv_addr, recv_id,
                rg_nr, jahr, monat, leistung_lines, total_czk,
                czk_eur=FX_CZK_EUR, draft=False, mwst_rev_charge=True):
    total_eur = round(total_czk / czk_eur, 2)
    rows = [["Pos.", "Bezeichnung", "Menge", "Einzelpreis CZK",
             "Gesamtpreis CZK"]]
    sum_calc = 0
    for i, (bez, menge, ep) in enumerate(leistung_lines, 1):
        gp = menge * ep
        sum_calc += gp
        rows.append([str(i), bez, f"{menge:,}".replace(",", "."),
                     num_cz(ep), num_cz(gp)])
    rows.append(["", "Zwischensumme netto", "", "", num_cz(sum_calc)])
    if mwst_rev_charge:
        rows.append(["", "MWSt (Reverse Charge §92a CZ-USt)", "", "", "0"])
    else:
        mwst = int(sum_calc * 0.21)
        rows.append(["", "MWSt 21 % CZ", "", "", num_cz(mwst)])
        sum_calc += mwst
    rows.append(["", "Rechnungsbetrag CZK", "", "", num_cz(sum_calc)])
    rows.append(["", f"Umrechnung CZK->EUR (Kurs {czk_eur:.3f})", "",
                 "", f"{num_cz(int(sum_calc))} CZK = {total_eur:,.2f} EUR"])

    sections = [
        ("Rechnungsdaten",
         f"Rechnungsnummer / Cislo faktury: {rg_nr}\n"
         f"Rechnungsdatum / Datum vystaveni: {15}. {MONTHS_DE[monat]} {jahr}\n"
         f"Leistungsdatum / Datum uskutecneneho zdanitelneho plneni (DUZP): "
         f"letzter Werktag {MONTHS_DE[monat]} {jahr}\n"
         f"Faelligkeit / Splatnost: 60 Tage netto (Konzernfristen / Group Terms)\n"
         f"Lieferbedingungen / Dodaci podminky: FCA Brno, Incoterms 2020\n"
         f"Variabilni symbol: {rg_nr.replace('-', '').replace('/', '')[-10:]}"),
        ("Leistender / Dodavatel",
         f"{sender}\n{sender_addr}\n{sender_id}\n\n"
         f"Bankverbindung: Komercni banka a.s., Brno\n"
         f"IBAN: CZ65 0100 0000 0011 2233 4455\nBIC: KOMBCZPP\n"
         f"Konto fuer EUR-Zahlungen: CZ22 0100 0000 0099 8877 6655"),
        ("Empfaenger / Odberatel",
         f"{recv}\n{recv_addr}\n{recv_id}\n\n"
         f"Bestell-Referenz / Objednavka: PO-{jahr}-{monat:02d}-REG-CZ"),
        ("Leistungsumfang / Predmet plneni",
         "Lieferung von Steckverbindern und Subassemblies fuer die Module "
         "ICP-3 / ADAS-V4D / ECU-900 / BMS-12 / LightCtrl-7 gemaess "
         "Rahmenliefervertrag REG-RCZ 2020 (interner Vertrag IC-RCZ-REG-2020). "
         "Verrechnung erfolgt zu Konzern-Verrechnungspreisen (Cost-Plus 7,5 %) "
         "gemaess Transfer-Pricing-Master-File 2022 und Local-File RCZ 2022. "
         "Die Bestellung ist Teil der monatlichen Roll-up-Belieferung zur Linie 1-4 "
         "des Werks Heilbronn (REG)."),
        ("Positionen", rows),
        ("MWSt / DPH-Behandlung",
         "Reverse Charge gemaess § 92a tschechisches Mehrwertsteuergesetz "
         "(Zakon c. 235/2004 Sb. o dani z pridane hodnoty) – innergemeinschaftliche "
         "Lieferung zwischen RCZ (DIC CZ28765432) und REG (USt-IdNr. DE312487901). "
         "Steuerschuld geht auf den Empfaenger ueber. Zusammenfassende Meldung "
         "(souhrnne hlaseni) erfolgt monatlich bis zum 25. des Folgemonats an "
         "Finanzamt Brno I. EC-Sales-List-Code: 31 (Warenlieferung). "
         "Intrastat-Meldung (Versendung CN-Code 8536 7000) erfolgt durch Logistik."),
        ("Zahlungsbedingungen",
         f"Zahlung in CZK auf das oben genannte Komercni banka Konto. "
         f"Wahlweise EUR-Zahlung zum Tageskurs CNB des Rechnungsdatums (Stand "
         f"{15}. {MONTHS_DE[monat]} {jahr}: {czk_eur:.3f} CZK/EUR, Quelle: "
         f"Ceska narodni banka). Zahlungsverzug: gesetzliche Zinsen nach NV "
         f"c. 351/2013 Sb. (CNB-Repo + 8 %-Pkte.). Skonto: nicht gewaehrt "
         f"(Konzern-Standardkonditionen)."),
        ("Buchhalterische Erfassung",
         "Erfassung im SAP S/4HANA (Mandant 700 RCZ, Buchungskreis CZ01) auf "
         "Erloeskonto 311000 (Intercompany-Erloese), Steuerkennzeichen RC. "
         "Im Empfaengersystem REG (Mandant 100, BUKR DE01) wird die Eingangsrechnung "
         "auf Konto 421000 (IC-Verbindlichkeiten RCZ) gebucht. "
         "Konzernabstimmung erfolgt monatlich via SAP IC-Matching (Group Reporting "
         "Tool, Florian Maier, Group Controller)."),
        ("Genehmigung / Schlussvermerk",
         f"Rechnungsausstellung durch Finanzcontroller RCZ. "
         f"Sachliche Pruefung: {PL}, Produktionsleiter (Mengen). "
         f"Rechnerische Pruefung: {FC}. Freigabe Werkleiter {WL}. "
         f"{'*** ENTWURF / NAVRH – noch nicht freigegeben ***' if draft else ''}"),
    ]
    write_doc(BASE / fname, header,
              f"Faktura / Rechnung Nr. {rg_nr}",
              sections,
              subtitle=f"RCZ -> REG | {MONTHS_DE[monat]} {jahr} | "
                       f"{num_cz(int(total_czk))} CZK (ca. {total_eur:,.0f} EUR)",
              draft=draft)


# Build IC-Rechnungen
def gen_lines(monat, jahr, base_factor=1.0):
    """Generate 4-6 line items for a monthly IC invoice."""
    # Steckverbinder volume scaled by month with some variability
    f = base_factor
    return [
        (f"Steckverbinder MQS-12 fuer ECU-900 (Pin-Header)",
         int(95000*f) + monat*150, 32),
        (f"Steckverbinder MLK-08 fuer ICP-3 (Infotainment)",
         int(48000*f) + monat*80, 41),
        (f"Steckverbinder HV-3-fach fuer BMS-12 (Hochvolt)",
         int(18000*f) + monat*30, 188),
        (f"Steckverbinder Coax FAKRA-Z fuer ADAS-V4D",
         int(22000*f) + monat*40, 75),
        (f"Subassembly Kabelsatz LightCtrl-7 (1,2 m)",
         int(12000*f) + monat*20, 240),
    ]


def total_from(lines):
    return sum(m * p for _, m, p in lines)


# RCZ_to_REG_IC_*  (2021-01..12, 2022-01..12, 2023-01..12 minus skips)
to_reg_months = []
for jahr, factor, skip in [(2021, 0.90, []),
                           (2022, 0.95, []),
                           (2023, 1.00, [9])]:
    for m in range(1, 13):
        if m in skip:
            continue
        to_reg_months.append((jahr, m, factor))

# Some are with FINAL/ENTWURF suffix; map filename style
to_reg_specials = {
    (2023, 1): "RCZ_to_REG_IC_2023_01_FINAL.docx",
    (2023, 6): "RCZ_to_REG_IC_2023_06_ENTWURF.docx",
}

for jahr, m, factor in to_reg_months:
    fname = to_reg_specials.get((jahr, m), f"RCZ_to_REG_IC_{jahr}_{m:02d}.docx")
    rg_nr = f"RCZ-REG-{jahr}-{m:02d}-001"
    lines = gen_lines(m, jahr, factor)
    total = total_from(lines)
    draft = "ENTWURF" in fname
    ic_rechnung(fname, RCZ_HDR,
                RCZ_NAME, RCZ_ADDR, RCZ_ICO,
                "Brennhagen Elektronik GmbH (REG)",
                "Heilbronner Strasse 88, 74072 Heilbronn, Deutschland",
                "HRB 221456, Amtsgericht Heilbronn | USt-IdNr. DE312487901",
                rg_nr, jahr, m, lines, total, draft=draft)


# RCZ_IC_Rechnung_*  (2020-01..12 minus 02,09; 2021-01..12; 2022-01..11 minus skips)
ic_specs = []
for jahr, skip in [(2020, [2, 9]),
                   (2021, []),
                   (2022, [12])]:
    for m in range(1, 13):
        if m in skip:
            continue
        ic_specs.append((jahr, m))

ic_specials = {
    (2022, 9): "RCZ_IC_Rechnung_2022_09_DRAFT.docx",
}

for jahr, m in ic_specs:
    fname = ic_specials.get((jahr, m), f"RCZ_IC_Rechnung_{jahr}_{m:02d}.docx")
    if not (BASE / fname).exists():
        # skip if file actually doesn't exist
        continue
    rg_nr = f"RCZ-IC-{jahr}-{m:02d}-{200+m:03d}"
    lines = gen_lines(m, jahr, 0.85 if jahr == 2020 else (0.90 if jahr == 2021 else 0.95))
    total = total_from(lines)
    draft = "DRAFT" in fname
    ic_rechnung(fname, RCZ_HDR,
                RCZ_NAME, RCZ_ADDR, RCZ_ICO,
                "Brennhagen Elektronik GmbH (REG)",
                "Heilbronner Strasse 88, 74072 Heilbronn, Deutschland",
                "HRB 221456, Amtsgericht Heilbronn | USt-IdNr. DE312487901",
                rg_nr, jahr, m, lines, total, draft=draft)


# ===========================================================================
# 4) RCZ_IC_Quartalsbericht_*  (2019 Q1-Q4, 2020 Q1, Q2, Q4) – 7 docs
# ===========================================================================

def ic_quartalsbericht(fname, jahr, quartal, umsatz_czk, ebit_czk, ppm, oee):
    m_start = (quartal - 1) * 3 + 1
    m_end   = m_start + 2
    sections = [
        ("Berichtszeitraum / Vykazovaci obdobi",
         f"Quartal: Q{quartal}/{jahr} ({MONTHS_DE[m_start]} - {MONTHS_DE[m_end]} {jahr})\n"
         f"Berichtsersteller: {FC} (Financni controller)\n"
         f"Berichtsempfaenger: Florian Maier (Group Controller REA), "
         f"Andreas Maier (Werkleiter REG)\n"
         f"Berichtswaehrung: CZK (mit Umrechnungstabelle EUR)\n"
         f"Vorlagestichtag: 20. Werktag nach Quartalsende\n"
         f"Reporting-Standard: IFRS-Paket (gem. Konzernrichtlinie REA-IFRS-2022) + "
         f"CZ-GAAP fuer lokale Statutory."),
        ("Executive Summary",
         f"Werk Brno hat in Q{quartal}/{jahr} einen Umsatz von "
         f"{num_cz(umsatz_czk)} TCZK (ca. "
         f"{num_cz(int(umsatz_czk/FX_CZK_EUR))} TEUR) erzielt, davon "
         f"~88 % aus Intercompany-Lieferungen an REG Heilbronn und ~12 % "
         f"aus externen Kunden (kleine OEM-Direktauftraege Skoda/Tatra "
         f"sowie Aftermarket). EBIT betraegt {num_cz(ebit_czk)} TCZK "
         f"(EBIT-Marge {ebit_czk/umsatz_czk*100:.1f} %). "
         f"Operative Performance solid: OEE {oee} %, "
         f"Customer Complaints PPM {ppm} (Target < 25 PPM)."),
        ("Umsatz nach Produktlinie (TCZK)",
         [["Produkt", "Q-Ist", "Q-Plan", "Abw.", "YTD-Ist"],
          ["Steckverbinder ECU-900", num_cz(int(umsatz_czk*0.34)),
           num_cz(int(umsatz_czk*0.32)), "+6 %",
           num_cz(int(umsatz_czk*0.34*quartal))],
          ["Steckverbinder ICP-3",   num_cz(int(umsatz_czk*0.22)),
           num_cz(int(umsatz_czk*0.24)), "-8 %",
           num_cz(int(umsatz_czk*0.22*quartal))],
          ["Steckverbinder ADAS-V4D",num_cz(int(umsatz_czk*0.16)),
           num_cz(int(umsatz_czk*0.15)), "+7 %",
           num_cz(int(umsatz_czk*0.16*quartal))],
          ["HV-Stecker BMS-12",      num_cz(int(umsatz_czk*0.14)),
           num_cz(int(umsatz_czk*0.12)), "+17 %",
           num_cz(int(umsatz_czk*0.14*quartal))],
          ["LightCtrl-7 Kabelsatz",  num_cz(int(umsatz_czk*0.08)),
           num_cz(int(umsatz_czk*0.09)), "-11 %",
           num_cz(int(umsatz_czk*0.08*quartal))],
          ["Externe / Aftermarket",  num_cz(int(umsatz_czk*0.06)),
           num_cz(int(umsatz_czk*0.08)), "-25 %",
           num_cz(int(umsatz_czk*0.06*quartal))]]),
        ("Ergebnis (TCZK)",
         [["Position", f"Q{quartal} Ist", "Plan", "Abw.", "Kommentar"],
          ["Umsatz",          num_cz(umsatz_czk), num_cz(int(umsatz_czk*0.97)),
           f"+{(umsatz_czk-int(umsatz_czk*0.97))/int(umsatz_czk*0.97)*100:.1f} %",
           "BMS-12 Hochlauf staerker als geplant"],
          ["Materialaufwand", num_cz(int(umsatz_czk*0.58)),
           num_cz(int(umsatz_czk*0.55)), "+5 %",
           "Kupfer- und Harz-Preise +12 %"],
          ["Personalaufwand", num_cz(int(umsatz_czk*0.18)),
           num_cz(int(umsatz_czk*0.18)), "0",
           "lt. Tarif OZ KOVO 2023 +6,5 %"],
          ["Abschreibungen",  num_cz(int(umsatz_czk*0.06)),
           num_cz(int(umsatz_czk*0.06)), "0", "planmaessig"],
          ["Sonstige Aufw.",  num_cz(int(umsatz_czk*0.10)),
           num_cz(int(umsatz_czk*0.10)), "0", "Energie, Wartung, IT"],
          ["EBIT",            num_cz(ebit_czk), num_cz(int(ebit_czk*1.05)),
           "-5 %",            "Materialdruck noch nicht ueberwaelzt"]]),
        ("Quality KPI",
         f"Customer Complaints PPM (extern): {ppm} (Target <25, Vorjahr {ppm+8})\n"
         f"Internal PPM: {ppm*4} (Target <120)\n"
         f"OEE Gesamtwerk: {oee} % (Target 85 %)\n"
         f"First Pass Yield (FPY) Linien 1-5: 96,8 % / 97,4 % / 95,2 % / 98,1 % / 97,7 %\n"
         f"Reklamationen Q{quartal}: 8 Stueck (davon 2 mit 8D-Report eingeleitet)\n"
         f"IATF-16949 Surveillance-Audit {jahr}: bestanden, 3 Minor / 0 Major."),
        ("Personal / Headcount",
         f"Stichtag Quartalsende: 680 MA (Plan 685). "
         f"Davon 540 in Produktion (3-Schicht), 80 Quality/Engineering, "
         f"35 Administration, 25 Logistik. "
         f"Krankenstand Q{quartal}: 4,1 % (Vorjahr 3,8 %, Branche CZ Auto 4,5 %). "
         f"Fluktuation YTD: 8,2 % p. a. (CZ Auto Benchmark 12 %). "
         f"Tarifverhandlungen mit OZ KOVO – Tarifabschluss +6,5 % Tabellenentgelte "
         f"ab 1.4. dieses Jahres erfolgreich abgeschlossen."),
        ("Investitionen",
         f"CapEx Q{quartal}: {num_cz(int(umsatz_czk*0.04))} TCZK "
         f"(geplant {num_cz(int(umsatz_czk*0.05))} TCZK). "
         f"Wichtigste Projekte: Erweiterung Linie 3 Crimp-Automat (8,5 Mio. CZK), "
         f"Energiemanagement ISO 50001 Vorbereitung (1,2 Mio. CZK), "
         f"SAP S/4HANA Migration Vorarbeiten (0,8 Mio. CZK)."),
        ("Ausblick",
         f"Q{quartal+1 if quartal<4 else 1}/{jahr if quartal<4 else jahr+1} "
         f"erwartet vergleichbares Umsatzniveau, weiterhin Druck auf Materialkosten. "
         f"Forecast-Update folgt mit RCZ_FC_Update_{jahr}_Q{quartal+1 if quartal<4 else 1}."),
        ("Schlussvermerk",
         f"Brno, {15+quartal}. {MONTHS_DE[m_end+1 if m_end<12 else 1]} "
         f"{jahr if m_end<12 else jahr+1}\n\n"
         f"{FC}, Financni controller    |    {WL}, Jednatel (sign-off)"),
    ]
    write_doc(BASE / fname, RCZ_HDR,
              f"Quartalsbericht Q{quartal}/{jahr} – Werk Brno (Intercompany Reporting)",
              sections,
              subtitle=f"Reporting an REA Group Controlling / REG Heilbronn | "
                       f"CZK / IFRS")


for jahr, quartals in [(2019, [1, 2, 3, 4]), (2020, [1, 2, 4])]:
    for q in quartals:
        umsatz = 380000 + (jahr - 2019) * 25000 + q * 6000  # TCZK
        ebit   = int(umsatz * 0.062)
        ppm    = 22 - q
        oee    = 82 + q // 2
        ic_quartalsbericht(f"RCZ_IC_Quartalsbericht_{jahr}_Q{q}.docx",
                           jahr, q, umsatz, ebit, ppm, oee)


# ===========================================================================
# 5) RCZ_ControlPlan_*  (4 docs)  ICP-3, BMS-12, ECU-900, ADAS-V4D
# ===========================================================================

CP_PROD = {
    "ICP-3":     ("Infotainment-Steckverbinder MLK-08 fuer ICP-3", "MLK-08-RCZ", 41,
                  "BMW G05/G07, Mercedes EQE/EQS"),
    "BMS-12":    ("HV-Steckverbinder 3-polig 200 A fuer BMS-12 EV", "HV3P-200A-RCZ", 188,
                  "VW ID.7, Hyundai Ioniq 6"),
    "ECU-900":   ("Pin-Header MQS-12 fuer Powertrain-ECU Gen3",      "MQS-12-RCZ", 32,
                  "VW MEB+/MQB-Evo, Stellantis STLA"),
    "ADAS-V4D":  ("FAKRA-Z Coax-Steckverbinder fuer ADAS Radar",      "FAKRA-Z-RCZ", 75,
                  "Mercedes EQS, Stellantis Cassino"),
}


def controlplan(fname, key, prod, partno, ep_czk, oem):
    sections = [
        ("Dokumentstamm",
         f"Control-Plan-Nr.: CP-RCZ-{partno}-Rev05\n"
         f"Status: in Serie (SOP erfolgt Q3/2021)\n"
         f"Produkt: {prod}\n"
         f"Konzernprodukt-Familie: {key}\n"
         f"OEM-Anwendung: {oem}\n"
         f"Verantwortliche: {QL} (Q-Leitung), {PL} (Produktionsleitung)\n"
         f"Genehmigt durch: {WL} (Werkleiter)\n"
         f"Naechste Pruefung: jaehrlich, Q4 {2024}\n"
         f"IATF-16949 Bezug: Abschnitt 8.5.1.1 + Anhang A (Control Plan)"),
        ("Geltungsbereich",
         f"Dieser Control-Plan beschreibt die durchgaengige Qualitaetslenkung "
         f"des Bauteils »{prod}« (Sachnummer RCZ {partno}) ueber alle Prozessschritte "
         f"von Wareneingang bis Versand zum Kunden REG Heilbronn (Linie 1-4) "
         f"sowie Direktbelieferung externer OEM-Tier-1 ({oem.split(',')[0]}). "
         f"Geltend fuer Linie 2 und Linie 4 des Werks Brno. Vorgaben gemaess "
         f"IATF 16949:2016, VDA Band 6.3, ISO/TS-Kundenanforderungen "
         f"(BMW QPN, VW Formel-Q, MB MBN, Stellantis CSR)."),
        ("Prozessschritte (Auszug)",
         [["Op.", "Prozess", "Maschine", "Pruefmerkmal", "Spez.", "Pruefmittel",
           "Frequenz", "CpK-Ziel"],
          ["10", "Wareneingang Cu-Litze",  "WE-Pruefstand",
           "Durchmesser, Festigkeit", "0,35 mm +/- 0,02",
           "Mikrometer, Zugpruefmaschine", "100 % Charge", ">=1,33"],
          ["20", "Stanzen Kontakte",       "Bruderer BSTA-40",
           "Kontur, Grat",            "Klasse A0 nach VDA",
           "Profilprojektor Mitutoyo PJ-A3000", "alle 500 Stk.", ">=1,67"],
          ["30", "Galvanik Sn/Ag",         "Inline-Galvanik",
           "Schichtdicke Sn",         "3,0-5,0 μm",
           "Roentgenfluoreszenz Fischerscope", "alle 1000 Stk.", ">=1,33"],
          ["40", "Spritzguss Gehaeuse",    "Arburg 470A",
           "Massmasse, Volumen",      "Zeichnung Rev. 5, Tol. ISO 2768-mK",
           "Messlehre, Waage Sartorius", "1/Stunde", ">=1,33"],
          ["50", "Crimpen Kontakte",       "Schaefer KSA 5500",
           "Crimp-Hoehe, Abzugskraft","H=1,15 +/- 0,03 mm, F>=120 N",
           "Mikroschliff CP-Lab, Zugpruefmaschine", "1/Schicht (5 Stk.)", ">=1,67"],
          ["60", "Montage",                "Manuelle Montagezelle",
           "Verriegelung, Optik",     "klick + 360° Sichtpruefung",
           "Klick-Tester, Sichtkontrolle 100 %", "100 %", "-"],
          ["70", "Endpruefung elektrisch", "InCircuit-Tester Test-Cell-CZ-04",
           "Isolation, Durchgang",    ">100 MOhm bei 500 V, R<10 mOhm",
           "Mega-Ohmmeter, Mikro-Ohmmeter", "100 %", "-"],
          ["80", "Verpackung",             "Tray-Stacker",
           "Korrekte Etikettierung",   "RDS-Label + DPM 2D-Datamatrix",
           "Scanner Datalogic Magellan", "100 %", "-"],
          ["90", "Versand",                "Logistikzone",
           "Vollstaendigkeit, ASN",    "SAP-EDI-ASN VDA 4905",
           "Wiegen, Scanner", "100 %", "-"]]),
        ("Reaktionsplan",
         "Bei Ueberschreitung der Spezifikationsgrenzen erfolgt sofortige Sperrung "
         "der Charge ueber SAP-QM (Q-Sperr-Status SK), Information an Schicht- und "
         "Q-Leitung. 8D-Report wird innerhalb 24 h bei Customer Complaint geoeffnet, "
         "Containment-Action D3 binnen 48 h, Root-Cause D4 binnen 10 Werktagen "
         "(Ishikawa + 5-Why), Permanente Korrekturmassnahmen D5/D6 binnen 30 Tagen. "
         "PPAP-Lieferung bei Aenderungen an Werkzeug oder Prozessparameter > 10 %."),
        ("Pruefmittelfaehigkeit (MSA)",
         "Alle Messmittel jaehrlich kalibriert durch CMI (Cesky metrologicky "
         "institut) bzw. interne Kalibrierwerkstatt. MSA Typ 1, 2, 3 nach VDA Band 5 "
         "(Type 1: Cg/Cgk >= 1,33; Type 2/3 Bias <= 10 % Toleranzbreite). "
         "Eskalationspfad bei MSA-Fail: Sperrung Pruefmittel + Re-Pruefung der "
         "Charge mit Backup-Equipment."),
        ("Freigabe / Schwenkung",
         f"Erstellt: {QL} (Manazerka kvality), {15}. Mai 2022\n"
         f"Geprueft: {PL} (Vedouci vyroby), {18}. Mai 2022\n"
         f"Freigegeben: {WL} (Jednatel), {20}. Mai 2022\n"
         f"Verteiler: Q-Wesen, Produktion Linie 2+4, Logistik, REG Heilbronn (Sabine Brand)\n"
         f"Aenderungshistorie Rev. 1-5: Rev.1 SOP-Vorbereitung 06/2021; "
         f"Rev.2 PPAP-Korrektur 09/2021; Rev.3 Werkzeug-Mod 02/2022; "
         f"Rev.4 Lieferantenwechsel Cu-Litze 11/2022; Rev.5 OEE-Optimierung 05/2023."),
    ]
    write_doc(BASE / fname, RCZ_HDR,
              f"Control Plan – {prod}",
              sections,
              subtitle=f"Sachnummer {partno} | OEM: {oem} | "
                       f"IATF 16949 / VDA 6.3")


for key, (prod, partno, ep, oem) in CP_PROD.items():
    controlplan(f"RCZ_ControlPlan_{key}_2023.docx", key, prod, partno, ep, oem)


# ===========================================================================
# 6) RCZ_Steuerbescheid_*  (3 docs) 2020, 2021, 2022
# ===========================================================================

def steuerbescheid(fname, jahr, ergebnis_czk, koerperschaftsteuer_pct=19):
    steuer = int(ergebnis_czk * koerperschaftsteuer_pct / 100)
    sections = [
        ("Adressat",
         f"{RCZ_NAME}\n{RCZ_ADDR}\n{RCZ_ICO}\n\n"
         f"Zu Haenden: {WL}, jednatel\nKopie: {FC} (Financni controller)"),
        ("Aussteller / Vydavatel",
         f"Specializovany financni urad\nUzemni pracoviste pro Jihomoravsky kraj\n"
         f"Trida Generala Pikya 11, 613 00 Brno\nReferent: Ing. Pavel Soukup, "
         f"Inspektor T. {2400+jahr}\nDatum vystaveni: 30. zari {jahr+1} "
         f"(30. September {jahr+1})\nCislo jednaci: SFU-BR-{jahr}-{int(ergebnis_czk/1000)%10000}/2024"),
        ("Rozhodnuti / Bescheid",
         f"Vec: Vymerove rozhodnuti k dani z prijmu pravnickych osob za zdanovaci "
         f"obdobi {jahr} (Koerperschaftsteuer-Festsetzung).\n\n"
         f"Auf Grundlage der eingereichten Steuererklaerung (DPPO) vom "
         f"30. cervna {jahr+1} und der durchgefuehrten Pruefung wird hiermit "
         f"die Koerperschaftsteuer (Dan z prijmu pravnickych osob, DPPO) der "
         f"Brennhagen CZ s.r.o. fuer das Jahr {jahr} verbindlich festgesetzt."),
        ("Festsetzung",
         [["Position", "Betrag CZK", "Kommentar"],
          ["Zu versteuerndes Einkommen (zaklad dane)", num_cz(ergebnis_czk),
           "nach Anpassungen § 23 Zakon c. 586/1992 Sb."],
          ["./. Verlustvortrag (ztrata)", "0", "kein Verlustvortrag vorhanden"],
          ["./. F&E-Abzug (odpocet vyzkum a vyvoj § 34)", num_cz(int(ergebnis_czk*0.04)),
           f"F&E-Aufwand {num_cz(int(ergebnis_czk*0.04))} CZK x 100 %"],
          ["./. Spenden (dary § 20)", "150.000",
           "Spende Masaryk-Universitaet"],
          ["Steuerbemessungsgrundlage nach Abzuegen", num_cz(int(ergebnis_czk*0.96)),
           "gerundet auf volle 1000 CZK"],
          [f"Steuersatz {koerperschaftsteuer_pct} %", num_cz(int(ergebnis_czk*0.96*koerperschaftsteuer_pct/100)),
           "§ 21 Zakon c. 586/1992 Sb."],
          ["Festgesetzte Steuer", num_cz(steuer), ""],
          ["./. Vorauszahlungen (zalohy) geleistet", num_cz(int(steuer*0.85)),
           "Quartals-Vorauszahlungen 2x p.a."],
          ["Verbleibender Zahlbetrag", num_cz(int(steuer*0.15)),
           f"faellig 30 Tage nach Zustellung"]]),
        ("Solidaritaetszuschlag / weitere Abgaben",
         f"In der Tschechischen Republik existiert kein Solidaritaetszuschlag "
         f"auf KSt; die KSt-Belastung umfasst 19 % (ab Steuerjahr 2024: 21 %). "
         f"Lokale Gewerbesteuer existiert nicht (anders als in DE). "
         f"Zusaetzlich faellig sind Grundsteuer (dan z nemovitych veci) ca. "
         f"180.000 CZK p. a. fuer das Werksgelaende, sowie KFZ-Steuer Flotte "
         f"ca. 95.000 CZK p. a."),
        ("Rechtsbehelfsbelehrung / Pouceni o opravnych prostredcich",
         "Gegen diesen Bescheid kann innerhalb von 30 Tagen ab Zustellung "
         "Beschwerde (odvolani) beim Specializovany financni urad – Generalni "
         "financni reditelstvi – eingelegt werden (§ 109 ff. Steuerverwaltungsordnung "
         "Zakon c. 280/2009 Sb., danovy rad). Die Beschwerde hat keine aufschiebende "
         "Wirkung; ein Antrag auf Aussetzung der Vollziehung kann gestellt werden."),
        ("Zahlungsabwicklung",
         f"Konto: 7704-67625001/0710 (Ceska narodni banka, fuer KSt)\n"
         f"Variabilni symbol: {RCZ_ICO.split()[1]}\nKonstantni symbol: 1148\n"
         f"Verwendungszweck: DPPO {jahr} Brennhagen CZ s.r.o."),
        ("Schlussvermerk Brennhagen CZ (internes Cover-Sheet)",
         f"Bescheid eingegangen am 5. rijna {jahr+1}. Pruefung durch {FC} "
         f"abgeschlossen am 12. rijna {jahr+1}. Abstimmung mit Group Tax "
         f"Director Dr. Heike Berger (REA Stuttgart) am 14. rijna {jahr+1}: "
         f"Festsetzung deckt sich mit Erwartung +/- 1,2 % (geringfuegige "
         f"Anpassung bei Abschreibung Anlagen IFRS vs. CZ-GAAP). Zahlung "
         f"freigegeben durch {WL} per Banking-Tool {15}. rijna {jahr+1}."),
    ]
    write_doc(BASE / fname, RCZ_HDR,
              f"Vymerove rozhodnuti k DPPO {jahr} / Koerperschaftsteuerbescheid {jahr}",
              sections,
              subtitle=f"Specializovany financni urad Brno | "
                       f"Festgesetzte KSt {num_cz(steuer)} CZK")


steuerbescheid("RCZ_Steuerbescheid_2020.docx", 2020, 41_200_000)
steuerbescheid("RCZ_Steuerbescheid_2021.docx", 2021, 46_800_000)
steuerbescheid("RCZ_Steuerbescheid_2022.docx", 2022, 52_300_000)


# ===========================================================================
# 7) RCZ_WP_Management_Letter_*  (3 docs) 2021, 2022, 2023
# ===========================================================================

def wp_management_letter(fname, jahr, findings_high, findings_med, kpi_text):
    sections = [
        ("Adressat",
         f"An die Geschaeftsleitung der {RCZ_NAME}\nZ. Hd. {WL}, Jednatel\n"
         f"Kopie: Group CFO Laura Bauer (REA), Group Controller Florian Maier (REA), "
         f"Chief Audit Executive Andreas Buehler (REA), KPMG AG WPG (Lead Partner "
         f"Dr. Maximilian Brand, Stuttgart)"),
        ("Aussteller",
         f"{WP_CZ}\nUnterzeichnender Partner: Ing. Tomas Holecek, statutarni "
         f"auditor (Oprav. KACR c. 2456)\nAuditteam Brno: Mgr. Hana Vesela "
         f"(Senior Manager), Bc. Jiri Slansky (Senior), Bc. Veronika Kratochvilova "
         f"(Assistant)"),
        ("Pruefungsgegenstand",
         f"Pruefung des Jahresabschlusses {jahr} (CZ-GAAP) sowie des IFRS-"
         f"Reporting-Pakets der Brennhagen CZ s.r.o. zum 31.12.{jahr}. "
         f"Pruefungsdurchfuehrung gemaess Mezinarodni auditorske standardy "
         f"(ISA – International Standards on Auditing) und tschechischem "
         f"Pruefungsgesetz (Zakon c. 93/2009 Sb., o auditorech). "
         f"Pruefungszeitraum vor Ort: 10. - 21. unora {jahr+1} (Field-Work). "
         f"Stichprobenbasis: 65 % der Erloese, 78 % der Materialaufwaende, "
         f"100 % Intercompany-Salden (IC-Matching mit REG / RHO)."),
        ("Bestaetigungsvermerk (Kurzfassung)",
         f"Auf Grundlage unserer Pruefung erteilen wir einen uneingeschraenkten "
         f"Bestaetigungsvermerk. Der Jahresabschluss vermittelt unter Beachtung "
         f"der tschechischen Bilanzierungsvorschriften (Zakon c. 563/1991 Sb., "
         f"o ucetnictvi, sowie Vyhlaska 500/2002 Sb.) ein den tatsaechlichen "
         f"Verhaeltnissen entsprechendes Bild der Vermoegens-, Finanz- und "
         f"Ertragslage. Auch das IFRS-Reporting-Paket entspricht in allen wesentlichen "
         f"Belangen den Konzernrichtlinien REA (IFRS 15/16/9/IAS 19)."),
        ("Pruefungsfeststellungen (Findings)",
         [["Nr", "Bereich", "Bewertung", "Feststellung", "Empfehlung"],
          *[[str(i+1), f["bereich"], "Hoch", f["finding"], f["empf"]]
            for i, f in enumerate(findings_high)],
          *[[str(len(findings_high)+i+1), f["bereich"], "Mittel", f["finding"], f["empf"]]
            for i, f in enumerate(findings_med)]]),
        ("Internes Kontrollsystem (IKS)",
         f"Das IKS der RCZ wird als grundsaetzlich angemessen eingeschaetzt. "
         f"Im Berichtsjahr wurden 14 Kontrollpunkte ueberprueft (SAP-Zugriffe, "
         f"4-Augen-Prinzip Wareneingang, Zahlungsfreigabe, Inventur, IC-Abstimmung). "
         f"Wesentliche Schwaeche: keine. Verbesserungspotenzial: Automatisierung "
         f"der monatlichen IC-Abstimmung mit RHO und REG (derzeit noch teilweise "
         f"Excel-basiert), avisiert mit SAP S/4HANA-Migration H2/{jahr+1}."),
        ("Wesentliche Schaetzungen / Risiken",
         kpi_text),
        ("Transfer Pricing",
         f"Die Konzernverrechnungspreise (Cost-Plus 7,5 % auf direkte und allokierte "
         f"indirekte Herstellkosten) wurden anhand der Master-File 2022 sowie "
         f"Local-File RCZ {jahr} ueberprueft. Die Cost-Plus-Marge liegt im "
         f"Interquartilsbereich der Benchmark-Studie (Amadeus-Database, "
         f"OECD-konforme Vergleichsstudie durch KPMG TP-Team Praha, Stand "
         f"{jahr}). Keine Beanstandungen."),
        ("Compliance / Steuerliche Risiken",
         f"DPPO-Vorauszahlungen wurden fristgerecht geleistet. Die DPH-"
         f"(Mehrwertsteuer-) Meldungen sowie das Kontrollni hlaseni wurden "
         f"monatlich termingerecht eingereicht. EORI-Nummer fuer Drittland-"
         f"Exporte (UK, CH, RS) liegt vor. Kein Pending Tax Audit bekannt."),
        ("Schlussvermerk",
         f"Praha / Brno, 28. unora {jahr+1}\n\n{WP_CZ}\n\n"
         f"Ing. Tomas Holecek, statutarni auditor (Partner)\n"
         f"Mgr. Hana Vesela, Senior Manager"),
    ]
    write_doc(BASE / fname, RCZ_HDR,
              f"Management Letter Jahresabschlusspruefung {jahr}",
              sections,
              subtitle=f"Prufer: KPMG Ceska republika, s.r.o. | "
                       f"uneingeschr. Bestaetigungsvermerk",
              confidential=True)


wp_management_letter("RCZ_WP_Management_Letter_2021.docx", 2021,
    findings_high=[
        {"bereich": "Vorraete",
         "finding": "Bewertung der Halbfertigerzeugnisse Linie 3 enthielt "
                    "veraltete Standardkosten (Kupferpreis 2019), Auswirkung +1,8 Mio. CZK.",
         "empf":    "Quartalsweise Anpassung der Standardkostensaetze ab 2022."}],
    findings_med=[
        {"bereich": "SAP-Berechtigungen",
         "finding": "Kritische SAP-Rolle SAP_ALL bei 2 Mitarbeitern aktiv (Notfall-Account).",
         "empf":    "Notfall-Konzept dokumentieren, Firefighter-Verfahren einfuehren."},
        {"bereich": "Inventur",
         "finding": "Inventurprotokolle Lager 4 nicht von zweiter Person abgezeichnet.",
         "empf":    "4-Augen-Prinzip strikt durchsetzen."},
    ],
    kpi_text=
    "Wesentliche Schaetzungen: (1) Garantierueckstellung (8,4 Mio. CZK) basierend "
    "auf 0,4 % vom Umsatz; (2) Rueckstellung Drohverlust Materiallieferungen "
    "Kupfer (1,2 Mio. CZK); (3) Forderungsabwertung Externer Aftermarket "
    "0,8 Mio. CZK. Risiken: Kupferpreisvolatilitaet, OEM-Single-Source-Risiko "
    "(75 % Umsatz auf 4 OEMs).")

wp_management_letter("RCZ_WP_Management_Letter_2022.docx", 2022,
    findings_high=[],
    findings_med=[
        {"bereich": "Bestaende",
         "finding": "Bestand an Cu-Litze um 18 % ueber Plan – Risiko Preisverfall.",
         "empf":    "Bestandsreduzierungsplan mit Einkauf / Materialdisposition."},
        {"bereich": "Forderungen",
         "finding": "DSO Externer Aftermarket 78 Tage (Ziel <60).",
         "empf":    "Mahnwesen straffen, Inkasso fuer Forderungen >120 Tage."},
        {"bereich": "Personal",
         "finding": "Bonus-Rueckstellung nicht abgegrenzt fuer Q4-Boni Werkleitung.",
         "empf":    "Anpassung des Abgrenzungsverfahrens HR/Finance."},
    ],
    kpi_text=
    "Garantierueckstellung 9,1 Mio. CZK (0,4 % vom Umsatz). Drohverluste 0,4 "
    "Mio. CZK (Erholung Materialmarkt). Forderungsabwertung 1,4 Mio. CZK (Anstieg). "
    "Risiken: anhaltende Energiepreisthematik in CZ, Personalkostendruck nach "
    "OZ-KOVO-Tarifabschluss +6,5 %.")

wp_management_letter("RCZ_WP_Management_Letter_2023.docx", 2023,
    findings_high=[],
    findings_med=[
        {"bereich": "IT-Sicherheit",
         "finding": "Patch-Management Backup-Server nicht durchgaengig dokumentiert.",
         "empf":    "ITIL-Aufstellung der Patch-Zyklen mit Verantwortlichkeiten."},
        {"bereich": "IC-Abstimmung",
         "finding": "IC-Differenzen zu REG kumuliert 0,9 Mio. CZK ueber Toleranz "
                    "(YE-Anpassung erfolgte).",
         "empf":    "Monatliche IC-Reconciliation, nicht erst zum YE."},
    ],
    kpi_text=
    "Garantierueckstellung 9,7 Mio. CZK. Bestandsrisiko Cu-Litze entschaerft "
    "(Preisabsicherung via REA Treasury / Markus Pflanzer). Risiken: BMS-12-"
    "Hochlauf VW ID.7 mit hohen QM-Anforderungen.")


# ===========================================================================
# 8) Misc single docs
# ===========================================================================

# RCZ_Compliance_Report_2023_v2
write_doc(BASE / "RCZ_Compliance_Report_2023_v2.docx", RCZ_HDR,
    "Compliance-Report Brennhagen CZ s.r.o. – Geschaeftsjahr 2023 (Version 2)",
    subtitle=f"Berichterstatter: Local Compliance Officer Brno | Stand: 28. Februar 2024",
    confidential=True,
    sections=[
        ("Einleitung",
         "Dieser Compliance-Report fasst die Compliance-Aktivitaeten der Brennhagen CZ "
         "s.r.o. im Geschaeftsjahr 2023 zusammen und dient als jaehrliche Berichterstattung "
         "an den Group Chief Compliance Officer der Brennhagen Elektronik AG (REA) sowie an "
         "den Konzern-Pruefungsausschuss (Vorsitz Prof. Dr.-Ing. Gerhard Voss). "
         "Version 2 enthaelt redaktionelle Korrekturen und Ergaenzungen zur Lohnsteuer-"
         "Compliance, nachdem die Erstversion vom 31.1.2024 unvollstaendig war."),
        ("Compliance-Organisation",
         f"Local Compliance Officer: Marcela Benesova (40 %), reports dotted line an "
         f"Chief Compliance Officer Brennhagen AG (REA Stuttgart). Compliance-Komitee "
         f"Brno tagt quartalsweise (Teilnehmer: {WL}, {QL}, {FC}, {HR}, Benesova). "
         f"Compliance-Hotline ueber Convercent / OneTrust (anonyme Whistleblower-"
         f"Plattform, betrieben durch Group). 2023 wurden 4 Hinweise eingereicht, "
         f"davon 2 unbegruendet, 1 disziplinarisch geahndet (Verstoss gegen "
         f"Reisekosten-Richtlinie), 1 noch in Bearbeitung (Vorwurf Lieferantenpraeferenz, "
         f"Untersuchung durch externe Kanzlei {ANW_CZ})."),
        ("Wesentliche Compliance-Felder",
         ("list", [
             "Korruptionspraevention (Zakon c. 418/2011 Sb., trestni odpovednost "
             "pravnickych osob – Strafverantwortlichkeit juristischer Personen): "
             "100 % Schulungsquote Indirect Personnel (Kontrolle via Cornerstone).",
             "Kartellrecht (Zakon c. 143/2001 Sb., o ochrane hospodarske souteze): "
             "keine Vorfaelle; jaehrliche Schulung Sales und Einkauf erfolgt.",
             "Exportkontrolle / Dual-Use: keine ECCN-Klassifizierungen ueber 3A001 "
             "betroffen, keine BAFA-Lizenzen erforderlich.",
             "Datenschutz (GDPR / OOU – Urad pro ochranu osobnich udaju): "
             "Verzeichnis der Verarbeitungstaetigkeiten aktualisiert; 1 Datenpanne "
             "(falsch adressierte E-Mail) gemeldet, keine Sanktion.",
             "Geldwaeschepraevention (AML, Zakon c. 253/2008 Sb.): nicht primaer "
             "betroffen (kein Finanzinstitut), Standard-KYC bei Neukunden via SAP.",
             "Arbeitsschutz / BOZP (Zakon c. 309/2006 Sb.): 8 Arbeitsunfaelle (Vj. 12), "
             "LTIFR 4,2 (Branche 6,8); keine schweren Unfaelle.",
         ])),
        ("Lohnsteuer-Compliance (Ergaenzung Version 2)",
         f"Die monatlichen Lohnsteuer-Meldungen (Mesicni vyuctovani zalohy na dan ze "
         f"zavisle cinnosti) wurden 2023 fristgerecht durch externen Lohnabrechnungs-"
         f"Dienstleister ADP CZ s.r.o. eingereicht. Jaehrliche Aufstellung der "
         f"Einkommen (rocni vyuctovani) und Bestaetigung fuer Mitarbeiter erfolgte "
         f"vor 15.3.2024. Solidaritaetszuschlag fuer Spitzeneinkommen >1.935.552 CZK "
         f"(Stand 2023) korrekt einbehalten. Keine Beanstandungen durch Specializovany "
         f"financni urad Brno-mesto. Sozialversicherung: Anmeldungen bei OSSZ Brno-"
         f"mesto vollstaendig, monatliche Beitragsmeldungen fristgerecht."),
        ("Krankenversicherung (VZP / Zamestnanecka)",
         f"Von 680 Mitarbeitern sind 482 bei VZP (Vseobecna zdravotni pojistovna CR, "
         f"kod 111), 142 bei Zamestnanecka pojistovna Skoda (kod 209), 56 bei Vojenska "
         f"zdravotni pojistovna (kod 201). Arbeitgeber-Anteile wurden monatlich "
         f"vollstaendig abgefuehrt (insgesamt 2023 ca. 48,2 Mio. CZK)."),
        ("Audit-Plan 2024",
         "Geplante interne Audits 2024: (a) IT-Berechtigungen und Patch-Management "
         "(Q2), (b) Lieferantenfreigabe-Prozess (Q3), (c) Reisekosten / Bewirtungen "
         "(Q4). Externe Audits: KPMG-JA-Pruefung (Feb 2024), IATF-16949 Rezertifizierung "
         "(Sep 2024) durch TUEV Sued / TUEV Nord CZ."),
        ("Schlussvermerk",
         f"Brno, 28. unora 2024\n\nMarcela Benesova, Local Compliance Officer\n"
         f"Freigabe: {WL}, Jednatel"),
    ])

# RCZ_Mietvertrag_Betriebsgelaende_2020
write_doc(BASE / "RCZ_Mietvertrag_Betriebsgelaende_2020.docx", RCZ_HDR,
    "Najemni smlouva / Mietvertrag Betriebsgelaende Brno-Slatina (CTP Park Brno)",
    subtitle="Vermieter: CTP Property XLVII, s.r.o. | Mieter: Brennhagen CZ s.r.o. | "
             "Laufzeit 1.7.2020 - 30.6.2030",
    sections=[
        ("Smluvni strany",
         "Pronajimatel / Vermieter:\nCTP Property XLVII, s.r.o.\n"
         "Central Trade Park D1 1571, 396 01 Humpolec, Ceska republika\n"
         "ICO 27845678, DIC CZ27845678\nzapsana KS v Ceskych Budejovicich, C 18234\n"
         "Zastoupena: Mgr. Petr Karlik (Asset Manager CEE)\n\n"
         "Najemce / Mieter:\nBrennhagen CZ s.r.o.\nTuzemska 47, 627 00 Brno-Slatina\n"
         "ICO 28765432, DIC CZ28765432\nZastoupena: Petr Novak (jednatel)"),
        ("Predmet najmu / Mietgegenstand",
         "Gewerbeimmobilie im CTP Park Brno (Brno-Slatina), bestehend aus:\n"
         "- Produktionshalle 9.800 m2 (Halle B-3)\n"
         "- Lagerflaeche 2.400 m2 (Lager B-3-L)\n"
         "- Buerogebaeude 1.600 m2 (3 Stockwerke)\n"
         "- Aussen-/Parkflaechen 4.500 m2 (Pkw-Stellplaetze 120, LKW-Andockstellen 8)\n"
         "Gesamt-Mietflaeche: 18.300 m2 vermietbare Flaeche. "
         "Adresse: Tuzemska 47, 627 00 Brno-Slatina. "
         "Kataster: Pozemek p. c. 1245/8, KU Slatina (okres Brno-mesto), "
         "List vlastnictvi 4567."),
        ("Mietzins / Najemne",
         "Grundmiete (zakladni najemne):\n"
         "- Produktion: 4,80 EUR/m2/Monat = 47.040 EUR/Monat\n"
         "- Lager: 3,20 EUR/m2/Monat = 7.680 EUR/Monat\n"
         "- Buero: 12,50 EUR/m2/Monat = 20.000 EUR/Monat\n"
         "- Aussenflaechen: 0,80 EUR/m2/Monat = 3.600 EUR/Monat\n"
         "Gesamt Grundmiete: 78.320 EUR/Monat (ca. 1.918.840 CZK/Monat)\n\n"
         "Nebenkosten (sluzby): pauschal 1,20 EUR/m2/Monat = 21.960 EUR/Monat. "
         "Endabrechnung jaehrlich bis 30.6. des Folgejahres. "
         "Indexierung: jaehrlich zum 1.1. mit HICP (CZ) +/- 100 %, Mindestindexierung 1,5 %, "
         "Hoechstindexierung 4,0 %. Erste Indexierung 1.1.2022."),
        ("Laufzeit / Kuendigung",
         "Mietbeginn: 1. cervence 2020 (1. Juli 2020).\n"
         "Festlaufzeit: 10 Jahre bis 30. cervna 2030.\n"
         "Verlaengerungsoption: 2 x 5 Jahre, Ausuebung bis 12 Monate vor Ablauf.\n"
         "Ausserordentliches Kuendigungsrecht: nur bei wesentlicher Vertragsverletzung "
         "des Vermieters (Heilungsfrist 60 Tage). Kein ordentliches Sonderkuendigungsrecht."),
        ("Mietsicherheit / Jistota",
         "Bankgarantie ueber 3 Monatsmieten = 300.840 EUR durch Komercni banka a.s. "
         "(KOMBCZPP) zugunsten CTP Property XLVII, s.r.o. Garantie wird jaehrlich "
         "an aktuelle Miethoehe angepasst. Rueckgabe bei vertragsgemaesser Beendigung."),
        ("Sonstige Bestimmungen",
         "Untervermietung nur mit schriftlicher Zustimmung des Vermieters. "
         "Bauliche Veraenderungen >50.000 EUR beduerfen Zustimmung. "
         "Schoenheitsreparaturen alle 5 Jahre durch Mieter. "
         "Anwendbares Recht: tschechisches Recht (Zakon c. 89/2012 Sb., obcansky zakonik). "
         "Gerichtsstand: Krajsky soud v Brne. "
         "Schlichtungsklausel: Schiedsgericht der Wirtschaftskammer der Tschechischen "
         "Republik (Rozhodci soud HK CR), 1 Schiedsrichter, Verfahrenssprache Tschechisch."),
        ("Unterschriften",
         signatures("Mgr. Petr Karlik", "Asset Manager", "CTP Property XLVII, s.r.o.",
                    WL, "Jednatel", RCZ_NAME,
                    place="Brno", date_str_="15. cervna 2020")),
    ])

# RCZ_Versicherungsnachweis_2023
write_doc(BASE / "RCZ_Versicherungsnachweis_2023.docx", RCZ_HDR,
    "Versicherungsnachweis 2023 – Brennhagen CZ s.r.o. (Brno)",
    subtitle="Stand: 1. Januar 2023 | Konzernpolicen REA + lokale CZ-Policen",
    sections=[
        ("Konzernpolicen via REA",
         "Die Brennhagen CZ s.r.o. ist als Tochter im Konzernversicherungsprogramm der "
         "Brennhagen Elektronik AG (REA Stuttgart) eingeschlossen. Versicherungsmakler "
         "ist Marsh GmbH (Frankfurt), Lead-Underwriter siehe Einzelpolicen. "
         "Folgende Konzernpolicen erfassen die RCZ:"),
        ("Policenuebersicht (Konzern + lokal)",
         [["Sparte", "Versicherer", "Police-Nr.", "Versicherungssumme",
           "Selbstbehalt", "Praemienanteil RCZ p.a."],
          ["Industrie-/Sach-/BU", "AXA XL (Frankfurt)", "AXA-IND-RG-2023-001",
           "850 Mio. EUR (PML, Konzern)", "250.000 EUR/Schaden", "ca. 380.000 EUR"],
          ["Betriebshaftpflicht (D&O excl.)", "HDI Global SE", "HDI-BHV-RG-2023-08",
           "75 Mio. EUR (Personenschaden)", "100.000 EUR", "ca. 95.000 EUR"],
          ["Produkthaftpflicht (Auto-OEM)", "AIG Europe SA", "AIG-PROD-RG-2023-12",
           "150 Mio. EUR (clean Recall)", "250.000 EUR", "ca. 220.000 EUR"],
          ["D&O / Managerhaftpflicht", "Allianz GCS SE", "AGCS-DO-RG-2023-03",
           "50 Mio. EUR", "150.000 EUR", "Konzern (RCZ-Allokation 28.000 EUR)"],
          ["Transport (Cargo)", "Zurich Insurance plc", "ZUR-CRG-RG-2023-06",
           "8 Mio. EUR/Sendung", "5.000 EUR", "ca. 42.000 EUR"],
          ["Cyber", "Munich Re Specialty", "MR-CYB-RG-2023-04",
           "25 Mio. EUR", "500.000 EUR", "Konzern (RCZ-Allokation 18.000 EUR)"],
          ["KFZ-Flotte CZ (lokal)", "Kooperativa pojistovna a.s.", "KOO-FL-2023-CZ-RCZ-12",
           "haftpflicht gesetzl.+ Voll", "10.000 CZK", "ca. 1,2 Mio. CZK"],
          ["Gebaeude (CTP Park Brno, lokal)", "Generali Ceska pojistovna a.s.", "GEN-GEB-2023-RCZ-44",
           "180 Mio. CZK (Sach+Inventar)", "50.000 CZK", "ca. 0,9 Mio. CZK"],
          ["Krankenversicherung Mitarbeiter", "VZP/Zamestnanecka/Vojenska",
           "gesetzlich/povinne", "gesetzl. Leistungsspektrum",
           "n/a", "Arbeitgeber-Anteil ca. 48 Mio. CZK"]]),
        ("Schadenfaelle 2022/2023",
         "2022: 1 Sachschaden Buerogebaeude (Wasserrohrbruch, 28.000 EUR, gedeckt via "
         "AXA XL nach Selbstbehalt). 2 KFZ-Schaeden (Parkschaeden, lokal Kooperativa, "
         "voll gedeckt). Keine Produkthaftpflicht-Faelle. Keine Cyber-Vorfaelle. "
         "2023 H1: 0 wesentliche Schaeden."),
        ("Compliance / Lokale Pflichten",
         "Die in der Tschechischen Republik vorgeschriebenen lokalen Versicherungen "
         "(Kfz-Haftpflicht POV, Gebaeude bei CTP-Mietvertrag-Pflicht, Unfallversicherung "
         "Arbeitgeber via Kooperativa) sind abgeschlossen. Konzernpolicen werden "
         "ueber Marsh-Wording in lokale Folgepolicen (Kooperativa, Generali Ceska) "
         "via Fronting-Vereinbarung eingebracht."),
        ("Ansprechpartner / Schadenmeldung",
         f"Schadenmeldung intern: {FC} (Finanzcontroller) als Schadenkoordinator. "
         f"Eskalation: Marsh Service Center DACH+CEE, marsh.frankfurt@marsh.com. "
         f"Konzern-Risikomanagement: Andreas Buehler, CAE (REA Stuttgart). "
         f"Lokale Notruf-Nr. Kooperativa 24/7: 957 105 105."),
    ])

# RCZ_to_REG_IC_2023_09 was excluded above – do nothing.


# ===========================================================================
# 9) PRJ-2024-003 SAP S/4HANA Rollout Status docs (2 files)
# ===========================================================================

def sap_status(fname, monat, jahr, milestones, risiken, next_steps):
    sections = [
        ("Projektrahmen",
         f"Projekt: PRJ-2024-003 SAP S/4HANA Rollout Polen / CZ (Wave 2)\n"
         f"Berichtsmonat: {MONTHS_DE[monat]} {jahr}\n"
         f"Projektleitung: Mag. Marek Wojciechowski (RPL) + {WL} (RCZ) Co-Lead CZ\n"
         f"Sponsor: Laura Bauer, CFO REA + Florian Maier, Group Controller\n"
         f"Integrationspartner: Capgemini Polska Sp. z o.o. (RPL Lead), Capgemini "
         f"CZ a.s. (RCZ-Stream)\n"
         f"Budget Gesamt: 8,4 Mio. EUR (RPL 4,8 + RCZ 3,6); "
         f"Ist {MONTHS_DE[monat]} {jahr}: ca. 62 % verbraucht"),
        ("Statusampel",
         [["Workstream", "Status", "Termin", "Kommentar"],
          ["Master Data Migration", "GELB", f"Q4/{jahr}", "Material-Stammdaten "
           "RCZ + RPL Harmonisierung in Verzug"],
          ["Finance (FI/CO)", "GRUEN", f"Q3/{jahr}", "Customizing FI/CO abgeschlossen, "
           "UAT laeuft"],
          ["Logistics (MM/SD)", "GRUEN", f"Q4/{jahr}", "Integration EDI VW/BMW getestet"],
          ["Production (PP/QM)", "GELB", f"Q1/{jahr+1}", "QM-Integration LIMS "
           "noch offen, Workshop {monat+1}/{jahr}"],
          ["HR (SAP SuccessFactors)", "ROT", f"Q2/{jahr+1}",
           "Schnittstelle ADP CZ Lohn aufwendiger als geschaetzt"],
          ["Cutover / Go-Live", "GRUEN", "1.1.2025", "Cutover-Plan v2 erstellt"]]),
        ("Erreichte Meilensteine im Berichtsmonat",
         ("list", milestones)),
        ("Top-Risiken",
         ("list", risiken)),
        ("Naechste Schritte",
         ("list", next_steps)),
        ("Kosten / Forecast",
         f"Budget Total 8,4 Mio. EUR | YTD Ist (per {MONTHS_DE[monat]}): "
         f"5,2 Mio. EUR | Forecast YE: 8,1 Mio. EUR | "
         f"Abweichung: -0,3 Mio. EUR Underrun (Verzoegerung HR-Workstream wirkt "
         f"liquiditaetsentlastend). External Spend Capgemini: 3,8 Mio. EUR YTD. "
         f"Hardware (HPE Server-Hosting AWS Frankfurt): 0,6 Mio. EUR YTD."),
        ("Lessons Learned (laufend)",
         "Wave 1 (Werk Heilbronn 2023) hat gezeigt, dass Master Data Cleansing "
         "vorgezogen werden muss; dies wurde fuer Wave 2 RPL/RCZ beruecksichtigt. "
         "Wave-2-spezifisch lessons: (1) Mehrsprachigkeit Trainingsunterlagen "
         "(PL, CZ, EN) frueher anstossen, (2) lokale Steuerprozesse (PL JPK_VAT, "
         "CZ Kontrolni hlaseni) benoetigen lokale Tax-Consultants ab Phase Design, "
         "(3) ADP-Schnittstellen sind landesspezifisch nicht standardisiert."),
        ("Schlussvermerk",
         f"Brno, {25}. {MONTHS_DE[monat]} {jahr}\n\n"
         f"{WL}, Co-Projektleiter RCZ\nMarek Wojciechowski, Projektleiter RPL"),
    ]
    write_doc(BASE / fname, REG_HDR,
              f"Statusbericht PRJ-2024-003 SAP S/4HANA Rollout – {MONTHS_DE[monat]} {jahr}",
              sections,
              subtitle="Wave 2: Werk Katowice (RPL) + Werk Brno (RCZ) | Go-Live 1.1.2025")


sap_status("PRJ-2024-003_Statusbericht_2024_06_SAP_S_4HANA_Rollout_Polan.docx",
    6, 2024,
    milestones=[
        "FI/CO-Customizing abgeschlossen, UAT-Start am 18.6.2024 mit Key-Usern Brno + Katowice",
        "EDI-Integration BMW (Tier-1-Anbindung VAB-Format) getestet ohne Befund",
        "Cutover-Plan v2 freigegeben durch Steering-Committee (12.6.2024)",
        "Trainingsmaterial PL/CZ/EN final – 240 Endnutzer adressiert",
        "Datenmigration Probelauf 1 fuer FI-Stammdaten (Konten, KSt) erfolgreich",
    ],
    risiken=[
        "HR-Workstream: ADP-CZ-Schnittstelle aufwendiger als geschaetzt, Verzug 8 Wochen",
        "Master Data RPL: 22 % der Material-Stammdaten noch nicht harmonisiert",
        "QM-Integration LIMS: Workshop in Juli noetig, sonst Risiko Go-Live PP/QM",
        "Capgemini-Ressourcen-Allokation H2/2024: Vertragsverlaengerung pending",
        "Schulungstermine September konkurrieren mit Werksferien CZ + IATF-Audit",
    ],
    next_steps=[
        "Eskalation HR-Workstream an Sponsor Laura Bauer (CFO) in KW27",
        "Master Data Cleansing Sprint RPL bis 31.7.2024",
        "QM-LIMS Workshop in Brno 9.-10.7.2024 mit RPL + RCZ + Capgemini",
        "Cutover-Probe (Dry Run 1) am 14.-16.9.2024",
        "Steering Committee Update am 5.7.2024 (Frankfurt)",
    ])

sap_status("PRJ-2024-003_Status_2024_09_SAP_S_4HANA_Rol.docx",
    9, 2024,
    milestones=[
        "Master Data Cleansing RPL abgeschlossen (Stand 9.9.2024)",
        "UAT FI/CO + MM/SD final abgeschlossen (Sign-off 18.9.2024)",
        "QM-LIMS Schnittstelle live im Test-Mandanten",
        "Dry Run 1 Cutover am 14.-16.9.2024 mit minor Findings (12 Defects, davon 0 Kritisch)",
        "Trainingsphase 1 (Endnutzer 180/240) absolviert",
    ],
    risiken=[
        "HR-Workstream nun GELB (verbessert), aber Cutover faellt mit ADP-Lohnabrechnung "
        "Januar zusammen – Fallback-Plan Parallel-Lauf vorbereitet",
        "Cutover-Wochenende 28.12.2024-1.1.2025 betrifft tschechische und polnische "
        "Feiertage – Risiko Verfuegbarkeit Capgemini-Berater",
        "VAB/EDI-Anbindung Stellantis verzoegert (CSR-Spezifikation 2024-Q3 vom OEM)",
    ],
    next_steps=[
        "Dry Run 2 Cutover am 26.-28.10.2024",
        "Defect-Beseitigung Dry Run 1 bis 15.10.2024",
        "Trainingsphase 2 (60 verbleibende Endnutzer) im November",
        "Go-/No-Go Entscheidung im Steering Committee 12.12.2024",
        "Hypercare-Plan ab 1.1.2025 fuer 8 Wochen vorbereitet",
    ])


# ===========================================================================
# 10) Misc cross-folder docs that landed in 06 (RHU_*, REG_BV, REA_*, Patent_*)
# ===========================================================================

# Patent_02_Jahresgebuehr_2022.docx
write_doc(BASE / "Patent_02_Jahresgebuehr_2022.docx", REG_HDR,
    "Patentamtliches Jahresgebuehrenschreiben – Patent EP3245187 (Steckverbinder-Crimp)",
    subtitle="Europaeisches Patentamt – Jahresgebuehr 2022 (6. Jahr)",
    sections=[
        ("Adressat",
         f"{R['name']}\n{R['addr']}\n{R['hrb']}\n\n"
         f"Zustellung an die Patentabteilung, z. Hd. Dr.-Ing. Annette Krueger, "
         f"Leiterin IP-Management Brennhagen Elektronik AG"),
        ("Aussteller",
         "Europaeisches Patentamt (EPA)\nBob-van-Benthem-Platz 1, 80469 Muenchen\n"
         "Bearbeiter: Herr M. Kuipers, Direktion 1.6.1 Gebuehren\n"
         "Aktenzeichen: EP 17/3245187.8 – Brennhagen Elektronik AG"),
        ("Sachverhalt",
         "Mit diesem Schreiben wird die Faelligkeit der Jahresgebuehr fuer das "
         "6. Patentjahr des europaeischen Patents EP3245187 (»Verfahren und "
         "Vorrichtung zum Crimpen von Steckverbinder-Kontakten mit reduzierter "
         "Crimphoehen-Streuung«) mitgeteilt. Patentinhaberin: Brennhagen Elektronik AG. "
         "Erfinder: Dipl.-Ing. Petr Novak (Werk Brno – RCZ), Dr.-Ing. Andreas Maier "
         "(Werk Heilbronn – REG), Dipl.-Ing. Marek Wojciechowski (RPL Katowice). "
         "Anmeldedatum: 14.5.2017, Erteilungsdatum: 22.11.2019."),
        ("Gebuehrenfestsetzung",
         [["Position", "Betrag EUR", "Faelligkeit"],
          ["Jahresgebuehr 6. Jahr (Art. 86 EPUe i.V.m. Regel 51 AOEPA)", "910,00",
           "31. Mai 2022 (letzter Tag des Anmeldemonats)"],
          ["Zuschlag bei verspaeteter Zahlung (Regel 51 (2))", "182,00",
           "innerhalb 6 Monaten nach Faelligkeit"],
          ["Validierungsgebuehren benannte Staaten (DE, FR, GB, IT, ES, CZ, PL, HU)",
           "0 (national bereits geleistet)", "n/a"]]),
        ("Validierung in Tschechien",
         "Das Patent ist national in der Tschechischen Republik validiert "
         "(Patent CZ EP3245187, Urad prumysloveho vlastnictvi – UPV, Praha). "
         "Lokaler Patentanwalt: Mgr. Marek Polacek, PatentServis Praha, a.s. "
         "Tschechische Jahresgebuehr 6. Jahr: 3.500 CZK, separat zahlbar an UPV "
         "bis 30. listopadu 2022."),
        ("Zahlung",
         "Konto EPA: Deutsche Bundesbank Filiale Muenchen\n"
         "IBAN: DE20 7000 0000 0070 0540 60 | BIC: MARKDEFF700\n"
         "Verwendungszweck: EP3245187 / 6. Jahr / Brennhagen AG"),
        ("Rechtsfolgen bei Nichtzahlung",
         "Bei Nichtzahlung der Jahresgebuehr binnen 6 Monaten nach Faelligkeit "
         "(auch mit Zuschlag) erlischt das Patent rueckwirkend zum Faelligkeitstag "
         "(Art. 86 (2) EPUe). Wiedereinsetzung in den vorigen Stand nach Art. 122 "
         "EPUe ist nur unter strengen Voraussetzungen moeglich."),
        ("Interne Bearbeitung Brennhagen",
         f"Eingegangen am 4.4.2022, weitergeleitet an Treasury (Markus Pflanzer) "
         f"zur Bezahlung. Zahlung erfolgte am 25.4.2022. "
         f"Tschechische Folgegebuehr koordiniert mit RCZ Brno ({FC}). "
         f"Patent-Portfolio-Status: aktiv, kommerzielle Verwertung in Steckverbinder-"
         f"Linien aller Konzernwerke (RCZ Brno Hauptstandort, REG Heilbronn, RPL Katowice)."),
    ])


# REA_Immobilienbewertung_Office_Shanghai_2023
write_doc(BASE / "REA_Immobilienbewertung_Office_Shanghai_2023.docx", REG_HDR,
    "Immobilienbewertung Office Shanghai (RCN) – Bewertungsstichtag 31.12.2023",
    subtitle="Gutachter: Cushman & Wakefield Shanghai | Auftraggeber: REA / Treasury",
    sections=[
        ("Auftraggeber / Auftragnehmer",
         "Auftraggeber: Brennhagen Elektronik AG, Vaihinger Strasse 120, 70567 Stuttgart\n"
         "Auftragnehmer: Cushman & Wakefield Shanghai Ltd., Plaza 66, Tower 2, "
         "1366 Nanjing Road West, Jing'an District, Shanghai 200040, PRC\n"
         "Gutachter: Mr. James Liu MRICS, Senior Director Valuation Greater China\n"
         "Auftrag erteilt: 12.10.2023 durch Markus Pflanzer (Group Treasurer REA)"),
        ("Bewertungsobjekt",
         "Bueroflaeche der Brennhagen (Shanghai) Co. Ltd. (RCN) im 18. Stock des "
         "Plaza 66 Tower 1, 1266 Nanjing Road West, Jing'an, Shanghai. "
         "Mietflaeche: 1.840 m2 NLA (Net Leasable Area). "
         "Nutzung: Regional HQ + Vertrieb + Aftermarket fuer Asien-Geschaeft. "
         "Hinweis: Bueroflaeche wird gemietet (Mieter RCN), nicht im Eigentum REA. "
         "Daher Bewertung im Sinne »Mietnutzungswert / Recognised Value of Lease"
         " Interest« sowie alternative Mietkostenvergleichswerte fuer ggf. "
         "anstehende Verlaengerungsverhandlung."),
        ("Methodik",
         "Bewertung nach RICS Valuation Standards (Red Book Global 2022). "
         "Methoden: (a) Vergleichswertverfahren (Comparable Method) mit "
         "Grade-A-Bueromieten Jing'an/Lujiazui CBD, (b) Diskontierter Cashflow "
         "(DCF) Mietzahlungen, (c) Anschlussmieten-Szenario 2026 bei Vertragsverlaengerung. "
         "Diskontierungssatz WACC RMB: 8,5 % p. a. Exit-Cap-Rate 6,0 %."),
        ("Marktdaten Shanghai Grade-A-Buero Jing'an Q4/2023",
         [["Kennzahl", "Q4/2023", "Vorjahr", "Trend"],
          ["Durchschnitts-Miete Grade-A Jing'an", "10,8 RMB/m2/Tag",
           "11,5 RMB/m2/Tag", "-6 % YoY"],
          ["Leerstand Grade-A Jing'an", "16,8 %", "13,2 %", "+3,6 pp"],
          ["Aktuell-Miete RCN Plaza 66", "9,2 RMB/m2/Tag",
           "8,8 RMB/m2/Tag", "+5 %"],
          ["Wechselkurs RMB/EUR YE2023", "7,82", "7,42", "+5,4 %"]]),
        ("Bewertungsergebnis",
         "Bewertungsstichtag: 31. Dezember 2023.\n"
         "Marktwert des Mietrechts (Lease-Hold Interest) zum Stichtag:\n"
         "- Net Present Value verbleibender Mietvertrag (Restlaufzeit 24 Monate): 14,2 Mio. RMB (ca. 1,82 Mio. EUR)\n"
         "- Wiederbeschaffungs-Mietkosten 2-Jahres-Vertrag bei Marktmiete: 16,8 Mio. RMB (ca. 2,15 Mio. EUR)\n"
         "Empfehlung: bei Anschlussmietvertrag Q1/2026 Neuverhandlung mit "
         "Zielkorridor 8,5 - 9,0 RMB/m2/Tag (Marktanpassung bei steigendem Leerstand)."),
        ("Empfehlungen an Treasury / Real Estate Committee",
         "(1) Vertragsverlaengerung 2026 ist wirtschaftlich; alternative Standortpruefung "
         "Lujiazui CBD (etwas hoehere Mieten, naeher zum Bankenviertel) ist nicht "
         "zwingend.\n"
         "(2) Hedging der RMB-Mietzahlungen ueber Konzern-Treasury (Markus Pflanzer) "
         "fuer 12-Monats-Horizont empfohlen.\n"
         "(3) Bilanzielle Behandlung: IFRS 16 ROU-Asset und Lease-Liability bereits "
         "korrekt erfasst (siehe Konzernabschluss REA 2023 Note 14)."),
        ("Schlussvermerk",
         "Shanghai/Stuttgart, 15. Januar 2024. Gutachten erstellt von James Liu MRICS, "
         "Cushman & Wakefield Shanghai Ltd. Gutachten ist vertraulich; nicht zur "
         "Weitergabe an Dritte. (Hinweis: dieses Dokument fand sich in der Brno-Akte; "
         "vermutlich Ablagefehler.)"),
    ],
    confidential=True)


# REA_MBZ_ECU-900_QBR_2022_Q1
write_doc(BASE / "REA_MBZ_ECU-900_QBR_2022_Q1.docx", REG_HDR,
    "Mercedes-Benz QBR ECU-900 – Q1/2022 (Quarterly Business Review)",
    subtitle="Kunde: Mercedes-Benz Group AG | Produkt: Powertrain-ECU Gen3 (ECU-900) | "
             "Werk: REG Heilbronn (Modulen) + RCZ Brno (Steckverbinder)",
    sections=[
        ("Teilnehmer",
         "Mercedes-Benz Group AG (MBZ):\n"
         "- Dr. Sebastian Vogt (Director Powertrain Procurement)\n"
         "- Daniel Wagner (Lead Buyer ECU)\n"
         "- Dipl.-Ing. Sarah Brettmann (Quality Engineering Powertrain)\n\n"
         "Brennhagen Elektronik AG / Toechter:\n"
         "- Stefan Richter (CMO/BD, REA)\n"
         "- Andreas Maier (Werkleiter REG, Hauptverantwortung ECU-900-Modul)\n"
         "- Petr Novak (Werkleiter RCZ, Steckverbinder-Zulieferung)\n"
         "- Eva Cerna (Q-Leitung RCZ)\n"
         "- Sabine Brand (Q-Leitung REG)"),
        ("Performance Q1/2022",
         [["KPI", "Q1/2022 Ist", "Target", "Status"],
          ["Lieferperformance OTIF", "98,4 %", ">98 %", "GRUEN"],
          ["Customer Complaints PPM", "12 PPM", "<25 PPM", "GRUEN"],
          ["Kostenentwicklung (Preisreduktion p.a.)", "-2,8 %", "-3,0 % (LTA)", "GELB"],
          ["Reklamationen 8D (Stk.)", "1", "<2", "GRUEN"],
          ["Qualitaetskosten (FCR + EFR)", "0,18 % Umsatz", "<0,25 %", "GRUEN"]]),
        ("Highlights / Themen",
         ("list", [
             "PPAP-Approval Linie 2 RCZ Brno fuer neue MQS-12 Variante (200 V) erfolgt am 14.3.2022",
             "Capacity-Aufbau fuer C-Klasse Hochlauf 2022/2023 (zusaetzliche 25 % "
             "Stueckzahl ab Q3/2022), CapEx 4,2 Mio. EUR in REG + 1,8 Mio. EUR in RCZ freigegeben",
             "Kupferpreis-Pass-Through-Klausel ueber Q-Index aktiviert (rueckwirkend Q4/2021)",
             "8D-Report 8D-2022-014 (Crimp-Hoehe out-of-spec, RCZ Brno) abgeschlossen, "
             "Root-Cause Werkzeugverschleiss, Werkzeug ersetzt, kein Field-Issue",
             "Vorbereitung Carbon-Footprint-Disclosure CDP-Supplier 2022 (Scope 1+2 Werks-Brno)",
         ])),
        ("Roadmap / Naechste Schritte",
         "Q2/2022: PPAP-Genehmigung 800 V-Variante fuer Hochleistungsmodelle. "
         "Q3/2022: Vertragsverlaengerung Rahmenliefervertrag ECU-900 (Laufzeit "
         "2023-2027). Q4/2022: Pruefung Erweiterung Lieferumfang auf Stellantis-"
         "Plattform STLA Large (Synergien fuer RCZ).\n\n"
         "Naechstes QBR: 24. Juni 2022 (Sindelfingen).\n\n"
         "Hinweis: dieses Dokument liegt versehentlich im RCZ-Brno-Ordner, "
         "obwohl primaerer Lead REG Heilbronn ist; betrifft RCZ als Sub-Tier-1-Lieferant."),
        ("Sign-off",
         signatures("Stefan Richter", "CMO/BD", R["name"],
                    "Dr. Sebastian Vogt", "Director Powertrain Procurement",
                    "Mercedes-Benz Group AG",
                    place="Heilbronn", date_str_="31. Maerz 2022")),
    ])


# REG_BV_Betriebliches_Vorschlagswesen_2022
write_doc(BASE / "REG_BV_Betriebliches_Vorschlagswesen_2022.docx", REG_HDR,
    "Betriebsvereinbarung Betriebliches Vorschlagswesen (BVW) 2022 – Werk Heilbronn",
    subtitle="Zwischen REG-Geschaeftsleitung und Betriebsrat REG Heilbronn (Klaus Bauer)",
    sections=[
        ("Praeambel",
         "Brennhagen Elektronik GmbH (REG), Werk Heilbronn, und der Betriebsrat des "
         "Werks Heilbronn (BR-Vorsitzender Klaus Bauer, stv. Konzernbetriebsrats-"
         "Vorsitz) vereinbaren das nachstehende Betriebliche Vorschlagswesen (BVW) "
         "als Teil des kontinuierlichen Verbesserungsprozesses (KVP / Kaizen). "
         "Geltungsbereich: alle 820 Mitarbeiter Werk Heilbronn (inkl. Auszubildende). "
         "Hinweis: identische BV besteht auch fuer Werk Brno (RCZ) mit OZ KOVO."),
        ("§ 1 Zweck",
         "Das BVW soll alle Mitarbeiter ermutigen, durch konkrete Verbesserungsvorschlaege "
         "zur Qualitaet, Effizienz, Arbeitssicherheit, Energieeinsparung und "
         "Nachhaltigkeit der REG beizutragen. Anerkannte Vorschlaege werden materiell "
         "und ideell honoriert. Das BVW ist Bestandteil des Total Productive Management "
         "(TPM) und IATF-16949-konformes KVP-Werkzeug."),
        ("§ 2 Einreichung",
         "Vorschlaege werden ueber die Cornerstone-BVW-Plattform online oder schriftlich "
         "(BVW-Briefkaesten an 4 Standorten im Werk) eingereicht. Jeder Vorschlag erhaelt "
         "eine eindeutige BVW-Nr. (Format BVW-REG-JJJJ-NNNN). Anonyme Vorschlaege werden "
         "nicht praemiert, koennen aber als Anregung weitergegeben werden."),
        ("§ 3 Bewertung",
         "Bewertungsgremium: 6 Personen (3 GL, 3 BR), Vorsitz wechselnd. Tagung "
         "monatlich. Bewertung erfolgt nach: (a) Einsparung in EUR/Jahr, (b) "
         "Sicherheits- oder Qualitaetswirkung, (c) Umsetzbarkeit, (d) Innovationsgrad. "
         "Beschlussfaehigkeit bei 4 Mitgliedern. Beschluss mit einfacher Mehrheit; "
         "bei Stimmengleichheit entscheidet GL-Vertreter."),
        ("§ 4 Praemierung",
         "(1) Nicht berechenbare Vorschlaege (z. B. Sicherheit, Ergonomie): Praemie "
         "50 - 500 EUR.\n"
         "(2) Berechenbare Vorschlaege: 25 % der bewerteten Netto-Jahresersparnis "
         "im 1. Jahr, gedeckelt auf 50.000 EUR (steuerbeguenstigt nach § 3 Nr. 16 "
         "EStG soweit anwendbar).\n"
         "(3) Sammelvorschlaege (Team) werden anteilig nach Beitrag praemiert.\n"
         "(4) Sonderpraemie »Vorschlag des Jahres«: 5.000 EUR + 2 Tage Sonderurlaub."),
        ("§ 5 Umsetzung",
         "Umsetzungsverantwortung liegt bei der jeweiligen Abteilungsleitung. "
         "Bei abteilungsuebergreifenden Vorschlaegen koordiniert der KVP-Beauftragte "
         "(Frank Hoffmann, REG). Status-Reporting an BVW-Gremium quartalsweise."),
        ("§ 6 Datenschutz / Vertraulichkeit",
         "Personenbezogene Daten der Einreicher werden gemaess DSGVO/BDSG verarbeitet. "
         "Auswertungen erfolgen anonymisiert. Aufbewahrungsfrist BVW-Vorgaenge "
         "10 Jahre. Zugriffsberechtigte: BVW-Gremium, HR, Vorgesetzte (lesend)."),
        ("§ 7 Inkrafttreten / Geltungsdauer",
         "Diese BV tritt am 1.4.2022 in Kraft und ersetzt die BV BVW vom 1.1.2017. "
         "Sie gilt unbefristet mit Kuendigungsfrist von 6 Monaten zum Jahresende. "
         "Beidseitige Evaluierung jaehrlich im Januar."),
        ("Unterschriften",
         signatures("Andreas Maier", "Werkleiter REG", "Brennhagen Elektronik GmbH",
                    "Klaus Bauer", "Betriebsratsvorsitzender", "BR Werk Heilbronn",
                    place="Heilbronn", date_str_="31. Maerz 2022")),
    ])


# RHU_IC_Rechnung_2020_07
write_doc(BASE / "RHU_IC_Rechnung_2020_07.docx", REG_HDR,
    "Faktura / Rechnung RHU -> REG Nr. RHU-IC-2020-07-207 (Sensorik Juli 2020)",
    subtitle="Brennhagen Hungary Kft. (Gyor) -> Brennhagen Elektronik GmbH (Heilbronn) | "
             "EUR (Konzernfakturierung HUF -> EUR Konzernkurs 350 HUF/EUR)",
    sections=[
        ("Rechnungsdaten",
         "Rechnungsnummer: RHU-IC-2020-07-207\nRechnungsdatum: 15. Juli 2020\n"
         "Leistungsdatum (Teljesites): 31. Juli 2020 (DUZP analogue HU)\n"
         "Faelligkeit: 60 Tage netto (Konzern-Group-Terms)\n"
         "Lieferbedingungen: FCA Gyor, Incoterms 2020"),
        ("Leistender",
         "Brennhagen Hungary Kft.\nGyari ut 5, 9027 Gyor, Magyarorszag\n"
         "Cegjegyzekszama (HRB): Cg.08-09-029876\n"
         "Adoszam: 24587431-2-08\nIBAN: HU42 1177 3016 1111 1018 0000 0000\n"
         "BIC: OTPVHUHB (OTP Bank Nyrt., Gyor-Belvaros)"),
        ("Empfaenger",
         "Brennhagen Elektronik GmbH (REG)\nHeilbronner Strasse 88, 74072 Heilbronn\n"
         "HRB 221456, AG Heilbronn | USt-IdNr. DE312487901"),
        ("Leistungsumfang",
         "Lieferung von Sensorik-Modulen (Drehratensensor, Beschleunigungssensor) "
         "fuer Verbau in ADAS-V4D und ECU-900 gemaess Rahmenliefervertrag RHU-REG-2019. "
         "Verrechnung zu Konzern-Verrechnungspreisen (Cost-Plus 8,0 %) gemaess "
         "Konzern-Transfer-Pricing-Richtlinie 2020. Hinweis: Diese Rechnung gehoert "
         "thematisch zu RHU Gyor und ist vermutlich versehentlich im RCZ-Brno-Ordner "
         "abgelegt worden."),
        ("Positionen",
         [["Pos.", "Bezeichnung", "Menge", "Einzelpreis EUR", "Gesamtpreis EUR"],
          ["1", "Drehratensensor 3-Achs ASD-302",      "18.000", "4,80",  "86.400"],
          ["2", "Beschleunigungssensor 3-Achs ASB-201","22.000", "3,90",  "85.800"],
          ["3", "Sensor-PCB 4-lagig 60x40 mm",         "12.000", "6,20",  "74.400"],
          ["4", "Testkalibrierung Endkontrolle",       "52.000", "0,18",   "9.360"],
          ["",  "Zwischensumme netto",                  "", "",         "255.960"],
          ["",  "MWSt (Reverse Charge §142 HU-USt)",    "", "",               "0"],
          ["",  "Rechnungsbetrag",                       "", "", "255.960 EUR"]]),
        ("Sonstiges",
         "Reverse Charge gemaess ungarischem MwSt-Gesetz § 142 (2007. evi CXXVII. tv.). "
         "Steuerschuld geht auf REG ueber. EC-Sales-List / Zusammenfassende Meldung "
         "monatlich an NAV. Intrastat-Meldung Versendung HU -> DE durch RHU-Logistik."),
        ("Buchhalterische Erfassung",
         "RHU (BUKR HU01) Konto 311000; REG (BUKR DE01) Konto 421000 RHU-IC-Vbk. "
         "Konzernabstimmung Florian Maier (Group Controller)."),
        ("Genehmigung",
         f"Erstellt: Lohnabteilung RHU.\nGeprueft: Lead Finance RHU Andrea Szabo.\n"
         f"Freigegeben: Werkleiter Laszlo Kovacs."),
    ])


# RHU_Legal_Compliance_Report_2022
write_doc(BASE / "RHU_Legal_Compliance_Report_2022.docx", REG_HDR,
    "Legal Compliance Report Brennhagen Hungary Kft. (Gyor) – Geschaeftsjahr 2022",
    subtitle="Berichterstatter: Local Compliance Officer Gyor | "
             "Hinweis: ggf. Fehlablage – Dokument betrifft RHU, nicht RCZ",
    sections=[
        ("Vorbemerkung zur Ablage",
         "Dieses Dokument betrifft die Brennhagen Hungary Kft. (RHU, Gyor) und ist "
         "vermutlich versehentlich im Ordner der Brennhagen CZ s.r.o. (Brno) abgelegt. "
         "Aus Konsistenzgruenden wird der Inhalt dennoch hier nochmals dokumentiert."),
        ("Compliance-Organisation Gyor",
         "Local Compliance Officer RHU: Dr. Eszter Toth (50 %), dotted line an "
         "Chief Compliance Officer REA. Quartalsweise Compliance-Komitee unter "
         "Vorsitz Werkleiter Laszlo Kovacs."),
        ("Wesentliche Compliance-Themen 2022",
         ("list", [
             "Anti-Korruption: Schulung 100 % Indirect Personnel, 0 Hinweise auf Hotline (Convercent).",
             "Kartellrecht: Schulung Sales/Einkauf erfolgt, keine Vorfaelle.",
             "Datenschutz (GDPR / NAIH): Verzeichnis Verarbeitungstaetigkeiten aktuell, "
             "1 Datenpanne (USB-Stick verloren, keine sensiblen Daten), gemeldet NAIH ohne Sanktion.",
             "Arbeitsschutz: LTIFR 5,1 (HU Branche 7,8); keine schweren Unfaelle.",
             "Steuerliche Compliance: Online-Faktura-Meldung (NAV Online Szamla) "
             "fristgerecht; keine Beanstandungen.",
             "Geldwaeschepraevention: nicht primaer betroffen; KYC-Standardprozess.",
             "EHS / Umwelt: ISO 14001 Surveillance-Audit bestanden (0 Major, 2 Minor).",
         ])),
        ("Personalrechtliche Themen",
         "Tarifabschluss mit lokaler Gewerkschaft VASAS am 1.4.2022: +9,5 % Tabellenentgelte "
         "(im Branchen-Mittel HU Auto), Erhoehung Cafeteria 660.000 HUF/Jahr. "
         "Keine arbeitsrechtlichen Klagen oder Eskalationen."),
        ("Audit-Plan 2023",
         "Geplante interne Audits 2023: Lieferantenfreigabe (Q1), Spesen/Reisekosten "
         "(Q2), IT-Berechtigungen (Q3), Q-Datenintegritaet IATF (Q4). "
         "Externe Audits: KPMG Magyarorszag JA-Pruefung (Feb 2023), IATF-Surveillance "
         "(Jun 2023)."),
        ("Schlussvermerk",
         "Gyor, 15.2.2023.\n\nDr. Eszter Toth, Local Compliance Officer.\n"
         "Freigabe: Laszlo Kovacs, Werkleiter."),
    ])


# ===========================================================================
# Final verification
# ===========================================================================

if __name__ == "__main__":
    from docx import Document
    total = thin = 0
    for p in sorted(BASE.rglob("*.docx")):
        d = Document(p)
        t = " ".join(par.text for par in d.paragraphs)
        for tbl in d.tables:
            for r in tbl.rows:
                for c in r.cells:
                    t += " " + c.text
        w = len(t.split())
        total += 1
        if w < 200:
            thin += 1
            print(f"  THIN ({w} w): {p.name}")
    print(f"\n{total} docx total, {thin} still thin")
