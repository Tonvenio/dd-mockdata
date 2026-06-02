"""Brennhagen AG / 07_Tochter_HU_Gyoer (RHU - Brennhagen Hungary Kft., Gyoer).

Re-generates all 102 thin .docx documents with realistic, length-appropriate
content. Subsidiary RHU: Sensorik-Produktion in Gyoer (HU), 540 MA,
Cg.08-09-029876, Werkleiter Laszlo Kovacs, HR Andrea Szabo.
Hungarian labor law; Sozialversicherung NAV (Nemzeti Adohivatal).
Currency HUF where appropriate, EUR for intercompany.
"""
# --- portable-paths-prelude --- (do not edit) ---
import sys
from pathlib import Path as _PathlibPath
_RP = _PathlibPath(__file__).resolve().parent
while not (_RP / "enhance_lib.py").exists() and _RP.parent != _RP:
    _RP = _RP.parent
_ROOT = _RP
sys.path.insert(0, str(_ROOT))
# --- end prelude ---
from enhance_lib import ROEHRIG_AG as R, ROEHRIG_SUBS as S, write_doc, signatures

BASE = f"{_ROOT}/roehrig_large/07_Tochter_HU_Gyoer"
H = {"name": R["name"], "addr": R["addr"], "hrb": R["hrb"]}

RHU = S["RHU"]
RHU_HDR = {
    "name": "Brennhagen Hungary Kft.",
    "addr": "Ipari Park, Bercsenyi liget 14, H-9027 Gyoer, Magyarorszag",
    "hrb":  "Cg.08-09-029876 (Gyoer-Moson-Sopron Megyei Cegbirosag), "
            "Adoszam HU 12876543-2-08",
}

# ----- helper for HUF amounts ------------------------------------------------
def huf(n):
    return f"{n:,} HUF".replace(",", " ")

def eur(n):
    return f"{n:,.0f} EUR".replace(",", ".")


# =============================================================================
# 1) ARBEITSVERTRAEGE RHU (5 docs) + AV_096 stray
# =============================================================================
def rhu_arbeitsvertrag(fname, position_de, position_hu, ma_name, monatslohn_huf,
                       eintritt, vorgesetzt):
    write_doc(f"{BASE}/{fname}", RHU_HDR,
        f"Munkaszerzodes / Arbeitsvertrag - {position_de}",
        subtitle=f"Brennhagen Hungary Kft., Gyoer - {position_hu}",
        sections=[
            ("Vertragsparteien / A szerzodo felek",
             "Munkaltato / Arbeitgeber:\n"
             "Brennhagen Hungary Kft., Ipari Park, Bercsenyi liget 14, H-9027 Gyoer\n"
             "Cegjegyzekszam: Cg.08-09-029876, Adoszam: HU 12876543-2-08\n"
             "vertreten durch / kepviseli: Laszlo Kovacs (uegyvezeto / Werkleiter)\n\n"
             f"Munkavallalo / Arbeitnehmer:\n{ma_name}\n"
             "(Adoazonosito jel und Sozialversicherungsnummer im Personalakt hinterlegt)"),
            ("Praeambel",
             f"Der Arbeitnehmer wird ab {eintritt} als {position_de} ({position_hu}) "
             f"in der ungarischen Produktionsstaette Gyoer der Brennhagen-Gruppe eingestellt. "
             f"Das Beschaeftigungsverhaeltnis unterliegt vollumfaenglich ungarischem Arbeitsrecht, "
             f"insbesondere dem Munka Toerveny Koenyve (Mt., Gesetz I/2012). Soweit der "
             f"Konzern-Tarifrahmen der Brennhagen Elektronik AG (Stuttgart) anwendbar ist, gehen "
             f"die ungarischen Mindeststandards vor; Konzernregelungen gelten nur, wenn sie "
             f"guenstiger sind."),
            ("Vertragsinhalt", ("clauses", [
                ("1. §  Taetigkeit",
                 [f"Der Arbeitnehmer uebernimmt die Position {position_de} im Werk Gyoer. "
                  f"Vorgesetzter ist {vorgesetzt}. Detaillierte Aufgabenbeschreibung "
                  f"(munkakoeri leiras) liegt als Anlage 1 bei und ist Bestandteil dieses "
                  f"Vertrages.",
                  "Aufgabenschwerpunkte: Sensorik-Produktion (Drehzahlsensoren, "
                  "Beschleunigungssensoren, Drucksensoren) gemaess IATF 16949 sowie der "
                  "kundenspezifischen Anforderungen (BMW, Mercedes, Audi, VW)."]),
                ("2. §  Beginn / Probezeit (probaido)",
                 [f"Vertragsbeginn: {eintritt}. Probezeit gem. Mt. § 45 (5): drei Monate, "
                  f"beidseitige Kuendigung jederzeit ohne Begruendung waehrend dieser Zeit.",
                  "Befristung: unbefristet (hatarozatlan idore szolo)."]),
                ("3. §  Arbeitszeit",
                 [f"Vollzeitbeschaeftigung 40 Stunden/Woche im Drei-Schicht-Modell "
                  f"(reggeli/dleutni/ejszakai musak). Zuschlaege gem. Mt.: 50 % "
                  f"Ueberstunde, 100 % Sonn-/Feiertag, 30 % Nachtarbeit (22:00-06:00)."]),
                ("4. §  Verguetung / Munkaber",
                 [f"Bruttomonatsgehalt: {huf(monatslohn_huf)} (zwoelf Monatsgehaelter "
                  f"plus 13. Monatsgehalt im Dezember als Bonus, sofern Kennzahlen erreicht).",
                  "Sozialversicherungsabzuege gem. ungarischem Recht: 18,5 % "
                  "Arbeitnehmerbeitrag (15 % SZJA + 18,5 % TB) plus 15 % "
                  "Einkommensteuer. Arbeitgeberseitig Szocho 13 % (Sozialbeitragssteuer) "
                  "und 1,5 % Berufsbildungsbeitrag - Abfuehrung an NAV (Nemzeti "
                  "Adohivatal, ungarische Steuer- und Sozialversicherungsbehoerde) "
                  "monatlich.",
                  "Zusatzleistungen: SZEP-Karte (Cafeteria) jaehrlich 450.000 HUF "
                  "steuerfrei; betriebliche Pruefungspraemie 12 % bei Q-Quote >= 99,5 %."]),
                ("5. §  Urlaub",
                 ["Jahresurlaub gem. Mt. § 116: 20 Werktage Grundurlaub plus "
                  "altersabhaengige Zusatzurlaubstage (+1 Tag pro 3 Jahre ab 25, max. "
                  "+10 Tage). Antrag schriftlich beim Vorgesetzten, Bewilligung durch "
                  "HR (Andrea Szabo)."]),
                ("6. §  Geheimhaltung / titoktartas",
                 ["Der Arbeitnehmer verpflichtet sich zur strikten Geheimhaltung "
                  "saemtlicher Geschaefts- und Betriebsgeheimnisse der Brennhagen-Gruppe, "
                  "insbesondere OEM-Spezifikationen (BMW, Mercedes, VW), Produktions- "
                  "und Pruefverfahren der Sensorik-Linien sowie Personaldaten. "
                  "Dauer: unbefristet ueber das Vertragsende hinaus."]),
                ("7. §  Wettbewerbsverbot",
                 ["Nachvertragliches Wettbewerbsverbot gem. Mt. § 228: 12 Monate fuer "
                  "Fuehrungspositionen, Karenzentschaedigung 50 % des letzten Brutto "
                  "(monatlich nachschuessig)."]),
                ("8. §  Kuendigung",
                 ["Kuendigungsfristen gem. Mt. § 69: gestaffelt 30 Tage (bis 3 Jahre "
                  "Betriebszugehoerigkeit) bis 90 Tage (ab 20 Jahre). Schriftform "
                  "(irasos formaba) zwingend. Bei verhaltensbedingter ausserordentlicher "
                  "Kuendigung (rendkivuli felmondas) Frist 15 Tage ab Kenntnis."]),
                ("9. §  Anwendbares Recht / Gerichtsstand",
                 ["Es gilt ungarisches Recht. Gerichtsstand ausschliesslich Gyoeri "
                  "Toervenyszek (Arbeitsgericht Gyoer). Bei konzernrechtlichen "
                  "Streitigkeiten mit der Brennhagen Elektronik AG ist - vorbehaltlich "
                  "zwingenden lokalen Rechts - ICC-Schiedsklausel mit Sitz Frankfurt "
                  "vereinbart."]),
            ])),
            ("Unterschriften",
             signatures(ma_name, position_de, "Brennhagen Hungary Kft.",
                        "Laszlo Kovacs", "uegyvezeto / Werkleiter", "Brennhagen Hungary Kft.",
                        place="Gyoer", date_str_=eintritt)),
        ])


rhu_arbeitsvertrag("RHU_Arbeitsvertrag_01_Geschäftsführer_in_R_2022.docx",
    "Geschaeftsfuehrerin Recht & Compliance", "jogi es compliance igazgato",
    "Dr. Eszter Nagy", 1_850_000, "1. April 2022",
    "Laszlo Kovacs (Werkleiter)")
rhu_arbeitsvertrag("RHU_Arbeitsvertrag_02_Produktionsleiter_in_2022_v2.docx",
    "Produktionsleiter Sensorik", "termelesi vezeto - szenzorika",
    "Tamas Horvath", 1_450_000, "1. Juni 2022",
    "Laszlo Kovacs (Werkleiter)")
rhu_arbeitsvertrag("RHU_Arbeitsvertrag_03_Qualitätsmanagerin_G_2022.docx",
    "Qualitaetsmanagerin Gyoer", "minosegiranyitasi vezeto",
    "Katalin Kiss", 1_280_000, "15. Mai 2022",
    "Laszlo Kovacs (Werkleiter)")
rhu_arbeitsvertrag("RHU_Arbeitsvertrag_04_Finanzcontroller_Győ_2022.docx",
    "Finanzcontroller Gyoer", "penzuegyi kontroller",
    "Janos Toth", 1_180_000, "1. September 2022",
    "Laszlo Kovacs (Werkleiter)")
rhu_arbeitsvertrag("RHU_Arbeitsvertrag_05_HR-Manager_Győr_2022.docx",
    "HR-Managerin Gyoer", "humaneroforras vezeto",
    "Andrea Szabo", 1_320_000, "1. Maerz 2022",
    "Laszlo Kovacs (Werkleiter)")


# AV_096 - misplaced cross-folder AV
write_doc(f"{BASE}/AV_096_Nicole_Lange_96_Head_of_Quality_IA.docx", H,
    "Arbeitsvertrag - Nicole Lange (Head of Quality, Interne Audit)",
    subtitle="REA-Konzern - Interaudit-Hinweis (urspruenglich abgelegt in 07_RHU); "
             "Vertrag selbst betrifft Konzernholding RHO Stuttgart",
    sections=[
        ("Vertragsparteien",
         "Arbeitgeber: Brennhagen Holding GmbH, Vaihinger Strasse 120, 70567 Stuttgart, "
         "HRB 726450, AG Stuttgart, vertreten durch CEO Anna Mueller.\n\n"
         "Arbeitnehmerin: Nicole Lange, geb. 1979, Tuebinger Strasse 22, 70178 Stuttgart."),
        ("Position / Funktion",
         "Frau Lange uebernimmt zum 1. September 2022 die Position Head of Quality "
         "Group / Interne Auditorin Q-Management mit Berichtslinie an Andreas Buehler "
         "(Chief Audit Executive). Funktional ist sie zustaendig fuer die Qualitaets- "
         "und Prozessaudits in allen Produktionswerken der Brennhagen-Gruppe (REG Heilbronn, "
         "RPL Katowice, RCZ Brno, RHU Gyoer, RCN Shanghai)."),
        ("Aufgaben",
         ("list", [
            "Jaehrliches Internal-Audit-Programm Q (Plan + Durchfuehrung) gem. IATF 16949",
            "Vor-Ort-Audits in den Werken inkl. RHU Gyoer (Sensorik-Linien)",
            "Eskalation an Group CEO/COO bei kritischen Findings (Schwellwert: 3 majors)",
            "VDA-6.3 Prozessaudits bei Lieferanten (Stichprobenkontrolle)",
            "Reporting an Pruefungsausschuss des Aufsichtsrats (halbjaehrlich)",
         ])),
        ("Verguetung",
         "Jahresgrundgehalt 118.000 EUR brutto (13x 9.077 EUR) zzgl. variabler Tantieme "
         "von bis zu 25 % des Grundgehalts. Dienstwagen Klasse Audi A4/BMW 3er, "
         "betriebliche Altersversorgung im Konzern-Standard (Heubeck-Rahmenwerk)."),
        ("Vertraulichkeit",
         "Mit Zugang zu hochsensiblen Audit-Findings ueber alle Werke verpflichtet sich "
         "Frau Lange zu absoluter Vertraulichkeit. Verstoesse koennen zur fristlosen "
         "Kuendigung sowie zu Schadensersatzanspruechen fuehren. Die Verpflichtung gilt "
         "ueber das Vertragsende hinaus unbefristet."),
        ("Unterschrift",
         signatures("Nicole Lange", "Head of Quality (designate)", "Brennhagen Holding GmbH",
                    "Anna Mueller", "Vorstandsvorsitzende", "Brennhagen Elektronik AG",
                    place="Stuttgart", date_str_="22. August 2022")),
    ])


# =============================================================================
# 2) CAPA reports (2)
# =============================================================================
def capa(fname, nr, datum, kunde, produkt, problem, ursache, massnahmen):
    write_doc(f"{BASE}/{fname}", RHU_HDR,
        f"CAPA-Report Nr. {nr} (Corrective and Preventive Action)",
        subtitle=f"Werk Gyoer - Sensorik-Produktion - {kunde} / {produkt}",
        sections=[
            ("Stammdaten", [
                ["Feld", "Wert"],
                ["CAPA-Nr.", nr],
                ["Eroeffnungsdatum", datum],
                ["Werk / Standort", "Brennhagen Hungary Kft., Gyoer (RHU)"],
                ["Linie", f"Sensorik {produkt}"],
                ["Kunde", kunde],
                ["CAPA-Owner", "Katalin Kiss (Qualitaetsmanagerin RHU)"],
                ["Eskalation an", "Sabine Brand (Group Q REA), Andreas Buehler (CAE)"],
                ["Status", "OPEN - in Abarbeitung"],
            ]),
            ("Problembeschreibung",
             problem),
            ("Sofortmassnahmen (Containment)",
             ("list", [
                "100 %-Sortierpruefung der WIP-Bestaende (ca. 18.400 Stueck) - Mehrarbeit Schicht 3",
                "Sperrlager-Verfuegung fuer alle Loose mit Verdacht (Q-Sperre Status: rot)",
                "Information Kunde durch Werkleiter Laszlo Kovacs binnen 24 h (telefonisch + 8D)",
                "Lieferstopp ab Folgetag bis Wirksamkeitsnachweis der Korrektur",
             ])),
            ("Ursachenanalyse (5-Why / Ishikawa)",
             ursache),
            ("Korrekturmassnahmen", ("list", massnahmen)),
            ("Wirksamkeitspruefung",
             "Die Wirksamkeit der Korrekturmassnahmen wird nach 30 / 60 / 90 Tagen geprueft. "
             "Pruefkriterien: ppm-Quote, Reklamationseingang, interne Sortierergebnisse. "
             "Audit durch Group Q (Sabine Brand) im naechsten Quartal. Schliessung der CAPA "
             "erfolgt erst nach dokumentiertem Wirksamkeitsnachweis durch Q-Leitung."),
            ("Freigabe / Unterschrift",
             signatures("Katalin Kiss", "Q-Leitung RHU", "Brennhagen Hungary Kft.",
                        "Laszlo Kovacs", "Werkleiter RHU", "Brennhagen Hungary Kft.",
                        place="Gyoer", date_str_=datum)),
        ])


capa("CAPA_2023_0010.docx", "RHU-CAPA-2023-010", "12. Juni 2023", "BMW Group",
     "Drehzahlsensor RHU-DRZ-4012 (BMW Plant Dingolfing)",
     "BMW meldete am 8.6.2023 eine signifikante Erhoehung der Rueckweiseraten in der "
     "Wareneingangspruefung (PPM ueber Grenzwert 50, gemessen 142). Betroffen sind die "
     "Liefermengen aus KW 21-23/2023, Charge L23-21-A bis L23-23-C. Symptom: "
     "Signalrauschen oberhalb 8.500 U/min, intermittierende Aussetzer im Pruefstand.",
     "5-Why-Analyse identifizierte Kalibrierdrift im automatischen Pruefstand TP-04. "
     "Die wochentliche Master-Kalibrierung war ausgeblieben, weil die zustaendige "
     "Q-Technikerin (Krankheitsfall) nicht ordnungsgemaess vertreten wurde. "
     "Root cause: Vertretungsplan nicht aktualisiert, fehlende automatisierte Erinnerung "
     "im MES-System.",
     [
        "Master-Kalibrierung des Pruefstandes TP-04 erfolgt (15.6.2023)",
        "Wochentliche automatische MES-Erinnerung an Schichtleiter implementiert",
        "Vertretungsplan Q-Technik aktualisiert (HR Andrea Szabo)",
        "Schulung 8 Q-Mitarbeiter (8 h) zur SOP Pruefstand-Kalibrierung",
        "Erweiterung der Selbstpruefung im Linientakt (Stichproben alle 100 Teile)",
        "Reporting an BMW (8D-Antwort D7 / D8 spaetestens 30.6.2023)",
     ])

capa("CAPA_2023_0016.docx", "RHU-CAPA-2023-016", "8. Oktober 2023", "Mercedes-Benz AG",
     "Beschleunigungssensor RHU-BES-5520 (MB-Werk Sindelfingen)",
     "Mercedes-Benz reklamierte am 2.10.2023 fehlerhafte Loetstellen am Anschlussstecker "
     "(intermittierender Kontaktverlust). Betroffen Lieferungen Q3/2023, geschaetzte "
     "Population ca. 7.200 Stueck verbaut. 0-km-Ausfall im MB-Werk: 23 Faelle.",
     "Ursache: temporaere Erhoehung der Loetbad-Temperatur an SMD-Linie 2 (Anstieg um "
     "8 Grad C ueber Sollwert) aufgrund defekter Temperatursonde. Wartungsintervall "
     "war ueberschritten (letzte Kalibrierung 9 Monate alt, SOP fordert 6 Monate).",
     [
        "Temperatursonde an SMD-Linie 2 ausgetauscht und neu kalibriert (5.10.2023)",
        "Selektion und 8D-Quarantaene aller verdaechtigen Loose (3.142 Stueck)",
        "100 %-Roentgenpruefung der Loetstellen fuer 6 Wochen (eskalierte Inspektion)",
        "Wartungsintervall aller Temperatursonden auf 3 Monate reduziert (SOP-Update)",
        "Reklamationsbearbeitung Mercedes inkl. Cost-of-Quality-Tragung (ca. 62.000 EUR)",
        "Lessons-Learned-Workshop mit RPL Katowice (ueberbetrieblicher Wissenstransfer)",
     ])


# =============================================================================
# 3) misplaced single-shots
# =============================================================================
write_doc(f"{BASE}/CZ_RCZ_Arbeitssicherheitsbericht_2021.docx",
    {"name": "Brennhagen CZ s.r.o.",
     "addr": "Prumyslova 145, 627 00 Brno, Czech Republic",
     "hrb":  "C 87654 vedena u Krajskeho soudu v Brne, IC 24876543"},
    "Arbeitssicherheitsbericht 2021 - Werk Brno (RCZ)",
    subtitle="Bericht gem. Zakon c. 309/2006 Sb. (CZ Arbeitsschutzgesetz) - "
             "Beigabe zur Konzern-EHS-Konsolidierung",
    sections=[
        ("Kennzahlen 2021", [
            ["Kennzahl", "Wert 2021", "Vorjahr 2020", "Trend"],
            ["LTIFR (lost-time injury freq. rate)", "3,2", "4,1", "verbessert"],
            ["Meldepflichtige Arbeitsunfaelle", "11", "14", "-3"],
            ["Ausfalltage (gesamt)", "182", "239", "-57"],
            ["Beinaheunfaelle (gemeldet)", "84", "62", "+22 (positives Reporting)"],
            ["Schulungsstunden Arbeitssicherheit", "4.680", "3.920", "+19 %"],
        ]),
        ("Wesentliche Ereignisse 2021",
         "Im Berichtsjahr ereigneten sich elf meldepflichtige Arbeitsunfaelle. Der "
         "schwerste Vorfall war ein Schnittverletzung an Linie 3 (Stanzwerkzeug) am "
         "14. Mai 2021 mit 28 Ausfalltagen. Untersuchung ergab unzureichende "
         "Schutzvorrichtung; Massnahme: Nachruestung Lichtgitter (28.000 EUR, "
         "abgeschlossen Juli 2021)."),
        ("Massnahmen und Audits",
         ("list", [
            "Externes Audit durch TUEV SUED Czech Republic (Maerz 2021) - bestanden",
            "EHS-Wochen-Walk Werkleiter Petr Novak: 52 Termine, 312 Findings",
            "Aktualisierung von 14 Standard-SOPs Arbeitssicherheit",
            "Schulungsoffensive Stapler-Verkehr (alle 124 Fahrer, jaehrlich)",
         ])),
        ("Fazit / Ausblick",
         "Die LTIFR verbessert sich kontinuierlich und liegt mit 3,2 unter dem "
         "tschechischen Branchenmittel (4,8). Fuer 2022 ist die ISO-45001-Zertifizierung "
         "geplant (Auditor TUEV NORD CZ)."),
        ("Bericht erstellt",
         signatures("Eva Cerna", "EHS-Beauftragte", "Brennhagen CZ s.r.o.",
                    "Petr Novak", "Werkleiter", "Brennhagen CZ s.r.o.",
                    place="Brno", date_str_="20. Januar 2022")),
    ])


write_doc(f"{BASE}/DSGVO_Datenpanne_2023_011.docx", H,
    "DSGVO-Meldung Datenpanne 2023-011",
    subtitle="Konzerndatenschutz Brennhagen-Gruppe - Vorfallsbericht gem. Art. 33 DSGVO",
    confidential=True,
    sections=[
        ("Vorfall-Stammdaten", [
            ["Feld", "Wert"],
            ["Vorfall-Nr.", "REA-DSGVO-2023-011"],
            ["Erstmeldung intern", "14. September 2023, 11:42 Uhr (MEZ)"],
            ["Vorfallseintritt (geschaetzt)", "10. - 13. September 2023"],
            ["Entdeckungsweg", "Mitarbeiterhinweis Service-Desk RHU Gyoer"],
            ["Betroffene Standorte", "RHU Gyoer (primaer); Konzern-VPN angebunden"],
            ["DPO-Eskalation", "Dr. Heike Berger (Group DPO), Hengeler Mueller (legal)"],
            ["Behoerden-Meldepflicht", "JA - 72 h Frist Art. 33 DSGVO eingehalten"],
        ]),
        ("Sachverhalt",
         "Im Werk Gyoer (RHU) wurde ein USB-Stick mit ca. 4.200 unverschluesselten "
         "Personaldaten-Datensaetzen (Name, Adresse, SV-Nummer, Bruttogehalt) aus dem "
         "Spind eines HR-Mitarbeiters entwendet. Der Stick wurde fuer eine Excel-basierte "
         "Lohnabrechnung verwendet (entgegen interner SOP). Eine Verbreitung der Daten "
         "konnte bislang nicht festgestellt werden; eine Anzeige bei der ungarischen "
         "Polizei (Gyoeri Rendor-Kapitanysag) wurde am 14.9.2023 erstattet."),
        ("Massnahmen",
         ("list", [
            "Information aller 4.200 Betroffenen schriftlich binnen 7 Tagen (Mt. § 11)",
            "Meldung an NAIH (Nemzeti Adatvedelmi es Informacioszabadsag Hatosag) am 16.9.2023",
            "Meldung an Konzern-LfDI Baden-Wuerttemberg ueber Group DPO",
            "USB-Stick-Nutzung in RHU vollstaendig durch Endpoint-Policy unterbunden",
            "Verpflichtende Datenschutz-Schulung aller 540 RHU-Mitarbeiter (Q4/2023)",
            "Spind-Sicherheitskonzept HR-Bereich (Schluesselverwaltung neu)",
         ])),
        ("Risikobewertung / Sanktionsrisiko",
         "Die DPIA-Nachbewertung stuft das Risiko als MITTEL ein (hoechste Sensitivitaet: "
         "Gehaltsdaten und SV-Nummer). Ein Bussgeld-Risiko nach Art. 83 DSGVO ist gegeben; "
         "der NAIH hat in vergleichbaren Faellen 5-15 Mio. HUF verhaengt. Die "
         "Konzern-Cyberversicherung (Allianz) wurde informiert."),
        ("Unterschrift",
         signatures("Dr. Heike Berger", "Group DPO / Tax Director", R["name"],
                    "Andrea Szabo", "HR-Leitung RHU", "Brennhagen Hungary Kft.",
                    place="Stuttgart / Gyoer", date_str_="16. September 2023")),
    ])


# REA_BOS_Langfristvertrag_LTA_2023
write_doc(f"{BASE}/REA_BOS_Langfristvertrag_LTA_2023.docx", H,
    "Langfristvertrag (LTA) Robert Bosch GmbH - Brennhagen Elektronik AG",
    subtitle="Rahmenliefervertrag Sensorik / Steuergeraete 2023-2027 "
             "(stray-filed im RHU-Ordner, betrifft Konzern-LTA)",
    sections=[
        ("Vertragsparteien",
         "Kunde: Robert Bosch GmbH, Robert-Bosch-Platz 1, 70839 Gerlingen-Schillerhoehe, "
         "HRB 14000, AG Stuttgart, vertreten durch Dr. Markus Heyn (CEO Mobility Solutions).\n\n"
         "Lieferant: Brennhagen Elektronik AG (Holding) mit Lieferwerk Brennhagen Hungary Kft., "
         "Gyoer (RHU) fuer Sensorik sowie Brennhagen Elektronik GmbH, Heilbronn (REG) fuer "
         "Steuergeraete-Plattformen."),
        ("Vertragsgegenstand", ("clauses", [
            ("1. §  Liefergegenstaende",
             ["Beschleunigungssensoren BES-5500-Familie (RHU Gyoer)",
              "Drucksensoren DRS-3300 (RHU Gyoer)",
              "Powertrain-ECU ECU-900 (REG Heilbronn)",
              "Anwendung: Bosch Mobility Solutions Plattformen ab MJ 2024 fuer "
              "OEM-Programme VW (MEB+), Stellantis (STLA Medium)."]),
            ("2. §  Volumen",
             ["Indikatives Gesamtvolumen 2023-2027: ca. 312 Mio. EUR.",
              "RHU-Anteil ca. 38 % (118 Mio. EUR), REG-Anteil ca. 62 % (194 Mio. EUR).",
              "Keine Mindestabnahmeverpflichtung; Bosch garantiert jedoch "
              "Volumenkorridor +/- 15 % gegenueber Rolling Forecast."]),
            ("3. §  Preise / Preisanpassung",
             ["Preise EUR-basiert. RHU-Sensorik-Preise unterliegen jaehrlicher "
              "Preisverhandlung (Productivity Sharing 2,5 % p. a. ab Y2).",
              "Rohstoff-Indexierung fuer Halbleiter (PHLX SOX), Edelmetalle (LME) und "
              "EUR/HUF-Wechselkurs (>5 % Schwankung loest Verhandlung aus)."]),
            ("4. §  Qualitaet / IATF 16949",
             ["Werk Gyoer (RHU) und Heilbronn (REG) sind nach IATF 16949 zertifiziert. "
              "Bosch-spezifische BPS-Anforderungen (Bosch Production System) sind im "
              "Anhang Q1 zu erfuellen. PPAP/PPF Level 3 fuer alle Erstmuster."]),
            ("5. §  Logistik",
             ["JIS/JIT-Anlieferung an Bosch-Werke (Stuttgart-Feuerbach, Bamberg, Reutlingen). "
              "Incoterms 2020: DDP Werk Bosch. Verpackung in Bosch-Mehrweg-Behaeltern (KLT)."]),
            ("6. §  Haftung",
             ["Produkthaftung gem. Bosch-Einkaufsbedingungen (BEB), Haftungshoechstgrenze "
              "10 Mio. EUR/Schadensfall, 20 Mio. EUR/Versicherungsjahr. "
              "Feldausfall-Risikoverteilung gem. ZVEI-Empfehlung Field-Failure-Cost-Sharing."]),
            ("7. §  Laufzeit / Kuendigung",
             ["1.1.2023 - 31.12.2027 mit Option auf Verlaengerung (zwei Jahre, einseitig Bosch). "
              "Kuendigung aus wichtigem Grund jederzeit moeglich (z.B. wiederholte 8D mit Klasse-A-Fehler)."]),
        ])),
        ("Unterschriften",
         signatures("Dr. Markus Heyn", "CEO Mobility Solutions", "Robert Bosch GmbH",
                    "Anna Mueller", "Vorstandsvorsitzende", "Brennhagen Elektronik AG",
                    place="Gerlingen / Stuttgart", date_str_="14. Februar 2023")),
    ])


# REA_IFRS_Bilanzierungsrichtlinie
write_doc(f"{BASE}/REA_IFRS_Bilanzierungsrichtlinie_IAS_36_Wertminderung_2023.docx", H,
    "Konzern-Bilanzierungsrichtlinie IAS 36 - Wertminderung von Vermoegenswerten",
    subtitle="Group Accounting Policy 2023 (CGU-Definition u.a. RHU Gyoer)",
    sections=[
        ("Anwendungsbereich",
         "Diese Richtlinie regelt die Anwendung von IAS 36 (Impairment of Assets) im "
         "Brennhagen-Konzern (REA und alle vollkonsolidierten Tochtergesellschaften REG, RSG, "
         "RPL, RCZ, RHU, RCN, RHO). Sie gilt fuer immaterielle Vermoegenswerte, "
         "Sachanlagen und cash-generating units (CGU)."),
        ("CGU-Festlegung im Konzern", [
            ["CGU-Code", "Beschreibung", "Buchwert 31.12.2023 (Mio. EUR)"],
            ["CGU-REG", "Werk Heilbronn (ICP-3, BMS-12)", "186"],
            ["CGU-RSG", "Software-Entwicklung Muenchen", "42"],
            ["CGU-RPL", "EMS Katowice", "128"],
            ["CGU-RCZ", "Steckverbinder Brno", "94"],
            ["CGU-RHU", "Sensorik Gyoer", "72"],
            ["CGU-RCN", "Vertrieb Shanghai", "31"],
        ]),
        ("Werthaltigkeitstest",
         "Werthaltigkeitstests werden jaehrlich zum 30.9. (Vor-Test) und 31.12. (final) "
         "durchgefuehrt. Der erzielbare Betrag (recoverable amount) ist der hoehere "
         "Betrag aus Nutzungswert (value in use) und beizulegendem Zeitwert abzueglich "
         "Veraeusserungskosten (FVLCS).\n\n"
         "Diskontierungszinssatz (WACC) je CGU: gewichtete Kapitalkosten unter "
         "Beruecksichtigung Laenderrisiko. WACC RHU 2023: 9,4 % (EUR-basiert), "
         "WACC RCN 2023: 11,8 % (CNY-basiert)."),
        ("Triggering Events",
         ("list", [
            "OEM-Kuendigung eines kritischen Programms (z.B. ICP-3 BMW)",
            "Wesentlicher Wechselkursverfall (HUF/EUR > 15 % Schwaeche)",
            "Strategische Strukturveraenderungen (z.B. Werksschliessung)",
            "Marktwert REA-Aktie < Buchwert pro Aktie ueber 6 Monate",
         ])),
        ("Verantwortlichkeiten",
         "Group Controlling (Florian Maier) erstellt die jaehrlichen Impairment-Tests. "
         "Pruefung durch KPMG (Dr. Maximilian Brand, Lead Partner). Dokumentation in "
         "der IFRS-Impairment-Datei je CGU mit Sensitivitaetsanalyse."),
        ("Inkrafttreten",
         "Diese Richtlinie tritt zum 1.1.2023 in Kraft und ersetzt die Version 2020. "
         "Naechste Ueberpruefung 2026."),
        ("Freigabe",
         signatures("Florian Maier", "Group Controller", R["name"],
                    "Laura Bauer", "CFO", R["name"],
                    place="Stuttgart", date_str_="15. Dezember 2022")),
    ])


# REA_STE ECR (Engineering Change Request)
write_doc(f"{BASE}/REA_STE_ADAS-V4D_ECR_2023_003.docx", H,
    "Engineering Change Request ECR-2023-003 ADAS-V4D",
    subtitle="Steuergeraete-Plattform ADAS-V4D - Aenderungsantrag (Software-Stack)",
    sections=[
        ("ECR-Stammdaten", [
            ["Feld", "Wert"],
            ["ECR-Nr.", "ADAS-V4D-ECR-2023-003"],
            ["Antragsteller", "Dr. Klaus Kessler (Werkleiter RSG Muenchen)"],
            ["Plattform", "ADAS-V4D (Radar Fusion Steuergeraet, Level-2/3)"],
            ["Programm", "Mercedes-Benz EQS-Refresh, Stellantis STLA Large"],
            ["Eskalation an", "CTO (S. Hoffmann; ab 1.7.2024 Dr. P. Hollmann)"],
            ["Priority", "P2 (geplant zur naechsten SOP-Phase)"],
        ]),
        ("Beschreibung der Aenderung",
         "Software-Update der Radar-Fusion-Library V4.2.1 -> V4.3.0 (basierend auf ASPICE "
         "Level 3 Reifegrad). Erweiterung der Sensorfusion um zusaetzlichen Eingangskanal "
         "fuer einen 4D-Imaging-Radar (Continental ARS640) der zur 2024-Plattform "
         "verfuegbar wird.\n\nHauptlieferant Sensorik: Brennhagen Hungary Kft. (RHU Gyoer) - "
         "kein Hardware-Change auf Sensorseite erforderlich, ausschliesslich Firmware-Update "
         "in den RHU-Beschleunigungssensoren BES-5500."),
        ("Auswirkungen", ("list", [
            "Software-Build RSG Muenchen: ca. 320 PT Mehraufwand (Lars Wittmann)",
            "Validierung: zusaetzliche 14.000 km Roadtest erforderlich (Q3/2023)",
            "RHU Gyoer Firmware-Update: 4-wochige PPAP-Wiederholung",
            "Stueckkosten-Aenderung: -1,40 EUR pro Steuergeraet (Skaleneffekt Fusion)",
            "Cybersecurity-Bewertung gem. ISO/SAE 21434 erforderlich (Re-Audit)",
        ])),
        ("Genehmigung",
         "Genehmigt durch CCB (Change Control Board) am 8. August 2023 nach Pruefung der "
         "FMEA-Auswirkungen, der Cybersecurity-Implikationen (TARA-Aktualisierung) und "
         "der Kunden-Notifikationsanforderungen (Mercedes-Benz Customer Specific "
         "Requirements MBN). Die ECR umfasst weiterhin eine Re-Validierung der "
         "Funktionssicherheits-Anforderungen gem. ISO 26262 (ASIL-B) durch das "
         "Sicherheits-Team bei RSG Muenchen. SOP-Termin der neuen Version V4.3.0: "
         "15. Januar 2024. Hardware-Lieferanten RHU und Continental wurden bilateral "
         "informiert; PPAP-Wiederholung fuer betroffene RHU-Sensorik startet "
         "15. September 2023."),
        ("Risikobewertung",
         "Restrisiken: (a) Verzoegerung der Continental-ARS640-Lieferung um >4 Wochen "
         "wuerde den SOP-Termin gefaehrden; Mitigation: Lieferantengespraeche taktisch "
         "verschaerft, Penalty-Klauseln vereinbart. (b) Cybersecurity-Reauditing durch "
         "TUEV SUED kann zusaetzliche Funde generieren; Pufferzeit 2 Wochen vorgesehen."),
        ("Unterschrift",
         signatures("Dr. Klaus Kessler", "Werkleiter RSG", "Brennhagen Software GmbH",
                    "Stefan Hoffmann", "CTO", R["name"],
                    place="Muenchen / Stuttgart", date_str_="8. August 2023")),
    ])


# RHO_BR_Protokoll
write_doc(f"{BASE}/RHO_BR_Protokoll_2023_11.docx", H,
    "Betriebsrats-Protokoll RHO Stuttgart - Sitzung November 2023",
    subtitle="Brennhagen Holding GmbH - ordentliche BR-Sitzung",
    sections=[
        ("Sitzungsdaten",
         "Datum: 21. November 2023, 09:00 - 12:30 Uhr\n"
         "Ort: RHO-Verwaltungsgebaeude Stuttgart, Konferenzraum K3\n"
         "Vorsitz: Marlies Duerr (Konzern-BR-Vorsitzende, IG Metall)\n"
         "Schriftfuehrung: Klaus Bauer (Stv. Vorsitzender)\n"
         "Beschlussfaehigkeit: festgestellt (9 von 9 Mitgliedern anwesend)"),
        ("Tagesordnung", ("list", [
            "TOP 1 - Protokoll der Vorsitzung (Oktober 2023)",
            "TOP 2 - Bericht des Vorstands (Anna Mueller, telefonisch zugeschaltet 15 Min.)",
            "TOP 3 - Personalsituation Tochterwerke (RHU Gyoer, RCZ Brno)",
            "TOP 4 - Bonusplan 2024 - Anhoerung gem. § 87 BetrVG",
            "TOP 5 - Auslandsentsendungen 2024 (insb. RCN Shanghai-Aufbau)",
            "TOP 6 - Verschiedenes",
        ])),
        ("TOP 3 - Tochterwerke RHU / RCZ",
         "Der Vorstand berichtete telefonisch ueber die positive Entwicklung in Gyoer "
         "(RHU): die Sensorik-Linie BES-5500 ist erfolgreich angelaufen, neue Kapazitaeten "
         "fuer den Hyundai-Auftrag 2024 sind gesichert. Werkleiter Laszlo Kovacs plant "
         "zusaetzlich 40 Mitarbeiter im Q1/2024. Der BR begruesst den Aufbau, mahnt aber "
         "an, dass die ungarischen Mitarbeiter Zugang zu Konzern-Bildungsangeboten "
         "erhalten muessen (Sprachkurse Deutsch B1 fuer Schichtleiter)."),
        ("TOP 4 - Bonusplan 2024",
         "Der Vorstand stellt den geplanten Konzern-Bonusplan 2024 vor: Komponente A "
         "(EBIT-Konzern, 40 %), Komponente B (EBIT-Tochter, 30 %), Komponente C "
         "(individuelle Ziele, 30 %). Der BR fordert eine ESG-Komponente (CO2-Reduktion "
         "Scope 1+2). Entscheidung vertagt auf Dezember-Sitzung."),
        ("TOP 6 - Verschiedenes",
         "Die naechste Sitzung findet am 12. Dezember 2023 statt. Konzern-BR-Klausur "
         "Januar 2024 in Heilbronn geplant."),
        ("Unterschrift",
         signatures("Marlies Duerr", "Vorsitzende Konzern-BR", "Brennhagen-Gruppe",
                    "Klaus Bauer", "Stv. Vorsitzender", "Brennhagen-Gruppe",
                    place="Stuttgart", date_str_="21. November 2023")),
    ])


# =============================================================================
# 4) RHU Compliance Report
# =============================================================================
write_doc(f"{BASE}/RHU_Compliance_Report_2023.docx", RHU_HDR,
    "Compliance-Jahresbericht 2023 - Brennhagen Hungary Kft.",
    subtitle="Bericht an Group Compliance / Group Audit (CAE Andreas Buehler)",
    sections=[
        ("Berichtsumfang",
         "Dieser Compliance-Bericht der Brennhagen Hungary Kft. (RHU) umfasst das "
         "Geschaeftsjahr 2023 (1.1. - 31.12.2023). Bericht erstattet an Group "
         "Compliance der Brennhagen Elektronik AG (Stuttgart) sowie an den "
         "Pruefungsausschuss-Vorsitzenden des Aufsichtsrats "
         "(Prof. Dr.-Ing. Gerhard Voss)."),
        ("Compliance-Organisation RHU", [
            ["Funktion", "Verantwortlich"],
            ["Werkleiter / Geschaeftsfuehrer", "Laszlo Kovacs"],
            ["Compliance-Koordinatorin RHU", "Dr. Eszter Nagy"],
            ["Datenschutz (DPO local)", "Andrea Szabo (HR + Datenschutz)"],
            ["Anti-Korruption-Ansprechperson", "Dr. Eszter Nagy"],
            ["Whistleblower-Channel", "Konzern-Hotline (EQS Integrity Line, 24/7, ungarisch)"],
        ]),
        ("Berichtsthemen 2023",
         ("list", [
            "Anti-Korruption / Geldwaesche: 4 Schulungen (alle 540 MA + 28 externe), 100 % Coverage",
            "Datenschutz / NAIH (ungarische Datenschutzbehoerde): 1 Vorfall gemeldet (DSGVO-2023-011, "
            "siehe gesondertes Dokument). Bussgeldverfahren laeuft, Allianz-Versicherung informiert.",
            "Kartellrecht: keine Vorfaelle. Konzern-Schulung in November 2023 absolviert.",
            "Steuer-Compliance (NAV - Nemzeti Adohivatal): Steuererklaerungen fristgerecht. "
            "Aussenpruefung KSt/USt fuer 2020-2022 in Q3/2023 abgeschlossen, Bescheid ohne Mehrsteuer.",
            "Exportkontrolle / Dual-Use: 12 Lieferungen China/Asien geprueft, alle freigegeben.",
            "Arbeitsrecht (Mt. - Munka Toerveny Koenyve): 2 individuelle "
            "Streitigkeiten, beide aussergerichtlich beigelegt.",
         ])),
        ("Whistleblower-Meldungen 2023",
         "Ueber den Konzern-Whistleblower-Channel (EQS Integrity Line) wurden 3 RHU-bezogene "
         "Meldungen empfangen: 2x Verdacht ueberhoehter Spesenabrechnung (Untersuchung "
         "unbestaetigt), 1x Hinweis auf Sicherheitsmangel (technische Massnahme umgesetzt). "
         "Keine Meldung mit strafrechtlicher Relevanz."),
        ("Schulungen 2023", [
            ["Schulungsthema", "Teilnehmer", "Anteil"],
            ["Code of Conduct (Online)", "538", "99,6 %"],
            ["Anti-Korruption (Praesenz Fuehrungskraefte)", "62", "100 %"],
            ["Datenschutz / DSGVO + NAIH", "540", "100 % (Nach Vorfall)"],
            ["Exportkontrolle", "28", "100 %"],
            ["Kartellrecht", "62", "100 %"],
        ]),
        ("Fazit / Empfehlungen",
         "Die Compliance-Lage der RHU ist insgesamt geordnet. Wesentlicher "
         "Verbesserungsbedarf besteht beim Datenschutz (USB-Stick-Vorfall): Endpoint-Policy "
         "wurde verschaerft, Datenschutz-Schulungen werden 2024 quartalsweise fortgesetzt. "
         "Empfehlung: lokale ISO-27001-Vorbereitung 2024 parallel zur Konzern-Initiative."),
        ("Freigabe",
         signatures("Dr. Eszter Nagy", "Compliance-Koord. RHU", "Brennhagen Hungary Kft.",
                    "Laszlo Kovacs", "Werkleiter RHU", "Brennhagen Hungary Kft.",
                    place="Gyoer", date_str_="29. Januar 2024")),
    ])


# =============================================================================
# 5) RHU IC Quartalsbericht (7)
# =============================================================================
def ic_quartal(fname, jahr, q, umsatz_huf, ebitda_huf, ftes, themen):
    write_doc(f"{BASE}/{fname}", RHU_HDR,
        f"IC-Quartalsbericht {q}/{jahr} - Brennhagen Hungary Kft.",
        subtitle=f"Intercompany-Reporting an Group Controlling (REA Stuttgart)",
        sections=[
            ("Berichtskennzahlen", [
                ["Kennzahl", "Q-Ist", "Vorjahres-Q", "Plan"],
                ["Umsatz (Mio. HUF)", str(umsatz_huf), str(int(umsatz_huf*0.92)), str(int(umsatz_huf*0.98))],
                ["EBITDA (Mio. HUF)", str(ebitda_huf), str(int(ebitda_huf*0.88)), str(int(ebitda_huf*1.05))],
                ["EBITDA-Marge", f"{ebitda_huf/umsatz_huf*100:.1f} %",
                 f"{ebitda_huf*0.88/(umsatz_huf*0.92)*100:.1f} %", "9,5 %"],
                ["Beschaeftigte (FTE)", str(ftes), str(ftes-12), str(ftes+5)],
                ["EUR/HUF Stichtagskurs", "385,2", "365,8", "390,0"],
            ]),
            ("Geschaeftsverlauf",
             f"Die Brennhagen Hungary Kft. (RHU) erzielte im {q}/{jahr} einen Umsatz von "
             f"{huf(umsatz_huf*1_000_000)} (entspricht ca. {umsatz_huf/385.2:.1f} Mio. EUR "
             f"zum Stichtagskurs). Hauptkunden im Sensorik-Bereich blieben BMW (28 %), "
             f"Mercedes (24 %), VW (18 %), Audi (12 %), uebrige Kunden 18 %. "
             f"Die Auslastung der drei Produktionslinien lag bei durchschnittlich 91 %."),
            ("Wesentliche Quartalsthemen",
             ("list", themen)),
            ("Risiken / Hinweise",
             "EUR/HUF-Wechselkursrisiko bleibt erhoeht (Schwankungsbreite +/- 8 % im Quartal). "
             "Hedging-Quote durch Group Treasury (Markus Pflanzer) bei 80 % der EUR-Exposures. "
             "Lieferantenrisiko: ein japanischer Halbleiter-Lieferant unter Beobachtung "
             "(Allokation seit Q1, alternative Quellen identifiziert). "
             "NAV-Aussenpruefung KSt 2020-2022 in Bearbeitung."),
            ("Ausblick",
             f"Erwartete Umsatz-Entwicklung Folgequartal: leichter Anstieg um 3-5 % "
             f"durch Anlauf neuer Mercedes-Programme. EBITDA-Margenziel: > 9 %. "
             f"Investitionen Q+1: rd. 280 Mio. HUF in eine zusaetzliche Pruefstandsstrasse."),
            ("Freigabe / Signatur",
             signatures("Janos Toth", "Finanzcontroller RHU", "Brennhagen Hungary Kft.",
                        "Laszlo Kovacs", "Werkleiter RHU", "Brennhagen Hungary Kft.",
                        place="Gyoer", date_str_=f"15. {['Januar','April','Juli','Oktober'][int(q[1])-1]} {jahr if q!='Q4' else int(jahr)+1}")),
        ])


_ic_quartal_data = [
    ("2019", "Q1", 3_120, 268, 498),
    ("2019", "Q3", 3_240, 281, 504),
    ("2019", "Q4", 3_410, 295, 512),
    ("2020", "Q1", 2_980, 215, 510),
    ("2020", "Q2", 2_640, 188, 502),  # COVID
    ("2020", "Q3", 3_080, 248, 508),
    ("2020", "Q4", 3_320, 286, 518),
]
_ic_quartal_themen = {
    "Q1": ["Hochlauf neuer Drehzahlsensor-Linie BMW", "Lieferantenwechsel ASIC (Infineon -> ST)",
           "IATF 16949 Surveillance-Audit bestanden", "Mehrarbeit im Maerz (Nachholbedarf)"],
    "Q2": ["COVID-19-Massnahmen (Kurzarbeit-Light gem. ungarischem KKV)", "Schutzkonzept Werk",
           "Lieferketten-Stoerungen behoben", "PPAP Audi-Auftrag erfolgreich"],
    "Q3": ["Sommerproduktion stabil, Rekord-Output September",
           "Schulungswoche Q-Personal", "Investitionsfreigabe SMT-Linie 4",
           "Hyundai-Auftrag (Pilotserie) angelaufen"],
    "Q4": ["Jahresendgeschaeft stark, Schichtbetrieb 6+1",
           "Inventur dezember regulaer abgeschlossen", "Bonusausschuettung 13. Monatsgehalt "
           "an alle Mitarbeiter (Q-Ziel uebertroffen)", "Strategie-Klausur Werkleitung 2024-2026"],
}

for j, q, u, e, f in _ic_quartal_data:
    ic_quartal(f"RHU_IC_Quartalsbericht_{j}_{q}.docx", j, q, u, e, f,
               _ic_quartal_themen[q])


# =============================================================================
# 6) IC-Rechnungen (RHU bills concerns / vendors) - 33 Stueck
# =============================================================================
def ic_rechnung(fname, datum, rechnungs_nr, leistung, betrag_huf, kunde="Brennhagen Elektronik AG"):
    is_draft = "DRAFT" in fname.upper()
    is_intern = "intern" in fname.lower() or "WIP" in fname
    write_doc(f"{BASE}/{fname}", RHU_HDR,
        f"Intercompany-Rechnung Nr. {rechnungs_nr}",
        subtitle=f"Brennhagen Hungary Kft. an {kunde}",
        draft=is_draft or is_intern,
        sections=[
            ("Rechnungssteller",
             "Brennhagen Hungary Kft.\n"
             "Ipari Park, Bercsenyi liget 14, H-9027 Gyoer, Magyarorszag\n"
             "Cg.08-09-029876, Adoszam: HU 12876543-2-08\n"
             "Bank: OTP Bank Nyrt., HU92 1177 3111 0098 7654 3210 0987\n"
             "Werkleiter: Laszlo Kovacs"),
            ("Rechnungsempfaenger",
             f"{kunde}\n"
             "Vaihinger Strasse 120, 70567 Stuttgart, Deutschland\n"
             "USt-IdNr.: DE 312 487 901, HRB 726451 AG Stuttgart"),
            ("Rechnungsdaten", [
                ["Feld", "Wert"],
                ["Rechnungsnr.", rechnungs_nr],
                ["Rechnungsdatum", datum],
                ["Leistungszeitraum", datum + " (Monatsleistung)"],
                ["Zahlungsbedingung", "30 Tage netto ab Rechnungsdatum"],
                ["Skonto", "2 % bei Zahlung binnen 10 Tagen"],
                ["Lieferschein-Nr.", "siehe Anlage (Sammellieferschein Monatsende)"],
            ]),
            ("Leistungsbeschreibung",
             leistung),
            ("Rechnungsbetrag", [
                ["Position", "HUF", "EUR (Tageskurs)"],
                ["Nettobetrag", huf(betrag_huf), f"{betrag_huf/385:,.0f} EUR".replace(",", ".")],
                ["Reverse-Charge (EU)", "0 (Art. 196 MwStSysRL - "
                 "Steuerschuldnerschaft Empfaenger)", "0"],
                ["Bruttobetrag", huf(betrag_huf), f"{betrag_huf/385:,.0f} EUR".replace(",", ".")],
            ]),
            ("Hinweise / Transferpreis",
             "Diese Intercompany-Rechnung wird auf Basis des konzernweiten "
             "Cost-Plus-Verrechnungspreismodells gem. Local-File RHU 2022 erstellt. "
             "Marge: Cost+8 % (Routinetaetigkeit Produktion). Dokumentation gem. "
             "ungarischen TP-Regeln (Art. 31. § und 32. § des KSz-Gesetzes) sowie "
             "BEPS-Vorgaben (OECD-Leitlinien 2022). Reverse-Charge Verfahren gem. "
             "Art. 196 EU-MwStSystemRL anwendbar."),
            ("Sachlich/rechnerisch geprueft",
             signatures("Janos Toth", "Finanzcontroller RHU", "Brennhagen Hungary Kft.",
                        "Laszlo Kovacs", "Werkleiter RHU", "Brennhagen Hungary Kft.",
                        place="Gyoer", date_str_=datum)),
        ])


_leistungen = [
    "Lieferung Sensorik-Halbfabrikate (Drehzahlsensoren BMW-Programm) gemaess Sammellieferschein. "
    "Mengen, Artikel-Nr. und Pruefzeugnisse im Anhang ZSB.",
    "Produktion und Lieferung Beschleunigungssensoren BES-5500-Familie an REG Heilbronn "
    "fuer Mercedes EQS-Plattform. EMV-Pruefzeugnisse beigefuegt.",
    "Drucksensoren DRS-3300 (1.000 Stueck) fuer VW-MEB-Anwendung. Konformitaetsbescheinigung "
    "gem. PPAP Level 3 beigefuegt.",
    "Lohnfertigung SMT-Bestueckung Steuergeraete-Leiterplatten REA-Auftrag SC-2240. "
    "Stueckzahl, Loesescheine und QPR siehe Anlage.",
    "Engineering Services Sensorik-Optimierung (Werkstech. Beratung 184 h Senior-Tarif).",
]

_invoice_months = []
for year in [2020, 2021, 2022]:
    for m in range(1, 13):
        _invoice_months.append((year, m))

# Map invoices in folder to data
_invoice_files = [
    # (filename, year, month, base_huf, leistung_idx, kunde)
    ("RHU_IC_Rechnung_2020_01.docx", 2020, 1, 142_500_000, 0, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2020_02.docx", 2020, 2, 138_200_000, 1, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2020_03.docx", 2020, 3, 145_700_000, 2, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2020_04.docx", 2020, 4, 98_400_000, 0, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2020_05.docx", 2020, 5, 102_300_000, 1, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2020_06.docx", 2020, 6, 124_800_000, 2, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2020_08_intern.docx", 2020, 8, 87_200_000, 4, "Brennhagen Holding GmbH"),
    ("RHU_IC_Rechnung_2020_09.docx", 2020, 9, 156_300_000, 0, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2020_10.docx", 2020, 10, 162_500_000, 1, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2020_12_WIP.docx", 2020, 12, 174_800_000, 2, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2021_01.docx", 2021, 1, 152_300_000, 0, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2021_02.docx", 2021, 2, 148_900_000, 1, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2021_03.docx", 2021, 3, 165_700_000, 2, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2021_04.docx", 2021, 4, 158_200_000, 0, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2021_05.docx", 2021, 5, 162_400_000, 1, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2021_06.docx", 2021, 6, 171_300_000, 2, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2021_07_DRAFT.docx", 2021, 7, 168_500_000, 3, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2021_08.docx", 2021, 8, 152_600_000, 0, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2021_09.docx", 2021, 9, 178_400_000, 1, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2021_10.docx", 2021, 10, 184_700_000, 2, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2021_11.docx", 2021, 11, 192_300_000, 0, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2021_12.docx", 2021, 12, 198_400_000, 1, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2022_01_intern.docx", 2022, 1, 92_300_000, 4, "Brennhagen Holding GmbH"),
    ("RHU_IC_Rechnung_2022_02.docx", 2022, 2, 187_900_000, 2, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2022_03.docx", 2022, 3, 198_500_000, 0, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2022_04.docx", 2022, 4, 192_400_000, 1, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2022_05.docx", 2022, 5, 204_700_000, 2, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2022_06.docx", 2022, 6, 212_300_000, 0, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2022_07.docx", 2022, 7, 197_800_000, 1, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2022_08.docx", 2022, 8, 184_200_000, 2, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2022_09.docx", 2022, 9, 218_700_000, 0, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2022_10.docx", 2022, 10, 224_800_000, 1, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2022_11.docx", 2022, 11, 231_200_000, 2, "Brennhagen Elektronik GmbH"),
    ("RHU_IC_Rechnung_2022_12.docx", 2022, 12, 238_400_000, 0, "Brennhagen Elektronik GmbH"),
]

for fn, yr, mn, amt, li, kd in _invoice_files:
    datum = f"{28}. {['Januar','Februar','Maerz','April','Mai','Juni','Juli','August','September','Oktober','November','Dezember'][mn-1]} {yr}"
    rn = f"RHU-IC-{yr}-{mn:02d}-001"
    ic_rechnung(fn, datum, rn, _leistungen[li], amt, kd)


# =============================================================================
# 7) RHU Mietvertrag
# =============================================================================
write_doc(f"{BASE}/RHU_Mietvertrag_Betriebsgelaende_2020.docx", RHU_HDR,
    "Mietvertrag Betriebsgelaende Gyoeri Ipari Park",
    subtitle="Brennhagen Hungary Kft. - Werks- und Verwaltungsgebaeude (Sensorik-Produktion)",
    sections=[
        ("Vertragsparteien",
         "Vermieter: Gyoeri Ipari Park Kft., Bercsenyi liget 12, H-9027 Gyoer, "
         "Cg.08-09-012345, vertreten durch Geschaeftsfuehrer Istvan Varga.\n\n"
         "Mieter: Brennhagen Hungary Kft., Bercsenyi liget 14, H-9027 Gyoer, "
         "Cg.08-09-029876, vertreten durch Laszlo Kovacs (uegyvezeto)."),
        ("Mietgegenstand", ("clauses", [
            ("1. §  Mietgegenstand",
             ["Werks- und Verwaltungsgebaeude im Industriepark Gyoer (Parzelle 14): "
              "ca. 18.400 qm Produktionsflaeche, 2.200 qm Verwaltungsflaeche, "
              "1.800 qm Lagerflaeche, plus 4.500 qm Parkflaeche.",
              "Industrielle Ausstattung gem. Anlage A (300A-Drehstromversorgung, "
              "Druckluft 8 bar, Klima-/Reinraumzonen ISO 8 fuer SMT-Linien, "
              "Sprinkleranlage, Hallenboden FM 5 Belastbarkeit 12 t/qm)."]),
            ("2. §  Mietzins",
             ["Grundmiete: 78.500.000 HUF p. a. (ca. 200.000 EUR p. a. zum mittleren "
              "EUR/HUF-Kurs), zahlbar quartalsweise im Voraus jeweils zum 15. des "
              "Quartalsbeginns.",
              "Nebenkostenvorauszahlung: 12.000.000 HUF p. a., Abrechnung jaehrlich "
              "bis 30.4. des Folgejahres.",
              "Wertsicherung: Anpassung jaehrlich um den ungarischen "
              "Verbraucherpreisindex (KSH-VPI), gedeckelt auf 5 % p. a."]),
            ("3. §  Laufzeit",
             ["Laufzeit: 1.1.2020 - 31.12.2034 (15 Jahre fest). Verlaengerungsoption "
              "Mieter 2 x 5 Jahre zu marktueblichen Konditionen, Anzeige 24 Monate "
              "vor Vertragsende."]),
            ("4. §  Verwendungszweck",
             ["Nutzung als Produktions- und Verwaltungsstaette fuer Automobil-Sensorik "
              "(Drehzahl-, Beschleunigungs- und Drucksensoren). Nutzungsaenderung "
              "bedarf der schriftlichen Zustimmung des Vermieters."]),
            ("5. §  Untervermietung",
             ["Untervermietung an verbundene Unternehmen der Brennhagen-Gruppe (insb. REG, RHO) "
              "ist ohne weitere Zustimmung gestattet; sonstige Untervermietung erfordert "
              "schriftliche Zustimmung des Vermieters (nicht unbillig zu verweigern)."]),
            ("6. §  Erhaltungspflichten",
             ["Schoenheitsreparaturen und Instandhaltung des Innenbereichs obliegen dem Mieter. "
              "Dach, Aussenwaende und tragende Bauteile sowie der industrielle Grundausbau "
              "verbleiben in der Verantwortung des Vermieters."]),
            ("7. §  Anwendbares Recht",
             ["Es gilt ungarisches Recht (insb. Ptk. 6:331. §). Gerichtsstand: "
              "Gyoeri Toervenyszek. Verhandlungssprache: Ungarisch / Deutsch."]),
        ])),
        ("Unterschrift",
         signatures("Istvan Varga", "GF Gyoeri Ipari Park Kft.", "Vermieter",
                    "Laszlo Kovacs", "uegyvezeto", "Brennhagen Hungary Kft.",
                    place="Gyoer", date_str_="18. Dezember 2019")),
    ])


# =============================================================================
# 8) RHU Steuerbescheid (4) - NAV
# =============================================================================
def steuerbescheid(fname, jahr, kst_huf, ust_saldo_huf, gewerbe_huf, bescheidnr,
                   anmerkungen):
    write_doc(f"{BASE}/{fname}", RHU_HDR,
        f"NAV-Steuerbescheid (Adobehoerdei Hatarozat) {jahr}",
        subtitle=f"Brennhagen Hungary Kft. - Koerperschaftsteuer, USt-Schlussbescheid, "
                 f"oertliche Gewerbesteuer (HIPA)",
        sections=[
            ("Bescheid-Daten", [
                ["Feld", "Wert"],
                ["Bescheid-Nr.", bescheidnr],
                ["Steuerjahr", str(jahr)],
                ["Steuerpflichtiger", "Brennhagen Hungary Kft., Gyoer"],
                ["Adoszam", "HU 12876543-2-08"],
                ["Festsetzende Behoerde", "NAV Gyoer-Moson-Sopron Megyei "
                 "Adoes Vamigazgatosaga"],
                ["Bescheid-Datum", f"15. Maerz {jahr+1}"],
                ["Rechtsmittelfrist", "30 Tage (gem. Art-tv. § 124)"],
            ]),
            ("Festsetzung",
             "Auf Grundlage der eingereichten Steuererklaerung der Brennhagen Hungary Kft. "
             f"fuer das Geschaeftsjahr {jahr} und der nachgereichten Unterlagen erfolgt "
             "folgende Festsetzung:"),
            ("Festsetzungsbetraege", [
                ["Steuerart", "Bemessungsgrundlage (HUF)", "Steuersatz", "Festsetzung (HUF)"],
                ["Tarsasagi Ado (KSt)", huf(kst_huf*100//9), "9,0 %", huf(kst_huf)],
                ["Helyi Iparuezesi Ado (HIPA - Gyoer)", huf(gewerbe_huf*100//2), "2,0 %", huf(gewerbe_huf)],
                ["Innovacios Jarulek", huf(gewerbe_huf*100//2*30//100), "0,3 %", huf(gewerbe_huf*30//100)],
                ["USt-Schlussbescheid (AFA)", "siehe Anlage", "0/5/27 %", huf(ust_saldo_huf)],
            ]),
            ("Begruendung / Anmerkungen",
             anmerkungen),
            ("NAV-Aussenpruefung",
             "Hinweis: fuer dieses Steuerjahr wurde keine Aussenpruefung durch die NAV "
             "ausgeloest. Die Pruefungsfristen gem. Art-tv. enden 5 Jahre nach Bescheid; "
             "die regulaere Aussenpruefung KSt 2020-2022 erfolgte 2023 ohne Beanstandung."),
            ("Rechtsbehelf",
             "Gegen diesen Bescheid kann binnen 30 Tagen nach Zustellung ein Einspruch "
             "(fellebbezes) bei der NAV Gyoer-Moson-Sopron eingelegt werden. Eine "
             "Klage vor dem Verwaltungsgericht (kozigazgatasi birosag) ist nach "
             "abgewiesener Beschwerde innerhalb von 30 Tagen moeglich."),
            ("Bestaetigung",
             signatures("Dr. Anna Kovacsne", "Adoellenor / NAV-Pruefer", "NAV Gyoer",
                        "Janos Toth", "Finanzcontroller (Empfang)", "Brennhagen Hungary Kft.",
                        place="Gyoer", date_str_=f"15. Maerz {jahr+1}")),
        ])


steuerbescheid("RHU_Steuerbescheid_2020.docx", 2020,
    214_000_000, -8_400_000, 84_000_000, "NAV-08-2020/RHU-04412",
    "Im Steuerjahr 2020 wurden COVID-19-bezogene Verlustverrechnungen geltend gemacht "
    "(Mehraufwand Schutzkonzept 124 Mio. HUF). Die KSt-Vorauszahlungen 2020 von "
    "240 Mio. HUF werden mit 214 Mio. HUF Jahresschuld saldiert; Erstattung von "
    "26 Mio. HUF erfolgt binnen 60 Tagen auf das hinterlegte OTP-Konto.")
steuerbescheid("RHU_Steuerbescheid_2021.docx", 2021,
    287_000_000, 4_200_000, 112_000_000, "NAV-08-2021/RHU-05218",
    "Erholung nach COVID, deutlicher Anstieg der Umsaetze und Margen. Innovacios Jarulek "
    "auf F&E-Aufwendungen fuer Sensorik-Entwicklung anteilig abzugsfaehig. USt-Saldo "
    "geringfuegig positiv durch erhoehten Exportanteil.")
steuerbescheid("RHU_Steuerbescheid_2022.docx", 2022,
    342_000_000, 12_800_000, 138_000_000, "NAV-08-2022/RHU-06034",
    "Starkes Umsatzwachstum durch Anlauf BMW-Programm. F&E-Steuerguthaben gem. "
    "Art. 7. § (k) Tao-tv. (KSt-Gesetz) in Hoehe von 18 Mio. HUF anerkannt (Sensorik-"
    "Entwicklungsprojekt 'BES-5500 NextGen').")
steuerbescheid("RHU_Steuerbescheid_KSt_2023.docx", 2023,
    398_000_000, 18_400_000, 162_000_000, "NAV-08-2023/RHU-06875",
    "Rekordergebnis 2023. NAV-Pruefung der Jahre 2020-2022 wurde im November 2023 ohne "
    "Beanstandung abgeschlossen (Bestaetigung NAV-Pruefungsprotokoll 2023/11/082). "
    "Verrechnungspreis-Dokumentation (TP-Local-File RHU 2023) wurde im Rahmen der "
    "Aussenpruefung anerkannt.")


# =============================================================================
# 9) RHU Versicherungsnachweis
# =============================================================================
write_doc(f"{BASE}/RHU_Versicherungsnachweis_2023.docx", RHU_HDR,
    "Versicherungsnachweis 2023 - Brennhagen Hungary Kft.",
    subtitle="Konzernversicherungsdeckung mit lokalen Ergaenzungen",
    sections=[
        ("Versicherungs-Uebersicht", [
            ["Sparte", "Versicherer", "Versicherungssumme", "Anmerkung"],
            ["Property-All-Risk", "Allianz Hungaria Zrt. (Lokalpolice + AGCS-Rueckdeckung)",
             "12.500 Mio. HUF (rd. 32 Mio. EUR)", "Werk Gyoer, Vorraete, BU 18 Monate"],
            ["Betriebshaftpflicht", "Allianz Global Corporate & Specialty SE (AGCS, Konzernpolice)",
             "50 Mio. EUR", "Konzern-Master-Policy, RHU als versicherte Tochter"],
            ["Produkthaftpflicht", "AIG Europe SA, Niederlassung Ungarn",
             "30 Mio. EUR", "Worldwide-Coverage inkl. USA Cap 10 Mio. EUR"],
            ["D&O (Manager-Haftpflicht)", "Allianz Global Corporate & Specialty SE",
             "50 Mio. EUR", "Konzernpolice Brennhagen-Gruppe (alle Tochter-GF inkludiert)"],
            ["Cyber-Versicherung", "Munich Re / HDI Cyber+", "15 Mio. EUR",
             "Konzernpolice mit Sublimit RHU 5 Mio. EUR"],
            ["Transport (Inland HU + EU)", "Generali Biztosito Zrt.",
             "200 Mio. HUF / Sendung", "JIT-Sendungen an Bosch/BMW abgesichert"],
            ["Kfz-Flotte (24 Fzg.)", "Allianz Hungaria Zrt.",
             "Vollkasko / Haftpflicht", "Werkverkehr, Aussendienst, GF-Dienstwagen"],
            ["Gesetzliche TB-Versicherung (NAV)", "NAV / Egeszsegbiztositasi Alap",
             "gesetzlich", "Arbeitgeber-Beitrag 13 % Szocho gem. Tbj."],
        ]),
        ("Konzernweite Bedingungen",
         "Die Konzernpolicen (AGCS, AIG, Munich Re) werden zentral durch das Group "
         "Treasury (Markus Pflanzer) und Group Risk Management (CRO-Office) verhandelt "
         "und administriert. Lokale Schadensmeldungen sind ueber das Schadensportal "
         "der jeweiligen Versicherung sowie parallel an Group Risk (rim@brennhagen-elektronik.de) "
         "zu melden.\n\nSelbstbehalte: Sachschaden 50.000 EUR / Schadensereignis (Konzern); "
         "Haftpflicht 25.000 EUR (Konzern); Cyber 100.000 EUR / Vorfall."),
        ("Lokale Versicherungsvermittler",
         "Lokaler Broker fuer RHU: Aon Magyarorszag Kft., Budapest, Ansprechperson "
         "Dr. Peter Szabo. Konzernbroker (global): Marsh McLennan."),
        ("Bestaetigung",
         signatures("Janos Toth", "Finanzcontroller RHU", "Brennhagen Hungary Kft.",
                    "Markus Pflanzer", "Group Treasurer", R["name"],
                    place="Gyoer / Stuttgart", date_str_="22. Januar 2023")),
    ])


# =============================================================================
# 10) WP-Management-Letter (3)
# =============================================================================
def wp_letter(fname, jahr, findings):
    write_doc(f"{BASE}/{fname}", RHU_HDR,
        f"Management Letter {jahr} - KPMG (RHU Pruefungsbericht-Begleitschreiben)",
        subtitle=f"Pruefer KPMG Hungaria Kft., Budapest - Local Audit Partner Dr. Tamas Hegedus",
        confidential=True,
        sections=[
            ("Adressat / Berichtsumfang",
             "An die Geschaeftsleitung der Brennhagen Hungary Kft. (Laszlo Kovacs) und "
             "an das Group Audit Office der Brennhagen Elektronik AG, Stuttgart (CFO Laura Bauer, "
             "CAE Andreas Buehler, Konzernabschlusspruefer KPMG AG WPG / Dr. Maximilian Brand).\n\n"
             f"Berichtsumfang: Audit-Findings des Jahresabschlusses {jahr} der RHU "
             "(IFRS-Reporting an Konzern + ungarische Lokalbilanz HAS - Hungarian "
             "Accounting Standards)."),
            ("Pruefungsumfang",
             f"Wir haben den Jahresabschluss der Brennhagen Hungary Kft. fuer das "
             f"Geschaeftsjahr {jahr} nach International Standards on Auditing (ISA) "
             f"sowie ungarischem Wirtschaftspruefungsgesetz (Szt. 2000/C) geprueft. "
             f"Pruefungstage: 42, davon 28 vor Ort im Werk Gyoer."),
            ("Wesentliche Findings", ("list", findings)),
            ("Bewertung",
             "Insgesamt bewerten wir das interne Kontrollsystem (IKS) der RHU als "
             "ANGEMESSEN. Keine wesentlichen Maengel im Sinne von ISA 265 wurden "
             "festgestellt. Die Berichterstattung an den Konzern (Group Reporting "
             "Package via SAP S/4HANA) ist zeitgerecht und qualitativ konsistent."),
            ("Empfehlungen",
             "Wir empfehlen, die identifizierten Findings im laufenden Geschaeftsjahr "
             "zu bearbeiten und uns im Vor-Audit-Meeting zur naechsten Jahrespruefung "
             "den Umsetzungsstand schriftlich zu bestaetigen."),
            ("Kontaktdaten KPMG",
             "KPMG Hungaria Kft., Vaci ut 31, H-1134 Budapest\n"
             "Local Audit Partner: Dr. Tamas Hegedus (tamas.hegedus@kpmg.hu)\n"
             "Local Senior Manager: Eva Toth\n\n"
             "Konzernpruefer: KPMG AG WPG, Theodor-Heuss-Strasse 5, 70174 Stuttgart, "
             "Lead Partner Dr. Maximilian Brand."),
            ("Unterschrift",
             signatures("Dr. Tamas Hegedus", "Audit Partner", "KPMG Hungaria Kft.",
                        "Eva Toth", "Senior Manager", "KPMG Hungaria Kft.",
                        place="Budapest", date_str_=f"31. Maerz {int(jahr)+1}")),
        ])


wp_letter("RHU_WP_Management_Letter_2021.docx", "2021", [
    "Bestandsbewertung: Methodischer Hinweis zur Bewertung der Slow-Moving-Items "
    "(>180 Tage). Empfehlung: pauschale Wertberichtigung 25 % statt bisher 15 %.",
    "Verrechnungspreis-Dokumentation TP-Local-File 2021 noch nicht final. "
    "Empfehlung: Abschluss bis 31.5.2022 (NAV-Frist).",
    "IT-Allgemeinkontrollen SAP-Berechtigungen: 4 SoD-Konflikte identifiziert "
    "(Lager + FiBu). Korrektur bereits initiiert.",
    "Konzern-Reporting Package: Reconciliation HAS -> IFRS war 1 Tag verspaetet. "
    "Empfehlung: Prozessbeschleunigung Monatsabschluss WD+3.",
])
wp_letter("RHU_WP_Management_Letter_2022.docx", "2022", [
    "Folgekontrolle Vorjahres-Findings: 3 von 4 umgesetzt; Slow-Moving Wertberichtigung "
    "auf 25 % erhoeht (Umsetzung bestaetigt).",
    "Neueintragungen Sensorik-Linie 4 (BES-5500): Anlagenzugaenge gem. IAS 16 sachgerecht "
    "aktiviert. Inbetriebnahme-Datum 12.9.2022 dokumentiert.",
    "Konzern-Goodwill keine Hinweise auf Wertminderung (Buchwert RHU CGU 68 Mio. EUR).",
    "DSGVO-Risikoinventur: empfehlen, HR-Prozesse staerker zu standardisieren (Hinweis "
    "auf USB-Stick-Praxis vor Ort - Empfehlung Endpoint-Policy).",
])
wp_letter("RHU_WP_Management_Letter_2023.docx", "2023", [
    "Datenschutz-Vorfall 2023-011: Bestaetigung der Wirksamkeit der Sofortmassnahmen "
    "(Endpoint-Policy, HR-Schulungen). Bussgeldverfahren NAIH offen, Rueckstellung "
    "8 Mio. HUF angemessen.",
    "TP-Dokumentation 2023: vollstaendig und konsistent mit Konzern-Master-File. "
    "NAV-Aussenpruefung 2020-2022 ohne Beanstandung bestaetigt (Pruefungsprotokoll "
    "NAV 2023/11/082).",
    "IFRS-Reporting Package: keine Korrekturbuchungen erforderlich. SAP S/4HANA "
    "Migrationsphase im Konzern fortgeschritten (RHU Go-Live geplant Q3/2024).",
    "Forderungsbestand 31.12.2023: Wertberichtigung 1,8 % angemessen, keine "
    "Ueberfaelligkeit > 90 Tage bei OEMs.",
])


# =============================================================================
# 11) RHU_to_REG_IC (37) - Intercompany invoices to REG Heilbronn
# =============================================================================
def rhu_to_reg(fname, jahr, monat):
    rn = f"RHU-REG-{jahr}-{monat:02d}-001"
    datum = f"28. {['Januar','Februar','Maerz','April','Mai','Juni','Juli','August','September','Oktober','November','Dezember'][monat-1]} {jahr}"
    # rotating mix of services
    base_amt = {2021: 168, 2022: 198, 2023: 224}[jahr]
    amt_huf = (base_amt + (monat % 4)*8) * 1_000_000
    write_doc(f"{BASE}/{fname}", RHU_HDR,
        f"Intercompany-Rechnung RHU -> REG Nr. {rn}",
        subtitle="Brennhagen Hungary Kft. an Brennhagen Elektronik GmbH, Heilbronn",
        sections=[
            ("Rechnungssteller / Rechnungsempfaenger",
             "Rechnungssteller: Brennhagen Hungary Kft.\n"
             "Ipari Park, Bercsenyi liget 14, H-9027 Gyoer\n"
             "Cg.08-09-029876, Adoszam HU 12876543-2-08\n"
             "Bank: OTP Bank Nyrt., HU92 1177 3111 0098 7654 3210 0987\n\n"
             "Rechnungsempfaenger: Brennhagen Elektronik GmbH\n"
             "Salzstrasse 124, 74076 Heilbronn, HRB 221456 AG Heilbronn\n"
             "USt-IdNr.: DE 312 487 850"),
            ("Rechnungsdaten", [
                ["Feld", "Wert"],
                ["Rechnungs-Nr.", rn],
                ["Rechnungsdatum", datum],
                ["Leistungszeitraum", f"01.{monat:02d}. - 30.{monat:02d}.{jahr}"],
                ["Bestell-Nr. REG", f"REG-PO-{jahr}-{monat:02d}-RHU-001"],
                ["Zahlungsbedingung", "30 Tage netto"],
            ]),
            ("Leistungsumfang",
             f"Lieferung Halbfabrikate und Komponenten Sensorik gemaess Konzern-Lieferprogramm "
             f"{jahr} (Drehzahlsensoren, Beschleunigungssensoren und Drucksensoren) fuer "
             f"die Endmontage am Werk Heilbronn (REG) in den Programmen ICP-3 (Infotainment) "
             f"und BMS-12 (Batteriemanagement). Sammellieferschein {monat:02d}/{jahr} liegt "
             f"als Anlage 1 bei. Pruefzeugnisse und PPAP-Konformitaetsbestaetigungen je "
             f"Charge auf Konzern-QM-Portal hinterlegt."),
            ("Rechnungsbetrag", [
                ["Position", "Betrag"],
                ["Nettobetrag", huf(amt_huf)],
                ["Reverse-Charge (Art. 196 EU-MwStSystemRL)", "0 HUF"],
                ["Gesamtbetrag (zu zahlen)", huf(amt_huf)],
                ["Zur Information: EUR-Aequivalent (EZB-Stichtagskurs)",
                 f"{amt_huf/385:,.0f} EUR".replace(",", ".")],
            ]),
            ("Verrechnungspreis-Hinweis",
             "Diese Intercompany-Lieferung erfolgt zu konzerneinheitlichen Verrechnungspreisen "
             "nach dem Cost-Plus-Verfahren (Cost+8 % auf vollkostige Produktion). Dokumentation "
             "im TP-Local-File RHU sowie im Konzern-Master-File hinterlegt; Konformitaet mit "
             "OECD-Leitlinien (2022) und ungarischem KSz-Gesetz (Art. 31-32. §) sichergestellt."),
            ("Sachlich/rechnerisch geprueft",
             signatures("Janos Toth", "Finanzcontroller RHU", "Brennhagen Hungary Kft.",
                        "Andreas Maier", "Werkleiter REG", "Brennhagen Elektronik GmbH",
                        place="Gyoer / Heilbronn", date_str_=datum)),
        ])


_rhu_to_reg = []
for y in [2021, 2022]:
    for m in range(1, 13):
        _rhu_to_reg.append((y, m))
for m in range(1, 13):
    if m != 11:  # 2023_11 not in file list
        _rhu_to_reg.append((2023, m))

for y, m in _rhu_to_reg:
    fn = f"RHU_to_REG_IC_{y}_{m:02d}.docx"
    rhu_to_reg(fn, y, m)


# =============================================================================
# 12) RPL_IC_Rechnung_2020_01 (misplaced from Polish folder)
# =============================================================================
RPL_HDR = {
    "name": "Brennhagen Polska Sp. z o.o.",
    "addr": "ul. Mikolowska 100, 40-065 Katowice, Polska",
    "hrb":  "KRS 0000543210, NIP PL 6342876541 (Krajowy Rejestr Sadowy)",
}
write_doc(f"{BASE}/RPL_IC_Rechnung_2020_01.docx", RPL_HDR,
    "Intercompany-Rechnung Nr. RPL-IC-2020-01-001",
    subtitle="Brennhagen Polska Sp. z o.o. an Brennhagen Elektronik GmbH (Heilbronn)",
    sections=[
        ("Rechnungssteller",
         "Brennhagen Polska Sp. z o.o., ul. Mikolowska 100, 40-065 Katowice\n"
         "KRS 0000543210, NIP PL 6342876541, REGON 243210987\n"
         "Bank: PKO BP S.A., PL98 1020 2892 0000 5102 0123 4567\n"
         "Werkleiter: Marek Wojciechowski"),
        ("Rechnungsempfaenger",
         "Brennhagen Elektronik GmbH, Salzstrasse 124, 74076 Heilbronn\n"
         "HRB 221456 AG Heilbronn, USt-IdNr. DE 312 487 850"),
        ("Rechnungsdaten", [
            ["Feld", "Wert"],
            ["Rechnungs-Nr.", "RPL-IC-2020-01-001"],
            ["Datum", "31. Januar 2020"],
            ["Leistungszeitraum", "01.01. - 31.01.2020"],
            ["Bestell-Nr.", "REG-PO-2020-01-RPL-001"],
            ["Zahlungsbedingung", "30 Tage netto"],
        ]),
        ("Leistungsbeschreibung",
         "EMS-/SMD-Bestueckungsleistung Januar 2020 fuer ICP-3-Plattform (REG Heilbronn). "
         "Auftragsvolumen 84.000 Boards inkl. AOI-Endpruefung und Vorbereitung Transport "
         "(JIT-Anlieferung Heilbronn)."),
        ("Rechnungsbetrag", [
            ["Position", "Netto", "VAT", "Brutto"],
            ["EMS-Leistung Januar 2020", "412.300,00 PLN", "0 (Reverse-Charge)", "412.300,00 PLN"],
            ["Verpackung / Transport", "18.400,00 PLN", "0", "18.400,00 PLN"],
            ["Gesamt", "430.700,00 PLN", "0", "430.700,00 PLN"],
        ]),
        ("Verrechnungspreis / Hinweise",
         "Cost-Plus 8 % gem. Master-File 2019. Polnische TP-Dokumentation (Local File) im "
         "Personalakt der RPL-Buchhaltung. Reverse-Charge Verfahren (EU-Lieferung mit "
         "Endmontage in DE)."),
        ("Sachlich/rechnerisch geprueft",
         signatures("Anna Kowalska", "HR / Finance RPL", "Brennhagen Polska Sp. z o.o.",
                    "Marek Wojciechowski", "Werkleiter RPL", "Brennhagen Polska Sp. z o.o.",
                    place="Katowice", date_str_="31. Januar 2020")),
    ])


# =============================================================================
# 13) RSG_BR_Protokoll (misplaced)
# =============================================================================
write_doc(f"{BASE}/RSG_BR_Protokoll_2020_05.docx",
    {"name": "Brennhagen Software GmbH",
     "addr": "Leopoldstrasse 184, 80804 Muenchen",
     "hrb":  "HRB 319872, AG Muenchen"},
    "BR-Sitzungsprotokoll Brennhagen Software GmbH - Mai 2020",
    subtitle="Ordentliche Sitzung des Betriebsrats RSG",
    sections=[
        ("Sitzungsdaten",
         "Datum: 14. Mai 2020, 13:30 - 16:00 Uhr\n"
         "Ort: Microsoft Teams (COVID-19-bedingt, alle Mitglieder remote)\n"
         "Vorsitz: Dipl.-Inf. Wolfgang Felser\n"
         "Stv. Vorsitz / Schriftfuehrung: Dr. Petra Henkel\n"
         "Anwesend: 7 von 7 Mitgliedern - Sitzung beschlussfaehig"),
        ("Tagesordnung", ("list", [
            "TOP 1 - Protokoll-Genehmigung Vorsitzung (April 2020)",
            "TOP 2 - COVID-19-Lagebericht und Schutzkonzept",
            "TOP 3 - Mobile Arbeit / Home-Office-Regelung dauerhaft (BV vorgeschlagen)",
            "TOP 4 - Personalentwicklung 2020 Q2/Q3 (insb. ASPICE-Schulungen)",
            "TOP 5 - Bonusplan 2020 - Auswirkung COVID",
            "TOP 6 - Verschiedenes",
        ])),
        ("TOP 2 - COVID-19-Lagebericht",
         "Werkleiter Dr. Klaus Kessler stellt die aktuelle Lage dar: 92 % der Mitarbeiter "
         "im Home-Office, Praesenzpflicht nur fuer 28 Mitarbeiter mit Labortaetigkeit "
         "(EMV-Kammer, HiL-Pruefstand). Schutzkonzept umgesetzt. Keine Infektionsfaelle "
         "bislang. Der BR begruesst die zuegige Umsetzung und schlaegt einen "
         "psychosozialen Beratungspool fuer dauerhaft remote arbeitende Mitarbeiter vor."),
        ("TOP 3 - Mobile Arbeit dauerhaft",
         "Der Vorstand (Anna Mueller) hat dem RSG-Werkleiter mitgeteilt, dass eine "
         "konzernweite Betriebsvereinbarung zur dauerhaften mobilen Arbeit erarbeitet "
         "wird (Federfuehrung KBR/Stuttgart). RSG-Variante: 80 % mobile Arbeit moeglich "
         "fuer Software-Entwicklung. Der BR stimmt zu, fordert aber explizite "
         "Anrechnung von Pausen und Schutz vor staendiger Erreichbarkeit (Recht auf "
         "Nicht-Erreichbarkeit ausserhalb Kernzeiten)."),
        ("TOP 5 - Bonusplan 2020 - COVID",
         "Der Vorstand schlaegt vor, den Bonusplan 2020 unveraendert beizubehalten, "
         "aber bei verschlechterter Konzern-EBIT-Lage einen Floor von 50 % der "
         "Zielboni zu garantieren. Der BR stimmt dem Vorschlag mehrheitlich zu."),
        ("TOP 6 - Verschiedenes",
         "Naechste Sitzung 11. Juni 2020 (Format wird je nach Pandemielage angepasst). "
         "Konzern-BR-Klausur Herbst 2020 vorlaeufig in Stuttgart geplant."),
        ("Unterschrift",
         signatures("Wolfgang Felser", "Vorsitzender BR RSG", "Brennhagen Software GmbH",
                    "Dr. Petra Henkel", "Stv. / Schriftfuehrung", "Brennhagen Software GmbH",
                    place="Muenchen (remote)", date_str_="14. Mai 2020")),
    ])


# =============================================================================
# Verification
# =============================================================================
if __name__ == "__main__":
    from docx import Document
    from pathlib import Path
    base = Path(BASE)
    total = thin = 0
    thin_names = []
    for p in sorted(base.rglob("*.docx")):
        d = Document(p)
        t = " ".join(par.text for par in d.paragraphs)
        for tbl in d.tables:
            for r in tbl.rows:
                for c in r.cells:
                    t += " " + c.text
        w = len(t.split())
        total += 1
        if w < 200:
            thin += 1
            thin_names.append((p.name, w))
    print(f"{total} total, {thin} still thin")
    for n, w in thin_names:
        print(f"  THIN: {n} ({w} words)")
