"""
enhance_lib.py
Shared helpers + canonical facts for re-generating dd-mockdata documents
with realistic length and consistent cross-references.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── CANONICAL FACTS ──────────────────────────────────────────────────────────

MUELLER = {
    "name":  "Halbreiter Maschinenbau GmbH",
    "short": "MMB",
    "addr":  "Industriestraße 12, 50829 Köln",
    "hrb":   "HRB 47312, Amtsgericht Köln",
    "ust":   "DE 198 765 432",
    "ceo":   "Klaus Müller",
    "cfo":   "Sandra Becker",
    "geschaeftszweck": "Entwicklung, Herstellung und Vertrieb von Sondermaschinen und Anlagen für die metallverarbeitende Industrie",
    "kunden":      ["ThyssenKrupp Steel Europe AG", "Bosch Rexroth AG", "Hella GmbH & Co. KGaA",
                    "Viessmann Climate Solutions SE", "Dürr AG"],
    "lieferanten": ["Schunk GmbH & Co. KG", "Igus GmbH", "Siemens AG – Antriebstechnik",
                    "Trumpf SE + Co. KG"],
    "produkte":    ["Pressenlinie PL-500 (hydraulische Stanzpresse)",
                    "Förderbandanlage FB-200 (modulares Transportsystem)",
                    "Laserschneidanlage LS-800 (5-Achs-CNC)",
                    "Montageroboter MR-150 (Kollaborationsroboter)"],
    "bank":  "Deutsche Bank AG, Filiale Köln",
    "iban":  "DE89 3007 0010 0123 4567 89",
    "wp":    "BDO AG Wirtschaftsprüfungsgesellschaft",
    "anwalt": "Heuking Kühn Lüer Wojtek Partnerschaft mbB",
    "employees_2023": 247,
    "revenue_2023":   48_630_000,
    "ebit_2023":      3_890_000,
    "ebitda_2023":    5_980_000,
}

BIOTECH = {
    "name":  "Sentavia Precision GmbH",
    "short": "BTP",
    "addr":  "Freimannstraße 45, 80939 München",
    "hrb":   "HRB 218934, Amtsgericht München",
    "ust":   "DE 287 432 109",
    "ceo":   "Dr. Katharina Berger",
    "cto":   "Dr. Marcus Vogt",
    "cfo":   "Thomas Müller",
    "cmo":   "Dr. Annika Schmidt",
    "qra":   "Dipl.-Ing. Stefan Hoffmann",
    "produkte": [
        ("Cardevio Pro",      "Klasse IIb (MDR Anhang VIII, Regel 10)", "MDR-UDI-DE/BTP/CS-PRO-001"),
        ("Ostevo Navigator",  "Klasse IIa (MDR Anhang VIII, Regel 7)",  "MDR-UDI-DE/BTP/OF-NAV-002"),
        ("Veridiq SARS-Flex", "IVD Klasse C (IVDR Anhang VIII, Regel 3)", "IVDR-UDI-DE/BTP/DK-SF-003"),
    ],
    "nb":     "TÜV SÜD Product Service GmbH",
    "nb_id":  "0123",
    "iso":    "DIN EN ISO 13485:2016",
    "distributoren": ["Siemens Healthineers AG", "Fresenius Medical Care AG & Co. KGaA", "B. Braun SE"],
    "kliniken":      ["Charité – Universitätsmedizin Berlin",
                      "Universitätsklinikum Hamburg-Eppendorf (UKE)",
                      "LMU Klinikum München", "Universitätsklinikum Heidelberg"],
    "lieferanten":   ["Texas Instruments Deutschland GmbH", "Sensirion AG (Schweiz)",
                      "Medi-Tec Supply GmbH", "Schreiner MediPharm GmbH"],
    "wp":     "PricewaterhouseCoopers GmbH WPG",
    "anwalt": "Noerr PartGmbB",
    "patent_anwalt": "Boehmert & Boehmert Partnerschaft",
    "investoren":    [("Sofinnova Partners (Paris)", "35 %"),
                      ("Lakestar (Zürich)",          "25 %"),
                      ("Management-Beteiligungs-GbR","15 %"),
                      ("HTGF (High-Tech Gründerfonds)","10 %"),
                      ("Weiterer Streubesitz",       "15 %")],
    "revenue_2023":   87_200_000,
    "ebitda_2023":    12_400_000,
    "employees_2023": 612,
    "stammkapital":   500_000,
}

ROEHRIG_AG = {
    "name":  "Brennhagen Elektronik AG",
    "short": "REA",
    "addr":  "Vaihinger Straße 120, 70567 Stuttgart",
    "hrb":   "HRB 726451, Amtsgericht Stuttgart",
    "ust":   "DE 312 487 901",
    "ceo":   "Anna Müller",                # Vorstandsvorsitzende
    "cfo":   "Laura Bauer",                # CFO
    "coo":   "Dr. Thomas Weber",
    "cto":   "Stefan Hoffmann",
    "ar_vorsitz": "Prof. Dr. Reinhold Köhler",
    "wkn":   "RHGRP1",
    "isin":  "DE000RHGRP12",
    "boerse": "Frankfurter Wertpapierbörse, Prime Standard",
    "wp":     "KPMG AG Wirtschaftsprüfungsgesellschaft",
    "anwalt": "Hengeler Mueller Partnerschaft mbB",
    "ir_bank": ["Deutsche Bank AG", "Commerzbank AG", "Berenberg"],
    "oem_kunden": ["BMW Group", "Volkswagen AG", "Mercedes-Benz Group AG",
                    "Stellantis N.V.", "Ford-Werke GmbH", "Continental AG",
                    "ZF Friedrichshafen AG", "Robert Bosch GmbH"],
    "produkte": [
        ("ICP-3",      "InfoConnect Pro (Infotainment-Modul)"),
        ("BMS-12",     "BatteryMS-12 (Batteriemanagementsystem EV)"),
        ("ADAS-V4D",   "Radar Fusion Steuergerät (Level-2/3 ADAS)"),
        ("ECU-900",    "Powertrain-ECU Gen3"),
        ("LightCtrl-7","Matrix-LED Steuermodul"),
    ],
    "revenue_2023":   612_000_000,
    "ebit_2023":      48_900_000,
    "ebitda_2023":    74_300_000,
    "employees_2023": 4180,
    "grundkapital":   62_500_000,   # AG
    "tochter": [
        # short, full name, city, country, HRB, employees, revenue (mio), focus
        ("REG", "Brennhagen Elektronik GmbH",  "Heilbronn", "DE", "HRB 221456, AG Heilbronn",  820, 280, "Produktion / Hauptwerk"),
        ("RSG", "Brennhagen Software GmbH",    "München",   "DE", "HRB 319872, AG München",    340,  62, "Embedded Software / ADAS"),
        ("RPL", "Brennhagen Polska Sp. z o.o.","Katowice",  "PL", "KRS 0000543210",            960,  98, "Produktion EMS / SMD"),
        ("RCZ", "Brennhagen CZ s.r.o.",        "Brno",      "CZ", "C 87654 KS Brno",           680,  72, "Produktion Steckverbinder"),
        ("RHU", "Brennhagen Hungary Kft.",     "Győr",      "HU", "Cg.08-09-029876",           540,  54, "Produktion Sensorik"),
        ("RCN", "Brennhagen (Shanghai) Co. Ltd.","Shanghai","CN", "913100007891234567",        320,  42, "Vertrieb / Aftermarket Asien"),
        ("RHO", "Brennhagen Holding GmbH",     "Stuttgart", "DE", "HRB 726450, AG Stuttgart",   45,   0, "Holding / Managementgesellschaft"),
    ],
}

# Mapping of subsidiary short → entry
ROEHRIG_SUBS = {t[0]: {"short": t[0], "name": t[1], "city": t[2], "country": t[3],
                       "hrb": t[4], "employees": t[5], "revenue_mio": t[6], "focus": t[7]}
                for t in ROEHRIG_AG["tochter"]}


# ── DOCX WRITING HELPERS ─────────────────────────────────────────────────────

MONTHS_DE = ["", "Januar", "Februar", "März", "April", "Mai", "Juni",
             "Juli", "August", "September", "Oktober", "November", "Dezember"]


def ds(y, m, d):
    return f"{d}. {MONTHS_DE[m]} {y}"


def eur(n):
    return f"{n:,.0f} EUR".replace(",", ".")


def num(n):
    return f"{n:,}".replace(",", ".")


def write_doc(path, header, title, sections, subtitle=None,
              confidential=False, draft=False):
    """
    Write a .docx to ``path`` with realistic legal/business structure.

    header     – dict with keys: name, addr, hrb (company identity bar)
    title      – main document title (e.g. "Anstellungsvertrag – Anna Müller")
    sections   – list of (heading, body). body may be:
                   * str  – paragraph text (split on blank line)
                   * list[list[str]] – table (first row is header)
                   * list[str] – bullet list (each str a bullet)
                   * ("clauses", [(label, [para, ...]), ...]) – numbered §
                   * ("list", [...]) – bullet list (alternative form)
    """
    doc = Document()

    if confidential:
        p = doc.add_paragraph("VERTRAULICH – STRENG VERTRAULICH / STRICTLY CONFIDENTIAL")
        p.runs[0].bold = True
        p.runs[0].font.color.rgb = RGBColor(0xCC, 0, 0)
    if draft:
        p = doc.add_paragraph("*** ENTWURF / DRAFT – NICHT ZUR WEITERGABE ***")
        p.runs[0].bold = True
        p.runs[0].font.color.rgb = RGBColor(0xCC, 0, 0)

    doc.add_heading(header["name"], level=1)
    p = doc.add_paragraph()
    p.add_run(header["addr"] + "  |  " + header["hrb"]).font.size = Pt(9)
    doc.add_paragraph("─" * 80)

    doc.add_heading(title, level=2)
    if subtitle:
        sub = doc.add_paragraph(subtitle)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for sec_title, body in sections:
        if sec_title:
            doc.add_heading(sec_title, level=3)
        _write_body(doc, body)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def _write_body(doc, body):
    if isinstance(body, tuple) and body and body[0] == "clauses":
        for label, paragraphs in body[1]:
            doc.add_heading(label, level=4)
            for i, t in enumerate(paragraphs, 1):
                p = doc.add_paragraph()
                p.add_run(f"({i}) ").bold = True
                p.add_run(t)
        return
    if isinstance(body, tuple) and body and body[0] == "list":
        for item in body[1]:
            doc.add_paragraph(item, style="List Bullet")
        return
    if isinstance(body, str):
        for para in body.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        return
    if isinstance(body, list):
        if body and isinstance(body[0], list):
            # table
            t = doc.add_table(rows=1, cols=len(body[0]))
            t.style = "Table Grid"
            for ci, hdr in enumerate(body[0]):
                t.rows[0].cells[ci].text = str(hdr)
            ncols = len(body[0])
            for row in body[1:]:
                r = t.add_row()
                for ci in range(ncols):
                    r.cells[ci].text = str(row[ci]) if ci < len(row) else ""
        else:
            for item in body:
                doc.add_paragraph(str(item), style="List Bullet")
        return


def signatures(left_name, left_role, left_company, right_name, right_role, right_company,
               place="Stuttgart", date_str_="—"):
    """Helper to format a two-column signature block as plain text section."""
    return (
        f"Ort, Datum: {place}, den {date_str_}\n\n"
        f"_________________________            _________________________\n"
        f"{left_name}                          {right_name}\n"
        f"{left_role}                          {right_role}\n"
        f"{left_company}                       {right_company}"
    )


def read_doc_title(path):
    """Extract the main title (level-2 heading) of an existing stub doc."""
    try:
        d = Document(path)
        for p in d.paragraphs:
            if p.style.name.startswith("Heading 2") and p.text.strip():
                return p.text.strip()
    except Exception:
        pass
    return Path(path).stem
